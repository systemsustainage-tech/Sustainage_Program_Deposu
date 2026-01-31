
from docx import Document
import sys

file_path = r"c:\SUSTAINAGESERVER\Kullanıcı Arayüzünü Modernleştir.docx"

try:
    doc = Document(file_path)
    print(f"--- Content of {file_path} ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    # Also print tables if any
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(" | ".join(row_text))

except Exception as e:
    print(f"Error reading docx: {e}")
