#!/usr/bin/env python3
# -*- coding: utf-8 -*-
r"""
SDG Program Modülleri Tanıtım Dokümanı Oluşturucu (Word .docx)
- İçerik: modüller, butonlar/bölümler, raporlama
- Resimler mevcutsa eklenir (c:\SDG\resimler)
"""
from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

ROOT_DIR = Path(__file__).resolve().parent.parent
OUTPUT_PATH = ROOT_DIR / "data/exports/SDG_Modul_Tanitim.docx"
# Önce resimler/sunum dizinini kullan, yoksa resimler’e düş
IMAGES_DIR_SUNUM = ROOT_DIR / "resimler/sunum"
IMAGES_DIR_MAIN = ROOT_DIR / "resimler"
IMAGES_DIR = IMAGES_DIR_SUNUM if IMAGES_DIR_SUNUM.exists() else IMAGES_DIR_MAIN

# Kurumsal renk paleti
PALETTE = {
    "Kirmizi": (0xE5, 0x39, 0x35),
    "Yesil": (0x43, 0xA0, 0x47),
    "Mavi": (0x1E, 0x88, 0xE5),
    "Sari": (0xFB, 0xC0, 0x2D),
    "Mor": (0x8E, 0x24, 0xAA),
    "Siyah": (0x21, 0x21, 0x21),
    "Turkuaz": (0x26, 0xC6, 0xDA),
}

MODULE_COLORS = {
    "SDG": PALETTE["Mavi"],
    "GRI": PALETTE["Mor"],
    "TSRS": PALETTE["Turkuaz"],
    "Karbon": PALETTE["Yesil"],
    "Eslesme": PALETTE["Sari"],
    "Raporlama": PALETTE["Kirmizi"],
    "Yonetim": PALETTE["Siyah"],
}


def add_heading(doc: Document, text: str, level: int = 1, rgb: tuple[int, int, int] | None = None) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if rgb:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*rgb)


def add_bullets(doc: Document, bullets: list[str]) -> None:
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p_format = p.paragraph_format
        p_format.space_after = Pt(2)


def add_image(doc: Document, image_name: str, width_inches: float = 4.5) -> None:
    img_path = IMAGES_DIR / image_name
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(width_inches))


def build_doc() -> None:
    doc = Document()
    # Başlık sayfası
    add_heading(doc, "SDG Program Modülleri Tanıtım Dokümanı", level=0)
    add_image(doc, "main.png", width_inches=5.5)

    # SDG
    add_heading(doc, "SDG", level=1, rgb=MODULE_COLORS["SDG"])
    sdg_bullets = [
        "Amaç: 17 hedef, 169 alt hedef ve 232 gösterge yönetimi",
        "Bölümler: Mevcut Hedefler, Seçilen Hedefler, İstatistik kartları",
        "Butonlar: İlerleme Takibi, Veri Toplama, Gelişmiş Analiz, Soru Bankası",
        "Butonlar: Veri Doğrulama, Raporlama, Detay Tablosu",
        "Raporlar: SDG ilerleme ve gösterge yanıt raporları",
    ]
    add_bullets(doc, sdg_bullets)
    add_image(doc, "SDGs1.jpeg")

    # GRI
    add_heading(doc, "GRI", level=1, rgb=MODULE_COLORS["GRI"])
    gri_bullets = [
        "Amaç: Global Reporting Initiative standartlarını inceleme ve yanıt",
        "Bölümler: Kategoriler (Universal/Economic/Environmental/Social), İçerik alanı",
        "Butonlar: SDG Seçimine Göre Filtrele, Filtreleri Temizle, CSV Dışa Aktar",
        "Fonksiyonlar: Standart/indikator kartları, yanıt kaydı",
        "Raporlar: CSV dışa aktarma ve GRI uyumlu kayıtlar",
    ]
    add_bullets(doc, gri_bullets)

    # TSRS
    add_heading(doc, "TSRS", level=1, rgb=MODULE_COLORS["TSRS"])
    tsrs_bullets = [
        "Amaç: TSRS standartları ve göstergeleri ile raporlama",
        "Bölümler: Yönetişim, Strateji, Risk Yönetimi, Metrikler",
        "Sekmeler: TSRS Standartları, Raporlama, GRI-SDG Entegrasyon",
        "Butonlar: Kategori butonları (Yönetişim/Strateji/Risk/Metrikler)",
        "Raporlar: TSRSReportingGUI ile TSRS raporlama",
    ]
    add_bullets(doc, tsrs_bullets)
    add_image(doc, "ESG1.png")

    # Karbon
    add_heading(doc, "Karbon", level=1, rgb=MODULE_COLORS["Karbon"])
    carbon_bullets = [
        "Amaç: Scope 1/2/3 emisyon veri girişi ve hesaplama",
        "Sekmeler: Scope1 (Sabit/Mobil/Kaçak), Scope2 (Elektrik/Isıtma), Scope3",
        "Sekmeler: Hedefler, Azaltma Girişimleri, Raporlar",
        "Butonlar: Kaydet ve Hesapla, CSV dışa aktarma (gerekli yerlerde)",
        "Raporlar: Özet rapor oluşturma ve DB’ye kaydetme",
    ]
    add_bullets(doc, carbon_bullets)
    add_image(doc, "esg.png")

    # Eşleştirme
    add_heading(doc, "Eşleştirme", level=1, rgb=MODULE_COLORS["Eslesme"])
    mapping_bullets = [
        "Amaç: SDG-GRI-TSRS eşleştirmelerini yönetme",
        "Sekmeler: SDG-GRI, SDG-TSRS, GRI-TSRS, CSV İşlemleri",
        "Butonlar: SDG seçimine göre filtrele, arama ve CSV import/export",
        "Raporlar: CSV export ile eşleştirme tabloları",
    ]
    add_bullets(doc, mapping_bullets)

    # Raporlama
    add_heading(doc, "Raporlama", level=1, rgb=MODULE_COLORS["Raporlama"])
    reporting_bullets = [
        "Amaç: SDG / SDG+GRI / SDG+GRI+TSRS raporlamaları",
        "Butonlar: SDG Raporu, SDG+GRI Raporu, SDG+GRI+TSRS Raporu",
        "Çıktılar: DOCX/Excel/PDF (modül raporlarına bağlı)",
    ]
    add_bullets(doc, reporting_bullets)

    # Yönetim
    add_heading(doc, "Yönetim", level=1, rgb=MODULE_COLORS["Yonetim"])
    mgmt_bullets = [
        "Amaç: Kullanıcı/Rol/Lisans/Modül kontrol",
        "Buton: Rol/İzin Matrisi (CSV) görüntüleme",
        "Güvenlik: Super-admin, lisans doğrulama, audit (ayrı modül)",
    ]
    add_bullets(doc, mgmt_bullets)
    add_image(doc, "login.jpg")

    # Kaydet
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


def main() -> None:
    out = build_doc()
    logging.info(f"Doküman oluşturuldu: {out}")


if __name__ == "__main__":
    main()
