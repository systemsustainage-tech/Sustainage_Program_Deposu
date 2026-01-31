import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX Rapor Render
Kullanım:
  python tools/render_report.py 
    --db ./sdg_desktop.sqlite 
    --type sdg 
    --out ./data/exports/rapor_sdg.docx 
    --company_id 1 
    --period 2024
"""
import argparse
import datetime
import json
import os
import sqlite3
from typing import Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from modules.esg.esg_manager import ESGManager
from modules.ungc.ungc_manager import UNGCManager

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

TEMPLATES = {
  "sdg": "raporlama/templates/sdg/template.docx",
  "sdg_gri": "raporlama/templates/sdg_gri/template.docx",
  "sdg_gri_tsrs": "raporlama/templates/sdg_gri_tsrs/template.docx",
  # İsteğe bağlı GRI şablonu; dosya yoksa boş doküman ile çalışacağız
  "gri": "raporlama/templates/gri/template.docx"
}

def get_company(cur, company_id) -> None:
    r = cur.execute("SELECT name FROM companies WHERE id=?", (company_id,)).fetchone()
    return r[0] if r else f"Şirket {company_id}"

def get_company_info(cur, company_id) -> None:
    row = cur.execute("SELECT * FROM company_info WHERE company_id=?", (company_id,)).fetchone()
    if not row:
        return {}
    # Şema: id, company_id, ...
    cols = [d[1] for d in cur.execute("PRAGMA table_info(company_info)").fetchall()]
    data = {}
    for idx, col in enumerate(cols):
        data[col] = row[idx]
    # Özet inşa et
    ozet_parts = []
    if data.get('sirket_adi'):
        ozet_parts.append("Şirket: " + str(data.get('sirket_adi')))
    loc = ", ".join([
        str(data.get('sehir') or '').strip(),
        str(data.get('ulke') or '').strip(),
    ]).strip(', ')
    if loc:
        ozet_parts.append("Konum: " + loc)
    if data.get('calisan_sayisi'):
        ozet_parts.append("Çalışan sayısı: " + str(data.get('calisan_sayisi')))
    if data.get('net_gelir'):
        ozet_parts.append("Net gelir: " + str(data.get('net_gelir')))
    if data.get('toplam_varliklar'):
        ozet_parts.append("Toplam varlıklar: " + str(data.get('toplam_varliklar')))
    if data.get('tedarik_zinciri'):
        tz = str(data.get('tedarik_zinciri') or '')
        cut = tz[:160]
        suffix = "..." if len(tz) > 160 else ""
        ozet_parts.append("Tedarik zinciri: " + cut + suffix)
    ozet = "; ".join(ozet_parts) if ozet_parts else "(Firma bilgileri girilmemiş)"
    return {
        'ozet': ozet,
        'adres': data.get('adres') or '',
        'ulke': data.get('ulke') or '',
        'sehir': data.get('sehir') or '',
        'web': data.get('web') or '',
        'eposta': data.get('eposta') or '',
        'irtibat': data.get('irtibat_noktasi') or ''
    }

def build_summaries(cur, company_id, period) -> None:
    """SDG, GRI ve TSRS özetlerini oluşturur.
    - SDG: responses bazlı genel özet
    - GRI: seçilen SDG göstergeleri üzerinden GRI açıklama sayımı
    - TSRS: hem SDG→TSRS hem de GRI→TSRS üzerinden TSRS metrik sayımı ve bölüm bazlı detaylar
    """

    # SDG özet (responses tablosu bazlı)
    cur2 = cur.execute(
        """SELECT COUNT(*), AVG(progress_pct) FROM responses WHERE company_id=? AND period=?""",
        (company_id, period)
    ).fetchone()
    total_resp = cur2[0] or 0
    avg_progress = round(cur2[1] or 0, 1)
    sdg_ozet = f"Toplam {total_resp} gösterge için veri girildi. Ortalama ilerleme: %{avg_progress}."

    # İlgili SDG gösterge kodlarını çıkar
    indicator_rows = cur.execute(
        """
        SELECT DISTINCT i.code
        FROM sdg_indicators i
        JOIN responses r ON r.indicator_id=i.id
        WHERE r.company_id=? AND r.period=?
        """,
        (company_id, period)
    ).fetchall()
    sdg_indicator_codes = [row[0] for row in indicator_rows]

    # GRI eşleştirmeleri (özet) – seçilen SDG göstergelerinden türet
    if sdg_indicator_codes:
        cur.execute("CREATE TEMP TABLE IF NOT EXISTS tmp_sdg_codes (code TEXT)")
        cur.execute("DELETE FROM tmp_sdg_codes")
        cur.executemany("INSERT INTO tmp_sdg_codes (code) VALUES (?)", [(c,) for c in sdg_indicator_codes])
        gri_disclosures = [
            row[0]
            for row in cur.execute(
                "SELECT DISTINCT gri_disclosure FROM map_sdg_gri "
                "WHERE sdg_indicator_code IN (SELECT code FROM tmp_sdg_codes)"
            ).fetchall()
        ]
        gri_count = len(gri_disclosures)
        gri_ozet = f"Seçili SDG göstergeleri ile ilişkili {gri_count} farklı GRI açıklaması bulundu."
    else:
        gri_disclosures = []
        gri_ozet = "GRI eşleştirme verisi bulunamadı."

    # TSRS özetleri
    tsrs_sections_summary = []
    tsrs_total = 0

    # 1) SDG→TSRS
    if sdg_indicator_codes:
        tsrs_rows_sdg = cur.execute(
            """
            SELECT tsrs_section, tsrs_metric
            FROM map_sdg_tsrs
            WHERE sdg_indicator_code IN (SELECT code FROM tmp_sdg_codes)
            """
        ).fetchall()
    else:
        tsrs_rows_sdg = []

    # 2) GRI→TSRS
    if gri_disclosures:
        cur.execute("CREATE TEMP TABLE IF NOT EXISTS tmp_gri_codes (code TEXT)")
        cur.execute("DELETE FROM tmp_gri_codes")
        cur.executemany("INSERT INTO tmp_gri_codes (code) VALUES (?)", [(c,) for c in gri_disclosures])
        tsrs_rows_gri = cur.execute(
            """
            SELECT tsrs_section, tsrs_metric
            FROM map_gri_tsrs
            WHERE gri_disclosure IN (SELECT code FROM tmp_gri_codes)
            """
        ).fetchall()
    else:
        tsrs_rows_gri = []

    # Bölüm bazında birleştir ve say
    from collections import defaultdict
    section_metrics = defaultdict(set)
    for sec, met in tsrs_rows_sdg + tsrs_rows_gri:
        if sec:
            section_metrics[sec].add(met)

    for sec, metrics in sorted(section_metrics.items()):
        cnt = len(metrics)
        tsrs_total += cnt
        sample = ", ".join(sorted(list(metrics))[:8])
        tsrs_sections_summary.append(f"{sec}: {cnt} metrik (örnekler: {sample}{' …' if len(metrics)>8 else ''})")

    if tsrs_sections_summary:
        tsrs_hdr = "TSRS metrikleri toplam " + str(tsrs_total) + " bağlantı bulundu. Bölüm bazında: "
        tsrs_ozet = tsrs_hdr + "; ".join(tsrs_sections_summary)
    else:
        tsrs_ozet = "TSRS eşleştirme verisi bulunamadı."

    return (sdg_ozet, gri_ozet, tsrs_ozet)

def get_gri_selected_rows(cur, company_id) -> None:
    """Seçilmiş GRI göstergelerini ve TSRS eşleştirmelerini döndürür."""
    rows = cur.execute(
        """
        SELECT gi.code AS indicator_code, gi.title AS indicator_title,
               gs.code AS standard_code, gs.title AS standard_title
        FROM gri_selections sel
        JOIN gri_indicators gi ON gi.id = sel.indicator_id
        JOIN gri_standards gs ON gs.id = gi.standard_id
        WHERE sel.company_id=? AND sel.selected=1
        ORDER BY gs.code, gi.code
        """,
        (company_id,)
    ).fetchall()
    out = []
    for icode, ititle, scode, stitle in rows:
        tsrs = cur.execute(
            """
            SELECT tsrs_section, tsrs_metric
            FROM map_gri_tsrs
            WHERE gri_disclosure = ?
            ORDER BY tsrs_section, tsrs_metric
            """,
            (icode,)
        ).fetchall()
        tsrs_summary = ", ".join([f"{sec}: {met}" for sec, met in tsrs]) if tsrs else "—"
        out.append((scode, icode, ititle, tsrs_summary))
    return out

def get_ceo_messages(cur, company_id, period) -> None:
    """CEO mesajlarını getir"""
    try:
        rows = cur.execute(
            """
            SELECT title, message_type, year, quarter, content, 
                   key_achievements, challenges, future_commitments,
                   signature_name, signature_title, signature_date
            FROM ceo_messages 
            WHERE company_id = ? AND year = ? AND is_published = 1
            ORDER BY year DESC, quarter DESC, created_at DESC
            """,
            (company_id, period)
        ).fetchall()
        
        messages = []
        for row in rows:
            messages.append({
                'title': row[0],
                'message_type': row[1],
                'year': row[2],
                'quarter': row[3],
                'content': row[4],
                'key_achievements': json.loads(row[5] or '[]'),
                'challenges': json.loads(row[6] or '[]'),
                'future_commitments': json.loads(row[7] or '[]'),
                'signature_name': row[8],
                'signature_title': row[9],
                'signature_date': row[10]
            })
        
        return messages
    except Exception as e:
        logging.info(f"[UYARI] CEO mesajları alınamadı: {e}")
        return []

def add_ceo_message_section(doc, cur, company_id, period) -> None:
    """DOCX dokümana CEO mesajı bölümü ekler"""
    try:
        messages = get_ceo_messages(cur, company_id, period)
        
        if not messages:
            return
        
        # Sayfa sonu ekle
        doc.add_page_break()
        doc.add_heading('CEO/Genel Müdür Mesajı', level=1)
        
        # En son mesajı kullan (yıllık varsa onu, yoksa çeyreklik)
        annual_message = next((msg for msg in messages if msg['message_type'] == 'annual'), None)
        latest_message = annual_message or messages[0]
        
        # Mesaj başlığı
        if latest_message['quarter']:
            title_line = (
                str(latest_message['title'])
                + " - "
                + str(latest_message['year'])
                + " Q"
                + str(latest_message['quarter'])
            )
            doc.add_heading(title_line, level=2)
        else:
            title_line = str(latest_message['title']) + " - " + str(latest_message['year'])
            doc.add_heading(title_line, level=2)
        
        # Mesaj içeriği
        if latest_message['content']:
            doc.add_paragraph(latest_message['content'])
            doc.add_paragraph()  # Boş satır
        
        # Ana başarılar
        if latest_message['key_achievements']:
            doc.add_heading('Ana Başarılar', level=3)
            for achievement in latest_message['key_achievements']:
                if achievement.strip():
                    p = doc.add_paragraph(achievement, style='List Bullet')
        
        # Karşılaşılan zorluklar
        if latest_message['challenges']:
            doc.add_heading('Karşılaşılan Zorluklar', level=3)
            for challenge in latest_message['challenges']:
                if challenge.strip():
                    p = doc.add_paragraph(challenge, style='List Bullet')
        
        # Gelecek taahhütleri
        if latest_message['future_commitments']:
            doc.add_heading('Gelecek Taahhütleri', level=3)
            for commitment in latest_message['future_commitments']:
                if commitment.strip():
                    p = doc.add_paragraph(commitment, style='List Bullet')
        
        # İmza
        if latest_message['signature_name'] or latest_message['signature_title']:
            doc.add_paragraph()
            doc.add_paragraph('Saygılarımla,')
            doc.add_paragraph()
            
            if latest_message['signature_name']:
                p = doc.add_paragraph(latest_message['signature_name'])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            if latest_message['signature_title']:
                p = doc.add_paragraph(latest_message['signature_title'])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            if latest_message['signature_date']:
                p = doc.add_paragraph(latest_message['signature_date'])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    except Exception as e:
        logging.info(f"[UYARI] CEO mesajı bölümü eklenemedi: {e}")

def add_gri_selected_table(doc, cur, company_id) -> None:
    """DOCX dokümana GRI seçilmiş göstergeler tablosu ekler."""
    try:
        rows = get_gri_selected_rows(cur, company_id)
    except Exception:
        rows = []
    # Başlık
    doc.add_page_break()
    doc.add_heading('GRI İçerik İndeksi – Seçilmiş Göstergeler', level=1)
    if not rows:
        doc.add_paragraph('Seçilmiş GRI göstergesi bulunamadı.')
        return
    table = doc.add_table(rows=len(rows)+1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "GRI Standart"
    hdr[1].text = "GRI Gösterge"
    hdr[2].text = "Başlık"
    hdr[3].text = "TSRS Eşleştirmeleri"
    for i, (scode, icode, ititle, tsrs) in enumerate(rows):
        cells = table.rows[i+1].cells
        cells[0].text = scode
        cells[1].text = icode
        cells[2].text = ititle
        cells[3].text = tsrs

def build_methodology_and_validation(cur, company_id) -> Dict[str, str]:
    """Metodoloji, kapsam ve doğrulama özeti için metinler üretir."""
    # Basit metodoloji metni
    metodoloji = (
        "Bu rapor, şirket tarafından sağlanan SDG göstergeleri yanıtları ve "
        "ilgili kanıt/doküman referansları temel alınarak hazırlanmıştır. Veri toplama, "
        "standart soru setleri üzerinden yürütülmüş, ilerleme yüzdeleri yanıtların "
        "tamamlanma durumuna göre hesaplanmıştır. Eşleştirmeler resmi SDG-GRI ve TSRS "
        "haritaları üzerinden türetilmiştir."
    )
    kapsam = (
        "Kapsam: Bu rapor yalnızca değerlendirilen dönem ve mevcut veri kaynakları "
        "ile sınırlıdır. Birim bazlı detaylar, şirketin beyan ettiği sorumluluk ve veri "
        "yöntemlerine göre derlenmiştir."
    )

    # Doğrulama özeti: sonuç ve kalite skorları
    # 1) sdg_validation_results üzerinden ihlal sayımı
    res_rows = cur.execute(
        """
        SELECT COALESCE(severity_level, 'info') AS sev, COUNT(*)
        FROM sdg_validation_results
        WHERE company_id = ?
        GROUP BY sev
        """,
        (company_id,)
    ).fetchall()
    sev_counts = {row[0]: row[1] for row in res_rows}
    sev_order = ['error', 'warning', 'info']
    sev_text = ", ".join([f"{s}: {sev_counts.get(s, 0)}" for s in sev_order])
    if not sev_text:
        sev_text = "Doğrulama bulgusu kaydı bulunamadı."

    # 2) sdg_data_quality_scores üzerinden ortalama skorlar
    qs = cur.execute(
        """
        SELECT AVG(completeness_score), AVG(accuracy_score), AVG(consistency_score),
               AVG(timeliness_score), AVG(overall_quality_score)
        FROM sdg_data_quality_scores
        WHERE company_id = ?
        """,
        (company_id,)
    ).fetchone() or (0,0,0,0,0)
    comp, acc, cons, time, overall = [round(v or 0, 1) for v in qs]
    kalite = (
        f"Kalite Skorları (ortalama): Eksiksizlik %{comp}, Doğruluk %{acc}, "
        f"Tutarlılık %{cons}, Güncellik %{time}, Genel %{overall}."
    )

    # Veri kaynakları ve yönetişim (firma_info üzerinden mevcutsa ekle)
    # company_info tablosunda ilgili alanlar opsiyonel olabilir.
    kaynaklar = []
    try:
        sql = (
            "SELECT data_sources, governance_notes, assurance_statement "
            "FROM company_info WHERE company_id=?"
        )
        info_row = cur.execute(sql, (company_id,)).fetchone()
        if info_row:
            ds, gov, ass = info_row
            if ds:
                kaynaklar.append("Veri Kaynakları: " + str(ds)[:300] + ("…" if len(str(ds)) > 300 else ""))
            if gov:
                kaynaklar.append("Yönetişim: " + str(gov)[:300] + ("…" if len(str(gov)) > 300 else ""))
            if ass:
                kaynaklar.append("Beyan/Assurance: " + str(ass)[:300] + ("…" if len(str(ass)) > 300 else ""))
    except Exception as e:
        logging.error(f"Silent error caught: {str(e)}")
    kaynak_text = " \n".join(kaynaklar) if kaynaklar else "Veri kaynakları ve yönetişim beyanları sisteme girilmemiş."

    full_text = (
        "Metodoloji:\n" + metodoloji + "\n\n" +
        "Kapsam:\n" + kapsam + "\n\n" +
        "Doğrulama Özeti:\n" + sev_text + "\n" + kalite + "\n\n" +
        kaynak_text
    )
    return {"metodoloji_ve_dogrulama": full_text}

def render(db, typ, out, company_id, period) -> None:
    tpl_path = TEMPLATES.get(typ)
    if tpl_path and os.path.exists(tpl_path):
        doc = Document(tpl_path)
    else:
        # Şablon yoksa boş doküman oluşturup başlık ekleyelim
        doc = Document()
        doc.add_heading(f"{typ.upper()} Raporu", level=1)
    
    # Türkçe karakter desteği için font ayarları
    def set_turkish_font_paragraph(paragraph, font_name='Calibri') -> None:
        for run in paragraph.runs:
            run.font.name = font_name
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            ascii_key = ns + 'ascii'
            hansi_key = ns + 'hAnsi'
            run._element.rPr.rFonts.set(ascii_key, font_name)
            run._element.rPr.rFonts.set(hansi_key, font_name)
    
    def set_turkish_font_cell(cell, font_name='Calibri') -> None:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                ascii_key = ns + 'ascii'
                hansi_key = ns + 'hAnsi'
                run._element.rPr.rFonts.set(ascii_key, font_name)
                run._element.rPr.rFonts.set(hansi_key, font_name)

    conn = sqlite3.connect(db)
    cur = conn.cursor()
    sirket = get_company(cur, company_id)
    firma = get_company_info(cur, company_id)
    sdg_ozet, gri_ozet, tsrs_ozet = build_summaries(cur, company_id, period)
    met_val = build_methodology_and_validation(cur, company_id)

    repl = {
        "{{sirket_adi}}": sirket,
        "{{donem}}": str(period),
        "{{rapor_tipi}}": typ.upper(),
        "{{olusturma_tarihi}}": datetime.date.today().isoformat(),
        "{{sdg_ozet}}": sdg_ozet,
        "{{gri_eslesme_ozet}}": gri_ozet,
        "{{tsrs_eslesme_ozet}}": tsrs_ozet,
        "{{ilerleme_grafigi_yeri}}": "(Grafik alanı)",
        "{{page}}": "",  # docx alanı yerine boş bırakıyoruz
        # Firma bilgileri entegrasyonu
        "{{firma_bilgileri_ozet}}": firma.get('ozet', ''),
        "{{firma_adres}}": firma.get('adres', ''),
        "{{firma_ulke}}": firma.get('ulke', ''),
        "{{firma_sehir}}": firma.get('sehir', ''),
        "{{firma_web}}": firma.get('web', ''),
        "{{firma_eposta}}": firma.get('eposta', ''),
        "{{firma_irtibat}}": firma.get('irtibat', '')
    }
    # İsteğe bağlı metodoloji/validasyon bölümü (placeholder yoksa sonradan ekleyeceğiz)
    repl["{{metodoloji_ve_dogrulama}}"] = met_val.get("metodoloji_ve_dogrulama", "")

    for p in doc.paragraphs:
        for k,v in repl.items():
            if k in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    inline[i].text = inline[i].text.replace(k, v)

    # tablolar varsa içinde de değiştir
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k,v in repl.items():
                    if k in cell.text:
                        for pr in cell.paragraphs:
                            for r in pr.runs:
                                r.text = r.text.replace(k, v)

    # Eğer şablonda metodoloji placeholder'ı yoksa, raporun sonuna ek bir bölüm ekleyelim
    if "{{metodoloji_ve_dogrulama}}" not in "\n".join([p.text for p in doc.paragraphs]):
        doc.add_page_break()
        doc.add_heading('Metodoloji ve Veri Kalitesi', level=1)
        for line in met_val.get("metodoloji_ve_dogrulama", "").split("\n"):
            if line.strip():
                doc.add_paragraph(line)

    # CEO mesajı bölümünü ekle (tüm rapor türleri için)
    try:
        add_ceo_message_section(doc, cur, company_id, period)
    except Exception as e:
        # CEO mesajı eklenemezse rapor yine de oluşturulsun
        logging.error(f"Silent error caught: {str(e)}")

    # sdg_gri ve sdg_gri_tsrs için: seçilmiş GRI göstergeleri tablosu ekle
    if typ in ("sdg_gri", "sdg_gri_tsrs", "gri"):
        try:
            add_gri_selected_table(doc, cur, company_id)
        except Exception as e:
            # Tablo eklenemese de raporu kaydedelim
            logging.error(f"Silent error caught: {str(e)}")

    # ESG Bölümü: E/S/G ve Genel skorları mevcut raporlara entegre
    try:
        base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
        esg_mgr = ESGManager(base_dir)
        esg_scores = esg_mgr.compute_scores(company_id, period)

        has_placeholder = any("{{esg_section}}" in p.text for p in doc.paragraphs)

        def add_esg_section(d) -> None:
            d.add_heading('ESG Skorları', level=1)
            d.add_paragraph(
                f"Genel ESG Skoru: %{esg_scores.get('overall', 0)} | "
                f"E %{esg_scores.get('E', 0)}, S %{esg_scores.get('S', 0)}, G %{esg_scores.get('G', 0)}"
            )
            # Detay tablosu: GRI ve TSRS yanıtlanan/toplam ve bonuslar
            tbl = d.add_table(rows=1, cols=3)
            hdr = tbl.rows[0].cells
            hdr[0].text = 'Kaynak'
            hdr[1].text = 'Yanıtlanan'
            hdr[2].text = 'Toplam'
            det = esg_scores.get('details', {})
            gri = det.get('gri', {})
            tsrs = det.get('tsrs', {})
            bonuses = det.get('bonuses', {})

            def add_row(name, answered, total) -> None:
                cells = tbl.add_row().cells
                cells[0].text = name
                cells[1].text = str(answered)
                cells[2].text = str(total)

            gri_ans = (gri.get('environmental', {}).get('answered', 0)
                       + gri.get('social', {}).get('answered', 0)
                       + gri.get('governance', {}).get('answered', 0))
            gri_tot = (gri.get('environmental', {}).get('total', 0)
                       + gri.get('social', {}).get('total', 0)
                       + gri.get('governance', {}).get('total', 0))
            tsrs_ans = (tsrs.get('environmental', {}).get('answered', 0)
                        + tsrs.get('social', {}).get('answered', 0)
                        + tsrs.get('governance', {}).get('answered', 0))
            tsrs_tot = (tsrs.get('environmental', {}).get('total', 0)
                        + tsrs.get('social', {}).get('total', 0)
                        + tsrs.get('governance', {}).get('total', 0))

            add_row('GRI (E/S/G2)', gri_ans, gri_tot)
            add_row('TSRS (E/S/G)', tsrs_ans, tsrs_tot)
            add_row('Bonus Kaynaklar (SDG Yanıtları)', bonuses.get('sdg_answers', 0), bonuses.get('sdg_answers', 0))

        if has_placeholder:
            # Placeholder'ı temizleyip bölüm ekle
            for p in doc.paragraphs:
                if '{{esg_section}}' in p.text:
                    p.text = ''
                    add_esg_section(doc)
                    break
        else:
            # Sona ekle
            doc.add_page_break()
            add_esg_section(doc)
    except Exception as e:
        # ESG bölümü eklenemezse rapor yine de oluşturulsun
        logging.error(f"Silent error caught: {str(e)}")

    # Tüm paragraflara Türkçe font uygula
    for paragraph in doc.paragraphs:
        set_turkish_font_paragraph(paragraph, 'Calibri')
    
    # Tüm tablo hücrelerine Türkçe font uygula
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_turkish_font_cell(cell, 'Calibri')

    # UNGC Uyum Bölümü: Mevcut raporlara entegre
    try:
        ungc = UNGCManager(db)
        ungc_res = ungc.compute_principle_status(company_id, period)
        # Şablonda placeholder varsa yerini doldur, yoksa sona ek bir bölüm ekle
        has_placeholder = any("{{ungc_section}}" in p.text for p in doc.paragraphs)
        section_title = 'UN Global Compact Uyum'
        section_text = f"Toplam UNGC Skoru: %{ungc_res.get('overall_score', 0)}"
        # Tablo: Principle, Category, Status, Score
        def add_ungc_table(d) -> None:
            d.add_heading(section_title, level=1)
            d.add_paragraph(section_text)
            tbl = d.add_table(rows=1, cols=4)
            hdr = tbl.rows[0].cells
            hdr[0].text = 'İlke'
            hdr[1].text = 'Kategori'
            hdr[2].text = 'Durum'
            hdr[3].text = 'Skor %'
            for pr in ungc_res.get('principles', []):
                row = tbl.add_row().cells
                row[0].text = pr.get('principle_id','')
                row[1].text = pr.get('category','')
                row[2].text = pr.get('status','')
                row[3].text = str(pr.get('score',0))

        if has_placeholder:
            # Basit yer değiştirme: placeholder'ı kaldırıp bölüm ekle
            for p in doc.paragraphs:
                if '{{ungc_section}}' in p.text:
                    p.text = ''
                    add_ungc_table(doc)
                    break
        else:
            # Sona ekle
            doc.add_page_break()
            add_ungc_table(doc)
    except Exception as e:
        # UNGC bölümü eklenemezse rapor yine de oluşturulsun
        logging.error(f"Silent error caught: {str(e)}")

    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    logging.info("Rapor oluşturuldu:", out)

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--db", required=True)
    ap.add_argument("--type", choices=["sdg","sdg_gri","sdg_gri_tsrs","gri"], required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--company_id", type=int, default=1)
    ap.add_argument("--period", default=str(datetime.date.today().year))
    args = ap.parse_args()
    render(args.db, args.type, args.out, args.company_id, args.period)
"""
"""
