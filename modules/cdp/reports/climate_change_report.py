#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CDP Climate Change Raporu
İklim değişikliği raporlaması için PDF/DOCX oluşturma
"""

import logging
import os
import sqlite3
from datetime import datetime
from typing import Any, Dict, Optional

import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Cm

matplotlib.use('Agg')  # GUI olmadan çalış

from ..ai_analyzer import CDPAIAnalyzer
from ..cdp_data_collector import CDPDataCollector


def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    para = _add_turkish_paragraph(doc, text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = _add_turkish_heading(doc, text, level=level)
    for run in heading.runs:
        run.font.name = font_name
    return heading


class ClimateChangeReport:
    """CDP Climate Change Raporu oluşturucu"""

    def __init__(self, use_ai: bool = True, api_key: Optional[str] = None):
        """
        Args:
            use_ai: AI analiz kullan
            api_key: OpenAI API key (opsiyonel)
        """
        self.collector = CDPDataCollector()
        self.analyzer = CDPAIAnalyzer(use_api=False if not api_key else True, api_key=api_key)
        self.use_ai = use_ai

    def generate_report(self, company_id: int, year: int, output_path: str) -> bool:
        """
        CDP Climate Change raporu oluştur
        
        Args:
            company_id: Şirket ID
            year: Raporlama yılı
            output_path: Çıktı dosyası yolu
        
        Returns:
            Başarılı ise True
        """
        try:
            logging.info(f"[CDP Climate Change] Rapor oluşturuluyor: {year}")

            # Veri toplama
            data = self._collect_all_data(company_id, year)

            # AI Analiz
            if self.use_ai:
                logging.info("[CDP Climate Change] AI analizi yapılıyor...")
                analysis = self.analyzer.analyze_climate_performance(data)
                score = self.analyzer.calculate_cdp_score_estimate(data)
            else:
                analysis = {}
                score = {}

            # Grafikler oluştur
            charts_dir = os.path.join(os.path.dirname(output_path), 'charts')
            os.makedirs(charts_dir, exist_ok=True)

            chart_files = self._create_charts(data, charts_dir)

            # DOCX raporu oluştur
            doc = self._create_docx_report(data, analysis, score, chart_files)
            doc.save(output_path)

            logging.info(f"[CDP Climate Change] Rapor kaydedildi: {output_path}")
            return True

        except Exception as e:
            logging.error(f"[HATA] CDP Climate Change raporu oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _collect_all_data(self, company_id: int, year: int) -> Dict[str, Any]:
        """Tüm gerekli verileri topla"""
        data = {}
        data['company_id'] = company_id
        data['company_info'] = self.collector.collect_company_info(company_id)
        data['emissions'] = self.collector.collect_climate_data(company_id, year).get('emissions', {})
        data['energy'] = self.collector.collect_climate_data(company_id, year).get('energy', {})
        data['comparison'] = self.collector.get_previous_year_comparison(company_id, year)
        data['intensity'] = self.collector.calculate_intensity_metrics(company_id, year)
        data['year'] = year

        return data

    def _add_logo(self, doc, company_id: int):
        """Raporun başına şirket logosunu ekle"""
        try:
            # 1. Veritabanından logo yolunu al
            conn = sqlite3.connect(self.collector.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT logo_path FROM company_profiles WHERE company_id = ?", (company_id,))
            result = cursor.fetchone()
            conn.close()
            
            logo_path = None
            if result and result[0] and os.path.exists(result[0]):
                logo_path = result[0]
            else:
                # 2. Varsayılan yolda ara
                data_dir = os.path.dirname(self.collector.db_path)
                possible_path = os.path.join(data_dir, "company_logos", f"company_{company_id}_logo.png")
                if os.path.exists(possible_path):
                    logo_path = possible_path
                else:
                    # jpg dene
                    possible_path = possible_path.replace(".png", ".jpg")
                    if os.path.exists(possible_path):
                        logo_path = possible_path

            if logo_path:
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, height=Cm(2.5))
        except Exception as e:
            logging.warning(f"Logo eklenirken hata: {e}")

    def _create_charts(self, data: Dict[str, Any], output_dir: str) -> Dict[str, str]:
        """Grafikleri oluştur"""
        chart_files = {}

        # 1. Scope bazında emisyon dağılımı (Pasta)
        emissions = data.get('emissions', {})
        if emissions:
            fig, ax = plt.subplots(figsize=(8, 6))

            scopes = ['Scope 1', 'Scope 2', 'Scope 3']
            values = [
                emissions.get('scope1', 0),
                emissions.get('scope2', 0),
                emissions.get('scope3', 0)
            ]

            colors = ['#FF6B6B', '#4ECDC4', '#45B7D1']

            ax.pie(values, labels=scopes, autopct='%1.1f%%', startangle=90, colors=colors)
            ax.set_title('Emisyon Dağılımı (Scope Bazında)', fontsize=14, fontweight='bold')

            chart_path = os.path.join(output_dir, 'emissions_by_scope.png')
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            chart_files['emissions_pie'] = chart_path

        # 2. Yıllık karşılaştırma (Bar)
        comparison = data.get('comparison', {})
        if comparison and 'emissions_change' in comparison:
            fig, ax = plt.subplots(figsize=(10, 6))

            years = [comparison['previous_year'], comparison['year']]
            values = [
                comparison['emissions_change']['previous'],
                comparison['emissions_change']['current']
            ]

            bars = ax.bar(years, values, color=['#95A5A6', '#2ECC71'], width=0.6)
            ax.set_ylabel('Emisyonlar (tCO2e)', fontsize=12)
            ax.set_title('Yıllık Emisyon Karşılaştırması', fontsize=14, fontweight='bold')
            ax.grid(axis='y', alpha=0.3)

            # Değerleri bar üzerinde göster
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{height:,.0f}',
                       ha='center', va='bottom', fontsize=10)

            chart_path = os.path.join(output_dir, 'emissions_comparison.png')
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            chart_files['emissions_comparison'] = chart_path

        return chart_files

    def _create_docx_report(self, data: Dict, analysis: Dict, score: Dict, charts: Dict) -> Document:
        """DOCX raporu oluştur"""
        doc = Document()
        
        # Logo ekle
        if 'company_id' in data:
            self._add_logo(doc, data['company_id'])

        # Sayfa başlığı
        title = _add_turkish_heading(doc, 'CDP CLIMATE CHANGE RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        company_name = data['company_info'].get('name', 'Şirket')
        year = data.get('year', datetime.now().year)

        subtitle = _add_turkish_paragraph(doc, f'{company_name} - {year}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.italic = True

        _add_turkish_paragraph(doc, )  # Boş satır

        # 1. YÖNETİCİ ÖZETİ
        _add_turkish_heading(doc, '1. YÖNETİCİ ÖZETİ', 1)

        if analysis and 'summary' in analysis:
            _add_turkish_paragraph(doc, analysis['summary'])
        else:
            emissions = data.get('emissions', {})
            summary_text = f"""
Bu rapor, {company_name}'in {year} yılı sera gazı emisyonlarını ve iklim değişikliği 
performansını CDP (Carbon Disclosure Project) standartlarına uygun olarak sunmaktadır.

Toplam Emisyonlar: {emissions.get('total', 0):,.0f} tCO2e
- Scope 1 (Doğrudan): {emissions.get('scope1', 0):,.0f} tCO2e
- Scope 2 (Satın Alınan Enerji): {emissions.get('scope2', 0):,.0f} tCO2e
- Scope 3 (Değer Zinciri): {emissions.get('scope3', 0):,.0f} tCO2e
            """.strip()
            _add_turkish_paragraph(doc, summary_text)

        # 2. ŞİRKET BİLGİLERİ
        doc.add_page_break()
        _add_turkish_heading(doc, '2. ŞİRKET BİLGİLERİ', 1)

        company_table = doc.add_table(rows=1, cols=2)
        company_table.style = 'Light Grid Accent 1'

        company_info = data['company_info']
        rows = [
            ('Şirket Adı', company_info.get('name', '-')),
            ('Sektör', company_info.get('sector', '-')),
            ('Ülke', company_info.get('country', 'Türkiye')),
            ('Çalışan Sayısı', f"{company_info.get('employees', 0):,}"),
            ('Raporlama Yılı', str(year))
        ]

        for label, value in rows:
            r = company_table.add_row().cells
            r[0].text = label
            r[1].text = str(value)

        # 3. SERA GAZI EMİSYONLARI
        doc.add_page_break()
        _add_turkish_heading(doc, '3. SERA GAZI EMİSYONLARI', 1)

        _add_turkish_heading(doc, '3.1. Scope Bazında Emisyonlar', 2)

        emissions = data.get('emissions', {})
        emissions_table = doc.add_table(rows=1, cols=3)
        emissions_table.style = 'Light Grid Accent 1'

        # Başlıklar
        headers = emissions_table.rows[0].cells
        headers[0].text = 'Scope'
        headers[1].text = 'Emisyon (tCO2e)'
        headers[2].text = 'Oran (%)'

        total = emissions.get('total', 0)

        scope_data = [
            ('Scope 1', emissions.get('scope1', 0)),
            ('Scope 2', emissions.get('scope2', 0)),
            ('Scope 3', emissions.get('scope3', 0)),
            ('TOPLAM', total)
        ]

        for scope, value in scope_data:
            row = emissions_table.add_row().cells
            row[0].text = scope
            row[1].text = f"{value:,.2f}"
            if scope != 'TOPLAM' and total > 0:
                row[2].text = f"{(value/total)*100:.1f}%"
            else:
                row[2].text = '-'

        # Grafik ekle
        if 'emissions_pie' in charts:
            _add_turkish_paragraph(doc, )
            doc.add_picture(charts['emissions_pie'], width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 4. YILLIK KARŞILAŞTIRMA
        doc.add_page_break()
        _add_turkish_heading(doc, '4. YILLIK KARŞILAŞTIRMA VE TRENDLER', 1)

        comparison = data.get('comparison', {})
        if comparison and 'emissions_change' in comparison:
            change_data = comparison['emissions_change']
            pct_change = change_data.get('percentage_change', 0)

            trend_text = f"""
{year} yılında toplam emisyonlarımız {change_data.get('current', 0):,.0f} tCO2e olarak gerçekleşmiştir.
Bu, {comparison['previous_year']} yılına göre {'%' + str(abs(pct_change)) + ' azalma' if pct_change < 0 else '%' + str(abs(pct_change)) + ' artış'} anlamına gelmektedir.
            """
            _add_turkish_paragraph(doc, trend_text.strip())

            # Grafik ekle
            if 'emissions_comparison' in charts:
                _add_turkish_paragraph(doc, )
                doc.add_picture(charts['emissions_comparison'], width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            _add_turkish_paragraph(doc, "Önceki yıl verisi bulunmamaktadır.")

        # 5. PERFORMANS ANALİZİ
        if analysis and 'performance' in analysis:
            doc.add_page_break()
            _add_turkish_heading(doc, '5. PERFORMANS ANALİZİ', 1)
            _add_turkish_paragraph(doc, analysis['performance'])

        # 6. RİSKLER VE FIRSATLAR
        if analysis and ('risks' in analysis or 'opportunities' in analysis):
            doc.add_page_break()
            _add_turkish_heading(doc, '6. RİSKLER VE FIRSATLAR', 1)

            if analysis.get('risks'):
                _add_turkish_heading(doc, '6.1. İklim Riskleri', 2)
                _add_turkish_paragraph(doc, analysis['risks'])

            if analysis.get('opportunities'):
                _add_turkish_heading(doc, '6.2. Fırsatlar', 2)
                _add_turkish_paragraph(doc, analysis['opportunities'])

        # 7. ÖNERİLER
        if analysis and 'recommendations' in analysis:
            doc.add_page_break()
            _add_turkish_heading(doc, '7. ÖNERİLER VE EYLEM PLANI', 1)
            _add_turkish_paragraph(doc, analysis['recommendations'])

        # 8. CDP SKOR TAHMİNİ
        if score:
            doc.add_page_break()
            _add_turkish_heading(doc, '8. CDP SKOR TAHMİNİ', 1)

            score_text = f"""
Mevcut performansınıza göre CDP Climate Change skorunuz tahmini: {score.get('grade', 'D')}

Skorlama Detayları:
- Disclosure (Açıklama): {score.get('disclosure', 0)}/100
- Awareness (Farkındalık): {score.get('awareness', 0)}/100
- Management (Yönetim): {score.get('management', 0)}/100
- Leadership (Liderlik): {score.get('leadership', 0)}/100

Toplam Skor: {score.get('total', 0)}/100
            """
            _add_turkish_paragraph(doc, score_text.strip())

        # Altbilgi
        doc.add_page_break()
        footer = _add_turkish_paragraph(doc, )
        footer.add_run(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y")}\n').font.size = Pt(9)
        footer.add_run('Bu rapor Sustainage CDP Modülü tarafından otomatik olarak oluşturulmuştur.').font.size = Pt(9)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return doc

