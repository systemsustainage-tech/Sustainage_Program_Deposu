import logging
import json
import os
from datetime import datetime
from typing import Dict, List, Optional

def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    from docx.shared import Pt
    para = doc.add_paragraph(text if text is not None else '', style=style)
    if not text:
        return para
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return heading


class WaterReporting:
    """Su Yönetimi Raporlama"""

    def __init__(self, water_manager) -> None:
        self.water_manager = water_manager
        # Handle case where water_manager might not have company_id attribute directly exposed
        # or if it's passed differently. In waste_reporting it was waste_manager.company_id.
        # Assuming usage pattern will be similar.
        self.company_id = getattr(water_manager, 'company_id', None)

    def generate_water_footprint_report(self, company_id: int, period: str,
                                      include_scope3: bool = False) -> Optional[str]:
        """
        Su ayak izi raporu oluştur (DOCX)
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, Inches
        except ImportError:
            logging.info("[UYARI] python-docx yüklü değil. DOCX raporlar oluşturulamaz.")
            return None

        try:
            # Su ayak izi hesapla
            footprint_data = self.water_manager.calculate_water_footprint(company_id, period)
            # Tüketim kayıtlarını al
            consumption_records = self.water_manager.get_water_consumption(company_id, period)

            if not footprint_data:
                logging.warning("Su ayak izi verisi bulunamadı")
                return None

            # Rapor dosyası oluştur
            report_filename = f"su_ayak_izi_raporu_{company_id}_{period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            report_dir = "reports"
            os.makedirs(report_dir, exist_ok=True)
            report_path = os.path.join(report_dir, report_filename)

            # Dokuman oluştur
            doc = Document()

            # Başlık
            title = _add_turkish_heading(doc, 'SU AYAK İZİ RAPORU', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Alt başlık
            subtitle = _add_turkish_paragraph(doc, f"Raporlama Dönemi: {period}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(14)

            _add_turkish_paragraph(doc, f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y')}")
            _add_turkish_paragraph(doc, f"Raporlama Standardı: ISO 14046")

            doc.add_page_break()

            # 1. Yönetici Özeti
            _add_turkish_heading(doc, "1. YÖNETİCİ ÖZETİ", 1)
            _add_turkish_paragraph(doc,
                f"Bu rapor, {period} yılı su ayak izi envanterini ISO 14046 standartlarına uygun olarak sunmaktadır."
            )

            _add_turkish_paragraph(doc, f"\nToplam Su Ayak İzi: {footprint_data['total_water_footprint']:.2f} m³")
            _add_turkish_paragraph(doc, f"  • Mavi Su: {footprint_data['total_blue_water']:.2f} m³")
            _add_turkish_paragraph(doc, f"  • Yeşil Su: {footprint_data['total_green_water']:.2f} m³")
            _add_turkish_paragraph(doc, f"  • Gri Su: {footprint_data['total_grey_water']:.2f} m³")

            # 2. Metodoloji
            _add_turkish_heading(doc, "2. METODOLOJİ", 1)
            _add_turkish_paragraph(doc,
                "Su ayak izi hesaplamaları ISO 14046:2014 standardına uygun olarak yapılmıştır. Water Footprint Network metodolojisi kullanılmıştır."
            )

            # 3. Su Ayak İzi Sonuçları
            _add_turkish_heading(doc, "3. SU AYAK İZİ SONUÇLARI", 1)

            # 3.1 Toplam Sonuçlar
            _add_turkish_heading(doc, "3.1 Toplam Sonuçlar", 2)

            results_table = doc.add_table(rows=1, cols=3)
            results_table.style = 'Table Grid'

            # Tablo başlıkları
            hdr_cells = results_table.rows[0].cells
            hdr_cells[0].text = 'Su Türü'
            hdr_cells[1].text = 'Miktar (m³)'
            hdr_cells[2].text = 'Yüzde (%)'

            # Veriler
            total = footprint_data['total_water_footprint']
            blue_pct = (footprint_data['total_blue_water'] / total * 100) if total > 0 else 0
            green_pct = (footprint_data['total_green_water'] / total * 100) if total > 0 else 0
            grey_pct = (footprint_data['total_grey_water'] / total * 100) if total > 0 else 0

            data_rows = [
                ['Mavi Su', f"{footprint_data['total_blue_water']:.2f}", f"{blue_pct:.1f}%"],
                ['Yeşil Su', f"{footprint_data['total_green_water']:.2f}", f"{green_pct:.1f}%"],
                ['Gri Su', f"{footprint_data['total_grey_water']:.2f}", f"{grey_pct:.1f}%"],
                ['TOPLAM', f"{total:.2f}", "100.0%"]
            ]

            for row_data in data_rows:
                row_cells = results_table.add_row().cells
                for j, cell_data in enumerate(row_data):
                    row_cells[j].text = cell_data

            # 3.2 Tüketim Türüne Göre Dağılım
            _add_turkish_heading(doc, "3.2 Tüketim Türüne Göre Dağılım", 2)

            breakdown_by_type = footprint_data.get('breakdown_by_type', {})
            if breakdown_by_type:
                type_table = doc.add_table(rows=len(breakdown_by_type) + 1, cols=5)
                type_table.style = 'Table Grid'

                # Başlıklar
                type_hdr = type_table.rows[0].cells
                type_hdr[0].text = 'Tüketim Türü'
                type_hdr[1].text = 'Mavi Su (m³)'
                type_hdr[2].text = 'Yeşil Su (m³)'
                type_hdr[3].text = 'Gri Su (m³)'
                type_hdr[4].text = 'Toplam (m³)'

                # Veriler
                for i, (cons_type, data) in enumerate(breakdown_by_type.items(), 1):
                    type_row = type_table.rows[i].cells
                    type_row[0].text = cons_type.title()
                    type_row[1].text = f"{data['blue_water']:.2f}"
                    type_row[2].text = f"{data['green_water']:.2f}"
                    type_row[3].text = f"{data['grey_water']:.2f}"
                    type_row[4].text = f"{data['total']:.2f}"

            # 3.3 Tüketim Kayıtları
            if consumption_records:
                _add_turkish_heading(doc, "3.3 Tüketim Kayıtları", 2)
                
                records_table = doc.add_table(rows=1, cols=6)
                records_table.style = 'Table Grid'
                
                hdr_cells = records_table.rows[0].cells
                hdr_cells[0].text = 'Tüketim Türü'
                hdr_cells[1].text = 'Miktar (m³)'
                hdr_cells[2].text = 'Kaynak'
                hdr_cells[3].text = 'Fatura Tarihi'
                hdr_cells[4].text = 'Son Ödeme'
                hdr_cells[5].text = 'Tedarikçi'
                
                for record in consumption_records:
                    row_cells = records_table.add_row().cells
                    row_cells[0].text = record.get('consumption_type', '')
                    row_cells[1].text = f"{record.get('total_water', 0):.2f}"
                    row_cells[2].text = record.get('source_type', '')
                    row_cells[3].text = record.get('invoice_date', '') or '-'
                    row_cells[4].text = record.get('due_date', '') or '-'
                    row_cells[5].text = record.get('supplier', '') or '-'

            # 4. Verimlilik Metrikleri
            _add_turkish_heading(doc, "4. VERİMLİLİK METRİKLERİ", 1)

            efficiency_metrics = footprint_data.get('efficiency_metrics', {})
            if efficiency_metrics:
                avg_eff_ratio = efficiency_metrics.get('average_efficiency_ratio', 0)
                avg_recycle = efficiency_metrics.get('average_recycling_rate', 0)
                _add_turkish_paragraph(doc, f"Ortalama Verimlilik Oranı: {avg_eff_ratio:.3f}")
                _add_turkish_paragraph(doc, f"Ortalama Geri Dönüşüm Oranı: {avg_recycle:.3f}")
                _add_turkish_paragraph(doc, f"Toplam Kayıt Sayısı: {efficiency_metrics.get('total_records', 0)}")

            # 5. SDG 6 İlerlemesi
            _add_turkish_heading(doc, "5. SDG 6 İLERLEMESİ", 1)
            _add_turkish_paragraph(doc,
                "Su ayak izi yönetimi, Birleşmiş Milletler Sürdürülebilir Kalkınma Hedefi 6 (Temiz Su ve Sanitasyon) ile doğrudan ilgilidir."
            )

            # 6. Öneriler
            _add_turkish_heading(doc, "6. ÖNERİLER", 1)

            recommendations = [
                "Su tüketimini azaltmak için verimlilik projeleri uygulayın",
                "Gri su geri dönüşüm sistemleri kurun",
                "Su kalitesi izleme programlarını güçlendirin",
                "Çalışanlara su tasarrufu konusunda eğitim verin",
                "Su ayak izi hesaplamalarını düzenli olarak güncelleyin"
            ]

            for i, rec in enumerate(recommendations, 1):
                _add_turkish_paragraph(doc, f"{i}. {rec}")

            # Dokumanı kaydet
            doc.save(report_path)
            logging.info(f"[OK] Su ayak izi raporu oluşturuldu: {report_path}")

            # Veritabanına kaydet
            self.water_manager.save_water_report(
                company_id, 
                period, 
                "DOCX", 
                footprint_data['total_blue_water'],
                footprint_data['total_green_water'],
                footprint_data['total_grey_water'],
                footprint_data['total_water_footprint'],
                report_file_path=report_path
            )

            return report_path

        except Exception as e:
            logging.error(f"[HATA] Su raporu oluşturma hatası: {e}")
            return None

    def generate_excel_report(self, company_id: int, period: str) -> Optional[str]:
        """Excel su ayak izi raporu oluştur"""
        try:
            import openpyxl
            from openpyxl.styles import Alignment, Font, PatternFill
        except ImportError:
            logging.info("[UYARI] openpyxl yüklü değil. Excel raporlar oluşturulamaz.")
            return None

        try:
            # Su ayak izi hesapla
            footprint_data = self.water_manager.calculate_water_footprint(company_id, period)
            # Tüketim kayıtlarını al
            consumption_records = self.water_manager.get_water_consumption(company_id, period)

            if not footprint_data:
                logging.warning("Su ayak izi verisi bulunamadı")
                return None

            # Excel dosyası oluştur
            report_filename = f"su_ayak_izi_raporu_{company_id}_{period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            report_dir = "reports"
            os.makedirs(report_dir, exist_ok=True)
            report_path = os.path.join(report_dir, report_filename)

            # Excel workbook oluştur
            wb = openpyxl.Workbook()

            # Stil tanımları
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center")

            # 1. Özet Sayfası
            ws_summary = wb.active
            ws_summary.title = "Su Ayak İzi Özeti"

            # Başlık
            ws_summary['A1'] = f"SU AYAK İZİ RAPORU - {period}"
            ws_summary['A1'].font = Font(size=16, bold=True)
            ws_summary.merge_cells('A1:D1')

            # Tarih
            ws_summary['A2'] = f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y')}"
            ws_summary['A2'].font = Font(size=10)

            # Boş satır
            ws_summary['A3'] = ""

            # Toplam sonuçlar
            ws_summary['A4'] = "TOPLAM SU AYAK İZİ SONUÇLARI"
            ws_summary['A4'].font = Font(size=12, bold=True)

            # Tablo başlıkları
            headers = [
                "Su Türü", 
                "Miktar (m³)", 
                "Yüzde (%)"
            ]
            for col, header in enumerate(headers, 1):
                cell = ws_summary.cell(row=5, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment

            # Veriler
            total = footprint_data['total_water_footprint']
            blue_pct = (footprint_data['total_blue_water'] / total * 100) if total > 0 else 0
            green_pct = (footprint_data['total_green_water'] / total * 100) if total > 0 else 0
            grey_pct = (footprint_data['total_grey_water'] / total * 100) if total > 0 else 0

            data_rows = [
                ["Mavi Su", footprint_data['total_blue_water'], f"{blue_pct:.1f}%"],
                ["Yeşil Su", footprint_data['total_green_water'], f"{green_pct:.1f}%"],
                ["Gri Su", footprint_data['total_grey_water'], f"{grey_pct:.1f}%"],
                ["TOPLAM", total, "100.0%"]
            ]

            for row, data in enumerate(data_rows, 6):
                for col, value in enumerate(data, 1):
                    ws_summary.cell(row=row, column=col, value=value)

            # 2. Tüketim Türü Dağılımı Sayfası
            ws_breakdown = wb.create_sheet("Tüketim Türü Dağılımı")

            # Başlık
            ws_breakdown['A1'] = "TÜKETİM TÜRÜNE GÖRE DAĞILIM"
            ws_breakdown['A1'].font = Font(size=14, bold=True)
            ws_breakdown.merge_cells('A1:F1')

            # Tablo başlıkları
            breakdown_headers = [
                "Tüketim Türü",
                "Mavi Su (m³)",
                "Yeşil Su (m³)",
                "Gri Su (m³)",
                "Toplam (m³)"
            ]
            for col, header in enumerate(breakdown_headers, 1):
                cell = ws_breakdown.cell(row=3, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment

            # Veriler
            breakdown_by_type = footprint_data.get('breakdown_by_type', {})
            row = 4
            for cons_type, data in breakdown_by_type.items():
                ws_breakdown.cell(row=row, column=1, value=cons_type.title())
                ws_breakdown.cell(row=row, column=2, value=data['blue_water'])
                ws_breakdown.cell(row=row, column=3, value=data['green_water'])
                ws_breakdown.cell(row=row, column=4, value=data['grey_water'])
                ws_breakdown.cell(row=row, column=5, value=data['total'])
                row += 1

            # 3. Verimlilik Metrikleri Sayfası
            ws_efficiency = wb.create_sheet("Verimlilik Metrikleri")

            # Başlık
            ws_efficiency['A1'] = "VERİMLİLİK METRİKLERİ"
            ws_efficiency['A1'].font = Font(size=14, bold=True)
            ws_efficiency.merge_cells('A1:B1')

            # Metrikler
            efficiency_metrics = footprint_data.get('efficiency_metrics', {})
            metrics_data = [
                ["Ortalama Verimlilik Oranı", efficiency_metrics.get('average_efficiency_ratio', 0)],
                ["Ortalama Geri Dönüşüm Oranı", efficiency_metrics.get('average_recycling_rate', 0)],
                ["Toplam Kayıt Sayısı", efficiency_metrics.get('total_records', 0)],
                ["Yüksek Kalite Veri", efficiency_metrics.get('high_quality_data', 0)],
                ["Orta Kalite Veri", efficiency_metrics.get('medium_quality_data', 0)],
                ["Düşük Kalite Veri", efficiency_metrics.get('low_quality_data', 0)]
            ]

            row = 3
            for metric_name, metric_value in metrics_data:
                ws_efficiency.cell(row=row, column=1, value=metric_name)
                ws_efficiency.cell(row=row, column=2, value=metric_value)
                row += 1

            # 4. Tüketim Kayıtları Sayfası
            if consumption_records:
                ws_records = wb.create_sheet("Tüketim Kayıtları")
                
                # Başlık
                ws_records['A1'] = "SU TÜKETİM KAYITLARI"
                ws_records['A1'].font = Font(size=14, bold=True)
                ws_records.merge_cells('A1:H1')
                
                # Tablo başlıkları
                record_headers = [
                    "Tüketim Türü",
                    "Miktar (m³)",
                    "Kaynak",
                    "Konum",
                    "Maliyet",
                    "Fatura Tarihi",
                    "Son Ödeme Tarihi",
                    "Tedarikçi"
                ]
                
                for col, header in enumerate(record_headers, 1):
                    cell = ws_records.cell(row=3, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment
                    
                # Veriler
                row = 4
                for record in consumption_records:
                    ws_records.cell(row=row, column=1, value=record.get('consumption_type', ''))
                    ws_records.cell(row=row, column=2, value=record.get('total_water', 0))
                    ws_records.cell(row=row, column=3, value=record.get('source_type', ''))
                    ws_records.cell(row=row, column=4, value=record.get('location', ''))
                    ws_records.cell(row=row, column=5, value=record.get('cost', 0))
                    ws_records.cell(row=row, column=6, value=record.get('invoice_date', ''))
                    ws_records.cell(row=row, column=7, value=record.get('due_date', ''))
                    ws_records.cell(row=row, column=8, value=record.get('supplier', ''))
                    row += 1

            # Sütun genişliklerini ayarla
            sheets_to_adjust = [ws_summary, ws_breakdown, ws_efficiency]
            if 'ws_records' in locals():
                sheets_to_adjust.append(ws_records)

            for ws in sheets_to_adjust:
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except Exception as e:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width

            # Excel dosyasını kaydet
            wb.save(report_path)
            logging.info(f"[OK] Excel su raporu oluşturuldu: {report_path}")

            return report_path

        except Exception as e:
            logging.error(f"[HATA] Excel raporu oluşturma hatası: {e}")
            return None

    def export_water_data_to_csv(self, company_id: int, period: str, data_type: str = 'consumption') -> str:
        """Su verilerini CSV olarak dışa aktar"""
        try:
            import csv
            
            data = []
            filename = ""
            
            if data_type == 'consumption':
                data = self.water_manager.get_water_consumption(company_id, period)
                filename = f"su_tuketimi_{company_id}_{period}.csv"
            elif data_type == 'targets':
                data = self.water_manager.get_water_targets(company_id)
                filename = f"su_hedefleri_{company_id}_{period}.csv"
            elif data_type == 'projects':
                data = self.water_manager.get_efficiency_projects(company_id)
                filename = f"su_projeleri_{company_id}_{period}.csv"
            elif data_type == 'quality':
                data = self.water_manager.get_quality_measurements(company_id)
                filename = f"su_kalitesi_{company_id}_{period}.csv"
            else:
                raise ValueError("Geçersiz veri türü")

            if not data:
                logging.warning("Dışa aktarılacak veri bulunamadı")
                return ""

            # CSV dosya yolu
            report_dir = "reports"
            os.makedirs(report_dir, exist_ok=True)
            csv_path = os.path.join(report_dir, filename)

            # CSV oluştur
            with open(csv_path, 'w', newline='', encoding='utf-8-sig') as csvfile:
                if data:
                    fieldnames = data[0].keys()
                    writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                    writer.writeheader()
                    writer.writerows(data)
            
            logging.info(f"[OK] CSV raporu oluşturuldu: {csv_path}")
            return csv_path

        except Exception as e:
            logging.error(f"CSV dışa aktarma hatası: {e}")
            return ""
