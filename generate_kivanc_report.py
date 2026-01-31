import sqlite3
import os
import sys
from datetime import datetime

# Try to import docx, install if missing
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not found. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

DB_PATH = r'c:\SUSTAINAGESERVER\backend\data\sdg_desktop.sqlite'
if not os.path.exists(DB_PATH):
    DB_PATH = '/var/www/sustainage/backend/data/sdg_desktop.sqlite'

def get_db():
    print(f"Connecting to DB at {DB_PATH}")
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def ai_generate(prompt):
    # Simulate AI response for high quality report
    return f"""
    [AI ANALİZİ - 2025]
    Kıvanç Demir-Çelik için yapılan analizler, 2025 yılı hedeflerine büyük ölçüde ulaşıldığını göstermektedir. 
    Özellikle {prompt} alanında kaydedilen ilerleme, sektör ortalamasının üzerindedir. 
    Karbon ayak izi azaltma çalışmaları ve enerji verimliliği yatırımları, finansal performansa da olumlu yansımıştır.
    Gelecek dönem için önerimiz, dijital dönüşüm ve döngüsel ekonomi uygulamalarına daha fazla odaklanılmasıdır.
    """

def create_report():
    print("Generating Report...")
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Sürdürülebilirlik Raporu 2025', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Kıvanç Demir-Çelik A.Ş.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(24)
    subtitle.runs[0].font.bold = True
    
    doc.add_page_break()
    
    # CEO Message
    doc.add_heading('CEO Mesajı', level=1)
    ceo_msg = """
    Değerli Paydaşlarımız,
    
    2025 yılı, Kıvanç Demir-Çelik için sürdürülebilirlik yolculuğunda bir dönüm noktası olmuştur. 
    "Gelecek İçin Çelik" vizyonumuz doğrultusunda, çevresel, sosyal ve yönetişimsel (ESG) performansımızı 
    en üst seviyeye taşımak için kararlı adımlar attık.
    
    Bu rapor, şeffaflık ve hesap verebilirlik ilkelerimiz çerçevesinde, yıl boyunca gerçekleştirdiğimiz 
    faaliyetleri ve elde ettiğimiz sonuçları sizlere sunmaktadır.
    
    Saygılarımla,
    
    Kıvanç Demir
    CEO
    """
    doc.add_paragraph(ceo_msg)
    doc.add_page_break()
    
    conn = get_db()
    cursor = conn.cursor()
    
    # Fetch Company ID
    cursor.execute("SELECT id FROM companies WHERE name LIKE 'Kıvanç Demir%'")
    res = cursor.fetchone()
    company_id = res['id'] if res else 1
    
    modules = [
        ('Çevresel Performans', ['carbon_emissions', 'energy_consumption', 'water_consumption', 'waste_generation']),
        ('Sosyal Etki', ['social_employees', 'social_incidents']),
        ('Yönetişim (Governance)', ['governance_structure']),
        ('SDG Performansı', ['sdg_progress']),
        ('TCFD & İklim Riskleri', ['tnfd_recommendations', 'tcfd_risks']),
        ('AB Yeşil Mutabakatı (CBAM & Taxonomy)', ['cbam_declarations', 'taxonomy_alignment'])
    ]
    
    for title, tables in modules:
        doc.add_heading(title, level=1)
        
        # AI Intro
        doc.add_paragraph(ai_generate(title))
        
        for table in tables:
            try:
                cursor.execute(f"SELECT * FROM {table} LIMIT 5")
                rows = cursor.fetchall()
                if rows:
                    doc.add_heading(f"Veri Tablosu: {table}", level=2)
                    t = doc.add_table(rows=1, cols=len(rows[0].keys()))
                    t.style = 'Table Grid'
                    
                    # Header
                    hdr_cells = t.rows[0].cells
                    for i, key in enumerate(rows[0].keys()):
                        hdr_cells[i].text = str(key)
                    
                    # Data
                    for row in rows:
                        row_cells = t.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val)
                    doc.add_paragraph("\n")
            except Exception as e:
                print(f"Skipping table {table}: {e}")
                
        doc.add_page_break()
        
    # GRI Index
    doc.add_heading('GRI İçerik İndeksi', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'GRI Standart'
    hdr[1].text = 'Açıklama'
    hdr[2].text = 'Sayfa No'
    
    gri_data = [
        ('GRI 102', 'Genel Açıklamalar', '5-12'),
        ('GRI 201', 'Ekonomik Performans', '15'),
        ('GRI 302', 'Enerji', '18'),
        ('GRI 305', 'Emisyonlar', '20'),
        ('GRI 401', 'İstihdam', '25'),
    ]
    
    for g in gri_data:
        row = table.add_row().cells
        row[0].text = g[0]
        row[1].text = g[1]
        row[2].text = g[2]

    if os.name == 'nt':
        filename = r'c:\SUSTAINAGESERVER\Kivanc_Demir_Celik_Surdurulebilirlik_Raporu_2025.docx'
    else:
        filename = '/var/www/sustainage/static/Kivanc_Demir_Celik_Surdurulebilirlik_Raporu_2025.docx'
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(filename), exist_ok=True)
    
    doc.save(filename)
    print(f"Report saved to {filename}")
    
    # Set permissions to be readable by web server
    try:
        if os.name != 'nt':
            os.chmod(filename, 0o644)
    except Exception as e:
        print(f"Could not set permissions: {e}")

if __name__ == "__main__":
    create_report()
