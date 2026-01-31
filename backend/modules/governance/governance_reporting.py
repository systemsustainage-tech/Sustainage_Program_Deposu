import logging
import os
import sqlite3
from datetime import datetime
from typing import Dict, List, Optional
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from docx import Document
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    if text is None: text = ""
    para = doc.add_paragraph(str(text), style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        try:
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        try:
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return heading

class GovernanceReporting:
    """Kurumsal Yönetim Raporlama"""

    def __init__(self, governance_manager) -> None:
        self.manager = governance_manager
        self.db_path = governance_manager.db_path

    def generate_governance_report(self, company_id: int, period: str,
                                 formats: List[str] = None) -> Dict[str, str]:
        """Yönetişim raporu oluştur"""
        if formats is None:
            formats = ['docx']
            
        generated_files = {}
        
        # DOCX Raporu
        if 'docx' in formats:
            docx_path = self._generate_docx_report(company_id, period)
            if docx_path:
                generated_files['docx'] = docx_path
                
        return generated_files

    def _add_logo(self, doc, company_id: int):
        """Raporun başına şirket logosunu ekle"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT logo_path FROM company_profiles WHERE company_id = ?", (company_id,))
            result = cursor.fetchone()
            conn.close()
            
            logo_path = None
            if result and result[0] and os.path.exists(result[0]):
                logo_path = result[0]
            else:
                data_dir = os.path.dirname(self.db_path)
                possible_path = os.path.join(data_dir, "company_logos", f"company_{company_id}_logo.png")
                if os.path.exists(possible_path):
                    logo_path = possible_path
                else:
                    possible_path = possible_path.replace(".png", ".jpg")
                    if os.path.exists(possible_path):
                        logo_path = possible_path

            if logo_path:
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, height=Cm(2.5))
        except Exception as e:
            logging.warning(f"Logo eklenirken hata: {e}")

    def _generate_docx_report(self, company_id: int, period: str) -> Optional[str]:
        if not DOCX_AVAILABLE:
            return None
            
        try:
            doc = Document()
            self._add_logo(doc, company_id)
            
            # Başlık
            title = _add_turkish_heading(doc, 'Kurumsal Yönetim Raporu', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            _add_turkish_paragraph(doc, f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y')}")
            _add_turkish_paragraph(doc, f"Dönem: {period or 'Tümü'}")
            
            # 1. Yönetim Kurulu
            doc.add_page_break()
            _add_turkish_heading(doc, '1. Yönetim Kurulu Yapısı', 1)
            
            conn = sqlite3.connect(self.db_path)
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            # Board Members
            cursor.execute("""
                SELECT member_name, position, member_type, independence_status, gender 
                FROM board_members 
                WHERE company_id = ? AND status = 'active'
            """, (company_id,))
            members = cursor.fetchall()
            
            if members:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Ad Soyad'
                hdr[1].text = 'Pozisyon'
                hdr[2].text = 'Üyelik Tipi'
                hdr[3].text = 'Bağımsızlık'
                hdr[4].text = 'Cinsiyet'
                
                for m in members:
                    row = table.add_row().cells
                    row[0].text = str(m['member_name'])
                    row[1].text = str(m['position'])
                    row[2].text = str(m['member_type'])
                    row[3].text = str(m['independence_status'] or '-')
                    row[4].text = str(m['gender'] or '-')
            else:
                _add_turkish_paragraph(doc, "Kayıtlı yönetim kurulu üyesi bulunmamaktadır.")

            # 2. Komiteler
            _add_turkish_heading(doc, '2. Yönetim Kurulu Komiteleri', 1)
            cursor.execute("""
                SELECT committee_name, committee_type, chair_person, member_count 
                FROM governance_committees 
                WHERE company_id = ? AND status = 'active'
            """, (company_id,))
            committees = cursor.fetchall()
            
            if committees:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Komite Adı'
                hdr[1].text = 'Tip'
                hdr[2].text = 'Başkan'
                hdr[3].text = 'Üye Sayısı'
                
                for c in committees:
                    row = table.add_row().cells
                    row[0].text = str(c['committee_name'])
                    row[1].text = str(c['committee_type'])
                    row[2].text = str(c['chair_person'] or '-')
                    row[3].text = str(c['member_count'] or '-')
            else:
                _add_turkish_paragraph(doc, "Kayıtlı komite bulunmamaktadır.")
                
            # 3. Politikalar
            _add_turkish_heading(doc, '3. Kurumsal Politikalar', 1)
            cursor.execute("""
                SELECT policy_name, version, effective_date, review_date 
                FROM governance_policies 
                WHERE company_id = ? AND status = 'active'
            """, (company_id,))
            policies = cursor.fetchall()
            
            if policies:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Politika'
                hdr[1].text = 'Versiyon'
                hdr[2].text = 'Yürürlük Tarihi'
                hdr[3].text = 'Gözden Geçirme'
                
                for p in policies:
                    row = table.add_row().cells
                    row[0].text = str(p['policy_name'])
                    row[1].text = str(p['version'] or '-')
                    row[2].text = str(p['effective_date'] or '-')
                    row[3].text = str(p['review_date'] or '-')
            else:
                _add_turkish_paragraph(doc, "Kayıtlı politika bulunmamaktadır.")

            conn.close()
            
            # Kaydet
            os.makedirs("reports", exist_ok=True)
            filename = f"yonetisim_raporu_{company_id}_{period or 'all'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join("reports", filename)
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            logging.error(f"Governance report generation error: {e}")
            return None
