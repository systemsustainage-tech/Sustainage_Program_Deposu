#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GELİŞMİŞ RAPORLAMA SİSTEMİ
==========================

Özelleştirilebilir rapor şablonları ve gelişmiş raporlama
"""

import logging
import json
import os
import sqlite3
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Dict, List

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, PageBreak

# Try to import font registration, if fails, define dummy
try:
    from ..reporting.font_utils import register_turkish_fonts_reportlab
except ImportError:
    def register_turkish_fonts_reportlab():
        pass

# Çeviri Sözlüğü
TRANSLATIONS = {
    'tr': {
        'company': 'Şirket',
        'period': 'Dönem',
        'report_date': 'Rapor Tarihi',
        'sdg_title': 'SDG İLERLEME RAPORU',
        'gri_title': 'GRI SÜRDÜRÜLEBİLİRLİK RAPORU',
        'carbon_title': 'KARBON AYAK İZİ RAPORU',
        'cbam_title': 'CBAM RAPORU',
        'tcfd_title': 'TCFD İKLİM RAPORU',
        'unknown': 'Bilinmeyen',
        'total_emissions': 'Toplam Emisyon',
        'scope': 'Kapsam',
        'emissions': 'Emisyon',
        'percentage': 'Yüzde',
        'summary': 'Rapor Özeti',
        'completed': 'Tamamlandı',
        'in_progress': 'Devam Ediyor',
        'started': 'Başlangıç',
        'goal': 'Hedef',
        'progress': 'İlerleme',
        'status': 'Durum',
        'notes': 'Notlar',
        'gri_topics': 'GRI Konuları',
        'description': 'Açıklama',
        'methodology': 'Metodoloji',
        'scope_breakdown': 'Emisyon Kapsamları',
        'cbam_declarant': 'Bildirim Yapan',
        'cbam_embedded': 'Gömülü Emisyonlar',
        'cbam_price': 'Karbon Fiyatı',
        'tcfd_governance': 'Yönetişim',
        'tcfd_strategy': 'Strateji',
        'tcfd_risk': 'Risk Yönetimi',
        'tcfd_metrics': 'Metrikler ve Hedefler',
        'data_source': 'Veri Kaynağı',
        'standard_ref': 'Standart Referans'
    },
    'en': {
        'company': 'Company',
        'period': 'Period',
        'report_date': 'Report Date',
        'sdg_title': 'SDG PROGRESS REPORT',
        'gri_title': 'GRI SUSTAINABILITY REPORT',
        'carbon_title': 'CARBON FOOTPRINT REPORT',
        'cbam_title': 'CBAM REPORT',
        'tcfd_title': 'TCFD CLIMATE REPORT',
        'unknown': 'Unknown',
        'total_emissions': 'Total Emissions',
        'scope': 'Scope',
        'emissions': 'Emissions',
        'percentage': 'Percentage',
        'summary': 'Report Summary',
        'completed': 'Completed',
        'in_progress': 'In Progress',
        'started': 'Started',
        'goal': 'Goal',
        'progress': 'Progress',
        'status': 'Status',
        'notes': 'Notes',
        'gri_topics': 'GRI Topics',
        'description': 'Description',
        'methodology': 'Methodology',
        'scope_breakdown': 'Emission Scopes',
        'cbam_declarant': 'Declarant',
        'cbam_embedded': 'Embedded Emissions',
        'cbam_price': 'Carbon Price',
        'tcfd_governance': 'Governance',
        'tcfd_strategy': 'Strategy',
        'tcfd_risk': 'Risk Management',
        'tcfd_metrics': 'Metrics and Targets'
    },
    'de': {
        'company': 'Unternehmen',
        'period': 'Zeitraum',
        'report_date': 'Berichtsdatum',
        'sdg_title': 'SDG FORTSCHRITTSBERICHT',
        'gri_title': 'GRI NACHHALTIGKEITSBERICHT',
        'carbon_title': 'CO2-FUSSABDRUCK BERICHT',
        'cbam_title': 'CBAM BERICHT',
        'tcfd_title': 'TCFD KLIMABERICHT',
        'unknown': 'Unbekannt',
        'total_emissions': 'Gesamtemissionen',
        'scope': 'Umfang',
        'emissions': 'Emissionen',
        'percentage': 'Prozentsatz',
        'summary': 'Berichtszusammenfassung',
        'completed': 'Abgeschlossen',
        'in_progress': 'In Bearbeitung',
        'started': 'Begonnen',
        'goal': 'Ziel',
        'progress': 'Fortschritt',
        'status': 'Status',
        'notes': 'Notizen',
        'gri_topics': 'GRI Themen',
        'description': 'Beschreibung',
        'methodology': 'Methodik',
        'scope_breakdown': 'Emissionsbereiche',
        'cbam_declarant': 'Anmelder',
        'cbam_embedded': 'Graue Emissionen',
        'cbam_price': 'CO2-Preis',
        'tcfd_governance': 'Unternehmensführung',
        'tcfd_strategy': 'Strategie',
        'tcfd_risk': 'Risikomanagement',
        'tcfd_metrics': 'Kennzahlen und Ziele',
        'data_source': 'Datenquelle',
        'standard_ref': 'Standardreferenz'
    }
}

@dataclass
class ReportTemplate:
    """Rapor şablonu veri yapısı"""
    template_id: str
    template_name: str
    template_type: str  # pdf, docx, xlsx, html
    category: str  # sdg, gri, carbon, cbam, tcfd, esrs, etc.
    language: str
    template_config: Dict[str, Any]
    is_active: bool
    created_at: str

@dataclass
class ReportSection:
    """Rapor bölümü veri yapısı"""
    section_id: str
    template_id: str  # Added template_id field
    section_name: str
    section_type: str  # header, content, table, chart, footer
    content: str
    order: int
    is_required: bool

def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    from docx.shared import Pt
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        from docx.oxml.ns import qn
        r = run._element
        r.rPr.rFonts.set(qn('w:ascii'), font_name)
        r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        r.rPr.rFonts.set(qn('w:cs'), font_name)
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        from docx.oxml.ns import qn
        r = run._element
        r.rPr.rFonts.set(qn('w:ascii'), font_name)
        r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        r.rPr.rFonts.set(qn('w:cs'), font_name)
    return heading


class AdvancedReportTemplates:
    """Gelişmiş rapor şablonları yöneticisi"""

    def __init__(self, db_path: str = None):
        self.db_path = db_path or 'data/sdg_desktop.db'
        self.templates = {}
        self._create_tables()
        self._load_default_templates()

    def _create_tables(self):
        """Raporlama tablolarını oluştur"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            # Rapor şablonları
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS report_templates (
                    id TEXT PRIMARY KEY,
                    template_name TEXT NOT NULL,
                    template_type TEXT NOT NULL,
                    category TEXT NOT NULL,
                    language TEXT DEFAULT 'tr',
                    template_config TEXT NOT NULL,
                    is_active INTEGER DEFAULT 1,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Rapor bölümleri
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS report_sections (
                    id TEXT PRIMARY KEY,
                    template_id TEXT NOT NULL,
                    section_name TEXT NOT NULL,
                    section_type TEXT NOT NULL,
                    content TEXT NOT NULL,
                    order_index INTEGER DEFAULT 0,
                    is_required INTEGER DEFAULT 1,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY(template_id) REFERENCES report_templates(id)
                )
            """)

            # Rapor oluşturma geçmişi
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS report_generation_log (
                    id TEXT PRIMARY KEY,
                    company_id INTEGER NOT NULL,
                    template_id TEXT NOT NULL,
                    report_name TEXT NOT NULL,
                    file_path TEXT NOT NULL,
                    generation_time REAL,
                    status TEXT DEFAULT 'success',
                    error_message TEXT,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY(company_id) REFERENCES companies(id),
                    FOREIGN KEY(template_id) REFERENCES report_templates(id)
                )
            """)

            # Rapor özelleştirmeleri
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS report_customizations (
                    id TEXT PRIMARY KEY,
                    company_id INTEGER NOT NULL,
                    template_id TEXT NOT NULL,
                    custom_config TEXT NOT NULL,
                    is_active INTEGER DEFAULT 1,
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY(company_id) REFERENCES companies(id),
                    FOREIGN KEY(template_id) REFERENCES report_templates(id)
                )
            """)

            conn.commit()
            conn.close()

        except Exception as e:
            logging.error(f"[HATA] Raporlama tabloları oluşturulamadı: {e}")

    def _load_default_templates(self):
        """Varsayılan şablonları yükle (TR, EN, DE)"""
        try:
            languages = ['tr', 'en', 'de']
            categories = [
                {'cat': 'sdg', 'name_tr': 'SDG İlerleme Raporu', 'name_en': 'SDG Progress Report', 'name_de': 'SDG Fortschrittsbericht', 'type': 'pdf'},
                {'cat': 'gri', 'name_tr': 'GRI Sürdürülebilirlik Raporu', 'name_en': 'GRI Sustainability Report', 'name_de': 'GRI Nachhaltigkeitsbericht', 'type': 'docx'},
                {'cat': 'carbon', 'name_tr': 'Karbon Ayak İzi Raporu', 'name_en': 'Carbon Footprint Report', 'name_de': 'CO2-Fußabdruck Bericht', 'type': 'pdf'},
                {'cat': 'cbam', 'name_tr': 'CBAM Raporu', 'name_en': 'CBAM Report', 'name_de': 'CBAM Bericht', 'type': 'pdf'},
                {'cat': 'tcfd', 'name_tr': 'TCFD İklim Raporu', 'name_en': 'TCFD Climate Report', 'name_de': 'TCFD Klimabericht', 'type': 'pdf'}
            ]

            base_config = {
                'page_size': 'A4',
                'orientation': 'portrait',
                'margins': {'top': 1, 'bottom': 1, 'left': 1, 'right': 1},
                'header': True,
                'footer': True,
                'page_numbers': True
            }

            for lang in languages:
                for cat_info in categories:
                    # PDF/DOCX Template
                    template_id = f"{cat_info['cat']}_template_{lang}"
                    template_name = cat_info.get(f'name_{lang}', cat_info['name_en'])
                    
                    template = ReportTemplate(
                        template_id=template_id,
                        template_name=template_name,
                        template_type=cat_info['type'],
                        category=cat_info['cat'],
                        language=lang,
                        template_config=base_config,
                        is_active=True,
                        created_at=datetime.now().isoformat()
                    )
                    self.templates[template_id] = template
                    self._save_template(template)
                    self._load_default_sections(template_id, lang)

                    # HTML Preview Template
                    html_template_id = f"{cat_info['cat']}_template_{lang}_html"
                    html_template = ReportTemplate(
                        template_id=html_template_id,
                        template_name=f"{template_name} (Web/HTML)",
                        template_type="html",
                        category=cat_info['cat'],
                        language=lang,
                        template_config=base_config,
                        is_active=True,
                        created_at=datetime.now().isoformat()
                    )
                    self.templates[html_template_id] = html_template
                    self._save_template(html_template)
                    self._load_default_sections(html_template_id, lang)

        except Exception as e:
            logging.error(f"[HATA] Varsayılan şablonlar yüklenemedi: {e}")

    def _load_default_sections(self, template_id: str, language: str):
        """Varsayılan zorunlu bölümleri ekle"""
        t = TRANSLATIONS.get(language, TRANSLATIONS['tr'])
        
        mandatory_sections = [
            {'id': 'data_source', 'name': t.get('data_source', 'Veri Kaynağı'), 'order': 900},
            {'id': 'methodology', 'name': t.get('methodology', 'Metodoloji'), 'order': 910},
            {'id': 'standard_ref', 'name': t.get('standard_ref', 'Standart Referans'), 'order': 920}
        ]
        
        for sec in mandatory_sections:
            section_id = f"{template_id}_{sec['id']}"
            section = ReportSection(
                section_id=section_id,
                template_id=template_id,
                section_name=sec['name'],
                section_type='text',
                content=f"[{sec['name']} İçeriği]",
                order=sec['order'],
                is_required=True
            )
            self._save_section(section)

    def _save_section(self, section: ReportSection):
        """Bölümü kaydet"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                INSERT OR REPLACE INTO report_sections
                (id, template_id, section_name, section_type, content, order_index, is_required)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
                section.section_id, section.template_id, section.section_name,
                section.section_type, section.content, section.order,
                1 if section.is_required else 0
            ))
            conn.commit()
            conn.close()
        except Exception as e:
            logging.error(f"[HATA] Bölüm kaydedilemedi: {e}")

    def delete_section(self, section_id: str) -> bool:
        """Bölümü sil (Zorunlu bölümler silinemez)"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Kontrol et
            cursor.execute("SELECT is_required FROM report_sections WHERE id = ?", (section_id,))
            result = cursor.fetchone()
            
            if not result:
                logging.warning(f"Silinecek bölüm bulunamadı: {section_id}")
                return False
                
            if result[0] == 1:
                logging.warning(f"Zorunlu bölüm silinemez: {section_id}")
                return False
                
            cursor.execute("DELETE FROM report_sections WHERE id = ?", (section_id,))
            conn.commit()
            conn.close()
            return True
            
        except Exception as e:
            logging.error(f"[HATA] Bölüm silinemedi: {e}")
            return False

    def _save_template(self, template: ReportTemplate):
        """Şablonu kaydet"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute("""
                INSERT OR REPLACE INTO report_templates
                (id, template_name, template_type, category, language, template_config, is_active, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                template.template_id, template.template_name, template.template_type,
                template.category, template.language, json.dumps(template.template_config),
                template.is_active, template.created_at
            ))

            conn.commit()
            conn.close()

        except Exception as e:
            logging.error(f"[HATA] Şablon kaydedilemedi: {e}")

    def generate_report(self, template_id: str, company_id: int, period: str, data: Dict[str, Any]) -> str:
        """Şablon ID'sine göre rapor oluştur"""
        try:
            if template_id not in self.templates:
                # Veritabanından yüklemeyi dene (hafızada yoksa)
                # Şimdilik varsayalım ki load_default_templates her şeyi yüklüyor
                # veya get_available_templates kullanıcısı bu ID'yi oradan aldı.
                # Basitçe self.templates'e güvenelim.
                logging.error(f"Şablon bulunamadı: {template_id}")
                return ""

            template = self.templates[template_id]
            lang = template.language
            
            if template.template_type == 'html':
                return self.create_html_report(template, company_id, period, data)
            
            if template.category == 'sdg':
                return self.create_sdg_report(company_id, period, data, lang)
            elif template.category == 'gri':
                return self.create_gri_report(company_id, period, data, lang)
            elif template.category == 'carbon':
                return self.create_carbon_report(company_id, period, data, lang)
            elif template.category == 'cbam':
                return self.create_cbam_report(company_id, period, data, lang)
            elif template.category == 'tcfd':
                return self.create_tcfd_report(company_id, period, data, lang)
            else:
                logging.error(f"Desteklenmeyen kategori: {template.category}")
                return ""

        except Exception as e:
            logging.error(f"[HATA] Rapor oluşturma hatası ({template_id}): {e}")
            return ""

    def create_html_report(self, template: ReportTemplate, company_id: int, period: str, data: Dict[str, Any]) -> str:
        """HTML formatında rapor (önizleme) oluştur"""
        try:
            report_name = f"{template.category.upper()}_{company_id}_{period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            file_path = f"raporlar/html/{report_name}.html"
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            t = TRANSLATIONS.get(template.language, TRANSLATIONS['tr'])
            
            # Basit HTML Şablonu
            html_content = f"""
            <!DOCTYPE html>
            <html lang="{template.language}">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{template.template_name}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 20px; }}
                    h1 {{ color: #2c3e50; border-bottom: 2px solid #eee; padding-bottom: 10px; }}
                    h2 {{ color: #34495e; margin-top: 30px; }}
                    .meta {{ background: #f9f9f9; padding: 15px; border-radius: 5px; margin-bottom: 30px; }}
                    table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
                    th, td {{ padding: 12px; border: 1px solid #ddd; text-align: left; }}
                    th {{ background-color: #f2f2f2; font-weight: bold; }}
                    .status-badge {{ padding: 4px 8px; border-radius: 4px; font-size: 0.9em; }}
                    .status-complete {{ background: #d4edda; color: #155724; }}
                    .status-progress {{ background: #fff3cd; color: #856404; }}
                </style>
            </head>
            <body>
                <h1>{template.template_name}</h1>
                
                <div class="meta">
                    <p><strong>{t['company']}:</strong> {data.get('company_name', t['unknown'])}</p>
                    <p><strong>{t['period']}:</strong> {period}</p>
                    <p><strong>{t['report_date']}:</strong> {datetime.now().strftime('%d.%m.%Y')}</p>
                </div>

                <div class="content">
            """

            # İçerik Ekleme (Kategoriye Göre)
            if template.category == 'sdg':
                html_content += f"<h2>{t['sdg_title']}</h2>"
                if 'sdg_data' in data:
                    html_content += f"<table><thead><tr><th>{t['goal']}</th><th>{t['progress']}</th><th>{t['status']}</th></tr></thead><tbody>"
                    for goal, progress in data['sdg_data'].items():
                        status_class = "status-complete" if progress >= 80 else "status-progress"
                        status_text = t['completed'] if progress >= 80 else t['in_progress']
                        html_content += f"<tr><td>{goal}</td><td>%{progress:.1f}</td><td><span class='status-badge {status_class}'>{status_text}</span></td></tr>"
                    html_content += "</tbody></table>"
            
            elif template.category == 'carbon' or template.category == 'cbam':
                title = t['carbon_title'] if template.category == 'carbon' else t['cbam_title']
                html_content += f"<h2>{title}</h2>"
                if 'emissions' in data:
                    emissions = data['emissions']
                    html_content += f"<p><strong>{t['total_emissions']}:</strong> {emissions.get('total', 0):.2f} tCO2e</p>"
                    if 'scope_breakdown' in emissions:
                        html_content += f"<h3>{t['scope_breakdown']}</h3><ul>"
                        for scope, val in emissions['scope_breakdown'].items():
                            html_content += f"<li>{scope.upper()}: {val:.2f} tCO2e</li>"
                        html_content += "</ul>"

            elif template.category == 'gri':
                html_content += f"<h2>{t['gri_title']}</h2>"
                if 'gri_topics' in data:
                    for topic, details in data['gri_topics'].items():
                        html_content += f"<h3>{topic}</h3>"
                        html_content += f"<p><strong>{t['description']}:</strong> {details.get('description', '-')}</p>"
                        html_content += f"<p><strong>{t['methodology']}:</strong> {details.get('methodology', '-')}</p>"

            elif template.category == 'tcfd':
                html_content += f"<h2>{t['tcfd_title']}</h2>"
                sections = ['governance', 'strategy', 'risk', 'metrics']
                for sec in sections:
                    key = f"tcfd_{sec}"
                    html_content += f"<h3>{t.get(key, sec.title())}</h3>"
                    html_content += f"<p>{data.get(sec, '-')}</p>"

            html_content += """
                </div>
            </body>
            </html>
            """

            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self._save_generation_log(company_id, template.template_id, report_name, file_path)
            logging.info(f"[OK] HTML rapor oluşturuldu: {file_path}")
            return file_path

        except Exception as e:
            logging.error(f"[HATA] HTML rapor oluşturulamadı: {e}")
            return ""

    def create_sdg_report(self, company_id: int, period: str, data: Dict[str, Any], language: str = 'tr') -> str:
        """SDG raporu oluştur"""
        try:
            t = TRANSLATIONS.get(language, TRANSLATIONS['tr'])
            template_id = f"sdg_template_{language}"
            report_name = f"SDG_{company_id}_{period}_{language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            file_path = f"raporlar/sdg/{report_name}.pdf"
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            register_turkish_fonts_reportlab()

            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=19, spaceAfter=30, alignment=TA_CENTER, fontName='NotoSans')
            story.append(Paragraph(t['sdg_title'], title_style))
            story.append(Spacer(1, 12))

            company_info = f"""
            <b>{t['company']}:</b> {data.get('company_name', t['unknown'])}<br/>
            <b>{t['period']}:</b> {period}<br/>
            <b>{t['report_date']}:</b> {datetime.now().strftime('%d.%m.%Y')}
            """
            story.append(Paragraph(company_info, styles['Normal']))
            story.append(Spacer(1, 20))

            if 'sdg_data' in data:
                sdg_data = data['sdg_data']
                table_data = [[t['goal'], t['progress'], t['status'], t['notes']]]

                for goal, progress in sdg_data.items():
                    status = t['completed'] if progress >= 80 else t['in_progress'] if progress >= 50 else t['started']
                    table_data.append([goal, f"{progress:.1f}%", status, ""])

                table = Table(table_data, colWidths=[2*inch, 1*inch, 1.5*inch, 2*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)

            if 'summary' in data:
                summary = data['summary']
                summary_text = f"""
                <b>{t['summary']}:</b><br/>
                • {t['total_emissions']}: {len(sdg_data) if 'sdg_data' in data else 0}<br/>
                • {t['completed']}: {summary.get('completed', 0)}
                """
                story.append(Spacer(1, 20))
                story.append(Paragraph(summary_text, styles['Normal']))

            doc.build(story)
            self._save_generation_log(company_id, template_id, report_name, file_path)
            logging.info(f"[OK] SDG raporu oluşturuldu ({language}): {file_path}")
            return file_path

        except Exception as e:
            logging.error(f"[HATA] SDG raporu oluşturulamadı: {e}")
            return ""

    def create_gri_report(self, company_id: int, period: str, data: Dict[str, Any], language: str = 'tr') -> str:
        """GRI raporu oluştur"""
        try:
            t = TRANSLATIONS.get(language, TRANSLATIONS['tr'])
            template_id = f"gri_template_{language}"
            report_name = f"GRI_{company_id}_{period}_{language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            file_path = f"raporlar/gri/{report_name}.docx"
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            doc = Document()
            title = _add_turkish_heading(doc, t['gri_title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            _add_turkish_paragraph(doc, f"{t['company']}: {data.get('company_name', t['unknown'])}")
            _add_turkish_paragraph(doc, f"{t['period']}: {period}")
            _add_turkish_paragraph(doc, f"{t['report_date']}: {datetime.now().strftime('%d.%m.%Y')}")
            _add_turkish_paragraph(doc, "")

            if 'gri_topics' in data:
                _add_turkish_heading(doc, t['gri_topics'], level=1)
                for topic, details in data['gri_topics'].items():
                    _add_turkish_heading(doc, topic, level=2)
                    _add_turkish_paragraph(doc, f"{t['description']}: {details.get('description', '')}")
                    _add_turkish_paragraph(doc, f"{t['scope']}: {details.get('scope', '')}")
                    _add_turkish_paragraph(doc, f"{t['methodology']}: {details.get('methodology', '')}")
                    _add_turkish_paragraph(doc, "")

            doc.save(file_path)
            self._save_generation_log(company_id, template_id, report_name, file_path)
            logging.info(f"[OK] GRI raporu oluşturuldu ({language}): {file_path}")
            return file_path

        except Exception as e:
            logging.error(f"[HATA] GRI raporu oluşturulamadı: {e}")
            return ""

    def create_carbon_report(self, company_id: int, period: str, data: Dict[str, Any], language: str = 'tr') -> str:
        """Karbon raporu oluştur"""
        try:
            t = TRANSLATIONS.get(language, TRANSLATIONS['tr'])
            template_id = f"carbon_template_{language}"
            report_name = f"Carbon_{company_id}_{period}_{language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            file_path = f"raporlar/carbon/{report_name}.pdf"
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            register_turkish_fonts_reportlab()

            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=30, alignment=TA_CENTER)
            story.append(Paragraph(t['carbon_title'], title_style))
            story.append(Spacer(1, 12))

            company_info = f"""
            <b>{t['company']}:</b> {data.get('company_name', t['unknown'])}<br/>
            <b>{t['period']}:</b> {period}<br/>
            <b>{t['report_date']}:</b> {datetime.now().strftime('%d.%m.%Y')}
            """
            story.append(Paragraph(company_info, styles['Normal']))
            story.append(Spacer(1, 20))

            if 'emissions' in data:
                emissions = data['emissions']
                total_emissions = emissions.get('total', 0)
                story.append(Paragraph(f"<b>{t['total_emissions']}:</b> {total_emissions:.2f} tCO2e", styles['Normal']))
                story.append(Spacer(1, 12))

                if 'scope_breakdown' in emissions:
                    scope_data = emissions['scope_breakdown']
                    table_data = [[t['scope'], f"{t['emissions']} (tCO2e)", t['percentage']]]

                    for scope, value in scope_data.items():
                        percentage = (value / total_emissions * 100) if total_emissions > 0 else 0
                        table_data.append([scope.upper(), f"{value:.2f}", f"{percentage:.1f}%"])

                    table = Table(table_data, colWidths=[2*inch, 1.5*inch, 1.5*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(Paragraph(t['scope_breakdown'], styles['Heading2']))
                    story.append(Spacer(1, 12))
                    story.append(table)

            doc.build(story)
            self._save_generation_log(company_id, template_id, report_name, file_path)
            logging.info(f"[OK] Karbon raporu oluşturuldu ({language}): {file_path}")
            return file_path

        except Exception as e:
            logging.error(f"[HATA] Karbon raporu oluşturulamadı: {e}")
            return ""

    def create_cbam_report(self, company_id: int, period: str, data: Dict[str, Any], language: str = 'tr') -> str:
        """CBAM raporu oluştur"""
        try:
            t = TRANSLATIONS.get(language, TRANSLATIONS['tr'])
            template_id = f"cbam_template_{language}"
            report_name = f"CBAM_{company_id}_{period}_{language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            file_path = f"raporlar/cbam/{report_name}.pdf"
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            register_turkish_fonts_reportlab()

            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=30, alignment=TA_CENTER)
            story.append(Paragraph(t['cbam_title'], title_style))
            story.append(Spacer(1, 12))

            company_info = f"""
            <b>{t['company']}:</b> {data.get('company_name', t['unknown'])}<br/>
            <b>{t['period']}:</b> {period}<br/>
            <b>{t['cbam_declarant']}:</b> {data.get('declarant_name', '-')}<br/>
            <b>{t['report_date']}:</b> {datetime.now().strftime('%d.%m.%Y')}
            """
            story.append(Paragraph(company_info, styles['Normal']))
            story.append(Spacer(1, 20))

            # CBAM Details
            if 'emissions' in data:
                emissions = data['emissions']
                story.append(Paragraph(f"<b>{t['cbam_embedded']}:</b> {emissions.get('total_embedded', 0):.2f} tCO2e", styles['Normal']))
                story.append(Paragraph(f"<b>{t['cbam_price']}:</b> {data.get('carbon_price', 0):.2f} EUR/t", styles['Normal']))
                
                # Mock table for goods
                if 'goods' in data:
                    table_data = [['CN Code', 'Quantity (t)', 'Embedded Emissions', 'Indirect Emissions']]
                    for good in data['goods']:
                        table_data.append([good.get('cn_code'), good.get('quantity'), good.get('embedded'), good.get('indirect')])
                    
                    table = Table(table_data)
                    table.setStyle(TableStyle([('GRID', (0,0), (-1,-1), 1, colors.black)]))
                    story.append(Spacer(1, 10))
                    story.append(table)

            doc.build(story)
            self._save_generation_log(company_id, template_id, report_name, file_path)
            logging.info(f"[OK] CBAM raporu oluşturuldu ({language}): {file_path}")
            return file_path

        except Exception as e:
            logging.error(f"[HATA] CBAM raporu oluşturulamadı: {e}")
            return ""

    def create_tcfd_report(self, company_id: int, period: str, data: Dict[str, Any], language: str = 'tr') -> str:
        """TCFD raporu oluştur"""
        try:
            t = TRANSLATIONS.get(language, TRANSLATIONS['tr'])
            template_id = f"tcfd_template_{language}"
            report_name = f"TCFD_{company_id}_{period}_{language}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            file_path = f"raporlar/tcfd/{report_name}.pdf"
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            register_turkish_fonts_reportlab()

            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=30, alignment=TA_CENTER)
            story.append(Paragraph(t['tcfd_title'], title_style))
            story.append(Spacer(1, 12))

            company_info = f"""
            <b>{t['company']}:</b> {data.get('company_name', t['unknown'])}<br/>
            <b>{t['period']}:</b> {period}<br/>
            <b>{t['report_date']}:</b> {datetime.now().strftime('%d.%m.%Y')}
            """
            story.append(Paragraph(company_info, styles['Normal']))
            story.append(Spacer(1, 20))

            # TCFD Sections
            sections = [
                ('governance', t['tcfd_governance']),
                ('strategy', t['tcfd_strategy']),
                ('risk', t['tcfd_risk']),
                ('metrics', t['tcfd_metrics'])
            ]

            for key, title in sections:
                story.append(Paragraph(title, styles['Heading2']))
                content = data.get(key, '...')
                story.append(Paragraph(content, styles['Normal']))
                story.append(Spacer(1, 10))

            doc.build(story)
            self._save_generation_log(company_id, template_id, report_name, file_path)
            logging.info(f"[OK] TCFD raporu oluşturuldu ({language}): {file_path}")
            return file_path

        except Exception as e:
            logging.error(f"[HATA] TCFD raporu oluşturulamadı: {e}")
            return ""

    def create_excel_report(self, company_id: int, period: str, data: Dict[str, Any]) -> str:
        """Excel raporu oluştur"""
        try:
            report_name = f"Excel_Raporu_{company_id}_{period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            file_path = f"raporlar/excel/{report_name}.xlsx"
            os.makedirs(os.path.dirname(file_path), exist_ok=True)

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Rapor Özeti"
            ws['A1'] = "SÜRDÜRÜLEBİLİRLİK RAPORU"
            ws['A1'].font = Font(size=16, bold=True)
            ws.merge_cells('A1:D1')

            ws['A3'] = "Şirket:"
            ws['B3'] = data.get('company_name', 'Bilinmeyen')
            ws['A4'] = "Dönem:"
            ws['B4'] = period

            if 'data_table' in data:
                table_data = data['data_table']
                headers = list(table_data[0].keys()) if table_data else []
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=7, column=col, value=header)
                    cell.font = Font(bold=True)
                for row, row_data in enumerate(table_data, 8):
                    for col, header in enumerate(headers, 1):
                        ws.cell(row=row, column=col, value=row_data.get(header, ''))

            wb.save(file_path)
            logging.info(f"[OK] Excel raporu oluşturuldu: {file_path}")
            return file_path

        except Exception as e:
            logging.error(f"[HATA] Excel raporu oluşturulamadı: {e}")
            return ""

    def _save_generation_log(self, company_id: int, template_id: str, report_name: str, file_path: str):
        """Rapor oluşturma geçmişini kaydet"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            import uuid
            log_id = f"log_{company_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{str(uuid.uuid4())[:8]}"
            cursor.execute("""
                INSERT INTO report_generation_log
                (id, company_id, template_id, report_name, file_path, status)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (log_id, company_id, template_id, report_name, file_path, 'success'))
            conn.commit()
            conn.close()
        except Exception as e:
            logging.error(f"[HATA] Rapor geçmişi kaydedilemedi: {e}")

    def get_available_templates(self) -> List[Dict[str, Any]]:
        """Mevcut şablonları getir"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM report_templates WHERE is_active = 1 ORDER BY category, language, template_name")
            templates = []
            for row in cursor.fetchall():
                templates.append({
                    'id': row[0],
                    'name': row[1],
                    'type': row[2],
                    'category': row[3],
                    'language': row[4],
                    'config': json.loads(row[5])
                })
            conn.close()
            return templates
        except Exception as e:
            logging.error(f"[HATA] Şablonlar alınamadı: {e}")
            return []

if __name__ == "__main__":
    # Test
    logging.basicConfig(level=logging.INFO)
    logging.info("[TEST] Gelişmiş Raporlama...")
    templates = AdvancedReportTemplates()
    
    # Test Data
    test_data = {
        'company_name': 'Test Corp',
        'sdg_data': {'SDG 1': 85.0, 'SDG 13': 45.0},
        'emissions': {'total': 1200.5, 'scope_breakdown': {'scope1': 500, 'scope2': 700.5}},
        'tcfd_risk': 'Climate risks are high.'
    }

    # Test Generations
    logging.info("Testing TR SDG...")
    templates.generate_report('sdg_template_tr', 1, '2024', test_data)
    
    logging.info("Testing EN Carbon...")
    templates.generate_report('carbon_template_en', 1, '2024', test_data)
    
    logging.info("Testing HTML Preview...")
    templates.generate_report('sdg_template_tr_html', 1, '2024', test_data)
    
    logging.info("[OK] Test completed")
