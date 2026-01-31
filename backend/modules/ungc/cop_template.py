#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
UNGC COP (Communication on Progress) Şablon Üretici
- Basit DOCX rapor oluşturur
- python-docx yoksa uyarı verir ve False döner
"""

import logging
import os
from typing import Any, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.info("[WARN] python-docx not available. UNGC COP DOCX disabled.")


def _add_heading(doc: Any, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Calibri'
        default_sizes = {0: 25, 1: 17, 2: 14, 3: 13}
        run.font.size = Pt(default_sizes.get(level, 13))


def _add_paragraph(doc: Any, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)


def render_cop_docx(company_name: str, year: int, output_path: str,
                    compliance_summary: Optional[dict] = None) -> bool:
    """
    UNGC COP dokümanını oluştur.

    Args:
        company_name: Şirket adı
        year: Rapor yılı
        output_path: Çıkış docx yolu
        compliance_summary: {'overall_score': float, 'Environment_score': float, ...}
    """
    if not DOCX_AVAILABLE:
        logging.error("[ERROR] python-docx not installed")
        return False

    try:
        doc = Document()
        _add_heading(doc, f"UNGC COP - {company_name} ({year})", level=0)

        # CEO Statement
        _add_heading(doc, "CEO Statement", level=1)
        _add_paragraph(doc, "Şirketimiz Birleşmiş Milletler Küresel İlkeler Sözleşmesi'ne"
                             " bağlılığını sürdürmektedir. İnsan hakları, işgücü, çevre"
                             " ve yolsuzlukla mücadele ilkelerinde sürekli gelişim hedeflenmektedir.")

        # Governance & Policies
        _add_heading(doc, "Governance & Policies", level=1)
        _add_paragraph(doc, "UNGC ilkeleri doğrultusunda politika ve prosedürler güncellenmiştir."
                             " Uyum komitesi ve ilgili çalışma grupları aktif olarak izleme yapmaktadır.")

        # Goals & KPIs
        _add_heading(doc, "Goals & KPIs", level=1)
        if compliance_summary:
            _add_paragraph(doc, f"Genel Uyum Skoru: {compliance_summary.get('overall_score', 0):.1f}")
            for k, v in compliance_summary.items():
                if k.endswith('_score') and k != 'overall_score':
                    _add_paragraph(doc, f"{k.replace('_score','')}: {v:.1f}")

            _add_heading(doc, "Kategori ve İlke Kırılımı", level=2)
            cats = [
                'Environment', 'Human Rights', 'Labour', 'Anti-Corruption'
            ]
            for cat in cats:
                s_key = f"{cat}_score"
                if s_key in compliance_summary:
                    _add_paragraph(doc, f"{cat}: {compliance_summary.get(s_key, 0):.1f}")

            if isinstance(compliance_summary.get('principles'), (list, tuple, dict)):
                _add_heading(doc, "İlke Bazlı Özet", level=2)
                principles = compliance_summary.get('principles')
                if isinstance(principles, dict):
                    items_list_dict: List[Tuple[Any, Any]] = list(principles.items())
                    items_list = items_list_dict
                else:
                    seq = principles if isinstance(principles, (list, tuple)) else []
                    items_list = [
                        (p.get('id') or p.get('code') or str(i+1), p) for i, p in enumerate(seq)
                    ]
                for pid, pdata in items_list:
                    ps = pdata.get('score', pdata.get('value', 0)) if isinstance(pdata, dict) else 0
                    _add_paragraph(doc, f"{pid}: {ps:.1f}")
                    kpis = pdata.get('kpis') if isinstance(pdata, dict) else None
                    if isinstance(kpis, (list, tuple)):
                        for kpi in kpis[:5]:
                            name = kpi.get('name') if isinstance(kpi, dict) else str(kpi)
                            val = kpi.get('current_value') if isinstance(kpi, dict) else None
                            tgt = kpi.get('target_value') if isinstance(kpi, dict) else None
                            _add_paragraph(doc, f"- {name} (değer: {val if val is not None else '-'}, hedef: {tgt if tgt is not None else '-'})")

            _add_heading(doc, "Veri Boşlukları Checklist", level=2)
            gaps = []
            for key in ('data_gaps', 'missing_kpis', 'missing_fields', 'empty_values', 'gaps'):
                v = compliance_summary.get(key)
                if isinstance(v, (list, tuple)):
                    gaps.extend([str(x) for x in v])
            if gaps:
                for g in gaps[:20]:
                    _add_paragraph(doc, f"- {g}")
            else:
                _add_paragraph(doc, "Veri boşluğu bildirilmedi.")
        else:
            _add_paragraph(doc, "KPI ve hedef verileri bu sürümde özet düzeydedir.")

        # Progress & Timeline
        _add_heading(doc, "Progress & Timeline", level=1)
        _add_paragraph(doc, "Yıl içinde kaydedilen gelişmeler ve gelecek dönem planları burada listelenir."
                             " Bu şablon temel bir iskele sunar ve detaylandırılabilir.")

        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.exists(out_dir):
            os.makedirs(out_dir, exist_ok=True)
        doc.save(output_path)
        logging.info(f"[UNGC] COP DOCX oluşturuldu: {output_path}")
        return True
    except Exception as e:
        logging.error(f"[ERROR] UNGC COP oluşturulamadı: {e}")
        return False
