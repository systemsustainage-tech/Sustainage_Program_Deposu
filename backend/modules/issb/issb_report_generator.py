import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ISSB Report Generator
- DOCX raporu oluşturma (ek: ISSB Gap Aksiyon Planı)
- Excel olarak aksiyon planı dışa aktarımı

Not: Bu basit oluşturucu, mevcut ISSB durum özetini (ISSBManager.generate_issb_report)
ve oluşturulmuş ISSB gap raporundaki "Tavsiye Edilen Aksiyonlar" listesini kullanır.
"""

import os
from datetime import datetime
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    from openpyxl.comments import Comment
    from openpyxl.styles import Alignment, Font
    from openpyxl.worksheet.datavalidation import DataValidation
    EXCEL_AVAILABLE = True
except Exception:
    EXCEL_AVAILABLE = False

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    MPL_AVAILABLE = True
except Exception:
    MPL_AVAILABLE = False


class ISSBReportGenerator:
    def __init__(self):
        # Utility class, no initialization required
        pass

    @staticmethod
    def _db_path() -> str:
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        return os.path.join(base_dir, "data", "sdg_desktop.sqlite")

    def _add_brand_header(self, doc, company_id: int) -> None:
        try:
            from modules.reporting.brand_identity_manager import BrandIdentityManager
            bim = BrandIdentityManager(self._db_path(), company_id)
            bi = bim.get_brand_identity(company_id)
            lp = bi.get('logo_path')
            if lp and os.path.exists(lp):
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(lp, width=Inches(1.6))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ht = (bi.get('texts') or {}).get('header')
            if ht:
                t = doc.add_paragraph(ht)
                t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    @staticmethod
    def _gap_report_path(company_id: int, year: int) -> str:
        return os.path.join("data", "reports", "issb", f"issb_gap_report_{company_id}_{year}.md")

    def parse_action_items(self, company_id: int, year: int) -> List[str]:
        """Markdown gap raporundan "Tavsiye Edilen Aksiyonlar" listesini çıkar."""
        path = self._gap_report_path(company_id, year)
        if not os.path.exists(path):
            return []

        items: List[str] = []
        in_actions = False
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line.startswith("## "):
                    # Başlığa göre aksiyon bölümünde olup olmadığımızı ayarla
                    in_actions = line.lower().startswith("## tavsiye edilen aksiyonlar")
                    continue
                if in_actions and line.startswith("- "):
                    items.append(line[2:].strip())
        return items

    def _action_plan_path(self, company_id: int, year: int) -> str:
        return os.path.join("data", "reports", "issb", f"issb_action_plan_{company_id}_{year}.xlsx")

    def get_action_plan_status_summary(self, company_id: int, year: int) -> Dict[str, int]:
        summary = {"tamamlanan": 0, "devam": 0, "beklemede": 0, "son_tarih_gecmis": 0}
        try:
            path = self._action_plan_path(company_id, year)
            if not EXCEL_AVAILABLE or not os.path.exists(path):
                return summary
            wb = openpyxl.load_workbook(path)
            ws = wb.active
            headers = [cell.value for cell in ws[1]]
            status_idx = headers.index("Durum") if "Durum" in headers else None
            due_idx = headers.index("Hedef Tarih") if "Hedef Tarih" in headers else None
            if status_idx is None:
                return summary
            from datetime import datetime
            today = datetime.now().date()
            for row in ws.iter_rows(min_row=2, values_only=True):
                status = (row[status_idx] or "").strip()
                if status == "Tamamlandı":
                    summary["tamamlanan"] += 1
                elif status == "Devam":
                    summary["devam"] += 1
                else:
                    summary["beklemede"] += 1
                if due_idx is not None and status != "Tamamlandı":
                    due_val = row[due_idx]
                    d = self._to_date_value(due_val)
                    if d and d < today:
                        summary["son_tarih_gecmis"] += 1
            return summary
        except Exception:
            return summary

    def get_action_plan_status_summary_with_filters(self, company_id: int, year: int, department: Optional[str] = None, responsible: Optional[str] = None, status_filter: Optional[str] = None, overdue_only: bool = False) -> Dict[str, int]:
        summary = {"tamamlanan": 0, "devam": 0, "beklemede": 0, "son_tarih_gecmis": 0}
        try:
            path = self._action_plan_path(company_id, year)
            if not EXCEL_AVAILABLE or not os.path.exists(path):
                return summary
            wb = openpyxl.load_workbook(path)
            ws = wb.active
            headers = [cell.value for cell in ws[1]]
            status_idx = headers.index("Durum") if "Durum" in headers else None
            due_idx = headers.index("Hedef Tarih") if "Hedef Tarih" in headers else None
            dep_idx = headers.index("Departman") if "Departman" in headers else None
            resp_idx = headers.index("Sorumlu") if "Sorumlu" in headers else None
            if status_idx is None:
                return summary
            from datetime import datetime
            today = datetime.now().date()
            for row in ws.iter_rows(min_row=2, values_only=True):
                st = (row[status_idx] or "").strip()
                dv = row[due_idx] if due_idx is not None else None
                depv = (row[dep_idx] if dep_idx is not None else None)
                respv = (row[resp_idx] if resp_idx is not None else None)
                if department and (not depv or str(depv).strip() != department):
                    continue
                if responsible and (not respv or str(respv).strip() != responsible):
                    continue
                if status_filter and st != status_filter:
                    continue
                if overdue_only:
                    d = self._to_date_value(dv)
                    if not d or d >= today:
                        continue
                if st == "Tamamlandı":
                    summary["tamamlanan"] += 1
                elif st == "Devam":
                    summary["devam"] += 1
                else:
                    summary["beklemede"] += 1
                if due_idx is not None and st != "Tamamlandı":
                    d2 = self._to_date_value(dv)
                    if d2 and d2 < today:
                        summary["son_tarih_gecmis"] += 1
            return summary
        except Exception:
            return summary

    def get_action_plan_monthly_completion(self, company_id: int, year: int) -> List[int]:
        vals = [0]*12
        try:
            path = self._action_plan_path(company_id, year)
            if not EXCEL_AVAILABLE or not os.path.exists(path):
                return vals
            wb = openpyxl.load_workbook(path)
            ws = wb.active
            headers = [cell.value for cell in ws[1]]
            status_idx = headers.index("Durum") if "Durum" in headers else None
            due_idx = headers.index("Hedef Tarih") if "Hedef Tarih" in headers else None
            if status_idx is None or due_idx is None:
                return vals
            for row in ws.iter_rows(min_row=2, values_only=True):
                st = (row[status_idx] or "").strip()
                dv = row[due_idx]
                if st != "Tamamlandı":
                    continue
                d = self._to_datetime_value(dv)
                if d and d.year == year:
                    vals[d.month-1] += 1
            return vals
        except Exception:
            return vals

    def get_action_plan_monthly_targets(self, company_id: int, year: int) -> List[int]:
        vals = [0]*12
        try:
            path = self._action_plan_path(company_id, year)
            if not EXCEL_AVAILABLE or not os.path.exists(path):
                return vals
            wb = openpyxl.load_workbook(path)
            ws = wb.active
            headers = [cell.value for cell in ws[1]]
            due_idx = headers.index("Hedef Tarih") if "Hedef Tarih" in headers else None
            if due_idx is None:
                return vals
            for row in ws.iter_rows(min_row=2, values_only=True):
                dv = row[due_idx]
                d = self._to_datetime_value(dv)
                if d and d.year == year:
                    vals[d.month-1] += 1
            return vals
        except Exception:
            return vals

    @staticmethod
    def _to_date_value(val: Any) -> Optional[datetime.date]:
        from datetime import date, datetime
        if isinstance(val, date):
            return val
        if isinstance(val, str) and val:
            try:
                return datetime.fromisoformat(val).date()
            except Exception:
                try:
                    return datetime.strptime(val, "%Y-%m-%d").date()
                except Exception:
                    return None
        return None

    @staticmethod
    def _to_datetime_value(val: Any) -> Optional[datetime]:
        from datetime import date, datetime
        if isinstance(val, datetime):
            return val
        if isinstance(val, date):
            return datetime(val.year, val.month, val.day)
        if isinstance(val, str) and val:
            try:
                return datetime.fromisoformat(val)
            except Exception:
                try:
                    return datetime.strptime(val, "%Y-%m-%d")
                except Exception:
                    return None
        return None

    def generate_docx_report(self, output_path: str, report: Dict, company_id: int, year: int) -> bool:
        if not DOCX_AVAILABLE:
            return False
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc = Document()

            self._add_brand_header(doc, company_id)

            # Başlık
            title = doc.add_paragraph(f"ISSB RAPORU — Şirket {company_id}, Yıl {year}")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Rapor Tarihi: {datetime.now().isoformat()}")
            doc.add_paragraph("")

            # Genel özet
            doc.add_heading("Genel Özet", level=1)
            doc.add_paragraph(f"Genel Uyumluluk: {report.get('overall_compliance', 0):.1f}%")
            doc.add_paragraph(f"Hazırlık Seviyesi: {report.get('readiness_level', 'N/A')}")

            # S1/S2 detayları
            s1 = report.get('ifrs_s1_compliance', {})
            s2 = report.get('ifrs_s2_compliance', {})
            doc.add_heading("IFRS S1 - Genel Gereksinimler", level=2)
            doc.add_paragraph(f"Uyumluluk Oranı: {s1.get('compliance_rate', 0):.1f}% | Tamamlanan: {s1.get('completed', 0)} | Devam Eden: {s1.get('in_progress', 0)} | Başlanmayan: {s1.get('not_started', 0)}")
            doc.add_heading("IFRS S2 - İklim Açıklamaları", level=2)
            doc.add_paragraph(f"Uyumluluk Oranı: {s2.get('compliance_rate', 0):.1f}% | Tamamlanan: {s2.get('completed', 0)} | Devam Eden: {s2.get('in_progress', 0)} | Başlanmayan: {s2.get('not_started', 0)}")

            # Durum Göstergesi (Aksiyon Planı)
            doc.add_heading("Durum Göstergesi (Aksiyon Planı)", level=1)
            summary = self.get_action_plan_status_summary(company_id, year)
            doc.add_paragraph(
                f"Tamamlanan: {summary['tamamlanan']} | Devam: {summary['devam']} | Beklemede: {summary['beklemede']} | Son Tarihi Geçmiş: {summary['son_tarih_gecmis']}"
            )

            try:
                if MPL_AVAILABLE:
                    ap_monthly = self.get_action_plan_monthly_completion(company_id, year)
                    fig, ax = plt.subplots(figsize=(6, 2))
                    ax.bar(list(range(1,13)), ap_monthly, color='#3b82f6')
                    ax.set_title('Aksiyon Planı Tamamlanan Aylık')
                    ax.set_xlabel('Ay')
                    ax.set_ylabel('Adet')
                    out_dir = os.path.dirname(output_path)
                    img_path = os.path.join(out_dir, f"issb_kpi_{company_id}_{year}.png")
                    plt.tight_layout()
                    fig.savefig(img_path)
                    plt.close(fig)
                    try:
                        from docx.shared import Inches
                        doc.add_picture(img_path, width=Inches(6))
                    except Exception:
                        doc.add_picture(img_path)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            doc.add_heading("Ek: ISSB Gap Aksiyon Planı", level=1)
            doc.add_paragraph("Doldurma talimatı:")
            doc.add_paragraph("- Eksik Kalem: gap raporundaki satırı ekleyin.")
            doc.add_paragraph("- Standard: S1 veya S2.")
            doc.add_paragraph("- Sorumlu Modül: örn. modules/issb.")
            doc.add_paragraph("- Önerilen Adım: kısa ve eylem odaklı.")
            doc.add_paragraph("- Hedef Tarih: YYYY-MM-DD.")
            doc.add_paragraph("- Durum: Beklemede, Devam, Tamamlandı.")
            doc.add_paragraph("- Bağımlılıklar: diğer görevler/veri/onaylar.")
            actions = self.parse_action_items(company_id, year)
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            headers = ["Eksik Kalem", "Standard", "Sorumlu Modül", "Önerilen Adım", "Hedef Tarih", "Durum", "Bağımlılıklar"]
            for i, h in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = h
                try:
                    cell.paragraphs[0].runs[0].bold = True
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")

            if not actions:
                row = table.add_row()
                row.cells[0].text = "(Gap raporunda aksiyon bulunamadı)"
                row.cells[1].text = ""
                row.cells[2].text = "modules/issb"
                row.cells[3].text = ""
                row.cells[4].text = ""
                row.cells[5].text = "Beklemede"
                row.cells[6].text = ""
            else:
                for act in actions:
                    row = table.add_row()
                    # Basit yerleştirme: Eksik Kalem/Önerilen Adım aynı; Standard bilinmiyor
                    row.cells[0].text = act
                    row.cells[1].text = "S1/S2"
                    row.cells[2].text = "modules/issb"
                    row.cells[3].text = act
                    row.cells[4].text = ""
                    row.cells[5].text = "Beklemede"
                    row.cells[6].text = ""

            doc.add_page_break()
            doc.add_heading("Dipnotlar", level=1)
            doc.add_paragraph("Veri Kaynağı: ISSB uyumluluk ve aksiyon planı verileri, yıl bazlı kayıtlar.")
            doc.add_paragraph("Varsayımlar: Eksik kalemlerin standart eşleşmeleri S1/S2 olarak yorumlanır.")
            doc.add_paragraph("Metodoloji: Gap raporundan aksiyon maddeleri tablolaştırılır; durum alanları manuel güncellenir.")

            doc.add_page_break()
            doc.add_heading("Metodoloji Genişletme", level=1)
            doc.add_paragraph("Financed emissions: varsa finans/kredi portföyleri için metodoloji ve veri alanları tanımlanır.")
            doc.add_paragraph("Internal carbon pricing: kurum içi karbon fiyatı bağlamı ve maliyet etkileri özetlenir.")
            doc.add_paragraph("Sektör metrikleri: GRI/SASB sektör standartlarına göre veri toplama metodolojisi ve kaynakları belirtilir.")

            doc.save(output_path)
            return True
        except Exception as e:
            logging.error(f"Error generating excel action plan: {e}")
            return False

    def export_excel_action_plan(self, output_path: str, company_id: int, year: int) -> bool:
        if not EXCEL_AVAILABLE:
            return False
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "ISSB Aksiyon Planı"

            headers = ["Eksik Kalem", "Standard", "Sorumlu Modül", "Departman", "Sorumlu", "Önerilen Adım", "Hedef Tarih", "Durum", "Bağımlılıklar"]
            ws.append(headers)

            bold = Font(bold=True)
            for col in range(1, len(headers) + 1):
                ws.cell(row=1, column=col).font = bold
                ws.cell(row=1, column=col).alignment = Alignment(horizontal='center')
            ws.freeze_panes = "A2"
            try:
                tips = {
                    1: "Gap raporundaki satır.",
                    2: "S1 veya S2.",
                    3: "Örn. modules/issb.",
                    4: "Departman.",
                    5: "Sorumlu kişi.",
                    6: "Kısa, eylem odaklı.",
                    7: "YYYY-MM-DD.",
                    8: "Beklemede, Devam, Tamamlandı.",
                    9: "Diğer görevler/veri/onaylar."
                }
                for c, tip in tips.items():
                    ws.cell(row=1, column=c).comment = Comment(tip, "SDG")
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
            try:
                dv = DataValidation(type="list", formula1='"Beklemede,Devam,Tamamlandı"', allow_blank=True)
                ws.add_data_validation(dv)
                dv.add("H2:H1048576")
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            actions = self.parse_action_items(company_id, year)
            if not actions:
                ws.append(["(Gap raporunda aksiyon bulunamadı)", "", "modules/issb", "", "", "", "", "Beklemede", ""])
            else:
                for act in actions:
                    ws.append([act, "S1/S2", "modules/issb", "", "", act, "", "Beklemede", ""])

            try:
                last_row = ws.max_row
                ws.auto_filter.ref = f"A1:I{last_row}"
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            # Ayrı sayfa: Gap Aksiyonları
            try:
                ws_gap = wb.create_sheet(title="Gap Aksiyonları")
                ws_gap.append(headers)
                for col in range(1, len(headers) + 1):
                    ws_gap.cell(row=1, column=col).font = bold
                    ws_gap.cell(row=1, column=col).alignment = Alignment(horizontal='center')
                ws_gap.freeze_panes = "A2"
                if not actions:
                    ws_gap.append(["(Gap raporunda aksiyon bulunamadı)", "", "modules/issb", "", "", "", "", "Beklemede", ""])
                else:
                    for act in actions:
                        ws_gap.append([act, "S1/S2", "modules/issb", "", "", act, "", "Beklemede", ""])
                try:
                    last_row = ws_gap.max_row
                    ws_gap.auto_filter.ref = f"A1:I{last_row}"
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            try:
                inst = wb.create_sheet(title="Talimatlar")
                inst.append(["ISSB Gap Aksiyon Planı – Doldurma Talimatları"])
                inst.append(["Eksik Kalem: gap raporundaki satır."])
                inst.append(["Standard: S1 veya S2."])
                inst.append(["Sorumlu Modül: örn. modules/issb."])
                inst.append(["Departman: ilgili birim."])
                inst.append(["Sorumlu: kişi/rol."])
                inst.append(["Önerilen Adım: kısa ve eylem odaklı."])
                inst.append(["Hedef Tarih: YYYY-MM-DD."])
                inst.append(["Durum: Beklemede, Devam, Tamamlandı."])
                inst.append(["Bağımlılıklar: diğer görevler/veri/onaylar."])
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            wb.save(output_path)
            return True
        except Exception:
            return False

    def export_pdf_action_plan(self, output_path: str, company_id: int, year: int) -> bool:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            elements = []
            elements.append(Paragraph(f"ISSB Aksiyon Planı — Şirket {company_id}, Yıl {year}", styles['Title']))
            elements.append(Paragraph(f"Oluşturulma Tarihi: {datetime.now().isoformat()}", styles['Normal']))
            elements.append(Spacer(1, 12))
            headers = ["Eksik Kalem", "Standard", "Sorumlu Modül", "Önerilen Adım", "Hedef Tarih", "Durum", "Bağımlılıklar"]
            actions = self.parse_action_items(company_id, year)
            data = [headers]
            if not actions:
                data.append(["(Gap raporunda aksiyon bulunamadı)", "", "modules/issb", "", "", "Beklemede", ""])
            else:
                for act in actions:
                    data.append([act, "S1/S2", "modules/issb", act, "", "Beklemede", ""])
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                ('TEXTCOLOR', (0,0), (-1,0), colors.black),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 10),
                ('BOTTOMPADDING', (0,0), (-1,0), 8),
                ('GRID', (0,0), (-1,-1), 0.25, colors.grey)
            ]))
            elements.append(table)
            doc.build(elements)
            return True
        except Exception:
            return False
