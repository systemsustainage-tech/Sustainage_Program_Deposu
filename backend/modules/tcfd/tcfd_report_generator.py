import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TCFD REPORT GENERATOR - Rapor Oluşturma
- PDF, DOCX, Excel formatlarında TCFD raporu
- 4 temel bileşen: Governance, Strategy, Risk Management, Metrics
- Grafikler ve tablolar
- Finansal etki analizi
"""

import json
import os
from datetime import datetime
from typing import Dict, Optional

# Rapor oluşturma için gerekli kütüphaneler
try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.pdfmetrics import registerFontFamily
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.info("[WARN] ReportLab not available. PDF generation disabled.")

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.info("[WARN] python-docx not available. DOCX generation disabled.")

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.info("[WARN] openpyxl not available. Excel generation disabled.")

# Basit görsel özet için ReportLab grafikler (opsiyonel)
try:
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.shapes import Drawing
    GRAPHICS_AVAILABLE = True
except Exception:
    GRAPHICS_AVAILABLE = False


def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=12):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri', font_size=None):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        # Varsayılan başlık boyutları: 0:25, 1:17, 2:14, 3:13
        default_sizes = {0: 25, 1: 17, 2: 14, 3: 13}
        size_pt = font_size if font_size is not None else default_sizes.get(level, 13)
        run.font.size = Pt(size_pt)
    return heading


class TCFDReportGenerator:
    """TCFD rapor oluşturucu"""

    def __init__(self, manager, calculator):
        """
        Args:
            manager: TCFDManager instance
            calculator: TCFDCalculator instance
        """
        self.manager = manager
        self.calculator = calculator
        self._fonts_registered = False
        # Senaryolar için isim listesi (öncelikli)
        self._default_scenarios = [
            "Net Zero 2050",
            "Below 2°C",
            "Current Policies",
            "Hot House World"
        ]

    def _load_scenario_data(self) -> Dict:
        """climate_scenarios.json dosyasını yükle"""
        try:
            module_dir = os.path.dirname(os.path.abspath(__file__))
            scenarios_path = os.path.join(module_dir, 'data', 'climate_scenarios.json')
            if os.path.exists(scenarios_path):
                with open(scenarios_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            logging.info(f"[WARN] Senaryo verileri yüklenemedi: {e}")
        return {}

    def _get_scenario_carbon_price(self, scenario_data: Dict, scenario_name: str, year: int) -> Optional[float]:
        """Verilen yıl için senaryo karbon fiyatını getir (USD/tCO2e)"""
        try:
            cmp = scenario_data.get('scenario_comparison_metrics', {})
            prices = cmp.get('carbon_price', {}).get('scenarios', {}).get(scenario_name)
            if not prices:
                return None
            # Yıl mevcutsa direkt al, değilse en yakın yıl
            ykey = str(year)
            if ykey in prices:
                return float(prices[ykey])
            # En yakın yıl (lineer yaklaşım yerine basit nearest)
            keys = sorted(int(k) for k in prices.keys())
            if not keys:
                return None
            closest = min(keys, key=lambda k: abs(k - year))
            return float(prices[str(closest)])
        except Exception:
            return None

    def _build_management_summary(self, company_id: int, year: int) -> Dict:
        """Yönetici özeti için senaryo finansal özet verilerini hazırla"""
        metrics = self.manager.get_metrics(company_id, year)
        scenario_data = self._load_scenario_data()

        if not metrics:
            return {'summary_rows': [], 'chart_data': []}

        # Emisyonlar ve finansal metrikler
        emissions = {
            'scope1': metrics.get('scope1_emissions') or 0,
            'scope2': metrics.get('scope2_emissions') or 0,
            'scope3': metrics.get('scope3_emissions') or 0,
        }
        climate_revenue_m = metrics.get('climate_related_revenue') or 0  # milyon
        climate_opex_m = metrics.get('climate_related_opex') or 0        # milyon
        climate_capex_m = metrics.get('climate_related_capex') or 0      # milyon

        climate_revenue = climate_revenue_m * 1_000_000
        climate_opex = climate_opex_m * 1_000_000
        climate_capex = climate_capex_m * 1_000_000

        rows = []
        chart_vals = []

        for scen in self._default_scenarios:
            carbon_price = self._get_scenario_carbon_price(scenario_data, scen, year)
            if carbon_price is None:
                carbon_price = metrics.get('internal_carbon_price') or 0

            impact = self.calculator.calculate_carbon_price_impact(
                emissions['scope1'], emissions['scope2'], emissions['scope3'], carbon_price, True
            )

            total_costs = impact['total_cost'] + climate_opex + climate_capex
            net_impact = climate_revenue - total_costs

            # Hassasiyet: -20%, +20%
            low_price = carbon_price * 0.8
            high_price = carbon_price * 1.2
            low_cost = self.calculator.calculate_carbon_price_impact(
                emissions['scope1'], emissions['scope2'], emissions['scope3'], low_price, True
            )['total_cost'] + climate_opex + climate_capex
            high_cost = self.calculator.calculate_carbon_price_impact(
                emissions['scope1'], emissions['scope2'], emissions['scope3'], high_price, True
            )['total_cost'] + climate_opex + climate_capex

            rows.append([
                scen,
                f"${carbon_price:,.0f}",
                f"${impact['total_cost']:,.0f}",
                f"${climate_revenue:,.0f}",
                f"${climate_opex:,.0f}",
                f"${climate_capex:,.0f}",
                f"${net_impact:,.0f}",
                f"${low_cost:,.0f} - ${high_cost:,.0f}"
            ])
            chart_vals.append(net_impact)

        return {'summary_rows': rows, 'chart_data': chart_vals, 'scenarios': self._default_scenarios}

    def _register_turkish_fonts(self):
        """Türkçe karakterleri destekleyen TTF fontları kaydet"""
        if self._fonts_registered:
            return

        try:
            # Proje kökünden raporlama/fonts klasörünü çöz
            module_dir = os.path.dirname(os.path.abspath(__file__))
            base_dir = os.path.abspath(os.path.join(module_dir, '..', '..'))
            fonts_dir = os.path.join(base_dir, 'raporlama', 'fonts')

            # Öncelik: NotoSans ailesi
            noto_regular = os.path.join(fonts_dir, 'NotoSans-Regular.ttf')
            noto_bold = os.path.join(fonts_dir, 'NotoSans-Bold.ttf')
            noto_italic = os.path.join(fonts_dir, 'NotoSans-Italic.ttf')

            if os.path.exists(noto_regular):
                pdfmetrics.registerFont(TTFont('NotoSans', noto_regular))
            if os.path.exists(noto_bold):
                pdfmetrics.registerFont(TTFont('NotoSans-Bold', noto_bold))
            if os.path.exists(noto_italic):
                pdfmetrics.registerFont(TTFont('NotoSans-Italic', noto_italic))

            # Aile eşlemeleri (bold/italic)
            registerFontFamily(
                'NotoSans',
                normal='NotoSans',
                bold='NotoSans-Bold',
                italic='NotoSans-Italic' if os.path.exists(noto_italic) else 'NotoSans',
                boldItalic='NotoSans-Bold'
            )

            # Alternatif: DejaVuSans (varsa)
            deja = os.path.join(fonts_dir, 'DejaVuSans.ttf')
            if os.path.exists(deja):
                pdfmetrics.registerFont(TTFont('DejaVuSans', deja))
                registerFontFamily('DejaVuSans', normal='DejaVuSans')

            self._fonts_registered = True
        except Exception as e:
            logging.info(f"[WARN] Türkçe font kaydı yapılamadı: {e}")

    # ========================================================================
    # PDF RAPOR
    # ========================================================================

    def generate_pdf_report(
        self,
        company_id: int,
        year: int,
        company_name: str,
        output_path: str
    ) -> bool:
        """
        PDF TCFD raporu oluştur
        
        Args:
            company_id: Firma ID
            year: Raporlama yılı
            company_name: Firma adı
            output_path: Çıktı dosya yolu
        
        Returns:
            Başarılı mı?
        """
        if not REPORTLAB_AVAILABLE:
            logging.error("[ERROR] ReportLab not installed")
            return False

        try:
            # Türkçe fontları yükle
            self._register_turkish_fonts()
            # PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )

            # Story (içerik)
            story = []
            styles = getSampleStyleSheet()

            # Custom styles - TÜRKÇE KARAKTER DESTEKLİ
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=25,
                textColor=colors.HexColor('#2E8B57'),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='NotoSans'
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=17,
                textColor=colors.HexColor('#2E8B57'),
                spaceAfter=12,
                spaceBefore=20,
                fontName='NotoSans'
            )

            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=12,
                fontName='NotoSans'
            )

            subheading_style = ParagraphStyle(
                'CustomSubHeading',
                parent=styles['Heading3'],
                fontSize=14,
                textColor=colors.HexColor('#2E8B57'),
                fontName='NotoSans'
            )

            # Marka logo/header
            try:
                from modules.reporting.brand_identity_manager import BrandIdentityManager
                base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                db_path = os.path.join(base_dir, "data", "sdg_desktop.sqlite")
                bim = BrandIdentityManager(db_path, company_id)
                bi = bim.get_brand_identity(company_id)
                lp = bi.get('logo_path')
                if lp and os.path.exists(lp):
                    try:
                        story.append(Image(lp, width=2.0*inch, height=2.0*inch, hAlign='CENTER'))
                        story.append(Spacer(1, 0.2*inch))
                    except Exception as e:
                        logging.error(f"Silent error caught: {str(e)}")
                ht = (bi.get('texts') or {}).get('header')
                if ht:
                    story.append(Paragraph(ht, ParagraphStyle('BrandHeader', parent=styles['Normal'], alignment=TA_CENTER, fontName='NotoSans')))
                    story.append(Spacer(1, 0.2*inch))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            # Başlık sayfası
            story.append(Spacer(1, 2*inch))
            story.append(Paragraph(f"<b>{company_name}</b>", title_style))
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph(
                "<b>TCFD Raporu</b>",
                ParagraphStyle('Subtitle', parent=styles['Normal'],
                             fontSize=19, alignment=TA_CENTER, fontName='NotoSans')
            ))
            story.append(Paragraph(
                "Task Force on Climate-related Financial Disclosures",
                ParagraphStyle('Subtitle2', parent=styles['Normal'],
                             fontSize=13, alignment=TA_CENTER, textColor=colors.grey, fontName='NotoSans')
            ))
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph(
                f"Raporlama Yılı: {year}",
                ParagraphStyle('Year', parent=styles['Normal'],
                             fontSize=15, alignment=TA_CENTER, textColor=colors.HexColor('#2E8B57'), fontName='NotoSans')
            ))
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph(
                f"Rapor Tarihi: {datetime.now().strftime('%d %B %Y')}",
                ParagraphStyle('Date', parent=styles['Normal'],
                             fontSize=11, alignment=TA_CENTER, textColor=colors.grey, fontName='NotoSans')
            ))

            story.append(PageBreak())

            # İçindekiler (basit)
            story.append(Paragraph("<b>İçindekiler</b>", heading_style))
            toc_data = [
                ["1.", "Yönetişim (Governance)", "3"],
                ["2.", "Strateji (Strategy)", "5"],
                ["3.", "Risk Yönetimi (Risk Management)", "8"],
                ["4.", "Metrikler ve Hedefler (Metrics and Targets)", "12"],
            ]

            toc_table = Table(toc_data, colWidths=[0.5*inch, 4*inch, 0.5*inch])
            toc_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), 'NotoSans', 12),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ]))
            story.append(toc_table)
            story.append(PageBreak())

            # Yönetici Özeti (Senaryo Finansal Özet)
            story.append(Paragraph("<b>Yönetici Özeti</b>", heading_style))
            story.append(Paragraph(
                "Seçilmiş senaryolar için karbon maliyeti, iklim ilişkili gelir/gider ve CAPEX etkilerinin özetidir.",
                body_style
            ))

            mgmt = self._build_management_summary(company_id, year)
            summary_data = [[
                "Senaryo", "Karbon Fiyatı (USD/tCO2e)", "Karbon Maliyeti", "İklim Geliri",
                "İklim OPEX", "İklim CAPEX", "Net Etki", "Hassasiyet (Toplam Maliyet)"
            ]] + (mgmt['summary_rows'] or [["Veri yok", "-", "-", "-", "-", "-", "-", "-"]])

            sum_table = Table(summary_data, colWidths=[1.6*inch, 1.3*inch, 1.3*inch, 1.3*inch, 1.1*inch, 1.1*inch, 1.2*inch, 2.0*inch])
            sum_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), 'NotoSans', 10),
                ('FONT', (0, 0), (-1, 0), 'NotoSans-Bold', 11),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E8F5E9')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FAFAFA')]),
                ('ALIGN', (1, 1), (-1, -1), 'RIGHT'),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ]))
            story.append(sum_table)

            # Basit görsel: Net etki çubuk grafiği
            if GRAPHICS_AVAILABLE and mgmt.get('chart_data'):
                try:
                    d = Drawing(400, 200)
                    bc = VerticalBarChart()
                    bc.x = 40
                    bc.y = 40
                    bc.height = 120
                    bc.width = 320
                    bc.data = [mgmt['chart_data']]
                    bc.categoryAxis.categoryNames = mgmt.get('scenarios', [])
                    bc.barLabels.nudge = 7
                    bc.barLabelFormat = '$%0.0f'
                    bc.valueAxis.labels.fontName = 'NotoSans'
                    bc.categoryAxis.labels.boxAnchor = 'ne'
                    bc.categoryAxis.labels.fontName = 'NotoSans'
                    bc.bars[0].fillColor = colors.HexColor('#66BB6A')
                    d.add(bc)
                    story.append(Spacer(1, 0.2*inch))
                    story.append(d)
                except Exception as ge:
                    logging.info(f"[WARN] Grafik oluşturulamadı: {ge}")

            story.append(PageBreak())

            # 1. Governance
            story.append(Paragraph("<b>1. Yönetişim (Governance)</b>", heading_style))
            gov_data = self.manager.get_governance(company_id, year)

            if gov_data:
                story.append(Paragraph(
                    "<b>Yönetim Kurulu Gözetimi:</b>",
                    subheading_style
                ))
                story.append(Paragraph(
                    gov_data.get('board_oversight', 'Veri yok'),
                    body_style
                ))
                story.append(Spacer(1, 0.2*inch))

                story.append(Paragraph(
                    "<b>Üst Yönetimin Rolü:</b>",
                    subheading_style
                ))
                story.append(Paragraph(
                    gov_data.get('management_role', 'Veri yok'),
                    body_style
                ))
            else:
                story.append(Paragraph(
                    "Yönetişim verileri henüz girilmemiş.",
                    body_style
                ))

            story.append(PageBreak())

            # 2. Strategy
            story.append(Paragraph("<b>2. Strateji (Strategy)</b>", heading_style))
            strategy_data = self.manager.get_strategy(company_id, year)

            if strategy_data:
                story.append(Paragraph(
                    "<b>Kısa Vadeli Riskler (0-3 yıl):</b>",
                    subheading_style
                ))
                story.append(Paragraph(
                    strategy_data.get('short_term_risks', 'Veri yok'),
                    body_style
                ))
                story.append(Spacer(1, 0.2*inch))

                story.append(Paragraph(
                    "<b>Fırsatlar:</b>",
                    subheading_style
                ))
                story.append(Paragraph(
                    strategy_data.get('short_term_opportunities', 'Veri yok'),
                    body_style
                ))
            else:
                story.append(Paragraph(
                    "Strateji verileri henüz girilmemiş.",
                    body_style
                ))

            story.append(PageBreak())

            # 3. Risk Management
            story.append(Paragraph("3. Risk Yönetimi", heading_style))
            risks = self.manager.get_climate_risks(company_id, year)

            if risks:
                # Risk tablosu
                risk_table_data = [
                    ["Risk Adı", "Kategori", "Derece", "Finansal Etki"]
                ]

                for risk in risks:
                    risk_table_data.append([
                        risk['risk_name'],
                        risk['risk_category'],
                        risk['risk_rating'],
                        f"${risk.get('financial_impact_high', 0):,.0f}" if risk.get('financial_impact_high') else "-"
                    ])

                risk_table = Table(risk_table_data, colWidths=[2.5*inch, 1.2*inch, 0.8*inch, 1*inch])
                risk_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2E8B57')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'NotoSans-Bold'),
                    ('FONTNAME', (0, 1), (-1, -1), 'NotoSans'),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('FONTSIZE', (0, 1), (-1, -1), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')])
                ]))
                story.append(risk_table)
                # Risk ısı haritası (derece dağılımı)
                if GRAPHICS_AVAILABLE:
                    try:
                        rating_counts = {'Low': 0, 'Medium': 0, 'High': 0}
                        for r in risks:
                            val = str(r.get('risk_rating', '')).strip().title()
                            if val in rating_counts:
                                rating_counts[val] += 1
                        data = [[rating_counts['Low'], rating_counts['Medium'], rating_counts['High']]]
                        d = Drawing(400, 180)
                        bc = VerticalBarChart()
                        bc.x = 40
                        bc.y = 30
                        bc.height = 120
                        bc.width = 320
                        bc.data = data
                        bc.categoryAxis.categoryNames = ['Low','Medium','High']
                        bc.valueAxis.labels.fontName = 'NotoSans'
                        bc.categoryAxis.labels.fontName = 'NotoSans'
                        # Renkler: Low=yeşil, Medium=turuncu, High= kırmızı
                        bc.bars[0].fillColor = colors.HexColor('#10B981')
                        # Bar set tek satır, segment renklerini ayarla
                        bc.barLabels.nudge = 6
                        bc.barLabelFormat = '%d'
                        # Farklı kategoriler için renkleri tek tek ata
                        for i, col in enumerate([colors.HexColor('#10B981'), colors.HexColor('#F59E0B'), colors.HexColor('#EF4444')]):
                            try:
                                bc.bars[i].fillColor = col
                            except Exception as e:
                                logging.error(f"Silent error caught: {str(e)}")
                        d.add(bc)
                        story.append(Spacer(1, 0.15*inch))
                        story.append(Paragraph("Risk Dereceleri Isı Grafiği", subheading_style))
                        story.append(d)
                    except Exception as ge:
                        logging.info(f"[WARN] Risk grafiği oluşturulamadı: {ge}")
            else:
                story.append(Paragraph(
                    "İklim riskleri henüz tanımlanmamış.",
                    body_style
                ))

            story.append(PageBreak())

            # 4. Metrics
            story.append(Paragraph("4. Metrikler ve Hedefler", heading_style))
            # Özet metrik tablosu (mevcut verilerden)
            metrics = self.manager.get_metrics(company_id, year)
            if metrics:
                met_table = Table([
                    ["Toplam Emisyon (tCO2e)", f"{(metrics.get('total_emissions') or 0):,.0f}"],
                    ["Enerji Tüketimi (MWh)", f"{(metrics.get('total_energy_consumption') or 0):,.0f}"],
                    ["Su Tüketimi (m³)", f"{(metrics.get('water_consumption') or 0):,.0f}"],
                    ["İç Karbon Fiyatı ($/tCO2e)", f"{(metrics.get('internal_carbon_price') or 0):,.0f}"],
                    ["İklim Geliri (USD)", f"${((metrics.get('climate_related_revenue') or 0)*1_000_000):,.0f}"],
                    ["İklim OPEX (USD)", f"${((metrics.get('climate_related_opex') or 0)*1_000_000):,.0f}"],
                    ["İklim CAPEX (USD)", f"${((metrics.get('climate_related_capex') or 0)*1_000_000):,.0f}"],
                ], colWidths=[3.0*inch, 2.2*inch])
                met_table.setStyle(TableStyle([
                    ('FONT', (0, 0), (-1, -1), 'NotoSans', 11),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#FAFAFA')])
                ]))
                story.append(met_table)
                # Metrik görselleştirme (bar grafik)
                if GRAPHICS_AVAILABLE:
                    try:
                        vals = [
                            float(metrics.get('total_emissions') or 0),
                            float(metrics.get('total_energy_consumption') or 0),
                            float(metrics.get('water_consumption') or 0),
                        ]
                        d2 = Drawing(400, 180)
                        bc2 = VerticalBarChart()
                        bc2.x = 40
                        bc2.y = 30
                        bc2.height = 120
                        bc2.width = 320
                        bc2.data = [vals]
                        bc2.categoryAxis.categoryNames = ['Emisyon','Enerji','Su']
                        bc2.valueAxis.labels.fontName = 'NotoSans'
                        bc2.categoryAxis.labels.fontName = 'NotoSans'
                        bc2.bars[0].fillColor = colors.HexColor('#6366F1')
                        bc2.barLabels.nudge = 6
                        bc2.barLabelFormat = '%0.0f'
                        d2.add(bc2)
                        story.append(Spacer(1, 0.15*inch))
                        story.append(Paragraph("Ana Metrikler Görselleştirmesi", subheading_style))
                        story.append(d2)
                    except Exception as ge:
                        logging.info(f"[WARN] Metrik grafiği oluşturulamadı: {ge}")
            else:
                story.append(Paragraph(
                    "Metrik verileri henüz girilmemiş.",
                    body_style
                ))

            story.append(PageBreak())
            story.append(Paragraph("Dipnotlar", heading_style))
            story.append(Paragraph("Veri Kaynağı: Şirket operasyonel metrikleri (tcfd_metrics), yıl bazlı kayıtlar.", body_style))
            story.append(Paragraph("Varsayımlar: Senaryo karbon fiyatları (modules/tcfd/data/climate_scenarios.json) yoksa dahili karbon fiyatı kullanılır.", body_style))
            story.append(Paragraph("Metodoloji: Karbon maliyeti = emisyon × fiyat; enerji maliyeti değişimi geçiş yüzdesi ve fiyat primi ile hesaplanır.", body_style))

            # PDF oluştur
            doc.build(story)

            logging.info(f"[OK] PDF raporu oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"[ERROR] PDF raporu oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return False

    # ========================================================================
    # DOCX RAPOR
    # ========================================================================

    def generate_docx_report(
        self,
        company_id: int,
        year: int,
        company_name: str,
        output_path: str
    ) -> bool:
        """
        Word DOCX raporu oluştur
        
        Returns:
            Başarılı mı?
        """
        if not DOCX_AVAILABLE:
            logging.error("[ERROR] python-docx not installed")
            return False

        try:
            doc = Document()

            try:
                from modules.reporting.brand_identity_manager import BrandIdentityManager
                base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
                db_path = os.path.join(base_dir, "data", "sdg_desktop.sqlite")
                bim = BrandIdentityManager(db_path, company_id)
                bi = bim.get_brand_identity(company_id)
                lp = bi.get('logo_path')
                if lp and os.path.exists(lp):
                    from docx.shared import Inches
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(lp, width=Inches(1.6))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ht = (bi.get('texts') or {}).get('header')
                if ht:
                    t = doc.add_paragraph(ht)
                    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            # Başlık
            title = _add_turkish_heading(doc, company_name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = _add_turkish_paragraph(doc, "TCFD Raporu")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(19)
            subtitle.runs[0].font.color.rgb = RGBColor(46, 139, 87)

            _add_turkish_paragraph(doc, f"Raporlama Yılı: {year}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_turkish_paragraph(doc, f"Rapor Tarihi: {datetime.now().strftime('%d %B %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_page_break()

            # Yönetici Özeti
            _add_turkish_heading(doc, "Yönetici Özeti", 1)
            _add_turkish_paragraph(doc, "Seçilmiş senaryolar için karbon maliyeti, iklim ilişkili gelir/gider ve CAPEX etkilerinin özeti.")

            mgmt = self._build_management_summary(company_id, year)
            rows = mgmt.get('summary_rows') or []
            if rows:
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Light Grid Accent 1'
                hdr = table.rows[0].cells
                hdr[0].text = 'Senaryo'
                hdr[1].text = 'Karbon Fiyatı'
                hdr[2].text = 'Karbon Maliyeti'
                hdr[3].text = 'İklim Geliri'
                hdr[4].text = 'İklim OPEX'
                hdr[5].text = 'İklim CAPEX'
                hdr[6].text = 'Net Etki'
                hdr[7].text = 'Hassasiyet (Toplam Maliyet)'

                for r in rows:
                    rc = table.add_row().cells
                    for i in range(8):
                        rc[i].text = str(r[i])
            else:
                _add_turkish_paragraph(doc, "Veri yok.")

            doc.add_page_break()

            # 1. Governance
            _add_turkish_heading(doc, "1. Yönetişim (Governance)", 1)
            gov_data = self.manager.get_governance(company_id, year)

            if gov_data:
                _add_turkish_heading(doc, "Yönetim Kurulu Gözetimi", 2)
                _add_turkish_paragraph(doc, gov_data.get('board_oversight', 'Veri yok'))

                _add_turkish_heading(doc, "Üst Yönetimin Rolü", 2)
                _add_turkish_paragraph(doc, gov_data.get('management_role', 'Veri yok'))
            else:
                _add_turkish_paragraph(doc, "Yönetişim verileri henüz girilmemiş.")

            doc.add_page_break()

            # 2. Strategy
            _add_turkish_heading(doc, "2. Strateji (Strategy)", 1)
            strategy_data = self.manager.get_strategy(company_id, year)

            if strategy_data:
                _add_turkish_heading(doc, "Kısa Vadeli Riskler", 2)
                _add_turkish_paragraph(doc, strategy_data.get('short_term_risks', 'Veri yok'))

                _add_turkish_heading(doc, "Fırsatlar", 2)
                _add_turkish_paragraph(doc, strategy_data.get('short_term_opportunities', 'Veri yok'))
            else:
                _add_turkish_paragraph(doc, "Strateji verileri henüz girilmemiş.")

            doc.add_page_break()

            # 3. Risk Management
            _add_turkish_heading(doc, "3. Risk Yönetimi", 1)
            risks = self.manager.get_climate_risks(company_id, year)

            if risks:
                # Risk tablosu
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'

                # Başlıklar
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Risk Adı'
                hdr_cells[1].text = 'Kategori'
                hdr_cells[2].text = 'Derece'
                hdr_cells[3].text = 'Finansal Etki'

                # Riskler
                for risk in risks:
                    row_cells = table.add_row().cells
                    row_cells[0].text = risk['risk_name']
                    row_cells[1].text = risk['risk_category']
                    row_cells[2].text = risk['risk_rating']
                    row_cells[3].text = f"${risk.get('financial_impact_high', 0):,.0f}" if risk.get('financial_impact_high') else "-"
            else:
                _add_turkish_paragraph(doc, "İklim riskleri henüz tanımlanmamış.")

            doc.add_page_break()

            # 4. Metrics
            _add_turkish_heading(doc, "4. Metrikler ve Hedefler", 1)
            _add_turkish_paragraph(doc, "Metrik verileri henüz girilmemiş.")
            doc.add_page_break()
            _add_turkish_heading(doc, "Dipnotlar", 1)
            _add_turkish_paragraph(doc, "Veri Kaynağı: Şirket operasyonel metrikleri (tcfd_metrics), yıl bazlı kayıtlar.")
            _add_turkish_paragraph(doc, "Varsayımlar: Senaryo karbon fiyatları (modules/tcfd/data/climate_scenarios.json) yoksa dahili karbon fiyatı kullanılır.")
            _add_turkish_paragraph(doc, "Metodoloji: Karbon maliyeti = emisyon × fiyat; enerji maliyeti değişimi geçiş yüzdesi ve fiyat primi ile hesaplanır.")

            # Kaydet
            doc.save(output_path)

            logging.info(f"[OK] DOCX raporu oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"[ERROR] DOCX raporu oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return False

    # ========================================================================
    # EXCEL RAPOR
    # ========================================================================

    def generate_excel_report(
        self,
        company_id: int,
        year: int,
        company_name: str,
        output_path: str
    ) -> bool:
        """
        Excel raporu oluştur (veri tabloları)
        
        Returns:
            Başarılı mı?
        """
        if not EXCEL_AVAILABLE:
            logging.error("[ERROR] openpyxl not installed")
            return False

        try:
            wb = openpyxl.Workbook()

            # Sheet 1: Overview
            ws1 = wb.active
            ws1.title = "Overview"
            ws1['A1'] = "TCFD Raporu"
            ws1['A1'].font = Font(size=19, bold=True, color="2E8B57")
            ws1['A2'] = company_name
            ws1['A3'] = f"Yıl: {year}"
            ws1['A4'] = f"Tarih: {datetime.now().strftime('%d/%m/%Y')}"

            # Sheet 2: Governance
            ws2 = wb.create_sheet("Governance")
            gov_data = self.manager.get_governance(company_id, year)

            if gov_data:
                ws2['A1'] = "Yönetişim Verileri"
                ws2['A1'].font = Font(bold=True, size=14)

                row = 3
                for key, value in gov_data.items():
                    if key not in ['id', 'company_id', 'reporting_year', 'created_at', 'updated_at', 'created_by']:
                        ws2[f'A{row}'] = key.replace('_', ' ').title()
                        ws2[f'B{row}'] = str(value) if value else ""
                        row += 1

            # Sheet 3: Risks
            ws3 = wb.create_sheet("Climate Risks")
            risks = self.manager.get_climate_risks(company_id, year)

            # Başlıklar
            headers = ["Risk Adı", "Kategori", "Tip", "Olasılık", "Etki", "Derece",
                      "Zaman Ufku", "Finansal Etki (Min)", "Finansal Etki (Maks)", "Durum"]
            for col, header in enumerate(headers, 1):
                cell = ws3.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="2E8B57", end_color="2E8B57", fill_type="solid")
                cell.font = Font(color="FFFFFF", bold=True)

            # Riskler
            for row, risk in enumerate(risks, 2):
                ws3.cell(row=row, column=1, value=risk['risk_name'])
                ws3.cell(row=row, column=2, value=risk['risk_category'])
                ws3.cell(row=row, column=3, value=risk['risk_type'])
                ws3.cell(row=row, column=4, value=risk.get('likelihood', ''))
                ws3.cell(row=row, column=5, value=risk.get('impact', ''))
                ws3.cell(row=row, column=6, value=risk['risk_rating'])
                ws3.cell(row=row, column=7, value=risk.get('time_horizon', ''))
                ws3.cell(row=row, column=8, value=risk.get('financial_impact_low', ''))
                ws3.cell(row=row, column=9, value=risk.get('financial_impact_high', ''))
                ws3.cell(row=row, column=10, value=risk.get('status', ''))

            # Sütun genişlikleri
            for col in range(1, 11):
                ws3.column_dimensions[openpyxl.utils.get_column_letter(col)].width = 15

            # Kaydet
            wb.save(output_path)

            logging.info(f"[OK] Excel raporu oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"[ERROR] Excel raporu oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return False


# Test
if __name__ == "__main__":
    logging.info(" TCFD Report Generator Test")
    logging.info("="*60)

    # Test için manager gerekli
    from tcfd_calculator import TCFDCalculator
    from tcfd_manager import TCFDManager

    db_path = "../../data/sdg_desktop.sqlite"

    if os.path.exists(db_path):
        manager = TCFDManager(db_path)
        calculator = TCFDCalculator()
        generator = TCFDReportGenerator(manager, calculator)

        logging.info("\n Report Generator başlatıldı")
        logging.info(f"   PDF: {REPORTLAB_AVAILABLE}")
        logging.info(f"   DOCX: {DOCX_AVAILABLE}")
        logging.info(f"   Excel: {EXCEL_AVAILABLE}")
    else:
        logging.info(" Veritabanı bulunamadı")

    logging.info("="*60)

