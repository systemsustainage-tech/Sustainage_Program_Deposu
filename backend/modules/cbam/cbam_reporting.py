#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import logging
import os
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .cbam_manager import CBAMManager


class CBAMReporting:
    def __init__(self, manager: CBAMManager) -> None:
        self.manager = manager

    def generate_cbam_report(self, company_id: int, period: str, formats: List[str] | None = None) -> Dict[str, str]:
        if formats is None:
            formats = ["docx"]

        generated: Dict[str, str] = {}

        if "docx" in formats:
            path = self._generate_docx_report(company_id, period)
            if path:
                generated["docx"] = path

        return generated

    def _generate_docx_report(self, company_id: int, period: str) -> str | None:
        try:
            summary = self.manager.calculate_cbam_liability(company_id, period)
            imports = self.manager.get_imports(company_id, period)

            base_dir = os.path.dirname(os.path.dirname(self.manager.db_path))
            output_dir = os.path.join(base_dir, "data", "companies", str(company_id), "reports")
            os.makedirs(output_dir, exist_ok=True)

            filename = f"cbam_report_{period}_{datetime.now().strftime('%Y%m%d')}.docx"
            filepath = os.path.join(output_dir, filename)

            doc = Document()

            title = doc.add_heading("CBAM RAPORU", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_paragraph(f"Raporlama Dönemi: {period}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(12)

            doc.add_paragraph(f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y')}")

            doc.add_page_break()

            doc.add_heading("1. ÖZET", level=1)
            doc.add_paragraph(f"Toplam Emisyon: {summary.get('total_emissions', 0):.2f} tCO2e")
            doc.add_paragraph(f"Toplam İthalat Miktarı: {summary.get('total_quantity', 0):.2f} ton")
            doc.add_paragraph(f"EU ETS Fiyatı: {summary.get('eu_ets_price', 0):.2f} €/tCO2")
            doc.add_paragraph(f"Ödenen Karbon Fiyatı: {summary.get('carbon_price_paid', 0):.2f} €")
            doc.add_paragraph(f"Ham CBAM Yükümlülüğü: {summary.get('cbam_liability_raw', 0):.2f} €")
            doc.add_paragraph(f"De-minimis Eşiği: {summary.get('de_minimis_threshold', 50.0):.1f} ton")
            doc.add_paragraph(f"Kapsanan Miktar: {summary.get('covered_quantity', 0):.2f} ton")
            below = summary.get("below_de_minimis", False)
            doc.add_paragraph(f"De-minimis Durumu: {'Eşik Altı (muaf)' if below else 'Eşik Üstü (yükümlü)'}")
            doc.add_paragraph(f"Nihai CBAM Yükümlülüğü: {summary.get('cbam_liability', 0):.2f} €")

            doc.add_page_break()

            doc.add_heading("2. ÜRÜN BAZINDA CBAM ÖZETİ", level=1)

            table = doc.add_table(rows=1, cols=6)
            hdr = table.rows[0].cells
            hdr[0].text = "Ürün"
            hdr[1].text = "Sektör"
            hdr[2].text = "Miktar"
            hdr[3].text = "Emisyon"
            hdr[4].text = "Ödenen Karbon"
            hdr[5].text = "Dönem"

            for imp in imports:
                row = table.add_row().cells
                row[0].text = str(imp.get("product_name") or imp.get("product_code") or "")
                row[1].text = str(imp.get("sector") or "")
                row[2].text = f"{float(imp.get('quantity') or 0):.2f}"
                row[3].text = f"{float(imp.get('embedded_emissions') or 0):.2f}"
                row[4].text = f"{float(imp.get('carbon_price_paid') or 0):.2f}"
                row[5].text = str(imp.get("import_period") or "")

            doc.save(filepath)
            logging.info(f"[OK] CBAM raporu oluşturuldu: {filepath}")
            return filepath
        except Exception as e:
            logging.error(f"CBAM DOCX raporu oluşturma hatası: {e}")
            return None

