import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SU RAPORLAMA SINIFI
Su ayak izi raporları ve analiz çıktıları
"""

import os
from datetime import datetime
from typing import Dict, List

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.info("[UYARI] python-docx yüklü değil. DOCX raporlar oluşturulamaz.")

try:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.info("[UYARI] openpyxl yuklu degil. Excel raporlar olusturulamaz.")

from config.settings import ensure_directories, get_db_path, get_export_dir

from utils.language_manager import LanguageManager
from .water_manager import WaterManager


def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    para = doc.add_paragraph(text if text is not None else '', style=style)
    if not text:
        return para
    for run in para.runs:
        try:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        try:
            run.font.name = font_name
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return heading


class WaterReporting:
    """Su raporlama sınıfı - ISO 14046 ve GRI uyumlu"""

    def __init__(self, db_path: str | None = None) -> None:
        if db_path:
            self.db_path = db_path
        else:
            ensure_directories()
            self.db_path = get_db_path()
        self.manager = WaterManager(self.db_path)
        self.lm = LanguageManager()

    def generate_water_footprint_report(self, company_id: int, period: str,
                                      include_scope3: bool = False) -> str:
        """
        Su ayak izi raporu oluştur (DOCX)
        
        Returns:
            Rapor dosya yolu
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kütüphanesi gerekli")

        # Su ayak izi hesapla
        footprint_data = self.manager.calculate_water_footprint(company_id, period)
        # Tüketim kayıtlarını al
        consumption_records = self.manager.get_water_consumption(company_id, period)

        if not footprint_data:
            raise ValueError("Su ayak izi verisi bulunamadı")

        # Rapor dosyası oluştur
        report_filename = f"water_footprint_report_{company_id}_{period}.docx"
        export_dir = str(get_export_dir(company_id))
        report_path = os.path.join(export_dir, report_filename)

        # Klasör yoksa oluştur
        os.makedirs(os.path.dirname(report_path), exist_ok=True)

        # Dokuman oluştur
        doc = Document()

        # Logo
        try:
            from modules.reporting.brand_identity_manager import BrandIdentityManager
            bim = BrandIdentityManager(self.db_path, company_id)
            bi = bim.get_brand_identity(company_id)
            logo_path = bi.get('logo_path')
            
            if logo_path and os.path.exists(logo_path):
                from docx.shared import Inches
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(logo_path, width=Inches(1.6))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            pass

        # Başlık
        title = _add_turkish_heading(doc, self.lm.tr('water_footprint_report', 'SU AYAK İZİ RAPORU'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Alt başlık
        subtitle = _add_turkish_paragraph(doc, f"{self.lm.tr('reporting_period', 'Raporlama Dönemi')}: {period}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)

        _add_turkish_paragraph(doc, f"{self.lm.tr('report_date', 'Rapor Tarihi')}: {datetime.now().strftime('%d.%m.%Y')}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('reporting_standard', 'Raporlama Standardı')}: ISO 14046")

        doc.add_page_break()

        # 1. Yönetici Özeti
        _add_turkish_heading(doc, f"1. {self.lm.tr('executive_summary', 'YÖNETİCİ ÖZETİ')}", 1)
        _add_turkish_paragraph(doc,
            self.lm.tr('water_report_intro', f"Bu rapor, {period} yılı su ayak izi envanterini ISO 14046 standartlarına uygun olarak sunmaktadır.").format(period=period)
        )

        _add_turkish_paragraph(doc, f"\n{self.lm.tr('total_water_footprint', 'Toplam Su Ayak İzi')}: {footprint_data['total_water_footprint']:.2f} m³")
        _add_turkish_paragraph(doc, f"  • {self.lm.tr('blue_water', 'Mavi Su')}: {footprint_data['total_blue_water']:.2f} m³")
        _add_turkish_paragraph(doc, f"  • {self.lm.tr('green_water', 'Yeşil Su')}: {footprint_data['total_green_water']:.2f} m³")
        _add_turkish_paragraph(doc, f"  • {self.lm.tr('grey_water', 'Gri Su')}: {footprint_data['total_grey_water']:.2f} m³")

        # 2. Metodoloji
        _add_turkish_heading(doc, f"2. {self.lm.tr('methodology', 'METODOLOJİ')}", 1)
        _add_turkish_paragraph(doc,
            self.lm.tr('water_methodology_desc', "Su ayak izi hesaplamaları ISO 14046:2014 standardına uygun olarak yapılmıştır. Water Footprint Network metodolojisi kullanılmıştır.")
        )

        # 3. Su Ayak İzi Sonuçları
        _add_turkish_heading(doc, f"3. {self.lm.tr('water_footprint_results', 'SU AYAK İZİ SONUÇLARI')}", 1)

        # 3.1 Toplam Sonuçlar
        _add_turkish_heading(doc, f"3.1 {self.lm.tr('total_results', 'Toplam Sonuçlar')}", 2)

        results_table = doc.add_table(rows=1, cols=3)
        results_table.style = 'Table Grid'

        # Tablo başlıkları
        hdr_cells = results_table.rows[0].cells
        hdr_cells[0].text = self.lm.tr('water_type', 'Su Türü')
        hdr_cells[1].text = f"{self.lm.tr('amount', 'Miktar')} (m³)"
        hdr_cells[2].text = f"{self.lm.tr('percentage', 'Yüzde')} (%)"

        # Veriler
        total = footprint_data['total_water_footprint']
        blue_pct = (footprint_data['total_blue_water'] / total * 100) if total > 0 else 0
        green_pct = (footprint_data['total_green_water'] / total * 100) if total > 0 else 0
        grey_pct = (footprint_data['total_grey_water'] / total * 100) if total > 0 else 0

        data_rows = [
            [self.lm.tr('blue_water', 'Mavi Su'), f"{footprint_data['total_blue_water']:.2f}", f"{blue_pct:.1f}%"],
            [self.lm.tr('green_water', 'Yeşil Su'), f"{footprint_data['total_green_water']:.2f}", f"{green_pct:.1f}%"],
            [self.lm.tr('grey_water', 'Gri Su'), f"{footprint_data['total_grey_water']:.2f}", f"{grey_pct:.1f}%"],
            [self.lm.tr('total', 'TOPLAM'), f"{total:.2f}", "100.0%"]
        ]

        for row_data in data_rows:
            row_cells = results_table.add_row().cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data

        # 3.2 Tüketim Türüne Göre Dağılım
        _add_turkish_heading(doc, f"3.2 {self.lm.tr('distribution_by_type', 'Tüketim Türüne Göre Dağılım')}", 2)

        breakdown_by_type = footprint_data.get('breakdown_by_type', {})
        if breakdown_by_type:
            type_table = doc.add_table(rows=len(breakdown_by_type) + 1, cols=5)
            type_table.style = 'Table Grid'

            # Başlıklar
            type_hdr = type_table.rows[0].cells
            type_hdr[0].text = self.lm.tr('consumption_type', 'Tüketim Türü')
            type_hdr[1].text = f"{self.lm.tr('blue_water', 'Mavi Su')} (m³)"
            type_hdr[2].text = f"{self.lm.tr('green_water', 'Yeşil Su')} (m³)"
            type_hdr[3].text = f"{self.lm.tr('grey_water', 'Gri Su')} (m³)"
            type_hdr[4].text = f"{self.lm.tr('total', 'Toplam')} (m³)"

            # Veriler
            for i, (cons_type, data) in enumerate(breakdown_by_type.items(), 1):
                type_row = type_table.rows[i].cells
                type_row[0].text = self.lm.tr(cons_type, cons_type.title())
                type_row[1].text = f"{data['blue_water']:.2f}"
                type_row[2].text = f"{data['green_water']:.2f}"
                type_row[3].text = f"{data['grey_water']:.2f}"
                type_row[4].text = f"{data['total']:.2f}"

        # 3.3 Tüketim Kayıtları
        if consumption_records:
            _add_turkish_heading(doc, f"3.3 {self.lm.tr('consumption_records', 'Tüketim Kayıtları')}", 2)
            
            records_table = doc.add_table(rows=1, cols=6)
            records_table.style = 'Table Grid'
            
            hdr_cells = records_table.rows[0].cells
            hdr_cells[0].text = self.lm.tr('consumption_type', 'Tüketim Türü')
            hdr_cells[1].text = f"{self.lm.tr('amount', 'Miktar')} (m³)"
            hdr_cells[2].text = self.lm.tr('source_type', 'Kaynak')
            hdr_cells[3].text = 'Fatura Tarihi'
            hdr_cells[4].text = 'Son Ödeme'
            hdr_cells[5].text = 'Tedarikçi'
            
            for record in consumption_records:
                row_cells = records_table.add_row().cells
                row_cells[0].text = record.get('consumption_type', '')
                row_cells[1].text = f"{record.get('total_water', 0):.2f}"
                row_cells[2].text = record.get('source_type', '')
                row_cells[3].text = record.get('invoice_date', '') or '-'
                row_cells[4].text = record.get('due_date', '') or '-'
                row_cells[5].text = record.get('supplier', '') or '-'

        # 4. Verimlilik Metrikleri
        _add_turkish_heading(doc, f"4. {self.lm.tr('efficiency_metrics', 'VERİMLİLİK METRİKLERİ')}", 1)

        efficiency_metrics = footprint_data.get('efficiency_metrics', {})
        if efficiency_metrics:
            avg_eff_ratio = efficiency_metrics.get('average_efficiency_ratio', 0)
            avg_recycle = efficiency_metrics.get('average_recycling_rate', 0)
            _add_turkish_paragraph(doc, f"{self.lm.tr('avg_efficiency_rate', 'Ortalama Verimlilik Oranı')}: {avg_eff_ratio:.3f}")
            _add_turkish_paragraph(doc, f"{self.lm.tr('avg_recycling_rate', 'Ortalama Geri Dönüşüm Oranı')}: {avg_recycle:.3f}")
            _add_turkish_paragraph(doc, f"{self.lm.tr('total_records', 'Toplam Kayıt Sayısı')}: {efficiency_metrics.get('total_records', 0)}")

        # 5. SDG 6 İlerlemesi
        _add_turkish_heading(doc, f"5. {self.lm.tr('sdg6_progress', 'SDG 6 İLERLEMESİ')}", 1)
        _add_turkish_paragraph(doc,
            self.lm.tr('sdg6_desc', "Su ayak izi yönetimi, Birleşmiş Milletler Sürdürülebilir Kalkınma Hedefi 6 (Temiz Su ve Sanitasyon) ile doğrudan ilgilidir.")
        )

        # 6. Öneriler
        _add_turkish_heading(doc, f"6. {self.lm.tr('recommendations', 'ÖNERİLER')}", 1)

        recommendations = [
            self.lm.tr('rec_efficiency', "Su tüketimini azaltmak için verimlilik projeleri uygulayın"),
            self.lm.tr('rec_recycle', "Gri su geri dönüşüm sistemleri kurun"),
            self.lm.tr('rec_quality', "Su kalitesi izleme programlarını güçlendirin"),
            self.lm.tr('rec_training', "Çalışanlara su tasarrufu konusunda eğitim verin"),
            self.lm.tr('rec_update', "Su ayak izi hesaplamalarını düzenli olarak güncelleyin")
        ]

        for i, rec in enumerate(recommendations, 1):
            _add_turkish_paragraph(doc, f"{i}. {rec}")

        # Dokumanı kaydet
        doc.save(report_path)

        return report_path

    def generate_excel_report(self, company_id: int, period: str) -> str:
        """Excel su ayak izi raporu oluştur"""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl kütüphanesi gerekli")

        # Su ayak izi hesapla
        footprint_data = self.manager.calculate_water_footprint(company_id, period)
        # Tüketim kayıtlarını al
        consumption_records = self.manager.get_water_consumption(company_id, period)

        if not footprint_data:
            raise ValueError("Su ayak izi verisi bulunamadı")

        # Excel dosyası oluştur
        report_filename = f"water_footprint_report_{company_id}_{period}.xlsx"
        report_path = os.path.join("data/exports", report_filename)

        # Klasör yoksa oluştur
        os.makedirs(os.path.dirname(report_path), exist_ok=True)

        # Excel workbook oluştur
        wb = openpyxl.Workbook()

        # Stil tanımları
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")

        # 1. Özet Sayfası
        ws_summary = wb.active
        ws_summary.title = self.lm.tr('water_footprint_summary', "Su Ayak İzi Özeti")

        # Başlık
        ws_summary['A1'] = f"{self.lm.tr('water_footprint_report', 'SU AYAK İZİ RAPORU').upper()} - {period}"
        ws_summary['A1'].font = Font(size=16, bold=True)
        ws_summary.merge_cells('A1:D1')

        # Tarih
        ws_summary['A2'] = f"{self.lm.tr('report_date', 'Rapor Tarihi')}: {datetime.now().strftime('%d.%m.%Y')}"
        ws_summary['A2'].font = Font(size=10)

        # Boş satır
        ws_summary['A3'] = ""

        # Toplam sonuçlar
        ws_summary['A4'] = self.lm.tr('total_water_footprint_results', "TOPLAM SU AYAK İZİ SONUÇLARI")
        ws_summary['A4'].font = Font(size=12, bold=True)

        # Tablo başlıkları
        headers = [
            self.lm.tr('water_type', "Su Türü"), 
            f"{self.lm.tr('amount', 'Miktar')} (m³)", 
            f"{self.lm.tr('percentage', 'Yüzde')} (%)"
        ]
        for col, header in enumerate(headers, 1):
            cell = ws_summary.cell(row=5, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        # Veriler
        total = footprint_data['total_water_footprint']
        blue_pct = (footprint_data['total_blue_water'] / total * 100) if total > 0 else 0
        green_pct = (footprint_data['total_green_water'] / total * 100) if total > 0 else 0
        grey_pct = (footprint_data['total_grey_water'] / total * 100) if total > 0 else 0

        data_rows = [
            [self.lm.tr('blue_water', "Mavi Su"), footprint_data['total_blue_water'], f"{blue_pct:.1f}%"],
            [self.lm.tr('green_water', "Yeşil Su"), footprint_data['total_green_water'], f"{green_pct:.1f}%"],
            [self.lm.tr('grey_water', "Gri Su"), footprint_data['total_grey_water'], f"{grey_pct:.1f}%"],
            [self.lm.tr('total', "TOPLAM"), total, "100.0%"]
        ]

        for row, data in enumerate(data_rows, 6):
            for col, value in enumerate(data, 1):
                ws_summary.cell(row=row, column=col, value=value)

        # 2. Tüketim Türü Dağılımı Sayfası
        ws_breakdown = wb.create_sheet(self.lm.tr('consumption_breakdown', "Tüketim Türü Dağılımı"))

        # Başlık
        ws_breakdown['A1'] = self.lm.tr('distribution_by_type_upper', "TÜKETİM TÜRÜNE GÖRE DAĞILIM")
        ws_breakdown['A1'].font = Font(size=14, bold=True)
        ws_breakdown.merge_cells('A1:F1')

        # Tablo başlıkları
        breakdown_headers = [
            self.lm.tr('consumption_type', "Tüketim Türü"),
            f"{self.lm.tr('blue_water', 'Mavi Su')} (m³)",
            f"{self.lm.tr('green_water', 'Yeşil Su')} (m³)",
            f"{self.lm.tr('grey_water', 'Gri Su')} (m³)",
            f"{self.lm.tr('total', 'Toplam')} (m³)"
        ]
        for col, header in enumerate(breakdown_headers, 1):
            cell = ws_breakdown.cell(row=3, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        # Veriler
        breakdown_by_type = footprint_data.get('breakdown_by_type', {})
        row = 4
        for cons_type, data in breakdown_by_type.items():
            ws_breakdown.cell(row=row, column=1, value=self.lm.tr(cons_type, cons_type.title()))
            ws_breakdown.cell(row=row, column=2, value=data['blue_water'])
            ws_breakdown.cell(row=row, column=3, value=data['green_water'])
            ws_breakdown.cell(row=row, column=4, value=data['grey_water'])
            ws_breakdown.cell(row=row, column=5, value=data['total'])
            row += 1

        # 3. Verimlilik Metrikleri Sayfası
        ws_efficiency = wb.create_sheet(self.lm.tr('efficiency_metrics', "Verimlilik Metrikleri"))

        # Başlık
        ws_efficiency['A1'] = self.lm.tr('efficiency_metrics_upper', "VERİMLİLİK METRİKLERİ")
        ws_efficiency['A1'].font = Font(size=14, bold=True)
        ws_efficiency.merge_cells('A1:B1')

        # Metrikler
        efficiency_metrics = footprint_data.get('efficiency_metrics', {})
        metrics_data = [
            [self.lm.tr('avg_efficiency_rate', "Ortalama Verimlilik Oranı"), efficiency_metrics.get('average_efficiency_ratio', 0)],
            [self.lm.tr('avg_recycling_rate', "Ortalama Geri Dönüşüm Oranı"), efficiency_metrics.get('average_recycling_rate', 0)],
            [self.lm.tr('total_records', "Toplam Kayıt Sayısı"), efficiency_metrics.get('total_records', 0)],
            [self.lm.tr('high_quality_data', "Yüksek Kalite Veri"), efficiency_metrics.get('high_quality_data', 0)],
            [self.lm.tr('medium_quality_data', "Orta Kalite Veri"), efficiency_metrics.get('medium_quality_data', 0)],
            [self.lm.tr('low_quality_data', "Düşük Kalite Veri"), efficiency_metrics.get('low_quality_data', 0)]
        ]

        row = 3
        for metric_name, metric_value in metrics_data:
            ws_efficiency.cell(row=row, column=1, value=metric_name)
            ws_efficiency.cell(row=row, column=2, value=metric_value)
            row += 1

        # 4. Tüketim Kayıtları Sayfası
        if consumption_records:
            ws_records = wb.create_sheet(self.lm.tr('consumption_records', "Tüketim Kayıtları"))
            
            # Başlık
            ws_records['A1'] = "SU TÜKETİM KAYITLARI"
            ws_records['A1'].font = Font(size=14, bold=True)
            ws_records.merge_cells('A1:H1')
            
            # Tablo başlıkları
            record_headers = [
                self.lm.tr('consumption_type', "Tüketim Türü"),
                f"{self.lm.tr('amount', 'Miktar')} (m³)",
                self.lm.tr('source_type', "Kaynak"),
                self.lm.tr('location', "Konum"),
                "Maliyet",
                "Fatura Tarihi",
                "Son Ödeme Tarihi",
                "Tedarikçi"
            ]
            
            for col, header in enumerate(record_headers, 1):
                cell = ws_records.cell(row=3, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                
            # Veriler
            row = 4
            for record in consumption_records:
                ws_records.cell(row=row, column=1, value=record.get('consumption_type', ''))
                ws_records.cell(row=row, column=2, value=record.get('total_water', 0))
                ws_records.cell(row=row, column=3, value=record.get('source_type', ''))
                ws_records.cell(row=row, column=4, value=record.get('location', ''))
                ws_records.cell(row=row, column=5, value=record.get('cost', 0))
                ws_records.cell(row=row, column=6, value=record.get('invoice_date', ''))
                ws_records.cell(row=row, column=7, value=record.get('due_date', ''))
                ws_records.cell(row=row, column=8, value=record.get('supplier', ''))
                row += 1

        # Sütun genişliklerini ayarla
        sheets_to_adjust = [ws_summary, ws_breakdown, ws_efficiency]
        if 'ws_records' in locals():
            sheets_to_adjust.append(ws_records)

        for ws in sheets_to_adjust:
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except Exception as e:
                        logging.error(f"Silent error caught: {str(e)}")
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

        # Excel dosyasını kaydet
        wb.save(report_path)

        return report_path

    def generate_sdg_6_progress_report(self, company_id: int, period: str) -> Dict:
        """SDG 6 ilerleme raporu oluştur"""
        try:
            # Su ayak izi hesapla
            footprint_data = self.manager.calculate_water_footprint(company_id, period)

            if not footprint_data:
                return {}

            # SDG 6 bileşenleri
            sdg_6_components = {
                'access_to_drinking_water': 95,    # Varsayılan değer
                'access_to_sanitation': 90,        # Varsayılan değer
                'water_quality': 85,               # Su kalitesi skorlarından
                'water_efficiency': 80,            # Verimlilik metriklerinden
                'water_governance': 75             # Varsayılan değer
            }

            # Verimlilik metriklerini SDG 6 skoruna dönüştür
            efficiency_metrics = footprint_data.get('efficiency_metrics', {})
            avg_efficiency = efficiency_metrics.get('average_efficiency_ratio', 0.5)
            avg_recycling = efficiency_metrics.get('average_recycling_rate', 0.2)

            # Verimlilik skorunu hesapla (0-100)
            efficiency_score = min(100, (avg_efficiency + avg_recycling) * 50)
            sdg_6_components['water_efficiency'] = efficiency_score

            # Su kalitesi skorunu hesapla (örnek)
            quality_score = 85  # Gerçek uygulamada kalite ölçümlerinden hesaplanır
            sdg_6_components['water_quality'] = quality_score

            # Toplam SDG 6 skoru
            total_score = sum(sdg_6_components.values()) / len(sdg_6_components)

            return {
                'period': period,
                'sdg_6_score': round(total_score, 1),
                'target_2030': 100.0,
                'progress_percentage': round(total_score, 1),
                'component_scores': sdg_6_components,
                'water_footprint_summary': {
                    'total_water_footprint': footprint_data['total_water_footprint'],
                    'blue_water': footprint_data['total_blue_water'],
                    'green_water': footprint_data['total_green_water'],
                    'grey_water': footprint_data['total_grey_water']
                },
                'efficiency_metrics': efficiency_metrics,
                'calculated_at': datetime.now().isoformat()
            }

        except Exception as e:
            logging.error(f"{self.lm.tr('sdg6_report_error', 'SDG 6 ilerleme raporu hatası')}: {e}")
            return {}

    def export_water_data_to_csv(self, company_id: int, period: str, data_type: str = 'consumption') -> str:
        """Su verilerini CSV olarak dışa aktar"""
        try:
            if data_type == 'consumption':
                data = self.manager.get_water_consumption(company_id, period)
                filename = f"water_consumption_{company_id}_{period}.csv"
            elif data_type == 'targets':
                data = self.manager.get_water_targets(company_id)
                filename = f"water_targets_{company_id}_{period}.csv"
            elif data_type == 'projects':
                data = self.manager.get_efficiency_projects(company_id)
                filename = f"water_projects_{company_id}_{period}.csv"
            elif data_type == 'quality':
                data = self.manager.get_quality_measurements(company_id)
                filename = f"water_quality_{company_id}_{period}.csv"
            else:
                raise ValueError(self.lm.tr('invalid_data_type', "Geçersiz veri türü"))

            if not data:
                raise ValueError(self.lm.tr('no_data_to_export', "Dışa aktarılacak veri bulunamadı"))

            # CSV dosya yolu (firma bazlı)
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            csv_path = os.path.join(base_dir, "data", "companies", str(company_id), "exports", filename)
            os.makedirs(os.path.dirname(csv_path), exist_ok=True)

            # CSV oluştur
            import csv
            with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
                if data:
                    fieldnames = data[0].keys()
                    writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                    writer.writeheader()
                    writer.writerows(data)

            return csv_path

        except Exception as e:
            logging.error(f"CSV dışa aktarma hatası: {e}")
            return ""

    def generate_water_trend_report(self, company_id: int, years: List[str]) -> Dict:
        """Su tüketimi trend raporu oluştur"""
        try:
            trend_data = []

            for year in years:
                footprint_data = self.manager.calculate_water_footprint(company_id, year)
                if footprint_data:
                    trend_data.append({
                        'year': year,
                        'total_consumption': footprint_data['total_water_footprint'],
                        'blue_water': footprint_data['total_blue_water'],
                        'green_water': footprint_data['total_green_water'],
                        'grey_water': footprint_data['total_grey_water']
                    })

            if len(trend_data) < 2:
                return {'error': self.lm.tr('insufficient_data_for_trend', 'Trend analizi için yeterli veri yok')}

            # Trend hesaplama
            total_change = 0
            annual_changes = {}

            for i in range(1, len(trend_data)):
                current_year = trend_data[i]['year']
                # previous_year = trend_data[i-1]['year']

                current_total = trend_data[i]['total_consumption']
                previous_total = trend_data[i-1]['total_consumption']

                if previous_total > 0:
                    change_percentage = ((current_total - previous_total) / previous_total) * 100
                else:
                    change_percentage = 0

                annual_changes[current_year] = round(change_percentage, 1)
                total_change += change_percentage

            # Ortalama trend
            avg_change = total_change / (len(trend_data) - 1)

            # Trend yönü
            if avg_change > 2:
                trend_direction = 'increasing'
            elif avg_change < -2:
                trend_direction = 'decreasing'
            else:
                trend_direction = 'stable'

            return {
                'trend_direction': trend_direction,
                'trend_percentage': round(avg_change, 1),
                'annual_changes': annual_changes,
                'trend_data': trend_data,
                'years_analyzed': years,
                'calculated_at': datetime.now().isoformat()
            }

        except Exception as e:
            logging.error(f"Trend raporu hatası: {e}")
            return {}
