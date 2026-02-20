#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Scope 3 Kategorileri Yöneticisi
GHG Protocol Scope 3 kategorileri için veri yönetimi
"""

import logging
import os
from datetime import datetime
from typing import Dict, List, Optional
from config.database import DB_PATH
from backend.core.base_manager import BaseTenantManager


class Scope3Manager(BaseTenantManager):
    """Scope 3 kategorileri yöneticisi - GHG Protocol uyumlu"""

    def __init__(self, db_path: str = DB_PATH, company_id: Optional[int] = None) -> None:
        super().__init__(db_path, company_id)
        self.create_tables()
        self.load_scope3_categories()

    def create_tables(self) -> None:
        """Gerekli tabloları oluştur"""
        try:
            # Scope 3 kategorileri tablosu
            self.execute_update("""
                CREATE TABLE IF NOT EXISTS scope3_categories (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    category_number INTEGER UNIQUE NOT NULL,
                    category_name TEXT NOT NULL,
                    description TEXT,
                    scope_type TEXT DEFAULT 'Indirect',
                    is_upstream BOOLEAN DEFAULT 1,
                    is_downstream BOOLEAN DEFAULT 0,
                    is_active BOOLEAN DEFAULT 1,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            # Scope 3 emisyon kayıtları tablosu
            self.execute_update("""
                CREATE TABLE IF NOT EXISTS scope3_emissions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    company_id INTEGER NOT NULL,
                    category_id INTEGER NOT NULL,
                    activity_data REAL,
                    activity_unit TEXT,
                    emission_factor REAL,
                    emission_factor_unit TEXT,
                    total_emissions REAL,
                    reporting_period TEXT,
                    data_source TEXT,
                    methodology TEXT,
                    uncertainty_level TEXT,
                    notes TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY(company_id) REFERENCES companies(id),
                    FOREIGN KEY(category_id) REFERENCES scope3_categories(id)
                )
            """)

            # Scope 3 hedefleri tablosu
            self.execute_update("""
                CREATE TABLE IF NOT EXISTS scope3_targets (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    company_id INTEGER NOT NULL,
                    category_id INTEGER,
                    target_type TEXT NOT NULL,
                    baseline_year INTEGER,
                    target_year INTEGER,
                    baseline_emissions REAL,
                    target_emissions REAL,
                    reduction_percentage REAL,
                    target_description TEXT,
                    status TEXT DEFAULT 'Active',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY(company_id) REFERENCES companies(id),
                    FOREIGN KEY(category_id) REFERENCES scope3_categories(id)
                )
            """)

            logging.info("[OK] Scope 3 tabloları oluşturuldu")

        except Exception as e:
            logging.error(f"[HATA] Scope 3 tablo oluşturma hatası: {e}")

    def add_emission_record(self, emission_data: Dict) -> bool:
        """Scope 3 emisyon kaydı ekle"""
        try:
            # Kategori ID'sini bul
            category_name = emission_data['category']
            # scope3_categories is global
            category_row = self.execute_query("SELECT id FROM scope3_categories WHERE category_name = ?", (category_name,))
            
            if not category_row:
                # Kategori yoksa ekle
                category_number = category_name.split('.')[0] if '.' in category_name else '1'
                self.execute_update("""
                    INSERT INTO scope3_categories (category_number, category_name, description)
                    VALUES (?, ?, ?)
                """, (int(category_number), category_name, f"Scope 3 Kategori {category_number}"))
                
                category_row = self.execute_query("SELECT id FROM scope3_categories WHERE category_name = ?", (category_name,))
                category_id = category_row[0]['id']
            else:
                category_id = category_row[0]['id']

            # Emisyon kaydını ekle
            self.execute_update("""
                INSERT INTO scope3_emissions (
                    company_id, category_id, activity_data, activity_unit,
                    emission_factor, total_emissions, reporting_period,
                    data_source, notes, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                emission_data['company_id'],
                category_id,
                emission_data['activity_data'],
                emission_data['unit'],
                emission_data['emission_factor'],
                emission_data['total_emission'],
                emission_data['period'],
                emission_data['source'],
                emission_data['notes'],
                datetime.now().isoformat(),
                datetime.now().isoformat()
            ))

            return True

        except Exception as e:
            logging.error(f"Scope 3 emisyon kaydı ekleme hatası: {e}")
            return False

    def add_target_record(self, target_data: Dict) -> bool:
        """Scope 3 hedef kaydı ekle"""
        try:
            # Kategori ID'sini bul
            category_name = target_data['category']
            category_row = self.execute_query("SELECT id FROM scope3_categories WHERE category_name = ?", (category_name,))
            
            if not category_row:
                # Kategori yoksa ekle
                category_number = category_name.split('.')[0] if '.' in category_name else '1'
                self.execute_update("""
                    INSERT INTO scope3_categories (category_number, category_name, description)
                    VALUES (?, ?, ?)
                """, (int(category_number), category_name, f"Scope 3 Kategori {category_number}"))
                
                category_row = self.execute_query("SELECT id FROM scope3_categories WHERE category_name = ?", (category_name,))
                category_id = category_row[0]['id']
            else:
                category_id = category_row[0]['id']

            # Hedef kaydını ekle
            self.execute_update("""
                INSERT INTO scope3_targets (
                    company_id, category_id, target_type, baseline_year, target_year,
                    baseline_emissions, target_emissions, reduction_percentage,
                    target_description, created_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                target_data['company_id'],
                category_id,
                target_data['target_type'],
                target_data['baseline_year'],
                target_data['target_year'],
                target_data['baseline_emissions'],
                target_data['target_emissions'],
                target_data['reduction_percentage'],
                target_data['target_description'],
                datetime.now().isoformat()
            ))

            return True

        except Exception as e:
            logging.error(f"Scope 3 hedef kaydı ekleme hatası: {e}")
            return False

    def generate_scope3_report(self, company_id: int, report_name: str, period: str,
                              format_type: str, content_options: Dict) -> str:
        """Scope 3 raporu oluştur"""
        try:
            import os
            import pandas as pd
            from docx import Document

            # Rapor klasörü oluştur
            report_dir = os.path.join(os.path.dirname(self.db_path), '..', 'reports', 'scope3')
            os.makedirs(report_dir, exist_ok=True)

            # Dosya adı oluştur
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = "".join(c for c in report_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"{safe_name}_{timestamp}.{format_type.lower()}"
            filepath = os.path.join(report_dir, filename)

            report_data = {}

            # Emisyon verileri
            if content_options.get('include_emissions', True):
                # skip_tenant_filter=True because we handle company_id manually in the query
                emissions_rows = self.execute_query("""
                    SELECT e.*, c.category_number, c.category_name, c.description
                    FROM scope3_emissions e
                    JOIN scope3_categories c ON e.category_id = c.id
                    WHERE e.company_id = ? AND e.reporting_period = ?
                    ORDER BY c.category_number
                """, (company_id, period), skip_tenant_filter=True)

                emissions_data = []
                for row in emissions_rows:
                    emissions_data.append({
                        'Kategori': f"{row['category_number']} - {row['category_name']}",
                        'Aktivite Verisi': row['activity_data'],
                        'Aktivite Birimi': row['activity_unit'],
                        'Emisyon Faktörü': row['emission_factor'],
                        'Toplam Emisyon (tCO2e)': row['total_emissions'],
                        'Raporlama Dönemi': row['reporting_period'],
                        'Veri Kaynağı': row['data_source'],
                        'Notlar': row['notes']
                    })
                report_data['emissions'] = emissions_data

            # Hedef verileri
            if content_options.get('include_targets', True):
                # skip_tenant_filter=True because we handle company_id manually in the query
                targets_rows = self.execute_query("""
                    SELECT t.*, c.category_number, c.category_name
                    FROM scope3_targets t
                    JOIN scope3_categories c ON t.category_id = c.id
                    WHERE t.company_id = ?
                    ORDER BY c.category_number
                """, (company_id,), skip_tenant_filter=True)

                targets_data = []
                for row in targets_rows:
                    targets_data.append({
                        'Kategori': f"{row['category_number']} - {row['category_name']}",
                        'Hedef Tipi': row['target_type'],
                        'Baz Yılı': row['baseline_year'],
                        'Hedef Yılı': row['target_year'],
                        'Baz Emisyon (tCO2e)': row['baseline_emissions'],
                        'Hedef Emisyon (tCO2e)': row['target_emissions'],
                        'Azaltım (%)': row['reduction_percentage'],
                        'Açıklama': row['target_description']
                    })
                report_data['targets'] = targets_data

            # Özet istatistikler
            if content_options.get('include_summary', True):
                # skip_tenant_filter=True because we handle company_id manually in the query
                summary_rows = self.execute_query("""
                    SELECT 
                        COUNT(*) as toplam_kategori,
                        SUM(total_emissions) as toplam_emisyon,
                        AVG(total_emissions) as ortalama_emisyon,
                        MAX(total_emissions) as max_emisyon,
                        MIN(total_emissions) as min_emisyon
                    FROM scope3_emissions e
                    WHERE e.company_id = ? AND e.reporting_period = ?
                """, (company_id, period), skip_tenant_filter=True)

                if summary_rows:
                    summary_row = summary_rows[0]
                    report_data['summary'] = {
                        'Toplam Kategori': summary_row['toplam_kategori'],
                        'Toplam Emisyon (tCO2e)': summary_row['toplam_emisyon'] or 0,
                        'Ortalama Emisyon (tCO2e)': summary_row['ortalama_emisyon'] or 0,
                        'Maksimum Emisyon (tCO2e)': summary_row['max_emisyon'] or 0,
                        'Minimum Emisyon (tCO2e)': summary_row['min_emisyon'] or 0
                    }

            # Rapor oluştur
            if format_type.lower() == 'excel':
                with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                    if 'emissions' in report_data:
                        df_emissions = pd.DataFrame(report_data['emissions'])
                        df_emissions.to_excel(writer, sheet_name='Emisyon Verileri', index=False)

                    if 'targets' in report_data:
                        df_targets = pd.DataFrame(report_data['targets'])
                        df_targets.to_excel(writer, sheet_name='Hedefler', index=False)

                    if 'summary' in report_data:
                        df_summary = pd.DataFrame([report_data['summary']])
                        df_summary.to_excel(writer, sheet_name='Özet', index=False)

            elif format_type.lower() == 'csv':
                if 'emissions' in report_data:
                    df_emissions = pd.DataFrame(report_data['emissions'])
                    df_emissions.to_csv(filepath, index=False, encoding='utf-8-sig')

            elif format_type.lower() == 'docx':
                doc = Document()
                doc.add_heading(f'{report_name} - {period}', 0)

                # Özet
                if 'summary' in report_data:
                    doc.add_heading('Özet İstatistikler', level=1)
                    summary = report_data['summary']
                    doc.add_paragraph(f"Toplam Kategori: {summary['Toplam Kategori']}")
                    doc.add_paragraph(f"Toplam Emisyon: {summary['Toplam Emisyon (tCO2e)']:.2f} tCO2e")
                    doc.add_paragraph(f"Ortalama Emisyon: {summary['Ortalama Emisyon (tCO2e)']:.2f} tCO2e")

                # Emisyon verileri
                if 'emissions' in report_data:
                    doc.add_heading('Emisyon Verileri', level=1)
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'

                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Kategori'
                    hdr_cells[1].text = 'Aktivite Verisi'
                    hdr_cells[2].text = 'Toplam Emisyon'
                    hdr_cells[3].text = 'Veri Kaynağı'
                    hdr_cells[4].text = 'Dönem'
                    hdr_cells[5].text = 'Notlar'

                    for emission in report_data['emissions']:
                        row_cells = table.add_row().cells
                        row_cells[0].text = emission['Kategori']
                        row_cells[1].text = str(emission['Aktivite Verisi'] or '')
                        row_cells[2].text = str(emission['Toplam Emisyon (tCO2e)'] or '')
                        row_cells[3].text = emission['Veri Kaynağı'] or ''
                        row_cells[4].text = emission['Raporlama Dönemi'] or ''
                        row_cells[5].text = (emission['Notlar'] or '')

                doc.save(filepath)

            return filepath

        except Exception as e:
            logging.error(f"Scope 3 rapor oluşturma hatası: {e}")
            return None

    def get_target_data(self, company_id: int) -> List[Dict]:
        """Scope 3 hedef verilerini getir"""
        try:
            # skip_tenant_filter=True because we handle company_id manually in the query
            targets_rows = self.execute_query("""
                SELECT t.*, c.category_number, c.category_name
                FROM scope3_targets t
                JOIN scope3_categories c ON t.category_id = c.id
                WHERE t.company_id = ?
                ORDER BY c.category_number
            """, (company_id,), skip_tenant_filter=True)

            targets = []
            for row in targets_rows:
                targets.append({
                    'id': row['id'],
                    'company_id': row['company_id'],
                    'category_id': row['category_id'],
                    'target_type': row['target_type'],
                    'baseline_year': row['baseline_year'],
                    'target_year': row['target_year'],
                    'baseline_emissions': row['baseline_emissions'],
                    'target_emissions': row['target_emissions'],
                    'reduction_percentage': row['reduction_percentage'],
                    'target_description': row['target_description'],
                    'status': row['status'],
                    'created_at': row['created_at'],
                    'category_number': row['category_number'],
                    'category_name': row['category_name']
                })

            return targets

        except Exception as e:
            logging.error(f"Scope 3 hedef verileri getirme hatası: {e}")
            return []

    def load_scope3_categories(self) -> None:
        """Scope 3 kategorilerini yükle"""
        try:
            # Kategorilerin zaten var olup olmadığını kontrol et
            count_result = self.execute_query("SELECT COUNT(*) as count FROM scope3_categories")
            count = count_result[0]['count'] if count_result else 0

            if count == 0:
                # Scope 3 kategorilerini ekle
                categories = [
                    (1, "Satın Alınan Mallar ve Hizmetler", "Üretim için satın alınan tüm mallar ve hizmetler", "Indirect", 1, 0),
                    (2, "Sermaye Malları", "Üretim tesisleri, ekipman ve altyapı yatırımları", "Indirect", 1, 0),
                    (3, "Yakıt ve Enerji Faaliyetleri", "Üretim dışı yakıt kullanımı", "Indirect", 1, 0),
                    (4, "Upstream Taşıma ve Dağıtım", "Satın alınan malların taşınması", "Indirect", 1, 0),
                    (5, "Operasyonlarda Oluşan Atık", "Üretim sürecinde oluşan atıkların bertarafı", "Indirect", 1, 0),
                    (6, "İş Seyahatleri", "Çalışanların iş amaçlı seyahatleri", "Indirect", 1, 0),
                    (7, "Çalışan İşe Gidiş-Geliş", "Çalışanların ev-iş arası ulaşımı", "Indirect", 1, 0),
                    (8, "Kiralanan Varlıklar", "Kiralık ofis, depo ve diğer varlıklar", "Indirect", 1, 0),
                    (9, "Downstream Taşıma ve Dağıtım", "Ürünlerin müşterilere ulaştırılması", "Indirect", 0, 1),
                    (10, "Satılan Ürünlerin İşlenmesi", "Müşterilerin ürünleri kullanımı", "Indirect", 0, 1),
                    (11, "Satılan Ürünlerin Kullanımı", "Ürünlerin müşteri tarafından kullanımı", "Indirect", 0, 1),
                    (12, "Satılan Ürünlerin Bertarafı", "Ürünlerin son kullanım sonrası bertarafı", "Indirect", 0, 1),
                    (13, "Kiralanan Varlıklar (Downstream)", "Müşterilere kiralanan varlıklar", "Indirect", 0, 1),
                    (14, "Franchise", "Franchise operasyonları", "Indirect", 0, 1),
                    (15, "Yatırımlar", "Finansal yatırımlar", "Indirect", 1, 0)
                ]

                for cat in categories:
                    self.execute_update("""
                        INSERT INTO scope3_categories 
                        (category_number, category_name, description, scope_type, is_upstream, is_downstream)
                        VALUES (?, ?, ?, ?, ?, ?)
                    """, cat)

                logging.info(f"[OK] {len(categories)} Scope 3 kategorisi eklendi")

        except Exception as e:
            logging.error(f"[HATA] Scope 3 kategorileri yükleme hatası: {e}")

    def get_categories(self) -> List[Dict]:
        """Tüm Scope 3 kategorilerini getir"""
        try:
            rows = self.execute_query("""
                SELECT id, category_number, category_name, description, 
                       scope_type, is_upstream, is_downstream, is_active
                FROM scope3_categories 
                WHERE is_active = 1
                ORDER BY category_number
            """)

            categories = []
            for row in rows:
                categories.append({
                    'id': row['id'],
                    'category_number': row['category_number'],
                    'category_name': row['category_name'],
                    'description': row['description'],
                    'scope_type': row['scope_type'],
                    'is_upstream': bool(row['is_upstream']),
                    'is_downstream': bool(row['is_downstream']),
                    'is_active': bool(row['is_active'])
                })

            return categories

        except Exception as e:
            logging.error(f"Scope 3 kategorileri getirme hatası: {e}")
            return []
