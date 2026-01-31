#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CEO Mesajı Oluşturma Dialog'u
Yeni CEO mesajları oluşturmak için interaktif arayüz
"""

import logging
import tkinter as tk
from datetime import datetime
from tkinter import filedialog, messagebox, scrolledtext, ttk
from typing import Callable, Optional

import PyPDF2
from docx import Document

from modules.strategic.ceo_message_manager import CEOMessageManager
from utils.language_manager import LanguageManager


class CEOMessageDialog:
    """CEO mesajı oluşturma dialog'u"""

    def __init__(self, parent, company_id: int, user_id: int, on_save: Optional[Callable] = None) -> None:
        self.parent = parent
        self.company_id = company_id
        self.user_id = user_id
        self.on_save = on_save
        self.ceo_manager = CEOMessageManager()

        self.create_dialog()
        self.load_templates()

    def create_dialog(self) -> None:
        """Dialog penceresini oluştur"""
        self.dialog = tk.Toplevel(self.parent)
        self.dialog.title("Yeni CEO Mesajı Oluştur")
        self.dialog.geometry("1200x800")
        self.dialog.configure(bg='#f8f9fa')
        self.dialog.resizable(True, True)

        # Modal yap
        self.dialog.transient(self.parent)
        self.dialog.grab_set()

        # Pencereyi ortala
        self.center_window()

        self.build_ui()

    def center_window(self) -> None:
        """Pencereyi ekranın ortasına yerleştir"""
        self.dialog.update_idletasks()
        width = self.dialog.winfo_width()
        height = self.dialog.winfo_height()
        x = (self.dialog.winfo_screenwidth() // 2) - (width // 2)
        y = (self.dialog.winfo_screenheight() // 2) - (height // 2)
        self.dialog.geometry(f'{width}x{height}+{x}+{y}')

    def build_ui(self) -> None:
        """Kullanıcı arayüzünü oluştur"""
        # Ana frame
        main_frame = tk.Frame(self.dialog, bg='#f8f9fa')
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)

        # Başlık
        header_frame = tk.Frame(main_frame, bg='#2c3e50', height=60)
        header_frame.pack(fill='x', pady=(0, 20))
        header_frame.pack_propagate(False)

        title_label = tk.Label(header_frame, text=" Yeni CEO Mesajı Oluştur",
                              font=('Segoe UI', 16, 'bold'), fg='white', bg='#2c3e50')
        title_label.pack(pady=15)

        # Sol panel - Form alanları (yatay scrollable)
        left_panel = tk.Frame(main_frame, bg='white', relief='solid', bd=1, width=600)
        left_panel.pack(side='left', fill='both', expand=True, padx=(0, 10))
        left_panel.pack_propagate(False)

        self.build_form_fields(left_panel)

        # Sağ panel - Önizleme
        right_panel = tk.Frame(main_frame, bg='white', relief='solid', bd=1)
        right_panel.pack(side='right', fill='both', expand=True, padx=(10, 0))

        self.build_preview_panel(right_panel)

        # Footer - Butonlar
        footer_frame = tk.Frame(main_frame, bg='#f8f9fa', height=60)
        footer_frame.pack(fill='x', pady=(20, 0))
        footer_frame.pack_propagate(False)

        self.build_footer_buttons(footer_frame)

    def build_form_fields(self, parent) -> None:
        """Form alanlarını oluştur"""
        # Scrollable frame
        canvas = tk.Canvas(parent, bg='white')
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='white')

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Form alanları
        form_frame = tk.Frame(scrollable_frame, bg='white')
        form_frame.pack(fill='both', expand=True, padx=20, pady=20)

        # Mesaj Türü
        tk.Label(form_frame, text="Mesaj Türü:", font=('Segoe UI', 11, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', pady=(0, 5))

        self.message_type_var = tk.StringVar(value='annual')
        message_type_frame = tk.Frame(form_frame, bg='white')
        message_type_frame.pack(fill='x', pady=(0, 15))

        message_types = [
            ('annual', 'Yıllık Sürdürülebilirlik Mesajı'),
            ('quarterly', 'Çeyreklik Performans Mesajı'),
            ('sustainability', 'Özel Sürdürülebilirlik Mesajı'),
            ('emergency', 'Acil Durum Mesajı')
        ]

        for value, text in message_types:
            rb = tk.Radiobutton(message_type_frame, text=text, variable=self.message_type_var,
                              value=value, bg='white', fg='#2c3e50',
                              command=self.on_message_type_change)
            rb.pack(anchor='w', pady=2)

        # Şablon Seçimi
        tk.Label(form_frame, text="Şablon:", font=('Segoe UI', 11, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', pady=(15, 5))

        template_frame = tk.Frame(form_frame, bg='white')
        template_frame.pack(fill='x', pady=(0, 15))

        self.template_var = tk.StringVar()
        self.template_combo = ttk.Combobox(template_frame, textvariable=self.template_var,
                                         font=('Segoe UI', 10), state='readonly')
        self.template_combo.pack(side='left', fill='x', expand=True, padx=(0, 10))
        self.template_combo.bind('<<ComboboxSelected>>', self.on_template_select)

        # Şablon indirme butonu
        download_template_btn = tk.Button(template_frame, text=self.lm.tr('download_template', " Şablon İndir"),
                                         font=('Segoe UI', 9, 'bold'), fg='white', bg='#3498db',
                                         relief='flat', cursor='hand2', padx=15, pady=5,
                                         command=self.download_template)
        download_template_btn.pack(side='right')

        # Dosya yükleme bölümü
        file_upload_frame = tk.LabelFrame(form_frame, text=self.lm.tr('file_upload_title', " Dosya Yükleme (Word/PDF)"),
                                        bg='white', font=('Segoe UI', 11, 'bold'),
                                        fg='#2c3e50')
        file_upload_frame.pack(fill='x', pady=(15, 15))

        # Dosya seçimi
        file_select_frame = tk.Frame(file_upload_frame, bg='white')
        file_select_frame.pack(fill='x', padx=10, pady=10)

        self.file_path_var = tk.StringVar()
        file_path_entry = tk.Entry(file_select_frame, textvariable=self.file_path_var,
                                 font=('Segoe UI', 10), bg='white', fg='#2c3e50',
                                 relief='solid', bd=1, state='readonly')
        file_path_entry.pack(side='left', fill='x', expand=True, padx=(0, 10))

        browse_btn = tk.Button(file_select_frame, text=self.lm.tr('select_file', " Dosya Seç"),
                              font=('Segoe UI', 9, 'bold'), fg='white', bg='#27ae60',
                              relief='flat', cursor='hand2', padx=15, pady=5,
                              command=self.browse_file)
        browse_btn.pack(side='right')

        # Dosya yükleme butonu
        upload_btn = tk.Button(file_upload_frame, text=self.lm.tr('upload_and_import', " Dosyayı Yükle ve İçeriği Aktar"),
                              font=('Segoe UI', 10, 'bold'), fg='white', bg='#e67e22',
                              relief='flat', cursor='hand2', padx=20, pady=8,
                              command=self.upload_file)
        upload_btn.pack(pady=(0, 10))

        # Başlık
        tk.Label(form_frame, text="Mesaj Başlığı:", font=('Segoe UI', 11, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', pady=(0, 5))

        self.title_var = tk.StringVar()
        title_entry = tk.Entry(form_frame, textvariable=self.title_var, font=('Segoe UI', 10),
                             bg='white', fg='#2c3e50', relief='solid', bd=1)
        title_entry.pack(fill='x', pady=(0, 15))

        # Yıl ve Çeyrek
        year_quarter_frame = tk.Frame(form_frame, bg='white')
        year_quarter_frame.pack(fill='x', pady=(0, 15))

        # Yıl
        year_frame = tk.Frame(year_quarter_frame, bg='white')
        year_frame.pack(side='left', fill='x', expand=True, padx=(0, 10))

        tk.Label(year_frame, text="Yıl:", font=('Segoe UI', 11, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w')

        self.year_var = tk.StringVar(value=str(datetime.now().year))
        year_entry = tk.Entry(year_frame, textvariable=self.year_var, font=('Segoe UI', 10),
                            bg='white', fg='#2c3e50', relief='solid', bd=1)
        year_entry.pack(fill='x', pady=(5, 0))

        # Çeyrek (sadece çeyreklik mesajlar için)
        quarter_frame = tk.Frame(year_quarter_frame, bg='white')
        quarter_frame.pack(side='right', fill='x', expand=True, padx=(10, 0))

        tk.Label(quarter_frame, text="Çeyrek:", font=('Segoe UI', 11, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w')

        self.quarter_var = tk.StringVar(value='1')
        self.quarter_combo = ttk.Combobox(quarter_frame, textvariable=self.quarter_var,
                                        values=['1', '2', '3', '4'], font=('Segoe UI', 10),
                                        state='readonly', width=5)
        self.quarter_combo.pack(fill='x', pady=(5, 0))

        # İmza Bilgileri
        signature_frame = tk.LabelFrame(form_frame, text="İmza Bilgileri",
                                       bg='white', font=('Segoe UI', 11, 'bold'),
                                       fg='#2c3e50')
        signature_frame.pack(fill='x', pady=(15, 15))

        # CEO Adı
        tk.Label(signature_frame, text="CEO/Genel Müdür Adı:", font=('Segoe UI', 10, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', padx=10, pady=(10, 5))

        self.ceo_name_var = tk.StringVar()
        ceo_name_entry = tk.Entry(signature_frame, textvariable=self.ceo_name_var,
                                font=('Segoe UI', 10), bg='white', fg='#2c3e50',
                                relief='solid', bd=1)
        ceo_name_entry.pack(fill='x', padx=10, pady=(0, 10))

        # CEO Unvanı
        tk.Label(signature_frame, text="CEO/Genel Müdür Unvanı:", font=('Segoe UI', 10, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', padx=10, pady=(0, 5))

        self.ceo_title_var = tk.StringVar()
        ceo_title_entry = tk.Entry(signature_frame, textvariable=self.ceo_title_var,
                                 font=('Segoe UI', 10), bg='white', fg='#2c3e50',
                                 relief='solid', bd=1)
        ceo_title_entry.pack(fill='x', padx=10, pady=(0, 10))

        # Ana Başarılar
        achievements_frame = tk.LabelFrame(form_frame, text="Ana Başarılar",
                                         bg='white', font=('Segoe UI', 11, 'bold'),
                                         fg='#2c3e50')
        achievements_frame.pack(fill='x', pady=(0, 15))

        self.achievements_text = scrolledtext.ScrolledText(achievements_frame, height=4,
                                                         font=('Segoe UI', 10), bg='white',
                                                         fg='#2c3e50', wrap='word')
        self.achievements_text.pack(fill='x', padx=10, pady=10)

        # Zorluklar
        challenges_frame = tk.LabelFrame(form_frame, text="Karşılaşılan Zorluklar",
                                       bg='white', font=('Segoe UI', 11, 'bold'),
                                       fg='#2c3e50')
        challenges_frame.pack(fill='x', pady=(0, 15))

        self.challenges_text = scrolledtext.ScrolledText(challenges_frame, height=4,
                                                       font=('Segoe UI', 10), bg='white',
                                                       fg='#2c3e50', wrap='word')
        self.challenges_text.pack(fill='x', padx=10, pady=10)

        # Gelecek Taahhütleri
        commitments_frame = tk.LabelFrame(form_frame, text="Gelecek Taahhütleri",
                                        bg='white', font=('Segoe UI', 11, 'bold'),
                                        fg='#2c3e50')
        commitments_frame.pack(fill='x', pady=(0, 15))

        self.commitments_text = scrolledtext.ScrolledText(commitments_frame, height=4,
                                                        font=('Segoe UI', 10), bg='white',
                                                        fg='#2c3e50', wrap='word')
        self.commitments_text.pack(fill='x', padx=10, pady=10)

        # Mesaj İçeriği
        content_frame = tk.LabelFrame(form_frame, text="Mesaj İçeriği",
                                    bg='white', font=('Segoe UI', 11, 'bold'),
                                    fg='#2c3e50')
        content_frame.pack(fill='both', expand=True)

        self.content_text = scrolledtext.ScrolledText(content_frame, height=8,
                                                    font=('Segoe UI', 10), bg='white',
                                                    fg='#2c3e50', wrap='word')
        self.content_text.pack(fill='both', expand=True, padx=10, pady=10)

        # İlk mesaj türü değişikliğini tetikle
        self.on_message_type_change()

    def build_preview_panel(self, parent) -> None:
        """Önizleme panelini oluştur"""
        # Önizleme başlığı ve butonu
        preview_header = tk.Frame(parent, bg='white')
        preview_header.pack(fill='x', padx=20, pady=(20, 10))

        tk.Label(preview_header, text=" Önizleme", font=('Segoe UI', 14, 'bold'),
                fg='#2c3e50', bg='white').pack(side='left')

        update_preview_btn = tk.Button(preview_header, text=" Önizlemeyi Güncelle",
                                     font=('Segoe UI', 10, 'bold'), fg='white', bg='#3498db',
                                     relief='flat', cursor='hand2',
                                     command=self.update_preview)
        update_preview_btn.pack(side='right')

        # Önizleme alanı
        preview_frame = tk.Frame(parent, bg='#f8f9fa', relief='solid', bd=1)
        preview_frame.pack(fill='both', expand=True, padx=20, pady=(0, 20))

        self.preview_text = scrolledtext.ScrolledText(preview_frame, height=15,
                                                    font=('Segoe UI', 10), bg='white',
                                                    fg='#2c3e50', wrap='word', state='disabled')
        self.preview_text.pack(fill='both', expand=True, padx=10, pady=10)

    def build_footer_buttons(self, parent) -> None:
        """Footer butonlarını oluştur"""
        button_frame = tk.Frame(parent, bg='#f8f9fa')
        button_frame.pack(expand=True)

        # İptal butonu
        cancel_btn = tk.Button(button_frame, text=" İptal",
                             font=('Segoe UI', 11, 'bold'), fg='white', bg='#95a5a6',
                             relief='flat', cursor='hand2', padx=30, pady=10,
                             command=self.cancel)
        cancel_btn.pack(side='left', padx=(0, 10))

        # Kaydet butonu
        save_btn = tk.Button(button_frame, text=" Kaydet",
                           font=('Segoe UI', 11, 'bold'), fg='white', bg='#27ae60',
                           relief='flat', cursor='hand2', padx=30, pady=10,
                           command=self.save_message)
        save_btn.pack(side='right', padx=(10, 0))

    def load_templates(self) -> None:
        """Mesaj şablonlarını yükle"""
        self.ceo_manager.create_default_templates()
        templates = self.ceo_manager.get_templates()

        template_options = []
        for template in templates:
            template_options.append(f"{template['name']} ({template['message_type']})")

        self.template_combo['values'] = template_options

    def on_message_type_change(self) -> None:
        """Mesaj türü değiştiğinde"""
        message_type = self.message_type_var.get()

        # Çeyrek seçimini göster/gizle
        if message_type == 'quarterly':
            self.quarter_combo.pack(fill='x', pady=(5, 0))
        else:
            self.quarter_combo.pack_forget()

        # Şablonları güncelle
        templates = self.ceo_manager.get_templates(message_type)
        template_options = [f"{t['name']} ({t['message_type']})" for t in templates]
        self.template_combo['values'] = template_options

        if template_options:
            self.template_combo.set(template_options[0])
            self.on_template_select()

    def on_template_select(self, event=None) -> None:
        """Şablon seçildiğinde"""
        selected = self.template_combo.get()
        if not selected:
            return

        # Şablon adını çıkar
        template_name = selected.split(' (')[0]

        # Şablonu bul ve içeriği oluştur
        templates = self.ceo_manager.get_templates()
        for template in templates:
            if template['name'] == template_name:
                self.generate_content_from_template(template)
                break

    def generate_content_from_template(self, template) -> None:
        """Şablondan içerik oluştur"""
        try:
            # Değişkenleri hazırla
            variables = {
                'year': self.year_var.get(),
                'quarter': self.quarter_var.get() if self.message_type_var.get() == 'quarterly' else '',
                'company_name': 'Şirket Adı',  # Bu bilgiyi veritabanından alabilirsiniz
                'ceo_name': self.ceo_name_var.get() or 'CEO Adı',
                'ceo_title': self.ceo_title_var.get() or 'Genel Müdür'
            }

            # Şablondan mesaj oluştur
            content = self.ceo_manager.generate_message_from_template(template['id'], variables)

            # İçeriği text alanına yerleştir
            self.content_text.delete(1.0, tk.END)
            self.content_text.insert(1.0, content)

            # Başlığı otomatik oluştur
            if not self.title_var.get():
                year = variables['year']
                if self.message_type_var.get() == 'quarterly':
                    quarter = variables['quarter']
                    self.title_var.set(f"{year} Yılı {quarter}. Çeyrek Performans Mesajı")
                elif self.message_type_var.get() == 'annual':
                    self.title_var.set(f"{year} Yıllık Sürdürülebilirlik Mesajı")
                elif self.message_type_var.get() == 'emergency':
                    self.title_var.set(f"{year} Acil Durum Mesajı")
                else:
                    self.title_var.set(f"{year} Sürdürülebilirlik Mesajı")

        except Exception as e:
            messagebox.showerror("Hata", f"Şablon işleme hatası: {e}")

    def update_preview(self) -> None:
        """Önizlemeyi güncelle"""
        try:
            # Mesaj içeriğini al
            content = self.content_text.get(1.0, tk.END).strip()

            # Başarıları al
            achievements = self.achievements_text.get(1.0, tk.END).strip()
            achievements_list = [line.strip() for line in achievements.split('\n') if line.strip()]

            # Zorlukları al
            challenges = self.challenges_text.get(1.0, tk.END).strip()
            challenges_list = [line.strip() for line in challenges.split('\n') if line.strip()]

            # Taahhütleri al
            commitments = self.commitments_text.get(1.0, tk.END).strip()
            commitments_list = [line.strip() for line in commitments.split('\n') if line.strip()]

            # Önizleme metnini oluştur
            preview_parts = []

            if content:
                preview_parts.append("=== MESAJ İÇERİĞİ ===")
                preview_parts.append(content)
                preview_parts.append("")

            if achievements_list:
                preview_parts.append("=== ANA BAŞARILAR ===")
                for i, achievement in enumerate(achievements_list, 1):
                    preview_parts.append(f"{i}. {achievement}")
                preview_parts.append("")

            if challenges_list:
                preview_parts.append("=== KARŞILAŞILAN ZORLUKLAR ===")
                for i, challenge in enumerate(challenges_list, 1):
                    preview_parts.append(f"{i}. {challenge}")
                preview_parts.append("")

            if commitments_list:
                preview_parts.append("=== GELECEK TAAHHÜTLERİ ===")
                for i, commitment in enumerate(commitments_list, 1):
                    preview_parts.append(f"{i}. {commitment}")
                preview_parts.append("")

            # İmza bilgileri
            if self.ceo_name_var.get() or self.ceo_title_var.get():
                preview_parts.append("=== İMZA ===")
                preview_parts.append("Saygılarımla,")
                preview_parts.append("")
                if self.ceo_name_var.get():
                    preview_parts.append(self.ceo_name_var.get())
                if self.ceo_title_var.get():
                    preview_parts.append(self.ceo_title_var.get())
                preview_parts.append("")
                preview_parts.append(f"Tarih: {datetime.now().strftime('%d.%m.%Y')}")

            # Önizlemeyi güncelle
            preview_text = '\n'.join(preview_parts)
            self.preview_text.config(state='normal')
            self.preview_text.delete(1.0, tk.END)
            self.preview_text.insert(1.0, preview_text)
            self.preview_text.config(state='disabled')

        except Exception as e:
            messagebox.showerror("Hata", f"Önizleme güncelleme hatası: {e}")

    def save_message(self) -> None:
        """Mesajı kaydet"""
        try:
            # Validasyon
            if not self.title_var.get().strip():
                messagebox.showerror("Hata", "Mesaj başlığı boş olamaz!")
                return

            if not self.content_text.get(1.0, tk.END).strip():
                messagebox.showerror("Hata", "Mesaj içeriği boş olamaz!")
                return

            # Başarıları al
            achievements = self.achievements_text.get(1.0, tk.END).strip()
            achievements_list = [line.strip() for line in achievements.split('\n') if line.strip()]

            # Zorlukları al
            challenges = self.challenges_text.get(1.0, tk.END).strip()
            challenges_list = [line.strip() for line in challenges.split('\n') if line.strip()]

            # Taahhütleri al
            commitments = self.commitments_text.get(1.0, tk.END).strip()
            commitments_list = [line.strip() for line in commitments.split('\n') if line.strip()]

            # Çeyrek bilgisini hazırla
            quarter = int(self.quarter_var.get()) if self.message_type_var.get() == 'quarterly' else None

            # Mesajı oluştur
            message_id = self.ceo_manager.create_message(
                company_id=self.company_id,
                title=self.title_var.get().strip(),
                message_type=self.message_type_var.get(),
                year=int(self.year_var.get()),
                quarter=quarter,
                content=self.content_text.get(1.0, tk.END).strip(),
                key_achievements=achievements_list,
                challenges=challenges_list,
                future_commitments=commitments_list,
                signature_name=self.ceo_name_var.get().strip(),
                signature_title=self.ceo_title_var.get().strip(),
                created_by=self.user_id
            )

            messagebox.showinfo("Başarılı", f"CEO mesajı başarıyla oluşturuldu!\nMesaj ID: {message_id}")

            # Callback'i çağır
            if self.on_save:
                self.on_save()

            # Dialog'u kapat
            self.dialog.destroy()

        except Exception as e:
            messagebox.showerror("Hata", f"Mesaj kaydetme hatası: {e}")

    def browse_file(self) -> None:
        """Dosya seç"""
        file_types = [
            (self.lm.tr('word_files', "Word Dosyaları"), "*.docx"),
            (self.lm.tr('pdf_files', "PDF Dosyaları"), "*.pdf"),
            (self.lm.tr('all_files', "Tüm Dosyalar"), "*.*")
        ]
        
        filename = filedialog.askopenfilename(
            title=self.lm.tr('select_file', "Dosya Seç"),
            filetypes=file_types
        )
        
        if filename:
            self.file_path_var.set(filename)

    def upload_file(self) -> None:
        """Dosyayı yükle ve içeriği aktar"""
        file_path = self.file_path_var.get()
        if not file_path:
            messagebox.showerror("Hata", "Lütfen önce bir dosya seçin!")
            return

        try:
            if file_path.lower().endswith('.docx'):
                self.upload_word_file(file_path)
            elif file_path.lower().endswith('.pdf'):
                self.upload_pdf_file(file_path)
            else:
                messagebox.showerror("Hata", "Desteklenen dosya formatları: .docx, .pdf")

        except Exception as e:
            messagebox.showerror("Hata", f"Dosya yükleme hatası: {e}")

    def upload_word_file(self, file_path) -> None:
        """Word dosyasını yükle ve içeriği aktar"""
        try:
            # Word dosyasını oku
            doc = Document(file_path)

            # Tüm paragrafları birleştir
            content_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            full_content = '\n\n'.join(content_parts)

            if not full_content:
                messagebox.showwarning("Uyarı", "Dosyada okunabilir metin bulunamadı!")
                return

            # İçeriği mesaj alanına aktar
            self.content_text.delete(1.0, tk.END)
            self.content_text.insert(1.0, full_content)

            # Başlık varsa otomatik doldur
            if not self.title_var.get() and len(content_parts) > 0:
                # İlk paragrafı başlık olarak kullan
                first_line = content_parts[0][:50]  # İlk 50 karakter
                if len(first_line) < len(content_parts[0]):
                    first_line += "..."
                self.title_var.set(first_line)

            messagebox.showinfo("Başarılı", "Word dosyası başarıyla yüklendi!\nİçerik mesaj alanına aktarıldı.")

            # Önizlemeyi güncelle
            self.update_preview()

        except Exception as e:
            messagebox.showerror("Hata", f"Word dosyası okuma hatası: {e}")

    def upload_pdf_file(self, file_path) -> None:
        """PDF dosyasını yükle ve içeriği aktar"""
        try:
            # PDF dosyasını oku
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Tüm sayfaları birleştir
                content_parts = []
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text.strip())

                full_content = '\n\n'.join(content_parts)

            if not full_content:
                messagebox.showwarning("Uyarı", "PDF dosyasında okunabilir metin bulunamadı!")
                return

            # İçeriği mesaj alanına aktar
            self.content_text.delete(1.0, tk.END)
            self.content_text.insert(1.0, full_content)

            # Başlık varsa otomatik doldur
            if not self.title_var.get() and len(content_parts) > 0:
                # İlk sayfanın ilk satırını başlık olarak kullan
                first_line = content_parts[0][:50]  # İlk 50 karakter
                if len(first_line) < len(content_parts[0]):
                    first_line += "..."
                self.title_var.set(first_line)

            messagebox.showinfo("Başarılı", f"PDF dosyası başarıyla yüklendi!\nİçerik mesaj alanına aktarıldı.\nSayfa sayısı: {len(pdf_reader.pages)}")

            # Önizlemeyi güncelle
            self.update_preview()

        except Exception as e:
            messagebox.showerror("Hata", f"PDF dosyası okuma hatası: {e}")

    def download_template(self) -> None:
        """Şablon indir"""
        selected_template = self.template_combo.get()
        if not selected_template:
            messagebox.showerror("Hata", "Lütfen önce bir şablon seçin!")
            return

        try:
            # Şablon adını çıkar
            template_name = selected_template.split(' (')[0]

            # Şablonu bul
            templates = self.ceo_manager.get_templates()
            selected_template_data = None

            for template in templates:
                if template['name'] == template_name:
                    selected_template_data = template
                    break

            if not selected_template_data:
                messagebox.showerror("Hata", "Şablon bulunamadı!")
                return

            # Dosya kaydetme dialog'u
            message_type = selected_template_data['message_type']
            file_name = f"CEO_Mesaj_Sablonu_{message_type}_{datetime.now().strftime('%Y%m%d')}.docx"

            file_path = filedialog.asksaveasfilename(
                title=self.lm.tr('save_template', "Şablonu Kaydet"),
                defaultextension=".docx",
                filetypes=[(self.lm.tr('word_files', "Word Dosyaları"), "*.docx")],
                initialfile=file_name
            )

            if file_path:
                self.create_template_document(selected_template_data, file_path)
                messagebox.showinfo("Başarılı", f"Şablon başarıyla kaydedildi!\n{file_path}")

        except Exception as e:
            messagebox.showerror("Hata", f"Şablon indirme hatası: {e}")

    def create_template_document(self, template_data, file_path) -> None:
        """Şablon Word belgesi oluştur"""
        try:
            # Yeni Word belgesi oluştur
            doc = Document()

            # Başlık
            title = doc.add_heading(f"{template_data['name']} Şablonu", 0)
            title.alignment = 1  # Ortala

            # Açıklama
            doc.add_paragraph("Bu şablon CEO mesajı oluşturmak için kullanılabilir.")
            doc.add_paragraph("")

            # Şablon içeriği
            template_content = template_data['template_content']

            # Değişkenler tablosu
            doc.add_heading("Kullanılacak Değişkenler:", level=1)

            variables_table = doc.add_table(rows=1, cols=2)
            variables_table.style = 'Table Grid'

            # Başlık satırı
            hdr_cells = variables_table.rows[0].cells
            hdr_cells[0].text = 'Değişken'
            hdr_cells[1].text = 'Açıklama'

            # Değişkenler
            variable_descriptions = {
                'year': 'Yıl (örn: 2024)',
                'quarter': 'Çeyrek (1-4)',
                'company_name': 'Şirket adı',
                'ceo_name': 'CEO/Genel Müdür adı',
                'ceo_title': 'CEO/Genel Müdür unvanı',
                'situation_description': 'Durum açıklaması (acil durum mesajları için)'
            }

            for variable in template_data['variables']:
                row_cells = variables_table.add_row().cells
                row_cells[0].text = f"{{{variable}}}"
                row_cells[1].text = variable_descriptions.get(variable, 'Değişken')

            doc.add_paragraph("")

            # Şablon içeriği
            doc.add_heading("Şablon İçeriği:", level=1)

            # Her bölümü ekle
            if 'greeting' in template_content:
                doc.add_paragraph(template_content['greeting'])
                doc.add_paragraph("")

            if 'intro' in template_content:
                doc.add_paragraph(template_content['intro'])
                doc.add_paragraph("")

            if 'achievements_section' in template_content:
                doc.add_heading(template_content['achievements_section'], level=2)
                doc.add_paragraph("[Ana başarılarınızı buraya yazın]")
                doc.add_paragraph("")

            if 'challenges_section' in template_content:
                doc.add_heading(template_content['challenges_section'], level=2)
                doc.add_paragraph("[Karşılaştığınız zorlukları buraya yazın]")
                doc.add_paragraph("")

            if 'commitments_section' in template_content:
                doc.add_heading(template_content['commitments_section'], level=2)
                doc.add_paragraph("[Gelecek taahhütlerinizi buraya yazın]")
                doc.add_paragraph("")

            if 'closing' in template_content:
                doc.add_paragraph(template_content['closing'])
                doc.add_paragraph("")

            if 'signature' in template_content:
                doc.add_paragraph(template_content['signature'])
                doc.add_paragraph("[CEO/Genel Müdür Adı]")
                doc.add_paragraph("[CEO/Genel Müdür Unvanı]")
                doc.add_paragraph(f"[Tarih: {datetime.now().strftime('%d.%m.%Y')}]")

            # Belgeyi kaydet
            doc.save(file_path)

        except Exception as e:
            raise Exception(f"Word belgesi oluşturma hatası: {e}")

    def cancel(self) -> None:
        """İptal et"""
        if messagebox.askyesno("İptal", "Mesaj oluşturmayı iptal etmek istediğinizden emin misiniz?"):
            self.dialog.destroy()


def show_ceo_message_dialog(parent, company_id: int, user_id: int, on_save: Optional[Callable] = None) -> None:
    """CEO mesajı dialog'unu göster"""
    CEOMessageDialog(parent, company_id, user_id, on_save)


if __name__ == "__main__":
    # Test
    root = tk.Tk()
    root.title("CEO Mesajı Dialog Test")
    root.geometry("1200x800")

    def on_save_callback() -> None:
        logging.info("Mesaj kaydedildi!")

    show_ceo_message_dialog(root, company_id=1, user_id=1, on_save=on_save_callback)

    root.mainloop()
