#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GRI DOCX/PDF Reports - Sprint 4
GRI Content Index ve özet raporları DOCX/PDF formatında
"""

import logging
import os
import sqlite3
from datetime import datetime
from typing import Dict, List

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                TableStyle)

from utils.language_manager import LanguageManager
from config.database import DB_PATH


def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return heading


class GRIDocxPDFReports:
    """GRI DOCX/PDF raporları sınıfı"""

    def __init__(self, db_path: str = DB_PATH) -> None:
        # db_path göreli ise proje köküne göre mutlak hale getir
        if not os.path.isabs(db_path):
            base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
            db_path = os.path.join(base_dir, db_path)
        self.db_path = db_path
        self.lm = LanguageManager()

    def get_connection(self) -> None:
        """Veritabanı bağlantısı"""
        return sqlite3.connect(self.db_path)

    def generate_gri_content_index_data(self, company_id: int = 1) -> Dict:
        """GRI Content Index verilerini oluştur"""
        conn = self.get_connection()
        cursor = conn.cursor()

        try:
            logging.info("GRI Content Index verileri oluşturuluyor...")

            # Tüm GRI göstergelerini al
            cursor.execute("""
                SELECT 
                    gs.code as standard_code,
                    gs.title as standard_title,
                    gs.category,
                    gi.code as disclosure_code,
                    gi.title as disclosure_title,
                    gi.description,
                    gi.unit,
                    gi.methodology,
                    gi.reporting_requirement,
                    gr.response_value,
                    gr.numerical_value,
                    gr.reporting_status,
                    gr.period,
                    gr.evidence_url,
                    gr.notes
                FROM gri_standards gs
                JOIN gri_indicators gi ON gs.id = gi.standard_id
                LEFT JOIN gri_responses gr ON gi.id = gr.indicator_id AND gr.company_id = ?
                ORDER BY gs.category, gs.code, gi.code
            """, (company_id,))

            indicators = cursor.fetchall()

            # Kategorilere göre organize et
            content_index = {
                'universal': [],
                'economic': [],
                'environmental': [],
                'social': [],
                'sector': [],
                'summary': {
                    'total_standards': 0,
                    'total_indicators': len(indicators),
                    'reported_indicators': 0,
                    'pending_indicators': 0,
                    'categories': {}
                }
            }

            processed_standards = set()

            for indicator in indicators:
                standard_code = indicator[0]
                category = indicator[2]
                has_response = indicator[9] is not None  # gr.response_value

                indicator_info = {
                    'standard_code': standard_code,
                    'standard_title': indicator[1],
                    'category': category,
                    'disclosure_code': indicator[3],
                    'disclosure_title': indicator[4],
                    'description': indicator[5],
                    'unit': indicator[6],
                    'methodology': indicator[7],
                    'reporting_requirement': indicator[8],
                    'response_value': indicator[9],
                    'numerical_value': indicator[10],
                    'reporting_status': indicator[11],
                    'period': indicator[12],
                    'evidence_url': indicator[13],
                    'notes': indicator[14],
                    'has_response': has_response,
                    'page_reference': f"{self.lm.tr('page', 'Sayfa')} {len(content_index.get(category.lower(), [])) + 1}"
                }

                # Kategoriye göre ekle
                category_key = category.lower().replace('-', '_').replace(' ', '_')
                if category_key in content_index:
                    content_index[category_key].append(indicator_info)
                else:
                    content_index['sector'].append(indicator_info)

                # Özet güncelle
                if standard_code not in processed_standards:
                    content_index['summary']['total_standards'] += 1
                    processed_standards.add(standard_code)

                if has_response:
                    content_index['summary']['reported_indicators'] += 1
                else:
                    content_index['summary']['pending_indicators'] += 1

                # Kategori bazında sayım
                if category not in content_index['summary']['categories']:
                    content_index['summary']['categories'][category] = {
                        'total': 0,
                        'reported': 0,
                        'pending': 0
                    }

                content_index['summary']['categories'][category]['total'] += 1
                if has_response:
                    content_index['summary']['categories'][category]['reported'] += 1
                else:
                    content_index['summary']['categories'][category]['pending'] += 1

            return content_index

        except Exception as e:
            logging.error(f"GRI Content Index veri oluşturma hatası: {e}")
            return {}
        finally:
            conn.close()

    def generate_docx_content_index(self, output_path: str, company_id: int = 1) -> bool:
        """GRI Content Index DOCX raporu oluştur"""
        try:
            logging.info(f"GRI Content Index DOCX raporu oluşturuluyor: {output_path}")

            # Veri oluştur
            content_data = self.generate_gri_content_index_data(company_id)
            if not content_data:
                logging.info("Content Index verisi oluşturulamadı!")
                return False

            # DOCX dokümanı oluştur
            doc = Document()

            # Başlık
            title = _add_turkish_heading(doc, self.lm.tr('gri_content_index', 'GRI Content Index'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Tarih
            date_para = _add_turkish_paragraph(doc, f"{self.lm.tr('report_date', 'Rapor Tarihi')}: {datetime.now().strftime('%d.%m.%Y')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            _add_turkish_paragraph(doc, '')  # Boş satır

            # Özet bölümü
            self.add_summary_section_docx(doc, content_data['summary'])

            # Kategori bölümleri
            categories = [
                ('universal', self.lm.tr('universal_standards', 'Universal Standards')),
                ('economic', self.lm.tr('economic_standards', 'Economic Standards')),
                ('environmental', self.lm.tr('environmental_standards', 'Environmental Standards')),
                ('social', self.lm.tr('social_standards', 'Social Standards')),
                ('sector', self.lm.tr('sector_standards', 'Sector-Specific Standards'))
            ]

            for category_key, category_title in categories:
                if content_data.get(category_key):
                    self.add_category_section_docx(doc, category_title, content_data[category_key])

            # Kaydet
            doc.save(output_path)
            logging.info(f"GRI Content Index DOCX başarıyla oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"GRI Content Index DOCX oluşturma hatası: {e}")
            return False

    def generate_pdf_content_index(self, output_path: str, company_id: int = 1) -> bool:
        """GRI Content Index PDF raporu oluştur"""
        try:
            logging.info(f"GRI Content Index PDF raporu oluşturuluyor: {output_path}")

            # Veri oluştur
            content_data = self.generate_gri_content_index_data(company_id)
            if not content_data:
                logging.info("Content Index verisi oluşturulamadı!")
                return False

            # PDF dokümanı oluştur
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            # Özel stiller
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=TA_CENTER
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=12
            )

            # Başlık
            story.append(Paragraph(self.lm.tr('gri_content_index', "GRI Content Index"), title_style))
            story.append(Paragraph(f"{self.lm.tr('report_date', 'Rapor Tarihi')}: {datetime.now().strftime('%d.%m.%Y')}", styles['Normal']))
            story.append(Spacer(1, 20))

            # Özet bölümü
            self.add_summary_section_pdf(story, content_data['summary'], styles)

            # Kategori bölümleri
            categories = [
                ('universal', self.lm.tr('universal_standards', 'Universal Standards')),
                ('economic', self.lm.tr('economic_standards', 'Economic Standards')),
                ('environmental', self.lm.tr('environmental_standards', 'Environmental Standards')),
                ('social', self.lm.tr('social_standards', 'Social Standards')),
                ('sector', self.lm.tr('sector_standards', 'Sector-Specific Standards'))
            ]

            for category_key, category_title in categories:
                if content_data.get(category_key):
                    self.add_category_section_pdf(story, category_title, content_data[category_key], heading_style, styles)

            # PDF oluştur
            doc.build(story)
            logging.info(f"GRI Content Index PDF başarıyla oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"GRI Content Index PDF oluşturma hatası: {e}")
            return False

    def generate_kpi_snapshot_pdf(self, snapshot: Dict, output_path: str) -> bool:
        try:
            logging.info(f"KPI Snapshot PDF raporu oluşturuluyor: {output_path}")

            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            title_style = ParagraphStyle(
                "KpiTitle",
                parent=styles["Heading1"],
                fontSize=16,
                spaceAfter=24,
                alignment=TA_CENTER,
            )

            story.append(Paragraph("KPI Snapshot", title_style))
            story.append(Paragraph(f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y')}", styles["Normal"]))
            story.append(Spacer(1, 16))

            company = snapshot.get("company") or {}
            period = snapshot.get("period") or {}
            kpis = snapshot.get("kpis") or []

            period_label = ""
            if isinstance(period, dict):
                period_label = period.get("label") or period.get("year") or ""
            else:
                period_label = str(period)

            meta_rows: List[List] = [
                ["Şirket", company.get("name") or ""],
                ["Sektör", company.get("sector") or ""],
                ["Ülke", company.get("country") or ""],
                ["Raporlama Dönemi", period_label],
                ["Toplam KPI", str(len(kpis))],
            ]

            meta_table = Table(meta_rows, colWidths=[140, 360])
            meta_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ]
                )
            )
            story.append(meta_table)
            story.append(Spacer(1, 20))

            if kpis:
                story.append(Paragraph("KPI Özeti", styles["Heading2"]))
                table_headers = ["Kod", "Ad", "Modül", "Değer", "Birim", "Yıl"]
                table_data: List[List] = [table_headers]
                for kpi in kpis[:200]:
                    table_data.append(
                        [
                            kpi.get("code") or "",
                            kpi.get("name") or "",
                            kpi.get("module") or "",
                            "" if kpi.get("value") is None else str(kpi.get("value")),
                            kpi.get("unit") or "",
                            "" if kpi.get("year") is None else str(kpi.get("year")),
                        ]
                    )
                kpi_table = Table(table_data, repeatRows=1)
                kpi_table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                            ("FONTSIZE", (0, 0), (-1, 0), 9),
                            ("FONTSIZE", (0, 1), (-1, -1), 7),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ]
                    )
                )
                story.append(kpi_table)

            doc.build(story)
            logging.info(f"KPI Snapshot PDF başarıyla oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"KPI Snapshot PDF oluşturma hatası: {e}")
            return False

    def add_summary_section_docx(self, doc, summary) -> None:
        """DOCX özet bölümü ekle"""
        _add_turkish_heading(doc, self.lm.tr('summary', 'Özet'), level=1)

        # Özet tablosu
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Başlık satırı
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = self.lm.tr('metric', 'Metrik')
        hdr_cells[1].text = self.lm.tr('value', 'Değer')

        # Özet verileri
        summary_data = [
            (self.lm.tr('total_standards', 'Toplam Standard'), str(summary['total_standards'])),
            (self.lm.tr('total_indicators', 'Toplam Gösterge'), str(summary['total_indicators'])),
            (self.lm.tr('reported_indicators', 'Raporlanan Gösterge'), str(summary['reported_indicators'])),
            (self.lm.tr('pending_indicators', 'Bekleyen Gösterge'), str(summary['pending_indicators'])),
            (self.lm.tr('reporting_ratio', 'Raporlama Oranı (%)'), f"{(summary['reported_indicators'] / summary['total_indicators'] * 100):.1f}" if summary['total_indicators'] > 0 else "0.0")
        ]

        for metric, value in summary_data:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value

        _add_turkish_paragraph(doc, '')  # Boş satır

    def add_category_section_docx(self, doc, category_title, indicators) -> None:
        """DOCX kategori bölümü ekle"""
        _add_turkish_heading(doc, category_title, level=1)

        # Kategori tablosu
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Başlık satırı
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = self.lm.tr('standard_code', 'Standard Code')
        hdr_cells[1].text = self.lm.tr('disclosure_code', 'Disclosure Code')
        hdr_cells[2].text = self.lm.tr('title', 'Title')
        hdr_cells[3].text = self.lm.tr('unit', 'Unit')
        hdr_cells[4].text = self.lm.tr('status', 'Status')
        hdr_cells[5].text = self.lm.tr('page_reference', 'Page Reference')

        # Gösterge verileri
        for indicator in indicators:
            row_cells = table.add_row().cells
            row_cells[0].text = indicator['standard_code']
            row_cells[1].text = indicator['disclosure_code']
            row_cells[2].text = indicator['disclosure_title']
            row_cells[3].text = indicator['unit'] or ''
            row_cells[4].text = self.lm.tr('reported', 'Reported') if indicator['has_response'] else self.lm.tr('pending', 'Pending')
            row_cells[5].text = indicator['page_reference']

        _add_turkish_paragraph(doc, '')  # Boş satır

    def add_summary_section_pdf(self, story, summary, styles) -> None:
        """PDF özet bölümü ekle"""
        story.append(Paragraph("Özet", styles['Heading1']))

        # Özet tablosu
        summary_data = [
            ['Metrik', 'Değer'],
            ['Toplam Standard', str(summary['total_standards'])],
            ['Toplam Gösterge', str(summary['total_indicators'])],
            ['Raporlanan Gösterge', str(summary['reported_indicators'])],
            ['Bekleyen Gösterge', str(summary['pending_indicators'])],
            ['Raporlama Oranı (%)', f"{(summary['reported_indicators'] / summary['total_indicators'] * 100):.1f}" if summary['total_indicators'] > 0 else "0.0"]
        ]

        summary_table = Table(summary_data)
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(summary_table)
        story.append(Spacer(1, 20))

    def add_category_section_pdf(self, story, category_title, indicators, heading_style, styles) -> None:
        """PDF kategori bölümü ekle"""
        story.append(Paragraph(category_title, heading_style))

        # Kategori tablosu
        table_data = [['Standard Code', 'Disclosure Code', 'Title', 'Unit', 'Status', 'Page Ref']]

        for indicator in indicators:
            table_data.append([
                indicator['standard_code'],
                indicator['disclosure_code'],
                indicator['disclosure_title'],
                indicator['unit'] or '',
                'Reported' if indicator['has_response'] else 'Pending',
                indicator['page_reference']
            ])

        category_table = Table(table_data)
        category_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
        ]))

        story.append(category_table)
        story.append(Spacer(1, 20))

    def generate_comprehensive_report_docx(self, output_path: str, company_id: int = 1) -> bool:
        """Kapsamlı GRI raporu DOCX oluştur"""
        try:
            logging.info(f"Kapsamlı GRI DOCX raporu oluşturuluyor: {output_path}")

            # Veri oluştur
            content_data = self.generate_gri_content_index_data(company_id)
            if not content_data:
                logging.info("Content Index verisi oluşturulamadı!")
                return False

            # DOCX dokümanı oluştur
            doc = Document()

            # Kapak sayfası
            self.add_cover_page_docx(doc)

            # İçindekiler
            self.add_table_of_contents_docx(doc, content_data)

            # Executive Summary
            self.add_executive_summary_docx(doc, content_data)

            # GRI Content Index
            self.add_gri_content_index_docx(doc, content_data)

            # Kategori detayları
            categories = [
                ('universal', 'Universal Standards'),
                ('economic', 'Economic Standards'),
                ('environmental', 'Environmental Standards'),
                ('social', 'Social Standards'),
                ('sector', 'Sector-Specific Standards')
            ]

            for category_key, category_title in categories:
                if content_data.get(category_key):
                    self.add_detailed_category_section_docx(doc, category_title, content_data[category_key])

            # Kaydet
            doc.save(output_path)
            logging.info(f"Kapsamlı GRI DOCX raporu başarıyla oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"Kapsamlı GRI DOCX oluşturma hatası: {e}")
            return False

    def add_cover_page_docx(self, doc) -> None:
        """DOCX kapak sayfası ekle"""
        title = _add_turkish_heading(doc, self.lm.tr('gri_sustainability_report', 'GRI Sustainability Report'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _add_turkish_paragraph(doc, '')  # Boş satır

        subtitle = _add_turkish_paragraph(doc, self.lm.tr('global_reporting_initiative', 'Global Reporting Initiative'))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _add_turkish_paragraph(doc, '')  # Boş satır

        company_info = _add_turkish_paragraph(doc, f"{self.lm.tr('company', 'Company')}: [Company Name]")
        company_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        report_date = _add_turkish_paragraph(doc, f"{self.lm.tr('report_date', 'Report Date')}: {datetime.now().strftime('%B %Y')}")
        report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def add_table_of_contents_docx(self, doc, content_data) -> None:
        """DOCX içindekiler ekle"""
        _add_turkish_heading(doc, self.lm.tr('table_of_contents', 'Table of Contents'), level=1)

        toc_items = [
            self.lm.tr('executive_summary', 'Executive Summary'),
            self.lm.tr('gri_content_index', 'GRI Content Index'),
            self.lm.tr('universal_standards', 'Universal Standards'),
            self.lm.tr('economic_standards', 'Economic Standards'),
            self.lm.tr('environmental_standards', 'Environmental Standards'),
            self.lm.tr('social_standards', 'Social Standards'),
            self.lm.tr('sector_standards', 'Sector-Specific Standards')
        ]

        for item in toc_items:
            _add_turkish_paragraph(doc, f'• {item}', style='List Bullet')

        doc.add_page_break()

    def add_executive_summary_docx(self, doc, content_data) -> None:
        """DOCX executive summary ekle"""
        _add_turkish_heading(doc, self.lm.tr('executive_summary', 'Executive Summary'), level=1)

        summary = content_data['summary']

        summary_text = f"""
        {self.lm.tr('gri_report_intro', 'This GRI sustainability report covers')} {summary['total_standards']} {self.lm.tr('standards', 'standards')} {self.lm.tr('and', 'and')} {summary['total_indicators']} {self.lm.tr('indicators', 'indicators')} {self.lm.tr('across_all_gri_categories', 'across all GRI categories.')}
        
        {self.lm.tr('of_the', 'Of the')} {summary['total_indicators']} {self.lm.tr('indicators', 'indicators')}, {summary['reported_indicators']} {self.lm.tr('have_been_reported', 'have been reported')} ({summary['reported_indicators'] / summary['total_indicators'] * 100:.1f}%), 
        {self.lm.tr('while', 'while')} {summary['pending_indicators']} {self.lm.tr('are_pending_reporting', 'are pending reporting.')}
        
        {self.lm.tr('gri_commitment_text', 'The report demonstrates our commitment to transparency and sustainability reporting in accordance with GRI standards.')}
        """

        _add_turkish_paragraph(doc, summary_text)
        doc.add_page_break()

    def add_gri_content_index_docx(self, doc, content_data) -> None:
        """DOCX GRI Content Index ekle"""
        _add_turkish_heading(doc, self.lm.tr('gri_content_index', 'GRI Content Index'), level=1)

        # Özet tablosu
        self.add_summary_section_docx(doc, content_data['summary'])

        # Kategori bölümleri
        categories = [
            ('universal', self.lm.tr('universal_standards', 'Universal Standards')),
            ('economic', self.lm.tr('economic_standards', 'Economic Standards')),
            ('environmental', self.lm.tr('environmental_standards', 'Environmental Standards')),
            ('social', self.lm.tr('social_standards', 'Social Standards')),
            ('sector', self.lm.tr('sector_standards', 'Sector-Specific Standards'))
        ]

        for category_key, category_title in categories:
            if content_data.get(category_key):
                self.add_category_section_docx(doc, category_title, content_data[category_key])

    def add_detailed_category_section_docx(self, doc, category_title, indicators) -> None:
        """DOCX detaylı kategori bölümü ekle"""
        _add_turkish_heading(doc, category_title, level=1)

        for indicator in indicators:
            _add_turkish_heading(doc, f"{indicator['standard_code']} - {indicator['disclosure_code']}", level=2)

            # Gösterge detayları
            details = _add_turkish_paragraph(doc, '')
            details.add_run("Title: ").bold = True
            details.add_run(indicator['disclosure_title'])

            if indicator['description']:
                details.add_run("\nDescription: ").bold = True
                details.add_run(indicator['description'])

            if indicator['unit']:
                details.add_run("\nUnit: ").bold = True
                details.add_run(indicator['unit'])

            if indicator['methodology']:
                details.add_run("\nMethodology: ").bold = True
                details.add_run(indicator['methodology'])

            details.add_run("\nStatus: ").bold = True
            details.add_run('Reported' if indicator['has_response'] else 'Pending')

            if indicator['has_response']:
                details.add_run("\nResponse Value: ").bold = True
                details.add_run(str(indicator['response_value']))

                if indicator['numerical_value']:
                    details.add_run("\nNumerical Value: ").bold = True
                    details.add_run(str(indicator['numerical_value']))

            _add_turkish_paragraph(doc, '')  # Boş satır

def export_gri_content_index_docx(output_path: str, company_id: int = 1) -> None:
    """GRI Content Index DOCX export fonksiyonu"""
    docx_reports = GRIDocxPDFReports()
    return docx_reports.generate_docx_content_index(output_path, company_id)

def export_gri_content_index_pdf(output_path: str, company_id: int = 1) -> None:
    """GRI Content Index PDF export fonksiyonu"""
    docx_reports = GRIDocxPDFReports()
    return docx_reports.generate_pdf_content_index(output_path, company_id)

def export_gri_comprehensive_report_docx(output_path: str, company_id: int = 1) -> None:
    """Kapsamlı GRI DOCX raporu export fonksiyonu"""
    docx_reports = GRIDocxPDFReports()
    return docx_reports.generate_comprehensive_report_docx(output_path, company_id)

if __name__ == "__main__":
    # Test exports
    docx_output = "gri/gri_content_index_test.docx"
    pdf_output = "gri/gri_content_index_test.pdf"
    comprehensive_docx = "gri/gri_comprehensive_report_test.docx"

    if export_gri_content_index_docx(docx_output):
        logging.info(f"GRI Content Index DOCX başarıyla oluşturuldu: {docx_output}")

    if export_gri_content_index_pdf(pdf_output):
        logging.info(f"GRI Content Index PDF başarıyla oluşturuldu: {pdf_output}")

    if export_gri_comprehensive_report_docx(comprehensive_docx):
        logging.info(f"Kapsamlı GRI DOCX raporu başarıyla oluşturuldu: {comprehensive_docx}")
