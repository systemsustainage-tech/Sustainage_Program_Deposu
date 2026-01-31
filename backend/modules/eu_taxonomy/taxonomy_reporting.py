#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
EU Taxonomy Reporting
Taxonomy raporları oluşturma
"""

import logging
import os
from datetime import datetime
from config.database import DB_PATH


def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    try:
        from docx.shared import Pt
    except Exception:
        Pt = None
    para = doc.add_paragraph(text if text is not None else "", style=style)
    if Pt is not None:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text if text is not None else "", level=level)
    for run in heading.runs:
        run.font.name = font_name
    return heading


class TaxonomyReporting:
    """EU Taxonomy raporlama"""

    def __init__(self, manager) -> None:
        self.manager = manager

    def generate_docx_report(self, company_id: int, period: str, output_dir: str = 'reports') -> str:
        """DOCX raporu oluştur"""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Rapor verilerini al
            report_data = self.manager.generate_taxonomy_report(company_id, period)

            # Çıktı dizini
            os.makedirs(output_dir, exist_ok=True)

            # Dosya adı
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"eu_taxonomy_report_{company_id}_{period}_{timestamp}.docx"
            filepath = os.path.join(output_dir, filename)

            # Doküman oluştur
            doc = Document()

            # Başlık
            title = _add_turkish_heading(doc, 'EU Taxonomy Uyum Raporu', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Genel bilgiler
            _add_turkish_heading(doc, 'Genel Bilgiler', level=1)

            info_table = doc.add_table(rows=1, cols=2)
            info_table.style = 'Light Grid Accent 1'

            for label, value in [
                ('Şirket ID', str(company_id)),
                ('Raporlama Dönemi', period),
                ('Rapor Tarihi', datetime.now().strftime('%d.%m.%Y'))
            ]:
                r = info_table.add_row().cells
                r[0].text = label
                r[1].text = value

            # Özet
            _add_turkish_heading(doc, 'Yönetici Özeti', level=1)

            summary = report_data.get('summary', {})

            p = _add_turkish_paragraph(doc, "")
            p.add_run("Toplam Faaliyet Sayısı: ").bold = True
            p.add_run(f"{summary.get('total_activities', 0)}\n")

            p.add_run("Uygun Faaliyet Sayısı: ").bold = True
            p.add_run(f"{summary.get('eligible_activities', 0)}\n")

            p.add_run("Uyumlu Faaliyet Sayısı: ").bold = True
            p.add_run(f"{summary.get('aligned_activities', 0)}\n")

            p.add_run("Uygunluk Oranı: ").bold = True
            p.add_run(f"{summary.get('eligibility_rate', 0):.1f}%\n")

            p.add_run("Uyumluluk Oranı: ").bold = True
            p.add_run(f"{summary.get('alignment_rate', 0):.1f}%")

            # KPI'lar
            _add_turkish_heading(doc, 'Taxonomy KPI\'ları', level=1)

            kpis = report_data.get('kpis', {})

            # Gelir KPI'ları
            _add_turkish_heading(doc, 'Gelir (Revenue)', level=2)
            revenue = kpis.get('revenue', {})

            revenue_table = doc.add_table(rows=1, cols=2)
            revenue_table.style = 'Light Grid Accent 1'

            for label, value in [
                ('Metrik', 'Değer'),
                ('Toplam Gelir Payı', f"{revenue.get('total', 0):.2f}%"),
                ('Uygun Gelir Oranı', f"{revenue.get('eligible_percentage', 0):.2f}%"),
                ('Uyumlu Gelir Oranı', f"{revenue.get('aligned_percentage', 0):.2f}%")
            ]:
                r = revenue_table.add_row().cells
                r[0].text = label
                r[1].text = value

            # CapEx KPI'ları
            _add_turkish_heading(doc, 'Sermaye Harcamaları (CapEx)', level=2)
            capex = kpis.get('capex', {})

            capex_table = doc.add_table(rows=1, cols=2)
            capex_table.style = 'Light Grid Accent 1'

            for label, value in [
                ('Metrik', 'Değer'),
                ('Toplam CapEx Payı', f"{capex.get('total', 0):.2f}%"),
                ('Uygun CapEx Oranı', f"{capex.get('eligible_percentage', 0):.2f}%"),
                ('Uyumlu CapEx Oranı', f"{capex.get('aligned_percentage', 0):.2f}%")
            ]:
                r = capex_table.add_row().cells
                r[0].text = label
                r[1].text = value

            # OpEx KPI'ları
            _add_turkish_heading(doc, 'Operasyonel Harcamalar (OpEx)', level=2)
            opex = kpis.get('opex', {})

            opex_table = doc.add_table(rows=1, cols=2)
            opex_table.style = 'Light Grid Accent 1'

            for label, value in [
                ('Metrik', 'Değer'),
                ('Toplam OpEx Payı', f"{opex.get('total', 0):.2f}%"),
                ('Uygun OpEx Oranı', f"{opex.get('eligible_percentage', 0):.2f}%"),
                ('Uyumlu OpEx Oranı', f"{opex.get('aligned_percentage', 0):.2f}%")
            ]:
                r = opex_table.add_row().cells
                r[0].text = label
                r[1].text = value

            # Faaliyetler
            _add_turkish_heading(doc, 'Taxonomy Faaliyetleri', level=1)

            activities = report_data.get('activities', [])

            if activities:
                # Faaliyet tablosu
                activity_table = doc.add_table(rows=len(activities) + 1, cols=6)
                activity_table.style = 'Light Grid Accent 1'

                # Başlıklar
                headers = ['Kod', 'Faaliyet', 'Sektör', 'Gelir %', 'Uygun', 'Uyumlu']
                for i, header in enumerate(headers):
                    activity_table.rows[0].cells[i].text = header

                # Veriler
                for i, activity in enumerate(activities, 1):
                    activity_table.rows[i].cells[0].text = activity.get('activity_code', '')
                    activity_table.rows[i].cells[1].text = activity.get('activity_name_tr', '')
                    activity_table.rows[i].cells[2].text = activity.get('sector', '')
                    activity_table.rows[i].cells[3].text = f"{activity.get('revenue_share', 0):.1f}"
                    activity_table.rows[i].cells[4].text = 'Evet' if activity.get('is_eligible') else 'Hayır'
                    activity_table.rows[i].cells[5].text = 'Evet' if activity.get('is_aligned') else 'Hayır'

            # Sonuç ve öneriler
            _add_turkish_heading(doc, 'Sonuç ve Öneriler', level=1)

            _add_turkish_paragraph(doc,
                "Bu rapor, şirketin EU Taxonomy uyumluluğunu değerlendirmektedir. "
                "Raporlanan KPI'lar, şirketin sürdürülebilir faaliyetlerinin finansal "
                "performansa katkısını göstermektedir."
            )

            # Belgeyi kaydet
            doc.save(filepath)

            logging.info(f"[OK] EU Taxonomy raporu oluşturuldu: {filepath}")
            return filepath

        except Exception as e:
            logging.error(f"[HATA] EU Taxonomy raporu oluşturma hatası: {e}")
            return None

    def generate_excel_report(self, company_id: int, period: str, output_dir: str = 'reports') -> str:
        """Excel raporu oluştur"""
        try:
            import pandas as pd

            # Rapor verilerini al
            report_data = self.manager.generate_taxonomy_report(company_id, period)

            # Çıktı dizini
            os.makedirs(output_dir, exist_ok=True)

            # Dosya adı
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"eu_taxonomy_report_{company_id}_{period}_{timestamp}.xlsx"
            filepath = os.path.join(output_dir, filename)

            # Excel writer
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                # Özet sayfası
                summary = report_data.get('summary', {})
                summary_df = pd.DataFrame([
                    ['Toplam Faaliyet', summary.get('total_activities', 0)],
                    ['Uygun Faaliyet', summary.get('eligible_activities', 0)],
                    ['Uyumlu Faaliyet', summary.get('aligned_activities', 0)],
                    ['Uygunluk Oranı (%)', f"{summary.get('eligibility_rate', 0):.1f}"],
                    ['Uyumluluk Oranı (%)', f"{summary.get('alignment_rate', 0):.1f}"]
                ], columns=['Metrik', 'Değer'])
                summary_df.to_excel(writer, sheet_name='Özet', index=False)

                # KPI sayfası
                kpis = report_data.get('kpis', {})
                kpi_data = []

                for kpi_type in ['revenue', 'capex', 'opex']:
                    kpi = kpis.get(kpi_type, {})
                    kpi_data.append([
                        kpi_type.upper(),
                        kpi.get('total', 0),
                        kpi.get('eligible', 0),
                        kpi.get('aligned', 0),
                        kpi.get('eligible_percentage', 0),
                        kpi.get('aligned_percentage', 0)
                    ])

                kpi_df = pd.DataFrame(kpi_data, columns=[
                    'KPI', 'Toplam', 'Uygun', 'Uyumlu', 'Uygun %', 'Uyumlu %'
                ])
                kpi_df.to_excel(writer, sheet_name='KPI', index=False)

                # Faaliyetler sayfası
                activities = report_data.get('activities', [])
                if activities:
                    activities_df = pd.DataFrame(activities)
                    activities_df.to_excel(writer, sheet_name='Faaliyetler', index=False)

            logging.info(f"[OK] EU Taxonomy Excel raporu oluşturuldu: {filepath}")
            return filepath

        except Exception as e:
            logging.error(f"[HATA] EU Taxonomy Excel raporu oluşturma hatası: {e}")
            return None

# Test fonksiyonu
if __name__ == '__main__':
    import sys
    sys.path.append(os.path.join(os.path.dirname(__file__), '..', '..'))

    from modules.eu_taxonomy.taxonomy_manager import EUTaxonomyManager

    db_path = DB_PATH
    manager = EUTaxonomyManager(db_path)
    reporting = TaxonomyReporting(manager)

    # Test rapor oluştur
    filepath = reporting.generate_docx_report(company_id=1, period='2024')
    logging.info(f"Rapor oluşturuldu: {filepath}")
