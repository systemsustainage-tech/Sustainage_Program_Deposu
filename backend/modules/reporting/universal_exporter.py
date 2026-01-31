import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Universal Export Sistemi
Tüm modüller için PNG, PDF, Excel export desteği
"""

import os
from datetime import datetime
from typing import Any, Dict, List, Optional

import pandas as pd


class UniversalExporter:
    """Evrensel export sistemi - tüm formatları destekler"""

    def __init__(self, export_dir: str = "exports") -> None:
        """
        Args:
            export_dir: Export klasörü
        """
        self.export_dir = export_dir
        os.makedirs(export_dir, exist_ok=True)

    # ============================================
    # EXCEL EXPORT
    # ============================================

    def export_to_excel(self, data: List[Dict], filename: str,
                       sheet_name: str = 'Veri',
                       metadata: Optional[Dict] = None,
                       multi_sheet: Optional[Dict[str, List[Dict]]] = None) -> str:
        """
        Excel'e export - tek veya çok sayfalı
        
        Args:
            data: Veri listesi (single sheet için)
            filename: Dosya adı
            sheet_name: Sayfa adı
            metadata: Metadata bilgileri
            multi_sheet: Çok sayfalı export {'Sayfa1': data1, 'Sayfa2': data2}
        
        Returns:
            str: Dosya yolu
        """
        try:
            filepath = os.path.join(self.export_dir, filename)
            if not filepath.endswith('.xlsx'):
                filepath += '.xlsx'

            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                # Çok sayfalı export
                if multi_sheet:
                    for sheet, sheet_data in multi_sheet.items():
                        if sheet_data:
                            df = pd.DataFrame(sheet_data)
                            df.to_excel(writer, sheet_name=sheet, index=False)
                        else:
                            # Boş sayfa
                            pd.DataFrame({'Bilgi': ['Veri bulunamadı']}).to_excel(
                                writer, sheet_name=sheet, index=False)
                else:
                    # Tek sayfalı export
                    if data:
                        df = pd.DataFrame(data)
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                    else:
                        pd.DataFrame({'Bilgi': ['Veri bulunamadı']}).to_excel(
                            writer, sheet_name=sheet_name, index=False)

                # Metadata sayfası ekle
                if metadata:
                    meta_df = pd.DataFrame([
                        {'Özellik': key, 'Değer': value}
                        for key, value in metadata.items()
                    ])
                    meta_df.to_excel(writer, sheet_name='Metadata', index=False)

            return filepath

        except Exception as e:
            raise Exception(f"Excel export hatası: {e}")

    def export_table_to_excel(self, table_data: List[tuple],
                             columns: List[str], filename: str) -> str:
        """
        Treeview veya tablo verilerini Excel'e export
        
        Args:
            table_data: Tablo satırları (tuple listesi)
            columns: Sütun isimleri
            filename: Dosya adı
        """
        data_dicts = []
        for row in table_data:
            row_dict = {}
            for i, col in enumerate(columns):
                row_dict[col] = row[i] if i < len(row) else ''
            data_dicts.append(row_dict)

        return self.export_to_excel(data_dicts, filename)

    # ============================================
    # PDF EXPORT
    # ============================================

    def export_to_pdf(self, content: str, filename: str,
                     title: Optional[str] = None, subtitle: Optional[str] = None,
                     company_info: Optional[Dict] = None,
                     add_toc: bool = True,
                     add_page_numbers: bool = True) -> str:
        """
        PDF'e export - profesyonel formatting
        
        Args:
            content: İçerik metni (Markdown veya plain text)
            filename: Dosya adı
            title: Rapor başlığı
            subtitle: Alt başlık
            company_info: Şirket bilgileri
            add_toc: İçindekiler ekle
            add_page_numbers: Sayfa numaraları ekle
        
        Returns:
            str: Dosya yolu
        """
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

            from .font_utils import register_turkish_fonts_reportlab

            filepath = os.path.join(self.export_dir, filename)
            if not filepath.endswith('.pdf'):
                filepath += '.pdf'

            # PDF oluştur
            doc = SimpleDocTemplate(filepath, pagesize=A4,
                                   rightMargin=2*cm, leftMargin=2*cm,
                                   topMargin=2*cm, bottomMargin=2*cm)
            # Türkçe fontları kaydet
            register_turkish_fonts_reportlab()

            story = []
            styles = getSampleStyleSheet()

            # Özel stiller
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=25,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='NotoSans'
            )

            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading2'],
                fontSize=15,
                textColor=colors.HexColor('#7f8c8d'),
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='NotoSans'
            )

            # Başlık
            if title:
                story.append(Paragraph(title, title_style))

            if subtitle:
                story.append(Paragraph(subtitle, subtitle_style))

            # Şirket bilgileri
            if company_info:
                story.append(Spacer(1, 0.5*cm))
                for key, value in company_info.items():
                    text = f"<b>{key}:</b> {value}"
                    story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 1*cm))

            # İçerik
            for line in content.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                    story.append(Spacer(1, 0.2*cm))

            # PDF oluştur
            doc.build(story)

            return filepath

        except ImportError:
            raise Exception("reportlab kütüphanesi yüklü değil. 'pip install reportlab' ile yükleyin.")
        except Exception as e:
            raise Exception(f"PDF export hatası: {e}")

    def export_chart_to_pdf(self, chart_figure, filename: str,
                           title: Optional[str] = None) -> str:
        """
        Matplotlib grafiğini PDF'e export
        
        Args:
            chart_figure: Matplotlib Figure nesnesi
            filename: Dosya adı
            title: Başlık
        """
        try:
            filepath = os.path.join(self.export_dir, filename)
            if not filepath.endswith('.pdf'):
                filepath += '.pdf'

            chart_figure.savefig(filepath, format='pdf', bbox_inches='tight', dpi=300)

            return filepath

        except Exception as e:
            raise Exception(f"Chart PDF export hatası: {e}")

    # ============================================
    # PNG/IMAGE EXPORT
    # ============================================

    def export_to_png(self, figure, filename: str, dpi: int = 300) -> str:
        """
        Matplotlib grafiğini PNG'ye export
        
        Args:
            figure: Matplotlib Figure
            filename: Dosya adı
            dpi: Çözünürlük
        """
        try:
            filepath = os.path.join(self.export_dir, filename)
            if not filepath.endswith('.png'):
                filepath += '.png'

            figure.savefig(filepath, format='png', bbox_inches='tight', dpi=dpi)

            return filepath

        except Exception as e:
            raise Exception(f"PNG export hatası: {e}")

    def export_widget_to_image(self, widget, filename: str) -> str:
        """
        Tkinter widget'ını görüntü olarak kaydet
        
        Args:
            widget: Tkinter widget
            filename: Dosya adı
        """
        try:
            from PIL import ImageGrab

            filepath = os.path.join(self.export_dir, filename)
            if not filepath.endswith('.png'):
                filepath += '.png'

            # Widget koordinatlarını al
            x = widget.winfo_rootx()
            y = widget.winfo_rooty()
            x1 = x + widget.winfo_width()
            y1 = y + widget.winfo_height()

            # Ekran görüntüsü al
            ImageGrab.grab(bbox=(x, y, x1, y1)).save(filepath)

            return filepath

        except Exception as e:
            raise Exception(f"Widget image export hatası: {e}")

    # ============================================
    # CSV EXPORT
    # ============================================

    def export_to_csv(self, data: List[Dict], filename: str,
                     encoding: str = 'utf-8-sig') -> str:
        """
        CSV'ye export (Excel uyumlu)
        
        Args:
            data: Veri listesi
            filename: Dosya adı
            encoding: Encoding (utf-8-sig Excel için)
        """
        try:
            filepath = os.path.join(self.export_dir, filename)
            if not filepath.endswith('.csv'):
                filepath += '.csv'

            if data:
                df = pd.DataFrame(data)
                df.to_csv(filepath, index=False, encoding=encoding)
            else:
                # Boş CSV
                pd.DataFrame({'Bilgi': ['Veri bulunamadı']}).to_csv(
                    filepath, index=False, encoding=encoding)

            return filepath

        except Exception as e:
            raise Exception(f"CSV export hatası: {e}")

    # ============================================
    # JSON EXPORT
    # ============================================

    def export_to_json(self, data: Any, filename: str,
                      indent: int = 2) -> str:
        """
        JSON'a export
        
        Args:
            data: Export edilecek veri
            filename: Dosya adı
            indent: JSON indent
        """
        try:
            import json

            filepath = os.path.join(self.export_dir, filename)
            if not filepath.endswith('.json'):
                filepath += '.json'

            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=indent)

            return filepath

        except Exception as e:
            raise Exception(f"JSON export hatası: {e}")

    def export_markdown_to_word(self, md_path: str, filename: Optional[str] = None,
                                title: Optional[str] = None) -> str:
        try:
            from docx import Document
            from docx.shared import Pt

            if not os.path.exists(md_path):
                raise Exception("Markdown dosyası bulunamadı")

            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()

            doc = Document()

            def add_paragraph(text: str, style: Optional[str] = None) -> None:
                p = doc.add_paragraph(text, style=style)
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

            def add_heading(text: str, level: int = 1) -> None:
                h = doc.add_heading(text, level=level)
                for run in h.runs:
                    run.font.name = 'Calibri'

            if title:
                add_heading(title, 0)

            in_code = False
            code_lang = ''
            for raw_line in content.split('\n'):
                line = raw_line.rstrip('\n')
                if line.strip().startswith('```'):
                    if not in_code:
                        in_code = True
                        code_lang = line.strip().replace('```', '').strip()
                        if code_lang:
                            add_paragraph(f"Kod Bloğu: {code_lang}")
                        else:
                            add_paragraph("Kod Bloğu")
                    else:
                        in_code = False
                        add_paragraph("")
                    continue

                if in_code:
                    p = doc.add_paragraph()
                    r = p.add_run(line)
                    r.font.name = 'Consolas'
                    r.font.size = Pt(10)
                    continue

                if line.startswith('# '):
                    add_heading(line[2:], 1)
                elif line.startswith('## '):
                    add_heading(line[3:], 2)
                elif line.startswith('### '):
                    add_heading(line[4:], 3)
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(line[2:])
                    p.style = 'List Bullet'
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                else:
                    add_paragraph(line)

            out_name = filename or os.path.splitext(os.path.basename(md_path))[0] + '.docx'
            out_path = os.path.join(self.export_dir, out_name)
            if not out_path.endswith('.docx'):
                out_path += '.docx'
            doc.save(out_path)
            return out_path
        except ImportError:
            raise Exception("python-docx kütüphanesi yüklü değil. 'pip install python-docx' ile yükleyin.")
        except Exception as e:
            raise Exception(f"Word export hatası: {e}")

    # ============================================
    # BATCH EXPORT
    # ============================================

    def batch_export(self, exports: List[Dict]) -> List[str]:
        """
        Toplu export - birden fazla format aynı anda
        
        Args:
            exports: Export tanımları listesi
                [{'type': 'excel', 'data': [...], 'filename': 'rapor.xlsx'}, ...]
        
        Returns:
            List[str]: Oluşturulan dosya yolları
        """
        results = []

        for export_config in exports:
            try:
                export_type = export_config.get('type', 'excel')
                data = export_config.get('data') or []
                filename = export_config.get('filename', f'export_{datetime.now().strftime("%Y%m%d_%H%M%S")}')

                if export_type == 'excel':
                    filepath = self.export_to_excel(data, filename)
                elif export_type == 'csv':
                    filepath = self.export_to_csv(data, filename)
                elif export_type == 'json':
                    filepath = self.export_to_json(data, filename)
                elif export_type == 'pdf':
                    content = export_config.get('content', '')
                    filepath = self.export_to_pdf(content, filename)
                else:
                    continue

                results.append(filepath)

            except Exception as e:
                logging.error(f"Batch export hatası ({export_type}): {e}")
                continue

        return results

    # ============================================
    # HELPER METHODS
    # ============================================

    def get_safe_filename(self, filename: str) -> str:
        """Güvenli dosya adı oluştur"""
        import re

        # Türkçe karakterleri değiştir
        replacements = {
            'ş': 's', 'Ş': 'S', 'ğ': 'g', 'Ğ': 'G',
            'ü': 'u', 'Ü': 'U', 'ı': 'i', 'İ': 'I',
            'ö': 'o', 'Ö': 'O', 'ç': 'c', 'Ç': 'C'
        }
        for old, new in replacements.items():
            filename = filename.replace(old, new)

        # Özel karakterleri temizle
        filename = re.sub(r'[^\w\s-]', '', filename)
        filename = re.sub(r'[-\s]+', '_', filename)

        return filename

    def add_timestamp_to_filename(self, filename: str) -> str:
        """Dosya adına timestamp ekle"""
        name, ext = os.path.splitext(filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return f"{name}_{timestamp}{ext}"

    def get_export_metadata(self, company_name: Optional[str] = None,
                           report_type: Optional[str] = None,
                           period: Optional[str] = None) -> Dict:
        """Export metadata oluştur"""
        return {
            'Şirket': company_name or 'N/A',
            'Rapor Türü': report_type or 'Genel',
            'Dönem': period or datetime.now().strftime('%Y'),
            'Oluşturma Tarihi': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'Sistem': 'SUSTAINAGE SDG v5.9.0'
        }


class AdvancedPDFExporter:
    """Gelişmiş PDF export - interaktif özelliklerle"""

    def __init__(self, export_dir: str = "exports") -> None:
        self.export_dir = export_dir
        os.makedirs(export_dir, exist_ok=True)

    def create_interactive_pdf(self, content: str, filename: str,
                             title: str, subtitle: Optional[str] = None,
                             company_info: Optional[Dict] = None,
                             add_toc: bool = True,
                             add_links: bool = True,
                             add_bookmarks: bool = True,
                             logo_path: Optional[str] = None,
                             brand_colors: Optional[Dict] = None,
                             add_page_numbers: bool = True) -> str:
        """
        İnteraktif PDF oluştur
        
        Features:
        - Otomatik içindekiler (TOC)
        - Hyperlink'ler
        - Bookmark'lar
        - Logo ve marka renkleri
        - Sayfa numaraları
        
        Args:
            content: İçerik (Markdown destekli)
            filename: Dosya adı
            title: Başlık
            subtitle: Alt başlık
            company_info: Şirket bilgileri
            add_toc: İçindekiler ekle
            add_links: Bağlantıları aktif et
            add_bookmarks: PDF bookmark ekle
            logo_path: Logo dosya yolu
            brand_colors: Marka renkleri {'primary': '#hexcode', 'secondary': '#hexcode'}
        """
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

            filepath = os.path.join(self.export_dir, filename)
            if not filepath.endswith('.pdf'):
                filepath += '.pdf'

            # PDF dökümanı
            doc = SimpleDocTemplate(filepath, pagesize=A4,
                                   rightMargin=2*cm, leftMargin=2*cm,
                                   topMargin=3*cm, bottomMargin=2.5*cm)

            story = []
            styles = getSampleStyleSheet()

            # Marka renkleri
            primary_color = colors.HexColor(brand_colors.get('primary', '#2c3e50')) if brand_colors else colors.HexColor('#2c3e50')
            secondary_color = colors.HexColor(brand_colors.get('secondary', '#3498db')) if brand_colors else colors.HexColor('#3498db')

            # Logo
            if logo_path and os.path.exists(logo_path):
                try:
                    logo = Image(logo_path, width=4*cm, height=2*cm)
                    story.append(logo)
                    story.append(Spacer(1, 0.5*cm))
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")

            # Başlık
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=28,
                textColor=primary_color,
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='DejaVuSans-Bold'  # Türkçe karakter desteği
            )
            story.append(Paragraph(title, title_style))

            if subtitle:
                subtitle_style = ParagraphStyle(
                    'CustomSubtitle',
                    parent=styles['Heading2'],
                    fontSize=16,
                    textColor=secondary_color,
                    spaceAfter=30,
                    alignment=TA_CENTER
                )
                story.append(Paragraph(subtitle, subtitle_style))

            story.append(Spacer(1, 1*cm))

            # Şirket bilgileri kutusu
            if company_info:
                company_data = [[key, value] for key, value in company_info.items()]
                company_table = Table(company_data, colWidths=[5*cm, 10*cm])
                company_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f9fa')),
                    ('TEXTCOLOR', (0, 0), (0, -1), primary_color),
                    ('TEXTCOLOR', (1, 0), (1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('PADDING', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
                ]))
                story.append(company_table)
                story.append(Spacer(1, 1*cm))

            # İçerik - Markdown parse
            lines = content.split('\n')
            for line in lines:
                line = line.strip()

                if not line:
                    story.append(Spacer(1, 0.3*cm))
                    continue

                # Başlıklar
                if line.startswith('# '):
                    text = line[2:]
                    story.append(PageBreak())
                    story.append(Paragraph(text, styles['Heading1']))
                elif line.startswith('## '):
                    text = line[3:]
                    story.append(Paragraph(text, styles['Heading2']))
                elif line.startswith('### '):
                    text = line[4:]
                    story.append(Paragraph(text, styles['Heading3']))
                # Liste
                elif line.startswith('- ') or line.startswith('* '):
                    text = '• ' + line[2:]
                    story.append(Paragraph(text, styles['Normal']))
                # Normal metin
                else:
                    story.append(Paragraph(line, styles['Normal']))

            # Sayfa numaraları ekle
            if add_page_numbers:
                def add_page_number(canvas, doc) -> None:
                    canvas.saveState()
                    canvas.setFont('Helvetica', 9)
                    page_num = canvas.getPageNumber()
                    text = f"Sayfa {page_num}"
                    canvas.drawRightString(A4[0] - 2*cm, 1.5*cm, text)
                    canvas.restoreState()

                doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
            else:
                doc.build(story)

            return filepath

        except Exception as e:
            raise Exception(f"İnteraktif PDF oluşturma hatası: {e}")

    # ============================================
    # UTILITY METHODS
    # ============================================

    def cleanup_old_exports(self, days: int = 30) -> int:
        """
        Eski export dosyalarını temizle
        
        Args:
            days: Kaç günden eski dosyalar silinsin
        
        Returns:
            int: Silinen dosya sayısı
        """
        from datetime import timedelta

        count = 0
        cutoff_time = datetime.now() - timedelta(days=days)

        try:
            for filename in os.listdir(self.export_dir):
                filepath = os.path.join(self.export_dir, filename)

                if os.path.isfile(filepath):
                    file_time = datetime.fromtimestamp(os.path.getmtime(filepath))

                    if file_time < cutoff_time:
                        os.remove(filepath)
                        count += 1

            return count

        except Exception as e:
            logging.error(f"Temizleme hatası: {e}")
            return count

