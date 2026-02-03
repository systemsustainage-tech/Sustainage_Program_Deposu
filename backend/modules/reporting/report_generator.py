#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Rapor Oluşturma Modülü
Rapor şablonları ve otomatik rapor üretimi
"""

import logging
import os
import sqlite3
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from config.database import DB_PATH


class ReportGenerator:
    """Rapor oluşturma ve şablon yönetimi"""

    def __init__(self, db_path: str = DB_PATH) -> None:
        if not os.path.isabs(db_path):
            base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
            db_path = os.path.join(base_dir, db_path)
        self.db_path = db_path
        self._init_db_tables()

    def _init_db_tables(self) -> None:
        """Raporlama tablolarını oluştur"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        try:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS report_templates (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    company_id INTEGER NOT NULL,
                    template_name TEXT NOT NULL,
                    template_type TEXT NOT NULL,
                    template_content TEXT NOT NULL,
                    template_variables TEXT,
                    language_code TEXT DEFAULT 'tr',
                    status TEXT DEFAULT 'active',
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (company_id) REFERENCES companies(id)
                )
            """)

            cursor.execute("""
                CREATE TABLE IF NOT EXISTS generated_reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    company_id INTEGER NOT NULL,
                    template_id INTEGER NOT NULL,
                    report_name TEXT NOT NULL,
                    report_period TEXT NOT NULL,
                    report_format TEXT NOT NULL,
                    file_path TEXT,
                    generation_date TEXT NOT NULL,
                    generated_by TEXT,
                    status TEXT DEFAULT 'generated',
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (company_id) REFERENCES companies(id),
                    FOREIGN KEY (template_id) REFERENCES report_templates(id)
                )
            """)

            conn.commit()
            logging.info("[OK] Raporlama modulu tablolari basariyla olusturuldu")

        except Exception as e:
            logging.error(f"[HATA] Raporlama modulu tablo olusturma: {e}")
            conn.rollback()
        finally:
            conn.close()

    def add_report_template(self, company_id: int, template_name: str, template_type: str,
                          template_content: str, template_variables: Optional[str] = None,
                          language_code: str = 'tr') -> bool:
        """Rapor şablonu ekle"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        try:
            cursor.execute("""
                INSERT INTO report_templates 
                (company_id, template_name, template_type, template_content,
                 template_variables, language_code)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (company_id, template_name, template_type, template_content,
                  template_variables, language_code))

            conn.commit()
            return True

        except Exception as e:
            logging.error(f"Rapor şablonu ekleme hatası: {e}")
            conn.rollback()
            return False
        finally:
            conn.close()

    def generate_report(self, company_id: int, template_id: int, report_name: str,
                       report_period: str, report_format: str, generated_by: Optional[str] = None) -> bool:
        """Rapor oluştur"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        try:
            generation_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            file_path = f"reports/{report_name}_{report_period}.{report_format}"

            cursor.execute("""
                INSERT INTO generated_reports 
                (company_id, template_id, report_name, report_period, report_format,
                 file_path, generation_date, generated_by)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (company_id, template_id, report_name, report_period, report_format,
                  file_path, generation_date, generated_by))

            conn.commit()
            return True

        except Exception as e:
            logging.error(f"Rapor oluşturma hatası: {e}")
            conn.rollback()
            return False
        finally:
            conn.close()

    def get_report_templates(self, company_id: int, template_type: Optional[str] = None) -> List[Dict]:
        """Rapor şablonlarını getir"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        try:
            if template_type:
                cursor.execute("""
                    SELECT id, template_name, template_type, template_content, language_code
                    FROM report_templates 
                    WHERE company_id = ? AND template_type = ? AND status = 'active'
                    ORDER BY template_name
                """, (company_id, template_type))
            else:
                cursor.execute("""
                    SELECT id, template_name, template_type, template_content, language_code
                    FROM report_templates 
                    WHERE company_id = ? AND status = 'active'
                    ORDER BY template_name
                """, (company_id,))

            templates = []
            for row in cursor.fetchall():
                templates.append({
                    'id': row[0],
                    'template_name': row[1],
                    'template_type': row[2],
                    'template_content': row[3],
                    'language_code': row[4]
                })

            return templates

        except Exception as e:
            logging.error(f"Rapor şablonları getirme hatası: {e}")
            return []
        finally:
            conn.close()

    def _get_sustainability_context(self, company_id: int, report_period: str) -> Tuple[Dict[str, List[Dict]], str]:
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row

        report_data: Dict[str, List[Dict]] = {}

        try:
            try:
                rows = conn.execute(
                    "SELECT * FROM carbon_emissions WHERE company_id = ? ORDER BY created_at DESC LIMIT 5",
                    (company_id,)
                ).fetchall()
                report_data["carbon"] = [dict(r) for r in rows]
            except Exception:
                report_data["carbon"] = []

            try:
                rows = conn.execute(
                    "SELECT * FROM energy_consumption WHERE company_id = ? ORDER BY created_at DESC LIMIT 5",
                    (company_id,)
                ).fetchall()
                report_data["energy"] = [dict(r) for r in rows]
            except Exception:
                report_data["energy"] = []

            try:
                rows = conn.execute(
                    "SELECT * FROM water_consumption WHERE company_id = ? ORDER BY created_at DESC LIMIT 5",
                    (company_id,)
                ).fetchall()
                report_data["water"] = [dict(r) for r in rows]
            except Exception:
                report_data["water"] = []

            try:
                rows = conn.execute(
                    "SELECT * FROM waste_generation WHERE company_id = ? ORDER BY created_at DESC LIMIT 5",
                    (company_id,)
                ).fetchall()
                report_data["waste"] = [dict(r) for r in rows]
            except Exception:
                report_data["waste"] = []
        finally:
            conn.close()

        try:
            from modules.ai.ai_manager import AIManager

            ai = AIManager(self.db_path)
            if ai.is_available():
                summary_text = ai.generate_summary(report_data, report_type="genel") or ""
            else:
                summary_text = self._simulate_ai_response(report_data)
        except Exception:
            summary_text = self._simulate_ai_response(report_data)

        return report_data, summary_text

    def generate_sustainability_report(self, company_id: int, report_period: str) -> str:
        report_data, summary = self._get_sustainability_context(company_id, report_period)

        md_content = f"""# Sürdürülebilirlik Raporu (AI Generated)
**Tarih:** {datetime.now().strftime('%Y-%m-%d %H:%M')}

## Yönetici Özeti
{summary}

## Veri Özeti
### Karbon Emisyonları
- Toplam Kayıt: {len(report_data.get('carbon', []))}
{self._table_to_md(report_data.get('carbon', []))}

### Enerji Tüketimi
- Toplam Kayıt: {len(report_data.get('energy', []))}
{self._table_to_md(report_data.get('energy', []))}

### Su Tüketimi
- Toplam Kayıt: {len(report_data.get('water', []))}
{self._table_to_md(report_data.get('water', []))}

### Atık Yönetimi
- Toplam Kayıt: {len(report_data.get('waste', []))}
{self._table_to_md(report_data.get('waste', []))}

---
*Bu rapor SustainAge AI tarafından oluşturulmuştur.*
"""
        return md_content

    def _simulate_ai_response(self, data: Dict[str, List[Dict]]) -> str:
        return (
            "Sistemde kayıtlı veriler analiz edildiğinde:\n"
            f"- Karbon emisyonları için {len(data.get('carbon', []))} adet kayıt bulundu.\n"
            f"- Enerji tüketimi için {len(data.get('energy', []))} adet kayıt bulundu.\n"
            "- Su ve atık yönetimi modülleri aktif olarak veri akışı sağlamaktadır.\n\n"
            "Genel olarak, kurumun sürdürülebilirlik performansı izlenebilir durumdadır. \n"
            "Veri girişlerinin düzenli yapılması, daha hassas analizler sağlayacaktır."
        )

    def _table_to_md(self, rows: List[Dict]) -> str:
        if not rows:
            return "_Veri yok._"

        keys = list(rows[0].keys())
        header = "| " + " | ".join(keys) + " |"
        sep = "| " + " | ".join(["---"] * len(keys)) + " |"

        lines = [header, sep]
        for row in rows:
            vals = [str(row.get(k, "")) for k in keys]
            lines.append("| " + " | ".join(vals) + " |")

        return "\n".join(lines)

    def generate_sustainability_docx(self, company_id: int, report_period: str, output_path: str) -> Optional[str]:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt
        except Exception as e:
            logging.error(f"Sustainability DOCX import error: {e}")
            return None

        report_data, summary = self._get_sustainability_context(company_id, report_period)

        doc = Document()
        title = doc.add_heading("Sürdürülebilirlik Raporu (AI Generated)", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"Tarih: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        meta_run.font.size = Pt(10)

        doc.add_heading("Yönetici Özeti", level=1)
        for part in summary.split("\n"):
            if part.strip():
                doc.add_paragraph(part.strip())

        sections = [
            ("carbon", "Karbon Emisyonları"),
            ("energy", "Enerji Tüketimi"),
            ("water", "Su Tüketimi"),
            ("waste", "Atık Yönetimi"),
        ]

        for key, title_text in sections:
            rows = report_data.get(key, [])
            doc.add_heading(title_text, level=2)
            doc.add_paragraph(f"Toplam Kayıt: {len(rows)}")
            if rows:
                keys = list(rows[0].keys())
                table = doc.add_table(rows=len(rows) + 1, cols=len(keys))
                header_cells = table.rows[0].cells
                for idx, k in enumerate(keys):
                    header_cells[idx].text = str(k)
                for i, row in enumerate(rows, start=1):
                    row_cells = table.rows[i].cells
                    for j, k in enumerate(keys):
                        row_cells[j].text = str(row.get(k, ""))

        doc.add_paragraph("Bu rapor SustainAge AI tarafından oluşturulmuştur.")

        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            return output_path
        except Exception as e:
            logging.error(f"Sustainability DOCX save error: {e}")
            return None

    def generate_sustainability_pdf(self, company_id: int, report_period: str, output_path: str) -> Optional[str]:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except Exception as e:
            logging.error(f"Sustainability PDF import error: {e}")
            return None

        try:
            from .font_utils import register_turkish_fonts_reportlab

            register_turkish_fonts_reportlab()
        except Exception as e:
            logging.error(f"Sustainability PDF font error: {e}")

        report_data, summary = self._get_sustainability_context(company_id, report_period)

        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            title_style = ParagraphStyle(
                "SustainabilityTitle",
                parent=styles["Heading1"],
                fontSize=22,
                spaceAfter=24,
                alignment=1,
                fontName="NotoSans",
            )

            story.append(Paragraph("Sürdürülebilirlik Raporu (AI Generated)", title_style))
            story.append(Spacer(1, 12))
            story.append(
                Paragraph(
                    f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 16))

            story.append(Paragraph("Yönetici Özeti", styles["Heading2"]))
            story.append(Spacer(1, 8))
            for part in summary.split("\n"):
                if part.strip():
                    story.append(Paragraph(part.strip(), styles["Normal"]))
            story.append(Spacer(1, 16))

            sections = [
                ("carbon", "Karbon Emisyonları"),
                ("energy", "Enerji Tüketimi"),
                ("water", "Su Tüketimi"),
                ("waste", "Atık Yönetimi"),
            ]

            for key, title_text in sections:
                rows = report_data.get(key, [])
                story.append(Paragraph(title_text, styles["Heading2"]))
                story.append(Spacer(1, 6))
                story.append(
                    Paragraph(f"Toplam Kayıt: {len(rows)}", styles["Normal"])
                )
                story.append(Spacer(1, 6))
                if rows:
                    keys = list(rows[0].keys())
                    data = [keys]
                    for row in rows:
                        data.append([str(row.get(k, "")) for k in keys])
                    table = Table(data)
                    table.setStyle(
                        TableStyle(
                            [
                                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                                ("FONTNAME", (0, 0), (-1, 0), "NotoSans-Bold"),
                                ("FONTNAME", (0, 1), (-1, -1), "NotoSans"),
                                ("FONTSIZE", (0, 0), (-1, -1), 8),
                                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                            ]
                        )
                    )
                    story.append(table)
                    story.append(Spacer(1, 12))

            story.append(
                Paragraph(
                    "Bu rapor SustainAge AI tarafından oluşturulmuştur.",
                    styles["Italic"],
                )
            )

            doc.build(story)
            return output_path
        except Exception as e:
            logging.error(f"Sustainability PDF generate error: {e}")
            return None
