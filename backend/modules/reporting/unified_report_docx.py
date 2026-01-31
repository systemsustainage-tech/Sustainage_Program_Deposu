#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
from datetime import datetime
from typing import Dict, List, Optional
import json
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import sqlite3
import logging


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_paragraph(doc: Document, text: str = "", bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


class UnifiedReportDocxGenerator:
    def __init__(self, base_dir: str) -> None:
        self.base_dir = base_dir
        # Ensure config can be imported
        import sys
        if base_dir not in sys.path:
            sys.path.append(base_dir)
        try:
            from config.database import DB_PATH
        except ImportError:
            # Fallback if config not found
            DB_PATH = os.path.join(base_dir, 'backend', 'data', 'sdg_desktop.sqlite')
        self.db_path = DB_PATH
        self.logger = logging.getLogger(__name__)

    def _build_target_summary(self, cur, company_id) -> tuple:
        """Hedef takibi özetini oluşturur."""
        try:
            # company_targets tablosu var mı kontrol et
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='company_targets'")
            if not cur.fetchone():
                return None, []

            cur.execute("""
                SELECT metric_type, baseline_year, target_year, baseline_value, target_value, status
                FROM company_targets 
                WHERE company_id=? 
                ORDER BY target_year ASC
            """, (company_id,))
            
            targets = cur.fetchall()
            if not targets:
                return None, []
                
            summary_text = f"Şirket, {len(targets)} adet sürdürülebilirlik hedefi belirlemiştir."
            return summary_text, targets
        except Exception as e:
            self.logger.error(f"Hedef özeti hatası: {e}")
            return None, []

    def _build_supply_chain_summary(self, cur, company_id, period) -> tuple:
        """Tedarik zinciri (Scope 3) verilerini özetler."""
        try:
            # supplier_environmental_data tablosu var mı kontrol et
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='supplier_environmental_data'")
            if not cur.fetchone():
                return None, {}

            # Bu şirkete bağlı tedarikçilerin verilerini çekmemiz lazım.
            # Ancak şu an supplier_environmental_data'da company_id yok, supplier_id var.
            # users tablosunda supplier'ın parent_company_id'si var mı?
            # Varsayım: supplier_portal üzerinden girilen verilerde supplier_id, users tablosundaki id'dir.
            # users tablosunda 'company_id' sütunu aslında kullanıcının bağlı olduğu ana şirketi gösterebilir (eğer supplier ise).
            # Veya 'suppliers' tablosu var mı?
            
            # Şimdilik basitçe supplier_environmental_data tablosundaki tüm verileri çekelim (tek şirket varsayımı veya demo için).
            # Gerçek senaryoda: SELECT * FROM supplier_environmental_data s JOIN users u ON s.supplier_id = u.id WHERE u.parent_company_id = ?
            
            cur.execute("""
                SELECT 
                    COUNT(DISTINCT supplier_id) as supplier_count,
                    SUM(energy_consumption) as total_energy,
                    SUM(emissions) as total_emissions,
                    SUM(waste_amount) as total_waste,
                    SUM(water_consumption) as total_water
                FROM supplier_environmental_data
                WHERE period = ?
            """, (period,))
            
            row = cur.fetchone()
            if not row or row[0] == 0:
                return None, {}
                
            data = {
                'supplier_count': row[0],
                'energy': row[1] or 0,
                'emissions': row[2] or 0,
                'waste': row[3] or 0,
                'water': row[4] or 0
            }
            
            summary_text = (f"{period} döneminde {data['supplier_count']} tedarikçiden çevresel veri toplanmıştır. "
                            f"Toplam Scope 3 (Tedarik Zinciri) Emisyonu: {data['emissions']:.2f} tCO2e.")
            return summary_text, data
        except Exception as e:
            self.logger.error(f"Tedarik zinciri özeti hatası: {e}")
            return None, {}

    def _build_summaries(self, cur, company_id, period) -> tuple:
        """SDG, GRI ve TSRS özetlerini oluşturur."""
        # SDG özet (responses tablosu bazlı)
        sdg_ozet = ""
        try:
            # responses tablosu var mı kontrol et
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='responses'")
            if cur.fetchone():
                cur2 = cur.execute(
                    """SELECT COUNT(*), AVG(progress_pct) FROM responses WHERE company_id=? AND period=?""",
                    (company_id, period)
                ).fetchone()
                total_resp = cur2[0] or 0
                avg_progress = round(cur2[1] or 0, 1)
                sdg_ozet = f"Toplam {total_resp} gösterge için veri girildi. Ortalama ilerleme: %{avg_progress}."
        except Exception as e:
            self.logger.error(f"SDG özet hatası: {e}")

        # GRI özet
        gri_ozet = ""
        try:
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='gri_selections'")
            if cur.fetchone():
                cur3 = cur.execute(
                    """SELECT COUNT(*) FROM gri_selections WHERE company_id=? AND selected=1""",
                    (company_id,)
                ).fetchone()
                gri_count = cur3[0] or 0
                gri_ozet = f"Seçilen GRI gösterge sayısı: {gri_count}."
        except Exception as e:
            self.logger.error(f"GRI özet hatası: {e}")

        # TSRS özet
        tsrs_ozet = ""
        if gri_ozet:
            tsrs_ozet = "TSRS uyumu, seçilen GRI göstergeleri üzerinden takip edilmektedir."
            
        return sdg_ozet, gri_ozet, tsrs_ozet

    def _get_gri_selected_rows(self, cur, company_id) -> list:
        """Seçilmiş GRI göstergelerini ve TSRS eşleştirmelerini döndürür."""
        try:
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='gri_selections'")
            if not cur.fetchone():
                return []
                
            rows = cur.execute(
                """
                SELECT gi.code AS indicator_code, gi.title AS indicator_title,
                       gs.code AS standard_code, gs.title AS standard_title
                FROM gri_selections sel
                JOIN gri_indicators gi ON gi.id = sel.indicator_id
                JOIN gri_standards gs ON gs.id = gi.standard_id
                WHERE sel.company_id=? AND sel.selected=1
                ORDER BY gs.code, gi.code
                """,
                (company_id,)
            ).fetchall()
            
            result = []
            for r in rows:
                icode = r[0]
                ititle = r[1]
                scode = r[2]
                stitle = r[3]
                # row_data: (scode, icode, ititle, tsrs_summary)
                result.append((scode, icode, ititle, ""))
            return result
        except Exception as e:
            self.logger.error(f"GRI satırları hatası: {e}")
            return []

    def _get_ceo_messages(self, cur, company_id, period) -> list:
        try:
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='ceo_messages'")
            if not cur.fetchone():
                return []

            cur.execute("SELECT * FROM ceo_messages WHERE company_id=? AND period=?", (company_id, period))
            if cur.description:
                cols = [description[0] for description in cur.description]
                rows = cur.fetchall()
                messages = []
                for row in rows:
                    msg = dict(zip(cols, row))
                    messages.append(msg)
                return messages
            return []
        except Exception as e:
            self.logger.error(f"CEO mesajları hatası: {e}")
            return []

    def _add_methodology_section(self, doc, cur, company_id):
        """Metodoloji ve Veri Kalitesi bölümünü ekler."""
        _add_heading(doc, "6. Metodoloji ve Veri Kalitesi")
        
        intro = (
            "Bu rapor, Sustainage platformu kullanılarak hazırlanmıştır. "
            "Veriler, şirket yetkilileri tarafından girilmiş ve otomatik doğrulama süreçlerinden geçirilmiştir."
        )
        _add_paragraph(doc, intro)
        
        # SDG Validation Results
        try:
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='sdg_validation_results'")
            if cur.fetchone():
                # Kalite puanlarını al (son doğrulama)
                # sdg_data_quality_scores tablosu varsa oradan, yoksa validation_results'dan özet çıkar
                
                # Önce sdg_data_quality_scores kontrolü
                cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='sdg_data_quality_scores'")
                if cur.fetchone():
                    scores = cur.execute(
                        "SELECT completeness_score, accuracy_score, consistency_score, timeliness_score, overall_quality_score "
                        "FROM sdg_data_quality_scores WHERE company_id=? ORDER BY validation_date DESC LIMIT 1",
                        (company_id,)
                    ).fetchone()
                    if scores:
                        _add_heading(doc, "Veri Kalitesi Puanları", level=2)
                        _add_paragraph(doc, f"Genel Kalite Skoru: {scores[4]:.2f}/100")
                        _add_paragraph(doc, f"- Tamlık: {scores[0]:.2f}")
                        _add_paragraph(doc, f"- Doğruluk: {scores[1]:.2f}")
                        _add_paragraph(doc, f"- Tutarlılık: {scores[2]:.2f}")
                        _add_paragraph(doc, f"- Zamanlılık: {scores[3]:.2f}")

                # Validation hataları özeti
                issues = cur.execute(
                    """
                    SELECT r.rule_name, COUNT(*) as issue_count 
                    FROM sdg_validation_results vr
                    JOIN sdg_validation_rules r ON vr.rule_id = r.id
                    WHERE vr.company_id=? AND vr.validation_status='failed'
                    GROUP BY r.rule_name
                    """, 
                    (company_id,)
                ).fetchall()
                
                if issues:
                    _add_heading(doc, "Tespit Edilen Veri İyileştirme Alanları", level=2)
                    for rule, count in issues:
                        _add_paragraph(doc, f"- {rule}: {count} kayıt incelenmeli")
                else:
                    _add_paragraph(doc, "Veri doğrulama sürecinde kritik bir hata tespit edilmemiştir.")
                    
        except Exception as e:
            self.logger.error(f"Metodoloji bölümü hatası: {e}")
            _add_paragraph(doc, "Veri kalitesi bilgileri alınırken bir hata oluştu.")

    def _add_materiality_section(self, doc, cur, company_id):
        """Çifte Önemlilik (Double Materiality) bölümünü ekler."""
        try:
            # materiality_topics tablosu var mı kontrol et
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='materiality_topics'")
            if not cur.fetchone():
                return

            cur.execute("""
                SELECT topic_name, category, stakeholder_impact, business_impact, priority_score
                FROM materiality_topics
                WHERE company_id=?
                ORDER BY priority_score DESC
                LIMIT 10
            """, (company_id,))
            
            topics = cur.fetchall()
            if not topics:
                return

            _add_heading(doc, "Çifte Önemlilik (Double Materiality) Analizi")
            _add_paragraph(doc, "Şirketin sürdürülebilirlik performansını etkileyen öncelikli konular, Çifte Önemlilik (Double Materiality) prensibine göre belirlenmiştir.")
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Konu"
            hdr[1].text = "Kategori"
            hdr[2].text = "Etki Önemliliği"
            hdr[3].text = "Finansal Önemlilik"
            
            for t in topics:
                row = table.add_row().cells
                row[0].text = str(t[0])
                row[1].text = str(t[1])
                row[2].text = f"{t[2]:.1f}" if t[2] else "-"
                row[3].text = f"{t[3]:.1f}" if t[3] else "-"
            
            _add_paragraph(doc, "Not: Etki Önemliliği (Impact Materiality) şirketin çevre ve toplum üzerindeki etkisini; Finansal Önemlilik (Financial Materiality) ise sürdürülebilirlik konularının şirketin finansal durumu üzerindeki etkisini ifade eder.")

        except Exception as e:
            self.logger.error(f"Önemlilik bölümü hatası: {e}")

    def _add_esrs_section(self, doc, cur, company_id):
        """ESRS Uyum bölümünü ekler."""
        try:
            has_data = False
            # Check esrs_disclosures (Genel istatistikler)
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='esrs_disclosures'")
            if cur.fetchone():
                cur.execute("""
                    SELECT 
                        COUNT(*) as total,
                        SUM(CASE WHEN completion_status = 'tamamlandi' THEN 1 ELSE 0 END) as completed,
                        SUM(CASE WHEN is_material = 1 THEN 1 ELSE 0 END) as material_count
                    FROM esrs_disclosures
                    WHERE company_id=?
                """, (company_id,))
                stats = cur.fetchone()
                
                if stats and stats[0] > 0:
                    has_data = True
                    _add_heading(doc, "ESRS (Avrupa Sürdürülebilirlik Raporlama Standartları) Uyumu")
                    total = stats[0]
                    completed = stats[1] or 0
                    material = stats[2] or 0
                    pct = (completed / total) * 100 if total > 0 else 0
                    
                    _add_paragraph(doc, f"Toplam ESRS Açıklaması: {total}")
                    _add_paragraph(doc, f"Tamamlanan: {completed} (%{pct:.1f})")
                    _add_paragraph(doc, f"Önemli (Material) Olarak Belirlenen: {material}")

            # Detaylı Değerlendirme (esrs_assessments)
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='esrs_assessments'")
            if cur.fetchone():
                # Yeni eklenen not alanlarını da çekiyoruz
                cur.execute("""
                    SELECT standard_code, status, 
                           governance_notes, strategy_notes, impact_risk_notes, metrics_notes
                    FROM esrs_assessments 
                    WHERE company_id=?
                    ORDER BY standard_code
                """, (company_id,))
                assessments = cur.fetchall()
                if assessments:
                    if not has_data:
                        _add_heading(doc, "ESRS (Avrupa Sürdürülebilirlik Raporlama Standartları) Uyumu")
                        has_data = True
                    
                    _add_heading(doc, "Standart Bazlı Değerlendirme ve Detaylar", level=2)
                    
                    for a in assessments:
                        std_code = a[0]
                        status = a[1]
                        gov_notes = a[2]
                        strat_notes = a[3]
                        risk_notes = a[4]
                        met_notes = a[5]
                        
                        # Başlık: ESRS 1 - Tamamlandı
                        p = doc.add_paragraph()
                        run = p.add_run(f"{std_code} - {status}")
                        run.bold = True
                        
                        # Varsa detay notları ekle
                        notes_added = False
                        if gov_notes:
                            _add_paragraph(doc, f"Yönetişim: {gov_notes}", style='List Bullet')
                            notes_added = True
                        if strat_notes:
                            _add_paragraph(doc, f"Strateji: {strat_notes}", style='List Bullet')
                            notes_added = True
                        if risk_notes:
                            _add_paragraph(doc, f"Risk ve Etki Yönetimi: {risk_notes}", style='List Bullet')
                            notes_added = True
                        if met_notes:
                            _add_paragraph(doc, f"Metrikler ve Hedefler: {met_notes}", style='List Bullet')
                            notes_added = True
                            
                        if not notes_added and status != 'not_started':
                            _add_paragraph(doc, "Detaylı açıklama girilmemiştir.", style='List Bullet')

        except Exception as e:
            self.logger.error(f"ESRS bölümü hatası: {e}")

    def _add_stakeholder_survey_section(self, doc, survey_data: Dict, section_prefix: str = ""):
        """Paydaş Anketi Sonuçları bölümünü ekler."""
        if not survey_data:
            return

        heading_text = "Paydaş Katılımı ve Anket Sonuçları"
        if section_prefix:
            heading_text = f"{section_prefix} {heading_text}"
        
        _add_heading(doc, heading_text)
        
        title = survey_data.get("survey_title", "Paydaş Anketi")
        desc = survey_data.get("survey_description", "")
        count = survey_data.get("response_count", 0)
        created_at = survey_data.get("created_at", "")
        insights = survey_data.get("insights_text", "")

        _add_paragraph(doc, f"Anket Başlığı: {title}", bold=True)
        if desc:
            _add_paragraph(doc, f"Açıklama: {desc}")
        _add_paragraph(doc, f"Toplam Katılımcı Sayısı: {count}")
        if created_at:
            _add_paragraph(doc, f"Oluşturulma Tarihi: {created_at}")
        
        if insights:
            _add_paragraph(doc, "Anket Analizi ve İçgörüler:", bold=True)
            _add_paragraph(doc, insights)

        # Top 3 / Bottom 3
        top3 = survey_data.get("top3", [])
        bottom3 = survey_data.get("bottom3", [])

        if top3:
            _add_heading(doc, "En Yüksek Öncelikli Konular (Top 3)", level=2)
            for item in top3:
                q_text = item.get("question") or item.get("title") or item.get("id")
                score = item.get("avg_score", 0)
                _add_paragraph(doc, f"- {q_text} (Skor: {score:.2f})")

        if bottom3:
            _add_heading(doc, "Gelişime Açık Konular (Bottom 3)", level=2)
            for item in bottom3:
                q_text = item.get("question") or item.get("title") or item.get("id")
                score = item.get("avg_score", 0)
                _add_paragraph(doc, f"- {q_text} (Skor: {score:.2f})")

        # Tüm Sorular Tablosu
        questions = survey_data.get("questions", [])
        if questions:
            _add_heading(doc, "Tüm Konu Başlıkları ve Skorlar", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "SDG Hedefi"
            hdr[1].text = "Konu / Soru"
            hdr[2].text = "Ortalama Skor"

            for q in questions:
                row = table.add_row().cells
                goal = q.get("goal")
                row[0].text = f"SDG {goal}" if goal else "-"
                row[1].text = str(q.get("question") or q.get("title") or "")
                row[2].text = f"{q.get('avg_score', 0):.2f}"

    def generate(
        self,
        company_id: int,
        reporting_period: str,
        report_name: str,
        selected_modules: List[str],
        module_reports: List[Dict],
        ai_comment: str,
        description: str,
        metrics: Dict,
    ) -> Optional[str]:
        try:
            doc = Document()

            db_path = self.db_path

            try:
                from modules.reporting.brand_identity_manager import BrandIdentityManager
                from docx.shared import Inches
                bim = BrandIdentityManager(db_path, company_id)
                bi = bim.get_brand_identity(company_id)
                logo_path = bi.get("logo_path")
                if logo_path and os.path.exists(logo_path):
                    logo_paragraph = doc.add_paragraph()
                    logo_run = logo_paragraph.add_run()
                    logo_run.add_picture(logo_path, width=Inches(1.6))
                    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_text = (bi.get("texts") or {}).get("header")
                if header_text:
                    header_paragraph = doc.add_paragraph(header_text)
                    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

            try:
                from docx.shared import Inches
                sdg_wheel = os.path.join(self.base_dir, "static", "images", "SDGs.jpeg")
                if os.path.exists(sdg_wheel):
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(sdg_wheel, width=Inches(2.5))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

            sdg_summary_text = ""
            gri_summary_text = ""
            tsrs_summary_text = ""
            methodology_text = ""
            ceo_message = None
            gri_selected_rows = []
            esg_scores = None
            ungc_result = None

            target_summary_text = None
            target_data = []
            supply_chain_summary_text = None
            supply_chain_data = {}

            try:
                conn = sqlite3.connect(db_path)
                cur = conn.cursor()
                sdg_ozet, gri_ozet, tsrs_ozet = self._build_summaries(cur, company_id, reporting_period)
                # sdg_validation_results üzerinden doğrulama özeti zaten _add_methodology_section içinde kullanılıyor
                # burada sadece özet metinleri alıyoruz
                
                sdg_summary_text = sdg_ozet or ""
                gri_summary_text = gri_ozet or ""
                tsrs_summary_text = tsrs_ozet or ""
                
                # Hedef ve Tedarik Zinciri Özetleri
                target_summary_text, target_data = self._build_target_summary(cur, company_id)
                supply_chain_summary_text, supply_chain_data = self._build_supply_chain_summary(cur, company_id, reporting_period)

                try:
                    gri_selected_rows = self._get_gri_selected_rows(cur, company_id) or []
                except Exception:
                    gri_selected_rows = []
                try:
                    ceo_messages = self._get_ceo_messages(cur, company_id, reporting_period)
                except Exception:
                    ceo_messages = []
                if ceo_messages:
                    annual_message = next((m for m in ceo_messages if m.get("message_type") == "annual"), None)
                    ceo_message = annual_message or ceo_messages[0]
            except Exception as e:
                self.logger.error(f"Özet oluşturma hatası: {e}")
            finally:
                try:
                    conn.close()
                except Exception:
                    pass

            try:
                from modules.esg.esg_manager import ESGManager
                esg_mgr = ESGManager(self.base_dir)
                esg_scores = esg_mgr.compute_scores(company_id, reporting_period)
            except Exception:
                esg_scores = None

            try:
                from modules.ungc.ungc_manager import UNGCManager
                ungc_mgr = UNGCManager(db_path)
                ungc_result = ungc_mgr.compute_principle_status(company_id, reporting_period)
            except Exception:
                ungc_result = None

            # Başlık ve meta bilgileri kapak sayfasına taşındı
            # Burası artık içerik başlangıcı


            _add_heading(doc, "1. Kapsam ve Seçilen Modüller")
            if selected_modules:
                for m in selected_modules:
                    _add_paragraph(doc, f"- {m}")
            else:
                _add_paragraph(doc, "Seçilen modül bulunmamaktadır.")

            if module_reports:
                _add_heading(doc, "2. Modül Bazlı Mevcut Raporlar")
                for r in module_reports:
                    line = (
                        f"- {r.get('module_code')}: {r.get('report_name')} "
                        f"({r.get('report_type')}, {r.get('reporting_period')}, {r.get('created_at')})"
                    )
                    _add_paragraph(doc, line)
                summary_by_module = {}
                for r in module_reports:
                    code = r.get("module_code") or ""
                    summary_by_module.setdefault(code, 0)
                    summary_by_module[code] += 1
                if summary_by_module:
                    _add_paragraph(doc)
                    _add_paragraph(doc, "Modül Bazında Rapor Sayıları:")
                    for code, count in sorted(summary_by_module.items(), key=lambda x: x[0]):
                        _add_paragraph(doc, f"- {code}: {count} rapor")

            if ai_comment:
                _add_heading(doc, "3. AI Yorumlu Sürdürülebilirlik Özeti")
                
                # Check if this is the structured analysis from AI Reports module
                if "1. Yönetici Özeti" in ai_comment:
                    # Parse the structured text manually to apply formatting
                    parts = ai_comment.split("**")
                    current_paragraph = []
                    
                    for part in parts:
                        part = part.strip()
                        if not part: continue
                        
                        # Headings usually start with numbers like "1. " or "2. "
                        if part[0].isdigit() and part[1:3] in [". ", ". \n"]:
                            # Flush previous paragraph
                            if current_paragraph:
                                _add_paragraph(doc, "".join(current_paragraph))
                                current_paragraph = []
                            # Add as heading
                            _add_heading(doc, part, level=2)
                        else:
                            # It's content, might contain bullet points (*)
                            if "*" in part:
                                bullets = part.split("*")
                                for b in bullets:
                                    b = b.strip()
                                    if b:
                                        if b.startswith("Not:"):
                                             _add_paragraph(doc, b, italic=True)
                                        else:
                                             _add_paragraph(doc, b, style='List Bullet')
                            else:
                                _add_paragraph(doc, part)
                else:
                    # Fallback for plain text
                    for part in ai_comment.split("\n"):
                         _add_paragraph(doc, part)

            if description:
                _add_heading(doc, "4. Kullanıcı Açıklaması")
                for part in description.split("\n"):
                    _add_paragraph(doc, part)

            # --- YENİ BÖLÜM: SDG, GRI ve TSRS Özetleri ---
            if sdg_summary_text or gri_summary_text or tsrs_summary_text:
                _add_heading(doc, "5. Entegrasyon ve İlerleme Özetleri")
                
                if sdg_summary_text:
                    _add_heading(doc, "SDG İlerleme Özeti", level=2)
                    _add_paragraph(doc, sdg_summary_text)
                
                if gri_summary_text:
                    _add_heading(doc, "GRI Eşleştirme Özeti", level=2)
                    _add_paragraph(doc, gri_summary_text)
                    
                    if gri_selected_rows:
                        _add_paragraph(doc) # Boşluk
                        _add_heading(doc, "GRI ve TSRS Eşleştirme Tablosu", level=3)
                        table = doc.add_table(rows=1, cols=4)
                        table.style = 'Table Grid'
                        hdr = table.rows[0].cells
                        hdr[0].text = "GRI Standart"
                        hdr[1].text = "GRI Gösterge"
                        hdr[2].text = "Başlık"
                        hdr[3].text = "TSRS Eşleşmesi"
                        for row_data in gri_selected_rows:
                            # row_data: (scode, icode, ititle, tsrs_summary)
                            cells = table.add_row().cells
                            cells[0].text = str(row_data[0] or "")
                            cells[1].text = str(row_data[1] or "")
                            cells[2].text = str(row_data[2] or "")
                            cells[3].text = str(row_data[3] or "")

                if tsrs_summary_text:
                    _add_heading(doc, "TSRS Eşleştirme Özeti", level=2)
                    _add_paragraph(doc, tsrs_summary_text)

            # --- YENİ BÖLÜM: Hedef ve Tedarik Zinciri Özetleri ---
            target_summary_text = ""
            target_list = []
            supply_chain_summary_text = ""
            supply_chain_data = {}
            
            try:
                conn = sqlite3.connect(db_path)
                cur = conn.cursor()
                target_summary_text, target_list = self._build_target_summary(cur, company_id)
                supply_chain_summary_text, supply_chain_data = self._build_supply_chain_summary(cur, company_id, reporting_period)
            except Exception as e:
                self.logger.error(f"Hedef/Tedarik özeti hatası: {e}")
            finally:
                try:
                    conn.close()
                except:
                    pass

            if target_summary_text or supply_chain_summary_text:
                 _add_heading(doc, "Kurumsal Performans ve Etki Analizi")
                 
                 if target_summary_text:
                     _add_heading(doc, "Sürdürülebilirlik Hedefleri", level=2)
                     _add_paragraph(doc, target_summary_text)
                     if target_list:
                         table = doc.add_table(rows=1, cols=4)
                         table.style = 'Table Grid'
                         hdr = table.rows[0].cells
                         hdr[0].text = "Kategori"
                         hdr[1].text = "Hedef Yılı"
                         hdr[2].text = "Hedef Değeri"
                         hdr[3].text = "Durum"
                         for t in target_list:
                             # t: category, baseline_year, target_year, baseline_value, target_value, status
                             row = table.add_row().cells
                             row[0].text = str(t[0] or "").capitalize()
                             row[1].text = str(t[2] or "")
                             row[2].text = str(t[4] or "")
                             row[3].text = str(t[5] or "")

                 if supply_chain_summary_text:
                     _add_heading(doc, "Tedarik Zinciri (Scope 3) Analizi", level=2)
                     _add_paragraph(doc, supply_chain_summary_text)
                     if supply_chain_data:
                         _add_paragraph(doc, "Detaylı Metrikler:", bold=True)
                         _add_paragraph(doc, f"- Toplam Enerji Tüketimi (Tedarikçi): {supply_chain_data.get('energy', 0):.2f}")
                         _add_paragraph(doc, f"- Toplam Su Tüketimi (Tedarikçi): {supply_chain_data.get('water', 0):.2f}")
                         _add_paragraph(doc, f"- Toplam Atık Miktarı (Tedarikçi): {supply_chain_data.get('waste', 0):.2f}")

            # 6. Çifte Önemlilik ve ESRS
            try:
                conn = sqlite3.connect(db_path)
                cur = conn.cursor()
                self._add_materiality_section(doc, cur, company_id)
                self._add_esrs_section(doc, cur, company_id)
            except Exception as e:
                self.logger.error(f"Önemlilik/ESRS bölümü eklenirken hata: {e}")
            finally:
                try:
                    conn.close()
                except Exception:
                    pass

            # 7. Metodoloji ve Veri Kalitesi
            try:
                conn = sqlite3.connect(db_path)
                cur = conn.cursor()
                self._add_methodology_section(doc, cur, company_id)
            except Exception as e:
                self.logger.error(f"Metodoloji bölümü eklenirken hata: {e}")
            finally:
                try:
                    conn.close()
                except Exception:
                    pass

            section_index = 7
            if ceo_message:
                heading_text = f"{section_index}. CEO/Genel Müdür Mesajı"
                _add_heading(doc, heading_text)
                title_parts = []
                msg_title = ceo_message.get("title")
                msg_year = ceo_message.get("year")
                msg_quarter = ceo_message.get("quarter")
                if msg_title:
                    title_parts.append(str(msg_title))
                if msg_year:
                    if msg_quarter:
                        title_parts.append(f"{msg_year} Q{msg_quarter}")
                    else:
                        title_parts.append(str(msg_year))
                if title_parts:
                    _add_heading(doc, " - ".join(title_parts), level=2)
                content_text = ceo_message.get("content")
                if content_text:
                    _add_paragraph(doc, content_text)
                    _add_paragraph(doc)
                key_achievements = ceo_message.get("key_achievements") or []
                if key_achievements:
                    _add_heading(doc, "Ana Başarılar", level=3)
                    for item in key_achievements:
                        if str(item).strip():
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(str(item))
                challenges = ceo_message.get("challenges") or []
                if challenges:
                    _add_heading(doc, "Karşılaşılan Zorluklar", level=3)
                    for item in challenges:
                        if str(item).strip():
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(str(item))
                future_commitments = ceo_message.get("future_commitments") or []
                if future_commitments:
                    _add_heading(doc, "Gelecek Taahhütleri", level=3)
                    for item in future_commitments:
                        if str(item).strip():
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(str(item))
                signature_name = ceo_message.get("signature_name")
                signature_title = ceo_message.get("signature_title")
                signature_date = ceo_message.get("signature_date")
                if signature_name or signature_title or signature_date:
                    _add_paragraph(doc)
                    _add_paragraph(doc, "Saygılarımla,")
                    _add_paragraph(doc)
                    if signature_name:
                        _add_paragraph(doc, str(signature_name))
                    if signature_title:
                        _add_paragraph(doc, str(signature_title))
                    if signature_date:
                        _add_paragraph(doc, str(signature_date))
                section_index += 1

            if metrics and isinstance(metrics, dict):
                survey = metrics.get('stakeholder_survey') or {}
                if survey:
                    try:
                        self._add_stakeholder_survey_section(doc, survey, section_prefix=f"{section_index}.")
                        section_index += 1
                    except Exception as e:
                        self.logger.error(f"Anket bölümü eklenirken hata: {e}")

                has_env = any(k in metrics for k in ['carbon', 'energy', 'water', 'waste'])
                if has_env:
                    _add_heading(doc, f"{section_index}. Çevresel Metrikler")
                    carbon = metrics.get('carbon') or {}
                    if carbon:
                        total_co2e = carbon.get('total_co2e', 0) or 0
                        _add_paragraph(doc, f"- Toplam CO2e emisyonu: {total_co2e:,.2f} tCO2e")
                        scope_breakdown = carbon.get('scope_breakdown') or []
                        if scope_breakdown:
                            _add_paragraph(doc, "  Scope bazında dağılım:", bold=True)
                            for item in scope_breakdown:
                                scope = item.get('scope') or ''
                                co2e_val = item.get('co2e', 0) or 0
                                _add_paragraph(doc, f"  - {scope}: {co2e_val:,.2f} tCO2e")
                    energy = metrics.get('energy') or {}
                    if energy:
                        total_energy = energy.get('total_consumption', 0) or 0
                        _add_paragraph(doc, f"- Toplam enerji tüketimi: {total_energy:,.2f}")
                        type_breakdown = energy.get('type_breakdown') or []
                        if type_breakdown:
                            _add_paragraph(doc, "  Enerji türüne göre dağılım:", bold=True)
                            for item in type_breakdown:
                                etype = item.get('energy_type') or ''
                                cons_val = item.get('consumption', 0) or 0
                                _add_paragraph(doc, f"  - {etype}: {cons_val:,.2f}")
                    water = metrics.get('water') or {}
                    if water:
                        total_water = water.get('total_consumption', 0) or 0
                        _add_paragraph(doc, f"- Toplam su tüketimi: {total_water:,.2f}")
                        source_breakdown = water.get('source_breakdown') or []
                        if source_breakdown:
                            _add_paragraph(doc, "  Su kaynağına göre dağılım:", bold=True)
                            for item in source_breakdown:
                                source = item.get('water_source') or ''
                                cons_val = item.get('consumption', 0) or 0
                                _add_paragraph(doc, f"  - {source}: {cons_val:,.2f}")
                    waste = metrics.get('waste') or {}
                    if waste:
                        total_waste = waste.get('total_amount', 0) or 0
                        _add_paragraph(doc, f"- Toplam atık miktarı: {total_waste:,.2f}")
                        type_breakdown = waste.get('type_breakdown') or []
                        if type_breakdown:
                            _add_paragraph(doc, "  Atık türüne göre dağılım:", bold=True)
                            for item in type_breakdown:
                                wtype = item.get('waste_type') or ''
                                amount_val = item.get('amount', 0) or 0
                                _add_paragraph(doc, f"  - {wtype}: {amount_val:,.2f}")
                    section_index += 1

            if metrics and isinstance(metrics, dict) and 'taxonomy' in metrics:
                tax = metrics.get('taxonomy') or {}
                _add_heading(doc, f"{section_index}. AB Taksonomisi Metrikleri")
                total_revenue = tax.get('total_revenue', 0) or 0
                aligned_revenue = tax.get('aligned_revenue', 0) or 0
                align_rev_pct = tax.get('alignment_percentage_revenue', 0) or 0
                total_capex = tax.get('total_capex', 0) or 0
                aligned_capex = tax.get('aligned_capex', 0) or 0
                align_capex_pct = tax.get('alignment_percentage_capex', 0) or 0
                total_opex = tax.get('total_opex', 0) or 0
                aligned_opex = tax.get('aligned_opex', 0) or 0
                align_opex_pct = tax.get('alignment_percentage_opex', 0) or 0
                _add_paragraph(doc, f"Toplam gelir: {total_revenue:,.2f} EUR")
                _add_paragraph(doc, f"Uyumlu gelir: {aligned_revenue:,.2f} EUR (%{align_rev_pct})")
                _add_paragraph(doc, f"Toplam CapEx: {total_capex:,.2f} EUR")
                _add_paragraph(doc, f"Uyumlu CapEx: {aligned_capex:,.2f} EUR (%{align_capex_pct})")
                _add_paragraph(doc, f"Toplam OpEx: {total_opex:,.2f} EUR")
                _add_paragraph(doc, f"Uyumlu OpEx: {aligned_opex:,.2f} EUR (%{align_opex_pct})")
                by_objective = tax.get('by_objective') or {}
                if by_objective:
                    _add_paragraph(doc, "Çevresel hedef bazında uyumlu gelir:", bold=True)
                    for code, amount in by_objective.items():
                        value = amount or 0
                        _add_paragraph(doc, f"- {code}: {value:,.2f} EUR")
                section_index += 1

            if esg_scores:
                _add_heading(doc, f"{section_index}. ESG Skorları")
                overall = esg_scores.get("overall", 0)
                e_score = esg_scores.get("E", 0)
                s_score = esg_scores.get("S", 0)
                g_score = esg_scores.get("G", 0)
                _add_paragraph(
                    doc,
                    f"Genel ESG Skoru: %{overall} | E %{e_score}, S %{s_score}, G %{g_score}",
                )
                details = esg_scores.get("details") or {}
                gri_det = details.get("gri") or {}
                tsrs_det = details.get("tsrs") or {}
                bonuses = details.get("bonuses") or {}
                table = doc.add_table(rows=1, cols=3)
                hdr = table.rows[0].cells
                hdr[0].text = "Kaynak"
                hdr[1].text = "Yanıtlanan"
                hdr[2].text = "Toplam"
                gri_ans = (
                    (gri_det.get("environmental", {}).get("answered", 0) or 0)
                    + (gri_det.get("social", {}).get("answered", 0) or 0)
                    + (gri_det.get("governance", {}).get("answered", 0) or 0)
                )
                gri_tot = (
                    (gri_det.get("environmental", {}).get("total", 0) or 0)
                    + (gri_det.get("social", {}).get("total", 0) or 0)
                    + (gri_det.get("governance", {}).get("total", 0) or 0)
                )
                tsrs_ans = (
                    (tsrs_det.get("environmental", {}).get("answered", 0) or 0)
                    + (tsrs_det.get("social", {}).get("answered", 0) or 0)
                    + (tsrs_det.get("governance", {}).get("answered", 0) or 0)
                )
                tsrs_tot = (
                    (tsrs_det.get("environmental", {}).get("total", 0) or 0)
                    + (tsrs_det.get("social", {}).get("total", 0) or 0)
                    + (tsrs_det.get("governance", {}).get("total", 0) or 0)
                )
                row = table.add_row().cells
                row[0].text = "GRI (E/S/G2)"
                row[1].text = str(gri_ans)
                row[2].text = str(gri_tot)
                row = table.add_row().cells
                row[0].text = "TSRS (E/S/G)"
                row[1].text = str(tsrs_ans)
                row[2].text = str(tsrs_tot)
                bonus_count = bonuses.get("sdg_answers", 0) or 0
                row = table.add_row().cells
                row[0].text = "Bonus Kaynaklar (SDG Yanıtları)"
                row[1].text = str(bonus_count)
                row[2].text = str(bonus_count)
                section_index += 1

            if ungc_result:
                _add_heading(doc, f"{section_index}. UN Global Compact Uyum")
                overall_ungc = ungc_result.get("overall_score", 0)
                _add_paragraph(doc, f"Toplam UNGC Skoru: %{overall_ungc}")
                principles = ungc_result.get("principles") or []
                if principles:
                    try:
                        from docx.shared import Inches
                        ungc_logo_path = os.path.join(self.base_dir, "static", "images", "ungc_logo.png")
                        if os.path.exists(ungc_logo_path):
                            p_logo = doc.add_paragraph()
                            r_logo = p_logo.add_run()
                            r_logo.add_picture(ungc_logo_path, width=Inches(1.5))
                            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
                    table = doc.add_table(rows=1, cols=4)
                    hdr = table.rows[0].cells
                    hdr[0].text = "İlke"
                    hdr[1].text = "Kategori"
                    hdr[2].text = "Durum"
                    hdr[3].text = "Skor %"
                    for pr in principles:
                        row = table.add_row().cells
                        row[0].text = str(pr.get("principle_id", ""))
                        row[1].text = str(pr.get("category", ""))
                        row[2].text = str(pr.get("status", ""))
                        row[3].text = str(pr.get("score", 0))
                    category_scores = ungc_result.get("category_scores") or {}
                    if category_scores:
                        _add_paragraph(doc)
                        _add_paragraph(doc, "Kategori Bazında Skorlar:")
                        for cat, score in category_scores.items():
                            _add_paragraph(doc, f"- {cat}: %{score}")
                    try:
                        from backend.modules.ungc.ungc_manager_enhanced import UNGCManagerEnhanced
                        enhanced_mgr = UNGCManagerEnhanced(db_path)
                        overall_data = enhanced_mgr.calculate_overall_score(company_id, reporting_period)
                        level_info = overall_data.get("level_info") or {}
                        level_name = level_info.get("level")
                        level_desc = level_info.get("description")
                        requirements = level_info.get("requirements") or []
                        if level_name or level_desc:
                            _add_paragraph(doc)
                            _add_paragraph(doc, "UNGC Seviye Sınıflandırması:")
                            if level_name:
                                _add_paragraph(doc, f"- Seviye: {level_name}")
                            if level_desc:
                                _add_paragraph(doc, f"- Açıklama: {level_desc}")
                            if requirements:
                                _add_paragraph(doc, "Öne çıkan gereklilikler:")
                                for req in requirements:
                                    _add_paragraph(doc, f"  - {req}")
                    except Exception:
                        pass
                section_index += 1



            try:
                from backend.modules.gri.gri_reporting import GRIReporting
                gri_reporting = GRIReporting(db_path=db_path)
                water_metrics = gri_reporting._get_gri303_water_metrics(company_id, reporting_period)
            except Exception:
                water_metrics = None

            if water_metrics:
                _add_heading(doc, f"{section_index}. GRI 303: Su ve Deşarjlar")
                table = doc.add_table(rows=1, cols=2)
                hdr = table.rows[0].cells
                hdr[0].text = "Metrik"
                hdr[1].text = "Değer"
                rows = [
                    ("Toplam Su Ayak İzi (m³)", str(water_metrics.get("total_water_footprint", 0))),
                    ("Mavi Su (m³)", str(water_metrics.get("total_blue_water", 0))),
                    ("Yeşil Su (m³)", str(water_metrics.get("total_green_water", 0))),
                    ("Gri Su (m³)", str(water_metrics.get("total_grey_water", 0))),
                    ("Ortalama Verimlilik Oranı", str(water_metrics.get("efficiency_metrics", {}).get("average_efficiency_ratio", 0))),
                    ("Ortalama Geri Dönüşüm Oranı", str(water_metrics.get("efficiency_metrics", {}).get("average_recycling_rate", 0))),
                ]
                for met, val in rows:
                    r = table.add_row().cells
                    r[0].text = met
                    r[1].text = val
                section_index += 1



            if gri_selected_rows:
                _add_heading(doc, f"{section_index}. GRI İçerik İndeksi – Seçilmiş Göstergeler")
                table = doc.add_table(rows=len(gri_selected_rows) + 1, cols=4)
                hdr = table.rows[0].cells
                hdr[0].text = "GRI Standart"
                hdr[1].text = "GRI Gösterge"
                hdr[2].text = "Başlık"
                hdr[3].text = "TSRS Eşleştirmeleri"
                for idx, row_data in enumerate(gri_selected_rows):
                    scode, icode, ititle, tsrs_text = row_data
                    cells = table.rows[idx + 1].cells
                    cells[0].text = str(scode)
                    cells[1].text = str(icode)
                    cells[2].text = str(ititle)
                    cells[3].text = str(tsrs_text)

            reports_dir = os.path.join(self.base_dir, "uploads", "reports")
            os.makedirs(reports_dir, exist_ok=True)

            safe_name = "".join(c if c.isalnum() or c in ["-", "_"] else "_" for c in report_name.strip()) or "unified_report"
            filename = f"unified_{safe_name}_{reporting_period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = os.path.join(reports_dir, filename)

            doc.save(file_path)
            return file_path
        except Exception as e:
            logging.error(f"UnifiedReportDocxGenerator Error: {e}")
            import traceback
            logging.error(traceback.format_exc())
            return None
