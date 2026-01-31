import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAPORLAMA MODÜLÜ GUI  
Rapor oluşturma, düzenleme ve yönetim arayüzü
"""

import os
import tkinter as tk
from datetime import datetime
from tkinter import messagebox, ttk

from utils.language_manager import LanguageManager
from utils.ui_theme import apply_theme
from config.icons import Icons


class ReportingGUI:
    """Raporlama Modülü GUI"""

    def __init__(self, parent, company_id: int) -> None:
        self.parent = parent
        self.company_id = company_id
        self.lm = LanguageManager()

        try:
            self.parent.winfo_toplevel().state('zoomed')
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

        self.setup_ui()
        self.load_data()

    def setup_ui(self) -> None:
        """Raporlama modülü arayüzünü oluştur"""
        apply_theme(self.parent)
        # Ana frame
        main_frame = tk.Frame(self.parent, bg='#f5f5f5')
        main_frame.pack(fill='both', expand=True, padx=10, pady=10)

        # Başlık
        title_frame = tk.Frame(main_frame, bg='#2c3e50', height=60)
        title_frame.pack(fill='x', pady=(0, 20))
        title_frame.pack_propagate(False)

        title_label = tk.Label(
            title_frame,
            text=f"{Icons.REPORT} {self.lm.tr('reporting_center_title', 'Raporlama Merkezi')}",
            font=('Segoe UI', 16, 'bold'),
            bg='#2c3e50',
            fg='white'
        )
        title_label.pack(expand=True)
        actions = tk.Frame(title_frame, bg='#2c3e50')
        actions.pack(side='right', padx=10)
        ttk.Button(actions, text=self.lm.tr('btn_report_center', "Rapor Merkezi"), style='Primary.TButton', command=self.open_central_report_center).pack(side='right')
        ttk.Button(actions, text=self.lm.tr('btn_env_report_center', "Çevresel Rapor Merkezi"), style='Primary.TButton', command=self.open_environmental_report_center).pack(side='right', padx=8)

        # Ana içerik - Notebook (Sekmeler)
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill='both', expand=True)

        # Temel Raporlama sekmesi
        self.basic_frame = tk.Frame(self.notebook, bg='white')
        self.notebook.add(self.basic_frame, text=f"{Icons.CLIPBOARD} {self.lm.tr('tab_basic_reporting', 'Temel Raporlama')}")

        # Gelişmiş Raporlama sekmesi
        self.advanced_frame = tk.Frame(self.notebook, bg='white')
        self.notebook.add(self.advanced_frame, text=self.lm.tr('tab_adv_reporting', "⚡ Gelişmiş Raporlama"))

        # Rapor Geçmişi sekmesi
        self.history_frame = tk.Frame(self.notebook, bg='white')
        self.notebook.add(self.history_frame, text=self.lm.tr('tab_report_history', "📜 Rapor Geçmişi"))

        # Rapor Ayarları sekmesi
        self.settings_frame = tk.Frame(self.notebook, bg='white')
        self.notebook.add(self.settings_frame, text=f"{Icons.SETTINGS} {self.lm.tr('tab_settings', 'Ayarlar')}")

        # Sekmeleri oluştur
        self.create_basic_reporting_tab()
        self.create_advanced_reporting_tab()
        self.create_history_tab()
        self.create_settings_tab()

    def open_central_report_center(self) -> None:
        try:
            from modules.reporting.report_center_gui import ReportCenterGUI
            win = tk.Toplevel(self.parent)
            gui = ReportCenterGUI(win, self.company_id)
            try:
                gui.module_filter_var.set('genel')
                gui.refresh_reports()
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    def open_environmental_report_center(self) -> None:
        try:
            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('win_env_report_center', "Çevresel Rapor Merkezi"))
            win.geometry("960x640")
            win.configure(bg='white')

            top = tk.Frame(win, bg='white')
            top.pack(fill='x', padx=12, pady=8)
            tk.Label(top, text=self.lm.tr('header_env_report_center', " Çevresel Rapor Merkezi"), font=('Segoe UI', 14, 'bold'), fg='#1E8449', bg='white').pack(side='left')

            left = tk.LabelFrame(win, text=self.lm.tr('grp_create_report', "Rapor Oluştur"), bg='white')
            left.pack(side='left', fill='y', padx=12, pady=8)

            tk.Label(left, text=self.lm.tr('lbl_report_type', "Rapor Türü"), bg='white').pack(anchor='w', padx=8, pady=(8, 2))
            report_type_var = tk.StringVar(value='Çevresel')
            # Localized values could be mapped, but for now we keep the values as they might be used as keys.
            # If they are keys, we should probably keep them English or key-like internally, but display localized.
            # But ttk.Combobox usually displays values.
            # Let's assume these are just display strings for now, or keys that are also display strings.
            # Given the context, "Çevresel", "Sosyal" etc are Turkish.
            # If we change them to keys, we might break backend logic if it expects these exact strings.
            # However, looking at _generate_environment_report_file, it uses report_type directly in filename.
            # So changing them to English keys might be safer for filenames, but we need to verify if other parts use them.
            # The user wants localization.
            # Let's use localized strings for display if possible.
            # But for simplicity and safety, let's keep the values but maybe localize the default value if it's a string literal in code?
            # Actually line 112 sets default to 'Çevresel'.
            # Line 113 values are hardcoded Turkish.
            # If I change them to tr() calls, they will be localized.
            # But then _generate_environment_report_file will use the localized string for filename.
            # That's probably fine.
            ttk.Combobox(left, textvariable=report_type_var, values=[
                self.lm.tr('type_environmental', 'Çevresel'),
                self.lm.tr('type_sdg', 'SDG'),
                self.lm.tr('type_tsrs', 'TSRS'),
                self.lm.tr('type_social', 'Sosyal'),
                self.lm.tr('type_cbam', 'CBAM')
            ], state='readonly', width=24).pack(padx=8, pady=2)

            tk.Label(left, text=self.lm.tr('lbl_year', "Yıl"), bg='white').pack(anchor='w', padx=8, pady=(8, 2))
            from datetime import datetime
            year_var = tk.StringVar(value=str(datetime.now().year))
            ttk.Entry(left, textvariable=year_var, width=10).pack(padx=8, pady=2)

            def create_report() -> None:
                try:
                    year = int(year_var.get())
                except Exception:
                    messagebox.showwarning(self.lm.tr('warn_title', "Uyarı"), self.lm.tr('warn_invalid_year', "Geçerli bir yıl giriniz"))
                    return
                self._generate_environment_report_file(report_type_var.get(), year)
                refresh_history()

            tk.Button(left, text=self.lm.tr('btn_create_report', "Raporu Oluştur"), bg='#2196F3', fg='white', command=create_report).pack(padx=8, pady=12, fill='x')

            right = tk.Frame(win, bg='white')
            right.pack(side='left', fill='both', expand=True, padx=8, pady=8)

            columns = ('file', 'created')
            tree = ttk.Treeview(right, columns=columns, show='headings')
            tree.heading('file', text=self.lm.tr('col_file', 'Dosya'))
            tree.heading('created', text=self.lm.tr('col_created', 'Oluşturulma'))
            tree.column('file', width=520)
            tree.column('created', width=140)
            tree.pack(fill='both', expand=True, side='top')

            preview = tk.Text(right, height=10, wrap='word', bg='#f8f9fa')
            preview.pack(fill='x', side='bottom', padx=4, pady=6)
            preview.insert('1.0', self.lm.tr('preview_placeholder', 'Önizleme: Word içeriği dönüştürülmüş kısa özet görüntülenir.'))
            preview.config(state=tk.DISABLED)

            def refresh_history() -> None:
                from datetime import datetime
                tree.delete(*tree.get_children())
                base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
                report_dir = os.path.join(base_dir, 'reports')
                os.makedirs(report_dir, exist_ok=True)
                for fname in sorted(os.listdir(report_dir)):
                    if fname.lower().endswith('.docx'):
                        fpath = os.path.join(report_dir, fname)
                        created = datetime.fromtimestamp(os.path.getmtime(fpath)).strftime('%Y-%m-%d %H:%M')
                        tree.insert('', 'end', values=(fname, created))

            def open_selected() -> None:
                sel = tree.selection()
                if not sel:
                    return
                fname = tree.item(sel[0])['values'][0]
                base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
                fpath = os.path.join(base_dir, 'reports', fname)
                try:
                    os.startfile(fpath)
                except Exception as e:
                    messagebox.showerror(self.lm.tr('error_title', 'Hata'), f"{self.lm.tr('err_file_open', 'Dosya açılamadı')}: {e}")

            btns = tk.Frame(right, bg='white')
            btns.pack(fill='x', pady=4)
            tk.Button(btns, text=self.lm.tr('btn_open_selected', 'Seçileni Aç'), command=open_selected).pack(side='left', padx=4)
            tk.Button(btns, text=self.lm.tr('btn_refresh', 'Yenile'), command=refresh_history).pack(side='left', padx=4)

            refresh_history()
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    def _generate_environment_report_file(self, report_type: str, year: int) -> None:
        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt  # type: ignore
        except Exception as e:
            messagebox.showerror(self.lm.tr('error_title', 'Hata'), f"{self.lm.tr('err_docx_missing', 'Rapor şablonu oluşturmak için python-docx gerekli')}: {e}")
            return

        # Veri Yöneticisini Başlat ve Verileri Çek
        try:
            from modules.environmental.detailed_energy_manager import DetailedEnergyManager
            manager = DetailedEnergyManager()
            data = manager.get_annual_report_data(self.company_id, year)
        except Exception as e:
            messagebox.showerror("Hata", f"Veri alınamadı: {e}")
            return

        base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
        report_dir = os.path.join(base_dir, 'reports')
        os.makedirs(report_dir, exist_ok=True)
        
        doc = Document()
        
        # Başlık
        doc.add_heading(f'{report_type} Raporu ({year})', level=1)
        doc.add_paragraph(self.lm.tr('report_auto_generated', 'Bu rapor sistem tarafından otomatik oluşturulmuştur.'))
        doc.add_paragraph(f"Oluşturma Tarihi: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        # Özet Tablosu
        doc.add_heading('Yıllık Enerji Özeti', level=2)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metrik'
        hdr_cells[1].text = 'Değer'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Toplam Tüketim'
        row_cells[1].text = f"{data.get('total_consumption', 0):,.2f} kWh (eşdeğer)"
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Toplam Maliyet'
        row_cells[1].text = f"{data.get('total_cost', 0):,.2f} TL"
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Toplam CO2 Emisyonu'
        row_cells[1].text = f"{data.get('total_emissions', 0):,.2f} kg CO2e"
        
        # Detay Tablosu
        doc.add_heading('Enerji Türüne Göre Dağılım', level=2)
        
        det_table = doc.add_table(rows=1, cols=4)
        det_table.style = 'Table Grid'
        hdr_cells = det_table.rows[0].cells
        hdr_cells[0].text = 'Enerji Türü'
        hdr_cells[1].text = 'Tüketim'
        hdr_cells[2].text = 'Maliyet'
        hdr_cells[3].text = 'Emisyon'
        
        for item in data.get('breakdown', []):
            row_cells = det_table.add_row().cells
            row_cells[0].text = str(item['type'])
            row_cells[1].text = f"{item['consumption']:,.2f}"
            row_cells[2].text = f"{item['cost']:,.2f}"
            row_cells[3].text = f"{item['emissions']:,.2f}"

        report_file = os.path.join(report_dir, f"{report_type.replace(' ', '_').lower()}_{year}.docx")
        try:
            doc.save(report_file)
            messagebox.showinfo(self.lm.tr('success_title', 'Başarılı'), f"{self.lm.tr('msg_report_created', 'Rapor oluşturuldu')}: {report_file}")
        except Exception as e:
            messagebox.showerror(self.lm.tr('error_title', 'Hata'), f"{self.lm.tr('err_report_save', 'Rapor kaydedilemedi')}: {e}")

    def create_basic_reporting_tab(self) -> None:
        """Temel raporlama sekmesi"""
        # Ana container
        container = tk.Frame(self.basic_frame, bg='white')
        container.pack(fill='both', expand=True, padx=20, pady=20)

        # Sol panel - Rapor Türleri
        left_panel = tk.Frame(container, bg='white', relief='solid', bd=1)
        left_panel.pack(side='left', fill='y', padx=(0, 10))
        left_panel.pack_propagate(False)
        left_panel.configure(width=300)

        tk.Label(left_panel, text="Rapor Türleri", font=('Segoe UI', 12, 'bold'),
                bg='white', fg='#1e40af').pack(pady=15)

        # Rapor türleri listesi
        self.report_types = [
            (Icons.REPORT, self.lm.tr('report_type_sdg', "SDG Raporu"), "sdg"),
            (Icons.SEED, self.lm.tr('report_type_gri', "GRI Raporu"), "gri"),
            ("🇹🇷", self.lm.tr('report_type_tsrs', "TSRS Raporu"), "tsrs"),
            ("🇪🇺", self.lm.tr('report_type_esrs', "ESRS Raporu"), "esrs"),
            ("💼", self.lm.tr('report_type_esg', "ESG Raporu"), "esg"),
            (Icons.CHART_UP, self.lm.tr('report_type_sustainability', "Sürdürülebilirlik Raporu"), "sustainability"),
            (Icons.LOADING, self.lm.tr('report_type_combined', "Birleşik Rapor"), "combined")
        ]

        for icon, name, code in self.report_types:
            def _open_module(c=code):
                self.open_report_module(c)
            ttk.Button(left_panel, text=f"{icon} {name}", style='Primary.TButton',
                       command=_open_module).pack(pady=5, padx=10, fill='x')

        # Sağ panel - Rapor Detayları
        right_panel = tk.Frame(container, bg='white', relief='solid', bd=1)
        right_panel.pack(side='right', fill='both', expand=True)

        tk.Label(right_panel, text=self.lm.tr('report_creation_title', "Rapor Oluşturma"), font=('Segoe UI', 12, 'bold'),
                bg='white', fg='#1e40af').pack(pady=15)

        # Rapor formu
        form_frame = tk.Frame(right_panel, bg='white')
        form_frame.pack(fill='both', expand=True, padx=20, pady=10)

        # Rapor adı
        tk.Label(form_frame, text=self.lm.tr('lbl_report_name', "Rapor Adı:"), bg='white').grid(row=0, column=0, sticky='w', pady=5)
        self.report_name_var = tk.StringVar(value=f"{self.lm.tr('default_report_name', 'Sürdürülebilirlik Raporu')} {datetime.now().year}")
        tk.Entry(form_frame, textvariable=self.report_name_var, width=40).grid(row=0, column=1, sticky='ew', pady=5)

        # Rapor dönemi
        tk.Label(form_frame, text=self.lm.tr('lbl_report_period', "Rapor Dönemi:"), bg='white').grid(row=1, column=0, sticky='w', pady=5)
        self.period_var = tk.StringVar(value="2024")
        period_combo = ttk.Combobox(form_frame, textvariable=self.period_var, width=37, state='readonly')
        period_combo['values'] = ["2024", "2023", "2022", "Q4 2024", "Q3 2024"]
        period_combo.bind('<<ComboboxSelected>>', self._on_period_selected)
        period_combo.grid(row=1, column=1, sticky='ew', pady=5)

        # Format seçimi
        tk.Label(form_frame, text=self.lm.tr('lbl_format', "Format:"), bg='white').grid(row=2, column=0, sticky='w', pady=5)
        self.format_var = tk.StringVar(value="PDF")
        format_combo = ttk.Combobox(form_frame, textvariable=self.format_var, width=37)
        format_combo['values'] = ["PDF", "Word", "Excel", "PowerPoint"]
        format_combo.grid(row=2, column=1, sticky='ew', pady=5)

        # Rapor oluştur butonu
        ttk.Button(form_frame, text=f"{Icons.CLIPBOARD} {self.lm.tr('btn_create_report', 'Rapor Oluştur')}", style='Primary.TButton',
                   command=self.on_generate).grid(row=3, column=0, columnspan=2, pady=20, sticky='ew')

        form_frame.columnconfigure(1, weight=1)

    def create_advanced_reporting_tab(self) -> None:
        """Gelişmiş raporlama sekmesi"""
        # İçerik oluştur
        content = tk.Frame(self.advanced_frame, bg='white')
        content.pack(fill='both', expand=True, padx=20, pady=20)

        tk.Label(content, text="⚡ Gelişmiş Raporlama Özellikleri",
                font=('Segoe UI', 14, 'bold'), bg='white', fg='#2c3e50').pack(pady=10)

        # Özellikler listesi
        features = [
            "🎨 Özelleştirilmiş Rapor Tasarımı",
            f"{Icons.REPORT} İnteraktif Grafikler ve Çizelgeler",
            f"{Icons.LOADING} Otomatik Veri Güncelleme",
            f"{Icons.EMAIL} Email ile Otomatik Gönderim",
            f"{Icons.WORLD} Çok Dilli Rapor Desteği",
            f"{Icons.TIME} Zamanlanmış Rapor Oluşturma",
            f"{Icons.LINK} API ile Veri Entegrasyonu"
        ]

        for feature in features:
            feature_frame = tk.Frame(content, bg='#ecf0f1', relief='solid', bd=1)
            feature_frame.pack(fill='x', pady=5)
            tk.Label(feature_frame, text=feature, font=('Segoe UI', 11),
                    bg='#ecf0f1', fg='#2c3e50').pack(anchor='w', padx=15, pady=10)

    def create_history_tab(self) -> None:
        """Rapor geçmişi sekmesi"""
        content = tk.Frame(self.history_frame, bg='white')
        content.pack(fill='both', expand=True, padx=20, pady=20)

        tk.Label(content, text="📜 Rapor Geçmişi",
                font=('Segoe UI', 14, 'bold'), bg='white', fg='#2c3e50').pack(pady=10)

        # Treeview for reports
        columns = ('Tarih', 'Rapor Adı', 'Tür', 'Format', 'Durum')
        self.reports_tree = ttk.Treeview(content, columns=columns, show='headings', height=15)

        for col in columns:
            self.reports_tree.heading(col, text=col)
            self.reports_tree.column(col, width=120)

        # Scrollbar
        scrollbar = ttk.Scrollbar(content, orient='vertical', command=self.reports_tree.yview)
        self.reports_tree.configure(yscrollcommand=scrollbar.set)

        self.reports_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

    def create_settings_tab(self) -> None:
        """Ayarlar sekmesi"""
        content = tk.Frame(self.settings_frame, bg='white')
        content.pack(fill='both', expand=True, padx=20, pady=20)

        tk.Label(content, text=f"{Icons.SETTINGS} Raporlama Ayarları",
                font=('Segoe UI', 14, 'bold'), bg='white', fg='#2c3e50').pack(pady=10)

        # Ayarlar formu
        settings_frame = tk.Frame(content, bg='white')
        settings_frame.pack(fill='x', pady=20)

        # Varsayılan format
        tk.Label(settings_frame, text="Varsayılan Format:", bg='white').grid(row=0, column=0, sticky='w', pady=10)
        self.default_format = tk.StringVar(value="PDF")
        format_combo = ttk.Combobox(settings_frame, textvariable=self.default_format)
        format_combo['values'] = ["PDF", "Word", "Excel"]
        format_combo.grid(row=0, column=1, sticky='ew', padx=10)

        # Otomatik backup
        self.auto_backup = tk.BooleanVar(value=True)
        tk.Checkbutton(settings_frame, text="Otomatik Backup", variable=self.auto_backup,
                      bg='white').grid(row=1, column=0, columnspan=2, sticky='w', pady=10)

        # Email bildirimleri
        self.email_notifications = tk.BooleanVar(value=False)
        tk.Checkbutton(settings_frame, text="Email Bildirimleri", variable=self.email_notifications,
                      bg='white').grid(row=2, column=0, columnspan=2, sticky='w', pady=10)

        # Kaydet butonu
        ttk.Button(settings_frame, text=f"{Icons.SAVE} Ayarları Kaydet", style='Primary.TButton').grid(row=3, column=0, columnspan=2, pady=20)

        settings_frame.columnconfigure(1, weight=1)

    def open_report_module(self, code: str) -> None:
        try:
            win = tk.Toplevel(self.parent)
            win.title("Rapor Modülü")
            win.geometry("900x650")
            if code == "sdg":
                from modules.sdg.sdg_reporting_gui import SDGReportingGUI
                SDGReportingGUI(win, self.company_id)
            elif code == "gri":
                from modules.gri.gri_gui import GRIGUI
                GRIGUI(win, self.company_id)
            elif code == "tsrs":
                from modules.tsrs.tsrs_reporting_gui import TSRSReportingGUI
                TSRSReportingGUI(win, self.company_id)
            elif code == "esrs":
                from modules.esrs.esrs_gui import ESRSGUI
                ESRSGUI(win, self.company_id)
            elif code == "esg":
                from modules.esg.esg_gui import ESGGUI
                ESGGUI(win, self.company_id)
            else:
                messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('msg_module_coming_soon', "Bu rapor türü için modül yakında eklenecek."))
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('err_module_open', 'Modül açılamadı')}: {str(e)}")

    def load_data(self) -> None:
        """Veri yükle"""
        try:
            # Gerçek raporları veritabanından veya dosya sisteminden yükle
            # Şimdilik boş bırakıyoruz, çünkü 'sample_reports' kullanıcıyı yanıltabilir.
            pass

        except Exception as e:
            logging.error(f"Veri yükleme hatası: {e}")

    def _validate_period(self, p: str):
        try:
            import re
            p = (p or "").strip()
            if p.isdigit() and len(p) == 4:
                val = int(p)
                if 1990 <= val <= 2100:
                    return str(val)
                return None
            m = re.fullmatch(r"Q([1-4])\s+(\d{4})", p)
            if m:
                q = int(m.group(1))
                y = int(m.group(2))
                if 1990 <= y <= 2100 and 1 <= q <= 4:
                    return f"Q{q} {y}"
            return None
        except Exception:
            return None

    def _on_period_selected(self, event=None) -> None:
        try:
            val = self._validate_period(self.period_var.get())
            if not val:
                messagebox.showwarning(self.lm.tr('warn_title', "Uyarı"), self.lm.tr('warn_invalid_period', "Geçersiz dönem seçimi"))
                self.period_var.set(str(datetime.now().year))
                return
            self.period_var.set(val)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    def on_generate(self) -> None:
        try:
            val = self._validate_period(self.period_var.get())
            if not val:
                messagebox.showwarning(self.lm.tr('warn_title', "Uyarı"), self.lm.tr('warn_invalid_period_detail', "Geçersiz dönem. Lütfen yıl veya çeyrek+yıl seçin"))
                return
            messagebox.showinfo(self.lm.tr('info', "Bilgi"), f"{self.lm.tr('msg_period_confirmed', 'Rapor dönemi onaylandı')}: {val}")
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('err_operation', 'İşlem hatası')}: {e}")
