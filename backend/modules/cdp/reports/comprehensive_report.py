#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CDP Kapsamlı Raporu
Tüm CDP modüllerini içeren kapsamlı rapor
"""

import logging
import os
import sqlite3
from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from .climate_change_report import ClimateChangeReport
from .forests_report import ForestsReport
from .water_security_report import WaterSecurityReport


def _add_turkish_paragraph(doc, text: str, style: str | None = None, font_name: str = 'Calibri', font_size: int = 11):
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        from docx.oxml.ns import qn
        r = run._element
        r.rPr.rFonts.set(qn('w:ascii'), font_name)
        r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        r.rPr.rFonts.set(qn('w:cs'), font_name)
    return para

def _add_turkish_heading(doc, text: str, level: int = 1, font_name: str = 'Calibri'):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        from docx.oxml.ns import qn
        r = run._element
        r.rPr.rFonts.set(qn('w:ascii'), font_name)
        r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
        r.rPr.rFonts.set(qn('w:cs'), font_name)
    return heading


class ComprehensiveReport:
    """CDP Kapsamlı Raporu oluşturucu"""

    def __init__(self, use_ai: bool = True, api_key: Optional[str] = None):
        self.climate_report = ClimateChangeReport(use_ai=use_ai, api_key=api_key)
        self.water_report = WaterSecurityReport(use_ai=use_ai, api_key=api_key)
        self.forests_report = ForestsReport(use_ai=use_ai, api_key=api_key)

    def generate_report(self, company_id: int, year: int, output_path: str) -> bool:
        """Kapsamlı CDP raporu oluştur"""
        try:
            logging.info(f"[CDP Comprehensive] Kapsamlı rapor oluşturuluyor: {year}")

            # Tüm raporları ayrı ayrı oluştur
            base_dir = os.path.dirname(output_path)

            # Climate
            climate_path = os.path.join(base_dir, f'CDP_Climate_Change_{year}.docx')
            self.climate_report.generate_report(company_id, year, climate_path)

            # Water
            water_path = os.path.join(base_dir, f'CDP_Water_Security_{year}.docx')
            self.water_report.generate_report(company_id, year, water_path)

            # Forests
            forests_path = os.path.join(base_dir, f'CDP_Forests_{year}.docx')
            self.forests_report.generate_report(company_id, year, forests_path)

            # Ana kapsamlı rapor
            doc = self._create_comprehensive_doc(company_id, year)
            doc.save(output_path)

            logging.info(f"[CDP Comprehensive] Rapor kaydedildi: {output_path}")
            logging.info("[CDP Comprehensive] Alt raporlar:")
            logging.info(f"  - {climate_path}")
            logging.info(f"  - {water_path}")
            logging.info(f"  - {forests_path}")

            return True

        except Exception as e:
            logging.error(f"[HATA] CDP Comprehensive raporu oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _add_logo(self, doc, company_id: int):
        """Raporun başına şirket logosunu ekle"""
        try:
            # 1. Veritabanından logo yolunu al
            conn = sqlite3.connect(self.climate_report.collector.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT logo_path FROM company_profiles WHERE company_id = ?", (company_id,))
            result = cursor.fetchone()
            conn.close()
            
            logo_path = None
            if result and result[0] and os.path.exists(result[0]):
                logo_path = result[0]
            else:
                # 2. Varsayılan yolda ara
                data_dir = os.path.dirname(self.climate_report.collector.db_path)
                possible_path = os.path.join(data_dir, "company_logos", f"company_{company_id}_logo.png")
                if os.path.exists(possible_path):
                    logo_path = possible_path
                else:
                    # jpg dene
                    possible_path = possible_path.replace(".png", ".jpg")
                    if os.path.exists(possible_path):
                        logo_path = possible_path

            if logo_path:
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, height=Cm(2.5))
        except Exception as e:
            logging.warning(f"Logo eklenirken hata: {e}")

    def _create_comprehensive_doc(self, company_id: int, year: int):
        """Kapsamlı rapor ana dokümanı"""
        doc = Document()
        
        # Logo ekle
        self._add_logo(doc, company_id)

        # Kapak sayfası
        title = _add_turkish_heading(doc, 'CDP KAPSAMLI SÜRDÜRÜLEBİLİRLİK RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        company_info = self.climate_report.collector.collect_company_info(company_id)
        company_name = company_info.get('name', 'Şirket')

        subtitle = _add_turkish_paragraph(doc, f'{company_name}\n{year}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.bold = True

        _add_turkish_paragraph(doc, '')
        _add_turkish_paragraph(doc, '')

        # İçindekiler
        _add_turkish_heading(doc, 'İÇİNDEKİLER', 1)
        toc_items = [
            '1. Giriş ve Genel Bakış',
            '2. CDP Climate Change Raporu',
            '3. CDP Water Security Raporu',
            '4. CDP Forests Raporu',
            '5. Entegre Performans Değerlendirmesi',
            '6. Sonuç ve Öneriler'
        ]

        for item in toc_items:
            _add_turkish_paragraph(doc, item, style='List Number')

        # 1. Giriş
        doc.add_page_break()
        _add_turkish_heading(doc, '1. GİRİŞ VE GENEL BAKIŞ', 1)

        intro_text = f"""
Bu rapor, {company_name}'in {year} yılı sürdürülebilirlik performansını 
CDP (Carbon Disclosure Project) standartları çerçevesinde kapsamlı bir şekilde sunmaktadır.

Rapor, aşağıdaki üç ana CDP modülünü kapsamaktadır:
• CDP Climate Change (İklim Değişikliği)
• CDP Water Security (Su Güvenliği)
• CDP Forests (Ormanlar)

Her modül için detaylı analizler, performans değerlendirmeleri ve öneriler 
ayrı raporlarda sunulmaktadır. Bu doküman, genel bir özet ve entegre 
değerlendirme sağlamaktadır.
        """.strip()

        _add_turkish_paragraph(doc, intro_text)

        # Alt raporlara referans
        doc.add_page_break()
        _add_turkish_heading(doc, 'DETAYLI RAPORLAR', 1)

        _add_turkish_paragraph(doc, """
Bu kapsamlı raporla birlikte aşağıdaki detaylı raporlar oluşturulmuştur:

1. CDP_Climate_Change_{year}.docx
   - Sera gazı emisyonları
   - Enerji tüketimi
   - İklim riskleri ve fırsatları
   - Azaltım stratejileri

2. CDP_Water_Security_{year}.docx
   - Su tüketimi ve çekimi
   - Su geri dönüşümü
   - Su riskleri
   - Su yönetimi stratejileri

3. CDP_Forests_{year}.docx
   - Orman ürünleri kullanımı
   - Ormansızlaşma riskleri
   - Sertifikasyon durumu
   - Tedarik zinciri yönetimi

Detaylı analiz ve veriler için lütfen ilgili raporlara bakınız.
        """.format(year=year).strip())

        # Altbilgi
        doc.add_page_break()
        footer = _add_turkish_paragraph(doc, '')
        footer.add_run(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y")}\n').font.size = Pt(9)
        footer.add_run('© Sustainage CDP Modülü - Kapsamlı Rapor').font.size = Pt(9)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return doc

