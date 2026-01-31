import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CDP Water Security Raporu
Su güvenliği raporlaması için PDF/DOCX oluşturma
"""

import logging
import os
import sqlite3
from datetime import datetime
from typing import Any, Dict, Optional

import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Cm

matplotlib.use('Agg')

from ..ai_analyzer import CDPAIAnalyzer
from ..cdp_data_collector import CDPDataCollector


def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    para = _add_turkish_paragraph(doc, text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = _add_turkish_heading(doc, text, level=level)
    for run in heading.runs:
        run.font.name = font_name
    return heading


class WaterSecurityReport:
    """CDP Water Security Raporu oluşturucu"""

    def __init__(self, use_ai: bool = True, api_key: Optional[str] = None):
        self.collector = CDPDataCollector()
        self.analyzer = CDPAIAnalyzer(use_api=False if not api_key else True, api_key=api_key)
        self.use_ai = use_ai

    def generate_report(self, company_id: int, year: int, output_path: str) -> bool:
        """CDP Water Security raporu oluştur"""
        try:
            logging.info(f"[CDP Water Security] Rapor oluşturuluyor: {year}")

            # Veri toplama
            data = self._collect_all_data(company_id, year)

            # AI Analiz
            if self.use_ai:
                analysis = self.analyzer.analyze_water_security(data)
            else:
                analysis = {}

            # Grafikler
            charts_dir = os.path.join(os.path.dirname(output_path), 'charts')
            os.makedirs(charts_dir, exist_ok=True)
            chart_files = self._create_charts(data, charts_dir)

            # DOCX raporu
            doc = self._create_docx_report(data, analysis, chart_files)
            doc.save(output_path)

            logging.info(f"[CDP Water Security] Rapor kaydedildi: {output_path}")
            return True

        except Exception as e:
            logging.error(f"[HATA] CDP Water Security raporu oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _collect_all_data(self, company_id: int, year: int) -> Dict[str, Any]:
        """Tüm gerekli verileri topla"""
        data = {}
        data['company_id'] = company_id
        data['company_info'] = self.collector.collect_company_info(company_id)
        water_data = self.collector.collect_water_data(company_id, year)
        data['water_consumption'] = water_data.get('water_consumption', {})
        data['year'] = year
        return data

    def _add_logo(self, doc, company_id: int):
        """Raporun başına şirket logosunu ekle"""
        try:
            # 1. Veritabanından logo yolunu al
            conn = sqlite3.connect(self.collector.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT logo_path FROM company_profiles WHERE company_id = ?", (company_id,))
            result = cursor.fetchone()
            conn.close()
            
            logo_path = None
            if result and result[0] and os.path.exists(result[0]):
                logo_path = result[0]
            else:
                # 2. Varsayılan yolda ara
                data_dir = os.path.dirname(self.collector.db_path)
                possible_path = os.path.join(data_dir, "company_logos", f"company_{company_id}_logo.png")
                if os.path.exists(possible_path):
                    logo_path = possible_path
                else:
                    # jpg dene
                    possible_path = possible_path.replace(".png", ".jpg")
                    if os.path.exists(possible_path):
                        logo_path = possible_path

            if logo_path:
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, height=Cm(2.5))
        except Exception as e:
            logging.warning(f"Logo eklenirken hata: {e}")

    def _create_charts(self, data: Dict[str, Any], output_dir: str) -> Dict[str, str]:
        """Su kullanımı grafiklerini oluştur"""
        chart_files = {}

        water = data.get('water_consumption', {})
        if water:
            fig, ax = plt.subplots(figsize=(10, 6))

            categories = ['Çekilen Su', 'Tüketilen Su', 'Deşarj Edilen', 'Geri Dönüştürülen']
            values = [
                water.get('withdrawn', 0),
                water.get('consumed', 0),
                water.get('discharged', 0),
                water.get('recycled', 0)
            ]

            colors = ['#3498DB', '#E74C3C', '#95A5A6', '#2ECC71']
            bars = ax.bar(categories, values, color=colors, width=0.6)

            ax.set_ylabel('Su Miktarı (m³)', fontsize=12)
            ax.set_title('Su Kullanımı Dağılımı', fontsize=14, fontweight='bold')
            ax.grid(axis='y', alpha=0.3)

            # Değerleri göster
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{height:,.0f}',
                       ha='center', va='bottom', fontsize=9)

            plt.xticks(rotation=15, ha='right')

            chart_path = os.path.join(output_dir, 'water_usage.png')
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            chart_files['water_usage'] = chart_path

        return chart_files

    def _create_docx_report(self, data: Dict, analysis: Dict, charts: Dict) -> Document:
        """DOCX raporu oluştur"""
        doc = Document()
        
        # Logo ekle
        if 'company_id' in data:
            self._add_logo(doc, data['company_id'])

        # Başlık
        title = _add_turkish_heading(doc, 'CDP WATER SECURITY RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        company_name = data['company_info'].get('name', 'Şirket')
        year = data.get('year', datetime.now().year)

        subtitle = _add_turkish_paragraph(doc, f'{company_name} - {year}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)

        _add_turkish_paragraph(doc, )

        # 1. Özet
        _add_turkish_heading(doc, '1. YÖNETİCİ ÖZETİ', 1)
        if analysis.get('summary'):
            _add_turkish_paragraph(doc, analysis['summary'])

        # 2. Su Kullanımı
        doc.add_page_break()
        _add_turkish_heading(doc, '2. SU KULLANIMI', 1)

        water = data.get('water_consumption', {})
        water_table = doc.add_table(rows=1, cols=2)
        water_table.style = 'Light Grid Accent 1'

        rows = [
            ('Çekilen Su', f"{water.get('withdrawn', 0):,.0f} m³"),
            ('Tüketilen Su', f"{water.get('consumed', 0):,.0f} m³"),
            ('Deşarj Edilen Su', f"{water.get('discharged', 0):,.0f} m³"),
            ('Geri Dönüştürülen Su', f"{water.get('recycled', 0):,.0f} m³")
        ]

        for label, value in rows:
            r = water_table.add_row().cells
            r[0].text = label
            r[1].text = value

        # Grafik
        if 'water_usage' in charts:
            _add_turkish_paragraph(doc, )
            doc.add_picture(charts['water_usage'], width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 3. Performans
        if analysis.get('performance'):
            doc.add_page_break()
            _add_turkish_heading(doc, '3. PERFORMANS DEĞERLENDİRMESİ', 1)
            _add_turkish_paragraph(doc, analysis['performance'])

        # 4. Öneriler
        if analysis.get('recommendations'):
            doc.add_page_break()
            _add_turkish_heading(doc, '4. ÖNERİLER', 1)
            _add_turkish_paragraph(doc, analysis['recommendations'])

        # Altbilgi
        doc.add_page_break()
        footer = _add_turkish_paragraph(doc, )
        footer.add_run(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y")}\n').font.size = Pt(9)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return doc

