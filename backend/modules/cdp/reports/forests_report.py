import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CDP Forests Raporu
Orman ürünleri raporlaması için PDF/DOCX oluşturma
"""

import os
import sqlite3
from datetime import datetime
from typing import Any, Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from ..ai_analyzer import CDPAIAnalyzer
from ..cdp_data_collector import CDPDataCollector


def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    para = _add_turkish_paragraph(doc, text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = _add_turkish_heading(doc, text, level=level)
    for run in heading.runs:
        run.font.name = font_name
    return heading


class ForestsReport:
    """CDP Forests Raporu oluşturucu"""

    def __init__(self, use_ai: bool = True, api_key: Optional[str] = None):
        self.collector = CDPDataCollector()
        self.analyzer = CDPAIAnalyzer(use_api=False if not api_key else True, api_key=api_key)
        self.use_ai = use_ai

    def generate_report(self, company_id: int, year: int, output_path: str) -> bool:
        """CDP Forests raporu oluştur"""
        try:
            logging.info(f"[CDP Forests] Rapor oluşturuluyor: {year}")

            # Veri toplama
            data = self._collect_all_data(company_id, year)

            # AI Analiz
            if self.use_ai:
                analysis = self.analyzer.analyze_forests(data)
            else:
                analysis = {}

            # DOCX raporu
            doc = self._create_docx_report(data, analysis)
            doc.save(output_path)

            logging.info(f"[CDP Forests] Rapor kaydedildi: {output_path}")
            return True

        except Exception as e:
            logging.error(f"[HATA] CDP Forests raporu oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _collect_all_data(self, company_id: int, year: int) -> Dict[str, Any]:
        """Tüm gerekli verileri topla"""
        data = {}
        data['company_id'] = company_id
        data['company_info'] = self.collector.collect_company_info(company_id)
        forests_data = self.collector.collect_forests_data(company_id, year)
        data['cdp_responses'] = forests_data.get('cdp_responses', [])
        data['year'] = year
        return data

    def _add_logo(self, doc, company_id: int):
        """Raporun başına şirket logosunu ekle"""
        try:
            # 1. Veritabanından logo yolunu al
            conn = sqlite3.connect(self.collector.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT logo_path FROM company_profiles WHERE company_id = ?", (company_id,))
            result = cursor.fetchone()
            conn.close()
            
            logo_path = None
            if result and result[0] and os.path.exists(result[0]):
                logo_path = result[0]
            else:
                # 2. Varsayılan yolda ara
                data_dir = os.path.dirname(self.collector.db_path)
                possible_path = os.path.join(data_dir, "company_logos", f"company_{company_id}_logo.png")
                if os.path.exists(possible_path):
                    logo_path = possible_path
                else:
                    # jpg dene
                    possible_path = possible_path.replace(".png", ".jpg")
                    if os.path.exists(possible_path):
                        logo_path = possible_path

            if logo_path:
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, height=Cm(2.5))
        except Exception as e:
            logging.warning(f"Logo eklenirken hata: {e}")

    def _create_docx_report(self, data: Dict, analysis: Dict) -> Document:
        """DOCX raporu oluştur"""
        doc = Document()
        
        # Logo ekle
        if 'company_id' in data:
            self._add_logo(doc, data['company_id'])

        # Başlık
        title = _add_turkish_heading(doc, 'CDP FORESTS RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        company_name = data['company_info'].get('name', 'Şirket')
        year = data.get('year', datetime.now().year)

        subtitle = _add_turkish_paragraph(doc, f'{company_name} - {year}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)

        _add_turkish_paragraph(doc, )

        # 1. Özet
        _add_turkish_heading(doc, '1. YÖNETİCİ ÖZETİ', 1)
        if analysis.get('summary'):
            _add_turkish_paragraph(doc, analysis['summary'])
        else:
            _add_turkish_paragraph(doc, "Bu rapor, orman ürünleri kullanımı ve ormansızlaşma risklerini değerlendirmektedir.")

        # 2. Orman Ürünleri
        doc.add_page_break()
        _add_turkish_heading(doc, '2. ORMAN ÜRÜNLERİ KULLANIMI', 1)
        _add_turkish_paragraph(doc, """
Şirketimiz, tedarik zincirinde aşağıdaki orman ürünlerini kullanmaktadır:
- Ahşap ve ahşap ürünleri
- Kağıt ve karton
- Palm yağı
- Sığır eti/deri
        """.strip())

        # 3. Riskler
        if analysis.get('risks'):
            doc.add_page_break()
            _add_turkish_heading(doc, '3. ORMANSIZLAŞMA RİSKLERİ', 1)
            _add_turkish_paragraph(doc, analysis['risks'])

        # 4. Öneriler
        if analysis.get('recommendations'):
            doc.add_page_break()
            _add_turkish_heading(doc, '4. ÖNERİLER', 1)
            _add_turkish_paragraph(doc, analysis['recommendations'])

        # Altbilgi
        doc.add_page_break()
        footer = _add_turkish_paragraph(doc, )
        footer.add_run(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y")}\n').font.size = Pt(9)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return doc

