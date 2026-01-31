#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SASB Report Generator - SASB raporu oluşturma sınıfı
- PDF raporu oluşturma
- DOCX raporu oluşturma
- Excel raporu oluşturma
- SASB disclosure raporu
"""

import logging
from datetime import datetime
from typing import Any, Dict, Optional

# Rapor oluşturma kütüphaneleri
try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.info("ReportLab yüklenmedi. PDF raporları oluşturulamayacak.")

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    logging.info("python-docx yüklenmedi. DOCX raporları oluşturulamayacak.")

try:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    logging.info("openpyxl yüklenmedi. Excel raporları oluşturulamayacak.")

import sqlite3

from ..reporting.font_utils import register_turkish_fonts_reportlab
from .sasb_calculator import SASBCalculator
from .sasb_manager import SASBManager


def _add_turkish_paragraph(doc, text=None, style=None, font_name='Calibri', font_size=12):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    para = doc.add_paragraph(text if text is not None else '', style=style)
    for run in para.runs:
        try:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        try:
            run.font.name = font_name
            if level == 0:
                run.font.size = Pt(19)
            elif level == 1:
                run.font.size = Pt(15)
            else:
                run.font.size = Pt(13)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
    return heading


class SASBReportGenerator:
    """SASB raporu oluşturma sınıfı"""

    def __init__(self, db_path: str):
        """
        SASB Report Generator'ı başlat
        
        Args:
            db_path: Veritabanı dosya yolu
        """
        self.db_path = db_path
        self.sasb_manager = SASBManager(db_path)
        self.sasb_calculator = SASBCalculator(db_path)
        self.logger = logging.getLogger(__name__)

    def _add_logo_to_docx(self, doc, company_id: int):
        """DOCX raporuna logo ekle"""
        try:
            from modules.reporting.brand_identity_manager import BrandIdentityManager
            bim = BrandIdentityManager(self.db_path, company_id)
            bi = bim.get_brand_identity(company_id)
            logo_path = bi.get('logo_path')
            
            if logo_path and os.path.exists(logo_path):
                from docx.shared import Inches
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(logo_path, width=Inches(1.6))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            self.logger.warning(f"Logo ekleme hatası (DOCX): {e}")

    def generate_pdf_report(self, company_id: int, output_path: str,
                          year: Optional[int] = None) -> bool:
        """
        PDF SASB raporu oluştur
        
        Args:
            company_id: Şirket ID'si
            output_path: Çıktı dosya yolu
            year: Raporlama yılı
            
        Returns:
            Başarı durumu
        """
        try:
            if not REPORTLAB_AVAILABLE:
                return False
            # Rapor verilerini al
            report_data = self._prepare_report_data(company_id, year)
            if not report_data:
                return False

            # PDF oluştur
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            # Türkçe fontları kaydet
            register_turkish_fonts_reportlab()

            # Özel stiller - TÜRKÇE KARAKTER DESTEKLİ
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=19,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.darkgreen,
                fontName='NotoSans'
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=15,
                spaceAfter=12,
                textColor=colors.darkblue,
                fontName='NotoSans'
            )
            ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=12,
                fontName='NotoSans'
            )

            # Rapor içeriği
            story = []

            # Logo
            try:
                from modules.reporting.brand_identity_manager import BrandIdentityManager
                from reportlab.platypus import Image
                bim = BrandIdentityManager(self.db_path, company_id)
                bi = bim.get_brand_identity(company_id)
                logo_path = bi.get('logo_path')
                
                if logo_path and os.path.exists(logo_path):
                    img = Image(logo_path, width=1.5*inch, height=1.5*inch, kind='proportional')
                    img.hAlign = 'CENTER'
                    story.append(img)
                    story.append(Spacer(1, 10))
            except Exception as e:
                self.logger.warning(f"Logo ekleme hatası (PDF): {e}")

            # Başlık
            story.append(Paragraph("SASB DISCLOSURE RAPORU", title_style))
            story.append(Spacer(1, 20))

            # Şirket bilgileri
            story.append(Paragraph("Şirket Bilgileri", heading_style))
            company_info = [
                ["Şirket Adı", report_data['company_name']],
                ["Sektör", report_data['sector_name']],
                ["Raporlama Yılı", str(report_data['year'])],
                ["Rapor Tarihi", datetime.now().strftime("%d.%m.%Y")]
            ]

            company_table = Table(company_info, colWidths=[2*inch, 4*inch])
            company_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'NotoSans-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(company_table)
            story.append(Spacer(1, 20))

            # Materiality skoru
            story.append(Paragraph("Finansal Materiality Skoru", heading_style))
            materiality_data = [
                ["Genel Skor", f"{report_data['materiality_score']['overall_score']:.1f}%"],
                ["Disclosure Oranı", f"{report_data['materiality_score']['disclosure_rate']:.1f}%"],
                ["Materiality Oranı", f"{report_data['materiality_score']['materiality_rate']:.1f}%"],
                ["Not", report_data['materiality_score']['grade']]
            ]

            materiality_table = Table(materiality_data, colWidths=[2*inch, 2*inch])
            materiality_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'NotoSans-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (1, 0), (1, -1), colors.lightcyan),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(materiality_table)
            story.append(Spacer(1, 20))

            # Disclosure verileri
            story.append(Paragraph("Disclosure Verileri", heading_style))

            # Boyut bazında veriler
            for dimension, data in report_data['dimension_scores'].items():
                story.append(Paragraph(f"{dimension} Boyutu", styles['Heading3']))

                dim_data = [
                    ["Toplam Topics", str(data['total_topics'])],
                    ["Material Topics", str(data['material_topics'])],
                    ["Disclose Edilen", str(data['disclosed_topics'])],
                    ["Disclose Edilen Material", str(data['disclosed_material'])],
                    ["Skor", f"{data['score']:.1f}%"]
                ]

                dim_table = Table(dim_data, colWidths=[2*inch, 1*inch])
                dim_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'NotoSans-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))

                story.append(dim_table)
                story.append(Spacer(1, 10))

            # Detaylı disclosure tablosu
            story.append(Paragraph("Detaylı Disclosure Verileri", heading_style))

            disclosure_headers = ["Topic Kodu", "Topic Adı", "Metrik", "Değer", "Birim"]
            disclosure_data = [disclosure_headers]

            for disclosure in report_data['disclosures']:
                disclosure_data.append([
                    disclosure['topic_code'],
                    disclosure['topic_name'][:30] + "..." if len(disclosure['topic_name']) > 30 else disclosure['topic_name'],
                    disclosure['metric_code'],
                    disclosure['metric_value'][:20] + "..." if len(disclosure['metric_value']) > 20 else disclosure['metric_value'],
                    disclosure['unit_of_measure']
                ])

            disclosure_table = Table(disclosure_data, colWidths=[1*inch, 2*inch, 1*inch, 1.5*inch, 0.8*inch])
            disclosure_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'NotoSans-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(disclosure_table)

            # PDF'i oluştur
            doc.build(story)
            return True

        except Exception as e:
            self.logger.error(f"PDF raporu oluşturulamadı: {e}")
            return False

    def generate_docx_report(self, company_id: int, output_path: str,
                           year: Optional[int] = None) -> bool:
        """
        DOCX SASB raporu oluştur
        
        Args:
            company_id: Şirket ID'si
            output_path: Çıktı dosya yolu
            year: Raporlama yılı
            
        Returns:
            Başarı durumu
        """
        try:
            # Rapor verilerini al
            report_data = self._prepare_report_data(company_id, year)
            if not report_data:
                return False

            # DOCX oluştur
            doc = Document()

            # Logo
            self._add_logo_to_docx(doc, company_id)

            # Başlık
            title = _add_turkish_heading(doc, 'SASB DISCLOSURE RAPORU', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Şirket bilgileri
            _add_turkish_heading(doc, 'Şirket Bilgileri', level=1)

            company_info = doc.add_table(rows=1, cols=2)
            company_info.style = 'Table Grid'

            company_data = [
                ("Şirket Adı", report_data['company_name']),
                ("Sektör", report_data['sector_name']),
                ("Raporlama Yılı", str(report_data['year'])),
                ("Rapor Tarihi", datetime.now().strftime("%d.%m.%Y"))
            ]

            for key, value in company_data:
                row = company_info.add_row().cells
                row[0].text = key
                row[1].text = value
                row[0].paragraphs[0].runs[0].bold = True

            # Materiality skoru
            _add_turkish_heading(doc, 'Finansal Materiality Skoru', level=1)

            materiality_table = doc.add_table(rows=1, cols=2)
            materiality_table.style = 'Table Grid'

            materiality_data = [
                ("Genel Skor", f"{report_data['materiality_score']['overall_score']:.1f}%"),
                ("Disclosure Oranı", f"{report_data['materiality_score']['disclosure_rate']:.1f}%"),
                ("Materiality Oranı", f"{report_data['materiality_score']['materiality_rate']:.1f}%"),
                ("Not", report_data['materiality_score']['grade'])
            ]

            for key, value in materiality_data:
                row = materiality_table.add_row().cells
                row[0].text = key
                row[1].text = value
                row[0].paragraphs[0].runs[0].bold = True

            # Boyut bazında skorlar
            _add_turkish_heading(doc, 'Boyut Bazında Skorlar', level=1)

            for dimension, data in report_data['dimension_scores'].items():
                _add_turkish_heading(doc, f"{dimension} Boyutu", level=2)

                dim_table = doc.add_table(rows=1, cols=2)
                dim_table.style = 'Table Grid'

                dim_data = [
                    ("Toplam Topics", str(data['total_topics'])),
                    ("Material Topics", str(data['material_topics'])),
                    ("Disclose Edilen", str(data['disclosed_topics'])),
                    ("Disclose Edilen Material", str(data['disclosed_material'])),
                    ("Skor", f"{data['score']:.1f}%")
                ]

                for key, value in dim_data:
                    row = dim_table.add_row().cells
                    row[0].text = key
                    row[1].text = value
                    row[0].paragraphs[0].runs[0].bold = True

            # Detaylı disclosure verileri
            _add_turkish_heading(doc, 'Detaylı Disclosure Verileri', level=1)

            disclosure_table = doc.add_table(rows=1, cols=5)
            disclosure_table.style = 'Table Grid'

            # Başlıklar
            headers = ["Topic Kodu", "Topic Adı", "Metrik", "Değer", "Birim"]
            for i, header in enumerate(headers):
                disclosure_table.cell(0, i).text = header
                disclosure_table.cell(0, i).paragraphs[0].runs[0].bold = True

            # Veriler
            for disclosure in report_data['disclosures']:
                row = disclosure_table.add_row()
                row.cells[0].text = disclosure['topic_code']
                row.cells[1].text = disclosure['topic_name']
                row.cells[2].text = disclosure['metric_code']
                row.cells[3].text = disclosure['metric_value']
                row.cells[4].text = disclosure['unit_of_measure']

            # DOCX'i kaydet
            doc.save(output_path)
            return True

        except Exception as e:
            self.logger.error(f"DOCX raporu oluşturulamadı: {e}")
            return False

    def generate_excel_report(self, company_id: int, output_path: str,
                            year: Optional[int] = None) -> bool:
        """
        Excel SASB raporu oluştur
        
        Args:
            company_id: Şirket ID'si
            output_path: Çıktı dosya yolu
            year: Raporlama yılı
            
        Returns:
            Başarı durumu
        """
        try:
            # Rapor verilerini al
            report_data = self._prepare_report_data(company_id, year)
            if not report_data:
                return False

            # Excel çalışma kitabı oluştur
            wb = openpyxl.Workbook()

            # Stil tanımları
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            data_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            center_alignment = Alignment(horizontal='center', vertical='center')

            # 1. Özet sayfası
            ws_summary = wb.active
            ws_summary.title = "Özet"

            # Başlık
            ws_summary['A1'] = "SASB DISCLOSURE RAPORU"
            ws_summary['A1'].font = Font(bold=True, size=17)
            ws_summary.merge_cells('A1:E1')

            # Şirket bilgileri
            ws_summary['A3'] = "Şirket Bilgileri"
            ws_summary['A3'].font = Font(bold=True, size=13)

            company_data = [
                ("Şirket Adı", report_data['company_name']),
                ("Sektör", report_data['sector_name']),
                ("Raporlama Yılı", str(report_data['year'])),
                ("Rapor Tarihi", datetime.now().strftime("%d.%m.%Y"))
            ]

            for i, (key, value) in enumerate(company_data, start=4):
                ws_summary[f'A{i}'] = key
                ws_summary[f'B{i}'] = value
                ws_summary[f'A{i}'].font = Font(bold=True)

            # Materiality skoru
            ws_summary['A9'] = "Finansal Materiality Skoru"
            ws_summary['A9'].font = Font(bold=True, size=12)

            materiality_data = [
                ("Genel Skor", f"{report_data['materiality_score']['overall_score']:.1f}%"),
                ("Disclosure Oranı", f"{report_data['materiality_score']['disclosure_rate']:.1f}%"),
                ("Materiality Oranı", f"{report_data['materiality_score']['materiality_rate']:.1f}%"),
                ("Not", report_data['materiality_score']['grade'])
            ]

            for i, (key, value) in enumerate(materiality_data, start=10):
                ws_summary[f'A{i}'] = key
                ws_summary[f'B{i}'] = value
                ws_summary[f'A{i}'].font = Font(bold=True)

            # 2. Boyut bazında skorlar sayfası
            ws_dimensions = wb.create_sheet("Boyut Skorları")

            # Başlıklar
            headers = ["Boyut", "Toplam Topics", "Material Topics", "Disclose Edilen", "Disclose Edilen Material", "Skor"]
            for i, header in enumerate(headers, start=1):
                cell = ws_dimensions.cell(row=1, column=i, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_alignment
                cell.border = border

            # Veriler
            row = 2
            for dimension, data in report_data['dimension_scores'].items():
                ws_dimensions.cell(row=row, column=1, value=dimension)
                ws_dimensions.cell(row=row, column=2, value=data['total_topics'])
                ws_dimensions.cell(row=row, column=3, value=data['material_topics'])
                ws_dimensions.cell(row=row, column=4, value=data['disclosed_topics'])
                ws_dimensions.cell(row=row, column=5, value=data['disclosed_material'])
                ws_dimensions.cell(row=row, column=6, value=f"{data['score']:.1f}%")

                # Stil uygula
                for col in range(1, 7):
                    cell = ws_dimensions.cell(row=row, column=col)
                    cell.fill = data_fill
                    cell.border = border
                    if col == 1:
                        cell.font = Font(bold=True)
                    else:
                        cell.alignment = center_alignment

                row += 1

            # Sütun genişliklerini ayarla
            for col in range(1, 7):
                ws_dimensions.column_dimensions[get_column_letter(col)].width = 20

            # 3. Detaylı disclosure verileri sayfası
            ws_disclosures = wb.create_sheet("Disclosure Verileri")

            # Başlıklar
            headers = ["Topic Kodu", "Topic Adı", "Metrik Kodu", "Metrik Adı", "Değer", "Birim", "Kaynak", "Notlar"]
            for i, header in enumerate(headers, start=1):
                cell = ws_disclosures.cell(row=1, column=i, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_alignment
                cell.border = border

            # Veriler
            row = 2
            for disclosure in report_data['disclosures']:
                ws_disclosures.cell(row=row, column=1, value=disclosure['topic_code'])
                ws_disclosures.cell(row=row, column=2, value=disclosure['topic_name'])
                ws_disclosures.cell(row=row, column=3, value=disclosure['metric_code'])
                ws_disclosures.cell(row=row, column=4, value=disclosure['metric_name'])
                ws_disclosures.cell(row=row, column=5, value=disclosure['metric_value'])
                ws_disclosures.cell(row=row, column=6, value=disclosure['unit_of_measure'])
                ws_disclosures.cell(row=row, column=7, value=disclosure['data_source'])
                ws_disclosures.cell(row=row, column=8, value=disclosure['notes'])

                # Stil uygula
                for col in range(1, 9):
                    cell = ws_disclosures.cell(row=row, column=col)
                    cell.fill = data_fill
                    cell.border = border
                    cell.alignment = center_alignment

                row += 1

            # Sütun genişliklerini ayarla
            column_widths = [15, 30, 15, 30, 15, 10, 20, 30]
            for i, width in enumerate(column_widths, start=1):
                ws_disclosures.column_dimensions[get_column_letter(i)].width = width

            # Excel dosyasını kaydet
            wb.save(output_path)
            return True

        except Exception as e:
            self.logger.error(f"Excel raporu oluşturulamadı: {e}")
            return False

    def _prepare_report_data(self, company_id: int, year: Optional[int] = None) -> Optional[Dict[str, Any]]:
        """
        Rapor verilerini hazırla
        
        Args:
            company_id: Şirket ID'si
            year: Raporlama yılı
            
        Returns:
            Rapor verileri
        """
        try:
            conn = sqlite3.connect(self.db_path)
            cur = conn.cursor()

            # Şirket bilgilerini al
            cur.execute("SELECT name FROM companies WHERE id = ?", (company_id,))
            company_result = cur.fetchone()
            
            # Year handling if not provided
            if year is None:
                cur.execute("SELECT MAX(year) FROM sasb_metric_responses WHERE company_id = ?", (company_id,))
                res = cur.fetchone()
                if res and res[0]:
                    year = res[0]
                else:
                    year = 2024
            
            conn.close()

            if not company_result:
                return None

            company_name = company_result[0]

            # Sektör bilgilerini al
            current_sector = self.sasb_manager.get_company_sector(company_id, year)
            if not current_sector:
                # Sektör verisi bulunamadı
                return None

            # Materiality skorunu hesapla
            materiality_score = self.sasb_calculator.calculate_materiality_score(
                company_id, 
                current_sector['id'], 
                year
            )
            
            # Disclosure verilerini al
            disclosures = self.sasb_manager.get_company_disclosures(company_id, year)

            return {
                'company_name': company_name,
                'sector_name': current_sector['sector_name'],
                'year': year,
                'materiality_score': materiality_score,
                'dimension_scores': materiality_score.get('dimension_scores', {}),
                'disclosures': disclosures
            }

        except Exception as e:
            self.logger.error(f"Rapor verileri hazırlanamadı: {e}")
            return None
