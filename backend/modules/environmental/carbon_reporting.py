import logging
import json
import os
import sqlite3
from datetime import datetime
from typing import Dict, List, Optional
from docx.shared import Cm

def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    from docx.shared import Pt
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return heading

class CarbonReporting:
    """Karbon Yönetimi Raporlama"""

    def __init__(self, carbon_manager) -> None:
        self.carbon_manager = carbon_manager
        # company_id might not be stored in manager, but passed in methods.
        # But EnergyReporting assumes it might be there. CarbonManager (env) doesn't store company_id in init.
        
    def generate_carbon_report(self, company_id: int, period: str,
                             include_details: bool = True, formats: List[str] = None) -> Dict[str, str]:
        """Karbon yönetimi raporu oluştur (DOCX ve Excel)"""
        if formats is None:
            formats = ['docx', 'excel']
            
        generated_files = {}
        
        # Veri toplama
        # Period should be year (int) for CarbonManager methods usually, but let's handle string 'YYYY'
        try:
            year = int(period)
        except (ValueError, TypeError):
            # If period is not a year, we might default to current year or handle error
            # For now, let's assume period is YYYY string
            logging.warning(f"Geçersiz dönem formatı: {period}. Yıl olarak varsayılıyor.")
            year = datetime.now().year

        records = self.carbon_manager.get_carbon_records(company_id, year)
        metrics = self.carbon_manager.get_total_carbon_footprint(company_id, year)
        
        # DOCX Raporu
        if 'docx' in formats:
            docx_path = self._generate_docx_report(company_id, period, records, metrics, include_details)
            if docx_path:
                generated_files['docx'] = docx_path
                
        # Excel Raporu
        if 'excel' in formats:
            excel_path = self._generate_excel_report(company_id, period, records, metrics)
            if excel_path:
                generated_files['excel'] = excel_path
                
        return generated_files

    def _add_logo(self, doc, company_id: int):
        """Raporun başına şirket logosunu ekle"""
        try:
            # 1. Veritabanından logo yolunu al
            conn = sqlite3.connect(self.carbon_manager.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT logo_path FROM company_profiles WHERE company_id = ?", (company_id,))
            result = cursor.fetchone()
            conn.close()
            
            logo_path = None
            if result and result[0] and os.path.exists(result[0]):
                logo_path = result[0]
            else:
                # 2. Varsayılan yolda ara
                data_dir = os.path.dirname(self.carbon_manager.db_path)
                possible_path = os.path.join(data_dir, "company_logos", f"company_{company_id}_logo.png")
                if os.path.exists(possible_path):
                    logo_path = possible_path
                else:
                    # jpg dene
                    possible_path = possible_path.replace(".png", ".jpg")
                    if os.path.exists(possible_path):
                        logo_path = possible_path

            if logo_path:
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, height=Cm(2.5))
        except Exception as e:
            logging.warning(f"Logo eklenirken hata: {e}")

    def _generate_docx_report(self, company_id: int, period: str, records: List[Dict], 
                            metrics: Dict, include_details: bool) -> Optional[str]:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Logo ekle
            self._add_logo(doc, company_id)
            
            # Başlık
            title = _add_turkish_heading(doc, 'Karbon Ayak İzi Raporu', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Rapor bilgileri
            _add_turkish_heading(doc, 'Rapor Bilgileri', level=1)
            report_info = _add_turkish_paragraph(doc, '')
            report_info.add_run('Şirket ID: ').bold = True
            report_info.add_run(f'{company_id}\n')
            report_info.add_run('Dönem (Yıl): ').bold = True
            report_info.add_run(f'{period}\n')
            report_info.add_run('Rapor Tarihi: ').bold = True
            report_info.add_run(f'{datetime.now().strftime("%d.%m.%Y %H:%M")}\n')
            
            # Özet Metrikler
            _add_turkish_heading(doc, 'Özet Metrikler', level=1)
            if metrics:
                metrics_table = doc.add_table(rows=1, cols=2)
                metrics_table.style = 'Table Grid'
                hdr_cells = metrics_table.rows[0].cells
                hdr_cells[0].text = 'Kapsam'
                hdr_cells[1].text = 'Emisyon (tCO2e)'
                
                metrics_data = [
                    ('Scope 1 (Doğrudan)', f"{metrics.get('scope1_total', 0):,.2f}"),
                    ('Scope 2 (Enerji)', f"{metrics.get('scope2_total', 0):,.2f}"),
                    ('Scope 3 (Diğer)', f"{metrics.get('scope3_total', 0):,.2f}"),
                    ('Toplam Karbon Ayak İzi', f"{metrics.get('total_footprint', 0):,.2f}")
                ]
                
                for metric_name, metric_value in metrics_data:
                    row_cells = metrics_table.add_row().cells
                    row_cells[0].text = metric_name
                    row_cells[1].text = metric_value

            # Emisyon Kayıtları
            if include_details and records:
                _add_turkish_heading(doc, 'Emisyon Kayıtları Detayı', level=1)
                
                # Tablo oluştur
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Kapsam'
                hdr_cells[1].text = 'Kaynak/Kategori'
                hdr_cells[2].text = 'Miktar'
                hdr_cells[3].text = 'Emisyon (tCO2e)'
                hdr_cells[4].text = 'Tedarikçi'
                hdr_cells[5].text = 'Fatura Tarihi'
                
                for record in records:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(record.get('scope', ''))
                    row_cells[1].text = f"{record.get('source', '')} / {record.get('type', '')}"
                    row_cells[2].text = f"{record.get('amount', 0)} {record.get('unit', '')}"
                    row_cells[3].text = f"{record.get('total_emissions', 0):.2f}"
                    row_cells[4].text = str(record.get('supplier', '') or '-')
                    row_cells[5].text = str(record.get('invoice_date', '') or '-')

            # Klasör oluştur ve kaydet
            base_dir = os.path.join("reports", f"company_{company_id}")
            os.makedirs(base_dir, exist_ok=True)
            filename = f"karbon_raporu_{company_id}_{period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(base_dir, filename)
            
            doc.save(filepath)
            logging.info(f"[OK] Karbon raporu oluşturuldu: {filepath}")
            return filepath
            
        except ImportError:
            logging.error("python-docx kütüphanesi yüklü değil.")
            return None
        except Exception as e:
            logging.error(f"DOCX rapor hatası: {e}")
            return None

    def _generate_excel_report(self, company_id: int, period: str, records: List[Dict], metrics: Dict) -> Optional[str]:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
            
            wb = Workbook()
            
            # Özet Sayfası
            ws_summary = wb.active
            ws_summary.title = "Özet"
            
            ws_summary['A1'] = "Karbon Ayak İzi Raporu"
            ws_summary['A1'].font = Font(size=16, bold=True)
            ws_summary['A2'] = f"Şirket ID: {company_id}"
            ws_summary['A3'] = f"Dönem: {period}"
            ws_summary['A4'] = f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            
            if metrics:
                row = 6
                ws_summary[f'A{row}'] = "Kapsam"
                ws_summary[f'B{row}'] = "Emisyon (tCO2e)"
                ws_summary[f'A{row}'].font = Font(bold=True)
                
                metrics_data = [
                    ('Scope 1 (Doğrudan)', metrics.get('scope1_total', 0)),
                    ('Scope 2 (Enerji)', metrics.get('scope2_total', 0)),
                    ('Scope 3 (Diğer)', metrics.get('scope3_total', 0)),
                    ('Toplam Karbon Ayak İzi', metrics.get('total_footprint', 0))
                ]
                
                for metric_name, metric_value in metrics_data:
                    row += 1
                    ws_summary[f'A{row}'] = metric_name
                    ws_summary[f'B{row}'] = metric_value
            
            # Kayıtlar Sayfası
            if records:
                ws_records = wb.create_sheet("Emisyon Kayıtları")
                
                headers = [
                    'Kapsam', 'Kaynak', 'Tür/Kategori', 'Miktar', 'Birim', 
                    'Emisyon (tCO2e)', 'Fatura Tarihi', 'Son Ödeme', 'Tedarikçi'
                ]
                
                for col, header in enumerate(headers, 1):
                    cell = ws_records.cell(row=1, column=col)
                    cell.value = header
                    cell.font = Font(bold=True)
                    
                for row_idx, record in enumerate(records, 2):
                    ws_records.cell(row=row_idx, column=1, value=record.get('scope'))
                    ws_records.cell(row=row_idx, column=2, value=record.get('source'))
                    ws_records.cell(row=row_idx, column=3, value=record.get('type'))
                    ws_records.cell(row=row_idx, column=4, value=record.get('amount'))
                    ws_records.cell(row=row_idx, column=5, value=record.get('unit'))
                    ws_records.cell(row=row_idx, column=6, value=record.get('total_emissions'))
                    ws_records.cell(row=row_idx, column=7, value=record.get('invoice_date'))
                    ws_records.cell(row=row_idx, column=8, value=record.get('due_date'))
                    ws_records.cell(row=row_idx, column=9, value=record.get('supplier'))

            # Klasör oluştur ve kaydet
            base_dir = os.path.join("reports", f"company_{company_id}")
            os.makedirs(base_dir, exist_ok=True)
            filename = f"karbon_raporu_{company_id}_{period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            filepath = os.path.join(base_dir, filename)
            
            wb.save(filepath)
            logging.info(f"[OK] Excel raporu oluşturuldu: {filepath}")
            return filepath
            
        except ImportError:
            logging.error("openpyxl kütüphanesi yüklü değil.")
            return None
        except Exception as e:
            logging.error(f"Excel rapor hatası: {e}")
            return None
