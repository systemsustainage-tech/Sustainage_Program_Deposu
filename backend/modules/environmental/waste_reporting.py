import logging
import json
import os
import sqlite3
from datetime import datetime
from typing import Dict, List, Optional
from docx.shared import Cm

def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    from docx.shared import Pt
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return heading

class WasteReporting:
    """Atık Yönetimi Raporlama"""

    def __init__(self, waste_manager) -> None:
        self.waste_manager = waste_manager
        self.company_id = waste_manager.company_id if hasattr(waste_manager, 'company_id') else None

    def generate_waste_report(self, company_id: int, period: str,
                             include_details: bool = True, formats: List[str] = None) -> Dict[str, str]:
        """Atık yönetimi raporu oluştur (DOCX ve Excel)"""
        if formats is None:
            formats = ['docx', 'excel']
            
        generated_files = {}
        
        # Veri toplama
        records = self.waste_manager.get_waste_records(company_id, period)
        
        # Metrics hesaplama
        metrics = {}
        if hasattr(self.waste_manager, 'calculate_waste_metrics'):
            metrics = self.waste_manager.calculate_waste_metrics(company_id)
        
        # DOCX Raporu
        if 'docx' in formats:
            docx_path = self._generate_docx_report(company_id, period, records, metrics, include_details)
            if docx_path:
                generated_files['docx'] = docx_path
                
        return generated_files

    def _add_logo(self, doc, company_id: int):
        """Raporun başına şirket logosunu ekle"""
        try:
            # 1. Veritabanından logo yolunu al
            conn = sqlite3.connect(self.waste_manager.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT logo_path FROM company_profiles WHERE company_id = ?", (company_id,))
            result = cursor.fetchone()
            conn.close()
            
            logo_path = None
            if result and result[0] and os.path.exists(result[0]):
                logo_path = result[0]
            else:
                # 2. Varsayılan yolda ara
                data_dir = os.path.dirname(self.waste_manager.db_path)
                possible_path = os.path.join(data_dir, "company_logos", f"company_{company_id}_logo.png")
                if os.path.exists(possible_path):
                    logo_path = possible_path
                else:
                    # jpg dene
                    possible_path = possible_path.replace(".png", ".jpg")
                    if os.path.exists(possible_path):
                        logo_path = possible_path

            if logo_path:
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, height=Cm(2.5))
        except Exception as e:
            logging.warning(f"Logo eklenirken hata: {e}")

    def _generate_docx_report(self, company_id: int, period: str, records: List[Dict], 
                            metrics: Dict, include_details: bool) -> Optional[str]:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Logo ekle
            self._add_logo(doc, company_id)
            
            # Başlık
            title = _add_turkish_heading(doc, 'Atık Yönetimi Raporu', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Rapor bilgileri
            _add_turkish_heading(doc, 'Rapor Bilgileri', level=1)
            report_info = _add_turkish_paragraph(doc, '')
            report_info.add_run('Şirket ID: ').bold = True
            report_info.add_run(f'{company_id}\n')
            report_info.add_run('Dönem: ').bold = True
            report_info.add_run(f'{period or "Tümü"}\n')
            report_info.add_run('Rapor Tarihi: ').bold = True
            report_info.add_run(f'{datetime.now().strftime("%d.%m.%Y %H:%M")}\n')
            
            # Özet Metrikler
            _add_turkish_heading(doc, 'Özet Metrikler', level=1)
            if metrics:
                metrics_table = doc.add_table(rows=1, cols=2)
                metrics_table.style = 'Table Grid'
                hdr_cells = metrics_table.rows[0].cells
                hdr_cells[0].text = 'Metrik'
                hdr_cells[1].text = 'Değer'
                
                metrics_data = [
                    ('Toplam Atık Miktarı', f"{metrics.get('total_waste', 0):,.2f} kg"),
                    ('Toplam Geri Dönüştürülen', f"{metrics.get('total_recycled', 0):,.2f} kg"),
                    ('Geri Dönüşüm Oranı', f"%{metrics.get('recycling_ratio', 0):.1f}"),
                ]
                
                for metric_name, metric_value in metrics_data:
                    row_cells = metrics_table.add_row().cells
                    row_cells[0].text = metric_name
                    row_cells[1].text = metric_value

            # Atık Kayıtları
            if include_details and records:
                _add_turkish_heading(doc, 'Atık Üretim ve Bertaraf Kayıtları', level=1)
                
                # Tablo oluştur
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Tarih'
                hdr_cells[1].text = 'Tür'
                hdr_cells[2].text = 'Kategori'
                hdr_cells[3].text = 'Miktar'
                hdr_cells[4].text = 'Yöntem'
                hdr_cells[5].text = 'Lokasyon'
                
                for record in records:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(record.get('date', '')[:10])
                    row_cells[1].text = str(record.get('type', ''))
                    row_cells[2].text = str(record.get('category', ''))
                    row_cells[3].text = f"{record.get('amount', 0)} {record.get('unit', '')}"
                    row_cells[4].text = str(record.get('method', '') or '-')
                    row_cells[5].text = str(record.get('location', '') or '-')

            # Klasör oluştur ve kaydet
            base_dir = os.path.join("reports", f"company_{company_id}")
            os.makedirs(base_dir, exist_ok=True)
            filename = f"atik_raporu_{company_id}_{period or 'all'}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(base_dir, filename)
            
            doc.save(filepath)
            logging.info(f"[OK] Atık raporu oluşturuldu: {filepath}")
            return filepath
            
        except ImportError:
            logging.error("python-docx kütüphanesi yüklü değil.")
            return None
        except Exception as e:
            logging.error(f"DOCX rapor hatası: {e}")
            return None
