import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
TSRS Raporlama Modülü
TSRS standartlarına uygun raporlar oluşturma (PDF, Word, Excel)
"""

import os
import sqlite3
from datetime import datetime
from typing import Dict, List

import pandas as pd
from config.database import DB_PATH

# Raporlama kütüphaneleri
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    # pdfmetrics/TTFont kullanılmıyor
except ImportError:
    logging.info("reportlab kütüphanesi bulunamadı. PDF raporları oluşturulamaz.")

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    logging.info("python-docx kütüphanesi bulunamadı. Word raporları oluşturulamaz.")

def _add_turkish_paragraph(doc, text=None, style=None, font_name='Calibri', font_size=12):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    para = doc.add_paragraph(text if text is not None else '', style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = font_name
        try:
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
        # Basit bir 1pt artış mantığı
        if level == 0:
            run.font.size = Pt(19)
        elif level == 1:
            run.font.size = Pt(15)
        else:
            run.font.size = Pt(13)
    return heading


class TSRSReporting:
    """TSRS Raporlama sınıfı"""

    def __init__(self, db_path: str = DB_PATH) -> None:
        if not os.path.isabs(db_path):
            repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
            db_path = os.path.join(repo_root, db_path)
        db_dir = os.path.dirname(db_path)
        try:
            os.makedirs(db_dir, exist_ok=True)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
        self.db_path = db_path
        self._setup_turkish_fonts()

    def _setup_turkish_fonts(self) -> None:
        """Türkçe font desteği ayarla"""
        try:
            from ..reporting.font_utils import register_turkish_fonts_reportlab
            register_turkish_fonts_reportlab()
            self.turkish_font = 'NotoSans'

        except Exception as e:
            logging.error(f"Font ayarlama hatası: {e}")
            self.turkish_font = 'Helvetica'

    def get_connection(self) -> sqlite3.Connection:
        """Veritabanı bağlantısı"""
        return sqlite3.connect(self.db_path)

    def get_company_info(self, company_id: int) -> Dict:
        """Şirket bilgilerini getir (şema uyumlu ve null-güvenli)"""
        conn = self.get_connection()
        cursor = conn.cursor()

        try:
            cursor.execute("PRAGMA table_info(company_info)")
            cols = [row[1] for row in cursor.fetchall()]

            cursor.execute("SELECT * FROM company_info WHERE company_id = ?", (company_id,))
            row = cursor.fetchone()
            if not row:
                return {
                    'company_name': 'Şirket Adı',
                    'sector': 'Sektör',
                    'address': 'Adres',
                    'phone': 'Telefon',
                    'email': 'E-posta',
                    'website': 'Web Sitesi',
                    'established_year': 'Kuruluş Yılı',
                    'employee_count': 'Çalışan Sayısı',
                    'revenue': 'Ciro'
                }

            data = dict(zip(cols, row))

            def pick(*keys):
                for k in keys:
                    v = data.get(k)
                    if v is not None and str(v).strip() != '':
                        return v
                return None

            return {
                'company_name': pick('ticari_unvan', 'sirket_adi') or 'Şirket Adı',
                'sector': pick('sektor') or 'Sektör',
                'address': pick('adres') or 'Adres',
                'phone': pick('telefon') or 'Telefon',
                'email': pick('email') or 'E-posta',
                'website': pick('website') or 'Web Sitesi',
                'established_year': pick('kurulusyili', 'kurulus_tarihi') or 'Kuruluş Yılı',
                'employee_count': pick('calisan_sayisi') or 'Çalışan Sayısı',
                'revenue': pick('yillik_ciro', 'revenue') or 'Ciro'
            }
        except Exception as e:
            logging.error(f"Şirket bilgileri getirilirken hata: {e}")
            return {
                'company_name': 'Şirket Adı',
                'sector': 'Sektör',
                'address': 'Adres',
                'phone': 'Telefon',
                'email': 'E-posta',
                'website': 'Web Sitesi',
                'established_year': 'Kuruluş Yılı',
                'employee_count': 'Çalışan Sayısı',
                'revenue': 'Ciro'
            }
        finally:
            conn.close()

    def get_tsrs_report_data(self, company_id: int, reporting_period: str) -> Dict:
        """TSRS rapor verilerini getir"""
        conn = self.get_connection()
        cursor = conn.cursor()

        try:
            # TSRS standartları ve yanıtları
            cursor.execute("""
                SELECT s.code, s.title, s.category, s.subcategory, s.requirement_level,
                       r.response_value, r.numerical_value, r.unit, r.reporting_status,
                       i.code as indicator_code, i.title as indicator_title
                FROM tsrs_standards s
                LEFT JOIN tsrs_indicators i ON s.id = i.standard_id
                LEFT JOIN tsrs_responses r ON i.id = r.indicator_id AND r.company_id = ? AND r.reporting_period = ?
                ORDER BY s.category, s.code, i.code
            """, (company_id, reporting_period))

            columns = [description[0] for description in cursor.description]
            data = []

            for row in cursor.fetchall():
                record = dict(zip(columns, row))
                data.append(record)

            # Kategorilere göre grupla
            categories: Dict[str, List[Dict]] = {}
            for record in data:
                category = record['category']
                if category not in categories:
                    categories[category] = []
                categories[category].append(record)

            return {
                'data': data,
                'categories': categories,
                'reporting_period': reporting_period
            }

        except Exception as e:
            logging.error(f"TSRS rapor verileri getirilirken hata: {e}")
            return {'data': [], 'categories': {}, 'reporting_period': reporting_period}
        finally:
            conn.close()

    def create_tsrs_pdf_report(self, company_id: int, reporting_period: str,
                              output_path: str) -> bool:
        """TSRS PDF raporu oluştur"""
        try:
            # Şirket bilgileri
            company_info = self.get_company_info(company_id)

            # TSRS verileri
            report_data = self.get_tsrs_report_data(company_id, reporting_period)

            # PDF oluştur
            doc = SimpleDocTemplate(output_path, pagesize=A4, topMargin=1*inch)
            story = []

            # Stil tanımlamaları
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=self.turkish_font,
                fontSize=19,
                spaceAfter=30,
                alignment=1  # Ortala
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontName=self.turkish_font,
                fontSize=15,
                spaceAfter=12
            )

            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=self.turkish_font,
                fontSize=11,
                spaceAfter=6
            )

            # Rapor içeriği
            story = []

            # Logo
            try:
                from modules.reporting.brand_identity_manager import BrandIdentityManager
                from reportlab.platypus import Image
                bim = BrandIdentityManager(self.db_path, company_id)
                bi = bim.get_brand_identity(company_id)
                logo_path = bi.get('logo_path')
                
                if logo_path and os.path.exists(logo_path):
                    img = Image(logo_path, width=1.5*inch, height=1.5*inch, kind='proportional')
                    img.hAlign = 'CENTER'
                    story.append(img)
                    story.append(Spacer(1, 10))
            except Exception as e:
                logging.warning(f"Logo ekleme hatası (PDF): {e}")

            # Başlık
            story.append(Paragraph("TSRS Sürdürülebilirlik Raporu", title_style))
            story.append(Spacer(1, 12))

            # Şirket bilgileri
            story.append(Paragraph(f"<b>Şirket:</b> {company_info['company_name']}", normal_style))
            story.append(Paragraph(f"<b>Sektör:</b> {company_info['sector']}", normal_style))
            story.append(Paragraph(f"<b>Raporlama Dönemi:</b> {reporting_period}", normal_style))
            story.append(Paragraph(f"<b>Rapor Tarihi:</b> {datetime.now().strftime('%d.%m.%Y')}", normal_style))
            story.append(Spacer(1, 20))

            # Kategoriler
            for category, records in report_data['categories'].items():
                story.append(Paragraph(f"<b>{category}</b>", heading_style))

                # Kategori tablosu
                table_data = [['Standart', 'Gösterge', 'Yanıt', 'Durum']]

                for record in records:
                    if record['indicator_code']:
                        table_data.append([
                            record['code'],
                            record['indicator_title'] or record['indicator_code'],
                            record['response_value'] or 'Yanıtlanmadı',
                            record['reporting_status'] or 'Beklemede'
                        ])

                if len(table_data) > 1:
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'NotoSans-Bold'),
                        ('FONTNAME', (0, 1), (-1, -1), 'NotoSans'),
                        ('FONTSIZE', (0, 0), (-1, 0), 11),
                        ('FONTSIZE', (0, 1), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))

                    story.append(table)
                    story.append(Spacer(1, 12))

            # PDF'i oluştur
            doc.build(story)
            logging.info(f"TSRS PDF raporu oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"TSRS PDF raporu oluşturulurken hata: {e}")
            return False

    def create_tsrs_docx_report(self, company_id: int, reporting_period: str,
                               output_path: str) -> bool:
        """TSRS Word raporu oluştur"""
        try:
            # Şirket bilgileri
            company_info = self.get_company_info(company_id)

            # TSRS verileri
            report_data = self.get_tsrs_report_data(company_id, reporting_period)

            # Word belgesi oluştur
            doc = Document()

            # Logo
            try:
                from modules.reporting.brand_identity_manager import BrandIdentityManager
                bim = BrandIdentityManager(self.db_path, company_id)
                bi = bim.get_brand_identity(company_id)
                logo_path = bi.get('logo_path')
                
                if logo_path and os.path.exists(logo_path):
                    from docx.shared import Inches
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(logo_path, width=Inches(1.6))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logging.warning(f"Logo ekleme hatası (DOCX): {e}")

            # Başlık
            title = _add_turkish_heading(doc, 'TSRS Sürdürülebilirlik Raporu', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Şirket bilgileri
            _add_turkish_heading(doc, 'Şirket Bilgileri', level=1)

            info_paragraph = _add_turkish_paragraph(doc, )
            info_paragraph.add_run('Şirket: ').bold = True
            info_paragraph.add_run(company_info['company_name'])

            info_paragraph = _add_turkish_paragraph(doc, )
            info_paragraph.add_run('Sektör: ').bold = True
            info_paragraph.add_run(company_info['sector'])

            info_paragraph = _add_turkish_paragraph(doc, )
            info_paragraph.add_run('Raporlama Dönemi: ').bold = True
            info_paragraph.add_run(reporting_period)

            info_paragraph = _add_turkish_paragraph(doc, )
            info_paragraph.add_run('Rapor Tarihi: ').bold = True
            info_paragraph.add_run(datetime.now().strftime('%d.%m.%Y'))

            # Kategoriler
            for category, records in report_data['categories'].items():
                _add_turkish_heading(doc, category, level=1)

                # Kategori tablosu
                if records:
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Başlık satırı
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Standart'
                    header_cells[1].text = 'Gösterge'
                    header_cells[2].text = 'Yanıt'
                    header_cells[3].text = 'Durum'

                    # Başlık hücrelerini kalın yap
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Veri satırları
                    for record in records:
                        if record['indicator_code']:
                            row_cells = table.add_row().cells
                            row_cells[0].text = record['code']
                            row_cells[1].text = record['indicator_title'] or record['indicator_code']
                            row_cells[2].text = record['response_value'] or 'Yanıtlanmadı'
                            row_cells[3].text = record['reporting_status'] or 'Beklemede'

                _add_turkish_paragraph(doc, )  # Boş satır

            # Belgeyi kaydet
            doc.save(output_path)
            logging.info(f"TSRS Word raporu oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"TSRS Word raporu oluşturulurken hata: {e}")
            return False

    def create_tsrs_excel_report(self, company_id: int, reporting_period: str,
                                output_path: str) -> bool:
        """TSRS Excel raporu oluştur"""
        try:
            # Şirket bilgileri
            company_info = self.get_company_info(company_id)

            # TSRS verileri
            report_data = self.get_tsrs_report_data(company_id, reporting_period)

            # Excel yazıcı oluştur
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                # Özet sayfası
                summary_data = {
                    'Bilgi': ['Şirket', 'Sektör', 'Raporlama Dönemi', 'Rapor Tarihi'],
                    'Değer': [
                        company_info['company_name'],
                        company_info['sector'],
                        reporting_period,
                        datetime.now().strftime('%d.%m.%Y')
                    ]
                }
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='Özet', index=False)

                # Ana veri sayfası
                if report_data['data']:
                    df = pd.DataFrame(report_data['data'])
                    df.to_excel(writer, sheet_name='TSRS Verileri', index=False)

                # Kategori bazlı sayfalar
                for category, records in report_data['categories'].items():
                    if records:
                        category_df = pd.DataFrame(records)
                        # Sayfa adını temizle
                        sheet_name = category.replace(' ', '_')[:31]  # Excel sayfa adı limiti
                        category_df.to_excel(writer, sheet_name=sheet_name, index=False)

                # İstatistik sayfası
                stats_data: Dict[str, List] = {
                    'Kategori': [],
                    'Toplam Standart': [],
                    'Yanıtlanan': [],
                    'Yanıt Oranı (%)': []
                }

                for category, records in report_data['categories'].items():
                    total_indicators = len([r for r in records if r['indicator_code']])
                    answered = len([r for r in records if r['response_value']])
                    answer_rate = (answered / total_indicators * 100) if total_indicators > 0 else 0

                    stats_data['Kategori'].append(category)
                    stats_data['Toplam Standart'].append(total_indicators)
                    stats_data['Yanıtlanan'].append(answered)
                    stats_data['Yanıt Oranı (%)'].append(round(answer_rate, 1))

                stats_df = pd.DataFrame(stats_data)
                stats_df.to_excel(writer, sheet_name='İstatistikler', index=False)

            logging.info(f"TSRS Excel raporu oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"TSRS Excel raporu oluşturulurken hata: {e}")
            return False

    def create_comprehensive_tsrs_report(self, company_id: int, reporting_period: str,
                                       output_dir: str) -> Dict[str, bool]:
        """Kapsamlı TSRS raporu oluştur (PDF, Word, Excel)"""
        results = {}

        # Çıktı dizinini oluştur
        os.makedirs(output_dir, exist_ok=True)

        # Dosya yolları
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        pdf_path = os.path.join(output_dir, f'TSRS_Raporu_{reporting_period}_{timestamp}.pdf')
        docx_path = os.path.join(output_dir, f'TSRS_Raporu_{reporting_period}_{timestamp}.docx')
        excel_path = os.path.join(output_dir, f'TSRS_Raporu_{reporting_period}_{timestamp}.xlsx')

        # PDF raporu
        try:
            results['pdf'] = self.create_tsrs_pdf_report(company_id, reporting_period, pdf_path)
        except Exception as e:
            logging.error(f"PDF raporu oluşturulurken hata: {e}")
            results['pdf'] = False

        # Word raporu
        try:
            results['docx'] = self.create_tsrs_docx_report(company_id, reporting_period, docx_path)
        except Exception as e:
            logging.error(f"Word raporu oluşturulurken hata: {e}")
            results['docx'] = False

        # Excel raporu
        try:
            results['excel'] = self.create_tsrs_excel_report(company_id, reporting_period, excel_path)
        except Exception as e:
            logging.error(f"Excel raporu oluşturulurken hata: {e}")
            results['excel'] = False

        return results
