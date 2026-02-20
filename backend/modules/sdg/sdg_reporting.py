import os
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional, Any

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # GUI olmayan backend

from core.base_manager import BaseTenantManager
from core.database_manager import DatabaseManager

# --- Türkçe Karakter Desteği için Font Ayarları (ReportLab) ---
def register_turkish_fonts_reportlab():
    """ReportLab için Türkçe fontları kaydet"""
    try:
        # Windows font yollarını kontrol et
        font_paths = [
            "C:\\Windows\\Fonts\\arial.ttf",
            "C:\\Windows\\Fonts\\calibri.ttf",
            "C:\\Windows\\Fonts\\tahoma.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",  # Linux fallback
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf" # Linux fallback
        ]
        
        font_path = None
        for path in font_paths:
            if os.path.exists(path):
                font_path = path
                break
        
        if font_path:
            pdfmetrics.registerFont(TTFont('NotoSans', font_path))
            # Bold versiyonu varsa onu da kaydet, yoksa normali kullan
            bold_path = font_path.replace(".ttf", "bd.ttf").replace("-Regular", "-Bold")
            if os.path.exists(bold_path):
                 pdfmetrics.registerFont(TTFont('NotoSans-Bold', bold_path))
            else:
                 pdfmetrics.registerFont(TTFont('NotoSans-Bold', font_path))
            return True
    except Exception as e:
        logging.warning(f"Font yükleme hatası: {e}")
    return False

# --- Türkçe Karakter Desteği (Python-Docx) ---
def _add_turkish_paragraph(doc, text='', style='Normal'):
    p = doc.add_paragraph(text, style=style)
    # Font ayarlarını zorla
    for run in p.runs:
        run.font.name = 'Arial'
    return p

def _add_turkish_heading(doc, text='', level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)  # Koyu mavi
    return h


class SDGReporting(BaseTenantManager):
    def __init__(self, db_path: str = None):
        super().__init__(db_path)
        self._ensure_reporting_schema()

    def _ensure_reporting_schema(self):
        """Raporlama için gerekli tabloları oluştur"""
        try:
            # Rapor şablonları tablosu
            self.execute_update("""
            CREATE TABLE IF NOT EXISTS report_templates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                template_key TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                description TEXT,
                formats TEXT DEFAULT '["PDF"]',  -- JSON array: ["PDF", "DOCX", "EXCEL"]
                content_path TEXT,               -- Özel şablon dosyası yolu (varsa)
                created_at TEXT DEFAULT (datetime('now')),
                updated_at TEXT DEFAULT (datetime('now')),
                is_active BOOLEAN DEFAULT 1
            )
            """, params=(), skip_tenant_filter=True)

            # Rapor geçmişi tablosu
            self.execute_update("""
            CREATE TABLE IF NOT EXISTS report_history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER NOT NULL,
                template_key TEXT NOT NULL,
                output_format TEXT NOT NULL,
                output_path TEXT NOT NULL,
                created_at TEXT DEFAULT (datetime('now')),
                meta TEXT,             -- JSON: ek bilgiler
                FOREIGN KEY (company_id) REFERENCES companies(id) ON DELETE CASCADE
            )
            """, params=(), skip_tenant_filter=True)
        except Exception as e:
            logging.error(f"Raporlama şeması oluşturulurken hata: {e}")

    def _get_logo_path(self, company_id: int) -> Optional[str]:
        """Şirket logosunun yolunu bul"""
        try:
            # 1. Veritabanından logo yolunu al
            result = self.execute_query("SELECT logo_path FROM company_profiles WHERE company_id = ?", (company_id,))
            
            if result and result[0]['logo_path'] and os.path.exists(result[0]['logo_path']):
                return result[0]['logo_path']
            
            # 2. Varsayılan yolda ara
            data_dir = os.path.dirname(self.db_path)
            possible_path = os.path.join(data_dir, "company_logos", f"company_{company_id}_logo.png")
            if os.path.exists(possible_path):
                return possible_path
            
            # jpg dene
            possible_path = possible_path.replace(".png", ".jpg")
            if os.path.exists(possible_path):
                return possible_path
                
        except Exception as e:
            logging.warning(f"Logo yolu aranırken hata: {e}")
            return None
        return None

    def _add_logo(self, doc, company_id: int):
        """Raporun başına şirket logosunu ekle"""
        try:
            logo_path = self._get_logo_path(company_id)

            if logo_path:
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, height=Cm(2.5))
        except Exception as e:
            logging.warning(f"Logo eklenirken hata: {e}")

    # --- Yardımcı: TSRS ve Doğrulama Özetleri ---
    def _compute_tsrs_summary(self, company_id: int) -> Dict:
        """Seçili SDG’ler ve bunlardan türeyen GRI açıklamaları üzerinden TSRS bölüm/metrik özetini çıkar."""
        # İlgili SDG gösterge kodları: company için yanıtlanmış indicator_code’lar
        sdg_indicator_codes_rows = self.execute_query(
            """
            SELECT DISTINCT si.code
            FROM sdg_indicators si
            JOIN sdg_targets st ON si.target_id=st.id
            JOIN sdg_goals sg ON st.goal_id=sg.id
            JOIN sdg_question_bank qb ON sg.code = qb.sdg_no AND si.code = qb.indicator_code
            JOIN sdg_question_responses qr ON qb.id = qr.question_id AND qr.company_id = ?
            WHERE sg.id IN (SELECT goal_id FROM user_sdg_selections WHERE company_id = ?)
            """,
            (company_id, company_id), company_id=company_id
        )
        sdg_indicator_codes = [row['code'] for row in sdg_indicator_codes_rows]

        # SDG→GRI açıklamaları
        gri_disclosures = []
        if sdg_indicator_codes:
            ph = ','.join('?' * len(sdg_indicator_codes))
            gri_disclosures_rows = self.execute_query(
                f"SELECT DISTINCT gri_disclosure FROM map_sdg_gri WHERE sdg_indicator_code IN ({ph})",
                tuple(sdg_indicator_codes), company_id=company_id
            )
            gri_disclosures = [row['gri_disclosure'] for row in gri_disclosures_rows]

        # TSRS kayıtları
        tsrs_rows_sdg = []
        tsrs_rows_gri = []
        if sdg_indicator_codes:
            ph = ','.join('?' * len(sdg_indicator_codes))
            tsrs_rows_sdg = self.execute_query(
                f"SELECT tsrs_section, tsrs_metric FROM map_sdg_tsrs WHERE sdg_indicator_code IN ({ph})",
                tuple(sdg_indicator_codes), company_id=company_id
            )

        if gri_disclosures:
            ph = ','.join('?' * len(gri_disclosures))
            tsrs_rows_gri = self.execute_query(
                f"SELECT tsrs_section, tsrs_metric FROM map_gri_tsrs WHERE gri_disclosure IN ({ph})",
                tuple(gri_disclosures), company_id=company_id
            )

        from collections import defaultdict
        section_metrics = defaultdict(set)
        for row in tsrs_rows_sdg + tsrs_rows_gri:
            sec = row['tsrs_section']
            met = row['tsrs_metric']
            if sec:
                section_metrics[sec].add(met)
        summary = []
        total = 0
        for sec, metrics in sorted(section_metrics.items()):
            cnt = len(metrics)
            total += cnt
            sample = ", ".join(sorted(list(metrics))[:8])
            summary.append((sec, cnt, sample, metrics))
        return {
            'total_tsrs_metrics': total,
            'by_section': summary
        }

    def _build_methodology_and_validation(self, company_id: int) -> Dict[str, str]:
        """Metodoloji metni ve doğrulama özeti üret."""
        metodoloji = (
            "Bu rapor, şirketin SDG soru yanıtları ve destekleyici evidansları "
            "temel alınarak hazırlanmıştır. İlerleme oranları soru bazında "
            "tamamlanma üzerinden hesaplanmış, eşleştirmeler resmi SDG-GRI ve TSRS "
            "haritaları üzerinden türetilmiştir."
        )
        kapsam = (
            "Kapsam: Değerlendirme, seçilen dönem ve mevcut veri kaynakları ile sınırlıdır. "
            "Birim ve sorumluluklar şirket beyanına dayanmaktadır."
        )

        # Doğrulama bulguları (severity sayımı)
        sev_counts_rows = self.execute_query("""
            SELECT COALESCE(severity_level,'info') AS sev, COUNT(*) as cnt
            FROM sdg_validation_results
            WHERE company_id = ?
            GROUP BY sev
        """, (company_id,), company_id=company_id)
        
        sev_counts = {row['sev']: row['cnt'] for row in sev_counts_rows}
        sev_order = ['error', 'warning', 'info']
        sev_text = ", ".join([f"{s}: {sev_counts.get(s, 0)}" for s in sev_order]) or "Doğrulama bulgusu kaydı bulunamadı."

        # Kalite skorları (ortalama)
        qs_rows = self.execute_query("""
            SELECT AVG(completeness_score) as comp, AVG(accuracy_score) as acc, AVG(consistency_score) as cons,
                   AVG(timeliness_score) as time, AVG(overall_quality_score) as overall
            FROM sdg_data_quality_scores
            WHERE company_id = ?
        """, (company_id,), company_id=company_id)
        
        qs = qs_rows[0] if qs_rows else {}
        comp = round(qs.get('comp') or 0, 1)
        acc = round(qs.get('acc') or 0, 1)
        cons = round(qs.get('cons') or 0, 1)
        time = round(qs.get('time') or 0, 1)
        overall = round(qs.get('overall') or 0, 1)
        
        kalite = (
            f"Kalite Skorları (ortalama): Eksiksizlik %{comp}, Doğruluk %{acc}, "
            f"Tutarlılık %{cons}, Güncellik %{time}, Genel %{overall}."
        )

        kaynaklar = []
        try:
            row_list = self.execute_query("SELECT data_sources, governance_notes, assurance_statement FROM company_info WHERE company_id=?", (company_id,), company_id=company_id)
            if row_list:
                row = row_list[0]
                ds, gov, ass = row['data_sources'], row['governance_notes'], row['assurance_statement']
                if ds:
                    txt = f"Veri Kaynakları: {str(ds)[:300]}"
                    if len(str(ds)) > 300:
                        txt += "…"
                    kaynaklar.append(txt)
                if gov:
                    txt = f"Yönetişim: {str(gov)[:300]}"
                    if len(str(gov)) > 300:
                        txt += "…"
                    kaynaklar.append(txt)
                if ass:
                    txt = f"Beyan/Assurance: {str(ass)[:300]}"
                    if len(str(ass)) > 300:
                        txt += "…"
                    kaynaklar.append(txt)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
        kaynak_text = "\n".join(kaynaklar) if kaynaklar else "Veri kaynakları ve yönetişim beyanları sisteme girilmemiş."

        full_text = (
            "Metodoloji:\n" + metodoloji + "\n\n" +
            "Kapsam:\n" + kapsam + "\n\n" +
            "Doğrulama Özeti:\n" + sev_text + "\n" + kalite + "\n\n" +
            kaynak_text
        )
        return {"metodoloji_ve_dogrulama": full_text}

    # --- Şema Yardımcıları (Şablonlar ve Geçmiş) ---
    # _ensure_reporting_schema metodu __init__ içinde çağrılıyor ve BaseTenantManager üzerinden çalışıyor.


    def get_company_data(self, company_id: int) -> Dict:
        """Şirket verilerini getir"""
        try:
            # Şirket bilgileri
            companies = self.execute_query("SELECT name FROM companies WHERE id = ?", (company_id,), company_id=company_id)
            
            if not companies:
                return {'error': 'Şirket bulunamadı'}
            
            company = companies[0]

            # SDG seçimleri
            selected_goals = self.execute_query("""
                SELECT sg.code, sg.title_tr, uss.selected_at
                FROM user_sdg_selections uss
                JOIN sdg_goals sg ON uss.goal_id = sg.id
                WHERE uss.company_id = ?
                ORDER BY sg.code
            """, (company_id,), company_id=company_id)

            # İlerleme verileri
            progress_data = self.execute_query("""
                SELECT sg.code, sg.title_tr, 
                       COUNT(DISTINCT qr.question_id) as answered_questions,
                       COUNT(DISTINCT qb.id) as total_questions,
                       ROUND(COUNT(DISTINCT qr.question_id) * 100.0 / COUNT(DISTINCT qb.id), 2) as completion_percentage
                FROM sdg_goals sg
                LEFT JOIN sdg_targets st ON sg.id = st.goal_id
                LEFT JOIN sdg_indicators si ON st.id = si.target_id
                LEFT JOIN sdg_question_bank qb ON sg.code = qb.sdg_no AND si.code = qb.indicator_code
                LEFT JOIN sdg_question_responses qr ON qb.id = qr.question_id AND qr.company_id = ?
                WHERE sg.id IN (SELECT goal_id FROM user_sdg_selections WHERE company_id = ?)
                GROUP BY sg.code, sg.title_tr
                ORDER BY sg.code
            """, (company_id, company_id), company_id=company_id)

            # GRI eşleştirmeleri
            gri_mappings = self.execute_query("""
                SELECT sg.code, COUNT(DISTINCT mg.gri_standard) as gri_count
                FROM sdg_goals sg
                LEFT JOIN sdg_targets st ON sg.id = st.goal_id
                LEFT JOIN sdg_indicators si ON st.id = si.target_id
                LEFT JOIN map_sdg_gri mg ON si.code = mg.sdg_indicator_code
                WHERE sg.id IN (SELECT goal_id FROM user_sdg_selections WHERE company_id = ?)
                GROUP BY sg.code
                ORDER BY sg.code
            """, (company_id,), company_id=company_id)

            return {
                'company': {
                    'name': company['name'],
                    'description': 'SDG Yönetim Sistemi'
                },
                'selected_goals': selected_goals,
                'progress_data': progress_data,
                'gri_mappings': gri_mappings,
                'report_date': datetime.now().strftime('%d.%m.%Y'),
                'report_time': datetime.now().strftime('%H:%M')
            }

        except Exception as e:
            logging.error(f"Şirket verileri getirilirken hata: {e}")
            return {'error': str(e)}

    # --- SDG 6 Su Metrikleri Yardımcıları ---
    def _get_sdg6_water_metrics(self, company_id: int, period: Optional[str] = None) -> Optional[Dict]:
        """WaterManager üzerinden SDG 6 su ayak izi ve verimlilik metriklerini getir.
           SDG 6.4: Su kullanım verimliliği ve tatlı su arzı
        """
        try:
            from modules.environmental.water_manager import WaterManager
        except Exception as e:
            logging.error(f"WaterManager import hatası: {e}")
            return None
        wm = WaterManager(db_path=self.db_path)
        # Öncelik: yıl bazlı (örn. 2025)
        year_period = datetime.now().strftime('%Y')
        metrics = wm.calculate_water_footprint(company_id, year_period) or {}
        if metrics.get('total_water_footprint', 0) == 0:
            # Yıl bazlı veri yoksa son kayıt dönemine göre dene
            try:
                records = wm.get_water_consumption(company_id)
                if records:
                    last_period = records[0].get('period')
                    if last_period:
                        metrics = wm.calculate_water_footprint(company_id, last_period) or {}
            except Exception as e:
                logging.error(f"Su kayıtları getirilirken hata: {e}")
        return metrics if metrics else None

    def create_pdf_report(self, company_id: int, output_path: str) -> bool:
        """PDF raporu oluştur"""
        try:
            # Verileri al
            data = self.get_company_data(company_id)
            if 'error' in data:
                logging.error(f"Veri alınırken hata: {data['error']}")
                return False

            # PDF oluştur
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            story = []

            # Stil tanımları
            styles = getSampleStyleSheet()
            # Türkçe fontları kaydet
            register_turkish_fonts_reportlab()

            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=25,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.darkblue,
                fontName='NotoSans'
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=17,
                spaceAfter=12,
                textColor=colors.darkgreen,
                fontName='NotoSans'
            )

            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                fontName='NotoSans'
            )

            # Logo
            logo_path = self._get_logo_path(company_id)
            if logo_path:
                try:
                    # Logoyu sol üst köşeye ekle (ya da ortaya)
                    # Aspect ratio'yu koruyarak max 1.5 inch yüksekliğe sığdır
                    img = Image(logo_path, width=1.5*inch, height=1.5*inch, kind='proportional')
                    img.hAlign = 'CENTER' 
                    story.append(img)
                    story.append(Spacer(1, 10))
                except Exception as e:
                    logging.warning(f"PDF'e logo eklenirken hata: {e}")

            # Başlık
            story.append(Paragraph("SDG İlerleme Raporu", title_style))
            story.append(Spacer(1, 20))

            # Şirket bilgileri
            story.append(Paragraph("Şirket Bilgileri", heading_style))
            company_info = f"""
            <b>Şirket Adı:</b> {data['company']['name']}<br/>
            <b>Açıklama:</b> {data['company']['description'] or 'Belirtilmemiş'}<br/>
            <b>Rapor Tarihi:</b> {data['report_date']} {data['report_time']}
            """
            story.append(Paragraph(company_info, normal_style))
            story.append(Spacer(1, 20))

            # TSRS eşleştirmeleri
            tsrs = self._compute_tsrs_summary(company_id)
            story.append(Paragraph("TSRS Eşleştirmeleri", heading_style))
            if tsrs.get('by_section'):
                tsrs_data = [['TSRS Bölüm', 'Metrik Sayısı', 'Örnek Metrikler']]
                for sec, cnt, sample, _metrics in tsrs['by_section']:
                    tsrs_data.append([sec, str(cnt), sample])
                tsrs_table = Table(tsrs_data, colWidths=[1.2*inch, 1*inch, 3.2*inch])
                tsrs_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(tsrs_table)
                story.append(Paragraph(f"Toplam TSRS metrik bağlantısı: {tsrs.get('total_tsrs_metrics', 0)}", normal_style))
            else:
                story.append(Paragraph("TSRS eşleştirme verisi bulunmamaktadır.", normal_style))

            story.append(Spacer(1, 20))

            # Metodoloji ve Veri Kalitesi
            mv = self._build_methodology_and_validation(company_id)
            story.append(Paragraph("Metodoloji ve Veri Kalitesi", heading_style))
            for line in mv.get('metodoloji_ve_dogrulama', '').split('\n'):
                if line.strip():
                    story.append(Paragraph(line, normal_style))


            # Seçilen SDG'ler
            story.append(Paragraph("Seçilen SDG Hedefleri", heading_style))
            if data['selected_goals']:
                goal_data = [['SDG No', 'Hedef Başlığı', 'Seçim Tarihi']]
                for goal in data['selected_goals']:
                    goal_data.append([
                        f"SDG {goal[0]}",
                        goal[1],
                        goal[2][:10] if goal[2] else 'Bilinmiyor'
                    ])

                goal_table = Table(goal_data, colWidths=[1*inch, 3*inch, 1.5*inch])
                goal_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(goal_table)
            else:
                story.append(Paragraph("Seçilen SDG hedefi bulunmamaktadır.", normal_style))

            story.append(Spacer(1, 20))

            # İlerleme durumu
            story.append(Paragraph("İlerleme Durumu", heading_style))
            if data['progress_data']:
                progress_data = [['SDG No', 'Hedef Başlığı', 'Cevaplanan', 'Toplam', 'Tamamlanma %']]
                for progress in data['progress_data']:
                    progress_data.append([
                        f"SDG {progress[0]}",
                        progress[1][:30] + "..." if len(progress[1]) > 30 else progress[1],
                        str(progress[2]),
                        str(progress[3]),
                        f"%{progress[4]}"
                    ])

                progress_table = Table(progress_data, colWidths=[0.8*inch, 2.5*inch, 0.8*inch, 0.8*inch, 1*inch])
                progress_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(progress_table)
            else:
                story.append(Paragraph("İlerleme verisi bulunmamaktadır.", normal_style))

            story.append(Spacer(1, 20))

            # GRI eşleştirmeleri
            story.append(Paragraph("GRI Eşleştirmeleri", heading_style))
            if data['gri_mappings']:
                gri_data = [['SDG No', 'GRI Standart Sayısı']]
                for gri in data['gri_mappings']:
                    gri_data.append([f"SDG {gri[0]}", str(gri[1])])

                gri_table = Table(gri_data, colWidths=[1*inch, 2*inch])
                gri_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(gri_table)
            else:
                story.append(Paragraph("GRI eşleştirme verisi bulunmamaktadır.", normal_style))

            story.append(Spacer(1, 20))

            # SDG 6: Su Metrikleri
            story.append(Paragraph("SDG 6: Su Metrikleri", heading_style))
            water = self._get_sdg6_water_metrics(company_id)
            if water:
                # Özet metrikler tablosu
                water_summary = [
                    ['Metrik', 'Değer'],
                    ['Toplam Mavi Su (m³)', str(water.get('total_blue_water', 0))],
                    ['Toplam Yeşil Su (m³)', str(water.get('total_green_water', 0))],
                    ['Toplam Gri Su (m³)', str(water.get('total_grey_water', 0))],
                    ['Toplam Su Ayak İzi (m³)', str(water.get('total_water_footprint', 0))],
                    ['Ortalama Verimlilik Oranı', str(water.get('efficiency_metrics', {}).get('average_efficiency_ratio', 0))],
                    ['Ortalama Geri Dönüşüm Oranı', str(water.get('efficiency_metrics', {}).get('average_recycling_rate', 0))]
                ]
                ws_table = Table(water_summary, colWidths=[2.5*inch, 2.5*inch])
                ws_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(ws_table)

                # Kaynağa göre dağılım (varsa)
                bsrc = water.get('breakdown_by_source') or {}
                if bsrc:
                    story.append(Spacer(1, 10))
                    story.append(Paragraph("Su Kaynağına Göre Dağılım", heading_style))
                    src_rows = [['Kaynak', 'Toplam (m³)', 'Mavi', 'Yeşil', 'Gri']]
                    for src, vals in bsrc.items():
                        src_rows.append([
                            src,
                            str(round(vals.get('total', 0), 2)),
                            str(round(vals.get('blue_water', 0), 2)),
                            str(round(vals.get('green_water', 0), 2)),
                            str(round(vals.get('grey_water', 0), 2))
                        ])
                    src_table = Table(src_rows, colWidths=[1.8*inch, 1.1*inch, 0.8*inch, 0.8*inch, 0.8*inch])
                    src_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(src_table)
            else:
                story.append(Paragraph("Su metrikleri bulunamadı veya su verisi girilmemiş.", normal_style))

            story.append(Spacer(1, 20))

            # Özet
            story.append(Paragraph("Özet", heading_style))
            total_goals = len(data['selected_goals'])
            total_questions = sum(p[3] for p in data['progress_data']) if data['progress_data'] else 0
            answered_questions = sum(p[2] for p in data['progress_data']) if data['progress_data'] else 0
            overall_completion = round((answered_questions / total_questions * 100) if total_questions > 0 else 0, 2)

            summary_text = f"""
            <b>Toplam Seçilen SDG Hedefi:</b> {total_goals}<br/>
            <b>Toplam Soru Sayısı:</b> {total_questions}<br/>
            <b>Cevaplanan Soru Sayısı:</b> {answered_questions}<br/>
            <b>Genel Tamamlanma Oranı:</b> %{overall_completion}
            """
            story.append(Paragraph(summary_text, normal_style))

            # PDF'i oluştur
            doc.build(story)
            return True

        except Exception as e:
            logging.error(f"PDF raporu oluşturulurken hata: {e}")
            return False

    def create_docx_report(self, company_id: int, output_path: str) -> bool:
        """DOCX raporu oluştur"""
        try:
            # Verileri al
            data = self.get_company_data(company_id)
            if 'error' in data:
                logging.error(f"Veri alınırken hata: {data['error']}")
                return False

            # DOCX oluştur
            doc = Document()

            # Logo ekle
            self._add_logo(doc, company_id)

            # Başlık
            title = _add_turkish_heading(doc, 'SDG İlerleme Raporu', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Şirket bilgileri
            _add_turkish_heading(doc, 'Şirket Bilgileri', level=1)
            company_info = _add_turkish_paragraph(doc, )
            company_info.add_run('Şirket Adı: ').bold = True
            company_info.add_run(data['company']['name'])
            company_info.add_run('\nAçıklama: ').bold = True
            company_info.add_run(data['company']['description'] or 'Belirtilmemiş')
            company_info.add_run('\nRapor Tarihi: ').bold = True
            company_info.add_run(f"{data['report_date']} {data['report_time']}")

            # Seçilen SDG'ler
            _add_turkish_heading(doc, 'Seçilen SDG Hedefleri', level=1)
            if data['selected_goals']:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Başlık satırı
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'SDG No'
                hdr_cells[1].text = 'Hedef Başlığı'
                hdr_cells[2].text = 'Seçim Tarihi'

                # Veri satırları
                for goal in data['selected_goals']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"SDG {goal[0]}"
                    row_cells[1].text = goal[1]
                    row_cells[2].text = goal[2] if goal[2] else 'Bilinmiyor'
            else:
                _add_turkish_paragraph(doc, 'Seçilen SDG hedefi bulunmamaktadır.')

            # İlerleme durumu
            _add_turkish_heading(doc, 'İlerleme Durumu', level=1)
            if data['progress_data']:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Başlık satırı
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'SDG No'
                hdr_cells[1].text = 'Hedef Başlığı'
                hdr_cells[2].text = 'Cevaplanan'
                hdr_cells[3].text = 'Toplam'
                hdr_cells[4].text = 'Tamamlanma %'

                # Veri satırları
                for progress in data['progress_data']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"SDG {progress[0]}"
                    row_cells[1].text = progress[1]
                    row_cells[2].text = str(progress[2])
                    row_cells[3].text = str(progress[3])
                    row_cells[4].text = f"%{progress[4]}"
            else:
                _add_turkish_paragraph(doc, 'İlerleme verisi bulunmamaktadır.')

            # GRI eşleştirmeleri
            _add_turkish_heading(doc, 'GRI Eşleştirmeleri', level=1)
            if data['gri_mappings']:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Başlık satırı
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'SDG No'
                hdr_cells[1].text = 'GRI Standart Sayısı'

                # Veri satırları
                for gri in data['gri_mappings']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"SDG {gri[0]}"
                    row_cells[1].text = str(gri[1])
            else:
                _add_turkish_paragraph(doc, 'GRI eşleştirme verisi bulunmamaktadır.')

            # TSRS eşleştirmeleri
            _add_turkish_heading(doc, 'TSRS Eşleştirmeleri', level=1)
            try:
                tsrs = self._compute_tsrs_summary(company_id)
                if tsrs.get('by_section'):
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hdr = table.rows[0].cells
                    hdr[0].text = 'TSRS Bölüm'
                    hdr[1].text = 'Metrik Sayısı'
                    hdr[2].text = 'Örnek Metrikler'
                    for sec, cnt, sample, _metrics in tsrs['by_section']:
                        row = table.add_row().cells
                        row[0].text = sec
                        row[1].text = str(cnt)
                        row[2].text = sample
                    _add_turkish_paragraph(doc, f"Toplam TSRS metrik bağlantısı: {tsrs.get('total_tsrs_metrics', 0)}")
                else:
                    _add_turkish_paragraph(doc, 'TSRS eşleştirme verisi bulunmamaktadır.')
            except Exception as e:
                logging.error(f"TSRS özeti oluşturulurken hata: {e}")

            # Metodoloji ve Veri Kalitesi
            _add_turkish_heading(doc, 'Metodoloji ve Veri Kalitesi', level=1)
            try:
                mv = self._build_methodology_and_validation(company_id)
                for line in mv.get('metodoloji_ve_dogrulama', '').split('\n'):
                    if line.strip():
                        p = _add_turkish_paragraph(doc, )
                        p.add_run(line)
            except Exception as e:
                logging.error(f"Metodoloji özeti oluşturulurken hata: {e}")

            # SDG 6: Su Metrikleri
            _add_turkish_heading(doc, 'SDG 6: Su Metrikleri', level=1)
            water = self._get_sdg6_water_metrics(company_id)
            if water:
                # Özet metrikler tablosu
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr = table.rows[0].cells
                hdr[0].text = 'Metrik'
                hdr[1].text = 'Değer'
                rows = [
                    ('Toplam Mavi Su (m³)', str(water.get('total_blue_water', 0))),
                    ('Toplam Yeşil Su (m³)', str(water.get('total_green_water', 0))),
                    ('Toplam Gri Su (m³)', str(water.get('total_grey_water', 0))),
                    ('Toplam Su Ayak İzi (m³)', str(water.get('total_water_footprint', 0))),
                    ('Ortalama Verimlilik Oranı', str(water.get('efficiency_metrics', {}).get('average_efficiency_ratio', 0))),
                    ('Ortalama Geri Dönüşüm Oranı', str(water.get('efficiency_metrics', {}).get('average_recycling_rate', 0)))
                ]
                for met, val in rows:
                    r = table.add_row().cells
                    r[0].text = met
                    r[1].text = val

                # Kaynağa göre dağılım varsa ek bir tablo
                bsrc = water.get('breakdown_by_source') or {}
                if bsrc:
                    table2 = doc.add_table(rows=1, cols=5)
                    table2.style = 'Table Grid'
                    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
                    h2 = table2.rows[0].cells
                    h2[0].text = 'Kaynak'
                    h2[1].text = 'Toplam (m³)'
                    h2[2].text = 'Mavi'
                    h2[3].text = 'Yeşil'
                    h2[4].text = 'Gri'
                    for src, vals in bsrc.items():
                        row = table2.add_row().cells
                        row[0].text = src
                        row[1].text = str(round(vals.get('total', 0), 2))
                        row[2].text = str(round(vals.get('blue_water', 0), 2))
                        row[3].text = str(round(vals.get('green_water', 0), 2))
                        row[4].text = str(round(vals.get('grey_water', 0), 2))
            else:
                _add_turkish_paragraph(doc, 'Su metrikleri bulunamadı veya su verisi girilmemiş.')

            # Özet
            _add_turkish_heading(doc, 'Özet', level=1)
            total_goals = len(data['selected_goals'])
            total_questions = sum(p[3] for p in data['progress_data']) if data['progress_data'] else 0
            answered_questions = sum(p[2] for p in data['progress_data']) if data['progress_data'] else 0
            overall_completion = round((answered_questions / total_questions * 100) if total_questions > 0 else 0, 2)

            summary = _add_turkish_paragraph(doc, )
            summary.add_run('Toplam Seçilen SDG Hedefi: ').bold = True
            summary.add_run(str(total_goals))
            summary.add_run('\nToplam Soru Sayısı: ').bold = True
            summary.add_run(str(total_questions))
            summary.add_run('\nCevaplanan Soru Sayısı: ').bold = True
            summary.add_run(str(answered_questions))
            summary.add_run('\nGenel Tamamlanma Oranı: ').bold = True
            summary.add_run(f"%{overall_completion}")

            # DOCX'i kaydet
            doc.save(output_path)
            return True

        except Exception as e:
            logging.error(f"DOCX raporu oluşturulurken hata: {e}")
            return False

    def create_progress_chart(self, company_id: int, output_path: str) -> bool:
        """İlerleme grafiği oluştur"""
        try:
            # Verileri al
            data = self.get_company_data(company_id)
            if 'error' in data or not data['progress_data']:
                return False

            # Grafik oluştur
            fig, ax = plt.subplots(figsize=(12, 8))

            sdg_numbers = [f"SDG {p[0]}" for p in data['progress_data']]
            completion_rates = [p[4] for p in data['progress_data']]

            # Bar grafik
            bars = ax.bar(sdg_numbers, completion_rates, color='skyblue', edgecolor='navy', alpha=0.7)

            # Grafik ayarları
            ax.set_xlabel('SDG Hedefleri', fontsize=12, fontweight='bold')
            ax.set_ylabel('Tamamlanma Oranı (%)', fontsize=12, fontweight='bold')
            ax.set_title('SDG İlerleme Durumu', fontsize=16, fontweight='bold', pad=20)
            ax.set_ylim(0, 100)

            # Değerleri bar üzerine yaz
            for bar, rate in zip(bars, completion_rates):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 1,
                       f'%{rate}', ha='center', va='bottom', fontweight='bold')

            # Grid
            ax.grid(True, alpha=0.3)
            ax.set_axisbelow(True)

            # X ekseni etiketlerini döndür
            plt.xticks(rotation=45, ha='right')

            # Layout ayarla
            plt.tight_layout()

            # Grafiği kaydet
            plt.savefig(output_path, dpi=300, bbox_inches='tight')
            plt.close()

            return True

        except Exception as e:
            logging.error(f"Grafik oluşturulurken hata: {e}")
            return False

    def get_report_templates(self) -> List[Dict]:
        """Rapor şablonlarını getir (DB’den veya yoksa varsayılanları yükle)"""
        self._ensure_reporting_schema()
        
        rows = self.execute_query("SELECT template_key, name, description, formats, content_path, is_active FROM report_templates ORDER BY name", skip_tenant_filter=True)
        if not rows:
            # Varsayılan şablonları seed et
            defaults = [
                ('basic', 'Temel Rapor', 'SDG ilerleme durumu ve temel bilgiler', json.dumps(['PDF','DOCX']), None, 1),
                ('detailed', 'Detaylı Rapor', 'Kapsamlı analiz ve grafiklerle detaylı rapor', json.dumps(['PDF','DOCX']), None, 1),
                ('executive', 'Yönetici Özeti', 'Yöneticiler için özet rapor', json.dumps(['PDF']), None, 1)
            ]
            for d in defaults:
                self.execute_update(
                    "INSERT OR IGNORE INTO report_templates (template_key, name, description, formats, content_path, is_active) VALUES (?,?,?,?,?,?)",
                    d, skip_tenant_filter=True
                )
            rows = self.execute_query("SELECT template_key, name, description, formats, content_path, is_active FROM report_templates ORDER BY name", skip_tenant_filter=True)
            
        templates = []
        for row in rows:
            # row is sqlite3.Row, supports index and key
            tkey = row['template_key']
            name = row['name']
            desc = row['description']
            formats_json = row['formats']
            cpath = row['content_path']
            active = row['is_active']
            
            try:
                formats = json.loads(formats_json) if formats_json else []
            except Exception:
                formats = ['PDF']
            templates.append({
                'id': tkey,
                'name': name,
                'description': desc or '',
                'formats': formats,
                'content_path': cpath,
                'is_active': bool(active)
            })
        return templates

    def add_report_template(self, template_key: str, name: str, description: str = '', formats: Optional[List[str]] = None, content_path: Optional[str] = None, is_active: bool = True) -> Optional[int]:
        self._ensure_reporting_schema()
        try:
            row_id = self.execute_update(
                "INSERT INTO report_templates (template_key, name, description, formats, content_path, is_active) VALUES (?,?,?,?,?,?)",
                (template_key, name, description, json.dumps(formats or ['PDF']), content_path, 1 if is_active else 0),
                skip_tenant_filter=True
            )
            return row_id
        except sqlite3.IntegrityError:
            return None

    def update_report_template(self, template_key: str, **updates) -> bool:
        self._ensure_reporting_schema()
        fields = []
        values = []
        for k, v in updates.items():
            if k == 'formats' and isinstance(v, list):
                v = json.dumps(v)
            fields.append(f"{k} = ?")
            values.append(v)
        if not fields:
            return False
        values.append(template_key)
        
        # execute_update returns lastrowid for INSERT, but rowcount isn't directly returned by execute_update in BaseTenantManager usually?
        # Let's check BaseTenantManager.execute_update.
        # It calls db_manager.execute_update which returns cursor.lastrowid usually.
        # But for UPDATE, we want rowcount.
        # BaseTenantManager execute_update returns cursor.lastrowid.
        # However, if I want rowcount, I might need execute_query or check if changed.
        # Or I can assume success if no error.
        
        # Checking BaseTenantManager code (I don't have it open but usually it returns lastrowid).
        # I'll just return True if no exception.
        try:
            self.execute_update(f"UPDATE report_templates SET {', '.join(fields)}, updated_at = datetime('now') WHERE template_key = ?", tuple(values), skip_tenant_filter=True)
            return True
        except Exception:
            return False

    def delete_report_template(self, template_key: str) -> bool:
        self._ensure_reporting_schema()
        try:
            self.execute_update("DELETE FROM report_templates WHERE template_key = ?", (template_key,), skip_tenant_filter=True)
            return True
        except Exception:
            return False

    # --- Geçmiş CRUD ---
    def add_report_history(self, company_id: int, template_key: str, output_format: str, output_path: str, meta: Optional[Dict] = None) -> Optional[int]:
        self._ensure_reporting_schema()
        try:
            rid = self.execute_update(
                "INSERT INTO report_history (company_id, template_key, output_format, output_path, meta) VALUES (?,?,?,?,?)",
                (company_id, template_key, output_format, output_path, json.dumps(meta or {})),
                skip_tenant_filter=True
            )
            return rid
        except Exception as e:
            logging.error(f"Rapor gecmisi ekleme hatasi: {e}")
            return None

    def get_report_history(self, company_id: Optional[int] = None) -> List[Dict]:
        self._ensure_reporting_schema()
        sql = "SELECT id, company_id, template_key, output_format, output_path, created_at, meta FROM report_history"
        params = ()
        if company_id is not None:
            sql += " WHERE company_id = ?"
            params = (company_id,)
        
        sql += " ORDER BY created_at DESC"
        
        try:
            rows = self.execute_query(sql, params, skip_tenant_filter=True)
            out = []
            for row in rows:
                try:
                    meta = json.loads(row['meta']) if row['meta'] else {}
                except Exception:
                    meta = {}
                out.append({
                    'id': row['id'],
                    'company_id': row['company_id'],
                    'template_key': row['template_key'],
                    'output_format': row['output_format'],
                    'output_path': row['output_path'],
                    'created_at': row['created_at'],
                    'meta': meta
                })
            return out
        except Exception as e:
            logging.error(f"Rapor gecmisi getirme hatasi: {e}")
            return []

    def delete_report_history(self, history_id: int) -> bool:
        self._ensure_reporting_schema()
        try:
            self.execute_update("DELETE FROM report_history WHERE id = ?", (history_id,), skip_tenant_filter=True)
            return True
        except Exception:
            return False

    def purge_report_history(self, company_id: Optional[int] = None) -> int:
        self._ensure_reporting_schema()
        sql = "DELETE FROM report_history"
        params = ()
        if company_id is not None:
            sql += " WHERE company_id = ?"
            params = (company_id,)
        
        try:
            self.execute_update(sql, params, skip_tenant_filter=True)
            return 1
        except Exception:
            return 0

if __name__ == "__main__":
    # Test
    reporting = SDGReporting()
    logging.info("SDG Raporlama Sistemi başlatıldı")

    # Test raporu oluştur
    success = reporting.create_pdf_report(1, "test_sdg_report.pdf")
    logging.info(f"PDF raporu oluşturuldu: {success}")

    success = reporting.create_docx_report(1, "test_sdg_report.docx")
    logging.info(f"DOCX raporu oluşturuldu: {success}")
