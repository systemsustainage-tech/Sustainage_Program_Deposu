import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SDG Veri Doğrulama GUI
Veri doğrulama ve kalite kontrol arayüzü
"""

import json
import os
import tkinter as tk
from datetime import datetime
from tkinter import messagebox, ttk

from utils.language_manager import LanguageManager

from .sdg_data_validation import SDGDataValidation


class SDGDataValidationGUI:
    """SDG Veri Doğrulama GUI"""

    def __init__(self, parent, company_id: int) -> None:
        self.parent = parent
        self.company_id = company_id
        self.data_validation = SDGDataValidation()
        self.lm = LanguageManager()

        try:
            self.parent.winfo_toplevel().state('zoomed')
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

        self.setup_ui()
        self.load_validation_rules()

    def setup_ui(self) -> None:
        """Veri doğrulama arayüzünü oluştur"""
        # Ana frame
        main_frame = tk.Frame(self.parent, bg='#f8f9fa')
        main_frame.pack(fill='both', expand=True, padx=15, pady=15)

        # Başlık
        header_frame = tk.Frame(main_frame, bg='#2c3e50', height=70)
        header_frame.pack(fill='x', pady=(0, 15))
        header_frame.pack_propagate(False)

        # Geri butonu
        back_btn = tk.Button(header_frame, text=f" {self.lm.tr('btn_back', 'Geri')}",
                             font=('Segoe UI', 10, 'bold'), bg='#1f2a36', fg='white',
                             relief='flat', bd=0, padx=15, pady=8,
                             command=self.back_to_sdg)
        back_btn.pack(side='left', padx=15, pady=15)

        title_frame = tk.Frame(header_frame, bg='#2c3e50')
        title_frame.pack(side='left', padx=20, pady=15)

        title_label = tk.Label(title_frame, text=self.lm.tr('title_sdg_data_validation', "SDG Veri Doğrulama"),
                              font=('Segoe UI', 16, 'bold'), fg='white', bg='#2c3e50')
        title_label.pack(side='left')

        subtitle_label = tk.Label(title_frame, text=self.lm.tr('subtitle_sdg_data_validation', "Veri kalitesi kontrolü ve doğrulama"),
                                 font=('Segoe UI', 11), fg='#bdc3c7', bg='#2c3e50')
        subtitle_label.pack(side='left', padx=(10, 0))

        # Ana içerik - Tabbed interface
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill='both', expand=True)

        # Doğrulama Kuralları sekmesi
        self.rules_frame = tk.Frame(self.notebook, bg='#f8f9fa')
        self.notebook.add(self.rules_frame, text=f" {self.lm.tr('tab_validation_rules', 'Doğrulama Kuralları')}")
        self.setup_rules_tab()

        # Veri Kontrolü sekmesi
        self.validation_frame = tk.Frame(self.notebook, bg='#f8f9fa')
        self.notebook.add(self.validation_frame, text=f" {self.lm.tr('tab_data_control', 'Veri Kontrolü')}")
        self.setup_validation_tab()

        # Kalite Skorları sekmesi
        self.quality_frame = tk.Frame(self.notebook, bg='#f8f9fa')
        self.notebook.add(self.quality_frame, text=f" {self.lm.tr('tab_quality_scores', 'Kalite Skorları')}")
        self.setup_quality_tab()

    def back_to_sdg(self) -> None:
        """SDG ana ekranına geri dön"""
        try:
            for widget in self.parent.winfo_children():
                widget.destroy()
            from .sdg_gui import SDGGUI
            SDGGUI(self.parent, self.company_id)
        except Exception as e:
            messagebox.showerror(self.lm.tr('msg_error', "Hata"), f"{self.lm.tr('msg_error_back', 'Geri dönme işleminde hata')}: {str(e)}")

    def setup_rules_tab(self) -> None:
        """Doğrulama kuralları sekmesini oluştur"""
        # Sol panel - Kural listesi
        left_panel = tk.Frame(self.rules_frame, bg='white', relief='raised', bd=1)
        left_panel.pack(side='left', fill='y', padx=(0, 10), pady=10, ipadx=10, ipady=10)

        tk.Label(left_panel, text=self.lm.tr('header_validation_rules', "Doğrulama Kuralları"), font=('Segoe UI', 12, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', pady=(0, 10))

        # Kural listesi
        self.rules_tree = ttk.Treeview(left_panel, columns=('name', 'type', 'status'),
                                      show='headings', height=15)

        self.rules_tree.heading('name', text=self.lm.tr('col_rule_name', 'Kural Adı'))
        self.rules_tree.heading('type', text=self.lm.tr('col_rule_type', 'Tür'))
        self.rules_tree.heading('status', text=self.lm.tr('col_rule_status', 'Durum'))

        self.rules_tree.column('name', width=200)
        self.rules_tree.column('type', width=100)
        self.rules_tree.column('status', width=80)

        self.rules_tree.pack(fill='both', expand=True, pady=(0, 10))

        # Scrollbar
        rules_scrollbar = ttk.Scrollbar(left_panel, orient="vertical", command=self.rules_tree.yview)
        self.rules_tree.configure(yscrollcommand=rules_scrollbar.set)
        rules_scrollbar.pack(side='right', fill='y')

        # Kural yönetimi butonları
        button_frame = tk.Frame(left_panel, bg='white')
        button_frame.pack(fill='x', pady=10)

        tk.Button(button_frame, text=f" {self.lm.tr('btn_new_rule', 'Yeni Kural')}", command=self.add_rule,
                 font=('Segoe UI', 10), bg='#27ae60', fg='white', relief='flat', padx=15).pack(side='left', padx=(0, 5))

        tk.Button(button_frame, text=f" {self.lm.tr('btn_edit_rule', 'Düzenle')}", command=self.edit_rule,
                 font=('Segoe UI', 10), bg='#3498db', fg='white', relief='flat', padx=15).pack(side='left', padx=5)

        tk.Button(button_frame, text=f" {self.lm.tr('btn_delete_rule', 'Sil')}", command=self.delete_rule,
                 font=('Segoe UI', 10), bg='#e74c3c', fg='white', relief='flat', padx=15).pack(side='left', padx=5)

        tk.Button(button_frame, text=f" {self.lm.tr('btn_toggle_active', 'Aktif/Pasif')}", command=self.toggle_rule_active,
                 font=('Segoe UI', 10), bg='#f39c12', fg='white', relief='flat', padx=15).pack(side='left', padx=5)

        # Sağ panel - Kural detayları
        right_panel = tk.Frame(self.rules_frame, bg='white', relief='raised', bd=1)
        right_panel.pack(side='right', fill='both', expand=True, padx=(10, 0), pady=10, ipadx=10, ipady=10)

        tk.Label(right_panel, text=self.lm.tr('header_rule_details', "Kural Detayları"), font=('Segoe UI', 12, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', pady=(0, 10))

        # Kural detayları metin alanı
        self.rule_details = tk.Text(right_panel, height=20, width=50,
                                   font=('Courier', 9), bg='#f8f9fa', wrap='word')
        self.rule_details.pack(fill='both', expand=True, pady=(0, 10))

        # Scrollbar
        details_scrollbar = ttk.Scrollbar(right_panel, orient="vertical", command=self.rule_details.yview)
        self.rule_details.configure(yscrollcommand=details_scrollbar.set)
        details_scrollbar.pack(side='right', fill='y')

        # Seçim değiştiğinde detayları göster
        self.rules_tree.bind('<<TreeviewSelect>>', self.on_rule_selected)

    def setup_validation_tab(self) -> None:
        """Veri kontrolü sekmesini oluştur"""
        # Kontrol paneli
        control_panel = tk.Frame(self.validation_frame, bg='white', relief='raised', bd=1)
        control_panel.pack(fill='x', padx=10, pady=10, ipadx=10, ipady=10)

        tk.Label(control_panel, text=self.lm.tr('header_data_control', "Veri Kontrolü"), font=('Segoe UI', 12, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', pady=(0, 10))

        # Kontrol seçenekleri
        options_frame = tk.Frame(control_panel, bg='white')
        options_frame.pack(fill='x', pady=(0, 10))

        self.check_completeness = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text=self.lm.tr('chk_completeness', "Eksik veri kontrolü"), variable=self.check_completeness,
                      bg='white', font=('Segoe UI', 10)).pack(side='left', padx=(0, 20))

        self.check_consistency = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text=self.lm.tr('chk_consistency', "Tutarlılık kontrolü"), variable=self.check_consistency,
                      bg='white', font=('Segoe UI', 10)).pack(side='left', padx=(0, 20))

        self.check_format = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text=self.lm.tr('chk_format', "Format kontrolü"), variable=self.check_format,
                      bg='white', font=('Segoe UI', 10)).pack(side='left')

        # Kontrol butonları
        button_frame = tk.Frame(control_panel, bg='white')
        button_frame.pack(fill='x', pady=10)

        tk.Button(button_frame, text=f" {self.lm.tr('btn_check', 'Kontrol Et')}", command=self.run_validation,
                 font=('Segoe UI', 12, 'bold'), bg='#3498db', fg='white', relief='flat', padx=20).pack(side='left', padx=(0, 10))

        tk.Button(button_frame, text=f" {self.lm.tr('btn_generate_report', 'Rapor Oluştur')}", command=self.generate_validation_report,
                 font=('Segoe UI', 12, 'bold'), bg='#27ae60', fg='white', relief='flat', padx=20).pack(side='left')

        # Sonuçlar paneli
        results_panel = tk.Frame(self.validation_frame, bg='white', relief='raised', bd=1)
        results_panel.pack(fill='both', expand=True, padx=10, pady=(0, 10), ipadx=10, ipady=10)

        tk.Label(results_panel, text=self.lm.tr('header_control_results', "Kontrol Sonuçları"), font=('Segoe UI', 12, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', pady=(0, 10))

        # Sonuçlar metin alanı
        self.validation_results = tk.Text(results_panel, height=15, width=80,
                                         font=('Courier', 9), bg='#f8f9fa', wrap='word')
        self.validation_results.pack(fill='both', expand=True)

        # Scrollbar
        results_scrollbar = ttk.Scrollbar(results_panel, orient="vertical", command=self.validation_results.yview)
        self.validation_results.configure(yscrollcommand=results_scrollbar.set)
        results_scrollbar.pack(side='right', fill='y')

    def setup_quality_tab(self) -> None:
        """Kalite skorları sekmesini oluştur"""
        # Kalite paneli
        quality_panel = tk.Frame(self.quality_frame, bg='white', relief='raised', bd=1)
        quality_panel.pack(fill='both', expand=True, padx=10, pady=10, ipadx=10, ipady=10)

        tk.Label(quality_panel, text=self.lm.tr('header_data_quality_scores', "Veri Kalite Skorları"), font=('Segoe UI', 12, 'bold'),
                bg='white', fg='#2c3e50').pack(anchor='w', pady=(0, 10))

        # Kalite skorları tablosu
        self.quality_tree = ttk.Treeview(quality_panel, columns=('sdg', 'completeness', 'consistency', 'accuracy', 'overall'),
                                        show='headings', height=15)

        self.quality_tree.heading('sdg', text=self.lm.tr('col_sdg_goal', 'SDG Hedefi'))
        self.quality_tree.heading('completeness', text=self.lm.tr('col_completeness', 'Eksiksizlik'))
        self.quality_tree.heading('consistency', text=self.lm.tr('col_consistency', 'Tutarlılık'))
        self.quality_tree.heading('accuracy', text=self.lm.tr('col_accuracy', 'Doğruluk'))
        self.quality_tree.heading('overall', text=self.lm.tr('col_overall_score', 'Genel Skor'))

        self.quality_tree.column('sdg', width=150)
        self.quality_tree.column('completeness', width=100)
        self.quality_tree.column('consistency', width=100)
        self.quality_tree.column('accuracy', width=100)
        self.quality_tree.column('overall', width=100)

        self.quality_tree.pack(fill='both', expand=True, pady=(0, 10))

        # Scrollbar
        quality_scrollbar = ttk.Scrollbar(quality_panel, orient="vertical", command=self.quality_tree.yview)
        self.quality_tree.configure(yscrollcommand=quality_scrollbar.set)
        quality_scrollbar.pack(side='right', fill='y')

        # Görünüm seçenekleri
        options_frame = tk.Frame(quality_panel, bg='white')
        options_frame.pack(fill='x', pady=(0, 6))
        self.quality_by_sdg = tk.BooleanVar(value=True)
        tk.Checkbutton(options_frame, text=self.lm.tr('chk_view_by_sdg', "SDG bazlı gösterim"), variable=self.quality_by_sdg, bg='white').pack(side='left')

        # Kalite yönetimi butonları
        button_frame = tk.Frame(quality_panel, bg='white')
        button_frame.pack(fill='x', pady=10)

        tk.Button(button_frame, text=f" {self.lm.tr('btn_refresh_scores', 'Skorları Yenile')}", command=self.refresh_quality_scores,
                 font=('Segoe UI', 10), bg='#3498db', fg='white', relief='flat', padx=15).pack(side='left', padx=(0, 5))

        tk.Button(button_frame, text=f" {self.lm.tr('btn_trend_analysis', 'Trend Analizi')}", command=self.show_trend_analysis,
                 font=('Segoe UI', 10), bg='#f39c12', fg='white', relief='flat', padx=15).pack(side='left', padx=5)

        tk.Button(button_frame, text=f" {self.lm.tr('btn_show_chart', 'Grafik Göster')}", command=self.show_quality_chart,
                 font=('Segoe UI', 10), bg='#9b59b6', fg='white', relief='flat', padx=15).pack(side='left', padx=5)

    def load_validation_rules(self) -> None:
        """Doğrulama kurallarını yükle"""
        try:
            # Mevcut verileri temizle
            for item in self.rules_tree.get_children():
                self.rules_tree.delete(item)

            # Veritabanından kuralları al
            rules = self.data_validation.get_rules()
            # Tree item -> rule_id eşlemesi
            self.rule_item_to_id = {}
            for row in rules:
                rule_id, rule_name, rule_type, is_active = row
                item_id = self.rules_tree.insert('', 'end', values=(rule_name, rule_type, 'Aktif' if is_active else 'Pasif'))
                self.rule_item_to_id[item_id] = rule_id

        except Exception as e:
            logging.error(f"{self.lm.tr('msg_error_loading_rules', 'Doğrulama kuralları yüklenirken hata')}: {e}")

    def add_rule(self) -> None:
        """Yeni kural ekle"""
        try:
            # Kural sihirbazı penceresi
            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('title_new_rule_wizard', "Yeni Kural Sihirbazı"))
            win.transient(self.parent)
            win.grab_set()
            frm = tk.Frame(win, bg='white', padx=10, pady=10)
            frm.pack(fill='both', expand=True)

            tk.Label(frm, text=self.lm.tr('lbl_rule_name', "Kural Adı"), bg='white').grid(row=0, column=0, sticky='w')
            name_var = tk.StringVar()
            tk.Entry(frm, textvariable=name_var, width=40).grid(row=0, column=1, sticky='w')

            tk.Label(frm, text=self.lm.tr('lbl_rule_type', "Tür"), bg='white').grid(row=1, column=0, sticky='w')
            rtype_var = tk.StringVar(value='completeness')
            rtype_cb = ttk.Combobox(frm, textvariable=rtype_var, values=['completeness', 'consistency', 'format'], state='readonly', width=37)
            rtype_cb.grid(row=1, column=1, sticky='w')

            tk.Label(frm, text=self.lm.tr('lbl_severity', "Ciddiyet"), bg='white').grid(row=2, column=0, sticky='w')
            sev_var = tk.StringVar(value='warning')
            sev_cb = ttk.Combobox(frm, textvariable=sev_var, values=['info', 'warning', 'error'], state='readonly', width=37)
            sev_cb.grid(row=2, column=1, sticky='w')

            tk.Label(frm, text=self.lm.tr('lbl_expression', "İfade (SQL WHERE)"), bg='white').grid(row=3, column=0, sticky='nw')
            expr_text = tk.Text(frm, height=5, width=50, bg='#f8f9fa')
            expr_text.grid(row=3, column=1, sticky='w')

            def preset_expression(*_) -> None:
                t = rtype_var.get()
                presets = {
                    'completeness': 'response_value IS NULL OR response_value = ""',
                    'consistency': 'sdg_no IS NULL OR indicator_code IS NULL',
                    'format': 'qt.type_name = "Sayı" AND (response_value IS NOT NULL AND (response_value NOT GLOB "-*[*0-9]*"))'
                }
                expr_text.delete('1.0', tk.END)
                expr_text.insert('1.0', presets.get(t, 'response_value IS NULL'))

            rtype_cb.bind('<<ComboboxSelected>>', preset_expression)
            preset_expression()

            tk.Label(frm, text=self.lm.tr('lbl_error_message', "Hata Mesajı"), bg='white').grid(row=4, column=0, sticky='w')
            errmsg_var = tk.StringVar(value='Kural ihlali')
            tk.Entry(frm, textvariable=errmsg_var, width=40).grid(row=4, column=1, sticky='w')

            active_var = tk.BooleanVar(value=True)
            tk.Checkbutton(frm, text=self.lm.tr('chk_active', "Aktif"), variable=active_var, bg='white').grid(row=5, column=1, sticky='w', pady=(5,0))

            btns = tk.Frame(frm, bg='white')
            btns.grid(row=6, column=1, sticky='e', pady=10)

            def submit() -> None:
                name = name_var.get().strip()
                if not name:
                    messagebox.showwarning(self.lm.tr('msg_warning', "Uyarı"), self.lm.tr('msg_rule_name_empty', "Kural adı boş olamaz"))
                    return
                rtype = rtype_var.get()
                expr = expr_text.get('1.0', tk.END).strip()
                rid = self.data_validation.add_rule(
                    rule_name=name, rule_type=rtype, validation_expression=expr,
                    error_message=errmsg_var.get().strip() or 'Kural ihlali',
                    severity_level=sev_var.get(), is_active=active_var.get()
                )
                if rid:
                    messagebox.showinfo(self.lm.tr('msg_success', "Başarılı"), self.lm.tr('msg_rule_added', "Kural eklendi"))
                    win.destroy()
                    self.load_validation_rules()
                else:
                    messagebox.showerror(self.lm.tr('msg_error', "Hata"), self.lm.tr('msg_rule_add_error', "Kural eklenemedi"))

            tk.Button(btns, text=self.lm.tr('btn_save', "Kaydet"), command=submit, bg='#27ae60', fg='white', relief='flat', padx=12).pack(side='left', padx=5)
            tk.Button(btns, text=self.lm.tr('btn_cancel', "İptal"), command=win.destroy).pack(side='left', padx=5)
        except Exception as e:
            messagebox.showerror(self.lm.tr('msg_error', "Hata"), f"{self.lm.tr('msg_error_adding_rule', 'Kural eklenirken hata')}: {str(e)}")

    def edit_rule(self) -> None:
        """Kural düzenle"""
        try:
            selection = self.rules_tree.selection()
            if not selection:
                messagebox.showwarning(self.lm.tr('msg_warning', "Uyarı"), self.lm.tr('msg_select_rule', "Lütfen bir kural seçin"))
                return
            item_id = selection[0]
            rule_id = getattr(self, 'rule_item_to_id', {}).get(item_id)
            if not rule_id:
                messagebox.showerror(self.lm.tr('msg_error', "Hata"), self.lm.tr('msg_rule_id_not_found', "Seçili kural ID’si bulunamadı"))
                return
            from tkinter import simpledialog
            new_name = simpledialog.askstring(self.lm.tr('title_new_rule_name', "Yeni Kural Adı"), self.lm.tr('prompt_new_rule_name', "Yeni adı girin (boş bırakılırsa değişmez)"))
            new_type = simpledialog.askstring(self.lm.tr('title_new_rule_type', "Yeni Tür"), self.lm.tr('prompt_new_rule_type', "Yeni tür (boş bırakılırsa değişmez)"))
            new_expr = simpledialog.askstring(self.lm.tr('title_new_rule_expr', "Yeni İfade"), self.lm.tr('prompt_new_rule_expr', "Yeni SQL WHERE ifadesi (boş bırakılırsa değişmez)"))
            updates = {}
            if new_name:
                updates['rule_name'] = new_name
            if new_type:
                updates['rule_type'] = new_type
            if new_expr:
                updates['validation_expression'] = new_expr
            if not updates:
                return
            ok = self.data_validation.update_rule(rule_id, **updates)
            if ok:
                messagebox.showinfo(self.lm.tr('msg_success', "Başarılı"), self.lm.tr('msg_rule_updated', "Kural güncellendi"))
                self.load_validation_rules()
            else:
                messagebox.showerror(self.lm.tr('msg_error', "Hata"), self.lm.tr('msg_rule_update_error', "Kural güncellenemedi"))
        except Exception as e:
            messagebox.showerror(self.lm.tr('msg_error', "Hata"), f"{self.lm.tr('msg_error_editing_rule', 'Kural düzenlenirken hata')}: {str(e)}")

    def delete_rule(self) -> None:
        """Kural sil"""
        try:
            selection = self.rules_tree.selection()
            if not selection:
                messagebox.showwarning(self.lm.tr('msg_warning', "Uyarı"), self.lm.tr('msg_select_rule', "Lütfen bir kural seçin"))
                return
            item_id = selection[0]
            rule_id = getattr(self, 'rule_item_to_id', {}).get(item_id)
            if not rule_id:
                messagebox.showerror(self.lm.tr('msg_error', "Hata"), self.lm.tr('msg_rule_id_not_found', "Seçili kural ID’si bulunamadı"))
                return
            if messagebox.askyesno(self.lm.tr('title_confirm', "Onay"), self.lm.tr('msg_confirm_delete_rule', "Seçili kuralı silmek istiyor musunuz?")):
                ok = self.data_validation.delete_rule(rule_id)
                if ok:
                    messagebox.showinfo(self.lm.tr('msg_success', "Başarılı"), self.lm.tr('msg_rule_deleted', "Kural silindi"))
                    self.load_validation_rules()
                else:
                    messagebox.showerror(self.lm.tr('msg_error', "Hata"), self.lm.tr('msg_rule_delete_error', "Kural silinemedi"))
        except Exception as e:
            messagebox.showerror(self.lm.tr('msg_error', "Hata"), f"{self.lm.tr('msg_error_deleting_rule', 'Kural silinirken hata')}: {str(e)}")

    def toggle_rule_active(self) -> None:
        """Seçili kuralı aktif/pasif yap"""
        try:
            selection = self.rules_tree.selection()
            if not selection:
                messagebox.showwarning(self.lm.tr('msg_warning', "Uyarı"), self.lm.tr('msg_select_rule', "Lütfen bir kural seçin"))
                return
            item_id = selection[0]
            rule_id = getattr(self, 'rule_item_to_id', {}).get(item_id)
            if not rule_id:
                messagebox.showerror(self.lm.tr('msg_error', "Hata"), self.lm.tr('msg_rule_id_not_found', "Seçili kural ID’si bulunamadı"))
                return
            # Mevcut durum
            values = self.rules_tree.item(item_id, 'values')
            current_status = values[2]
            new_active = 0 if current_status == 'Aktif' else 1
            ok = self.data_validation.update_rule(rule_id, is_active=new_active)
            if ok:
                status_str = 'aktif' if new_active else 'pasif'
                messagebox.showinfo(self.lm.tr('msg_success', "Başarılı"), self.lm.tr('msg_rule_status_changed', f"Kural {status_str} yapıldı").format(status=status_str))
                self.load_validation_rules()
            else:
                messagebox.showerror(self.lm.tr('msg_error', "Hata"), self.lm.tr('msg_rule_update_error', "Kural güncellenemedi"))
        except Exception as e:
            messagebox.showerror(self.lm.tr('msg_error', "Hata"), f"{self.lm.tr('msg_error_changing_status', 'Durum değiştirilirken hata')}: {str(e)}")

    def run_validation(self) -> None:
        """Veri doğrulaması çalıştır"""
        try:
            # Kontrol seçeneklerini al
            checks = []
            if self.check_completeness.get():
                checks.append('completeness')
            if self.check_consistency.get():
                checks.append('consistency')
            if self.check_format.get():
                checks.append('format')

            if not checks:
                messagebox.showwarning("Uyarı", "Lütfen en az bir kontrol türü seçin")
                return

            # Doğrulama çalıştır
            results = self.data_validation.validate_responses(self.company_id, checks)

            # Sonuçları göster
            self.validation_results.delete('1.0', tk.END)

            if results:
                result_text = "VERİ DOĞRULAMA SONUÇLARI\n"
                result_text += "=" * 50 + "\n\n"

                for check_type in checks:
                    result_text += f"{check_type.upper()} KONTROLÜ:\n"
                    result_text += "-" * 30 + "\n"

                    if check_type in results:
                        for issue in results[check_type]:
                            result_text += f" {issue}\n"
                    else:
                        result_text += " Sorun bulunamadı\n"

                    result_text += "\n"

                self.validation_results.insert('1.0', result_text)
            else:
                self.validation_results.insert('1.0', "Doğrulama sonucu bulunamadı")

            messagebox.showinfo("Başarılı", "Veri doğrulaması tamamlandı")

        except Exception as e:
            messagebox.showerror("Hata", f"Doğrulama çalıştırılırken hata: {str(e)}")

    def generate_validation_report(self) -> None:
        """Doğrulama raporu oluştur ve dışa aktar (JSON ve TXT)"""
        try:
            results = self.data_validation.validate_company_data(self.company_id)
            export_dir = os.path.join(os.path.dirname(__file__), '..', 'data', 'exports')
            os.makedirs(export_dir, exist_ok=True)
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            json_path = os.path.join(export_dir, f"validation_report_{self.company_id}_{ts}.json")
            txt_path = os.path.join(export_dir, f"validation_report_{self.company_id}_{ts}.txt")

            with open(json_path, 'w', encoding='utf-8') as f:
                f.write(json.dumps(results, ensure_ascii=False, indent=2))

            # İnsan okunur TXT
            lines = []
            lines.append(f"Şirket ID: {results.get('company_id')}")
            lines.append(f"Tarih: {results.get('validation_date')}")
            lines.append(f"Toplam Kural: {results.get('total_rules')}")
            lines.append(f"Geçen: {results.get('passed_rules')} - Başarısız: {results.get('failed_rules')}")
            lines.append("")
            lines.append("Detaylar:")
            for d in results.get('details', []):
                lines.append(f"- [{d.get('severity')}] {d.get('rule_name')} | SDG {d.get('sdg_no')} {d.get('indicator_code')} | {d.get('error_message')}")
            lines.append("")
            qs = results.get('quality_scores', {})
            lines.append("Kalite Skorları:")
            lines.append(f"  Eksiksizlik: %{qs.get('completeness_score', 0)}")
            lines.append(f"  Doğruluk: %{qs.get('accuracy_score', 0)}")
            lines.append(f"  Tutarlılık: %{qs.get('consistency_score', 0)}")
            lines.append(f"  Güncellik: %{qs.get('timeliness_score', 0)}")
            lines.append(f"  Genel: %{qs.get('overall_quality_score', 0)}")

            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write("\n".join(lines))

            # DOCX opsiyonel (python-docx varsa)
            try:
                from docx import Document
                doc = Document()
                doc.add_heading('SDG Veri Doğrulama Raporu', level=1)
                doc.add_paragraph(f"Şirket ID: {results.get('company_id')}")
                doc.add_paragraph(f"Tarih: {results.get('validation_date')}")
                doc.add_heading('Detaylar', level=2)
                for d in results.get('details', []):
                    doc.add_paragraph(f"[{d.get('severity')}] {d.get('rule_name')} | SDG {d.get('sdg_no')} {d.get('indicator_code')}")
                    doc.add_paragraph(f"  Hata: {d.get('error_message')}")
                    doc.add_paragraph(f"  Öneri: {d.get('suggested_fix')}")
                doc.add_heading('Kalite Skorları', level=2)
                doc.add_paragraph(f"Eksiksizlik: %{qs.get('completeness_score', 0)}")
                doc.add_paragraph(f"Doğruluk: %{qs.get('accuracy_score', 0)}")
                doc.add_paragraph(f"Tutarlılık: %{qs.get('consistency_score', 0)}")
                doc.add_paragraph(f"Güncellik: %{qs.get('timeliness_score', 0)}")
                doc.add_paragraph(f"Genel: %{qs.get('overall_quality_score', 0)}")
                docx_path = os.path.join(export_dir, f"validation_report_{self.company_id}_{ts}.docx")
                doc.save(docx_path)
                messagebox.showinfo("Başarılı", f"Raporlar oluşturuldu:\n{json_path}\n{txt_path}\n{docx_path}")
            except Exception:
                # python-docx yoksa sadece JSON/TXT bildir
                messagebox.showinfo("Başarılı", f"Raporlar oluşturuldu:\n{json_path}\n{txt_path}")
        except Exception as e:
            messagebox.showerror("Hata", f"Rapor oluşturulurken hata: {str(e)}")

    def refresh_quality_scores(self) -> None:
        """Kalite skorlarını yenile (tek satır: genel skorlar)"""
        try:
            for item in self.quality_tree.get_children():
                self.quality_tree.delete(item)
            if getattr(self, 'quality_by_sdg', None) and self.quality_by_sdg.get():
                # SDG bazlı kırılım
                sdg_scores = self.data_validation.calculate_quality_scores_by_sdg(self.company_id)
                if not sdg_scores:
                    messagebox.showinfo("Bilgi", "SDG bazlı skor bulunamadı")
                for sdg_no, s in sorted(sdg_scores.items(), key=lambda x: x[0]):
                    self.quality_tree.insert('', 'end', values=(
                        f"SDG {sdg_no}",
                        f"%{s.get('completeness', 0):.1f}",
                        f"%{s.get('consistency', 0):.1f}",
                        f"%{s.get('accuracy', 0):.1f}",
                        f"%{s.get('overall', 0):.1f}"
                    ))
            else:
                # Genel skor
                scores = self.data_validation.calculate_quality_scores(self.company_id)
                self.quality_tree.insert('', 'end', values=(
                    'Genel',
                    f"%{scores.get('completeness_score', 0):.1f}",
                    f"%{scores.get('consistency_score', 0):.1f}",
                    f"%{scores.get('accuracy_score', 0):.1f}",
                    f"%{scores.get('overall_quality_score', 0):.1f}"
                ))
            messagebox.showinfo("Başarılı", "Kalite skorları güncellendi")
        except Exception as e:
            messagebox.showerror("Hata", f"Kalite skorları yenilenirken hata: {str(e)}")

    def show_trend_analysis(self) -> None:
        """Basit trend analizi penceresi"""
        try:
            summary = self.data_validation.get_validation_summary(self.company_id)
            win = tk.Toplevel(self.parent)
            win.title("Trend Analizi")
            text = tk.Text(win, height=20, width=80, bg='#f8f9fa')
            text.pack(fill='both', expand=True)
            text.insert('1.0', json.dumps(summary, ensure_ascii=False, indent=2))
        except Exception as e:
            messagebox.showerror("Hata", f"Trend analizi oluşturulamadı: {str(e)}")

    def show_quality_chart(self) -> None:
        """Basit ASCII grafik penceresi"""
        try:
            scores = self.data_validation.calculate_quality_scores(self.company_id)
            win = tk.Toplevel(self.parent)
            win.title("Kalite Grafiği")
            text = tk.Text(win, height=15, width=60, bg='#f8f9fa', font=('Courier', 10))
            text.pack(fill='both', expand=True)
            def bar(label, value) -> str:
                filled = int((value/100)*40)
                return f"{label:12} | " + ('#'*filled).ljust(40, ' ') + f" | %{value:.1f}\n"
            out = ''
            out += bar('Eksiksizlik', scores.get('completeness_score', 0))
            out += bar('Doğruluk', scores.get('accuracy_score', 0))
            out += bar('Tutarlılık', scores.get('consistency_score', 0))
            out += bar('Güncellik', scores.get('timeliness_score', 0))
            out += bar('Genel', scores.get('overall_quality_score', 0))
            text.insert('1.0', out)
        except Exception as e:
            messagebox.showerror("Hata", f"Grafik gösterilirken hata: {str(e)}")

    def on_rule_selected(self, _event=None) -> None:
        """Seçili kural detaylarını göster"""
        try:
            selection = self.rules_tree.selection()
            if not selection:
                return
            item_id = selection[0]
            rule_id = getattr(self, 'rule_item_to_id', {}).get(item_id)
            if not rule_id:
                return
            detail = self.data_validation.get_rule_by_id(rule_id)
            self.rule_details.delete('1.0', tk.END)
            if detail:
                self.rule_details.insert('1.0', json.dumps(detail, ensure_ascii=False, indent=2))
            else:
                self.rule_details.insert('1.0', 'Detay bulunamadı')
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

if __name__ == "__main__":
    # Test
    root = tk.Tk()
    root.title("SDG Veri Doğrulama")
    app = SDGDataValidationGUI(root, 1)
    root.mainloop()
