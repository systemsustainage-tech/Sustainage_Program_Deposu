#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Değer Yaratma Hikayesi Raporu
"""

import logging
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..iirc_manager import IIRCManager


def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    para = _add_turkish_paragraph(doc, text, style=style)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = _add_turkish_heading(doc, text, level=level)
    for run in heading.runs:
        run.font.name = font_name
    return heading


class ValueCreationReportGenerator:
    """Değer Yaratma Hikayesi Raporu"""

    def __init__(self):
        self.manager = IIRCManager()

    def generate_value_report(self, company_id: int, year: int, output_path: str) -> bool:
        """Değer yaratma raporu oluştur"""
        try:
            logging.info(f"[IIRC Value] Değer yaratma raporu oluşturuluyor: {year}")

            # Veri toplama
            value_elements = self.manager.get_value_creation_story(company_id, year)
            connections = self.manager.get_connectivity_map(company_id, year)

            # DOCX oluştur
            doc = self._create_docx(company_id, year, value_elements, connections)
            doc.save(output_path)

            logging.info(f"[IIRC Value] Rapor kaydedildi: {output_path}")
            return True

        except Exception as e:
            logging.error(f"[HATA] Değer yaratma raporu oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _create_docx(self, company_id: int, year: int, value_elements: list, connections: list) -> Document:
        """DOCX oluştur"""
        doc = Document()

        # Başlık
        title = _add_turkish_heading(doc, 'DEĞER YARATMA HİKAYESİ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = _add_turkish_paragraph(doc, f'Sustainage - {year}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)

        _add_turkish_paragraph(doc, )

        # Giriş
        _add_turkish_heading(doc, 'GİRİŞ', 1)
        intro = """
Bu rapor, organizasyonumuzun nasıl değer yarattığını hikaye eden entegre bir anlatımdır.
6 sermaye modeli kullanılarak, girdilerden çıktılara ve sonuçlara uzanan değer 
yaratma sürecimiz detaylı olarak açıklanmaktadır.
        """.strip()
        _add_turkish_paragraph(doc, intro)

        # Değer yaratma öğeleri
        doc.add_page_break()
        _add_turkish_heading(doc, 'DEĞER YARATMA SÜRECİ', 1)

        # Varsayılan hikaye (eğer veri yoksa)
        if not value_elements:
            value_elements = self._get_default_story_elements()

        for element in value_elements:
            title_text = element.get('element_title', element.get('baslik', ''))
            content_text = element.get('content', element.get('aciklama', ''))

            _add_turkish_heading(doc, title_text, 2)
            _add_turkish_paragraph(doc, content_text if content_text else 'İçerik eklenmemiş.')
            _add_turkish_paragraph(doc, )

        # Sermaye bağlantıları
        doc.add_page_break()
        _add_turkish_heading(doc, 'SERMAYE BAĞLANTILARI', 1)

        _add_turkish_paragraph(doc, """
Aşağıdaki tablo, 6 sermaye arasındaki bağlantıları göstermektedir.
Bu bağlantılar, bir sermayenin diğerine nasıl dönüştüğünü veya 
nasıl katkı sağladığını gösterir.
        """.strip())

        if connections:
            conn_table = doc.add_table(rows=len(connections) + 1, cols=4)
            conn_table.style = 'Light Grid Accent 1'

            # Başlıklar
            headers = conn_table.rows[0].cells
            headers[0].text = 'Kaynak Sermaye'
            headers[1].text = 'Hedef Sermaye'
            headers[2].text = 'Bağlantı Türü'
            headers[3].text = 'Güç'

            # Veriler
            for i, conn in enumerate(connections, 1):
                row = conn_table.rows[i].cells
                row[0].text = conn.get('source_name', '-')
                row[1].text = conn.get('target_name', '-')
                row[2].text = conn.get('connection_type', '-')
                row[3].text = conn.get('connection_strength', '-')

        # Altbilgi
        doc.add_page_break()
        footer = _add_turkish_paragraph(doc, )
        footer.add_run(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y")}\n').font.size = Pt(9)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return doc

    def _get_default_story_elements(self) -> list:
        """Varsayılan hikaye öğeleri"""
        return [
            {
                'element_title': 'Organizasyona Genel Bakış',
                'content': 'Şirketimiz sürdürülebilirlik alanında değer yaratmaktadır.'
            },
            {
                'element_title': 'İş Modeli',
                'content': '6 sermaye kullanarak paydaşlarımız için değer yaratıyoruz.'
            },
            {
                'element_title': 'Performans',
                'content': 'Tüm sermaye türlerinde olumlu gelişmeler kaydedilmiştir.'
            }
        ]

