import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
IIRC Entegre Rapor Oluşturucu
Tam entegre rapor (DOCX/PDF) oluşturma
"""

import os
from datetime import datetime
from typing import Any, Dict

import matplotlib
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

matplotlib.use('Agg')

from ..iirc_manager import IIRCManager


def _add_turkish_paragraph(doc, text=None, style=None, font_name='Calibri', font_size=11):
    para = doc.add_paragraph(text if text is not None else '', style=style)
    for run in para.runs:
        try:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        try:
            run.font.name = font_name
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
    return heading


class IntegratedReportGenerator:
    """IIRC Entegre Rapor Oluşturucu"""

    def __init__(self):
        self.manager = IIRCManager()

    def generate_full_report(self, company_id: int, year: int, output_path: str) -> bool:
        """Tam entegre rapor oluştur"""
        try:
            logging.info(f"[IIRC] Entegre rapor oluşturuluyor: {year}")

            # Tüm verileri topla
            data = self._collect_all_data(company_id, year)

            # DOCX oluştur
            doc = self._create_integrated_report(data)
            doc.save(output_path)

            logging.info(f"[IIRC] Rapor kaydedildi: {output_path}")
            return True

        except Exception as e:
            logging.error(f"[HATA] IIRC raporu oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _collect_all_data(self, company_id: int, year: int) -> Dict[str, Any]:
        """Tüm verileri topla"""
        data = {
            'company_id': company_id,
            'year': year,
            'company_info': self.manager.collector.collect_company_info(company_id) if hasattr(self.manager, 'collector') else {'name': 'Şirket'},
            'capitals': self.manager.get_six_capitals(company_id, year),
            'value_story': self.manager.get_value_creation_story(company_id, year),
            'connections': self.manager.get_connectivity_map(company_id, year),
            'report_sections': self.manager.get_report_content(company_id, year),
            'summary': self.manager.get_company_summary(company_id, year)
        }
        return data

    def _create_integrated_report(self, data: Dict) -> Document:
        """Entegre rapor DOCX oluştur"""
        doc = Document()
        company_id = data.get('company_id')

        # Logo
        try:
            # db_path bul
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
            db_path = os.path.join(base_dir, "data", "sdg_desktop.sqlite")
            
            from modules.reporting.brand_identity_manager import BrandIdentityManager
            bim = BrandIdentityManager(db_path, company_id)
            bi = bim.get_brand_identity(company_id)
            logo_path = bi.get('logo_path')
            
            if logo_path and os.path.exists(logo_path):
                from docx.shared import Inches
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(logo_path, width=Inches(1.6))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logging.warning(f"Logo ekleme hatası (IIRC): {e}")

        # KAPAK SAYFASI
        title = _add_turkish_heading(doc, 'ENTEGRE RAPOR', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        company_name = data.get('company_info', {}).get('name', 'Şirket')
        year = data.get('year', datetime.now().year)

        subtitle = _add_turkish_paragraph(doc, f'{company_name}\n{year}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(18)
        subtitle.runs[0].font.bold = True

        _add_turkish_paragraph(doc, )
        _add_turkish_paragraph(doc, )

        slogan = _add_turkish_paragraph(doc, 'Değer Yaratma Hikayemiz')
        slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
        slogan.runs[0].font.size = Pt(14)
        slogan.runs[0].font.italic = True
        slogan.runs[0].font.color.rgb = RGBColor(106, 27, 154)  # Mor

        # İÇİNDEKİLER
        doc.add_page_break()
        _add_turkish_heading(doc, 'İÇİNDEKİLER', 1)

        toc_items = [
            '1. Organizasyona Genel Bakış',
            '2. Yönetişim',
            '3. İş Modeli ve Değer Yaratma',
            '4. 6 Sermaye Modeli',
            '5. Riskler ve Fırsatlar',
            '6. Strateji ve Kaynak Tahsisi',
            '7. Performans',
            '8. Gelecek Görünümü',
            '9. Raporun Hazırlanma Temeli'
        ]

        for item in toc_items:
            _add_turkish_paragraph(doc, item, style='List Number')

        # 1. ORGANİZASYONA GENEL BAKIŞ
        doc.add_page_break()
        _add_turkish_heading(doc, '1. ORGANİZASYONA GENEL BAKIŞ', 1)

        company_info = data.get('company_info', {})
        org_overview = f"""
{company_name}, {company_info.get('sector', 'sürdürülebilirlik')} sektöründe faaliyet gösteren, 
{company_info.get('employees', 0)} çalışanıyla değer yaratan bir organizasyondur.

MİSYON
Sürdürülebilir bir gelecek için değer yaratmak.

VİZYON
Sektöründe sürdürülebilirlik lideri olmak.

DEĞERLER
• Şeffaflık
• Paydaş Odaklılık
• Sürdürülebilirlik
• İnovasyon
• Mükemmellik
        """.strip()

        _add_turkish_paragraph(doc, org_overview)

        # 2. YÖNETİŞİM
        doc.add_page_break()
        _add_turkish_heading(doc, '2. YÖNETİŞİM', 1)

        governance_text = """
Yönetim Kurulu, şirketin stratejik yönünü belirler ve tüm paydaşların çıkarlarını dengeler.
Sürdürülebilirlik konuları, kurumsal yönetişimimizin merkezinde yer almaktadır.

YÖNETİM YAPISI
• Yönetim Kurulu: Stratejik gözetim
• Üst Yönetim: Operasyonel yönetim
• Sürdürülebilirlik Komitesi: ESG gözetimi
        """.strip()

        _add_turkish_paragraph(doc, governance_text)

        # 3. İŞ MODELİ VE DEĞER YARATMA
        doc.add_page_break()
        _add_turkish_heading(doc, '3. İŞ MODELİ VE DEĞER YARATMA', 1)

        business_model = """
İş modelimiz, 6 sermaye türünü kullanarak paydaşlarımız için değer yaratmayı amaçlar.

GİRDİLER (6 Sermaye)
• Mali Sermaye: Finansal kaynaklar
• İmalat Sermayesi: Fiziksel varlıklar
• Entelektüel Sermaye: Bilgi ve inovasyon
• İnsan Sermayesi: Yetenekli çalışanlar
• Sosyal Sermaye: Paydaş ilişkileri
• Doğal Sermaye: Çevresel kaynaklar

FAALİYETLER
• Ürün/Hizmet Üretimi
• Sürdürülebilir Operasyonlar
• İnovasyon ve Ar-Ge
• Paydaş Katılımı

ÇIKTILAR
• Ürünler ve Hizmetler
• İstihdam ve Gelir
• İnovasyon ve Gelişim
• Çevresel Performans

SONUÇLAR (Değer Yaratma)
• Ekonomik Değer: Gelir, kar, ekonomik katkı
• Sosyal Değer: İstihdam, eğitim, topluluk katkısı
• Çevresel Değer: Emisyon azaltımı, kaynak verimliliği
        """.strip()

        _add_turkish_paragraph(doc, business_model)

        # 4. 6 SERMAYE MODELİ
        doc.add_page_break()
        _add_turkish_heading(doc, '4. 6 SERMAYE MODELİ', 1)

        capitals = data.get('capitals', [])

        for capital in capitals:
            _add_turkish_heading(doc, f"4.{capitals.index(capital) + 1}. {capital.get('capital_name', 'Sermaye')}", 2)

            capital_table = doc.add_table(rows=1, cols=2)
            capital_table.style = 'Light Grid Accent 1'

            rows_data = [
                ('Açıklama', capital.get('description', '-')),
                ('Mevcut Değer', f"{capital.get('current_value', 0):.2f}"),
                ('Trend', capital.get('trend', 'Stabil')),
                ('Durum', capital.get('status', 'Aktif'))
            ]

            for label, value in rows_data:
                r = capital_table.add_row().cells
                r[0].text = label
                r[1].text = str(value)

            _add_turkish_paragraph(doc, )

        # 5. RİSKLER VE FIRSATLAR
        doc.add_page_break()
        _add_turkish_heading(doc, '5. RİSKLER VE FIRSATLAR', 1)

        risks_opps = """
BAŞLICA RİSKLER
• İklim değişikliği riskleri
• Düzenleyici değişiklikler
• Rekabet baskısı
• Tedarik zinciri kesintileri

BAŞLICA FIRSATLAR
• Yenilenebilir enerji geçişi
• Dijital dönüşüm
• Sürdürülebilir ürün talebi
• Yeni pazarlar
        """.strip()

        _add_turkish_paragraph(doc, risks_opps)

        # 6. STRATEJİ VE KAYNAK TAHSİSİ
        doc.add_page_break()
        _add_turkish_heading(doc, '6. STRATEJİ VE KAYNAK TAHSİSİ', 1)

        strategy = """
Stratejimiz, 6 sermayemizi optimal şekilde kullanarak uzun vadeli değer yaratmaya odaklanır.

STRATEJİK ÖNCE LİKLER
1. Sürdürülebilir büyüme
2. İklim etkisini azaltma
3. İnovasyon ve dijitalleşme
4. Paydaş değeri yaratma

KAYNAK TAHSİSİ
• Ar-Ge ve İnovasyon: %15
• Sürdürülebilirlik Projeleri: %10
• Çalışan Gelişimi: %5
• Topluluk Yatırımları: %3
        """.strip()

        _add_turkish_paragraph(doc, strategy)

        # 7. PERFORMANS
        doc.add_page_break()
        _add_turkish_heading(doc, '7. PERFORMANS', 1)

        summary = data.get('summary', {})

        performance = f"""
6 SERMAYE PERFORMANSI

Toplam Sermaye Sayısı: {summary.get('total_capitals', 0)}
Aktif Sermaye: {summary.get('active_capitals', 0)}
Entegrasyon Skoru: {summary.get('integration_score', 0):.1f}%

En Güçlü Sermaye: {summary.get('strongest_capital', 'Belirsiz')}
        """.strip()

        _add_turkish_paragraph(doc, performance)

        # 8. GELECEK GÖRÜNÜMÜ
        doc.add_page_break()
        _add_turkish_heading(doc, '8. GELECEK GÖRÜNÜMÜ', 1)

        outlook = """
Gelecek dönemde, sürdürülebilirlik hedeflerimize ulaşmak için 6 sermayemizi 
stratejik olarak geliştirmeye devam edeceğiz.

KISA VADELİ HEDEFLER (1 Yıl)
• Karbon emisyonlarında %10 azaltım
• Çalışan memnuniyetinde %5 artış
• Yenilenebilir enerji kullanımında %20 artış

ORTA VADELİ HEDEFLER (3 Yıl)
• Net-Zero yol haritası oluşturma
• Tam dijital dönüşüm
• Döngüsel ekonomi modeline geçiş

UZUN VADELİ HEDEFLER (5+ Yıl)
• Karbon nötr olma (2030)
• Sıfır atık hedefi
• %100 sürdürülebilir tedarik zinciri
        """.strip()

        _add_turkish_paragraph(doc, outlook)

        # 9. RAPORUN HAZIRLANMA TEMELİ
        doc.add_page_break()
        _add_turkish_heading(doc, '9. RAPORUN HAZIRLANMA TEMELİ', 1)

        basis = """
Bu entegre rapor, IIRC (International Integrated Reporting Council) Çerçevesine 
uygun olarak hazırlanmıştır.

RAPORLAMA STANDARTLARI
• IIRC Entegre Raporlama Çerçevesi
• GRI Standartları
• TCFD Önerileri
• ESRS (European Sustainability Reporting Standards)

VERİ TOPLAMA VE DOĞRULAMA
• Veriler şirket içi sistemlerden toplanmıştır
• Finansal veriler bağımsız denetimden geçmiştir
• ESG verileri iç kontrol süreçlerinden geçmiştir

RAPORLAMA KAPSAMI
• Raporlama Dönemi: 1 Ocak - 31 Aralık {year}
• Coğrafi Kapsam: Türkiye operasyonları
• Konsolidasyon: Tam konsolidasyon
        """.format(year=year).strip()

        _add_turkish_paragraph(doc, basis)

        # ALTBILGI
        doc.add_page_break()
        footer = _add_turkish_paragraph(doc, )
        footer.add_run(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y")}\n').font.size = Pt(9)
        footer.add_run('Bu rapor Sustainage IIRC Modülü tarafından otomatik olarak oluşturulmuştur.\n').font.size = Pt(9)
        footer.add_run('© IIRC Framework - Entegre Raporlama').font.size = Pt(9)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return doc

