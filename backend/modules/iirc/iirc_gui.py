import logging
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
IIRC (Entegre Raporlama) GUI - TAMAMEN TÜRKÇE
6 Sermaye Modeli ve Değer Yaratma Hikayesi için kullanıcı arayüzü
"""

import tkinter as tk
from tkinter import messagebox, ttk

from utils.language_manager import LanguageManager
from .iirc_manager import IIRCManager
from .six_capitals import SixCapitalsManager
from .value_creation import ValueCreationStory
from config.icons import Icons

try:
    from openpyxl.styles import Font
except ImportError:
    Font = None

class IIRCGUI:
    """IIRC (Entegre Raporlama) GUI - TÜRKÇE"""

    def __init__(self, parent, company_id: int) -> None:
        self.parent = parent
        self.company_id = company_id
        self.lm = LanguageManager()
        self.manager = IIRCManager()
        self.capitals_mgr = SixCapitalsManager()
        self.value_story = ValueCreationStory()

        try:
            self.parent.winfo_toplevel().state('zoomed')
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

        self.setup_ui()
        self.load_data()

    def setup_ui(self) -> None:
        """IIRC arayüzünü oluştur"""
        # Ana frame
        main_frame = tk.Frame(self.parent, bg='#f5f5f5')
        main_frame.pack(fill='both', expand=True, padx=10, pady=10)

        # Başlık
        title_frame = tk.Frame(main_frame, bg='#6A1B9A', height=60)
        title_frame.pack(fill='x', pady=(0, 10))
        title_frame.pack_propagate(False)

        title_label = tk.Label(title_frame, text=self.lm.tr('iirc_title', "️ IIRC - Entegre Raporlama"),
                              font=('Segoe UI', 16, 'bold'), fg='white', bg='#6A1B9A')
        title_label.pack(expand=True)

        # Ana içerik alanı
        content_frame = tk.Frame(main_frame, bg='#f5f5f5')
        content_frame.pack(fill='both', expand=True)

        # Notebook oluştur
        self.notebook = ttk.Notebook(content_frame)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=10)

        # Sekmeleri oluştur
        self.create_dashboard_tab()
        self.create_six_capitals_tab()
        self.create_value_creation_tab()
        self.create_connectivity_tab()
        self.create_report_tab()

    def create_dashboard_tab(self) -> None:
        """Dashboard sekmesi"""
        dashboard_frame = tk.Frame(self.notebook, bg='white')
        self.notebook.add(dashboard_frame, text=f" {self.lm.tr('overview', 'Genel Bakış')}")

        # Başlık
        tk.Label(dashboard_frame, text=self.lm.tr('iirc_dashboard', "Entegre Raporlama Dashboard"),
                font=('Segoe UI', 14, 'bold'), bg='white').pack(pady=10)

        # İstatistik kartları
        self.stats_frame = tk.Frame(dashboard_frame, bg='white')
        self.stats_frame.pack(fill='x', padx=20, pady=10)

        # 6 Sermaye özet kartları
        self.capitals_frame = tk.Frame(dashboard_frame, bg='white')
        self.capitals_frame.pack(fill='both', expand=True, padx=20, pady=10)

    def create_six_capitals_tab(self) -> None:
        """6 Sermaye sekmesi"""
        capitals_frame = tk.Frame(self.notebook, bg='white')
        self.notebook.add(capitals_frame, text=f" {self.lm.tr('six_capitals', '6 Sermaye')}")

        # Başlık
        tk.Label(capitals_frame, text=self.lm.tr('six_capitals_model', "6 Sermaye Modeli"),
                font=('Segoe UI', 14, 'bold'), bg='white').pack(pady=10)

        # Alt başlık ve açıklama
        info_text = self.lm.tr('iirc_framework_info', """
IIRC (International Integrated Reporting Council) Çerçevesi, organizasyonların 
değer yaratmak için 6 sermaye türünü nasıl kullandığını gösterir.
        """).strip()

        tk.Label(capitals_frame, text=info_text,
                font=('Segoe UI', 9), bg='white', fg='#666',
                wraplength=800, justify='center').pack(pady=(0, 10))

        # 2 kolonlu layout: Sol = Sermaye listesi, Sağ = Detay
        main_container = tk.Frame(capitals_frame, bg='white')
        main_container.pack(fill='both', expand=True, padx=10, pady=10)

        # Sol panel - Sermaye kartları
        left_panel = tk.Frame(main_container, bg='white', width=400)
        left_panel.pack(side='left', fill='both', expand=True)

        # Canvas ve scrollbar (sol)
        canvas = tk.Canvas(left_panel, bg='white')
        scrollbar = ttk.Scrollbar(left_panel, orient="vertical", command=canvas.yview)
        self.capitals_content = tk.Frame(canvas, bg='white')

        self.capitals_content.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=self.capitals_content, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Sağ panel - Özet bilgi
        right_panel = tk.Frame(main_container, bg='#f8f9fa', width=400, relief='solid', bd=1)
        right_panel.pack(side='right', fill='both', expand=True, padx=(10, 0))

        # Sağ panel başlık
        tk.Label(right_panel, text=self.lm.tr('six_capitals_summary', "6 SERMAYE ÖZETİ"),
                font=('Segoe UI', 12, 'bold'), bg='#6A1B9A', fg='white').pack(fill='x', pady=0)

        # İçerik frame
        self.capitals_info_frame = tk.Frame(right_panel, bg='#f8f9fa')
        self.capitals_info_frame.pack(fill='both', expand=True, padx=15, pady=15)

    def create_value_creation_tab(self) -> None:
        """Değer yaratma hikayesi sekmesi"""
        value_frame = tk.Frame(self.notebook, bg='white')
        self.notebook.add(value_frame, text=f" {self.lm.tr('value_creation_story', 'Değer Yaratma Hikayesi')}")

        # Başlık
        tk.Label(value_frame, text=self.lm.tr('value_creation_story', "Değer Yaratma Hikayesi"),
                font=('Segoe UI', 14, 'bold'), bg='white').pack(pady=10)

        # Canvas ve scrollbar
        canvas = tk.Canvas(value_frame, bg='white')
        scrollbar = ttk.Scrollbar(value_frame, orient="vertical", command=canvas.yview)
        self.value_content = tk.Frame(canvas, bg='white')

        self.value_content.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=self.value_content, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True, padx=10)
        scrollbar.pack(side="right", fill="y")

    def create_connectivity_tab(self) -> None:
        """Bağlantı haritası sekmesi"""
        connectivity_frame = tk.Frame(self.notebook, bg='white')
        self.notebook.add(connectivity_frame, text=f" {self.lm.tr('connectivity_map', 'Bağlantı Haritası')}")

        # Başlık
        tk.Label(connectivity_frame, text=self.lm.tr('capital_connections', "Sermaye Bağlantıları"),
                font=('Segoe UI', 14, 'bold'), bg='white').pack(pady=10)

        # Bağlantı listesi
        self.connectivity_tree = ttk.Treeview(connectivity_frame,
                                             columns=('Kaynak', 'Hedef', 'Tür', 'Güç'),
                                             show='headings', height=15)

        self.connectivity_tree.heading('Kaynak', text=self.lm.tr('source_capital', 'Kaynak Sermaye'))
        self.connectivity_tree.heading('Hedef', text=self.lm.tr('target_capital', 'Hedef Sermaye'))
        self.connectivity_tree.heading('Tür', text=self.lm.tr('connection_type', 'Bağlantı Türü'))
        self.connectivity_tree.heading('Güç', text=self.lm.tr('strength', 'Güç'))

        self.connectivity_tree.pack(fill='both', expand=True, padx=20, pady=10)

    def create_report_tab(self) -> None:
        """Entegre rapor sekmesi"""
        report_frame = tk.Frame(self.notebook, bg='white')
        self.notebook.add(report_frame, text=f" {self.lm.tr('integrated_report', 'Entegre Rapor')}")

        # Başlık
        tk.Label(report_frame, text=self.lm.tr('create_integrated_report', "Entegre Rapor Oluştur"),
                font=('Segoe UI', 14, 'bold'), bg='white').pack(pady=10)

        # Rapor butonları
        btn_frame = tk.Frame(report_frame, bg='white')
        btn_frame.pack(fill='x', padx=20, pady=10)

        report_buttons = [
            (f" {self.lm.tr('integrated_report_template', 'Entegre Rapor Şablonu')}", self.generate_template, '#6A1B9A'),
            (f" {self.lm.tr('six_capitals_report', '6 Sermaye Raporu')}", self.generate_capitals_report, '#8E24AA'),
            (f" {self.lm.tr('value_creation_report', 'Değer Yaratma Raporu')}", self.generate_value_report, '#AB47BC'),
            (f" {self.lm.tr('connectivity_analysis', 'Bağlantı Analizi')}", self.generate_connectivity_report, '#BA68C8'),
            (f" {self.lm.tr('full_integrated_report_pdf', 'Tam Entegre Rapor (PDF)')}", self.generate_full_report, '#CE93D8'),
        ]

        for i, (text, command, color) in enumerate(report_buttons):
            btn = ttk.Button(btn_frame, text=text, command=command,
                             style='Primary.TButton')
            btn.grid(row=i//2, column=i%2, padx=10, pady=10, sticky='ew')

        btn_frame.grid_columnconfigure(0, weight=1)
        btn_frame.grid_columnconfigure(1, weight=1)

    def load_data(self) -> None:
        """Verileri yükle"""
        try:
            self.load_dashboard_data()
            self.load_capitals_data()
            self.load_value_creation_data()
            self.load_connectivity_data()
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('err_data_load', 'Veri yüklenirken hata')}: {e}")

    def load_dashboard_data(self) -> None:
        """Dashboard verilerini yükle"""
        try:
            summary = self.manager.get_company_summary(self.company_id)

            # İstatistik kartlarını temizle
            for widget in self.stats_frame.winfo_children():
                widget.destroy()

            stats_data = [
                (self.lm.tr('total_capital', "Toplam Sermaye"), summary.get('total_capitals', 0), '#6A1B9A'),
                (self.lm.tr('active_capital', "Aktif Sermaye"), summary.get('active_capitals', 0), '#8E24AA'),
                (self.lm.tr('connections', "Bağlantılar"), summary.get('total_connections', 0), '#AB47BC'),
                (self.lm.tr('integration_score', "Entegrasyon Skoru"), f"{summary.get('integration_score', 0):.1f}%", '#BA68C8'),
            ]

            for i, (title, value, color) in enumerate(stats_data):
                self.create_stat_card(self.stats_frame, title, value, color, 0, i)

            # 6 Sermaye özet kartları
            for widget in self.capitals_frame.winfo_children():
                widget.destroy()

            capitals = summary.get('capitals_data', [])
            for i, capital in enumerate(capitals):
                row = i // 3
                col = i % 3
                self.create_capital_summary_card(self.capitals_frame, capital, row, col)

        except Exception as e:
            logging.error(f"Dashboard verileri yüklenirken hata: {e}")

    def create_stat_card(self, parent, title, value, color, row, col):
        """İstatistik kartı oluştur"""
        card = tk.Frame(parent, bg=color, relief='raised', bd=2)
        card.grid(row=row, column=col, padx=10, pady=10, sticky='nsew')
        parent.grid_columnconfigure(col, weight=1)

        tk.Label(card, text=title, font=('Segoe UI', 10, 'bold'),
                bg=color, fg='white').pack(pady=5)
        tk.Label(card, text=str(value), font=('Segoe UI', 16, 'bold'),
                bg=color, fg='white').pack(pady=5)

    def create_capital_summary_card(self, parent, capital, row, col):
        """Sermaye özet kartı oluştur"""
        card = tk.Frame(parent, bg='#f8f9fa', relief='solid', bd=1)
        card.grid(row=row, column=col, padx=5, pady=5, sticky='nsew')
        parent.grid_columnconfigure(col, weight=1)
        parent.grid_rowconfigure(row, weight=1)

        # İkon ve başlık
        icon = capital.get('icon', '')
        name = capital.get('capital_name', self.lm.tr('capital', 'Sermaye'))

        tk.Label(card, text=f"{icon} {name}", font=('Segoe UI', 12, 'bold'),
                bg='#f8f9fa').pack(pady=5)

        # Değer
        value = capital.get('current_value', 0)
        tk.Label(card, text=f"{self.lm.tr('value', 'Değer')}: {value:.2f}", font=('Segoe UI', 10),
                bg='#f8f9fa').pack()

        # Trend
        trend = capital.get('trend', 'Stabil')
        trend_map = {
            'Artan': self.lm.tr('increasing', 'Artan'),
            'Azalan': self.lm.tr('decreasing', 'Azalan'),
            'Stabil': self.lm.tr('stable', 'Stabil')
        }
        display_trend = trend_map.get(trend, trend)
        
        trend_color = {'Artan': '#27ae60', 'Azalan': '#e74c3c', 'Stabil': '#f39c12'}.get(trend, '#7f8c8d')
        tk.Label(card, text=f"{self.lm.tr('trend', 'Trend')}: {display_trend}", font=('Segoe UI', 9),
                fg=trend_color, bg='#f8f9fa').pack(pady=(0, 5))

    def load_capitals_data(self) -> None:
        """6 Sermaye verilerini yükle"""
        try:
            capitals = self.manager.get_six_capitals(self.company_id)

            # Sol panel - Sermaye kartları
            for widget in self.capitals_content.winfo_children():
                widget.destroy()

            for capital in capitals:
                self.create_capital_detail_card(self.capitals_content, capital)

            # Sağ panel - Özet bilgiler
            if hasattr(self, 'capitals_info_frame'):
                for widget in self.capitals_info_frame.winfo_children():
                    widget.destroy()

                # Genel bilgi
                info_text = self.lm.tr('iirc_capitals_info', """
6 SERMAYE MODELİ HAKKINDA

Organizasyonlar, değer yaratmak için 6 farklı 
sermaye türünü kullanır:

1️⃣ Mali Sermaye
   Finansal kaynaklar ve yatırımlar

2️⃣ İmalat Sermayesi
   Fiziksel varlıklar ve altyapı

3️⃣ Entelektüel Sermaye
   Bilgi, patent ve markalar

4️⃣ İnsan Sermayesi
   Çalışan yetkinlikleri ve deneyim

5️⃣ Sosyal Sermaye
   Paydaş ilişkileri ve ağlar

6️⃣ Doğal Sermaye
   Çevresel kaynaklar
                """).strip()

                tk.Label(self.capitals_info_frame, text=info_text,
                        font=('Courier New', 9), bg='#f8f9fa',
                        justify='left').pack(anchor='w', pady=10)
                
                # İstatistikler
                tk.Label(self.capitals_info_frame, text="─────────────────────────",
                        font=('Segoe UI', 9), bg='#f8f9fa').pack(pady=5)
                
                stats_text = f"""
{self.lm.tr('current_status', 'MEVCUT DURUM')}:

{self.lm.tr('total_capital', 'Toplam Sermaye')}: {len(capitals)}
{self.lm.tr('active_capital', 'Aktif Sermaye')}: {len([c for c in capitals if c.get('status') == 'Aktif'])}

{self.lm.tr('capital_values', 'Sermaye Değerleri')}:
                """.strip()
                
                tk.Label(self.capitals_info_frame, text=stats_text,
                        font=('Courier New', 9), bg='#f8f9fa',
                        justify='left').pack(anchor='w', pady=5)

                # Her sermaye için değer göstergesi
                for capital in capitals:
                    capital_frame = tk.Frame(self.capitals_info_frame, bg='#f8f9fa')
                    capital_frame.pack(fill='x', pady=3)

                    name = capital.get('capital_name', '')[:20]  # Kısa isim
                    value = capital.get('current_value', 0)

                    tk.Label(capital_frame, text=f"{name}:",
                            font=('Segoe UI', 9), bg='#f8f9fa',
                            width=20, anchor='w').pack(side='left')

                    # Progress bar
                    progress = ttk.Progressbar(capital_frame, length=150,
                                              mode='determinate', value=min(value, 100))
                    progress.pack(side='left', padx=(5, 5))

                    tk.Label(capital_frame, text=f"{value:.1f}",
                            font=('Segoe UI', 9, 'bold'), bg='#f8f9fa',
                            fg='#6A1B9A').pack(side='left')

                # Rapor butonu
                tk.Label(self.capitals_info_frame, text="─────────────────────────",
                        font=('Segoe UI', 9), bg='#f8f9fa').pack(pady=10)

                ttk.Button(self.capitals_info_frame, text=f"{Icons.REPORT} {self.lm.tr('create_six_capitals_report', '6 Sermaye Raporu Oluştur')}",
                           command=self.generate_capitals_report,
                           style='Primary.TButton').pack(pady=5)

                ttk.Button(self.capitals_info_frame, text=f"{Icons.CHART_UP} {self.lm.tr('capital_charts', 'Sermaye Grafikleri')}",
                           command=self.show_capitals_charts,
                           style='Primary.TButton').pack(pady=5)

        except Exception as e:
            logging.error(f"Sermaye verileri yüklenirken hata: {e}")

    def create_capital_detail_card(self, parent, capital):
        """Sermaye detay kartı oluştur"""
        card = tk.Frame(parent, bg='white', relief='solid', bd=1)
        card.pack(fill='x', padx=10, pady=5)

        # Başlık
        header_frame = tk.Frame(card, bg='#6A1B9A')
        header_frame.pack(fill='x')

        icon = capital.get('icon', '')
        name = capital.get('capital_name', self.lm.tr('capital', 'Sermaye'))

        tk.Label(header_frame, text=f"{icon} {name}",
                font=('Segoe UI', 14, 'bold'), fg='white', bg='#6A1B9A').pack(side='left', padx=10, pady=5)

        # İçerik
        content_frame = tk.Frame(card, bg='white')
        content_frame.pack(fill='x', padx=15, pady=10)

        description = capital.get('description', '')
        tk.Label(content_frame, text=description, font=('Segoe UI', 10),
                bg='white', wraplength=800, justify='left').pack(anchor='w')

        # Detayları görüntüle butonu
        def show_details():
            self.show_capital_details(capital)

        ttk.Button(content_frame, text=f" {self.lm.tr('view_details', 'Detayları Görüntüle')}", command=show_details,
                   style='Primary.TButton').pack(anchor='w', pady=(10, 0))

    def show_capital_details(self, capital):
        """Sermaye detaylarını göster"""
        # Detay penceresi oluştur
        detail_window = tk.Toplevel(self.parent)
        detail_window.title(f"{capital.get('capital_name', self.lm.tr('capital', 'Sermaye'))} - {self.lm.tr('details', 'Detaylar')}")
        detail_window.geometry("600x500")

        # Başlık
        header = tk.Frame(detail_window, bg='#6A1B9A')
        header.pack(fill='x')

        tk.Label(header, text=f"{capital.get('icon', '')} {capital.get('capital_name', self.lm.tr('capital', 'Sermaye'))}",
                font=('Segoe UI', 14, 'bold'), fg='white', bg='#6A1B9A').pack(pady=10)

        # İçerik
        content = tk.Frame(detail_window, bg='white')
        content.pack(fill='both', expand=True, padx=20, pady=20)

        # Bilgiler
        info_items = [
            (self.lm.tr('description', 'Açıklama'), capital.get('description', '-')),
            (self.lm.tr('current_value', 'Mevcut Değer'), f"{capital.get('current_value', 0):.2f}"),
            (self.lm.tr('change_from_prev', 'Önceki Dönem Değişimi'), f"{capital.get('change_from_previous', 0):.2f}"),
            (self.lm.tr('trend', 'Trend'), capital.get('trend', 'Stabil')),
            (self.lm.tr('status', 'Durum'), capital.get('status', 'Aktif'))
        ]

        for label, value in info_items:
            frame = tk.Frame(content, bg='white')
            frame.pack(fill='x', pady=5)

            tk.Label(frame, text=f"{label}:", font=('Segoe UI', 10, 'bold'),
                    bg='white').pack(side='left')
            tk.Label(frame, text=value, font=('Segoe UI', 10),
                    bg='white').pack(side='left', padx=(10, 0))

        # Kapat butonu
        ttk.Button(detail_window, text=self.lm.tr('btn_close', "Kapat"), command=detail_window.destroy,
                   style='Primary.TButton').pack(pady=10)

    def load_value_creation_data(self) -> None:
        """Değer yaratma hikayesi verilerini yükle"""
        try:
            template = self.value_story.get_story_template()

            # İçeriği temizle
            for widget in self.value_content.winfo_children():
                widget.destroy()

            for element in template:
                self.create_value_element_card(self.value_content, element)

        except Exception as e:
            logging.error(f"Değer yaratma verileri yüklenirken hata: {e}")

    def create_value_element_card(self, parent, element):
        """Değer yaratma öğesi kartı oluştur"""
        card = tk.Frame(parent, bg='#f8f9fa', relief='solid', bd=1)
        card.pack(fill='x', padx=10, pady=5)

        # Başlık
        tk.Label(card, text=element['baslik'], font=('Segoe UI', 12, 'bold'),
                bg='#f8f9fa', fg='#6A1B9A').pack(anchor='w', padx=10, pady=5)

        # Açıklama
        tk.Label(card, text=element['aciklama'], font=('Segoe UI', 10),
                bg='#f8f9fa', wraplength=800, justify='left').pack(anchor='w', padx=10)

        # Sorular
        sorular_text = "\n".join([f"• {s}" for s in element['sorular']])
        tk.Label(card, text=sorular_text, font=('Segoe UI', 9),
                bg='#f8f9fa', fg='#7f8c8d', justify='left').pack(anchor='w', padx=20, pady=(5, 10))

    def load_connectivity_data(self) -> None:
        """Bağlantı haritası verilerini yükle"""
        try:
            connections = self.manager.get_connectivity_map(self.company_id)

            # Treeview'ı temizle
            for item in self.connectivity_tree.get_children():
                self.connectivity_tree.delete(item)

            for conn in connections:
                self.connectivity_tree.insert('', 'end', values=(
                    f"{conn.get('source_icon', '')} {conn.get('source_name', '')}",
                    f"{conn.get('target_icon', '')} {conn.get('target_name', '')}",
                    conn.get('connection_type', ''),
                    conn.get('connection_strength', '')
                ))

        except Exception as e:
            logging.error(f"Bağlantı verileri yüklenirken hata: {e}")

    def generate_template(self) -> None:
        """Entegre rapor şablonu oluştur"""
        try:
            from datetime import datetime
            from tkinter import filedialog

            # Basit şablon dökümanı oluştur
            filepath = filedialog.asksaveasfilename(
                title=self.lm.tr('save_template', "Şablon Kaydet"),
                defaultextension=".docx",
                filetypes=[(self.lm.tr('word_file', "Word Dosyası"), "*.docx")],
                initialfile=f'IIRC_Sablon_{datetime.now().year}.docx'
            )

            if not filepath:
                return

            from docx import Document
            doc = Document()
            doc.add_heading(self.lm.tr('integrated_report_template_title', 'ENTEGRE RAPOR ŞABLONU'), 0)
            doc.add_paragraph(self.lm.tr('integrated_report_template_desc', 'Bu şablon, IIRC Çerçevesine uygun entegre rapor hazırlamak için kullanılabilir.'))

            sections = [
                self.lm.tr('section_org_overview', '1. Organizasyona Genel Bakış'),
                self.lm.tr('section_governance', '2. Yönetişim'),
                self.lm.tr('section_business_model', '3. İş Modeli'),
                self.lm.tr('section_risks', '4. Riskler ve Fırsatlar'),
                self.lm.tr('section_strategy', '5. Strateji ve Kaynak Tahsisi'),
                self.lm.tr('section_performance', '6. Performans'),
                self.lm.tr('section_future', '7. Gelecek Görünümü'),
                self.lm.tr('section_basis', '8. Raporun Hazırlanma Temeli')
            ]

            for section in sections:
                doc.add_heading(section, 1)
                doc.add_paragraph('[İçerik buraya girilecek]')
                doc.add_paragraph()

            doc.save(filepath)
            messagebox.showinfo(self.lm.tr('success', "Başarılı"), f"{self.lm.tr('template_created', 'Şablon oluşturuldu!')}\n\n{filepath}")

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('err_template_create', 'Şablon oluşturulamadı')}: {str(e)}")

    def generate_capitals_report(self) -> None:
        """6 Sermaye raporu oluştur"""
        try:
            import threading
            from datetime import datetime
            from tkinter import filedialog

            from .reports.capitals_report import CapitalsReportGenerator

            year = datetime.now().year
            filepath = filedialog.asksaveasfilename(
                title=self.lm.tr('save_report', "Raporu Kaydet"),
                defaultextension=".docx",
                filetypes=[(self.lm.tr('word_file', "Word Dosyası"), "*.docx")],
                initialfile=f'IIRC_6_Sermaye_{year}.docx'
            )

            if not filepath:
                return

            def generate():
                try:
                    reporter = CapitalsReportGenerator()
                    success = reporter.generate_capitals_report(self.company_id, year, filepath)

                    if success:
                        messagebox.showinfo(self.lm.tr('success', "Başarılı"), f"{self.lm.tr('capitals_report_created', '6 Sermaye raporu oluşturuldu!')}\n\n{filepath}")
                    else:
                        messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_report_create', "Rapor oluşturulamadı!"))
                except Exception as e:
                    messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('error', 'Hata')}: {str(e)}")

            thread = threading.Thread(target=generate)
            thread.start()

        except Exception as e:
            messagebox.showerror("Hata", f"Hata: {str(e)}")


    def generate_value_report(self) -> None:
        """Değer yaratma raporu oluştur"""
        try:
            import threading
            from datetime import datetime
            from tkinter import filedialog

            from .reports.value_creation_report import ValueCreationReportGenerator

            year = datetime.now().year
            filepath = filedialog.asksaveasfilename(
                title=self.lm.tr('save_report', "Raporu Kaydet"),
                defaultextension=".docx",
                filetypes=[(self.lm.tr('word_file', "Word Dosyası"), "*.docx")],
                initialfile=f'IIRC_Deger_Yaratma_{year}.docx'
            )

            if not filepath:
                return

            def generate():
                try:
                    reporter = ValueCreationReportGenerator()
                    success = reporter.generate_value_report(self.company_id, year, filepath)

                    if success:
                        messagebox.showinfo(self.lm.tr('success', "Başarılı"), f"{self.lm.tr('value_report_created', 'Değer yaratma raporu oluşturuldu!')}\n\n{filepath}")
                    else:
                        messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_report_create', "Rapor oluşturulamadı!"))
                except Exception as e:
                    messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('error', 'Hata')}: {str(e)}")

            thread = threading.Thread(target=generate)
            thread.start()

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('error', 'Hata')}: {str(e)}")


    def generate_connectivity_report(self) -> None:
        """Bağlantı analizi raporu oluştur"""
        try:
            from datetime import datetime
            from tkinter import filedialog

            year = datetime.now().year
            filepath = filedialog.asksaveasfilename(
                title=self.lm.tr('save_report', "Raporu Kaydet"),
                defaultextension=".xlsx",
                filetypes=[(self.lm.tr('excel_file', "Excel Dosyası"), "*.xlsx")],
                initialfile=f'IIRC_Baglanti_Analizi_{year}.xlsx'
            )

            if not filepath:
                return

            # Excel ile bağlantı analizi
            from openpyxl import Workbook
            from openpyxl.styles import Font as ExcelFont
            wb = Workbook()
            ws = wb.active
            ws.title = self.lm.tr('connectivity_analysis', "Bağlantı Analizi")

            # Başlıklar
            ws['A1'] = self.lm.tr('capital_connectivity_analysis', 'SERMAYE BAĞLANTI ANALİZİ')
            ws['A1'].font = ExcelFont(size=14, bold=True)

            connections = self.manager.get_connectivity_map(self.company_id, year)

            headers = [
                self.lm.tr('source_capital', 'Kaynak Sermaye'),
                self.lm.tr('target_capital', 'Hedef Sermaye'),
                self.lm.tr('connection_type', 'Bağlantı Türü'),
                self.lm.tr('strength', 'Güç'),
                self.lm.tr('description', 'Açıklama')
            ]
            for col, header in enumerate(headers, 1):
                ws.cell(row=3, column=col, value=header)

            for row_idx, conn in enumerate(connections, 4):
                ws.cell(row=row_idx, column=1, value=conn.get('source_name', '-'))
                ws.cell(row=row_idx, column=2, value=conn.get('target_name', '-'))
                ws.cell(row=row_idx, column=3, value=conn.get('connection_type', '-'))
                ws.cell(row=row_idx, column=4, value=conn.get('connection_strength', '-'))
                ws.cell(row=row_idx, column=5, value=conn.get('description', '-'))

            # Sütun genişlikleri
            ws.column_dimensions['A'].width = 25
            ws.column_dimensions['B'].width = 25
            ws.column_dimensions['C'].width = 20
            ws.column_dimensions['D'].width = 15
            ws.column_dimensions['E'].width = 50

            wb.save(filepath)
            messagebox.showinfo(self.lm.tr('success', "Başarılı"), f"{self.lm.tr('connectivity_report_created', 'Bağlantı analizi oluşturuldu!')}\n\n{filepath}")

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('error', 'Hata')}: {str(e)}")

    def generate_full_report(self) -> None:
        """Tam entegre rapor oluştur"""
        try:
            import threading
            from datetime import datetime
            from tkinter import filedialog

            from .reports.integrated_report import IntegratedReportGenerator

            year = datetime.now().year
            filepath = filedialog.asksaveasfilename(
                title=self.lm.tr('save_report', "Raporu Kaydet"),
                defaultextension=".docx",
                filetypes=[(self.lm.tr('word_file', "Word Dosyası"), "*.docx")],
                initialfile=f'IIRC_Entegre_Rapor_{year}.docx'
            )

            if not filepath:
                return

            def generate():
                try:
                    reporter = IntegratedReportGenerator()
                    success = reporter.generate_full_report(self.company_id, year, filepath)

                    if success:
                        messagebox.showinfo(self.lm.tr('success', "Başarılı"),
                            f"{self.lm.tr('full_report_created', 'Tam Entegre Rapor oluşturuldu!')}\n\n{filepath}\n\n" +
                            self.lm.tr('report_compliance_msg', "Rapor IIRC Çerçevesine uygun olarak hazırlanmıştır."))
                    else:
                        messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_report_create', "Rapor oluşturulamadı!"))
                except Exception as e:
                    messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('error', 'Hata')}: {str(e)}")

            thread = threading.Thread(target=generate)
            thread.start()

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('error', 'Hata')}: {str(e)}")

    def show_capitals_charts(self) -> None:
        """Sermaye grafiklerini göster"""
        try:
            from datetime import datetime

            import matplotlib.pyplot as plt

            year = datetime.now().year
            capitals = self.manager.get_six_capitals(self.company_id, year)

            if not capitals:
                messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('no_capital_data', "Henüz sermaye verisi bulunmamaktadır."))
                return

            # Radar chart oluştur
            fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(projection='polar'))

            categories = [c.get('capital_name', '') for c in capitals]
            values = [c.get('current_value', 0) for c in capitals]

            # Normalize
            max_val = max(values) if values else 1
            normalized_values = [(v / max_val) * 100 if max_val > 0 else 0 for v in values]

            # Kapalı şekil için
            categories + [categories[0]]
            values_closed = normalized_values + [normalized_values[0]]

            # Açılar
            import numpy as np
            angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False).tolist()
            angles_closed = angles + [angles[0]]

            # Çiz
            ax.plot(angles_closed, values_closed, 'o-', linewidth=2, color='#6A1B9A', label=self.lm.tr('current_value', 'Mevcut Değer'))
            ax.fill(angles_closed, values_closed, alpha=0.25, color='#6A1B9A')

            ax.set_xticks(angles)
            ax.set_xticklabels(categories, size=9)
            ax.set_ylim(0, 100)
            ax.set_title(self.lm.tr('six_capitals_analysis', '6 Sermaye Değer Analizi'), size=14, fontweight='bold', pad=20)
            ax.grid(True)
            ax.legend(loc='upper right')

            plt.tight_layout()
            plt.show()

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('err_chart_create', 'Grafik oluşturulamadı')}: {str(e)}")
