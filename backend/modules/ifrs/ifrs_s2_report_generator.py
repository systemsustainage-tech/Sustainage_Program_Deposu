#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
IFRS S2 REPORT GENERATOR - TCFD Verilerinden IFRS S2 DOCX
- Governance, Strategy, Risk Management, Metrics & Targets bölümleri
- TCFDManager üzerinden veri çekerek IFRS S2 başlıklarına yerleştirir
"""

import logging
import os

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.info("[WARN] python-docx not available. DOCX generation disabled.")


class IFRSS2ReportGenerator:
    """IFRS S2 rapor üretici"""

    def __init__(self, manager):
        """
        Args:
            manager: TCFDManager instance (TCFD verilerini sağlar)
        """
        self.manager = manager

    def _add_heading(self, doc: Document, text: str, level: int = 1) -> None:
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Calibri'
            # Varsayılan başlık boyutları: 0:25, 1:17, 2:14, 3:13
            default_sizes = {0: 25, 1: 17, 2: 14, 3: 13}
            run.font.size = Pt(default_sizes.get(level, 13))

    def _add_paragraph(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph(text)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    def generate_docx(
        self,
        company_id: int,
        year: int,
        company_name: str,
        output_path: str,
    ) -> bool:
        """
        IFRS S2 DOCX raporu üret

        Returns:
            Başarılı mı?
        """
        if not DOCX_AVAILABLE:
            logging.error("[ERROR] python-docx not installed")
            return False

        try:
            doc = Document()
            self._add_heading(doc, f"IFRS S2 - İklim İlişkili Açıklamalar: {company_name} ({year})", level=0)

            # Governance
            self._add_heading(doc, "Governance", level=1)
            gov = self.manager.get_governance(company_id, year)
            if gov:
                self._add_paragraph(doc, f"Yönetim Kurulu Gözetimi: {gov.get('board_oversight', '')}")
                self._add_paragraph(doc, f"Yönetim Rolü: {gov.get('management_role', '')}")
                self._add_paragraph(doc, f"Komite: {gov.get('committee_name', '')} (Var: {gov.get('climate_committee', False)})")
            else:
                self._add_paragraph(doc, "Governance verisi bulunamadı.")

            # Strategy
            self._add_heading(doc, "Strategy", level=1)
            strat = self.manager.get_strategy(company_id, year)
            if strat:
                self._add_paragraph(doc, f"Kısa Vadeli Riskler: {strat.get('short_term_risks', '')}")
                self._add_paragraph(doc, f"Kısa Vadeli Fırsatlar: {strat.get('short_term_opportunities', '')}")
                self._add_paragraph(doc, f"Dayanıklılık/Senaryolar: {strat.get('resilience_scenarios', '')}")
            else:
                self._add_paragraph(doc, "Strategy verisi bulunamadı.")

            # Risk Management
            self._add_heading(doc, "Risk Management", level=1)
            risks = self.manager.get_climate_risks(company_id, year) or []
            if risks:
                self._add_paragraph(doc, f"İklim risk kayıt sayısı: {len(risks)}")
                for r in risks[:5]:  # ilk 5 risk için özet
                    self._add_paragraph(doc, f"- {r.get('risk_name', '')} ({r.get('risk_category', '')}/{r.get('risk_type', '')}) - Etki: {r.get('impact', '')}, Olasılık: {r.get('likelihood', '')}")
            else:
                self._add_paragraph(doc, "Risk verisi bulunamadı.")

            # Metrics & Targets
            self._add_heading(doc, "Metrics and Targets", level=1)
            metrics = self.manager.get_metrics(company_id, year)
            if metrics:
                self._add_paragraph(doc, f"Scope 1 Emisyonlar: {metrics.get('scope1_emissions', 'n/a')} tCO2e")
                self._add_paragraph(doc, f"Scope 2 Emisyonlar: {metrics.get('scope2_emissions', 'n/a')} tCO2e")
                self._add_paragraph(doc, f"Toplam Emisyonlar: {metrics.get('total_emissions', 'n/a')} tCO2e")
                self._add_paragraph(doc, f"Emisyon Yoğunluğu: {metrics.get('emissions_intensity', 'n/a')} ({metrics.get('intensity_metric', '')})")
                self._add_paragraph(doc, f"Yenilenebilir Enerji Oranı: {metrics.get('renewable_energy_pct', 'n/a')} %")
                self._add_paragraph(doc, f"İç Karbon Fiyatı: {metrics.get('internal_carbon_price', 'n/a')} TRY/tCO2e")
            else:
                self._add_paragraph(doc, "Metrics & Targets verisi bulunamadı.")

            # Kaydet
            out_dir = os.path.dirname(output_path)
            if out_dir and not os.path.exists(out_dir):
                os.makedirs(out_dir, exist_ok=True)
            doc.save(output_path)
            logging.info(f"[IFRS S2] DOCX rapor oluşturuldu: {output_path}")
            return True

        except Exception as e:
            logging.error(f"[ERROR] IFRS S2 rapor oluşturulamadı: {e}")
            return False
