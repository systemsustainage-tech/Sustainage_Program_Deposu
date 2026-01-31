#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SOSYAL RAPORLAMA SINIFI
İK, İSG ve Eğitim raporları
"""

import logging
import os
import sqlite3
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.info("[UYARI] python-docx yüklü değil. DOCX raporlar oluşturulamaz.")

try:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logging.info("[UYARI] openpyxl yuklu degil. Excel raporlar olusturulamaz.")

from config.settings import ensure_directories, get_db_path, get_export_dir
from utils.language_manager import LanguageManager

from .hr_metrics import HRMetrics
from .ohs_metrics import OHSMetrics
from .training_metrics import TrainingMetrics


def _add_turkish_paragraph(doc, text, style=None, font_name='Calibri', font_size=11):
    """Türkçe karakter destekli paragraf ekle"""
    para = doc.add_paragraph(text if text is not None else '', style=style)
    if not text:
        return para
    for run in para.runs:
        try:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakter destekli başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        try:
            run.font.name = font_name
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return heading


class SocialReporting:
    """Sosyal performans raporlama sınıfı"""

    def __init__(self, db_path: str | None = None) -> None:
        if db_path:
            self.db_path = db_path
        else:
            ensure_directories()
            self.db_path = get_db_path()
        
        self.lm = LanguageManager()
        self.hr = HRMetrics()
        self.ohs = OHSMetrics()
        self.training = TrainingMetrics()

    def generate_social_report(self, company_id: int, period: str, 
                             formats: List[str] = None) -> Dict[str, str]:
        """Web uygulaması için birleştirilmiş raporlama fonksiyonu"""
        if formats is None:
            formats = ['docx']
            
        generated_files = {}
        try:
            year = int(period) if period and period.isdigit() else datetime.now().year
        except ValueError:
            year = datetime.now().year
        
        if 'docx' in formats:
            # Varsayılan olarak kapsamlı raporu kullan
            path = self.generate_comprehensive_report(company_id, year)
            if path:
                generated_files['docx'] = path
                
        if 'excel' in formats:
            path = self.generate_excel_export(company_id, year)
            if path:
                generated_files['excel'] = path
                
        return generated_files

    def _add_logo(self, doc, company_id: int):
        """Raporun başına şirket logosunu ekle"""
        try:
            # 1. Veritabanından logo yolunu al
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT logo_path FROM company_profiles WHERE company_id = ?", (company_id,))
            result = cursor.fetchone()
            conn.close()
            
            logo_path = None
            if result and result[0] and os.path.exists(result[0]):
                logo_path = result[0]
            else:
                # 2. Varsayılan yolda ara
                # self.db_path genellikle .../data/sdg_desktop.sqlite
                data_dir = os.path.dirname(self.db_path)
                possible_path = os.path.join(data_dir, "company_logos", f"company_{company_id}_logo.png")
                if os.path.exists(possible_path):
                    logo_path = possible_path
                else:
                    # jpg dene
                    possible_path = possible_path.replace(".png", ".jpg")
                    if os.path.exists(possible_path):
                        logo_path = possible_path

            if logo_path:
                section = doc.sections[0]
                header = section.header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(logo_path, height=Cm(2.5))
        except Exception as e:
            logging.warning(f"Logo eklenirken hata: {e}")

    def generate_hr_report(self, company_id: int, year: int) -> str:
        """İK Metrikleri Raporu (DOCX)"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kütüphanesi gerekli")

        hr_data = self.hr.get_workforce_summary(company_id, year)
        
        filename = f"IK_Metrikleri_{company_id}_{year}.docx"
        export_dir = str(get_export_dir(company_id))
        filepath = os.path.join(export_dir, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        doc = Document()
        
        # Logo ekle
        self._add_logo(doc, company_id)
        
        # Başlık
        title = _add_turkish_heading(doc, self.lm.tr('hr_report_title', 'İNSAN KAYNAKLARI METRİKLERİ RAPORU'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        _add_turkish_paragraph(doc, f"{self.lm.tr('report_date', 'Rapor Tarihi')}: {datetime.now().strftime('%d.%m.%Y')}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('reporting_period', 'Raporlama Yılı')}: {year}")

        doc.add_page_break()

        # 1. Çalışan Profili
        _add_turkish_heading(doc, f"1. {self.lm.tr('section_employee_profile', 'Çalışan Profili')}", 1)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = self.lm.tr('metric', 'Metrik')
        hdr[1].text = self.lm.tr('value', 'Değer')

        metrics = [
            (self.lm.tr('total_employees', 'Toplam Çalışan'), f"{hr_data.get('total_employees', 0)}"),
            (self.lm.tr('female_employees', 'Kadın Çalışan'), f"{hr_data.get('female_count', 0)} (%{hr_data.get('female_percentage', 0):.1f})"),
            (self.lm.tr('male_employees', 'Erkek Çalışan'), f"{hr_data.get('male_count', 0)} (%{hr_data.get('male_percentage', 0):.1f})"),
            (self.lm.tr('turnover_rate', 'Turnover Oranı'), f"%{hr_data.get('turnover_rate', 0):.1f}")
        ]

        for label, value in metrics:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

        # 2. İşe Alım ve Ayrılma
        _add_turkish_heading(doc, f"2. {self.lm.tr('recruitment_separation', 'İşe Alım ve Ayrılma')}", 1)
        _add_turkish_paragraph(doc, f"{self.lm.tr('new_hires', 'Yeni İşe Alım')}: {hr_data.get('new_hires', 0)}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('terminations', 'Ayrılma')}: {hr_data.get('terminations', 0)}")

        # Finansal/Maliyet Bilgileri (Varsa)
        # Not: Şu an için veritabanında bu detaylar yok, ancak yapı olarak ekliyoruz.
        _add_turkish_heading(doc, f"3. {self.lm.tr('hr_financials', 'İK Maliyetleri ve Dış Kaynaklar')}", 1)
        _add_turkish_paragraph(doc, self.lm.tr('no_financial_data', "Bu dönem için detaylı finansal veri bulunmamaktadır."))

        doc.save(filepath)
        return filepath

    def generate_comprehensive_report(self, company_id: int, year: int) -> str:
        """Kapsamlı Sosyal Performans Raporu (DOCX)"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kütüphanesi gerekli")

        filename = f"Sosyal_Performans_{company_id}_{year}.docx"
        export_dir = str(get_export_dir(company_id))
        filepath = os.path.join(export_dir, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        doc = Document()
        
        # Logo ekle
        self._add_logo(doc, company_id)
        
        title = _add_turkish_heading(doc, self.lm.tr('social_report_title', 'SOSYAL PERFORMANS RAPORU'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = _add_turkish_paragraph(doc, f"Sustainage - {year}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        
        _add_turkish_paragraph(doc, f"{self.lm.tr('report_date', 'Rapor Tarihi')}: {datetime.now().strftime('%d.%m.%Y')}")

        # İK Bölümü
        doc.add_page_break()
        _add_turkish_heading(doc, f"1. {self.lm.tr('section_hr_metrics', 'İNSAN KAYNAKLARI METRİKLERİ')}", 1)
        hr_data = self.hr.get_workforce_summary(company_id, year)
        _add_turkish_paragraph(doc, f"{self.lm.tr('total_employees', 'Toplam Çalışan')}: {hr_data.get('total_employees', 0)}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('female_employees_rate', 'Kadın Çalışan Oranı')}: %{hr_data.get('female_percentage', 0):.1f}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('turnover_rate', 'Turnover Oranı')}: %{hr_data.get('turnover_rate', 0):.1f}")

        # İSG Bölümü
        doc.add_page_break()
        _add_turkish_heading(doc, f"2. {self.lm.tr('section_ohs', 'İŞ SAĞLIĞI VE GÜVENLİĞİ')}", 1)
        ohs_data = self.ohs.get_summary(company_id, year)
        _add_turkish_paragraph(doc, f"{self.lm.tr('total_incidents', 'Toplam Kaza')}: {ohs_data.get('total_incidents', 0)}")
        _add_turkish_paragraph(doc, f"LTIFR: {ohs_data.get('ltifr', 0):.2f}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('total_lost_days', 'Toplam Kayıp Gün')}: {ohs_data.get('total_lost_days', 0)}")

        # Eğitim Bölümü
        doc.add_page_break()
        _add_turkish_heading(doc, f"3. {self.lm.tr('section_training_dev', 'EĞİTİM VE GELİŞİM')}", 1)
        training_data = self.training.get_summary(company_id, year)
        _add_turkish_paragraph(doc, f"{self.lm.tr('total_programs', 'Toplam Eğitim Programı')}: {training_data.get('total_programs', 0)}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('avg_hours_per_person', 'Kişi Başı Eğitim Saati')}: {training_data.get('avg_hours_per_employee', 0):.1f}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('completion_rate', 'Tamamlanma Oranı')}: %{training_data.get('completion_rate', 0):.1f}")

        doc.save(filepath)
        return filepath

    def generate_excel_export(self, company_id: int, year: int) -> str:
        """Tüm verileri Excel'e aktar"""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl kütüphanesi gerekli")

        filename = f"Sosyal_Metrikler_{company_id}_{year}.xlsx"
        export_dir = str(get_export_dir(company_id))
        filepath = os.path.join(export_dir, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        wb = openpyxl.Workbook()
        wb.remove(wb.active)

        # İK Sheet
        ws_hr = wb.create_sheet(self.lm.tr('sheet_hr_metrics', "İK Metrikleri"))
        # ... (İK kodu aynı kalacak) ...
        ws_hr['A1'] = self.lm.tr('hr_metrics_title', 'İNSAN KAYNAKLARI METRİKLERİ')
        ws_hr['A1'].font = Font(size=14, bold=True)

        hr_data = self.hr.get_workforce_summary(company_id, year)
        rows = [
            (self.lm.tr('total_employees', 'Toplam Çalışan'), hr_data.get('total_employees', 0)),
            (self.lm.tr('female', 'Kadın'), hr_data.get('female_count', 0)),
            (self.lm.tr('male', 'Erkek'), hr_data.get('male_count', 0)),
            (self.lm.tr('turnover_rate', 'Turnover Oranı') + ' (%)', hr_data.get('turnover_rate', 0))
        ]

        for i, (label, value) in enumerate(rows, 3):
            ws_hr.cell(row=i, column=1, value=label)
            ws_hr.cell(row=i, column=2, value=value)

        # İSG Sheet
        ws_ohs = wb.create_sheet(self.lm.tr('sheet_ohs', "İSG"))
        ws_ohs['A1'] = self.lm.tr('ohs_title', 'İŞ SAĞLIĞI VE GÜVENLİĞİ')
        ws_ohs['A1'].font = Font(size=14, bold=True)

        ohs_data = self.ohs.get_summary(company_id, year)
        rows = [
            (self.lm.tr('total_incidents', 'Toplam Kaza'), ohs_data.get('total_incidents', 0)),
            ('LTIFR', ohs_data.get('ltifr', 0)),
            ('TRIR', ohs_data.get('trir', 0)),
            (self.lm.tr('total_lost_days', 'Kayıp Gün'), ohs_data.get('total_lost_days', 0))
        ]

        for i, (label, value) in enumerate(rows, 3):
            ws_ohs.cell(row=i, column=1, value=label)
            ws_ohs.cell(row=i, column=2, value=value)

        # Eğitim Sheet
        ws_train = wb.create_sheet(self.lm.tr('sheet_training', "Eğitim"))
        ws_train['A1'] = self.lm.tr('training_dev_title', 'EĞİTİM VE GELİŞİM')
        ws_train['A1'].font = Font(size=14, bold=True)

        training_data = self.training.get_summary(company_id, year)
        rows = [
            (self.lm.tr('total_programs', 'Toplam Program'), training_data.get('total_programs', 0)),
            (self.lm.tr('total_hours', 'Toplam Saat'), training_data.get('total_hours', 0)),
            (self.lm.tr('hours_per_person', 'Kişi Başı Saat'), training_data.get('avg_hours_per_employee', 0)),
            (self.lm.tr('completion_pct', 'Tamamlanma (%)'), training_data.get('completion_rate', 0))
        ]

        for i, (label, value) in enumerate(rows, 3):
            ws_train.cell(row=i, column=1, value=label)
            ws_train.cell(row=i, column=2, value=value)

        # Finansal Veri Sheet
        ws_fin = wb.create_sheet(self.lm.tr('sheet_financials', "Finansal Veriler"))
        ws_fin['A1'] = self.lm.tr('financial_data', 'FİNANSAL VERİLER VE MALİYETLER')
        ws_fin['A1'].font = Font(size=14, bold=True)
        
        headers = ['Program', 'Tedarikçi', 'Tutar', 'Para Birimi', 'Fatura Tarihi']
        for col, header in enumerate(headers, 1):
            cell = ws_fin.cell(row=3, column=col, value=header)
            cell.font = Font(bold=True)

        # Veritabanından finansal verileri çek
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT program_name, supplier, total_cost, currency, invoice_date 
                FROM training_programs 
                WHERE company_id = ?
            """, (company_id,))
            financial_records = cursor.fetchall()
            conn.close()
        except Exception as e:
            logging.error(f"Excel finansal veri hatası: {e}")
            financial_records = []
            
        if financial_records:
            for i, record in enumerate(financial_records, 4):
                ws_fin.cell(row=i, column=1, value=record[0])
                ws_fin.cell(row=i, column=2, value=record[1])
                ws_fin.cell(row=i, column=3, value=record[2])
                ws_fin.cell(row=i, column=4, value=record[3])
                ws_fin.cell(row=i, column=5, value=record[4])
        else:
            ws_fin.cell(row=4, column=1, value=self.lm.tr('no_data', "Veri Bulunamadı"))

        wb.save(filepath)
        return filepath

    def generate_ohs_report(self, company_id: int, year: int) -> str:
        """İSG Raporu (DOCX)"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kütüphanesi gerekli")

        ohs_data = self.ohs.get_summary(company_id, year)
        
        filename = f"ISG_Raporu_{company_id}_{year}.docx"
        export_dir = str(get_export_dir(company_id))
        filepath = os.path.join(export_dir, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        doc = Document()
        
        # Logo ekle
        self._add_logo(doc, company_id)
        
        title = _add_turkish_heading(doc, self.lm.tr('ohs_report_title', 'İŞ SAĞLIĞI VE GÜVENLİĞİ RAPORU'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        _add_turkish_paragraph(doc, f"{self.lm.tr('report_date', 'Rapor Tarihi')}: {datetime.now().strftime('%d.%m.%Y')}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('reporting_period', 'Raporlama Yılı')}: {year}")

        doc.add_page_break()

        # 1. Kaza İstatistikleri
        _add_turkish_heading(doc, f"1. {self.lm.tr('section_incident_stats', 'Kaza İstatistikleri')}", 1)
        
        stats = [
            (self.lm.tr('total_incidents', 'Toplam Kaza'), ohs_data.get('total_incidents', 0)),
            (self.lm.tr('lost_time_incidents', 'Kayıp Günlü Kaza'), ohs_data.get('lost_time_incidents', 0)),
            (self.lm.tr('fatalities', 'Ölümlü Kaza'), ohs_data.get('fatalities', 0)),
            (self.lm.tr('total_lost_days', 'Toplam Kayıp Gün'), ohs_data.get('total_lost_days', 0))
        ]
        
        for label, val in stats:
            _add_turkish_paragraph(doc, f"• {label}: {val}")

        # 2. Performans Metrikleri
        _add_turkish_heading(doc, f"2. {self.lm.tr('section_ohs_metrics', 'Performans Metrikleri')}", 1)
        _add_turkish_paragraph(doc, f"LTIFR: {ohs_data.get('ltifr', 0):.2f}")
        _add_turkish_paragraph(doc, f"TRIR: {ohs_data.get('trir', 0):.2f}")

        doc.save(filepath)
        return filepath

    def generate_training_report(self, company_id: int, year: int) -> str:
        """Eğitim Raporu (DOCX)"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kütüphanesi gerekli")

        training_data = self.training.get_summary(company_id, year)
        
        filename = f"Egitim_Raporu_{company_id}_{year}.docx"
        export_dir = str(get_export_dir(company_id))
        filepath = os.path.join(export_dir, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        doc = Document()
        
        # Logo ekle
        self._add_logo(doc, company_id)
        
        title = _add_turkish_heading(doc, self.lm.tr('training_report_title', 'EĞİTİM VE GELİŞİM RAPORU'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        _add_turkish_paragraph(doc, f"{self.lm.tr('report_date', 'Rapor Tarihi')}: {datetime.now().strftime('%d.%m.%Y')}")
        _add_turkish_paragraph(doc, f"{self.lm.tr('reporting_period', 'Raporlama Yılı')}: {year}")

        doc.add_page_break()

        # 1. Eğitim Özeti
        _add_turkish_heading(doc, f"1. {self.lm.tr('section_training_stats', 'Eğitim İstatistikleri')}", 1)
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = self.lm.tr('metric', 'Metrik')
        hdr[1].text = self.lm.tr('value', 'Değer')

        metrics = [
            (self.lm.tr('total_programs', 'Toplam Program'), f"{training_data.get('total_programs', 0)}"),
            (self.lm.tr('total_participants', 'Toplam Katılımcı'), f"{training_data.get('total_participants', 0)}"),
            (self.lm.tr('total_training_hours', 'Toplam Eğitim Saati'), f"{training_data.get('total_hours', 0):.1f}"),
            (self.lm.tr('avg_hours_per_person', 'Kişi Başı Saat'), f"{training_data.get('avg_hours_per_employee', 0):.1f}"),
            (self.lm.tr('completion_rate', 'Tamamlanma Oranı'), f"%{training_data.get('completion_rate', 0):.1f}")
        ]

        for label, value in metrics:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

        # 2. Finansal Detaylar (Eğitim Yatırımları)
        _add_turkish_heading(doc, f"2. {self.lm.tr('training_financials', 'Eğitim Yatırımları')}", 1)
        
        fin_table = doc.add_table(rows=1, cols=5)
        fin_table.style = 'Table Grid'
        hdr = fin_table.rows[0].cells
        hdr[0].text = self.lm.tr('program_name', 'Program Adı')
        hdr[1].text = self.lm.tr('provider', 'Tedarikçi')
        hdr[2].text = self.lm.tr('cost', 'Tutar')
        hdr[3].text = self.lm.tr('currency', 'Para Birimi')
        hdr[4].text = self.lm.tr('invoice_date', 'Fatura Tarihi')

        # Veritabanından finansal verileri çek
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("""
                SELECT program_name, supplier, total_cost, currency, invoice_date 
                FROM training_programs 
                WHERE company_id = ?
            """, (company_id,))
            financial_records = cursor.fetchall()
            conn.close()
        except Exception as e:
            logging.error(f"Finansal veri çekme hatası: {e}")
            financial_records = []

        if financial_records:
            for record in financial_records:
                row = fin_table.add_row().cells
                row[0].text = str(record[0] or '-')
                row[1].text = str(record[1] or '-')
                row[2].text = str(record[2] or '-')
                row[3].text = str(record[3] or '-')
                row[4].text = str(record[4] or '-')
        else:
            row = fin_table.add_row().cells
            row[0].text = self.lm.tr('no_data', "Veri Yok")
            for i in range(1, 5):
                row[i].text = "-"
        
        doc.save(filepath)
        return filepath
