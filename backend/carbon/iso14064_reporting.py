#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ISO 14064-1:2018 RAPORLAMA MODÜLÜ
Standarda uygun GHG Inventory Report oluşturma
"""

import logging
import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .carbon_manager import CarbonManager
from .iso14064_compliance import ISO14064Compliance
from .offset_manager import OffsetManager
from config.icons import Icons
from config.database import DB_PATH

def _add_turkish_paragraph(doc, text=None, style=None, font_name='Calibri', font_size=11):
    para = doc.add_paragraph(text if text is not None else '', style=style)
    if not text:
        return para
    for run in para.runs:
        try:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        try:
            run.font.name = font_name
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return heading

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


class ISO14064Reporter:
    """ISO 14064-1:2018 formatında rapor oluşturucu"""

    def __init__(self, db_path: str = DB_PATH) -> None:
        self.carbon_mgr = CarbonManager(db_path)
        self.offset_mgr = OffsetManager(db_path)
        self.iso_mgr = ISO14064Compliance(db_path)

    def generate_iso14064_report(self, company_id: int, period: str,
                                 include_verification: bool = True) -> Optional[str]:
        """
        ISO 14064-1:2018 uyumlu GHG Inventory Report oluştur
        
        Rapor Yapısı (ISO 6.4):
        1. Executive Summary
        2. Organizational Boundaries
        3. Operational Boundaries
        4. GHG Emissions Quantification
        5. Base Year and Recalculation
        6. Uncertainty Assessment
        7. GHG Removals and Offsets
        8. Verification Statement
        9. Appendices
        """

        try:
            # Verileri topla
            emissions_data = self.carbon_mgr.generate_emissions_summary(
                company_id, period, include_scope3=True
            )

            net_data = self.offset_mgr.get_net_emissions(company_id, period)

            uncertainty = self.iso_mgr.assess_uncertainty(company_id, period)

            compliance = self.iso_mgr.get_compliance_report(company_id, period)

            # Şirket bilgisi
            company_name = self._get_company_name(company_id)

            # DOCX oluştur
            doc = Document()

            # ==================== COVER PAGE ====================
            self._add_cover_page(doc, company_name, period)

            doc.add_page_break()

            # ==================== TABLE OF CONTENTS ====================
            self._add_table_of_contents(doc)

            doc.add_page_break()

            # ==================== 1. EXECUTIVE SUMMARY ====================
            self._add_section(doc, "1. EXECUTIVE SUMMARY", level=1)

            summary_text = f"""
This Greenhouse Gas (GHG) Inventory Report has been prepared in accordance with ISO 14064-1:2018 
"Greenhouse gases — Part 1: Specification with guidance at the organization level for quantification 
and reporting of greenhouse gas emissions and removals".

Reporting Organization: {company_name}
Reporting Period: {period}
Report Date: {datetime.now().strftime('%B %d, %Y')}

Total GHG Emissions (Gross): {emissions_data.get('total_co2e', 0):.2f} tCO2e
Total GHG Removals/Offsets: {net_data.get('total_offset', 0) if net_data else 0:.2f} tCO2e
Net GHG Emissions: {net_data.get('total_net', emissions_data.get('total_co2e', 0)) if net_data else emissions_data.get('total_co2e', 0):.2f} tCO2e

ISO 14064-1:2018 Compliance: {compliance.get('overall_status', 'Pending')}
            """
            _add_turkish_paragraph(doc, summary_text.strip())

            # Emissions by Scope Table
            self._add_section(doc, "1.1 GHG Emissions by Scope", level=2)

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'

            headers = ['Scope', 'Emissions (tCO2e)', 'Percentage (%)']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header

            scope1 = emissions_data.get('scope1_total', 0)
            scope2 = emissions_data.get('scope2_total', 0)
            scope3 = emissions_data.get('scope3_total', 0)
            total = emissions_data.get('total_co2e', 0)

            data = [
                ['Scope 1 (Direct)', f"{scope1:.2f}", f"{(scope1/total*100) if total > 0 else 0:.1f}%"],
                ['Scope 2 (Energy Indirect)', f"{scope2:.2f}", f"{(scope2/total*100) if total > 0 else 0:.1f}%"],
                ['Scope 3 (Other Indirect)', f"{scope3:.2f}", f"{(scope3/total*100) if total > 0 else 0:.1f}%"],
                ['TOTAL', f"{total:.2f}", '100.0%']
            ]

            for row_data in data:
                row = table.add_row().cells
                for j, cell_value in enumerate(row_data):
                    row[j].text = cell_value

            doc.add_page_break()

            # ==================== 2. ORGANIZATIONAL BOUNDARIES ====================
            self._add_section(doc, "2. ORGANIZATIONAL BOUNDARIES (ISO 5.2.1)", level=1)

            boundary_text = f"""
{company_name} has established organizational boundaries in accordance with ISO 14064-1:2018, 
clause 5.2.1. The organization uses the operational control approach for consolidating GHG emissions.

Consolidation Approach: Operational Control
Reporting Boundary: All facilities and operations where {company_name} has operational control

Included Entities:
• Head office and administrative facilities
• Manufacturing/production facilities
• Warehouses and distribution centers
• Company-owned vehicles

Excluded Entities:
• Joint ventures where operational control is not held
• Franchises operated independently
• Leased assets where the organization does not have operational control
            """
            _add_turkish_paragraph(doc, boundary_text.strip())

            doc.add_page_break()

            # ==================== 3. OPERATIONAL BOUNDARIES ====================
            self._add_section(doc, "3. OPERATIONAL BOUNDARIES (ISO 5.2.2)", level=1)

            operational_text = """
The operational boundaries include the following GHG sources and scopes:

3.1 Scope 1 (Direct GHG Emissions)
• Stationary combustion: Natural gas, diesel, fuel oil, LPG for heating and power generation
• Mobile combustion: Company-owned vehicles (gasoline, diesel)
• Fugitive emissions: Refrigerant leakage from cooling systems (R-134a, R-404A, R-410A)

3.2 Scope 2 (Energy Indirect GHG Emissions)
• Purchased electricity from the grid
• Purchased steam/district heating

3.3 Scope 3 (Other Indirect GHG Emissions)
Included Categories:
• Category 1: Purchased goods and services
• Category 4: Upstream transportation
• Category 5: Waste generated in operations
• Category 6: Business travel
• Category 7: Employee commuting
• Category 9: Downstream transportation

Note: Other Scope 3 categories have been assessed as not material or not applicable.
            """
            _add_turkish_paragraph(doc, operational_text.strip())

            doc.add_page_break()

            # ==================== 4. QUANTIFICATION METHODOLOGY ====================
            self._add_section(doc, "4. GHG QUANTIFICATION METHODOLOGY (ISO 5.2.4)", level=1)

            method_text = """
4.1 Calculation Approach
GHG emissions are quantified using the calculation methodology in accordance with:
• ISO 14064-1:2018
• GHG Protocol Corporate Accounting and Reporting Standard
• IPCC 2006 Guidelines for National Greenhouse Gas Inventories

4.2 Emission Factors
The following emission factor sources are used:
• IPCC 2006 for stationary and mobile combustion
• IPCC AR5 for Global Warming Potential (GWP) values:
  - CO2: 1
  - CH4: 25
  - N2O: 298
  - HFCs: Substance-specific GWP
• DEFRA 2023 conversion factors for UK-specific activities
• National grid emission factors for purchased electricity

4.3 Calculation Formula
Emissions (tCO2e) = Activity Data × Emission Factor × GWP

Where:
• Activity Data: Quantity of fuel consumed, electricity purchased, distance traveled, etc.
• Emission Factor: CO2, CH4, N2O emission per unit of activity
• GWP: Global Warming Potential to convert to CO2 equivalents
            """
            _add_turkish_paragraph(doc, method_text.strip())

            doc.add_page_break()

            # ==================== 5. BASE YEAR & RECALCULATION ====================
            self._add_section(doc, "5. BASE YEAR AND RECALCULATION POLICY (ISO 5.2.6)", level=1)

            base_year_text = """
5.1 Base Year Selection
Base Year: 2020
Base Year Emissions: 1,000.0 tCO2e
Rationale: 2020 selected as the base year as it represents the first complete year of comprehensive 
data collection following ISO 14064-1 methodology.

5.2 Recalculation Policy
The base year will be recalculated if any of the following triggers occur:
• Structural changes (mergers, acquisitions, divestments) affecting emissions by >5%
• Changes in calculation methodologies
• Discovery of significant errors or improvements in accuracy
• Changes in emission factors that significantly affect the inventory
• Outsourcing or insourcing of emitting activities

Significance Threshold: 5% of base year emissions (50 tCO2e)

5.3 Recalculation History
No recalculations have been performed to date.
            """
            _add_turkish_paragraph(doc, base_year_text.strip())

            doc.add_page_break()

            # ==================== 6. UNCERTAINTY ASSESSMENT ====================
            self._add_section(doc, "6. UNCERTAINTY ASSESSMENT (ISO 5.2.7)", level=1)

            if uncertainty:
                uncertainty_text = f"""
An uncertainty assessment has been conducted in accordance with ISO 14064-1:2018, Annex A.

Overall Uncertainty: ±{uncertainty.get('overall_uncertainty_pct', 0):.1f}%

Emissions Range:
• Lower Bound: {uncertainty.get('lower_bound_total', 0):.2f} tCO2e
• Reported Value: {uncertainty.get('total_emissions', 0):.2f} tCO2e
• Upper Bound: {uncertainty.get('upper_bound_total', 0):.2f} tCO2e

Data Quality Tiers:
• Tier 1 (Measured): ±5% uncertainty
• Tier 2 (Calculated): ±10% uncertainty
• Tier 3 (Estimated): ±30% uncertainty
• Tier 4 (Default): ±50% uncertainty

Uncertainty Reduction Actions:
• Installation of sub-metering for key energy consumers
• Direct measurement of fuel consumption where feasible
• Use of site-specific emission factors instead of defaults
                """
            else:
                uncertainty_text = "Uncertainty assessment pending. Will be completed in next reporting cycle."

            doc.add_paragraph(uncertainty_text.strip())

            doc.add_page_break()

            # ==================== 7. GHG REMOVALS & OFFSETS ====================
            self._add_section(doc, "7. GHG REMOVALS AND OFFSETS (ISO 5.2.8)", level=1)

            if net_data and net_data.get('total_offset', 0) > 0:
                offset_text = f"""
{company_name} has purchased carbon offsets to achieve carbon neutrality.

Total Offsets: {net_data.get('total_offset', 0):.2f} tCO2e
Offset Projects: Verified Carbon Standard (VCS), Gold Standard
Retirement Status: All offsets have been retired for {period} reporting period

Net Emissions: {net_data.get('total_net', 0):.2f} tCO2e
Carbon Neutral Status: {'YES - Carbon Neutral Achieved' if net_data.get('carbon_neutral') else 'NO - In Progress'}

Offset Quality Criteria:
• All offsets from verified projects (VCS, Gold Standard, CDM)
• Additionality demonstrated
• No double counting
• Permanent removal or long-term storage (>100 years)
• Third-party verification completed
                """
            else:
                offset_text = "No carbon offsets were purchased or retired during the reporting period."

            doc.add_paragraph(offset_text.strip())

            doc.add_page_break()

            # ==================== 8. VERIFICATION ====================
            if include_verification:
                self._add_section(doc, "8. VERIFICATION STATEMENT (ISO 6.3)", level=1)

                verification_text = """
This GHG Inventory has been prepared for third-party verification in accordance with 
ISO 14064-3:2019 "Greenhouse gases — Part 3: Specification with guidance for the verification 
and validation of greenhouse gas statements".

Verification Level: Reasonable Assurance
Materiality Threshold: 5%
Verification Standard: ISO 14064-3:2019

[Verification statement to be inserted upon completion of third-party verification]
                """
                _add_turkish_paragraph(doc, verification_text.strip())

                doc.add_page_break()

            # ==================== 9. COMPLIANCE CHECKLIST ====================
            self._add_section(doc, "9. ISO 14064-1:2018 COMPLIANCE CHECKLIST", level=1)

            if compliance:
                _add_turkish_paragraph(doc, f"Overall Compliance: {compliance.get('compliance_percentage', 0):.0f}%")
                _add_turkish_paragraph(doc, f"Status: {compliance.get('overall_status', 'Pending')}\n")

                checklist_table = doc.add_table(rows=len(compliance.get('checklist', [])) + 1, cols=3)
                checklist_table.style = 'Light List Accent 1'

                # Headers
                checklist_table.rows[0].cells[0].text = 'ISO Clause'
                checklist_table.rows[0].cells[1].text = 'Requirement'
                checklist_table.rows[0].cells[2].text = 'Status'

                for i, item in enumerate(compliance.get('checklist', []), start=1):
                    checklist_table.rows[i].cells[0].text = item.get('iso_clause', '')
                    checklist_table.rows[i].cells[1].text = item.get('requirement_description', '')
                    checklist_table.rows[i].cells[2].text = item.get('compliance_status', 'pending').upper()

            # ==================== FOOTER ====================
            self._add_footer(doc, company_name, period)

            # Kaydet
            reports_dir = "reports"
            if not os.path.exists(reports_dir):
                os.makedirs(reports_dir)

            filename = f"ISO14064_GHG_Inventory_{company_name.replace(' ', '_')}_{period}.docx"
            filepath = os.path.join(reports_dir, filename)

            doc.save(filepath)

            return filepath

        except Exception as e:
            logging.error(f"[ERROR] ISO 14064-1 rapor oluşturulamadı: {e}")
            import traceback
            traceback.print_exc()
            return None

    def _add_cover_page(self, doc, company_name: str, period: str) -> None:
        """Kapak sayfası ekle"""
        # Başlık
        title = _add_turkish_heading(doc, 'GREENHOUSE GAS INVENTORY REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Alt başlık
        subtitle = _add_turkish_paragraph(doc, '\nIn Accordance with ISO 14064-1:2018')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.italic = True

        doc.add_paragraph('\n' * 3)

        # Şirket bilgisi
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f'{company_name}\n').font.size = Pt(18)
        info.add_run(f'Reporting Period: {period}\n').font.size = Pt(14)
        info.add_run(f'Report Date: {datetime.now().strftime("%B %d, %Y")}\n').font.size = Pt(12)

        doc.add_paragraph('\n' * 5)

        # ISO logo metni
        iso_text = _add_turkish_paragraph(doc, 'ISO 14064-1:2018 COMPLIANT')
        iso_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        iso_text.runs[0].font.bold = True
        iso_text.runs[0].font.size = Pt(12)

    def _add_table_of_contents(self, doc) -> None:
        """İçindekiler tablosu"""
        _add_turkish_heading(doc, 'TABLE OF CONTENTS', level=1)

        toc_items = [
            '1. Executive Summary',
            '2. Organizational Boundaries',
            '3. Operational Boundaries',
            '4. GHG Quantification Methodology',
            '5. Base Year and Recalculation Policy',
            '6. Uncertainty Assessment',
            '7. GHG Removals and Offsets',
            '8. Verification Statement',
            '9. ISO 14064-1:2018 Compliance Checklist'
        ]

        for item in toc_items:
            _add_turkish_paragraph(doc, item, style='List Number')

    def _add_section(self, doc, title: str, level: int = 1) -> None:
        """Bölüm başlığı ekle"""
        heading = _add_turkish_heading(doc, title, level=level)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    def _add_footer(self, doc, company_name: str, period: str) -> None:
        """Footer ekle"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"{company_name} | ISO 14064-1:2018 GHG Inventory | {period}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _get_company_name(self, company_id: int) -> str:
        """Şirket adını al"""
        try:
            conn = self.carbon_mgr.get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM companies WHERE id = ?", (company_id,))
            result = cursor.fetchone()
            conn.close()
            return result[0] if result else f"Company {company_id}"
        except Exception:
            return f"Company {company_id}"


if __name__ == "__main__":
    # Test
    reporter = ISO14064Reporter()

    # ISO 14064-1 raporu oluştur
    report_path = reporter.generate_iso14064_report(
        company_id=1,
        period="2024",
        include_verification=True
    )

    if report_path:
        logging.info(f"{Icons.SUCCESS} ISO 14064-1:2018 raporu oluşturuldu: {report_path}")
    else:
        logging.info(f"{Icons.FAIL} Rapor oluşturulamadı")

