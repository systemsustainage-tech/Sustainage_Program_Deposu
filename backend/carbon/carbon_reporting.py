#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import logging
"""
KARBON RAPORLAMA MODÜLÜ
GHG Protocol uyumlu karbon raporları oluşturma
"""

import os
from datetime import datetime

from .carbon_manager import CarbonManager
from config.database import DB_PATH
try:
    from docx.shared import Pt
except ImportError:
    pass

def _add_turkish_paragraph(doc, text=None, style=None, font_name='Calibri', font_size=12):
    """Türkçe karakterleri destekleyen paragraf ekle"""
    para = doc.add_paragraph(text if text is not None else '', style=style)
    if not text:
        return para
    for run in para.runs:
        try:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return para

def _add_turkish_heading(doc, text, level=1, font_name='Calibri'):
    """Türkçe karakterleri destekleyen başlık ekle"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        try:
            run.font.name = font_name
            from docx.oxml.ns import qn
            r = run._element
            r.rPr.rFonts.set(qn('w:ascii'), font_name)
            r.rPr.rFonts.set(qn('w:hAnsi'), font_name)
            r.rPr.rFonts.set(qn('w:cs'), font_name)
        except Exception:
            pass
    return heading

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.info("[UYARI] python-docx yuklu degil. Rapor olusturma devre disi.")

class CarbonReporting:
    """Karbon raporlama sınıfı"""

    def __init__(self, db_path: str = DB_PATH) -> None:
        self.db_path = db_path
        self.manager = CarbonManager(db_path)

    def generate_ghg_inventory_report(self, company_id: int, period: str,
                                     include_scope3: bool = False) -> str:
        """
        GHG Inventory raporu oluştur (DOCX)
        
        Returns:
            Rapor dosya yolu
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kütüphanesi gerekli")

        # Emisyon özetini al
        summary = self.manager.generate_emissions_summary(
            company_id=company_id,
            period=period,
            include_scope3=include_scope3
        )

        # Hedefleri al
        targets = self.manager.get_carbon_targets(company_id)

        # Trend analizi
        try:
            current_year = int(period)
            self.manager.get_emissions_trend(
                company_id=company_id,
                start_year=current_year - 4,
                end_year=current_year
            )
        except Exception as e:
            logging.error(f'Silent error in carbon_reporting.py: {str(e)}')

        # Dokuman oluştur
        doc = Document()

        # Başlık
        title = doc.add_heading('KARBON EMİSYON ENVANTERİ RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Alt başlık
        subtitle = doc.add_paragraph(f'Raporlama Dönemi: {period}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)

        doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y")}')
        doc.add_paragraph('Raporlama Standardı: GHG Protocol')

        doc.add_page_break()

        # 1. Yönetici Özeti
        doc.add_heading('1. YÖNETİCİ ÖZETİ', 1)
        doc.add_paragraph(
            f"Bu rapor, {period} yılı karbon emisyon envanterini GHG Protocol "
            f"standartlarına uygun olarak sunmaktadır."
        )

        doc.add_paragraph(f"\nToplam Emisyon (Scope 1+2): {summary['scope1_2_total']:.2f} tCO2e")
        doc.add_paragraph(f"  • Scope 1 (Doğrudan): {summary['scope1_total']:.2f} tCO2e")
        doc.add_paragraph(f"  • Scope 2 (Dolaylı): {summary['scope2_total']:.2f} tCO2e")

        if include_scope3 and summary.get('scope3_total'):
            doc.add_paragraph(f"  • Scope 3 (Değer Zinciri): {summary['scope3_total']:.2f} tCO2e")
            doc.add_paragraph(f"\nToplam Emisyon (Scope 1+2+3): {summary['total_co2e']:.2f} tCO2e")

        # 2. Metodoloji
        doc.add_heading('2. METODOLOJİ', 1)
        doc.add_paragraph(
            "Emisyon hesaplamaları GHG Protocol Corporate Standard'a uygun olarak yapılmıştır. "
            "Operasyonel kontrol yaklaşımı kullanılmıştır."
        )

        # 3. Scope 1 Detayı
        doc.add_heading('3. SCOPE 1 - DOĞRUDAN EMİSYONLAR', 1)

        if 'scope1_breakdown' in summary:
            s1_detail = summary['scope1_breakdown']

            doc.add_paragraph(f"Toplam Scope 1: {s1_detail['total_co2e']:.2f} tCO2e\n")

            if 'stationary' in s1_detail:
                doc.add_heading('3.1 Sabit Yakma Kaynakları', 2)
                for detail in s1_detail['stationary'].get('details', []):
                    doc.add_paragraph(
                        f"  • {detail['fuel_type']}: {detail['quantity']:.1f} {detail['unit']} "
                        f"→ {detail['co2e']:.2f} tCO2e",
                        style='List Bullet'
                    )

            if 'mobile' in s1_detail:
                doc.add_heading('3.2 Mobil Yakma (Araç Filosu)', 2)
                doc.add_paragraph(f"Toplam: {s1_detail['mobile']['total_co2e']:.2f} tCO2e")

        # 4. Scope 2 Detayı
        doc.add_heading('4. SCOPE 2 - DOLAYILI EMİSYONLAR', 1)

        if 'scope2_breakdown' in summary:
            s2_detail = summary['scope2_breakdown']
            doc.add_paragraph(f"Toplam Scope 2: {s2_detail['total_co2e']:.2f} tCO2e")
            doc.add_paragraph("Hesaplama Metodu: Lokasyon Bazlı (Location-based)")

        # 5. Hedefler
        if targets:
            doc.add_heading('5. KARBON AZALTMA HEDEFLERİ', 1)
            for target in targets:
                if target['status'] == 'active':
                    doc.add_paragraph(
                        f"Hedef: {target['target_name']}\n"
                        f"  Baz Yıl: {target['baseline_year']} - {target['baseline_co2e']:.1f} tCO2e\n"
                        f"  Hedef Yıl: {target['target_year']} - {target['target_co2e']:.1f} tCO2e\n"
                        f"  Azaltma Hedefi: %{target['target_reduction_pct']:.0f}"
                    )

        # Dosyayı kaydet
        # Firma bazlı rapor klasörü
        base_dir = os.path.dirname(os.path.dirname(self.db_path))
        output_dir = os.path.join(base_dir, 'data', 'companies', str(company_id), 'reports')
        os.makedirs(output_dir, exist_ok=True)

        filename = f"carbon_inventory_{period}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(output_dir, filename)

        doc.save(filepath)
        return filepath

    def generate_tsrs_compliance_report(self, company_id: int, period: str) -> str:
        """
        TSRS (TSRS 2 İklim Standardı) Uyum Raporu Oluştur
        
        TSRS 2 Standardı Gereklilikleri:
        1. Yönetişim (Governance)
        2. Strateji (Strategy)
        3. Risk Yönetimi (Risk Management)
        4. Metrikler ve Hedefler (Metrics and Targets)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kütüphanesi gerekli")

        # Emisyon verilerini al
        summary = self.manager.generate_emissions_summary(company_id, period, include_scope3=True)
        targets = self.manager.get_carbon_targets(company_id)
        
        # Doküman oluştur
        doc = Document()
        
        # Başlık
        title = _add_turkish_heading(doc, 'TSRS 2 İKLİM STANDARDI UYUM RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Rapor Bilgileri
        subtitle = _add_turkish_paragraph(doc, f'Raporlama Dönemi: {period}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y")}')
        
        doc.add_page_break()
        
        # 1. YÖNETİŞİM
        _add_turkish_heading(doc, '1. YÖNETİŞİM', 1)
        _add_turkish_paragraph(doc, 
            "Şirket yönetimi, iklim kaynaklı risk ve fırsatları düzenli olarak değerlendirmektedir. "
            "Yönetim kurulu seviyesinde sürdürülebilirlik komitesi oluşturulmuş ve sorumluluklar tanımlanmıştır."
        )
        
        # 2. STRATEJİ
        _add_turkish_heading(doc, '2. STRATEJİ', 1)
        _add_turkish_paragraph(doc,
            "İklim değişikliğinin şirket stratejisi üzerindeki etkileri kısa, orta ve uzun vadeli olarak analiz edilmektedir. "
            "Düşük karbonlu ekonomiye geçiş senaryoları (1.5°C ve 2°C) dikkate alınmaktadır."
        )
        
        # 3. RİSK YÖNETİMİ
        _add_turkish_heading(doc, '3. RİSK YÖNETİMİ', 1)
        _add_turkish_paragraph(doc,
            "İklim riskleri (fiziksel ve geçiş riskleri) kurumsal risk yönetimi süreçlerine entegre edilmiştir. "
            "Tedarik zinciri riskleri ve karbon fiyatlandırma riskleri izlenmektedir."
        )
        
        # 4. METRİKLER VE HEDEFLER
        _add_turkish_heading(doc, '4. METRİKLER VE HEDEFLER', 1)
        
        # 4.1 Scope 1-2-3 Emisyonları
        _add_turkish_heading(doc, '4.1 Sera Gazı Emisyonları (TSRS Gerekliliği)', 2)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Kapsam'
        hdr_cells[1].text = 'Emisyon (tCO2e)'
        hdr_cells[2].text = 'Açıklama'
        
        # Scope 1
        row = table.add_row().cells
        row[0].text = 'Kapsam 1 (Doğrudan)'
        row[1].text = f"{summary['scope1_total']:.2f}"
        row[2].text = 'Sabit ve mobil yakma kaynaklı'
        
        # Scope 2
        row = table.add_row().cells
        row[0].text = 'Kapsam 2 (Enerji Dolaylı)'
        row[1].text = f"{summary['scope2_total']:.2f}"
        row[2].text = 'Satın alınan elektrik'
        
        # Scope 3
        row = table.add_row().cells
        row[0].text = 'Kapsam 3 (Diğer Dolaylı)'
        row[1].text = f"{summary.get('scope3_total', 0):.2f}"
        row[2].text = 'Değer zinciri emisyonları'
        
        # Toplam
        row = table.add_row().cells
        row[0].text = 'TOPLAM'
        row[1].text = f"{summary['total_co2e']:.2f}"
        row[2].text = '-'

        # 4.2 Hedefler
        if targets:
            _add_turkish_heading(doc, '4.2 İklim Hedefleri', 2)
            for target in targets:
                if target['status'] == 'active':
                    _add_turkish_paragraph(doc, 
                        f"• {target['target_name']}: {target['baseline_year']} baz yılına göre "
                        f"{target['target_year']} yılına kadar %{target['target_reduction_pct']:.0f} azaltım."
                    )
        else:
            _add_turkish_paragraph(doc, "\nHenüz tanımlanmış bir azaltım hedefi bulunmamaktadır.")

        # Kaydet
        base_dir = os.path.dirname(os.path.dirname(self.db_path))
        output_dir = os.path.join(base_dir, 'data', 'companies', str(company_id), 'reports')
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"tsrs_uyum_raporu_{period}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(output_dir, filename)
        
        doc.save(filepath)
        return filepath

    def generate_excel_report(self, company_id: int, period: str) -> str:
        """Excel dashboard raporu oluştur"""
        try:
            import pandas as pd
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Font, PatternFill
        except ImportError:
            raise ImportError("pandas ve openpyxl kütüphaneleri gerekli")

        # Emisyon verilerini al
        emissions = self.manager.get_emissions(company_id, period)

        # DataFrame oluştur
        df = pd.DataFrame(emissions)

        # Excel dosyası oluştur
        # Firma bazlı rapor klasörü
        base_dir = os.path.dirname(os.path.dirname(self.db_path))
        output_dir = os.path.join(base_dir, 'data', 'companies', str(company_id), 'reports')
        os.makedirs(output_dir, exist_ok=True)

        filename = f"carbon_dashboard_{period}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        filepath = os.path.join(output_dir, filename)

        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Emisyon Verileri', index=False)

            # Özet sayfa
            summary = self.manager.generate_emissions_summary(company_id, period, True)
            summary_df = pd.DataFrame([{
                'Scope': 'Scope 1',
                'Emisyon (tCO2e)': summary['scope1_total']
            }, {
                'Scope': 'Scope 2',
                'Emisyon (tCO2e)': summary['scope2_total']
            }, {
                'Scope': 'Scope 3',
                'Emisyon (tCO2e)': summary.get('scope3_total', 0)
            }, {
                'Scope': 'TOPLAM',
                'Emisyon (tCO2e)': summary['total_co2e']
            }])

            summary_df.to_excel(writer, sheet_name='Özet', index=False)

        return filepath
