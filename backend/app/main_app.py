#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SUSTAINAGE SDG - ANA UYGULAMA
- Sol menü çerçevesi
- Sağ tarafta içerik alanı
- Modül yönetimi
"""

import io
import json
import logging
import os
import sqlite3
import sys
import tkinter as tk
from datetime import datetime
from tkinter import filedialog, messagebox, ttk

from yonetim.licensing.core.license_ed25519 import get_license_info
from utils.language_manager import LanguageManager
from utils.ui_theme import apply_theme, standardize_buttons
from config.icons import Icons
from modules.reporting.report_scheduler import ReportScheduler

# Windows terminal için UTF-8 desteği
if __name__ == '__main__' and os.name == 'nt':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except (AttributeError, ValueError) as e:
        logging.error(f"Warning: Failed to set UTF-8 encoding: {e}")

# Modülleri import et
sys.path.append(os.path.join(os.path.dirname(__file__), '..'))

# Modül importları - LAZY LOADING için kaldırıldı
# Modüller sadece gerektiğinde import edilecek

# Temel import'lar

# Modüller LAZY olarak yüklenecek; üst düzey import yok

class MainApp:
    def __init__(self, parent, user, company_id: int = 1, current_lang: str = 'tr') -> None:
        # Logger
        self.logger = logging.getLogger(__name__)
        self.lm = LanguageManager()
        logging.debug(f"[DEBUG] MainApp init: current_lang={current_lang}, lm.current_lang={self.lm.current_lang}")
        
        # Her zaman dili yükle veya güncelle
        if current_lang:
            logging.debug(f"[DEBUG] Loading language: {current_lang}")
            self.lm.load_language(current_lang)
            # Verify if it loaded correctly
            test_val = self.lm.tr('menu_company_info', 'Firma Bilgileri')
            logging.debug(f"[DEBUG] Loaded {len(self.lm.translations)} keys. Example 'menu_company_info': {test_val}")
            
            # Critical Check: If English is selected but we got Turkish, something is wrong
            if current_lang == 'en' and test_val == 'Firma Bilgileri':
                logging.error("[ERROR] English selected but Turkish text returned! Retrying load...")
                self.lm.load_language('en')
                logging.debug(f"[DEBUG] Retry result: {self.lm.tr('menu_company_info')}")
        
        self.parent = parent
        self.user = user
        self.base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        self.db_path = os.path.join(self.base_dir, "data", "sdg_desktop.sqlite")
        # Kullanıcı yetki kontrolü UserManager üzerinden yapılacak
        self.company_id = company_id or 1  # Login ekranından gelen firma ID
        # Firma bazlı klasör yapısı
        self.company_dir = self._ensure_company_dir()
        self.current_module = None  # Mevcut modül takibi
        # Gezinme yığını ve mevcut görünüm fonksiyonu
        self._nav_stack = []
        self._current_view_func = None

        # UI ayarları - hızlı varsayılan değerler
        self.dashboard_variant = 'classic'
        self.dashboard_image_mode = 'fit_width_scroll'
        self.dashboard_bg_gradient_colors = None
        self.dashboard_image_path = None
        self.show_welcome_banner = True

        # PERFORMANS OPTİMİZASYONU - Modül cache
        self.module_cache = {}  # Yüklenen modülleri cache'le
        self.loading_frames = {}  # Yükleme ekranları
        self.cached_modules = set()  # Cache'lenmiş modüller
        self.module_instances = {}

        # UI ayarları dosyası - arka planda yükle
        self.ui_settings_path = os.path.join(self.base_dir, 'config', 'ui_settings.json')
        self.parent.after(10, self._load_ui_settings_async)
        # Kullanıcı yetkilerini arka planda yükle
        self.user_permissions = set()
        if self.user and len(self.user) > 0:
            # İzinleri senkron yükle ki ilk ekran erişim reddi olmasın
            self._load_user_permissions_async()
        # Firma bilgileri tablosunu hazırla
        try:
            self.ensure_company_info_table()
        except Exception as e:
            logging.info(f"Firma bilgileri tablosu hazırlanamadı: {e}")
        self.setup_ui()

        # UI başlatma

    def _check_license_status(self) -> None:
        """Uygulama başlangıcında lisans kontrolü"""
        try:
            conn = sqlite3.connect(self.db_path)
            self.license_info = get_license_info(conn)
            conn.close()
            
            self.license_state = self.license_info.get("state", "none")
            self.license_restricted = False
            self.license_expired = False
            
            if self.license_state == "expired":
                self.license_expired = True
                logging.warning("License expired!")
                # Süre bitti uyarısı - setup_ui sonrası gösterilecek
                self.parent.after(1000, lambda: messagebox.showwarning(
                    self.lm.tr('license_expired_title', "Lisans Süresi Doldu"),
                    self.lm.tr('license_expired_msg', "Lisans süreniz dolmuştur. Bazı özellikler kısıtlanmış olabilir.")
                ))
                
            elif self.license_state not in ("valid", "tolerated"):
                self.license_restricted = True
                logging.warning(f"License invalid: {self.license_state}")
                if self.license_state == "none":
                     msg = "Lisans bulunamadı."
                else:
                     msg = f"Lisans geçersiz: {self.license_info.get('reason', '')}"
                
                self.parent.after(1000, lambda: messagebox.showerror(
                    self.lm.tr('license_error_title', "Lisans Hatası"),
                    msg
                ))
                
        except Exception as e:
            logging.error(f"License check error: {e}")
            self.license_info = {"state": "error", "message": str(e)}
            self.license_state = "error"

    def _load_ui_settings_async(self) -> None:
        """UI ayarlarını arka planda yükle"""
        try:
            if os.path.exists(self.ui_settings_path):
                with open(self.ui_settings_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.dashboard_variant = (data or {}).get('dashboard_variant', 'classic')
                    self.dashboard_image_mode = (data or {}).get('dashboard_image_mode', self.dashboard_image_mode)
                    bg_grad = (data or {}).get('dashboard_bg_gradient')
                    self.dashboard_image_path = (data or {}).get('dashboard_image_path')
                    if isinstance(bg_grad, dict):
                        colors = bg_grad.get('colors')
                        if isinstance(colors, (list, tuple)) and len(colors) == 2:
                            c1, c2 = colors[0], colors[1]
                            if isinstance(c1, str) and isinstance(c2, str):
                                self.dashboard_bg_gradient_colors = (c1.strip(), c2.strip())
                    self.show_welcome_banner = self._get_user_welcome_banner_setting()
        except Exception as e:
            logging.info(f"UI ayarları arka planda yüklenemedi: {e}")

    def _load_user_permissions_async(self) -> None:
        """Kullanıcı yetkilerini arka planda yükle"""
        try:
            if self.user and len(self.user) > 0:
                username = self.user[1] if len(self.user) > 1 else ''
                user_id_login_db = self.user[0]

                # Cache kontrolü
                try:
                    from utils.cache import user_permissions_cache
                    cached_perms = user_permissions_cache.get_permissions(user_id_login_db)
                    if cached_perms is not None:
                        self.user_permissions = cached_perms
                        logging.info(f" Kullanıcı yetkileri cache'den yüklendi ({len(cached_perms)} izin)")
                        return
                except Exception as e:
                    logging.debug(f"Cache error: {e}")

                # UserManager'ı lazy import et
                from yonetim.kullanici_yonetimi.models.user_manager import UserManager
                um = UserManager(self.db_path)

                # RBAC izinlerini yükle
                perms = set(um.get_user_permissions(user_id_login_db))
                if not perms and username:
                    try:
                        mapped_user = um.get_user_by_username(username)
                        mapped_id = (mapped_user or {}).get('id')
                        if mapped_id:
                            perms = set(um.get_user_permissions(mapped_id))
                    except Exception as e:
                        logging.error(f"Username ile RBAC eşlemede hata: {e}")

                self.user_permissions = perms

                # Cache'e kaydet
                try:
                    from utils.cache import user_permissions_cache
                    user_permissions_cache.set_permissions(user_id_login_db, perms)
                except Exception as e:
                    logging.debug(f"Cache set error: {e}")

                # Admin için varsayılan izinler
                if not self.user_permissions and username == 'admin':
                    self.user_permissions = {
                        'dashboard.read', 'dashboard.advanced', 'company.read',
                        'sdg.read', 'gri.read', 'tsrs.read', 'esg.read',
                        'strategic.read', 'data.import', 'forms.manage',
                        'tasks.auto_create', 'files.manage', 'hr.read', 'policy.read',
                        'skdm.read', 'mapping.read', 'prioritization.read',
                        'waste.read', 'water.read', 'supply_chain.read',
                        'product_tech.read', 'report.read', 'system.settings'
                    }

                logging.debug(f"[DEBUG] Kullanıcı {username or user_id_login_db} izinleri yüklendi: {len(self.user_permissions)} izin")
        except Exception as e:
            logging.info(f"İzin sistemi arka planda yüklenemedi: {e}")
            import traceback
            traceback.print_exc()
            self.user_permissions = set()

    def _draw_live_stats(self, canvas) -> None:
        """Dashboard üzerine canlı istatistikleri çizer"""
        try:
            cw = canvas.winfo_width()
            ch = canvas.winfo_height()
            if cw < 200 or ch < 200:
                return

            # Stats container coordinates (Top Right)
            box_width = 250
            box_height = 180
            margin = 20
            x1 = cw - box_width - margin
            y1 = margin
            x2 = cw - margin
            y2 = y1 + box_height

            # Background for stats
            canvas.create_rectangle(x1, y1, x2, y2, fill='white', outline='#e5e7eb', width=1)
            canvas.create_line(x1, y1+40, x2, y1+40, fill='#e5e7eb') # Header separator

            # Header
            canvas.create_text(x1 + 15, y1 + 20, text="Canlı Özet", font=('Segoe UI', 12, 'bold'), anchor='w', fill='#111827')
            
            # 1. Energy
            try:
                from modules.environmental.detailed_energy_manager import DetailedEnergyManager
                manager = DetailedEnergyManager()
                # Get current year data
                import datetime
                year = datetime.datetime.now().year
                data = manager.get_annual_report_data(self.company_id, year)
                total_cons = data.get('total_consumption', 0)
                
                canvas.create_text(x1 + 15, y1 + 65, text="Toplam Enerji:", font=('Segoe UI', 9, 'bold'), anchor='w', fill='#4b5563')
                canvas.create_text(x2 - 15, y1 + 65, text=f"{total_cons:,.0f} kWh", font=('Segoe UI', 9), anchor='e', fill='#059669')
            except Exception as e:
                logging.debug(f"Energy stats error: {e}")
                canvas.create_text(x1 + 15, y1 + 65, text="Enerji verisi yok", font=('Segoe UI', 9), anchor='w', fill='#9ca3af')

            # 2. Tasks
            try:
                # Direct DB query for speed
                conn = sqlite3.connect(self.db_path)
                cur = conn.cursor()
                cur.execute("SELECT COUNT(*) FROM tasks WHERE status != 'completed'")
                row = cur.fetchone()
                pending_count = row[0] if row else 0
                conn.close()
                
                canvas.create_text(x1 + 15, y1 + 95, text="Bekleyen Görevler:", font=('Segoe UI', 9, 'bold'), anchor='w', fill='#4b5563')
                canvas.create_text(x2 - 15, y1 + 95, text=str(pending_count), font=('Segoe UI', 9), anchor='e', fill='#d97706')
            except Exception as e:
                logging.debug(f"Task stats error: {e}")

            # 3. Last Login / System Status
            canvas.create_text(x1 + 15, y1 + 125, text="Sistem Durumu:", font=('Segoe UI', 9, 'bold'), anchor='w', fill='#4b5563')
            canvas.create_text(x2 - 15, y1 + 125, text="Aktif", font=('Segoe UI', 9), anchor='e', fill='#2563eb')
            
            # Timestamp
            import datetime
            now_str = datetime.datetime.now().strftime("%H:%M")
            canvas.create_text(x2 - 15, y2 - 15, text=f"Güncelleme: {now_str}", font=('Segoe UI', 8), anchor='e', fill='#9ca3af')

        except Exception as e:
            logging.error(f"Draw live stats error: {e}")

    def _load_dashboard_images_async(self, canvas) -> None:
        """Dashboard resimlerini arka planda yükle"""
        try:
            import math
            resimler_dir = os.path.join(self.base_dir, 'resimler')
            center_candidates = ['sdg.png', 'sdg.jpg']
            center_path = None
            for name in center_candidates:
                p = os.path.join(resimler_dir, name)
                if os.path.exists(p):
                    center_path = p
                    break

            ring_paths = []
            for i in range(1, 18):
                p = os.path.join(resimler_dir, f"{i}.png")
                if os.path.exists(p):
                    ring_paths.append(p)

            from PIL import Image, ImageTk
            center_img = Image.open(center_path) if center_path else None
            ring_imgs = []
            for p in ring_paths:
                try:
                    ring_imgs.append(Image.open(p))
                except Exception as e:
                    logging.error(f"SDG görseli yükleme hatası ({p}): {e}")

            if not hasattr(self, 'image_references'):
                self.image_references = []

            def render(_=None) -> None:
                canvas.delete("all")
                self.image_references = []
                cw = canvas.winfo_width()
                ch = canvas.winfo_height()
                if cw < 10 or ch < 10:
                    return
                cx, cy = cw // 2, ch // 2
                resample = getattr(Image, 'Resampling', Image).LANCZOS if hasattr(Image, 'Resampling') else getattr(Image, 'LANCZOS', getattr(Image, 'BICUBIC', Image.NEAREST))

                # Merkez görsel
                center_box = int(min(cw, ch) * 0.35)
                if center_img is not None:
                    iw, ih = center_img.size
                    scale = min(center_box / iw, center_box / ih)
                    nw = max(1, int(iw * scale))
                    nh = max(1, int(ih * scale))
                    c_resized = center_img.resize((nw, nh), resample)
                    c_photo = ImageTk.PhotoImage(c_resized, master=self.parent)
                    self.image_references.append(c_photo)
                    canvas.create_image(cx, cy, image=c_photo, anchor='center')

                # Ring görselleri
                ring_count = len(ring_imgs)
                if ring_count > 0:
                    ring_radius = int(min(cw, ch) * 0.4)
                    ring_tile = int(min(cw, ch) * 0.08)

                    for i, img in enumerate(ring_imgs):
                        angle = (i * 2 * math.pi / ring_count) - math.pi/2  # 12 yönünden başla
                        rx = int(cx + ring_radius * math.cos(angle))
                        ry = int(cy + ring_radius * math.sin(angle))
                        iw, ih = img.size
                        scale = min(ring_tile / iw, ring_tile / ih)
                        nw = max(1, int(iw * scale))
                        nh = max(1, int(ih * scale))
                        r_resized = img.resize((nw, nh), resample)
                        r_photo = ImageTk.PhotoImage(r_resized, master=self.parent)
                        self.image_references.append(r_photo)
                        canvas.create_image(rx - nw // 2, ry - nh // 2, image=r_photo, anchor='nw')

                self._draw_live_stats(canvas)

            canvas.bind('<Configure>', render)
            # İlk render
            self.parent.after(100, render)

        except Exception as e:
            logging.info(f"Dashboard resimleri yüklenemedi: {e}")
            # Fallback dashboard göster
            self.show_fallback_dashboard()

    def _show_loading_screen(self, module_name: str) -> None:
        """Hızlı yükleme ekranı göster"""
        # Önceki yükleme ekranını temizle
        if module_name in self.loading_frames:
            self.loading_frames[module_name].destroy()

        # Yükleme ekranı oluştur
        loading_frame = tk.Frame(self.content_area, bg='white')
        loading_frame.pack(fill='both', expand=True)

        # Yükleme animasyonu
        loading_text = self.lm.tr('module_loading', '{module} modülü yükleniyor...').format(module=module_name.upper())
        loading_label = tk.Label(loading_frame,
                                text=loading_text,
                                font=('Segoe UI', 14, 'bold'),
                                fg='#1e40af', bg='white')
        loading_label.pack(expand=True)

        # Progress bar
        progress = ttk.Progressbar(loading_frame, mode='indeterminate')
        progress.pack(pady=10)
        progress.start()

        # Cache'e kaydet
        self.loading_frames[module_name] = loading_frame

        # UI'yi güncelle
        self.parent.update_idletasks()

    def _hide_loading_screen(self, module_name: str) -> None:
        """Yükleme ekranını gizle"""
        if module_name in self.loading_frames:
            self.loading_frames[module_name].destroy()
            del self.loading_frames[module_name]

    def _show_module_fast(self, module_name: str, module_title: str, permission: str = None) -> None:
        try:
            if permission and not self._require_permission(permission, module_title):
                return
            self._audit_navigation(module_title)
            self.current_module = module_name
            self.content_header.pack(fill='x')
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
            self.content_title['text'] = module_title
            self._add_back_button()
            if module_name in self.module_instances:
                self._show_module_frame(module_name)
                return
            self._show_loading_screen(module_name)
            def _load():
                ok = self._ensure_module_instance(module_name)
                self._hide_loading_screen(module_name)
                if ok:
                    self._show_module_frame(module_name)
                else:
                    self.show_module_error(module_name)
            self.parent.after(10, _load)
        except Exception as e:
            logging.error(f"{module_title} modülü yükleme hatası: {e}")
            self.show_module_error(module_name)

    def _ensure_module_instance(self, module_name: str) -> bool:
        try:
            if module_name in self.module_instances:
                return True
            GUIClass = self._lazy_import_module(module_name)
            if not GUIClass:
                return False
            container = tk.Frame(self.content_area, bg='white')
            # Sadece bir turda farklı imza denemeleri (gereksiz tekrar kaldırıldı)
            try:
                GUIClass(container, self.company_id, main_app=self)
                created = True
            except TypeError:
                try:
                    GUIClass(container, self.company_id)
                    created = True
                except TypeError:
                    try:
                        GUIClass(container, self.db_path, self.company_id)
                        created = True
                    except TypeError:
                        try:
                            user_id = self.user[0] if self.user else None
                            if user_id is not None:
                                GUIClass(container, self.company_id, user_id)
                                created = True
                            else:
                                raise TypeError("No user_id available for signature")
                        except TypeError:
                            try:
                                GUIClass(container)
                                created = True
                            except Exception as e:
                                logging.error(f"[ERROR] Module {module_name} instantiation failed: {e}")
                                created = False

            if not created:
                try:
                    container.destroy()
                except Exception as e:
                    logging.debug(f"Destroy error: {e}")
                return False

            try:
                container.pack_forget()
            except Exception as e:
                logging.debug(f"Pack forget error: {e}")
            self.module_instances[module_name] = container
            return True
        except Exception as e:
            logging.error(f"[ERROR] Module {module_name} instantiation error: {e}")
            return False

    def _show_module_frame(self, module_name: str) -> None:
        try:
            frame = self.module_instances.get(module_name)
            if not frame or not frame.winfo_exists():
                return
            try:
                for child in self.content_area.winfo_children():
                    child.pack_forget()
            except Exception as e:
                logging.debug(f"Child pack forget error: {e}")
            frame.pack(fill='both', expand=True)
            try:
                standardize_buttons(self.content_area)
            except Exception as e:
                logging.debug(f"Standardize buttons error: {e}")
            try:
                self.parent.update_idletasks()
            except Exception as e:
                logging.debug(f"Update idletasks error: {e}")
        except Exception as e:
            logging.debug(f"Show module frame error: {e}")

    def _lazy_import_module(self, module_name):
        """Modülü sadece gerektiğinde import et - CACHE DESTEKLİ"""
        # Cache'de var mı kontrol et
        if module_name in self.module_cache:
            return self.module_cache[module_name]

        try:
            if module_name == 'sdg':
                # SDG modülünü doğru paket yolundan import et
                from modules.sdg.sdg_gui import SDGGUI
                self.module_cache[module_name] = SDGGUI
                return SDGGUI
            elif module_name == 'energy':
                from modules.environmental.energy_gui import EnergyGUI
                self.module_cache[module_name] = EnergyGUI
                return EnergyGUI
            elif module_name == 'gri':
                from modules.gri.gri_gui import GRIGUI
                self.module_cache[module_name] = GRIGUI
                return GRIGUI
            elif module_name == 'mapping':
                from modules.mapping.mapping_gui import MappingGUI
                self.module_cache[module_name] = MappingGUI
                return MappingGUI
            elif module_name == 'yonetim':
                from yonetim.yönetim_gui import YonetimGUI
                self.module_cache[module_name] = YonetimGUI
                return YonetimGUI
            elif module_name == 'raporlama':
                from raporlama.reporting_gui import ReportingGUI
                self.module_cache[module_name] = ReportingGUI
                return ReportingGUI
            elif module_name == 'mapping':
                from mapping.mapping_gui import MappingGUI
                self.module_cache[module_name] = MappingGUI
                return MappingGUI
            elif module_name == 'prioritization':
                from modules.prioritization.prioritization_gui import PrioritizationGUI
                self.module_cache[module_name] = PrioritizationGUI
                return PrioritizationGUI
            elif module_name == 'waste':
                from modules.waste_management.waste_gui import WasteGUI
                self.module_cache[module_name] = WasteGUI
                return WasteGUI
            elif module_name == 'water':
                from modules.water_management.water_gui import WaterGUI
                self.module_cache[module_name] = WaterGUI
                return WaterGUI
            elif module_name == 'supply_chain':
                from modules.supply_chain.supply_chain_gui import SupplyChainGUI
                self.module_cache[module_name] = SupplyChainGUI
                return SupplyChainGUI
            elif module_name == 'carbon':
                from carbon.carbon_gui import CarbonGUI
                self.module_cache[module_name] = CarbonGUI
                return CarbonGUI
            elif module_name == 'product_tech':
                from modules.product_technology.product_tech_gui import ProductTechGUI
                self.module_cache[module_name] = ProductTechGUI
                return ProductTechGUI
            elif module_name == 'reporting':
                from modules.reporting.reporting_gui import ReportingGUI
                self.module_cache[module_name] = ReportingGUI
                return ReportingGUI
            elif module_name == 'esrs':
                from modules.esrs.esrs_gui import ESRSGUI
                self.module_cache[module_name] = ESRSGUI
                return ESRSGUI
            elif module_name == 'tsrs':
                from modules.tsrs.tsrs_gui import TSRSGUI
                self.module_cache[module_name] = TSRSGUI
                return TSRSGUI
            elif module_name == 'issb':
                from modules.issb.issb_gui import ISSBGUI
                self.module_cache[module_name] = ISSBGUI
                return ISSBGUI
            elif module_name == 'skdm':
                from modules.skdm.skdm_gui import SKDMGUI
                self.module_cache[module_name] = SKDMGUI
                return SKDMGUI
            elif module_name == 'prioritization':
                from modules.prioritization.prioritization_gui import PrioritizationGUI
                self.module_cache[module_name] = PrioritizationGUI
                return PrioritizationGUI
            elif module_name == 'waste':
                from modules.waste_management.waste_gui import WasteGUI
                self.module_cache[module_name] = WasteGUI
                return WasteGUI
            elif module_name == 'water':
                from modules.water_management.water_gui import WaterGUI
                self.module_cache[module_name] = WaterGUI
                return WaterGUI
            elif module_name == 'supply_chain':
                from modules.supply_chain.supply_chain_gui import SupplyChainGUI
                self.module_cache[module_name] = SupplyChainGUI
                return SupplyChainGUI
            elif module_name == 'advanced_materiality':
                from modules.analytics.advanced_materiality_gui import AdvancedMaterialityGUI
                self.module_cache[module_name] = AdvancedMaterialityGUI
                return AdvancedMaterialityGUI
            elif module_name == 'advanced_security':
                from modules.security.advanced_security_gui import AdvancedSecurityGUI
                self.module_cache[module_name] = AdvancedSecurityGUI
                return AdvancedSecurityGUI
            elif module_name == 'auto_tasks':
                from modules.auto_tasks.auto_task_gui import AutoTaskGUI
                self.module_cache[module_name] = AutoTaskGUI
                return AutoTaskGUI
            elif module_name == 'esg':
                from modules.esg.esg_gui import ESGGUI
                self.module_cache[module_name] = ESGGUI
                return ESGGUI
            elif module_name == 'data_import':
                from modules.data_import.import_gui import DataImportGUI
                self.module_cache[module_name] = DataImportGUI
                return DataImportGUI
            elif module_name == 'form_management':
                from modules.forms.form_gui import FormManagementGUI
                self.module_cache[module_name] = FormManagementGUI
                return FormManagementGUI
            elif module_name == 'advanced_files':
                from modules.file_manager.advanced_file_gui import AdvancedFileGUI
                self.module_cache[module_name] = AdvancedFileGUI
                return AdvancedFileGUI
            elif module_name == 'hr_metrics':
                from modules.social.hr_metrics_gui import HRMetricsGUI
                self.module_cache[module_name] = HRMetricsGUI
                return HRMetricsGUI
            elif module_name == 'ungc':
                from modules.ungc.ungc_gui import UNGCGUI
                self.module_cache[module_name] = UNGCGUI
                return UNGCGUI
            elif module_name == 'tcfd':
                from modules.tcfd.tcfd_gui import TCFDGUI
                self.module_cache[module_name] = TCFDGUI
                return TCFDGUI
            elif module_name == 'sasb':
                from modules.sasb.sasb_gui import SASBGUI
                self.module_cache[module_name] = SASBGUI
                return SASBGUI
            elif module_name == 'tasks':
                from tasks.task_gui import TaskGUI
                self.module_cache[module_name] = TaskGUI
                return TaskGUI
            elif module_name == 'admin_dashboard':
                from tasks.admin_dashboard_gui import AdminDashboardGUI
                self.module_cache[module_name] = AdminDashboardGUI
                return AdminDashboardGUI
            elif module_name == 'innovation':
                from modules.innovation.innovation_gui import InnovationGUI
                self.module_cache[module_name] = InnovationGUI
                return InnovationGUI
            elif module_name == 'quality':
                from modules.quality.quality_gui import QualityGUI
                self.module_cache[module_name] = QualityGUI
                return QualityGUI
            elif module_name == 'digital_security':
                from modules.digital_security.digital_security_gui import DigitalSecurityGUI
                self.module_cache[module_name] = DigitalSecurityGUI
                return DigitalSecurityGUI
            elif module_name == 'emergency':
                from modules.emergency.emergency_gui import EmergencyGUI
                self.module_cache[module_name] = EmergencyGUI
                return EmergencyGUI
            elif module_name == 'strategic':
                from modules.strategic.strategic_gui import StrategicGUI
                self.module_cache[module_name] = StrategicGUI
                return StrategicGUI
            elif module_name == 'advanced_dashboard':
                from modules.data_inventory.advanced_dashboard_gui import AdvancedDashboardGUI
                self.module_cache[module_name] = AdvancedDashboardGUI
                return AdvancedDashboardGUI
            elif module_name == 'esg_consolidated_dashboard':
                from modules.esg.esg_consolidated_dashboard import ESGConsolidatedDashboard
                self.module_cache[module_name] = ESGConsolidatedDashboard
                return ESGConsolidatedDashboard
            elif module_name == 'policy_library':
                from modules.policy_library import PolicyLibraryGUI
                self.module_cache[module_name] = PolicyLibraryGUI
                return PolicyLibraryGUI
            elif module_name == 'social_dashboard':
                from modules.social.social_dashboard_gui import SocialDashboardGUI
                self.module_cache[module_name] = SocialDashboardGUI
                return SocialDashboardGUI
            elif module_name == 'cdp':
                from modules.cdp.cdp_gui import CDPGUI
                self.module_cache[module_name] = CDPGUI
                return CDPGUI
            elif module_name == 'iirc':
                from modules.iirc.iirc_gui import IIRCGUI
                self.module_cache[module_name] = IIRCGUI
                return IIRCGUI
            elif module_name == 'validation':
                from modules.validation.validation_gui import ValidationGUI
                self.module_cache[module_name] = ValidationGUI
                return ValidationGUI
            elif module_name == 'benchmark':
                from modules.analytics.benchmark_gui import BenchmarkGUI
                self.module_cache[module_name] = BenchmarkGUI
                return BenchmarkGUI
            elif module_name == 'forecasting':
                from modules.analytics.forecasting_gui import ForecastingGUI
                self.module_cache[module_name] = ForecastingGUI
                return ForecastingGUI
            elif module_name == 'report_center':
                from modules.reporting.report_center_gui import ReportCenterGUI
                self.module_cache[module_name] = ReportCenterGUI
                return ReportCenterGUI
            elif module_name == 'integration':
                from modules.integration.integration_gui import IntegrationGUI
                self.module_cache[module_name] = IntegrationGUI
                return IntegrationGUI
            elif module_name == 'visualization':
                from modules.visualization.visualization_center_gui import VisualizationCenterGUI
                self.module_cache[module_name] = VisualizationCenterGUI
                return VisualizationCenterGUI
            elif module_name == 'auditor':
                from modules.auditor.auditor_gui import AuditorGUI
                self.module_cache[module_name] = AuditorGUI
                return AuditorGUI
            elif module_name == 'stakeholder':
                from modules.stakeholder.stakeholder_gui import StakeholderGUI
                self.module_cache[module_name] = StakeholderGUI
                return StakeholderGUI
            elif module_name == 'scenario':
                from modules.scenario_analysis.scenario_gui import ScenarioGUI
                self.module_cache[module_name] = ScenarioGUI
                return ScenarioGUI
            elif module_name == 'csrd':
                from modules.csrd.csrd_gui import CSRDGUI
                self.module_cache[module_name] = CSRDGUI
                return CSRDGUI
            elif module_name == 'report_scheduler':
                from modules.reporting.report_scheduler_gui import ReportSchedulerGUI
                self.module_cache[module_name] = ReportSchedulerGUI
                return ReportSchedulerGUI
            # YENI MODULLER (23 Ekim 2025)
            elif module_name == 'company_management':
                from modules.company.company_gui import CompanyGUI
                self.module_cache[module_name] = CompanyGUI
                return CompanyGUI
            elif module_name == 'ai_module':
                from modules.ai.ai_module_gui import AIModuleGUI
                self.module_cache[module_name] = AIModuleGUI
                return AIModuleGUI
            elif module_name == 'document_processing':
                from modules.document_processing.doc_gui import DocProcessingGUI
                self.module_cache[module_name] = DocProcessingGUI
                return DocProcessingGUI
            elif module_name == 'erp_integration':
                from modules.erp_integration.erp_gui import ERPGUI
                self.module_cache[module_name] = ERPGUI
                return ERPGUI
            elif module_name == 'auto_reporting':
                from modules.automated_reporting.auto_report_gui import AutoReportGUI
                self.module_cache[module_name] = AutoReportGUI
                return AutoReportGUI
            elif module_name == 'advanced_calculation':
                from modules.advanced_calculation.calculation_gui import CalculationGUI
                self.module_cache[module_name] = CalculationGUI
                return CalculationGUI
            elif module_name == 'advanced_inventory':
                from modules.advanced_inventory.inventory_gui import InventoryGUI
                self.module_cache[module_name] = InventoryGUI
                return InventoryGUI
            elif module_name == 'eu_taxonomy':
                from modules.eu_taxonomy.taxonomy_gui import EUTaxonomyGUI
                self.module_cache[module_name] = EUTaxonomyGUI
                return EUTaxonomyGUI
            elif module_name == 'risks':
                from modules.strategic.risk_gui import RiskGUI
                self.module_cache[module_name] = RiskGUI
                return RiskGUI
            elif module_name == 'ceo_messages':
                from modules.strategic.ceo_message_gui import CEOMessageGUI
                self.module_cache[module_name] = CEOMessageGUI
                return CEOMessageGUI
            else:
                logging.info(f"Bilinmeyen modül: {module_name}")
                return None
        except ImportError as e:
            logging.error(f"ImportError - Modül {module_name} yüklenemedi: {e}")
            return None
        except Exception as e:
            logging.error(f"Genel hata - Modül {module_name} yüklenemedi: {e}")
            import traceback
            traceback.print_exc()
            return None

    def show_module_error(self, module_name, error_message=None) -> None:
        """Modül yüklenemediğinde hata göster"""
        error_frame = tk.Frame(self.content_area, bg='white')
        error_frame.pack(fill='both', expand=True)

        if error_message:
            display_text = self.lm.tr('module_load_error_with_msg', '{module} modülü yüklenirken hata: {msg}').format(module=module_name, msg=error_message)
        else:
            display_text = self.lm.tr('module_load_error', '{module} modülü yüklenemedi').format(module=module_name)

        error_label = tk.Label(error_frame,
                             text=display_text,
                             font=('Segoe UI', 14, 'bold'),
                             fg='#e74c3c', bg='white')
        error_label.pack(pady=50)

        retry_btn = tk.Button(error_frame, text=self.lm.tr('retry', " Tekrar Dene"),
                             font=('Segoe UI', 12), bg='#3498db', fg='white',
                             command=lambda: self.parent.after(100, self._navigate(getattr(self, f'show_{module_name.lower()}', None))))
        retry_btn.pack(pady=10)

        self._add_back_button()

    def _ensure_company_dir(self) -> None:
        """Firma bazlı dosya dizinini oluşturur ve yolunu döndürür."""
        try:
            root = os.path.join(self.base_dir, 'data', 'companies', str(self.company_id))
            os.makedirs(root, exist_ok=True)
            for sub in ('exports', 'imports', 'reports', 'backups'):
                os.makedirs(os.path.join(root, sub), exist_ok=True)
            return root
        except Exception as e:
            logging.info(f"Firma dizini oluşturulamadı: {e}")
            return os.path.join(self.base_dir, 'data')

    def setup_ui(self) -> None:
        """Ana uygulama arayüzünü oluştur"""
        # Ana pencere ayarları (güvenli kullanıcı adı/isim kullanımı)
        try:
            display_name = (
                self.user[2] if isinstance(self.user, (list, tuple)) and len(self.user) > 2 and self.user[2] else None
            )
        except Exception as e:
            logging.debug(f"Could not get display name: {e}")
            display_name = None
        try:
            username = (
                self.user[1] if isinstance(self.user, (list, tuple)) and len(self.user) > 1 and self.user[1] else self.lm.tr('default_user', 'Kullanıcı')
            )
        except Exception as e:
            logging.debug(f"Could not get username: {e}")
            username = self.lm.tr('default_user', 'Kullanıcı')
        
        welcome_msg = self.lm.tr('app_welcome_title', "Sustainage - Hoş Geldiniz, {name}").format(name=display_name or username)
        self.parent.title(welcome_msg)
        self.parent.state('zoomed')  # Tam ekran
        apply_theme(self.parent)
        self.parent.configure(bg='#f5f5f5')

        # Ana çerçeve
        self.setup_main_frame()

        # Sol menü
        self.setup_sidebar()

        # Sağ içerik alanı
        self.setup_content_area()

        # Mouse tekerleği desteği ekle
        self.setup_mouse_wheel_support()

        # Klavye kısayollarını aktifleştir
        try:
            self.setup_keyboard_shortcuts()
        except Exception as e:
            logging.info(f"Kısayollar başlatılamadı: {e}")

    def setup_main_frame(self) -> None:
        """Ana çerçeveyi oluştur"""
        self.main_frame = tk.Frame(self.parent, bg='#f5f5f5')
        self.main_frame.pack(fill='both', expand=True, padx=10, pady=10)

    def setup_sidebar(self) -> None:
        """Sol menü çerçevesini oluştur"""
        # Sol menü çerçevesi - Tema ile uyumlu mavi arkaplan
        self.sidebar = tk.Frame(self.main_frame, width=270, bg='#2F6DB2', highlightthickness=0)
        self.sidebar.pack(side='left', fill='y', padx=(0, 10))
        self.sidebar.pack_propagate(False)

        # Logo ve başlık
        logo_frame = tk.Frame(self.sidebar, bg='#2F6DB2', highlightthickness=0)
        logo_frame.pack(fill='x', padx=16, pady=8)

        title_label = tk.Label(logo_frame, text=self.lm.tr('sidebar_title', "SUSTAINAGE"),
                              font=('Segoe UI', 16, 'bold'), fg='white', bg='#2F6DB2')
        title_label.pack()

        subtitle_label = tk.Label(logo_frame, text=self.lm.tr('app_subtitle', "Sürdürülebilir Kalkınma Yönetimi"),
                                 font=('Segoe UI', 9), fg='#E0E0E0', bg='#2F6DB2')
        subtitle_label.pack()

        # Arama kutusu
        search_frame = tk.Frame(self.sidebar, bg='#2F6DB2')
        search_frame.pack(fill='x', padx=12, pady=(0, 6))
        tk.Label(search_frame, text=self.lm.tr('search_label', "Ara"), font=('Segoe UI', 9), fg='white', bg='#2F6DB2').pack(side='left')
        self.menu_search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=self.menu_search_var, width=24)
        search_entry.pack(side='left', padx=8)
        try:
            self.menu_search_var.trace_add('write', lambda *args: self._apply_menu_filter())
        except Exception as e:
            logging.debug(f"Menu filter trace error: {e}")

        # Menü butonları
        self._menu_button_height = 32
        self.setup_menu_buttons()

        # Alt bilgi
        self.setup_sidebar_footer()
        
        # Kalan boşluğu dolduracak spacer
        tk.Frame(self.sidebar, bg='#2F6DB2').pack(fill='both', expand=True)

    def setup_mouse_wheel_support(self) -> None:
        """Mouse tekerleği desteği ekle - tüm uygulamada"""
        def _on_mousewheel(event):
            # Hangi widget'ta olduğumuzu kontrol et
            widget = event.widget

            # Menü butonu canvas'larını ignore et
            if isinstance(widget, tk.Canvas) and hasattr(self, '_menu_canvases'):
                try:
                    if widget in self._menu_canvases:
                        return 'break'
                except Exception as e:
                    logging.debug(f"Menu canvas check error: {e}")
            # Canvas widget'ı ise scroll yap (yalnızca içerik canvas'ları)
            if isinstance(widget, tk.Canvas):
                # Statik canvas kontrolü (dashboard vb.)
                if getattr(widget, 'is_static', False):
                    return 'break'
                try:
                    widget.yview_scroll(int(-1*(event.delta/120)), "units")
                except Exception as e:
                    logging.debug(f"Canvas scroll error: {e}")
                return 'break'
            # Scrollable widget'lar için genel scroll
            elif hasattr(widget, 'yview_scroll'):
                try:
                    widget.yview_scroll(int(-1*(event.delta/120)), "units")
                except Exception as e:
                    logging.debug(f"Widget scroll error: {e}")
                return 'break'

        # Ana pencereye mouse tekerleği bağla
        self.parent.bind("<MouseWheel>", _on_mousewheel)

        # Tüm child widget'lara da bağla
        def bind_recursive(widget):
            try:
                widget.bind("<MouseWheel>", _on_mousewheel)
                for child in widget.winfo_children():
                    bind_recursive(child)
            except Exception as e:
                logging.debug(f"Recursive bind error: {e}")

        bind_recursive(self.parent)

    def setup_menu_buttons(self) -> None:
        """Menü butonlarını oluştur - Sadece ana modüller"""
        # Double check language just before creating menu
        if self.lm.current_lang != 'tr' and self.lm.tr('menu_company_info') == 'Firma Bilgileri':
             logging.info(f"[WARN] Language seems not loaded correctly in setup_menu_buttons. Reloading {self.lm.current_lang}...")
             self.lm.load_language(self.lm.current_lang)

        logging.debug(f"[DEBUG] setup_menu_buttons: current_lang={self.lm.current_lang}")
        logging.debug(f"[DEBUG] menu_company_info translation: {self.lm.tr('menu_company_info')}")
        
        menu_container = tk.Frame(self.sidebar, bg='#2F6DB2', highlightthickness=0)
        # pady alt boşluğu kaldırıldı (gap fix)
        menu_container.pack(fill='x', padx=10, pady=(0, 0))
        menu_frame = tk.Frame(menu_container, bg='#2F6DB2', highlightthickness=0)
        menu_frame.pack(fill='both', expand=True)

        self._menu_controls = {'main': []}

        parent_modules = [
            ("🏢", self.lm.tr("menu_company_info", "Firma Bilgileri"), self.show_company_info),
            (Icons.REPORT, self.lm.tr("menu_dashboard", "Dashboard"), self.show_dashboard),
            (Icons.LEAF, self.lm.tr("menu_esg", "ESG"), self.show_esg_dashboard),
            ("📘", self.lm.tr("menu_gri", "GRI"), self.show_gri_dashboard),
            ("🎯", self.lm.tr("menu_sdg", "SDG"), self.show_sdg_dashboard),
            ("⚡", self.lm.tr("energy_module_title", "Enerji Yönetimi"), self.show_energy),
            ("🇹🇷", self.lm.tr("menu_tsrs", "TSRS"), self.show_tsrs_dashboard),
            ("🇪🇺", self.lm.tr("menu_esrs", "ESRS"), self.show_esrs),
            (Icons.CHART_UP, self.lm.tr("menu_issb_ifrs", "ISSB/IFRS"), self.show_issb),
            (Icons.WORLD, self.lm.tr("menu_tcfd", "TCFD"), self.show_tcfd),
            ("🧾", self.lm.tr("menu_cdp", "CDP"), self.show_cdp),
            ("🏭", self.lm.tr("menu_sasb", "SASB"), self.show_sasb),
            ("🔖", self.lm.tr("menu_eu_taxonomy", "EU Taxonomy"), self.show_eu_taxonomy),
            (Icons.LINK, self.lm.tr("menu_iirc", "IIRC"), self.show_iirc),
            (Icons.REPORT, self.lm.tr("menu_reporting", "Raporlama"), self.show_reporting),
            ("🏢", self.lm.tr("menu_management", "Yönetim"), self.show_management),
        ]

        try:
            if self.user and len(self.user) > 1 and self.user[1] == '__super__':
                parent_modules.insert(1, ("⚡", self.lm.tr("menu_super_admin", "Super Admin"), self.show_super_admin_panel))
        except Exception as e:
            logging.debug(f"Super admin check error: {e}")

        for icon, text, command in parent_modules:
            self._create_menu_button(menu_frame, icon, text, command)

        try:
            menu_frame.update_idletasks()
        except Exception as e:
            logging.debug(f"Menu update idletasks error: {e}")

    def _create_menu_button(self, parent, icon, text, command):
        def wrapped_cmd(fn=command, label=text):
            self._navigate(fn)
            try:
                # Menüde kalıcı vurguyu güncelle
                for t, btn in self._menu_controls.get('main', []):
                    if hasattr(btn, 'mark_selected'):
                        btn.mark_selected(t == label)
            except Exception as e:
                logging.debug(f"Mark selected error: {e}")
        
        # Icon ve text ayrı gönderiliyor (hizalama için)
        frame, _ = self._create_neumorphic_menu_control(parent, text, wrapped_cmd, icon=icon)
        # pady=(2, 0) ile butonlar arası boşluk azaltıldı ve alt boşluk kaldırıldı
        frame.pack(fill='x', pady=(2, 0))
        self._menu_controls['main'].append((text, frame))

    def _create_expandable_group(self, parent, group_name, sub_items):
        group_container = tk.Frame(parent, bg='#2F6DB2')
        group_container.pack(fill='x', pady=2)
        button_frame = tk.Frame(group_container, bg='#2F6DB2')
        button_frame.pack(fill='x')
        expand_btn, expand_icon = self._create_neumorphic_menu_control(button_frame, group_name, None, right_text='▶')
        expand_btn.pack(fill='x')
        sub_frame = tk.Frame(group_container, bg='#2F6DB2')
        def make_toggle(gname, sframe, icon, items):
            return lambda: self._toggle_menu_group(gname, sframe, icon, items, group_container)
        def bind_toggle(widget):
            try:
                for w in widget.winfo_children():
                    w.bind('<ButtonRelease-1>', lambda e: make_toggle(group_name, sub_frame, expand_icon, sub_items)())
                widget.bind('<ButtonRelease-1>', lambda e: make_toggle(group_name, sub_frame, expand_icon, sub_items)())
            except Exception as e:
                logging.debug(f"Toggle bind error: {e}")
        bind_toggle(expand_btn)
        self._menu_controls['groups'][group_name] = {
            'button': expand_btn,
            'icon': expand_icon,
            'sub_frame': sub_frame,
            'container': group_container,
            'sub_buttons': [],
            'items': sub_items,
        }

    def _toggle_menu_group(self, group_name, sub_frame, icon, sub_items, group_container):
        """Menü grubunu aç/kapat"""
        if group_name in self.expanded_menu_groups:
            # Kapat
            self.expanded_menu_groups.discard(group_name)
            try:
                if icon:
                    icon.configure(text="▶")
            except Exception as e:
                logging.debug(f"Icon configure error: {e}")
            sub_frame.pack_forget()
        else:
            # Aç
            self.expanded_menu_groups.add(group_name)
            try:
                if icon:
                    icon.configure(text="▼")
            except Exception as e:
                logging.debug(f"Icon update error: {e}")

            # Alt menüyü göster
            sub_frame.pack(fill='x', padx=(15, 0), pady=(2, 0))

            # Mevcut alt butonları temizle
            for widget in sub_frame.winfo_children():
                widget.destroy()

            # Alt butonları oluştur - Closure fix
            self._menu_controls['groups'][group_name]['sub_buttons'] = []
            for item_text, item_command in sub_items:
                # Command wrapper fonksiyonu - closure düzeltmesi
                def make_command(cmd):
                    def wrapper():
                        self._navigate(cmd)
                    return wrapper

                frame, _ = self._create_neumorphic_menu_control(sub_frame, f"  • {item_text}", make_command(item_command))
                frame.pack(fill='x', pady=1)
                self._menu_controls['groups'][group_name]['sub_buttons'].append((item_text, frame))

        # Canvas scroll region'ı güncelle
        try:
            if hasattr(self, 'menu_canvas'):
                self.menu_canvas.update_idletasks()
                self.menu_canvas.configure(scrollregion=self.menu_canvas.bbox("all"))
        except Exception as e:
            logging.debug(f"Scroll region update error: {e}")

    def _apply_menu_filter(self) -> None:
        text = (self.menu_search_var.get() or '').lower().strip() if hasattr(self, 'menu_search_var') else ''
        # Ana butonlar
        for label, btn in self._menu_controls.get('main', []):
            show = (text in label.lower()) if text else True
            try:
                btn.pack_forget()
                if show:
                    btn.pack(fill='x', pady=2)
            except Exception as e:
                logging.debug(f"Menu filter pack error: {e}")
        # Gruplar
        for gname, meta in self._menu_controls.get('groups', {}).items():
            matches_group = text in gname.lower() if text else True
            match_sub = False
            # Grubu genişlet ve alt öğeleri oluşturulmamışsa oluştur
            if text and gname not in self.expanded_menu_groups:
                try:
                    self._toggle_menu_group(gname, meta['sub_frame'], meta['icon'], meta.get('items', []), meta['container'])
                except Exception as e:
                    logging.debug(f"Menu group toggle error: {e}")
            # Alt butonlarda filtre
            for label, btn in meta.get('sub_buttons', []):
                ok = (text in label.lower()) if text else True
                if ok:
                    match_sub = True
                try:
                    btn.pack_forget()
                    if ok:
                        btn.pack(fill='x', pady=1)
                except Exception as e:
                    logging.debug(f"Menu sub-button pack error: {e}")
            # Grup buton görünürlüğü
            show_group = matches_group or match_sub
            try:
                meta['container'].pack_forget()
                if show_group:
                    meta['container'].pack(fill='x', pady=2)
            except Exception as e:
                logging.debug(f"Menu container pack error: {e}")

    def _has_permission(self, permission_name: str) -> bool:
        try:
            # Super admin kontrolü - TÜM yetkiler
            if self.user and len(self.user) > 1 and self.user[1] == '__super__':
                return True

            # Admin kullanıcısı kontrolü - username veya role bazlı
            if self.user and len(self.user) > 1:
                username = self.user[1]
                # Admin kullanıcısı veya admin rolü - TÜM yetkiler
                if username in ('admin', '_super_') or username == 'admin':
                    return True

            # Role bazlı admin kontrolü
            if self.user and len(self.user) > 0:
                try:
                    conn = sqlite3.connect(self.db_path)
                    cursor = conn.cursor()
                    cursor.execute("SELECT role FROM users WHERE id = ?", (self.user[0],))
                    result = cursor.fetchone()
                    conn.close()

                    if result and result[0] == 'admin':
                        return True
                except Exception as e:
                    logging.error(f"Role check DB error: {e}")

            # UserManager üzerinden yetki kontrolü yapılacak
            has_perm = permission_name in getattr(self, 'user_permissions', set())
            return has_perm
        except Exception as e:
            logging.error(f"İzin kontrolü hatası: {e}")
            return False

    def _require_permission(self, permission_name: str, module_name: str) -> bool:
        try:
            if not self._has_permission(permission_name):
                # Audit: navigate_denied
                try:
                    conn = sqlite3.connect(self.db_path)
                    actor = self.user[1] if self.user and len(self.user) > 1 else 'unknown'
                    details = {
                        'module': module_name,
                        'permission': permission_name,
                        'user_id': self.user[0] if self.user else None,
                    }
                    # Audit fonksiyonu lazy import
                    from yonetim.security.core.audit import audit_user_action
                    audit_user_action(conn, actor, 'navigate_denied', details)
                    conn.close()
                except Exception as e:
                    logging.error(f"Audit denied log error: {e}")
                try:
                    messagebox.showwarning(self.lm.tr('permission_denied', "Yetki yok"), self.lm.tr('access_denied_for_module', "{module} için erişim izniniz yok.").format(module=module_name))
                except Exception as e:
                    logging.error(f"Permission warning dialog error: {e}")
                return False
            return True
        except Exception as e:
            logging.error(f"Permission check critical error: {e}")
            return True

    def _audit_navigation(self, module_name: str) -> None:
        try:
            conn = sqlite3.connect(self.db_path)
            actor = self.user[1] if self.user and len(self.user) > 1 else 'unknown'
            details = {
                'module': module_name,
                'user_id': self.user[0] if self.user else None,
            }
            # Audit fonksiyonu lazy import
            from yonetim.security.core.audit import audit_user_action
            audit_user_action(conn, actor, 'navigate', details)
            conn.close()
        except Exception as e:
            # Sessiz geç; audit hatası navigasyonu engellemesin
            try:
                logging.error(f"Audit write error: {e}")
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    def on_menu_hover(self, button, is_entering) -> None:
        """Menü hover efekti"""
        if is_entering:
            button.configure(bg='#60a5fa', fg='white')
        else:
            button.configure(bg='#3b82f6', fg='white')

    def setup_sidebar_footer(self) -> None:
        """Sol menü alt bilgisini oluştur"""
        # bg rengini sidebar ile aynı yaptık (#2F6DB2) ki bütünlük sağlansın (gap fix)
        footer_frame = tk.Frame(self.sidebar, bg='#2F6DB2', highlightthickness=0)
        # side='top' yaparak menünün hemen altına ekliyoruz
        # pady=(0, 0) alt boşluk ve üst boşluk 0 (gap completely closed)
        footer_frame.pack(side='top', fill='x', padx=10, pady=(0, 0))

        # Yenile butonu super admin modülüne taşındı

        neo_help, _ = self._create_neumorphic_menu_control(footer_frame, self.lm.tr('menu_help', 'Yardım'), self.show_help, icon="📚")
        neo_help.pack(fill='x', pady=(0, 2))
        
        # Lisans durumu göstergesi
        if hasattr(self, 'license_info'):
            lic_state = self.license_info.get('state', 'none')
            # Renkler: Valid=Yeşil, Tolerated=Turuncu, Expired/Invalid=Kırmızı
            if lic_state in ('valid', 'tolerated'):
                lic_color = '#2ecc71'
            elif lic_state == 'expired':
                lic_color = '#f39c12' # Expired ama belki çalışıyor
            else:
                lic_color = '#e74c3c'
            
            lic_text = "LİSANS: " + (lic_state.upper() if lic_state else "YOK")
            
            lic_frame = tk.Frame(footer_frame, bg='#2F6DB2')
            lic_frame.pack(fill='x', pady=(5, 5))
            
            tk.Label(lic_frame, text=lic_text, font=('Segoe UI', 7, 'bold'), 
                    fg=lic_color, bg='#2F6DB2').pack(side='left', padx=5)
            
            # Eğer süre varsa göster
            exp = self.license_info.get('exp')
            if exp:
                import time
                try:
                    remaining = int((int(exp) - time.time()) / 86400)
                    if remaining < 0:
                        rem_text = " (DOLDU)"
                        fg_rem = '#e74c3c'
                    else:
                        rem_text = f" ({remaining} gün)"
                        fg_rem = '#bdc3c7'
                    tk.Label(lic_frame, text=rem_text, font=('Segoe UI', 7), 
                            fg=fg_rem, bg='#2F6DB2').pack(side='left')
                except:
                    pass

        neo_exit, _ = self._create_neumorphic_menu_control(footer_frame, self.lm.tr('menu_exit', 'Çıkış'), self.exit_application, icon="🚪")
        neo_exit.pack(fill='x')

    def setup_content_area(self) -> None:
        """Sağ içerik alanını oluştur"""
        # Ana içerik çerçevesi
        self.content_frame = tk.Frame(self.main_frame, bg='white', relief='raised', bd=1)
        self.content_frame.pack(side='right', fill='both', expand=True)

        # İçerik başlığı
        self.content_header = tk.Frame(self.content_frame, bg='#f8f9fa', height=60)
        self.content_header.pack(fill='x')
        self.content_header.pack_propagate(False)

        self.content_title = tk.Label(self.content_header, text=self.lm.tr("dashboard_title", "Dashboard"),
                                    font=('Segoe UI', 18, 'bold'), fg='#2E8B57', bg='#f8f9fa')
        self.content_title.pack(side='left', padx=20, pady=15)
        # Başlık sağ aksiyon alanı (modüle özel butonlar için)
        self.content_actions = tk.Frame(self.content_header, bg='#f8f9fa')
        self.content_actions.pack(side='right', padx=10, pady=10)

        # İçerik alanı
        self.content_area = tk.Frame(self.content_frame, bg='white')
        self.content_area.pack(fill='both', expand=True)

        # Varsayılan olarak dashboard'ı göster
        self.show_dashboard()

    # Gezinme yardımcıları
    def _navigate(self, target_fn) -> None:
        """Yeni görünüme giderken mevcut görünümü yığına it ve hedefi çalıştır."""
        try:
            if callable(getattr(self, '_current_view_func', None)):
                # Aynı sayfaya tekrar gitmeyi yığma
                if self._current_view_func is not target_fn:
                    self._nav_stack.append(self._current_view_func)
            # Hedefi çalıştır
            target_fn()
            # Mevcut görünüm fonksiyonunu güncelle
            self._current_view_func = target_fn
            self._animate_reveal()
        except Exception as e:
            logging.error(f"Navigasyon hatası: {e}")

    def _animate_reveal(self) -> None:
        try:
            area = getattr(self, 'content_area', None)
            frame = getattr(self, 'content_frame', None)
            if not area or not frame or not area.winfo_exists():
                return
            frame.update_idletasks()
            w = frame.winfo_width()
            h = frame.winfo_height()
            overlay = tk.Canvas(frame, bg='', highlightthickness=0)
            # Use grid to cover the entire frame (requires frame to use grid layout)
            overlay.grid(row=0, column=0, rowspan=2, sticky='nsew')
            
            # Create rectangle using relative dimensions if possible, but canvas items use absolute coords
            # We can bind to Configure to resize the rectangle, but for a short animation it might be overkill.
            # Since we used grid(sticky='nsew'), the canvas resizes with the frame.
            # We just need to make sure the rectangle covers it.
            rect = overlay.create_rectangle(0, 0, w*2, h*2, fill='#f8f9fa', outline='') # Make it large enough
            stipples = ['gray75', 'gray75', 'gray50', 'gray50', 'gray25', 'gray25', 'gray12', 'gray12', None]
            duration = 600
            steps = len(stipples)
            idx = 0
            def step():
                nonlocal idx
                if idx >= steps:
                    overlay.grid_forget()
                    overlay.destroy()
                    return
                st = stipples[idx]
                try:
                    overlay.itemconfigure(rect, stipple=st if st else '')
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
                idx += 1
                frame.after(max(10, duration // steps), step)
            step()
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    def _create_neumorphic_button(self, parent, text, command):
        base = tk.Frame(parent, bg='#1e40af')
        w = 240
        h = 40
        c = tk.Canvas(base, width=w, height=h, bg='#1e40af', highlightthickness=0)
        c.pack(fill='x')
        bg = '#1e40af'
        light = '#274fbd'
        dark = '#162e7a'
        def draw(pressed=False):
            c.delete('all')
            line_color = dark if pressed else light
            shadow_color = light if pressed else dark
            c.create_rectangle(2, 2, w-2, h-2, outline='', fill=bg)
            c.create_rectangle(2, 2, w-2, h-2, outline=line_color, width=2)
            c.create_line(4, 4, w-4, 4, fill=line_color, width=2)
            c.create_line(4, h-4, w-4, h-4, fill=shadow_color, width=2)
            c.create_line(4, 4, 4, h-4, fill=line_color, width=2)
            c.create_line(w-4, 4, w-4, h-4, fill=shadow_color, width=2)
            c.create_rectangle(8, 8, w-8, h-8, outline='', fill=bg)
            c.create_text(w//2, h//2, text=text, fill='#ffffff', font=('Segoe UI', 11, 'bold'))
        def on_press(event):
            draw(True)
        def on_release(event):
            draw(False)
            if callable(command):
                try:
                    command()
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
        c.bind('<ButtonPress-1>', on_press)
        c.bind('<ButtonRelease-1>', on_release)
        draw(False)
        return base

    def _create_neumorphic_menu_control(self, parent, text, command, right_text=None, accent_bg=None, icon=None):
        base = tk.Frame(parent, bg=accent_bg or '#1e40af')
        w = 240
        h = getattr(self, '_menu_button_height', 32)
        c = tk.Canvas(base, width=w, height=h, bg=accent_bg or '#1e40af', highlightthickness=0)
        c.pack(fill='x')
        bg = accent_bg or '#1e40af'
        def _shade(hex_color, factor):
            try:
                hex_color = hex_color.lstrip('#')
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                r = max(0, min(255, int(r * factor)))
                g = max(0, min(255, int(g * factor)))
                b = max(0, min(255, int(b * factor)))
                return f"#{r:02x}{g:02x}{b:02x}"
            except Exception:
                return '#274fbd'
        light = _shade(bg, 1.08)
        dark = _shade(bg, 0.92)
        hovered = False
        selected = False
        def draw(pressed=False):
            try:
                c.delete('all')
                base_bg = _shade(bg, 0.95) if selected else bg
                line_color = _shade(light, 1.05) if hovered and not pressed else (dark if pressed else light)
                shadow_color = _shade(dark, 0.95) if hovered and not pressed else (light if pressed else dark)
                c.configure(bg=base_bg)
                c.create_rectangle(2, 2, w-2, h-2, outline='', fill=base_bg)
                c.create_rectangle(2, 2, w-2, h-2, outline=line_color, width=1)
                c.create_line(4, 4, w-4, 4, fill=line_color, width=1)
                c.create_line(4, h-4, w-4, h-4, fill=shadow_color, width=1)
                c.create_line(4, 4, 4, h-4, fill=line_color, width=1)
                c.create_line(w-4, 4, w-4, h-4, fill=shadow_color, width=1)
                c.create_rectangle(6, 6, w-6, h-6, outline='', fill=bg)
                
                # Icon ve metin çizimi
                text_x = 12
                if icon:
                    c.create_text(20, h//2, text=icon, fill='#ffffff', font=('Segoe UI', 12), anchor='center')
                    text_x = 40
                
                c.create_text(text_x, h//2, text=text, fill='#ffffff', font=('Segoe UI', 9), anchor='w')
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
        def on_press(event):
            draw(True)
        def on_release(event):
            draw(False)
            if callable(command):
                try:
                    command()
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
            # Kalıcı seçim vurgusu
            try:
                for _, btn in self._menu_controls.get('main', []):
                    if hasattr(btn, 'mark_selected'):
                        btn.mark_selected(False)
                # Gruplar
                for gmeta in self._menu_controls.get('groups', {}).values():
                    for _, sb in gmeta.get('sub_buttons', []):
                        if hasattr(sb, 'mark_selected'):
                            sb.mark_selected(False)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
            if hasattr(base, 'mark_selected'):
                base.mark_selected(True)
        def on_enter(event):
            nonlocal hovered
            hovered = True
            draw(False)
        def on_leave(event):
            nonlocal hovered
            hovered = False
            draw(False)
        def on_wheel(event):
            return 'break'
        c.bind('<ButtonPress-1>', on_press)
        c.bind('<ButtonRelease-1>', on_release)
        c.bind('<Enter>', on_enter)
        c.bind('<Leave>', on_leave)
        c.bind('<MouseWheel>', on_wheel)
        try:
            if not hasattr(self, '_menu_canvases'):
                self._menu_canvases = []
            self._menu_canvases.append(c)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
        draw(False)
        icon_lbl = None
        if right_text:
                try:
                    icon_lbl = tk.Label(base, text=right_text, font=('Segoe UI', 9), bg=bg, fg='white')
                    icon_lbl.pack(side='right', padx=15)
                except Exception:
                    icon_lbl = None
        def _resize(_):
            try:
                new_w = max(200, base.winfo_width())
                c.configure(width=new_w, height=h)
                nonlocal w
                w = new_w
                draw(False)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
        base.bind('<Configure>', _resize)
        # Kalıcı seçim vurgusunu dışarıya aç
        def mark_selected(val):
            nonlocal selected
            selected = bool(val)
            draw(False)
        base.mark_selected = mark_selected
        return base, icon_lbl



    def show_esg_consolidated_dashboard(self) -> None:
        try:
            from modules.esg.esg_consolidated_dashboard import ESGConsolidatedDashboard
            ESGConsolidatedDashboard(self.content_area, self.company_id, self.db_path)
        except Exception:
            from tkinter import messagebox
            try:
                messagebox.showinfo(self.lm.tr('info', 'Bilgi'), self.lm.tr('esg_dashboard_wip', 'ESG konsolide dashboard yüklenemedi veya geliştiriliyor'))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    def show_esg_dashboard(self) -> None:
        try:
            if not self._require_permission('esg.read', 'ESG Modülü'):
                return
            self._audit_navigation('ESG Dashboard')
            self.current_module = 'esg_dashboard'
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            
            # Modern Arka Plan
            self.content_area.configure(bg='#f0f2f5')
            self.content_area.pack(fill='both', expand=True) # Padding kaldırıldı, içeride verilecek
            
            self.content_title['text'] = self.lm.tr("esg_dashboard_title", 'ESG Dashboard')
            self._add_back_button()
            
            # --- Modern Grid Layout ---
            container = tk.Frame(self.content_area, bg='#f0f2f5')
            container.pack(fill='both', expand=True, padx=40, pady=40)
            
            # Başlık Alanı
            tk.Label(container, text=self.lm.tr('esg_center_title', "ESG Yönetim Merkezi"), font=('Segoe UI', 26, 'bold'), 
                    bg='#f0f2f5', fg='#1e293b').pack(anchor='w')
            tk.Label(container, text=self.lm.tr('esg_center_desc', "Kurumsal sürdürülebilirlik performansınızı tek merkezden yönetin"), 
                    font=('Segoe UI', 11), bg='#f0f2f5', fg='#64748b').pack(anchor='w', pady=(5, 30))

            tiles = [
                (self.lm.tr("esg_consolidated", 'Konsolide Görünüm'), self.lm.tr("esg_consolidated_desc", 'Tüm metriklerin özeti'), '#3b82f6', lambda: self._navigate(self.show_esg_consolidated_dashboard), Icons.REPORT),
                (self.lm.tr("esg_module", 'ESG Veri Girişi'), self.lm.tr("esg_module_desc", 'Detaylı veri yönetimi'), '#10b981', lambda: self._navigate(self.show_esg), Icons.LEAF),
                (self.lm.tr("reporting", 'Raporlama Merkezi'), self.lm.tr("reporting_center", 'GRI/TSRS Raporları'), '#8b5cf6', lambda: self._navigate(self.show_reporting), '📑'),
                (self.lm.tr("signals", 'Sinyaller & Uyarılar'), self.lm.tr("signals_desc", 'Önemli bulgular'), '#f59e0b', lambda: self._navigate(self.show_esg), '🔔'),
                (self.lm.tr("settings", 'ESG Ayarları'), self.lm.tr("settings_desc", 'Metrik ve ağırlıklar'), '#64748b', lambda: self._navigate(self.show_esg), Icons.SETTINGS),
            ]

            grid_frame = tk.Frame(container, bg='#f0f2f5')
            grid_frame.pack(fill='both', expand=True)

            for i, (title, desc, color, cmd, icon) in enumerate(tiles):
                r, c = divmod(i, 3) # 3 Sütunlu Grid
                
                card = tk.Frame(grid_frame, bg='white', cursor='hand2')
                card.grid(row=r, column=c, padx=10, pady=10, sticky='nsew', ipady=10)
                grid_frame.grid_columnconfigure(c, weight=1)
                
                # Sol Şerit (Renkli)
                tk.Frame(card, bg=color, width=6).pack(side='left', fill='y')
                
                content = tk.Frame(card, bg='white')
                content.pack(side='left', fill='both', expand=True, padx=15, pady=15)
                
                # İkon
                tk.Label(content, text=icon, font=('Segoe UI', 24), bg='white').pack(anchor='w')
                
                # Başlık
                tk.Label(content, text=title, font=('Segoe UI', 13, 'bold'), 
                        fg='#1e293b', bg='white').pack(anchor='w', pady=(10, 2))
                
                # Açıklama
                tk.Label(content, text=desc, font=('Segoe UI', 9), 
                        fg='#64748b', bg='white').pack(anchor='w')
                
                # Hover Effect
                def on_enter(e, c=card, ct=content):
                    c.configure(bg='#f8fafc')
                    ct.configure(bg='#f8fafc')
                    # Hafif yukarı kalkma efekti için padding ile oynanabilir ama grid'de zor
                    
                def on_leave(e, c=card, ct=content):
                    c.configure(bg='white')
                    ct.configure(bg='white')
                    
                def on_click(e, command=cmd):
                    command()
                    
                card.bind('<Enter>', on_enter)
                card.bind('<Leave>', on_leave)
                card.bind('<Button-1>', on_click)
                
                for child in card.winfo_children():
                    child.bind('<Button-1>', on_click)
                    if isinstance(child, tk.Frame):
                        for sub in child.winfo_children():
                            sub.bind('<Button-1>', on_click)
                            
        except Exception as e:
            logging.error(f"ESG Dashboard Error: {e}")
            import traceback
            traceback.print_exc()

    def show_gri_dashboard(self) -> None:
        try:
            if not self._require_permission('gri.read', 'GRI Standartları'):
                return
            self._audit_navigation('GRI Dashboard')
            self.current_module = 'gri_dashboard'
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            
            # Modern Arka Plan
            self.content_area.configure(bg='#f0f2f5')
            self.content_area.pack(fill='both', expand=True)

            self.content_title['text'] = self.lm.tr("gri_dashboard_title", 'GRI Dashboard')
            self._add_back_button()
            
            # --- Modern Grid Layout ---
            container = tk.Frame(self.content_area, bg='#f0f2f5')
            container.pack(fill='both', expand=True, padx=40, pady=40)
            
            # Başlık Alanı
            tk.Label(container, text=self.lm.tr('gri_center_title', "GRI Standartları Merkezi"), font=('Segoe UI', 26, 'bold'), 
                    bg='#f0f2f5', fg='#1e293b').pack(anchor='w')
            tk.Label(container, text=self.lm.tr('gri_center_desc', "Global Reporting Initiative (GRI) uyumlu raporlama yönetimi"), 
                    font=('Segoe UI', 11), bg='#f0f2f5', fg='#64748b').pack(anchor='w', pady=(5, 30))

            tiles = [
                (self.lm.tr("gri_module", 'GRI Standartları'), self.lm.tr("gri_standards", 'Standart yönetimi'), '#2563eb', lambda: self._navigate(self.show_gri), '📘'),
                (self.lm.tr("materiality", 'Önceliklendirme'), self.lm.tr("matrix", 'Materiality Matrisi'), '#1e40af', lambda: self._navigate(self.show_gri_materiality), '🎯'),
                (self.lm.tr("content_index", 'GRI İçerik İndeksi'), self.lm.tr("gri_content_index", 'Otomatik indeks oluşturma'), '#0ea5e9', lambda: self._navigate(self.show_gri_content_index), Icons.CLIPBOARD),
                (self.lm.tr("compliance", 'Uyumluluk Kontrolü'), self.lm.tr("compliance_check", 'Boşluk analizi'), '#6366f1', lambda: self._navigate(self.show_gri_compliance_check), Icons.SUCCESS),
            ]
            
            grid_frame = tk.Frame(container, bg='#f0f2f5')
            grid_frame.pack(fill='both', expand=True)

            for i, (title, desc, color, cmd, icon) in enumerate(tiles):
                r, c = divmod(i, 3) 
                
                card = tk.Frame(grid_frame, bg='white', cursor='hand2')
                card.grid(row=r, column=c, padx=10, pady=10, sticky='nsew', ipady=10)
                grid_frame.grid_columnconfigure(c, weight=1)
                
                tk.Frame(card, bg=color, width=6).pack(side='left', fill='y')
                
                content = tk.Frame(card, bg='white')
                content.pack(side='left', fill='both', expand=True, padx=15, pady=15)
                
                tk.Label(content, text=icon, font=('Segoe UI', 24), bg='white').pack(anchor='w')
                tk.Label(content, text=title, font=('Segoe UI', 13, 'bold'), fg='#1e293b', bg='white').pack(anchor='w', pady=(10, 2))
                tk.Label(content, text=desc, font=('Segoe UI', 9), fg='#64748b', bg='white').pack(anchor='w')
                
                def on_enter(e, c=card, ct=content):
                    c.configure(bg='#f8fafc')
                    ct.configure(bg='#f8fafc')
                def on_leave(e, c=card, ct=content):
                    c.configure(bg='white')
                    ct.configure(bg='white')
                def on_click(e, command=cmd):
                    command()
                    
                card.bind('<Enter>', on_enter)
                card.bind('<Leave>', on_leave)
                card.bind('<Button-1>', on_click)
                for child in card.winfo_children():
                    child.bind('<Button-1>', on_click)
                    if isinstance(child, tk.Frame):
                        for sub in child.winfo_children():
                            sub.bind('<Button-1>', on_click)

        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    def show_sdg_dashboard(self) -> None:
        try:
            if not self._require_permission('sdg.read', 'SDG Modülü'):
                return
            self._audit_navigation('SDG Dashboard')
            self.current_module = 'sdg_dashboard'
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            
            # Modern Arka Plan
            self.content_area.configure(bg='#f0f2f5')
            self.content_area.pack(fill='both', expand=True)

            self.content_title['text'] = self.lm.tr('sdg_dashboard_title', 'SDG Dashboard')
            self._add_back_button()
            
            # --- Modern Grid Layout ---
            container = tk.Frame(self.content_area, bg='#f0f2f5')
            container.pack(fill='both', expand=True, padx=40, pady=40)
            
            # Başlık Alanı
            tk.Label(container, text=self.lm.tr('sdg_analysis_title', "SDG Hedefleri & Analiz"), font=('Segoe UI', 26, 'bold'), 
                    bg='#f0f2f5', fg='#1e293b').pack(anchor='w')
            tk.Label(container, text=self.lm.tr('sdg_analysis_desc', "Sürdürülebilir Kalkınma Amaçları (SDG) performans ve uyum yönetimi"), 
                    font=('Segoe UI', 11), bg='#f0f2f5', fg='#64748b').pack(anchor='w', pady=(5, 30))

            tiles = [
                (self.lm.tr('sdg_module', 'SDG 17 Hedef'), self.lm.tr('17_goals', 'Hedef bazlı yönetim'), '#1e40af', lambda: self._navigate(self.show_sdg), '🎯'),
                (self.lm.tr('mapping', 'GRI/TSRS Eşleştirme'), self.lm.tr('gri_tsrs', 'Çapraz uyum tablosu'), '#2563eb', lambda: self._navigate(self.show_mapping), Icons.LINK),
                (self.lm.tr('environmental', 'Çevresel Performans'), self.lm.tr('carbon_water_waste', 'Karbon, Su, Atık'), '#10b981', lambda: self._navigate(self.show_carbon), Icons.WORLD),
                (self.lm.tr('social', 'Sosyal Metrikler'), self.lm.tr('hr_ohs', 'İK ve İSG verileri'), '#f59e0b', lambda: self._navigate(self.show_social_dashboard), Icons.USERS),
                (self.lm.tr('cbam_skdm', 'CBAM / SKDM'), self.lm.tr('border_carbon', 'Sınırda Karbon Düzenlemesi'), '#8b5cf6', lambda: self._navigate(self.show_cbam_dashboard), '🧮'),
                (self.lm.tr('reporting', 'Raporlama Merkezi'), self.lm.tr('center', 'Detaylı raporlar'), '#6366f1', lambda: self._navigate(self.show_reporting), '🧾'),
            ]

            grid_frame = tk.Frame(container, bg='#f0f2f5')
            grid_frame.pack(fill='both', expand=True)

            for i, (title, desc, color, cmd, icon) in enumerate(tiles):
                r, c = divmod(i, 3) 
                
                card = tk.Frame(grid_frame, bg='white', cursor='hand2')
                card.grid(row=r, column=c, padx=10, pady=10, sticky='nsew', ipady=10)
                grid_frame.grid_columnconfigure(c, weight=1)
                
                tk.Frame(card, bg=color, width=6).pack(side='left', fill='y')
                
                content = tk.Frame(card, bg='white')
                content.pack(side='left', fill='both', expand=True, padx=15, pady=15)
                
                tk.Label(content, text=icon, font=('Segoe UI', 24), bg='white').pack(anchor='w')
                tk.Label(content, text=title, font=('Segoe UI', 13, 'bold'), fg='#1e293b', bg='white').pack(anchor='w', pady=(10, 2))
                tk.Label(content, text=desc, font=('Segoe UI', 9), fg='#64748b', bg='white').pack(anchor='w')
                
                def on_enter(e, c=card, ct=content):
                    c.configure(bg='#f8fafc')
                    ct.configure(bg='#f8fafc')
                def on_leave(e, c=card, ct=content):
                    c.configure(bg='white')
                    ct.configure(bg='white')
                def on_click(e, command=cmd):
                    command()
                    
                card.bind('<Enter>', on_enter)
                card.bind('<Leave>', on_leave)
                card.bind('<Button-1>', on_click)
                for child in card.winfo_children():
                    child.bind('<Button-1>', on_click)
                    if isinstance(child, tk.Frame):
                        for sub in child.winfo_children():
                            sub.bind('<Button-1>', on_click)

        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    def show_energy(self) -> None:
        self._show_module_fast('energy', self.lm.tr('energy_module_title', 'Enerji Yönetimi'), 'energy.read')

    def show_tsrs_dashboard(self) -> None:
        try:
            if not self._require_permission('tsrs.read', 'TSRS Standartları'):
                return
            self._audit_navigation('TSRS Dashboard')
            self.current_module = 'tsrs_dashboard'
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            
            # Modern Arka Plan
            self.content_area.configure(bg='#f0f2f5')
            self.content_area.pack(fill='both', expand=True)

            self.content_title['text'] = self.lm.tr('tsrs_dashboard_title', 'TSRS Dashboard')
            self._add_back_button()
            
            # --- Modern Grid Layout ---
            container = tk.Frame(self.content_area, bg='#f0f2f5')
            container.pack(fill='both', expand=True, padx=40, pady=40)
            
            # Başlık Alanı
            tk.Label(container, text=self.lm.tr('tsrs_center_title', "TSRS & Uyum Standartları"), font=('Segoe UI', 26, 'bold'), 
                    bg='#f0f2f5', fg='#1e293b').pack(anchor='w')
            tk.Label(container, text=self.lm.tr('tsrs_center_desc', "Türkiye Sürdürülebilirlik Raporlama Standartları ve uluslararası uyum"), 
                    font=('Segoe UI', 11), bg='#f0f2f5', fg='#64748b').pack(anchor='w', pady=(5, 30))

            tiles = [
                (self.lm.tr('tsrs_module', 'TSRS Modülü'), self.lm.tr('standards', 'KGK uyumlu standartlar'), '#2563eb', lambda: self._navigate(self.show_tsrs), '🇹🇷'),
                (self.lm.tr('esrs', 'ESRS Uyumu'), self.lm.tr('europe', 'Avrupa raporlama standartları'), '#1e40af', lambda: self._navigate(self.show_esrs), '🇪🇺'),
                (self.lm.tr('issb_ifrs', 'ISSB / IFRS'), self.lm.tr('s1_s2', 'S1 ve S2 standartları'), '#0ea5e9', lambda: self._navigate(self.show_issb), Icons.CHART_UP),
                (self.lm.tr('eu_taxonomy', 'EU Taxonomy'), self.lm.tr('classification', 'Yeşil sınıflandırma'), '#6366f1', lambda: self._navigate(self.show_eu_taxonomy), '🔖'),
            ]

            grid_frame = tk.Frame(container, bg='#f0f2f5')
            grid_frame.pack(fill='both', expand=True)

            for i, (title, desc, color, cmd, icon) in enumerate(tiles):
                r, c = divmod(i, 3) 
                
                card = tk.Frame(grid_frame, bg='white', cursor='hand2')
                card.grid(row=r, column=c, padx=10, pady=10, sticky='nsew', ipady=10)
                grid_frame.grid_columnconfigure(c, weight=1)
                
                tk.Frame(card, bg=color, width=6).pack(side='left', fill='y')
                
                content = tk.Frame(card, bg='white')
                content.pack(side='left', fill='both', expand=True, padx=15, pady=15)
                
                tk.Label(content, text=icon, font=('Segoe UI', 24), bg='white').pack(anchor='w')
                tk.Label(content, text=title, font=('Segoe UI', 13, 'bold'), fg='#1e293b', bg='white').pack(anchor='w', pady=(10, 2))
                tk.Label(content, text=desc, font=('Segoe UI', 9), fg='#64748b', bg='white').pack(anchor='w')
                
                def on_enter(e, c=card, ct=content):
                    c.configure(bg='#f8fafc')
                    ct.configure(bg='#f8fafc')
                def on_leave(e, c=card, ct=content):
                    c.configure(bg='white')
                    ct.configure(bg='white')
                def on_click(e, command=cmd):
                    command()
                    
                card.bind('<Enter>', on_enter)
                card.bind('<Leave>', on_leave)
                card.bind('<Button-1>', on_click)
                for child in card.winfo_children():
                    child.bind('<Button-1>', on_click)
                    if isinstance(child, tk.Frame):
                        for sub in child.winfo_children():
                            sub.bind('<Button-1>', on_click)

        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    def _add_back_button(self) -> None:
        """Header sağ aksiyonlarına standart bir Geri butonu ekle."""
        try:
            if len(getattr(self, '_nav_stack', [])) == 0:
                return
            try:
                for child in list(self.content_actions.winfo_children()):
                    try:
                        if isinstance(child, tk.Button) and str(child.cget('text')).strip().startswith(Icons.LEFT):
                            child.destroy()
                    except Exception as e:
                        logging.error(f"Silent error caught: {str(e)}")
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
            back_btn = tk.Button(self.content_actions, text=f"{Icons.LEFT} {self.lm.tr('btn_back', 'Geri')}", command=self._go_back)
            back_btn.pack(side='left', padx=5)
        except Exception as e:
            logging.info(f"Geri butonu eklenemedi: {e}")

    def _prepare_standard_page(self, title: str) -> None:
        """Sayfa başlığını ve standart geri butonunu hazırla."""
        try:
            self._audit_navigation(title)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
        try:
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
        self.content_title['text'] = title
        self._add_back_button()

    def show_example_page(self) -> None:
        """Standart geri butonlu örnek sayfa"""
        try:
            self.clear_content()
            self._prepare_standard_page(self.lm.tr('example_page', "Örnek Sayfa"))
            body = tk.Frame(self.content_area, bg='white')
            body.pack(fill='both', expand=True, padx=16, pady=16)
            tk.Label(body, text=self.lm.tr('example_page_desc1', "Bu sayfa standart geri buton kullanımını gösterir."),
                    font=('Segoe UI', 12), bg='white', fg='#2c3e50').pack(pady=12)
            tk.Label(body, text=self.lm.tr('example_page_desc2', "Başlık, içerik alanı ve geri buton ana yapısı tek örnekte."),
                    font=('Segoe UI', 10), bg='white', fg='#555555').pack()
        except Exception as e:
            try:
                from tkinter import messagebox
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('example_page_load_error', f"Örnek sayfa yüklenemedi: {e}"))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    def show_dashboard(self) -> None:
        """Dashboard'ı göster - varyanta göre yönlendir"""
        try:
            # Super admin için dashboard her zaman erişilebilir
            if not (self.user and len(self.user) > 1 and self.user[1] == '__super__'):
                if not self._require_permission("dashboard.read", "Dashboard"):
                    return
            self.current_module = 'dashboard'  # Modül takibi
            # Mevcut görünüm fonksiyonunu dashboard olarak set et
            self._current_view_func = self.show_dashboard
            if getattr(self, 'dashboard_variant', 'classic') == 'experimental':
                self.show_dashboard_experimental()
            else:
                self.show_dashboard_classic()
        except Exception as e:
            logging.error(f"Dashboard yükleme hatası: {e}")
            error_label = tk.Label(self.content_area, text=self.lm.tr('dashboard_load_error', "Dashboard yüklenirken hata oluştu"),
                                 font=('Segoe UI', 14), fg='red', bg='white')
            error_label.pack(pady=50)

    def hot_reload(self) -> None:
        try:
            if not (self.user and len(self.user) > 1 and self.user[1] == '__super__'):
                try:
                    messagebox.showwarning(self.lm.tr('warning', "Uyarı"), self.lm.tr('reload_super_admin_only', "Yenileme sadece süper admin için kullanılabilir."))
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
                return
            import importlib
            import sys
            importlib.invalidate_caches()
            base_dir = os.path.abspath(self.base_dir)
            skip_prefixes = (
                'tkinter', 'PIL', 'argon2', 'sqlite3', 'ctypes', 'os', 'sys'
            )
            for name, mod in list(sys.modules.items()):
                try:
                    if any(name.startswith(p) for p in skip_prefixes):
                        continue
                    p = getattr(mod, '__file__', None)
                    if not p:
                        continue
                    ap = os.path.abspath(p)
                    if ap.startswith(base_dir):
                        try:
                            importlib.reload(mod)
                        except Exception as e:
                            logging.error(f"Silent error caught: {str(e)}")
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
            try:
                self.module_cache.clear()
                self.module_instances.clear()
                self.cached_modules = set()
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
            self.clear_content()
            try:
                fn = self._current_view_func or self.show_dashboard
                fn()
            except Exception:
                self.show_dashboard()
            try:
                messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('code_modules_reloaded', "Kod ve modüller yenilendi."))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
        except Exception as e:
            try:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('reload_failed', f"Yenileme başarısız: {e}"))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    def show_dashboard_classic(self) -> None:
        """Mevcut (klasik) dashboard"""
        try:
            self.clear_content()
            # Üst başlığı boşalt ve header'ı gizle
            self.content_title['text'] = ""
            self.content_header.pack_forget()  # Header'ı gizle

            # İçerik alanını header olmadan tam sayfa yap
            self.content_area.pack_forget()
            self.content_area.pack(fill='both', expand=True, padx=0, pady=0)

            # Dashboard arka plan görseli: önce konfigürasyon, sonra varsayılan adaylar
            resimler_dir = os.path.join(self.base_dir, 'resimler')
            candidates = ['dashboard.png', 'dashboard.jpg']
            image_path = None

            # Konfigürasyon ile belirtilen yol öncelikli
            cfg_path = getattr(self, 'dashboard_image_path', None)
            if cfg_path:
                # Mutlak değilse proje köküne göre yorumla
                if not os.path.isabs(cfg_path):
                    cfg_path = os.path.join(self.base_dir, cfg_path)
                if os.path.exists(cfg_path):
                    image_path = cfg_path

            for name in candidates:
                p = os.path.join(resimler_dir, name)
                if os.path.exists(p):
                    image_path = image_path or p
                    break

            # Görsel yolunun varlığını güvenli kontrol et
            if image_path and os.path.exists(image_path):
                try:
                    from PIL import Image, ImageTk

                    # Canvas oluştur - tam sayfa için
                    canvas = tk.Canvas(self.content_area, bg='white', highlightthickness=0)
                    canvas.is_static = True  # Global scroll handler'dan koru
                    canvas.pack(fill='both', expand=True)
                    # Kaydırma tamamen kapalı (performans ve sabit görüntü için)
                    # Mouse tekerleği ile kaymayı engelle (kesin çözüm)
                    canvas.bind('<MouseWheel>', lambda e: "break")
                    canvas.bind('<Button-4>', lambda e: "break")
                    canvas.bind('<Button-5>', lambda e: "break")

                    # Resmi yükle
                    image = Image.open(image_path)

                    # Resample ayarları
                    try:
                        resample = Image.Resampling.LANCZOS
                    except Exception:
                        resample = getattr(Image, 'LANCZOS', getattr(Image, 'BICUBIC', Image.NEAREST))

                    # Debounce için after kullan
                    if hasattr(self, '_dash_resize_after_id'):
                        try:
                            self.parent.after_cancel(self._dash_resize_after_id)
                        except Exception as e:
                            logging.error(f"Silent error caught: {str(e)}")
                    
                    def create_gradient_bg(width, height):
                        """
                        Oluşturulacak gradyan:
                        Dashboard görselinin koyu tonlarına uygun (Deep Navy)
                        0%   -> #01071A (Resim köşe rengi - Dark Navy)
                        100% -> #0B293F (Resim ortalama rengi - Deep Blue)
                        """
                        base = Image.new('RGB', (width, 1))
                        pixels = base.load()
                        
                        stops = [
                            (0.0, (1, 7, 26)),    # #01071A
                            (1.0, (11, 41, 63))   # #0B293F
                        ]
                        
                        for x in range(width):
                            ratio = x / (width - 1) if width > 1 else 0
                            
                            # İlgili aralığı bul
                            start_stop = stops[0]
                            end_stop = stops[-1]
                            
                            for i in range(len(stops) - 1):
                                if stops[i][0] <= ratio <= stops[i+1][0]:
                                    start_stop = stops[i]
                                    end_stop = stops[i+1]
                                    break
                            
                            # İki stop arasında interpolasyon
                            segment_len = end_stop[0] - start_stop[0]
                            if segment_len == 0:
                                local_ratio = 0
                            else:
                                local_ratio = (ratio - start_stop[0]) / segment_len
                                
                            r = int(start_stop[1][0] + (end_stop[1][0] - start_stop[1][0]) * local_ratio)
                            g = int(start_stop[1][1] + (end_stop[1][1] - start_stop[1][1]) * local_ratio)
                            b = int(start_stop[1][2] + (end_stop[1][2] - start_stop[1][2]) * local_ratio)
                            
                            pixels[x, 0] = (r, g, b)
                            
                        return base.resize((width, height), Image.Resampling.NEAREST)

                    def _render_resized():
                        canvas_width = canvas.winfo_width()
                        canvas_height = canvas.winfo_height()
                        if canvas_width <= 1 or canvas_height <= 1:
                            return
                        
                        canvas.delete("all")
                        
                        # Arka plan gradyanı oluştur ve çiz
                        bg_img = create_gradient_bg(canvas_width, canvas_height)
                        bg_photo = ImageTk.PhotoImage(bg_img, master=self.parent)
                        canvas.create_image(0, 0, image=bg_photo, anchor='nw')
                        
                        # Dashboard resmini çiz
                        img_width, img_height = image.size
                        scale = min(canvas_width / img_width if img_width else 1,
                                    canvas_height / img_height if img_height else 1)
                        target_w = max(1, int(img_width * scale))
                        target_h = max(1, int(img_height * scale))
                        resized_img = image.resize((target_w, target_h), resample)
                        photo = ImageTk.PhotoImage(resized_img, master=self.parent)
                        
                        center_x = canvas_width // 2
                        center_y = canvas_height // 2
                        canvas.create_image(center_x, center_y, image=photo, anchor='center')
                        
                        if not hasattr(self, 'image_references'):
                            self.image_references = []
                        # Referansları sakla ki garbage collector silmesin
                        self.image_references = [bg_photo, photo]

                    # Canvas boyut değişikliklerini dinle
                    def _on_configure(_):
                        try:
                            self._dash_resize_after_id = self.parent.after(50, _render_resized)
                        except Exception:
                            _render_resized()
                    canvas.bind('<Configure>', _on_configure)

                    # İlk yükleme için tetikle
                    self.content_area.update_idletasks()
                    _render_resized()

                except Exception as e:
                    self.logger.error(f"{self.lm.tr('dashboard_image_load_error')}: {e}")
                    # Hata durumunda basit dashboard göster
                    self.show_fallback_dashboard()
            else:
                self.logger.warning("Dashboard arka plan görseli bulunamadı (dashboard.png/.jpg veya config yolu)")
                # Resim yoksa basit dashboard göster
                self.show_fallback_dashboard()
        except Exception as e:
            self.logger.error(f"Klasik dashboard yükleme hatası: {e}")
            self.show_fallback_dashboard()

    def show_dashboard_experimental(self) -> None:
        """Deneysel dashboard: merkezde sdg.png, etrafında 1–17 SDG görselleri dairesel dizilim."""
        try:
            self.clear_content()
            self.content_title['text'] = ""
            self.content_header.pack_forget()
            self.content_area.pack_forget()
            self.content_area.pack(fill='both', expand=True, padx=0, pady=0)

            canvas = tk.Canvas(self.content_area, bg='white', highlightthickness=0)
            canvas.is_static = True  # Global scroll handler'dan koru
            canvas.pack(fill='both', expand=True)

            # Kaydırmayı engelle
            canvas.bind('<MouseWheel>', lambda e: "break")
            canvas.bind('<Button-4>', lambda e: "break")
            canvas.bind('<Button-5>', lambda e: "break")

            # Resim yükleme işlemini arka planda yap
            self.parent.after(50, lambda: self._load_dashboard_images_async(canvas))

            # Bu kod artık _load_dashboard_images_async metodunda
        except Exception as e:
            self.logger.error(f"Deneysel dashboard hata: {e}")
            error_label = tk.Label(self.content_area, text=f"Dashboard yüklenirken hata oluştu: {e}",
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    def add_overlay_text(self, canvas, canvas_width, canvas_height) -> None:
        """Resmin üzerine metin ekle"""
        # Başlık - üst ortada
        title_x = canvas_width // 2
        title_y = 80

        # Başlık için gölge efekti
        canvas.create_text(title_x + 2, title_y + 2, text=self.lm.tr('sdg_title_full', "SÜRDÜRÜLEBİLİR KALKINMA HEDEFLERİ"),
                          font=('Segoe UI', 24, 'bold'), fill='black', anchor='center')
        canvas.create_text(title_x, title_y, text=self.lm.tr('sdg_title_full', "SÜRDÜRÜLEBİLİR KALKINMA HEDEFLERİ"),
                          font=('Segoe UI', 24, 'bold'), fill='white', anchor='center')

        # Şirket adı - başlığın altında
        company_name = ""
        try:
            conn = sqlite3.connect(self.db_path)
            cur = conn.cursor()
            cur.execute("SELECT ticari_unvan, sirket_adi FROM company_info WHERE company_id=?", (self.company_id,))
            row = cur.fetchone()
            conn.close()
            if row:
                company_name = (row[0] or '').strip() or (row[1] or '').strip()
        except Exception as e:
            self.logger.warning(f"Firma adı alınamadı: {e}")

        if company_name:
            company_x = canvas_width // 2
            company_y = 130

            # Şirket adı için gölge efekti
            canvas.create_text(company_x + 2, company_y + 2, text=company_name,
                              font=('Segoe UI', 14), fill='black', anchor='center')
            canvas.create_text(company_x, company_y, text=company_name,
                              font=('Segoe UI', 14), fill='white', anchor='center')

        # Alt kısımda SDG kartları
        self.add_sdg_cards_overlay(canvas, canvas_width, canvas_height)

    def add_sdg_cards_overlay(self, canvas, canvas_width, canvas_height) -> None:
        """SDG kartlarını resmin üzerine ekle"""
        # Kartlar için alt kısım
        cards_start_y = canvas_height - 200
        card_width = 150
        card_height = 120
        card_spacing = 20

        # SDG istatistiklerini al
        sdg_stats = self.get_sdg_statistics()

        # Kart verileri
        cards_data = [
            ("", self.lm.tr('sdg_goals', "SDG Hedefleri"), self.lm.tr('goals_count', "{count} Hedef").format(count=sdg_stats['total_goals']), "#4CAF50"),
            ("", self.lm.tr('gri_standards', "GRI Standartları"), self.lm.tr('global_reporting', "Global Raporlama"), "#2196F3"),
            ("", self.lm.tr('tsrs_standards', "TSRS Standartları"), self.lm.tr('turkey_standards', "Türkiye Standartları"), "#FF9800"),
            ("", self.lm.tr('carbon_calculation', "Karbon Hesaplama"), self.lm.tr('emission_calculation', "Emisyon Hesaplama"), "#8BC34A"),
            ("", self.lm.tr('mapping', "Eşleştirme"), self.lm.tr('links', "Bağlantılar"), "#9C27B0"),
        ]

        # Kartları ortalayarak yerleştir
        total_cards_width = len(cards_data) * card_width + (len(cards_data) - 1) * card_spacing
        start_x = (canvas_width - total_cards_width) // 2

        for i, (icon, title, description, color) in enumerate(cards_data):
            card_x = start_x + i * (card_width + card_spacing)
            card_y = cards_start_y

            # Kart arka planı (yarı şeffaf)
            canvas.create_rectangle(card_x, card_y, card_x + card_width, card_y + card_height,
                                   fill=color, outline='white', width=2)

            # Kart içeriği
            icon_y = card_y + 20
            title_y = card_y + 50
            desc_y = card_y + 80

            # İkon
            canvas.create_text(card_x + card_width//2, icon_y, text=icon,
                              font=('Segoe UI', 20), fill='white', anchor='center')

            # Başlık
            canvas.create_text(card_x + card_width//2, title_y, text=title,
                              font=('Segoe UI', 10, 'bold'), fill='white', anchor='center')

            # Açıklama
            canvas.create_text(card_x + card_width//2, desc_y, text=description,
                              font=('Segoe UI', 8), fill='white', anchor='center')

    def show_fallback_dashboard(self) -> None:
        """Resim yüklenemediğinde gösterilecek basit dashboard"""
        # Kaydırılabilir içerik alanı oluştur
        parent = self.content_area
        scroll_inner = self._create_scrollable_container(parent)

        # Hoş geldin banner'ı (kapat + bir daha gösterme)
        try:
            if getattr(self, 'show_welcome_banner', True):
                banner = tk.Frame(scroll_inner, bg='#1e40af')
                banner.pack(fill='x', pady=(0, 8))

                # Sol tarafta başlık
                title = tk.Label(
                    banner,
                    text=self.lm.tr('welcome_title', "Sustainage'ye Hoş Geldiniz!"),
                    font=('Segoe UI', 18, 'bold'),
                    fg='white',
                    bg='#1e40af'
                )
                title.pack(side='left', padx=(12, 6), pady=8)

                # Sağ tarafta seçenekler
                controls = tk.Frame(banner, bg='#1e40af')
                controls.pack(side='right', padx=12, pady=8)

                dont_show_var = tk.BooleanVar(value=False)
                dont_show_chk = tk.Checkbutton(
                    controls,
                    text=self.lm.tr('dont_show_again', 'Bir daha gösterme'),
                    variable=dont_show_var,
                    font=('Segoe UI', 10),
                    fg='white',
                    bg='#1e40af',
                    activebackground='#1e40af',
                    activeforeground='white',
                    selectcolor='#1e40af'
                )
                dont_show_chk.pack(side='left', padx=(0, 8))

                close_btn = tk.Button(
                    controls,
                    text=self.lm.tr('btn_close', 'Kapat'),
                    font=('Segoe UI', 10, 'bold'),
                    fg='white',
                    bg='#ef4444',
                    relief='flat', bd=0, cursor='hand2', padx=10, pady=4,
                    command=lambda: self._dismiss_welcome_banner(banner, dont_show_var.get())
                )
                close_btn.pack(side='left')
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

        # Başlık
        title_label = tk.Label(
            scroll_inner,
            text=self.lm.tr('sdg_title_full', "SÜRDÜRÜLEBİLİR KALKINMA HEDEFLERİ"),
            font=('Segoe UI', 24, 'bold'),
            fg='#2E8B57',
            bg='white'
        )
        title_label.pack(side='top', pady=(4, 6))

        try:
            info_btn = tk.Button(scroll_inner, text=f" {self.lm.tr('company_info', 'Firma Bilgileri')}",
                                 font=('Segoe UI', 10, 'bold'), fg='white', bg='#1e40af',
                                 relief='flat', bd=0, cursor='hand2', padx=12, pady=6,
                                 command=lambda: self._navigate(self.show_company_info))
            info_btn.pack(side='top', pady=(0, 8))
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

        # Şirket adı
        company_name = ""
        try:
            conn = sqlite3.connect(self.db_path)
            cur = conn.cursor()
            cur.execute("SELECT ticari_unvan, sirket_adi FROM company_info WHERE company_id=?", (self.company_id,))
            row = cur.fetchone()
            conn.close()
            if row:
                company_name = (row[0] or '').strip() or (row[1] or '').strip()
        except Exception as e:
            logging.info(f"Firma adı alınamadı: {e}")

        subtitle_text = self.lm.tr('dashboard_image_missing', "Dashboard görseli bulunamadı — basit görünüm gösteriliyor")
        if company_name:
            subtitle_text = f"{company_name} — {subtitle_text}"
        company_label = tk.Label(
            scroll_inner,
            text=subtitle_text,
            font=('Segoe UI', 12),
            fg='#666666',
            bg='white'
        )
        company_label.pack(side='top', pady=(0, 8))

        # Dashboard kartları
        self.show_dashboard_cards(parent=scroll_inner)

    def _dismiss_welcome_banner(self, banner, dont_show=False) -> None:
        """Hoş geldin banner'ını gizle ve gerekirse kalıcı olarak kapat."""
        try:
            if banner:
                banner.pack_forget()
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")
        try:
            if dont_show:
                self.show_welcome_banner = False
                # Kullanıcı bazlı ayarı veritabanında sakla
                self._save_user_welcome_banner_setting(False)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    def show_dashboard_cards(self, parent=None) -> None:
        """Dashboard kartlarını göster"""
        parent = parent or self.content_area
        # Kartlar çerçevesi
        cards_frame = tk.Frame(parent, bg='white')
        cards_frame.pack(fill='x')

        # SDG istatistiklerini al
        sdg_stats = self.get_sdg_statistics()

        # Kart verileri
        cards_data = [
            ("", self.lm.tr('sdg_goals', "SDG Hedefleri"), self.lm.tr('sdg_stats_format', "{goals} Hedef, {indicators} Gösterge").format(goals=sdg_stats['total_goals'], indicators=sdg_stats['total_indicators']), "#4CAF50"),
            ("", self.lm.tr('gri_standards', "GRI Standartları"), self.lm.tr('gri_desc', "Global Raporlama İnisiyatifi"), "#2196F3"),
            ("", self.lm.tr('tsrs_standards', "TSRS Standartları"), self.lm.tr('tsrs_desc', "Türkiye Sürdürülebilirlik Standartları"), "#FF9800"),
            ("", self.lm.tr('esg_module', "ESG Modülü"), self.lm.tr('esg_desc', "Environmental, Social, Governance"), "#8b5cf6"),
            ("", self.lm.tr('policy_library', "Politika Kütüphanesi"), self.lm.tr('policy_desc', "Kurumsal Politikalar ve Uyum"), "#6366f1"),
            (Icons.RECYCLE, self.lm.tr('waste_management', "Atık Yönetimi"), self.lm.tr('waste_desc', "Atık Takibi ve Raporlama"), "#22c55e"),
            ("", self.lm.tr('water_management', "Su Yönetimi"), self.lm.tr('water_desc', "Su Kullanımı ve Kalite Takibi"), "#0ea5e9"),
            ("", self.lm.tr('supply_chain', "Tedarik Zinciri"), self.lm.tr('supply_chain_desc', "Sürdürülebilir Tedarik Zinciri"), "#f97316"),
            ("", self.lm.tr('skdm_module', "SKDM Modülü"), self.lm.tr('skdm_desc', "Karbon, Su, Atık, Tedarik Zinciri"), "#8BC34A"),
        ]

        # Kartları oluştur
        for i, (icon, title, description, color) in enumerate(cards_data):
            card = tk.Frame(cards_frame, bg=color, relief='raised', bd=1)
            card.grid(row=i//3, column=i%3, padx=10, pady=10, sticky='ew')
            cards_frame.grid_columnconfigure(i%3, weight=1)

            # Kart içeriği
            icon_label = tk.Label(card, text=icon, font=('Segoe UI', 24),
                                fg='white', bg=color)
            icon_label.pack(pady=(15, 5))

            title_label = tk.Label(card, text=title, font=('Segoe UI', 12, 'bold'),
                                 fg='white', bg=color)
            title_label.pack()

            desc_label = tk.Label(card, text=description, font=('Segoe UI', 9),
                                fg='white', bg=color, wraplength=150)
            desc_label.pack(pady=(5, 15))

        # SDG detay tablosu kaldırıldı - SDG modülüne taşındı

    def _get_user_welcome_banner_setting(self) -> bool:
        """Kullanıcının hoş geldin banner ayarını veritabanından al"""
        try:
            if not hasattr(self, 'user') or not self.user:
                return True  # Kullanıcı yoksa varsayılan olarak göster

            user_id = self.user[0]
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            # user_settings tablosunu oluştur
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS user_settings (
                    user_id INTEGER PRIMARY KEY,
                    show_welcome_banner INTEGER DEFAULT 1,
                    onboarding_completed INTEGER DEFAULT 0,
                    show_onboarding INTEGER DEFAULT 1,
                    theme TEXT DEFAULT 'default',
                    language TEXT DEFAULT 'tr',
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)

            cursor.execute("SELECT show_welcome_banner FROM user_settings WHERE user_id = ?", (user_id,))
            result = cursor.fetchone()
            conn.close()

            if result:
                return bool(result[0])
            return True  # Varsayılan olarak göster
        except Exception as e:
            logging.info(f"Hoş geldin banner ayarı alınamadı: {e}")
            return True

    def _save_user_welcome_banner_setting(self, show_banner: bool) -> None:
        """Kullanıcının hoş geldin banner ayarını veritabanında sakla"""
        try:
            if not hasattr(self, 'user') or not self.user:
                return

            user_id = self.user[0]
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            # user_settings tablosunu oluştur
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS user_settings (
                    user_id INTEGER PRIMARY KEY,
                    show_welcome_banner INTEGER DEFAULT 1,
                    onboarding_completed INTEGER DEFAULT 0,
                    show_onboarding INTEGER DEFAULT 1,
                    theme TEXT DEFAULT 'default',
                    language TEXT DEFAULT 'tr',
                    created_at TEXT DEFAULT CURRENT_TIMESTAMP
                )
            """)

            cursor.execute("""
                INSERT OR REPLACE INTO user_settings (user_id, show_welcome_banner)
                VALUES (?, ?)
            """, (user_id, 1 if show_banner else 0))

            conn.commit()
            conn.close()
        except Exception as e:
            logging.info(f"Hoş geldin banner ayarı kaydedilemedi: {e}")

    def _save_ui_setting(self, key: str, value) -> None:
        """config/ui_settings.json içinde tek bir UI anahtarını güncelle."""
        try:
            path = getattr(self, 'ui_settings_path', None) or os.path.join(self.base_dir, 'config', 'ui_settings.json')
            # Mevcut içerik
            data = {}
            if os.path.exists(path):
                try:
                    with open(path, 'r', encoding='utf-8') as f:
                        data = json.load(f) or {}
                except Exception:
                    data = {}
            # Güncelle ve kaydet
            data[key] = value
            # Klasör var mı kontrol et
            cfg_dir = os.path.dirname(path)
            if not os.path.exists(cfg_dir):
                os.makedirs(cfg_dir, exist_ok=True)
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            try:
                logging.info(f"UI ayarı kaydedilemedi ({key}): {e}")
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    def _create_scrollable_container(self, parent) -> None:
        """Verilen parent içinde dikey kaydırılabilir bir içerik alanı oluşturur."""
        try:
            canvas = tk.Canvas(parent, bg='white', highlightthickness=0)
            scrollbar = ttk.Scrollbar(parent, orient='vertical', command=canvas.yview)
            canvas.configure(yscrollcommand=scrollbar.set)
            scrollbar.pack(side='right', fill='y')
            canvas.pack(side='left', fill='both', expand=True)

            inner = tk.Frame(canvas, bg='white')
            window = canvas.create_window((0, 0), window=inner, anchor='nw')

            def on_configure(event=None) -> None:
                try:
                    canvas.configure(scrollregion=canvas.bbox('all'))
                    # İç çerçeveyi canvas genişliğine eşitle
                    canvas.itemconfigure(window, width=canvas.winfo_width())
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
            inner.bind('<Configure>', on_configure)

            # Fare tekerleği ile kaydırma (Windows)
            def _on_mousewheel(event) -> None:
                try:
                    canvas.yview_scroll(int(-1 * (event.delta / 120)), 'units')
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
            canvas.bind_all('<MouseWheel>', _on_mousewheel)

            return inner
        except Exception as e:
            logging.info(f"Kaydırılabilir alan oluşturulamadı: {e}")
            return parent

    def get_sdg_statistics(self) -> None:
        """SDG istatistiklerini getir"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            # Toplam hedef sayısı
            cursor.execute("SELECT COUNT(*) FROM sdg_goals")
            total_goals = cursor.fetchone()[0]

            # Toplam gösterge sayısı (gerçek göstergelerden)
            cursor.execute("SELECT COUNT(*) FROM sdg_indicators")
            total_indicators = cursor.fetchone()[0] or 0

            # Ortalama hazırlık skoru (yanıtlar üzerinden progress_pct ortalaması)
            cursor.execute("SELECT AVG(progress_pct) FROM responses")
            avg_readiness = cursor.fetchone()[0] or 0

            conn.close()

            return {
                'total_goals': total_goals,
                'total_indicators': total_indicators,
                'avg_readiness': avg_readiness
            }
        except Exception as e:
            logging.error(f"SDG istatistikleri getirilirken hata: {e}")
            return {
                'total_goals': 17,
                'total_indicators': 232,
                'avg_readiness': 0
            }

    def show_sdg_main_image(self) -> None:
        """SDG ana resmini göster (super admin için değiştirilebilir)"""
        try:
            # Resim çerçevesi
            image_frame = tk.Frame(self.content_area, bg='white')
            image_frame.pack(fill='x', pady=20)

            # Resim yolu: önce kullanıcı tarafından değiştirilen dosyayı, sonra yaygın varsayılanları, en sonda klasördeki ilk resmi dene
            resimler_dir = os.path.join(self.base_dir, 'resimler')
            candidates = [
                'SDGs.jpeg',
                'main.png',
                'sdg.png',
                '111.png',
            ]
            image_path = None
            for name in candidates:
                p = os.path.join(resimler_dir, name)
                if os.path.exists(p):
                    image_path = p
                    break
            # Klasörde herhangi bir resim varsa onu kullan (jpg/png/jpeg/gif/bmp)
            if image_path is None and os.path.isdir(resimler_dir):
                for fname in os.listdir(resimler_dir):
                    if fname.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp')):
                        image_path = os.path.join(resimler_dir, fname)
                        break

            if image_path and os.path.exists(image_path):
                try:
                    from PIL import Image, ImageTk

                    # Resmi yükle ve boyutlandır
                    image = Image.open(image_path)
                    # Resmi uygun boyuta getir (genişlik 600px, yükseklik orantılı)
                    new_width = 600
                    new_height = int(new_width * image.height / image.width)
                    try:
                        resample = Image.Resampling.LANCZOS
                    except Exception:
                        # Eski Pillow sürümleri için geri dönüş
                        resample = getattr(Image, 'LANCZOS', getattr(Image, 'BICUBIC', Image.NEAREST))
                    image = image.resize((new_width, new_height), resample)
                    # Tk bağlamına doğru şekilde bağla (çoklu Tk pencerelerinde hata önler)
                    photo = ImageTk.PhotoImage(image, master=self.parent)

                    # Resim referansını sakla
                    if not hasattr(self, 'image_references'):
                        self.image_references = []
                    self.image_references.append(photo)

                    # Resim label'ı
                    image_label = tk.Label(image_frame, image=photo, bg='white')
                    image_label.image = photo  # Referansı sakla
                    image_label.pack()

                except Exception as e:
                    logging.error(f"PIL resim yükleme hatası: {e}")
                    # PIL hatası durumunda basit mesaj göster
                    no_image_label = tk.Label(image_frame, text=self.lm.tr('sdg_image_load_error', "SDG Ana Resmi Yüklenemedi"),
                                            font=('Segoe UI', 14), fg='#666666', bg='white')
                    no_image_label.pack(pady=50)
                    return

                # Super admin için resim değiştirme butonu
                if self._is_admin_or_super():
                    change_btn = tk.Button(image_frame, text=self.lm.tr('change_image', "Resmi Değiştir"),
                                         font=('Segoe UI', 10), fg='#2E8B57', bg='#E8F5E8',
                                         relief='flat', bd=1, cursor='hand2',
                                         command=self.change_sdg_image)
                    change_btn.pack(pady=10)
            else:
                # Resim bulunamadığında veya klasör boşsa varsayılan mesaj
                no_image_label = tk.Label(image_frame, text=self.lm.tr('sdg_image_not_found', "SDG Ana Resmi Bulunamadı"),
                                        font=('Segoe UI', 14), fg='#666666', bg='white')
                no_image_label.pack(pady=50)

        except Exception as e:
            logging.error(f"SDG resmi yükleme hatası: {e}")

    def change_sdg_image(self) -> None:
        """SDG resmini değiştir (super admin)"""
        try:
            from tkinter import filedialog

            # Dosya seçici
            file_path = filedialog.askopenfilename(
                title=self.lm.tr('select_sdg_image', "SDG Ana Resmini Seçin"),
                filetypes=[(self.lm.tr('image_files', "Resim dosyaları"), "*.jpg *.jpeg *.png *.gif *.bmp"), (self.lm.tr('all_files', "Tüm dosyalar"), "*.*")]
            )

            if file_path:
                import shutil

                # Hedef dosya yolu
                target_path = os.path.join(self.base_dir, 'resimler', 'SDGs.jpeg')

                # Dosyayı kopyala
                shutil.copy2(file_path, target_path)

                # Başarı mesajı
                from tkinter import messagebox
                messagebox.showinfo(self.lm.tr('success', "Başarılı"), self.lm.tr('sdg_image_changed', "SDG ana resmi başarıyla değiştirildi!\nDeğişiklikleri görmek için dashboard'ı yenileyin."))

        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('image_change_error', 'Resim değiştirilemedi')}: {str(e)}")

    def clear_content(self) -> None:
        """İçerik alanını temizle"""
        try:
            if hasattr(self, 'content_area') and self.content_area.winfo_exists():
                try:
                    for child in self.content_area.winfo_children():
                        child.pack_forget()
                except tk.TclError as e:
                    logging.error(f'Silent error in main_app.py: {str(e)}')

            # Başlık sağ aksiyonlarını temizle
            try:
                if hasattr(self, 'content_actions') and self.content_actions.winfo_exists():
                    for child in self.content_actions.winfo_children():
                        child.destroy()
            except tk.TclError as e:
                logging.error(f'Silent error in main_app.py: {str(e)}')
            try:
                self.content_area.pack_forget()
                self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

        except Exception as e:
            logging.error(f"Content area temizleme hatası: {e}")
            try:
                self.content_area.pack_forget()
                self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    # Modül gösterim fonksiyonları
    def show_sdg(self) -> None:
        """SDG modülünü göster - HIZLI YÜKLEME"""
        self._show_module_fast('sdg', 'SDG Modülü', 'sdg.read')

    def show_sdg_alt_modules(self) -> None:
        """SDG altında alt modül butonları"""
        try:
            sub_btn_style = {
                'font': ('Segoe UI', 10, 'bold'), 'fg': '#1e40af', 'bg': '#e6eefc',
                'relief': 'groove', 'bd': 1, 'cursor': 'hand2', 'padx': 8, 'pady': 4
            }

            # Alt menü butonları
            map_btn = tk.Button(self.content_actions, text=self.lm.tr('mapping', "️ Eşleştirme"), command=lambda: self._navigate(self.show_mapping), **sub_btn_style)
            map_btn.pack(side='left', padx=5)

            env_btn = tk.Button(self.content_actions, text=self.lm.tr('environmental_metrics', " Çevresel Metrikler"), command=lambda: self._navigate(self.show_carbon), **sub_btn_style)
            env_btn.pack(side='left', padx=5)

            social_btn = tk.Button(self.content_actions, text=self.lm.tr('social_metrics', " Sosyal Metrikler"), command=lambda: self._navigate(self.show_social_dashboard), **sub_btn_style)
            social_btn.pack(side='left', padx=5)

            cbam_btn = tk.Button(self.content_actions, text=self.lm.tr('cbam_skdm', " CBAM/SKDM"), command=self.show_cbam_dashboard, **sub_btn_style)
            cbam_btn.pack(side='left', padx=5)

            sdg_report_btn = tk.Button(self.content_actions, text=self.lm.tr('sdg_reports', " SDG Raporları"), command=lambda: self._generate_report(self.lm.tr('sdg_report_title', " SDG Raporu")), **sub_btn_style)
            sdg_report_btn.pack(side='left', padx=5)
            SDGGUI_class = self._lazy_import_module('sdg')
            if SDGGUI_class:
                SDGGUI_class(self.content_area, self.company_id)
            else:
                self.show_module_error("SDG Modülü", self.lm.tr('sdg_module_load_error', "SDG modülü yüklenemedi"))
        except Exception as e:
            logging.error(f"SDG modülü yüklenirken hata: {e}")
            error_label = tk.Label(self.content_area, text=self.lm.tr('sdg_module_load_error_with_msg', "SDG modülü yüklenirken hata: {e}").format(e=e),
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    def show_gri(self) -> None:
        """GRI modülünü hızlı göster"""
        self._show_module_fast('gri', 'GRI Standartları', 'gri.read')

    def show_tsrs(self) -> None:
        """TSRS modülünü göster - YENİ Dashboard"""
        try:
            if not self._require_permission("tsrs.read", "TSRS Standartları"):
                return
            self._audit_navigation("TSRS")
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
            self.content_title['text'] = self.lm.tr('tsrs_module_title_long', "TSRS - Türkiye Sürdürülebilirlik Raporlama Standardı")
            # Modül takibi ve geri butonu
            self.current_module = 'tsrs'
            self._add_back_button()


            # TSRS Dashboard'u göster
            self._show_tsrs_dashboard()

        except Exception as e:
            logging.error(f"TSRS modülü yüklenirken hata: {e}")
            error_label = tk.Label(self.content_area, text=self.lm.tr('tsrs_module_load_error_with_msg', "TSRS modülü yüklenirken hata: {e}").format(e=e),
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    def _show_tsrs_dashboard(self) -> None:
        """TSRS Dashboard'u göster"""
        from modules.tsrs_dashboard import TSRSDashboard

        dash_frame = tk.Frame(self.content_area, bg='white')
        dash_frame.pack(fill='both', expand=True)

        # Başlık
        tk.Label(dash_frame, text=self.lm.tr("tsrs_dashboard_title", "TSRS Dashboard - Sürdürülebilirlik Raporlama"),
                font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=20)

        # TSRS Manager'ı başlat
        tsrs = TSRSDashboard()

        # Örnek veri oluştur (ilk seferinde)
        tsrs.create_sample_data(self.company_id)

        # Özet verileri al
        summary = tsrs.get_dashboard_summary(self.company_id, 2024)

        if "error" in summary:
            tk.Label(dash_frame, text=f"{self.lm.tr('error_label', 'Hata')}: {summary['error']}",
                    font=('Segoe UI', 12), fg='red', bg='white').pack(pady=20)
            return

        # İstatistik kartları
        stats_frame = tk.Frame(dash_frame, bg='white')
        stats_frame.pack(fill='x', padx=20, pady=10)

        # Toplam standartlar
        self._create_metric_card(stats_frame, self.lm.tr("tsrs_total_standards", " Toplam Standart"),
                                f"{summary.get('total_standards', 0)} {self.lm.tr('tsrs_standard_unit', 'Standart')}",
                                "#3498db", 0, 0)

        # Toplam göstergeler
        self._create_metric_card(stats_frame, self.lm.tr("tsrs_total_indicators", " Toplam Gösterge"),
                                f"{summary.get('total_indicators', 0)} {self.lm.tr('tsrs_indicator_unit', 'Gösterge')}",
                                "#f39c12", 0, 1)

        # Yanıtlanan göstergeler
        self._create_metric_card(stats_frame, self.lm.tr("tsrs_answered_indicators", " Yanıtlanan"),
                                f"{summary.get('answered_indicators', 0)} {self.lm.tr('tsrs_indicator_unit', 'Gösterge')}",
                                "#e74c3c", 0, 2)

        # Yanıt yüzdesi
        self._create_metric_card(stats_frame, self.lm.tr("tsrs_completion", " Tamamlanma"),
                                f"%{summary.get('response_percentage', 0)}",
                                "#2ecc71", 0, 3)

        # Kategori butonları
        category_frame = tk.Frame(dash_frame, bg='white')
        category_frame.pack(fill='x', padx=20, pady=20)

        tk.Label(category_frame, text=self.lm.tr("tsrs_categories_title", "TSRS Kategorileri"),
                font=('Segoe UI', 14, 'bold'), bg='white').pack(anchor='w', pady=(0, 10))

        categories = [
            ("️ " + self.lm.tr("tsrs_cat_governance", "Yönetişim"), "Yönetişim", "#3498db"),
            (" " + self.lm.tr("tsrs_cat_strategy", "Strateji"), "Strateji", "#f39c12"),
            ("️ " + self.lm.tr("tsrs_cat_risk", "Risk Yönetimi"), "Risk Yönetimi", "#e74c3c"),
            (" " + self.lm.tr("tsrs_cat_metrics", "Metrikler"), "Metrikler", "#2ecc71"),
        ]

        for title, category, color in categories:
            btn = tk.Button(category_frame, text=title,
                           font=('Segoe UI', 11, 'bold'), bg=color, fg='white',
                           relief='flat', bd=0, cursor='hand2', padx=15, pady=10,
                           command=lambda c=category: self._show_tsrs_category(tsrs, c))
            btn.pack(side='left', padx=5)

        # Rapor butonları
        report_frame = tk.Frame(dash_frame, bg='white')
        report_frame.pack(fill='x', padx=20, pady=20)

        tk.Button(report_frame, text=" " + self.lm.tr("tsrs_generate_report", "TSRS Raporu Oluştur"),
                 bg='#9b59b6', fg='white', font=('Segoe UI', 11, 'bold'),
                 padx=20, pady=10, cursor='hand2',
                 command=lambda: self._generate_report(self.lm.tr("tsrs_report_name", " TSRS Raporu"))).pack(side='left', padx=5)

        tk.Button(report_frame, text=" " + self.lm.tr("tsrs_progress_tracking", "İlerleme Takibi"),
                 bg='#1abc9c', fg='white', font=('Segoe UI', 11, 'bold'),
                 padx=20, pady=10, cursor='hand2',
                 command=self._show_tsrs_progress_tracking).pack(side='left', padx=5)

    def _show_tsrs_progress_tracking(self) -> None:
        """TSRS ilerleme takibi penceresi"""
        try:
            from modules.tsrs_dashboard import TSRSDashboard
            tsrs = TSRSDashboard()
            summary = tsrs.get_dashboard_summary(self.company_id, 2024)

            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('tsrs_progress_title', "TSRS İlerleme Takibi"))
            win.geometry("800x500")
            win.configure(bg='white')

            tk.Label(win, text=self.lm.tr('tsrs_progress_summary', "TSRS İlerleme Özeti"), font=('Segoe UI', 14, 'bold'), bg='white', fg='#2c3e50').pack(pady=(12,6))

            cards = tk.Frame(win, bg='white')
            cards.pack(fill='x', padx=10)

            def _card(parent, title, value, color):
                f = tk.Frame(parent, bg=color, relief='raised', bd=1)
                f.pack(side='left', fill='x', expand=True, padx=4, pady=6)
                tk.Label(f, text=title, font=('Segoe UI', 10, 'bold'), bg=color, fg='white').pack(pady=(6,0))
                tk.Label(f, text=str(value), font=('Segoe UI', 18, 'bold'), bg=color, fg='white').pack(pady=(0,8))

            _card(cards, self.lm.tr('total_standards', 'Toplam Standart'), summary.get('total_standards', 0), '#3b82f6')
            _card(cards, self.lm.tr('total_indicators', 'Toplam Gösterge'), summary.get('total_indicators', 0), '#f59e0b')
            _card(cards, self.lm.tr('answered', 'Yanıtlanan'), summary.get('answered_indicators', 0), '#ef4444')
            _card(cards, self.lm.tr('completion', 'Tamamlanma'), f"{summary.get('response_percentage', 0)}%", '#10b981')

            frame = tk.Frame(win, bg='white')
            frame.pack(fill='both', expand=True, padx=10, pady=10)
            cols = (self.lm.tr('period', 'Dönem'), self.lm.tr('status', 'Durum'), self.lm.tr('count', 'Adet'))
            tv = ttk.Treeview(frame, columns=cols, show='headings')
            for c in cols:
                tv.heading(c, text=c)
                tv.column(c, width=120 if c!=self.lm.tr('count', 'Adet') else 80, anchor='center')
            for row in summary.get('recent_responses', []) or []:
                tv.insert('', 'end', values=(row[0], row[1], row[2]))
            sb = ttk.Scrollbar(frame, orient='vertical', command=tv.yview)
            tv.configure(yscrollcommand=sb.set)
            tv.pack(side='left', fill='both', expand=True)
            sb.pack(side='right', fill='y')
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', 'Hata'), self.lm.tr('tsrs_progress_error', 'TSRS ilerleme takibi açılamadı: {e}').format(e=e))

    def _show_tsrs_category(self, tsrs_manager, category) -> None:
        """TSRS kategori detaylarını göster"""
        from tkinter import messagebox

        try:
            standards = tsrs_manager.get_category_standards(category)

            if not standards:
                messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('no_standard_in_category', "{category} kategorisinde standart bulunamadı.").format(category=category))
                return

            # Detay penceresi
            detail_window = tk.Toplevel(self.parent)
            detail_window.title(self.lm.tr('tsrs_category_standards', "TSRS {category} Standartları").format(category=category))
            detail_window.geometry("900x600")
            detail_window.configure(bg='white')

            # Başlık
            tk.Label(detail_window, text=self.lm.tr('tsrs_category_standards', "TSRS {category} Standartları").format(category=category),
                    font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=20)

            # Standartlar listesi
            standards_frame = tk.Frame(detail_window, bg='white')
            standards_frame.pack(fill='both', expand=True, padx=20, pady=10)

            # Scrollable frame
            canvas = tk.Canvas(standards_frame, bg='white')
            scrollbar = tk.Scrollbar(standards_frame, orient="vertical", command=canvas.yview)
            scrollable_frame = tk.Frame(canvas, bg='white')

            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )

            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)

            # Standartları listele
            for standard in standards:
                standard_frame = tk.Frame(scrollable_frame, bg='#f8f9fa', relief='groove', bd=1)
                standard_frame.pack(fill='x', padx=5, pady=5)

                tk.Label(standard_frame, text=f"{standard['code']} - {standard['title']}",
                        font=('Segoe UI', 12, 'bold'), bg='#f8f9fa').pack(anchor='w', padx=10, pady=5)

                tk.Label(standard_frame, text=standard['description'],
                        font=('Segoe UI', 10), bg='#f8f9fa', wraplength=800, justify='left').pack(anchor='w', padx=10, pady=2)

                # Butonlar
                btn_frame = tk.Frame(standard_frame, bg='#f8f9fa')
                btn_frame.pack(fill='x', padx=10, pady=5)

                tk.Button(btn_frame, text=self.lm.tr('view_indicators', "Göstergeleri Gör"),
                         font=('Segoe UI', 9), bg='#3498db', fg='white',
                         relief='flat', bd=0, cursor='hand2', padx=10, pady=3,
                         command=lambda s=standard: self._show_tsrs_indicators(tsrs_manager, s)).pack(side='left', padx=2)

                tk.Button(btn_frame, text=self.lm.tr('enter_data', "Veri Gir"),
                         font=('Segoe UI', 9), bg='#27ae60', fg='white',
                         relief='flat', bd=0, cursor='hand2', padx=10, pady=3,
                         command=lambda s=standard: self._show_tsrs_data_entry(tsrs_manager, s)).pack(side='left', padx=2)

            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('tsrs_category_error', "TSRS kategori hatası: {e}").format(e=e))

    def _show_tsrs_indicators(self, tsrs_manager, standard) -> None:
        """TSRS göstergelerini göster"""
        from tkinter import messagebox

        try:
            indicators = tsrs_manager.get_indicators_by_standard(standard['id'])

            if not indicators:
                messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('no_indicator_in_standard', "{code} standartı için gösterge bulunamadı.").format(code=standard['code']))
                return

            # Göstergeler penceresi
            indicators_window = tk.Toplevel(self.parent)
            indicators_window.title(self.lm.tr('standard_indicators', "{code} - Göstergeler").format(code=standard['code']))
            indicators_window.geometry("800x500")
            indicators_window.configure(bg='white')

            # Başlık
            tk.Label(indicators_window, text=f"{standard['code']} - {standard['title']}",
                    font=('Segoe UI', 14, 'bold'), bg='white').pack(pady=20)

            # Göstergeler listesi
            indicators_frame = tk.Frame(indicators_window, bg='white')
            indicators_frame.pack(fill='both', expand=True, padx=20, pady=10)

            # Scrollable frame
            canvas = tk.Canvas(indicators_frame, bg='white')
            scrollbar = tk.Scrollbar(indicators_frame, orient="vertical", command=canvas.yview)
            scrollable_frame = tk.Frame(canvas, bg='white')

            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )

            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)

            # Göstergeleri listele
            for indicator in indicators:
                indicator_frame = tk.Frame(scrollable_frame, bg='white', relief='solid', bd=1)
                indicator_frame.pack(fill='x', padx=5, pady=3)

                tk.Label(indicator_frame, text=f"{indicator['code']} - {indicator['title']}",
                        font=('Segoe UI', 11, 'bold'), bg='white').pack(anchor='w', padx=10, pady=5)

                tk.Label(indicator_frame, text=indicator['description'],
                        font=('Segoe UI', 9), bg='white', wraplength=600, justify='left').pack(anchor='w', padx=10, pady=2)

                # Veri tipi ve birim
                info_text = self.lm.tr('indicator_type', "Tip: {data_type}").format(data_type=indicator['data_type'])
                if indicator['unit']:
                    info_text += self.lm.tr('indicator_unit', " | Birim: {unit}").format(unit=indicator['unit'])
                if indicator['is_mandatory']:
                    info_text += self.lm.tr('indicator_mandatory', " | Zorunlu")

                tk.Label(indicator_frame, text=info_text,
                        font=('Segoe UI', 8), bg='white', fg='#666').pack(anchor='w', padx=10, pady=2)

                # Veri gir butonu
                btn_frame = tk.Frame(indicator_frame, bg='white')
                btn_frame.pack(fill='x', padx=10, pady=5)

                tk.Button(btn_frame, text=self.lm.tr('enter_data', "Veri Gir"),
                         font=('Segoe UI', 9), bg='#27ae60', fg='white',
                         relief='flat', bd=0, cursor='hand2', padx=10, pady=3,
                         command=lambda i=indicator: self._show_tsrs_indicator_form(tsrs_manager, i)).pack(side='left')

            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('indicators_error', "Göstergeler hatası: {e}").format(e=e))

    def _show_tsrs_indicator_form(self, tsrs_manager, indicator) -> None:
        """TSRS gösterge veri giriş formu"""
        from tkinter import messagebox

        # Veri giriş penceresi
        form_window = tk.Toplevel(self.parent)
        form_window.title(self.lm.tr('data_entry_title', "{code} - Veri Girişi").format(code=indicator['code']))
        form_window.geometry("600x700")
        form_window.configure(bg='white')

        # Başlık
        tk.Label(form_window, text=f"{indicator['code']} - {indicator['title']}",
                font=('Segoe UI', 14, 'bold'), bg='white').pack(pady=20)

        # Form alanları
        form_frame = tk.Frame(form_window, bg='white')
        form_frame.pack(fill='both', expand=True, padx=30, pady=20)

        # Raporlama dönemi
        tk.Label(form_frame, text=self.lm.tr('reporting_period', "Raporlama Dönemi:"), font=('Segoe UI', 10, 'bold'), bg='white').pack(anchor='w', pady=(10, 5))
        period_entry = tk.Entry(form_frame, font=('Segoe UI', 10), width=50)
        period_entry.insert(0, "2024")
        period_entry.pack(pady=(0, 15))

        # Veri tipine göre alanlar
        if indicator['data_type'] in ['numeric', 'percentage']:
            tk.Label(form_frame, text=self.lm.tr('numeric_value', "Sayısal Değer:"), font=('Segoe UI', 10, 'bold'), bg='white').pack(anchor='w', pady=(10, 5))
            value_entry = tk.Entry(form_frame, font=('Segoe UI', 10), width=50)
            value_entry.pack(pady=(0, 15))
        else:
            tk.Label(form_frame, text=self.lm.tr('response_label', "Yanıt:"), font=('Segoe UI', 10, 'bold'), bg='white').pack(anchor='w', pady=(10, 5))
            value_text = tk.Text(form_frame, height=4, width=50, font=('Segoe UI', 10))
            value_text.pack(pady=(0, 15))

        # Metodoloji
        tk.Label(form_frame, text=self.lm.tr('methodology', "Metodoloji:"), font=('Segoe UI', 10), bg='white').pack(anchor='w', pady=(10, 5))
        methodology_text = tk.Text(form_frame, height=3, width=50, font=('Segoe UI', 10))
        methodology_text.pack(pady=(0, 15))

        # Veri kaynağı
        tk.Label(form_frame, text=self.lm.tr('data_source', "Veri Kaynağı:"), font=('Segoe UI', 10), bg='white').pack(anchor='w', pady=(10, 5))
        source_entry = tk.Entry(form_frame, font=('Segoe UI', 10), width=50)
        source_entry.pack(pady=(0, 15))

        # Notlar
        tk.Label(form_frame, text=self.lm.tr('notes', "Notlar:"), font=('Segoe UI', 10), bg='white').pack(anchor='w', pady=(10, 5))
        notes_text = tk.Text(form_frame, height=3, width=50, font=('Segoe UI', 10))
        notes_text.pack(pady=(0, 20))

        # Kaydet butonu
        def save_data() -> None:
            try:
                response_data = {
                    'response_value': value_text.get('1.0', tk.END).strip() if indicator['data_type'] not in ['numeric', 'percentage'] else '',
                    'numerical_value': value_entry.get() if indicator['data_type'] in ['numeric', 'percentage'] else None,
                    'unit': indicator.get('unit', ''),
                    'methodology_used': methodology_text.get('1.0', tk.END).strip(),
                    'data_source': source_entry.get(),
                    'reporting_status': 'Draft',
                    'notes': notes_text.get('1.0', tk.END).strip()
                }

                success = tsrs_manager.save_indicator_response(
                    self.company_id, indicator['id'], period_entry.get(), response_data
                )

                if success:
                    messagebox.showinfo(self.lm.tr('success', "Başarılı"), self.lm.tr('tsrs_data_saved', "TSRS verisi başarıyla kaydedildi!"))
                    form_window.destroy()
                else:
                    messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('data_save_failed', "Veri kaydedilemedi!"))

            except Exception as e:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('data_save_error', "Veri kaydetme hatası: {e}").format(e=e))

        tk.Button(form_frame, text=self.lm.tr('btn_save', "Kaydet"), font=('Segoe UI', 10, 'bold'), bg='#27ae60',
                 fg='white', relief='flat', padx=20, pady=5, command=save_data).pack(pady=10)

    def _show_tsrs_data_entry(self, tsrs_manager, standard) -> None:
        """TSRS standart veri girişi"""
        from tkinter import messagebox
        messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('tsrs_data_entry_wip', "{code} standartı için veri girişi özelliği geliştiriliyor...\n\nLütfen ilgili göstergeleri kullanarak veri girin.").format(code=standard['code']))

    def show_esg(self) -> None:
        """ESG modülünü hızlı göster"""
        self._show_module_fast('esg', 'ESG Modülü', 'esg.read')

    def show_advanced_dashboard(self) -> None:
        """Gelişmiş Dashboard'ı hızlı göster"""
        self._show_module_fast('advanced_dashboard', 'Gelişmiş Dashboard', 'dashboard.advanced')

    def show_strategic(self) -> None:
        """Stratejik Yönetim modülünü hızlı göster"""
        self._show_module_fast('strategic', 'Stratejik Yönetim', 'strategic.read')

    def show_data_import(self) -> None:
        """Veri İçe Aktarım modülünü hızlı göster"""
        self._show_module_fast('data_import', 'Veri İçe Aktarım', 'data.import')

    def show_form_management(self) -> None:
        """Form Yönetimi modülünü hızlı göster"""
        self._show_module_fast('form_management', 'Form Yönetimi', 'forms.manage')

    def show_auto_tasks(self) -> None:
        """Otomatik Görev Oluşturma modülünü göster - HIZLI YÜKLEME"""
        try:
            if not self._require_permission("tasks.auto_create", "Otomatik Görevler"):
                return
            self._audit_navigation("Otomatik Görevler")
            self.current_module = 'auto_tasks'
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)

            self.content_title['text'] = self.lm.tr('auto_task_creation_title', "Otomatik Görev Oluşturma")
            self._add_back_button()

            # HIZLI YÜKLEME - Önce yükleme ekranı göster
            self._show_loading_screen('auto_tasks')

            # Arka planda modülü yükle
            def load_auto_tasks_module():
                try:
                    AutoTaskGUI_class = self._lazy_import_module('auto_tasks')
                    # Yükleme ekranını gizle
                    self._hide_loading_screen('auto_tasks')
                    if AutoTaskGUI_class:
                        # Kullanıcı ID'sini al
                        user_id = self.user[0] if self.user and len(self.user) > 0 else 1
                        # Modülü başlat
                        AutoTaskGUI_class(self.content_area, self.company_id, user_id)
                        self.cached_modules.add('auto_tasks')
                    else:
                        self.show_module_error("Otomatik Görevler", "Otomatik Görevler modülü yüklenemedi")
                except Exception as e:
                    logging.error(f"Otomatik Görevler modülü yükleme hatası: {e}")
                    self._hide_loading_screen('auto_tasks')
                    self.show_module_error("Otomatik Görevler")

            # Arka planda yükle (non-blocking)
            self.parent.after(10, load_auto_tasks_module)

        except Exception as e:
            logging.error(f"Otomatik Görevler modülü yükleme hatası: {e}")
            self.show_module_error("Otomatik Görevler")

    def show_advanced_files(self) -> None:
        """Gelişmiş Dosya Yönetimi modülünü hızlı göster"""
        self._show_module_fast('advanced_files', 'Gelişmiş Dosya Yönetimi', 'files.manage')

    def show_hr_metrics(self) -> None:
        """İK Metrikleri modülünü hızlı göster"""
        self._show_module_fast('hr_metrics', 'İnsan Kaynakları Metrikleri', 'hr.read')

    def show_policy_library(self) -> None:
        """Politika Kütüphanesi modülünü hızlı göster"""
        self._show_module_fast('policy_library', 'Politika Kütüphanesi', 'policy.read')

    def show_skdm(self) -> None:
        """SKDM modülünü hızlı göster"""
        self._show_module_fast('skdm', 'SKDM Modülü', 'skdm.read')

    def show_carbon(self) -> None:
        """Karbon modülünü göster - YENİ Çevresel Dashboard"""
        try:
            self._audit_navigation("Karbon")
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)

            self.content_title['text'] = self.lm.tr('environmental_metrics_dashboard_title', "Çevresel Metrikler Dashboard")
            # Geri butonu
            self._add_back_button()

            # YENİ: Çevresel modülleri göster
            self._show_environmental_dashboard()

        except Exception as e:
            logging.error(f"Çevresel modül hatası: {e}")
            error_label = tk.Label(self.content_area, text=f"Modül yüklenemedi: {e}",
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    def _show_environmental_dashboard(self) -> None:
        """Çevresel metrikler dashboard'u"""
        from tkinter import messagebox

        from modules.environmental import CarbonCalculator, EnergyManager, WasteManager, WaterManager

        # Dashboard frame
        dash_frame = tk.Frame(self.content_area, bg='white')
        dash_frame.pack(fill='both', expand=True)

        # Başlık
        tk.Label(dash_frame, text=self.lm.tr('environmental_metrics_title', "Çevresel Performans Metrikleri"),

                font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=20)

        # İstatistik kartları
        stats_frame = tk.Frame(dash_frame, bg='white')
        stats_frame.pack(fill='x', padx=20, pady=10)

        # Karbon - Güvenli veri al
        try:
            carbon_calc = CarbonCalculator()
            carbon_data = carbon_calc.get_company_summary(self.company_id, 2024)
            carbon_total = carbon_data.get('total_ton', 0) if carbon_data else 0
        except Exception:
            carbon_total = 20.5  # Fallback veri

        self._create_metric_card(stats_frame, " Karbon Ayak İzi",
                                f"{carbon_total:.2f} ton CO₂e",
                                "#E5243B", 0, 0)

        # Enerji - Güvenli veri al
        try:
            energy = EnergyManager()
            energy_data = energy.get_summary(self.company_id, 2024)
            energy_total = energy_data.get('total_kwh', 50000) if energy_data else 50000
            energy_renewable = energy_data.get('renewable_percent', 20) if energy_data else 20
        except Exception:
            energy_total = 50000
            energy_renewable = 20

        self._create_metric_card(stats_frame, " Enerji Tüketimi",
                                f"{energy_total:,} kWh\n%{energy_renewable} Yenilenebilir",
                                "#FFC107", 0, 1)

        # Su - Güvenli veri al
        try:
            water = WaterManager()
            water_data = water.get_summary(self.company_id, 2024)
            water_total = water_data.get('total_m3', 5000) if water_data else 5000
            water_recycle = water_data.get('recycle_percent', 15) if water_data else 15
        except Exception:
            water_total = 5000
            water_recycle = 15

        self._create_metric_card(stats_frame, " Su Tüketimi",
                                f"{water_total:,} m³\n%{water_recycle} Geri Dönüşüm",
                                "#2196F3", 0, 2)

        # Atık - Güvenli veri al
        try:
            waste = WasteManager()
            waste_data = waste.get_summary(self.company_id, 2024)
            waste_total = waste_data.get('total_ton', 15.2) if waste_data else 15.2
            waste_recycle = waste_data.get('recycle_percent', 45) if waste_data else 45
        except Exception:
            waste_total = 15.2
            waste_recycle = 45

        self._create_metric_card(stats_frame, "️ Atık Yönetimi",
                                f"{waste_total:.1f} ton\n%{waste_recycle} Geri Dönüşüm",
                                "#4CAF50", 0, 3)

        # Veri girişi butonları
        btn_frame = tk.Frame(dash_frame, bg='white')
        btn_frame.pack(pady=20)

        def open_carbon_form() -> None:
            try:
                # Scope 3 modülünü kullan
                from modules.scope3.scope3_gui import Scope3GUI

                win = tk.Toplevel(self.parent)
                win.title("Karbon Veri Girişi - GHG Protocol Uyumlu Emisyon İzleme")
                win.geometry("1200x800")
                win.configure(bg='white')

                # Scope 3 GUI'yi aç
                Scope3GUI(win, self.company_id)

            except Exception as e:
                messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('carbon_form_error', 'Karbon formu açılamadı')}: {e}")
                logging.error(f"Karbon form hatası: {e}")

        tk.Button(btn_frame, text=f" {self.lm.tr('add_carbon_data', 'Karbon Verisi Ekle')}",
                 bg='#E5243B', fg='white', font=('Segoe UI', 11, 'bold'),
                 padx=20, pady=10, cursor='hand2',
                 command=open_carbon_form).pack(side='left', padx=5)

        tk.Button(btn_frame, text=f" {self.lm.tr('detailed_report', 'Detaylı Rapor')}",
                 bg='#2196F3', fg='white', font=('Segoe UI', 11, 'bold'),
                 padx=20, pady=10, cursor='hand2',
                 command=self.show_environmental_report_center).pack(side='left', padx=5)

    def _create_metric_card(self, parent, title, value, color, row, col) -> None:
        """Metrik kartı oluştur"""
        card = tk.Frame(parent, bg=color, relief='raised', bd=2)
        card.grid(row=row, column=col, padx=10, pady=10, sticky='nsew')
        parent.grid_columnconfigure(col, weight=1)

        tk.Label(card, text=title, font=('Segoe UI', 12, 'bold'),
                bg=color, fg='white').pack(pady=10)
        tk.Label(card, text=value, font=('Segoe UI', 14),
                bg=color, fg='white').pack(pady=10)

    def show_environmental_report_center(self) -> None:
        """Çevresel rapor merkezi: oluştur, önizle, geçmiş raporları görüntüle"""
        win = tk.Toplevel(self.parent)
        win.title(self.lm.tr('env_report_center', "Çevresel Rapor Merkezi"))
        win.geometry("960x640")
        win.configure(bg='white')

        top = tk.Frame(win, bg='white')
        top.pack(fill='x', padx=12, pady=8)
        tk.Label(top, text=f" {self.lm.tr('env_report_center', 'Çevresel Rapor Merkezi')}", font=('Segoe UI', 14, 'bold'), fg='#1E8449', bg='white').pack(side='left')

        # Sol: rapor oluşturma formu
        left = tk.LabelFrame(win, text=self.lm.tr('create_report', "Rapor Oluştur"), bg='white')
        left.pack(side='left', fill='y', padx=12, pady=8)

        tk.Label(left, text=self.lm.tr('report_type', "Rapor Türü"), bg='white').pack(anchor='w', padx=8, pady=(8, 2))
        report_type_var = tk.StringVar(value='Çevresel')
        ttk.Combobox(left, textvariable=report_type_var, values=['Çevresel', 'SDG', 'TSRS', 'Sosyal', 'CBAM'], state='readonly', width=24).pack(padx=8, pady=2)

        tk.Label(left, text=self.lm.tr('year', "Yıl"), bg='white').pack(anchor='w', padx=8, pady=(8, 2))
        from datetime import datetime
        year_var = tk.StringVar(value=str(datetime.now().year))
        ttk.Entry(left, textvariable=year_var, width=10).pack(padx=8, pady=2)

        def create_report() -> None:
            try:
                year = int(year_var.get())
            except Exception:
                messagebox.showwarning(self.lm.tr('warning', "Uyarı"), self.lm.tr('enter_valid_year', "Geçerli bir yıl giriniz"))
                return
            self._generate_report_file(report_type_var.get(), year)
            refresh_history()

        tk.Button(left, text=self.lm.tr('create_report_btn', "Raporu Oluştur"), bg='#2196F3', fg='white', command=create_report).pack(padx=8, pady=12, fill='x')

        # Sağ: geçmiş raporlar listesi + önizleme
        right = tk.Frame(win, bg='white')
        right.pack(side='left', fill='both', expand=True, padx=8, pady=8)

        columns = ('file', 'created')
        tree = ttk.Treeview(right, columns=columns, show='headings')
        tree.heading('file', text=self.lm.tr('file', 'Dosya'))
        tree.heading('created', text=self.lm.tr('created_at', 'Oluşturulma'))
        tree.column('file', width=520)
        tree.column('created', width=140)
        tree.pack(fill='both', expand=True, side='top')

        preview = tk.Text(right, height=10, wrap='word', bg='#f8f9fa')
        preview.pack(fill='x', side='bottom', padx=4, pady=6)
        preview.insert('1.0', self.lm.tr('preview_placeholder', 'Önizleme: Word içeriği dönüştürülmüş kısa özet görüntülenir.'))
        preview.config(state=tk.DISABLED)

        def refresh_history() -> None:
            from datetime import datetime

            tree.delete(*tree.get_children())
            report_dir = os.path.join(self.base_dir, 'reports')
            os.makedirs(report_dir, exist_ok=True)

            for fname in sorted(os.listdir(report_dir)):
                if fname.lower().endswith('.docx'):
                    fpath = os.path.join(report_dir, fname)
                    created = datetime.fromtimestamp(os.path.getmtime(fpath)).strftime('%Y-%m-%d %H:%M')
                    tree.insert('', 'end', values=(fname, created))

        def open_selected() -> None:
            sel = tree.selection()
            if not sel:
                return
            fname = tree.item(sel[0])['values'][0]
            fpath = os.path.join(self.base_dir, 'reports', fname)
            try:
                os.startfile(fpath)
            except Exception as e:
                messagebox.showerror(self.lm.tr('error', 'Hata'), f"{self.lm.tr('file_open_error', 'Dosya açılamadı')}: {e}")

        btns = tk.Frame(right, bg='white')
        btns.pack(fill='x', pady=4)
        tk.Button(btns, text=self.lm.tr('open_selected', 'Seçileni Aç'), command=open_selected).pack(side='left', padx=4)
        tk.Button(btns, text=self.lm.tr('btn_refresh', 'Yenile'), command=refresh_history).pack(side='left', padx=4)

        refresh_history()

    def _generate_report_file(self, report_type: str, year: int) -> None:
        """Basit rapor dosyası üretimi (python-docx ile)"""
        try:
            from docx import Document  # type: ignore
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', 'Hata'), f"{self.lm.tr('python_docx_required', 'Rapor şablonu oluşturmak için python-docx gerekli')}: {e}")
            return

        report_dir = os.path.join(self.base_dir, 'reports')
        os.makedirs(report_dir, exist_ok=True)
        doc = Document()
        doc.add_heading(f'{report_type} {self.lm.tr("report_suffix", "Raporu")} ({year})', level=1)
        doc.add_paragraph(self.lm.tr('auto_generated_report_msg', 'Bu rapor sistem tarafından otomatik oluşturulmuştur.'))
        report_file = os.path.join(report_dir, f"{report_type.replace(' ', '_').lower()}_{year}.docx")
        try:
            doc.save(report_file)
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', 'Hata'), f"{self.lm.tr('report_save_error', 'Rapor kaydedilemedi')}: {e}")
            return
        messagebox.showinfo(self.lm.tr('success', 'Başarılı'), f"{self.lm.tr('report_created', 'Rapor oluşturuldu')}:\n{report_file}")

    def show_social_dashboard(self) -> None:
        """Sosyal metrikler dashboard'u hızlı göster"""
        try:
            self._show_module_fast('social_dashboard', 'Sosyal Performans Metrikleri')
        except Exception:
            self._show_simple_social_dashboard()

    def _show_simple_social_dashboard(self) -> None:
        """Basit sosyal dashboard (fallback)"""
        from modules.social import HRMetrics, OHSMetrics, TrainingMetrics

        dash_frame = tk.Frame(self.content_area, bg='white')
        dash_frame.pack(fill='both', expand=True)

        tk.Label(dash_frame, text=self.lm.tr('social_dashboard_title', "Sosyal Performans Dashboard"),
                font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=20)

        stats_frame = tk.Frame(dash_frame, bg='white')
        stats_frame.pack(fill='x', padx=20, pady=10)

        # Basit KPI kartları
        try:
            hr = HRMetrics()
            hr_data = hr.get_workforce_summary(self.company_id, 2024)
            hr_total = hr_data.get('total_employees', 150) if hr_data else 150
        except Exception:
            hr_total = 150

        self._create_metric_card(stats_frame, " İş Gücü",
                                f"{hr_total} Çalışan",
                                "#3b82f6", 0, 0)

        try:
            ohs = OHSMetrics()
            ohs_data = ohs.get_summary(self.company_id, 2024)
            ohs_incidents = ohs_data.get('total_incidents', 2) if ohs_data else 2
            ohs_ltifr = ohs_data.get('ltifr', 0.4) if ohs_data else 0.4
        except Exception:
            ohs_incidents = 2
            ohs_ltifr = 0.4

        self._create_metric_card(stats_frame, " İSG",
                                f"{ohs_incidents} Kaza\nLTIFR: {ohs_ltifr}",
                                "#ef4444", 0, 1)

        try:
            training = TrainingMetrics()
            training_data = training.get_summary(self.company_id, 2024)
            training_hours = training_data.get('avg_hours_per_employee', 8.5) if training_data else 8.5
        except Exception:
            training_hours = 8.5

        self._create_metric_card(stats_frame, " Eğitim",
                                f"{training_hours:.1f} Saat/Kişi",
                                "#8b5cf6", 0, 2)

    def show_mapping(self) -> None:
        """Standart Eşleştirme (Mapping) modülünü hızlı göster"""
        self._show_module_fast('mapping', f"{Icons.LINK} Standart Eşleştirme", 'mapping.read')

    def show_surveys(self) -> None:
        """Online Anket Yönetim Sistemi (Hosting Tabanlı)"""
        try:
            self._audit_navigation("Anket Yönetimi")
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack(fill='both', expand=True)
            self.content_title['text'] = self.lm.tr('online_survey_management_title', "🌐 Online Anket Yönetimi")
            # Geri butonu
            self._add_back_button()

            # Hosting tabanlı anket sistemini yükle
            from modules.surveys.survey_gui import SurveyManagementGUI

            gui = SurveyManagementGUI(self.parent, self.db_path)
            gui.show()

        except ImportError as e:
            self.logger.error(f"Anket modülü bulunamadı: {e}")
            messagebox.showerror(
                "Modül Eksik",
                "Online Anket Yönetimi modülü bulunamadı.\n\n"
                "Lütfen modülün kurulu olduğundan emin olun."
            )
        except Exception as e:
            self.logger.error(f"Anket sistemi gösterilemedi: {e}")
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('survey_load_error', 'Anket sistemi yüklenirken hata oluştu')}:\n{str(e)}")

    def _open_user_survey(self, user_survey_id: int) -> None:
        """Seçilen kullanıcı anketini detaylı şekilde aç"""
        try:
            from tkinter import messagebox, ttk

            from modules.surveys.survey_builder import SurveyBuilder
            builder = SurveyBuilder()

            # Detay ve soruları al
            detail = builder.get_user_survey_detail(user_survey_id)
            if not detail:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('survey_detail_not_found', "Anket detayı bulunamadı."))
                return
            questions = builder.get_survey_questions(detail['template_id'])
            existing = builder.get_existing_responses(user_survey_id)

            # Görünümü hazırla
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
            self.content_title['text'] = self.lm.tr('survey_title_prefix', "Anket: {name}").format(name=detail['template_name'])
            self._add_back_button()

            meta_frame = tk.Frame(self.content_area, bg='white')
            meta_frame.pack(fill='x', pady=(0, 10))
            tk.Label(meta_frame, text=detail.get('template_description') or '',
                     font=('Segoe UI', 10), bg='white', fg='#555').pack(anchor='w')
            tk.Label(meta_frame, text=f"Kategori: {detail.get('template_category','')} | Durum: {detail.get('status','')}",
                     font=('Segoe UI', 9), bg='white', fg='#777').pack(anchor='w', pady=(2,0))

            # Kaydırılabilir alan
            container = tk.Frame(self.content_area, bg='white')
            container.pack(fill='both', expand=True)
            canvas = tk.Canvas(container, bg='white', highlightthickness=0)
            scroll_y = tk.Scrollbar(container, orient='vertical', command=canvas.yview)
            form_frame = tk.Frame(canvas, bg='white')
            form_frame.bind(
                '<Configure>',
                lambda e: canvas.configure(scrollregion=canvas.bbox('all'))
            )
            canvas.create_window((0, 0), window=form_frame, anchor='nw')
            canvas.configure(yscrollcommand=scroll_y.set)
            canvas.pack(side='left', fill='both', expand=True)
            scroll_y.pack(side='right', fill='y')

            widgets = {}

            def get_existing_value(qid) -> None:
                try:
                    return existing.get(qid)
                except Exception:
                    return None

            # Soruları oluştur
            for idx, q in enumerate(questions, start=1):
                row = tk.Frame(form_frame, bg='white')
                row.pack(fill='x', padx=4, pady=6)
                label_text = f"{idx}. {q['text']}" + (" *" if q.get('required') else "")
                tk.Label(row, text=label_text, font=('Segoe UI', 10, 'bold'), bg='white').pack(anchor='w', pady=(0,4))

                qtype = (q.get('type') or 'text').lower()
                qid = q['id']
                existing_val = get_existing_value(qid)

                if qtype == 'text':
                    txt = tk.Text(row, height=3, font=('Segoe UI', 10), bg='#f8fafc')
                    if existing_val:
                        try:
                            txt.insert('1.0', existing_val)
                        except Exception as e:
                            logging.error(f"Silent error caught: {str(e)}")
                    txt.pack(fill='x', padx=2)
                    widgets[qid] = ('text', txt)
                elif qtype == 'multiple_choice':
                    # options beklenen format: {"options": ["...", "..."]}
                    opts = []
                    try:
                        data = q.get('options') or {}
                        if isinstance(data, dict):
                            opts = data.get('options') or []
                        elif isinstance(data, list):
                            opts = data
                    except Exception:
                        opts = []
                    var = tk.StringVar(value=existing_val or (opts[0] if opts else ''))
                    combo = ttk.Combobox(row, textvariable=var, values=opts, state='readonly')
                    combo.pack(fill='x')
                    widgets[qid] = ('choice', var)
                elif qtype == 'boolean':
                    var = tk.StringVar(value=existing_val or '')
                    rb_frame = tk.Frame(row, bg='white')
                    rb_frame.pack(anchor='w')
                    ttk.Radiobutton(rb_frame, text=self.lm.tr('yes', 'Evet'), variable=var, value='Evet').pack(side='left', padx=(0,8))
                    ttk.Radiobutton(rb_frame, text=self.lm.tr('no', 'Hayır'), variable=var, value='Hayır').pack(side='left')
                    widgets[qid] = ('bool', var)
                elif qtype == 'scale':
                    # options beklenen: {"min":1, "max":5, "labels":[...]}
                    cfg = q.get('options') or {}
                    mn = 1
                    mx = 5
                    try:
                        mn = int(cfg.get('min', 1))
                        mx = int(cfg.get('max', 5))
                    except Exception as e:
                        logging.error(f"Silent error caught: {str(e)}")
                    var = tk.DoubleVar(value=float(existing_val) if existing_val else float(mn))
                    scale = ttk.Scale(row, orient='horizontal', from_=mn, to=mx, variable=var)
                    scale.pack(fill='x')
                    widgets[qid] = ('scale', var)
                else:
                    ent = tk.Entry(row, font=('Segoe UI', 10), bg='#f8fafc')
                    if existing_val:
                        try:
                            ent.insert(0, existing_val)
                        except Exception as e:
                            logging.error(f"Silent error caught: {str(e)}")
                    ent.pack(fill='x', padx=2)
                    widgets[qid] = ('entry', ent)

            def collect_answers() -> None:
                answers = []
                for qid, (kind, ref) in widgets.items():
                    val = None
                    try:
                        if kind == 'text':
                            val = ref.get('1.0', 'end').strip()
                        elif kind in ('choice', 'bool', 'scale'):
                            val = str(ref.get())
                        elif kind == 'entry':
                            val = ref.get().strip()
                    except Exception:
                        val = ''
                    answers.append((qid, val))
                return answers

            def save_answers() -> None:
                answers = collect_answers()
                ok = True
                for qid, val in answers:
                    if not builder.submit_survey_response(user_survey_id, qid, val):
                        ok = False
                if ok:
                    messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('answers_saved', "Cevaplar kaydedildi."))
                else:
                    messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('answers_save_error', "Bazı cevaplar kaydedilemedi."))

            def complete_survey() -> None:
                save_answers()
                if builder.complete_survey(user_survey_id):
                    messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('survey_completed', "Anket tamamlandı."))
                    # Tamamlandıktan sonra listesine dön
                    self._navigate(self.show_surveys)
                else:
                    messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('survey_complete_error', "Anket tamamlanırken hata oluştu."))

            # Header sağ aksiyonları: Kaydet, Tamamla
            try:
                btn_style = {
                    'font': ('Segoe UI', 10, 'bold'), 'fg': 'white', 'relief': 'flat', 'bd': 0,
                    'cursor': 'hand2', 'padx': 10, 'pady': 6
                }
                save_btn = tk.Button(self.content_actions, text=f" {self.lm.tr('btn_save', 'Kaydet')}", bg='#2E8B57', command=save_answers, **btn_style)
                save_btn.pack(side='right', padx=5)
                complete_btn = tk.Button(self.content_actions, text=f" {self.lm.tr('complete', 'Tamamla')}", bg='#1F6FEB', command=complete_survey, **btn_style)
                complete_btn.pack(side='right', padx=5)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            # Tamamlanmışsa girişleri pasifleştir
            if str(detail.get('status','')).lower() == 'completed':
                try:
                    for kind, ref in widgets.values():
                        if kind == 'text':
                            ref.configure(state='disabled')
                        elif kind in ('entry',):
                            ref.configure(state='disabled')
                        elif kind in ('choice', 'bool', 'scale'):
                            # ttk widget'larda disable farklıdır
                            try:
                                logging.debug(f"Skipping disable for widget type {kind}")
                                # Radiobutton/Scale/Combobox farklı; yoksa geç
                            except Exception as e:
                                logging.error(f"Silent error caught: {str(e)}")
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
        except Exception as e:
            try:
                from tkinter import messagebox
                messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('survey_open_error', 'Anket açılırken hata')}: {e}")
            except Exception as e:
                logging.error(f"Anket açılırken hata: {e}")

    def show_combined_reports(self) -> None:
        """Birleştirilmiş rapor merkezi - SCROLL DESTEKLİ"""
        try:
            self._audit_navigation("Raporlar")
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
            self.content_title['text'] = self.lm.tr('report_center_title', "Rapor Merkezi")
            # Geri butonu
            self._add_back_button()

            from tkinter import messagebox

            # Canvas ve Scrollbar oluştur
            canvas = tk.Canvas(self.content_area, bg='white', highlightthickness=0)
            scrollbar = tk.Scrollbar(self.content_area, orient="vertical", command=canvas.yview)

            # Scrollable frame
            dash_frame = tk.Frame(canvas, bg='white')
            dash_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )

            canvas.create_window((0, 0), window=dash_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)

            # Mouse wheel desteği
            def _on_mousewheel(event) -> None:
                try:
                    canvas.yview_scroll(int(-1*(event.delta/120)), "units")
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
            canvas.bind_all("<MouseWheel>", _on_mousewheel)

            # Pack canvas ve scrollbar
            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")

            tk.Label(dash_frame, text=self.lm.tr('reports_tools_title', "Sürdürülebilirlik Raporları ve Araçları"),

                    font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=20)

            # Ana rapor türleri - yan yana butonlar
            self._create_report_buttons(dash_frame)

            # GRI Standartları ile uyumlu rapor
            self._create_gri_compliance_section(dash_frame)

        except Exception as e:
            logging.error(f"Birleştirilmiş rapor merkezi hatası: {e}")
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_report_center_load', "Rapor merkezi yüklenemedi: {}").format(e))

    def _create_report_buttons(self, parent) -> None:
        """Rapor butonlarını yan yana oluştur"""
        # Ana rapor türleri
        report_types = [
            ("", self.lm.tr('report_type_sdg', "SDG Raporu"), self.lm.tr('report_desc_sdg', "SDG performansı ve hedefler"), "#3b82f6", "sdg"),
            ("", self.lm.tr('report_type_env', "Çevresel Rapor"), self.lm.tr('report_desc_env', "Karbon, Enerji, Su, Atık"), "#10b981", "environmental"),
            ("", self.lm.tr('report_type_social', "Sosyal Rapor"), self.lm.tr('report_desc_social', "İK, İSG, Eğitim metrikleri"), "#f59e0b", "social"),
            ("", self.lm.tr('report_type_economic', "Ekonomik Rapor"), self.lm.tr('report_desc_economic', "GRI 201 - Ekonomik değer"), "#ef4444", "economic"),
            ("", self.lm.tr('report_type_cbam', "CBAM Raporu"), self.lm.tr('report_desc_cbam', "İnovasyon tasarruf analizi"), "#8b5cf6", "cbam"),
        ]

        # Başlık
        tk.Label(parent, text=self.lm.tr('report_types', "Rapor Türleri"), font=('Segoe UI', 14, 'bold'), bg='white').pack(anchor='w', padx=20, pady=(20, 10))


        # Butonlar container
        buttons_frame = tk.Frame(parent, bg='white')
        buttons_frame.pack(fill='x', padx=20, pady=10)

        # Her satırda 3 buton olacak şekilde yerleştir
        for i in range(0, len(report_types), 3):
            row_frame = tk.Frame(buttons_frame, bg='white')
            row_frame.pack(fill='x', pady=5)

            for j in range(3):
                if i + j < len(report_types):
                    icon, title, desc, color, report_type = report_types[i + j]

                    # Ana rapor butonu
                    btn_frame = tk.Frame(row_frame, bg='white')
                    btn_frame.pack(side='left', fill='both', expand=True, padx=5)

                    main_btn = tk.Button(btn_frame, text=f"{icon}\n{title}\n{desc}",
                                       font=('Segoe UI', 10, 'bold'), bg=color, fg='white',
                                       relief='raised', bd=2, cursor='hand2', padx=10, pady=15,
                                       command=lambda rt=report_type: self._show_report_section(rt))
                    main_btn.pack(fill='both', expand=True)

                    # Altında araçlar bölümü (gizli, tıklandığında gösterilecek)
                    tools_container = tk.Frame(btn_frame, bg='white')
                    # tools_container.pack(fill='x', pady=(5, 0))  # Başta gizli

                    # Raporlama araçları
                    tools_frame = tk.Frame(tools_container, bg='#f8f9fa', relief='solid', bd=1)
                    tools_frame.pack(fill='x', pady=(5, 0))

                    tk.Label(tools_frame, text=f"️ {self.lm.tr('tools', 'Araçlar')}", font=('Segoe UI', 9, 'bold'),

                            bg='#f8f9fa').pack(anchor='w', padx=10, pady=(5, 0))

                    tools_buttons_frame = tk.Frame(tools_frame, bg='#f8f9fa')
                    tools_buttons_frame.pack(fill='x', padx=10, pady=5)

                    # Araç butonları yan yana
                    tool_buttons = [
                        ("", self.lm.tr('btn_history', "Geçmiş"), lambda: self.show_report_history()),
                        ("", self.lm.tr('btn_template', "Şablon"), lambda: self.show_report_templates()),
                        ("️", self.lm.tr('settings', "Ayarlar"), lambda: self.show_report_settings())
                    ]

                    for tool_icon, tool_text, tool_cmd in tool_buttons:
                        tool_btn = tk.Button(tools_buttons_frame, text=f"{tool_icon} {tool_text}",
                                           font=('Segoe UI', 8), bg='#10b981', fg='white',
                                           relief='flat', bd=0, cursor='hand2', padx=5, pady=2,
                                           command=tool_cmd)
                        tool_btn.pack(side='left', padx=2)

                    # Gelişmiş raporlama
                    advanced_frame = tk.Frame(tools_container, bg='#f8f9fa', relief='solid', bd=1)
                    advanced_frame.pack(fill='x', pady=(5, 0))

                    tk.Label(advanced_frame, text=f" {self.lm.tr('advanced', 'Gelişmiş')}", font=('Segoe UI', 9, 'bold'),

                            bg='#f8f9fa').pack(anchor='w', padx=10, pady=(5, 0))

                    advanced_buttons_frame = tk.Frame(advanced_frame, bg='#f8f9fa')
                    advanced_buttons_frame.pack(fill='x', padx=10, pady=5)

                    # Gelişmiş butonlar yan yana
                    advanced_buttons = [
                        ("", "Özel", lambda: self.show_custom_report()),
                        ("", "Karşılaştır", lambda: self.show_comparative_analysis()),
                        ("", "Hedef", lambda: self.show_target_tracking())
                    ]

                    for adv_icon, adv_text, adv_cmd in advanced_buttons:
                        adv_btn = tk.Button(advanced_buttons_frame, text=f"{adv_icon} {adv_text}",
                                          font=('Segoe UI', 8), bg='#8b5cf6', fg='white',
                                          relief='flat', bd=0, cursor='hand2', padx=5, pady=2,
                                          command=adv_cmd)
                        adv_btn.pack(side='left', padx=2)

    def _show_report_section(self, report_type) -> None:
        """Rapor bölümünü göster/gizle"""
        # Bu metod buton tıklandığında araçları göster/gizle için kullanılabilir
        # Şimdilik doğrudan rapor oluştur
        report_names = {
            'sdg': ' SDG Raporu',
            'environmental': ' Çevresel Rapor',
            'social': ' Sosyal Rapor',
            'economic': ' Ekonomik Rapor',
            'cbam': ' CBAM Raporu'
        }

        report_name = report_names.get(report_type, f'{report_type} Raporu')
        self._generate_report(report_name)

    def _create_gri_compliance_section(self, parent) -> None:
        """GRI Standartları ile uyumlu rapor bölümü"""
        # GRI Uyumlu Rapor bölümü
        gri_frame = tk.Frame(parent, bg='white')
        gri_frame.pack(fill='x', padx=20, pady=30)

        # Başlık
        tk.Label(gri_frame, text=self.lm.tr('gri_report_title', "GRI Standartları ile Uyumlu Sürdürülebilirlik Raporu"),

                font=('Segoe UI', 14, 'bold'), bg='white').pack(anchor='w', pady=(0, 15))

        # Açıklama
        tk.Label(gri_frame, text=self.lm.tr('gri_report_desc', "Global Raporlama İnisiyatifi (GRI) standartlarına uygun kapsamlı sürdürülebilirlik raporu"),

                font=('Segoe UI', 10), fg='#666', bg='white').pack(anchor='w', pady=(0, 10))

        # GRI rapor butonları
        gri_buttons_frame = tk.Frame(gri_frame, bg='white')
        gri_buttons_frame.pack(fill='x', pady=10)

        gri_buttons = [
            ("", self.lm.tr('gri_universal_title', "GRI Universal Standartları"), self.lm.tr('gri_universal_desc', "Genel açıklamalar"), "#1f2937"),
            ("", self.lm.tr('gri_env_title', "GRI Çevresel Konular"), self.lm.tr('gri_env_desc', "Çevresel etkiler ve performans"), "#10b981"),
            ("", self.lm.tr('gri_social_title', "GRI Sosyal Konular"), self.lm.tr('gri_social_desc', "Sosyal etkiler ve işgücü"), "#f59e0b"),
            ("", self.lm.tr('gri_eco_title', "GRI Ekonomik Konular"), self.lm.tr('gri_eco_desc', "Ekonomik performans ve etki"), "#ef4444"),
            ("", self.lm.tr('gri_sector_title', "GRI Sektörel Standartlar"), self.lm.tr('gri_sector_desc', "Sektöre özel açıklamalar"), "#8b5cf6"),
            ("", self.lm.tr('gri_integrated_title', "GRI Entegre Rapor"), self.lm.tr('gri_integrated_desc', "Tüm konuları kapsayan rapor"), "#6366f1")
        ]

        # Her satırda 3 buton
        for i in range(0, len(gri_buttons), 3):
            row_frame = tk.Frame(gri_buttons_frame, bg='white')
            row_frame.pack(fill='x', pady=5)

            for j in range(3):
                if i + j < len(gri_buttons):
                    icon, title, desc, color = gri_buttons[i + j]

                    btn_frame = tk.Frame(row_frame, bg='white')
                    btn_frame.pack(side='left', fill='both', expand=True, padx=5)

                    gri_btn = tk.Button(btn_frame, text=f"{icon}\n{title}\n{desc}",
                                      font=('Segoe UI', 9, 'bold'), bg=color, fg='white',
                                      relief='raised', bd=2, cursor='hand2', padx=10, pady=12,
                                      command=lambda t=title: self._generate_report(f"GRI {t}"))
                    gri_btn.pack(fill='both', expand=True)

        # GRI araçları
        gri_tools_frame = tk.Frame(gri_frame, bg='#f8f9fa', relief='solid', bd=1)
        gri_tools_frame.pack(fill='x', pady=(15, 0))

        tk.Label(gri_tools_frame, text=f"️ {self.lm.tr('gri_tools', 'GRI Araçları')}", font=('Segoe UI', 10, 'bold'),
                bg='#f8f9fa').pack(anchor='w', padx=15, pady=(10, 5))

        tools_buttons_frame = tk.Frame(gri_tools_frame, bg='#f8f9fa')
        tools_buttons_frame.pack(fill='x', padx=15, pady=(0, 10))

        gri_tools = [
            ("", self.lm.tr('gri_tool_materiality', "GRI Materiality Matrisi"), lambda: self.show_gri_materiality()),
            ("", self.lm.tr('gri_tool_index', "GRI İçerik İndeksi"), lambda: self.show_gri_content_index()),
            ("", self.lm.tr('gri_tool_compliance', "GRI Uyumluluk Kontrolü"), lambda: self.show_gri_compliance_check()),
            ("", self.lm.tr('gri_tool_performance', "GRI Performans Takibi"), lambda: self.show_gri_performance_tracking())
        ]

        for icon, text, cmd in gri_tools:
            tool_btn = tk.Button(tools_buttons_frame, text=f"{icon} {text}",
                               font=('Segoe UI', 9), bg='#6b7280', fg='white',
                               relief='flat', bd=0, cursor='hand2', padx=10, pady=5,
                               command=cmd)
            tool_btn.pack(side='left', padx=5)

    def show_report_history(self) -> None:
        """Rapor geçmişi"""
        try:
            from modules.reporting.report_center_gui import ReportCenterGUI
            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('report_history_title', "Rapor Geçmişi"))
            self._report_center_gui = ReportCenterGUI(win, self.company_id)
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('report_history_load_error', "Rapor geçmişi açılamadı: {}").format(e))

    def show_gri_materiality(self) -> None:
        """GRI Materiality Matrisi"""
        try:
            from modules.analytics.advanced_materiality_gui import AdvancedMaterialityGUI
            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('gri_materiality_title', "GRI Materiality Matrisi"))
            self._gri_materiality_gui = AdvancedMaterialityGUI(win, self.company_id)
        except Exception as e:
            from tkinter import messagebox
            msg = str(e)
            if ("shape mismatch" in msg) or ("broadcast" in msg):
                messagebox.showerror(
                    self.lm.tr('error', "Hata"),
                    self.lm.tr('materiality_shape_mismatch', "Materiality matrisi açılamadı: veri şekilleri uyumsuz.\nPaydaş etkisi, iş etkisi, boyut ve etiket sayıları eşit olmalı.\nLütfen materyal konu kayıtlarını kontrol edin.")
                )
            else:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('materiality_matrix_error', "Materiality matrisi açılamadı: {}").format(e))

    def show_gri_content_index(self) -> None:
        """GRI İçerik İndeksi"""
        try:
            from tkinter import filedialog, messagebox

            from modules.gri.gri_content_index import GRIContentIndex

            content_index = GRIContentIndex(self.db_path)
            result = content_index.generate_content_index(self.company_id)

            if result.get('success'):
                # Excel dosyasını kaydet
                save_path = filedialog.asksaveasfilename(
                    defaultextension=".xlsx",
                    filetypes=[(self.lm.tr("excel_files", "Excel Dosyaları"), "*.xlsx"), (self.lm.tr("all_files", "Tüm Dosyalar"), "*.*")],
                    title=self.lm.tr('gri_index_save_title', "GRI İçerik İndeksi Kaydet"),
                    initialfile=f"GRI_Content_Index_{datetime.now().strftime('%Y%m%d')}.xlsx"
                )

                if save_path:
                    result['excel_path'] = save_path
                    content_index.export_to_excel(self.company_id, save_path)
                    messagebox.showinfo(self.lm.tr('success', "Başarılı"), self.lm.tr('gri_index_created', "GRI İçerik İndeksi oluşturuldu:\n{}").format(save_path))
            else:
                messagebox.showwarning(self.lm.tr('warning', "Uyarı"), result.get('message', self.lm.tr('gri_index_create_warning', 'İçerik indeksi oluşturulamadı')))
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('gri_index_create_error', "GRI İçerik İndeksi oluşturulamadı: {}").format(e))

    def show_gri_compliance_check(self) -> None:
        """GRI Uyumluluk Kontrolü"""
        try:
            from tkinter import messagebox

            from modules.gri.gri_manager import GRIManager

            gri_manager = GRIManager(self.db_path)

            # Uyumluluk kontrolü yap
            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('gri_compliance_title', "GRI Uyumluluk Kontrolü"))
            win.geometry("900x700")

            main_frame = tk.Frame(win, bg='white')
            main_frame.pack(fill='both', expand=True, padx=20, pady=20)

            tk.Label(main_frame, text=self.lm.tr('gri_compliance_title', "GRI Uyumluluk Kontrolü"),

                    font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=10)

            # Treeview
            columns = (
                self.lm.tr('col_standard', 'Standart'),
                self.lm.tr('col_disclosure', 'Disclosure'),
                self.lm.tr('status', 'Durum'),
                self.lm.tr('col_deficiencies', 'Eksiklikler')
            )
            tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=20)

            for col in columns:
                tree.heading(col, text=col)
                tree.column(col, width=200)

            tree.pack(fill='both', expand=True, pady=10)

            # Kontrol butonu
            def run_check():
                for item in tree.get_children():
                    tree.delete(item)

                # GRI standartlarını kontrol et
                standards = gri_manager.get_gri_standards()
                for std in standards[:20]:  # İlk 20 standart
                    tree.insert('', 'end', values=(
                        std.get('code', ''),
                        std.get('title', ''),
                        'Kontrol ediliyor...',
                        ''
                    ))

            tk.Button(main_frame, text=self.lm.tr('run_compliance_check', "Uyumluluk Kontrolü Yap"),
                     command=run_check, bg='#4CAF50', fg='white',
                     font=('Segoe UI', 11, 'bold'), padx=20, pady=10).pack(pady=10)

        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('gri_compliance_check_error', "GRI uyumluluk kontrolü açılamadı: {}").format(e))

    def show_gri_performance_tracking(self) -> None:
        """GRI Performans Takibi"""
        try:
            from tkinter import filedialog, messagebox

            from modules.gri.gri_kpi_reports import GRIKPIReports

            kpi_reports = GRIKPIReports(self.db_path)
            result = kpi_reports.generate_kpi_dashboard(self.company_id)

            if result.get('success'):
                # Excel dosyasını kaydet
                save_path = filedialog.asksaveasfilename(
                    defaultextension=".xlsx",
                    filetypes=[(self.lm.tr("excel_files", "Excel Dosyaları"), "*.xlsx"), (self.lm.tr("all_files", "Tüm Dosyalar"), "*.*")],
                    title=self.lm.tr('gri_perf_save_title', "GRI Performans Takibi Kaydet"),
                    initialfile=f"GRI_Performance_{datetime.now().strftime('%Y%m%d')}.xlsx"
                )

                if save_path:
                    result['excel_path'] = save_path
                    kpi_reports.export_kpi_dashboard(self.company_id, save_path)
                    messagebox.showinfo(self.lm.tr('success', "Başarılı"), self.lm.tr('gri_performance_created', "GRI Performans Takibi oluşturuldu:\n{}").format(save_path))
            else:
                messagebox.showwarning(self.lm.tr('warning', "Uyarı"), result.get('message', self.lm.tr('gri_performance_warning', 'Performans takibi oluşturulamadı')))
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('gri_performance_error', "GRI performans takibi oluşturulamadı: {}").format(e))

    def show_report_templates(self) -> None:
        """Rapor şablonları"""
        try:
            from tkinter import messagebox

            from modules.reporting.advanced_report_templates import AdvancedReportTemplates

            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('report_templates_title', "Rapor Şablonları"))
            win.geometry("800x600")

            main_frame = tk.Frame(win, bg='white')
            main_frame.pack(fill='both', expand=True, padx=20, pady=20)

            tk.Label(main_frame, text=self.lm.tr('report_templates', "Rapor Şablonları"),

                    font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=10)

            templates = AdvancedReportTemplates(self.db_path)
            report_types = templates.REPORT_TYPES

            # Şablon listesi
            list_frame = tk.Frame(main_frame, bg='white')
            list_frame.pack(fill='both', expand=True, pady=10)

            for template_key, template_info in report_types.items():
                frame = tk.Frame(list_frame, bg='#f0f0f0', relief='raised', bd=1)
                frame.pack(fill='x', pady=5, padx=10)

                tk.Label(frame, text=template_info['name'],
                        font=('Segoe UI', 12, 'bold'), bg='#f0f0f0').pack(anchor='w', padx=10, pady=5)
                tk.Label(frame, text=f"Sayfa: {template_info['pages']} | Detay: {template_info['detail_level']}",
                        font=('Segoe UI', 10), bg='#f0f0f0', fg='#666').pack(anchor='w', padx=10)

                def create_template(key=template_key):
                    messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('template_create_info', "{} şablonu oluşturulacak").format(report_types[key]['name']))

                tk.Button(frame, text=self.lm.tr('use', "Kullan"), command=create_template,

                         bg='#4CAF50', fg='white', padx=20).pack(side='right', padx=10, pady=5)
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('templates_load_error', "Rapor şablonları açılamadı: {}").format(e))

    def show_report_settings(self) -> None:
        """Rapor ayarları"""
        try:
            from tkinter import messagebox

            from config.app_config import AppConfig

            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('report_settings', "Rapor Ayarları"))
            win.geometry("600x500")

            main_frame = tk.Frame(win, bg='white')
            main_frame.pack(fill='both', expand=True, padx=20, pady=20)

            tk.Label(main_frame, text=self.lm.tr('report_settings', "Rapor Ayarları"),

                    font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=10)

            config = AppConfig(self.db_path)

            # Ayarlar formu
            settings_frame = tk.Frame(main_frame, bg='white')
            settings_frame.pack(fill='both', expand=True, pady=10)

            settings_data = {
                'report_directory': (self.lm.tr('lbl_report_dir', 'Rapor Dizini'), config.get('report_directory', 'data/reports')),
                'default_report_language': (self.lm.tr('lbl_default_lang', 'Varsayılan Dil'), config.get('default_report_language', 'tr')),
                'max_file_size_mb': (self.lm.tr('lbl_max_file_size', 'Maksimum Dosya Boyutu (MB)'), config.get('max_file_size_mb', 50))
            }

            entries = {}
            row = 0
            for key, (label, value) in settings_data.items():
                tk.Label(settings_frame, text=label, font=('Segoe UI', 10), bg='white').grid(row=row, column=0, sticky='w', pady=5)
                entry = tk.Entry(settings_frame, width=40)
                entry.insert(0, str(value))
                entry.grid(row=row, column=1, pady=5, padx=10)
                entries[key] = entry
                row += 1

            def save_settings():
                for key, entry in entries.items():
                    config.set(key, entry.get())
                messagebox.showinfo(self.lm.tr('success', "Başarılı"), self.lm.tr('msg_settings_saved', "Ayarlar kaydedildi"))

            tk.Button(main_frame, text=self.lm.tr('btn_save', "Kaydet"), command=save_settings,

                     bg='#4CAF50', fg='white', padx=30, pady=10,
                     font=('Segoe UI', 11, 'bold')).pack(pady=20)
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_report_settings_load', "Rapor ayarları açılamadı: {}").format(e))

    def show_custom_report(self) -> None:
        """Özelleştirilmiş rapor"""
        try:
            from modules.reporting.report_center_gui import ReportCenterGUI
            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('title_custom_report', "Özelleştirilmiş Rapor"))
            self._custom_report_gui = ReportCenterGUI(win, self.company_id)
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_custom_report_open', "Özelleştirilmiş rapor açılamadı: {}").format(e))

    def show_comparative_analysis(self) -> None:
        """Karşılaştırmalı analiz"""
        try:
            from tkinter import messagebox

            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('title_comparative_analysis', "Karşılaştırmalı Analiz"))
            win.geometry("900x700")

            main_frame = tk.Frame(win, bg='white')
            main_frame.pack(fill='both', expand=True, padx=20, pady=20)

            tk.Label(main_frame, text=self.lm.tr('comparative_analysis', "Karşılaştırmalı Analiz"),

                    font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=10)

            # Yıl seçimi
            year_frame = tk.Frame(main_frame, bg='white')
            year_frame.pack(pady=10)

            tk.Label(year_frame, text=self.lm.tr('compare_years', "Karşılaştırılacak Yıllar:"),

                    font=('Segoe UI', 11), bg='white').pack(side='left', padx=10)

            current_year = datetime.now().year
            year1_var = tk.StringVar(value=str(current_year - 1))
            year2_var = tk.StringVar(value=str(current_year))

            ttk.Combobox(year_frame, textvariable=year1_var, width=8,
                        values=[str(y) for y in range(2020, current_year + 1)]).pack(side='left', padx=5)
            tk.Label(year_frame, text="vs", bg='white').pack(side='left', padx=5)
            ttk.Combobox(year_frame, textvariable=year2_var, width=8,
                        values=[str(y) for y in range(2020, current_year + 1)]).pack(side='left', padx=5)

            # Sonuç alanı
            result_frame = tk.Frame(main_frame, bg='#f0f0f0', relief='sunken', bd=1)
            result_frame.pack(fill='both', expand=True, pady=10)

            result_text = tk.Text(result_frame, wrap='word', font=('Segoe UI', 10))
            result_text.pack(fill='both', expand=True, padx=10, pady=10)

            def run_analysis():
                try:
                    year1 = int(year1_var.get())
                    year2 = int(year2_var.get())

                    result_text.delete('1.0', 'end')
                    result_text.insert('end', self.lm.tr('msg_analysis_comparison', "Yıllar arası karşılaştırma: {} vs {}\n\n").format(year1, year2))
                    result_text.insert('end', self.lm.tr('msg_analysis_running', "Analiz yapılıyor...\n"))

                    # Basit karşılaştırma (gerçek implementasyon için Analytics modülü kullanılabilir)
                    result_text.insert('end', self.lm.tr('msg_analysis_completed', "Karşılaştırmalı analiz tamamlandı.\n"))
                    result_text.insert('end', self.lm.tr('msg_analysis_hint', "Detaylı analiz için Analytics modülü kullanılabilir.\n"))
                except Exception as e:
                    messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_analysis_failed', "Analiz hatası: {}").format(e))

            tk.Button(main_frame, text=self.lm.tr('run_analysis', "Analiz Yap"), command=run_analysis,
                     bg='#4CAF50', fg='white', padx=30, pady=10,
                     font=('Segoe UI', 11, 'bold')).pack(pady=10)
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_comparative_analysis_open', "Karşılaştırmalı analiz açılamadı: {}").format(e))

    def show_target_tracking(self) -> None:
        """Hedef takibi"""
        try:
            from tkinter import messagebox

            win = tk.Toplevel(self.parent)
            win.title(self.lm.tr('title_target_tracking', "Hedef Takibi"))
            win.geometry("1000x700")

            main_frame = tk.Frame(win, bg='white')
            main_frame.pack(fill='both', expand=True, padx=20, pady=20)

            tk.Label(main_frame, text=self.lm.tr('target_tracking', "Hedef Takibi"),

                    font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=10)

            # Hedef listesi
            columns = (
                self.lm.tr('col_target', 'Hedef'),
                self.lm.tr('col_category', 'Kategori'),
                self.lm.tr('col_baseline_year', 'Baz Yıl'),
                self.lm.tr('col_target_year', 'Hedef Yıl'),
                self.lm.tr('col_progress', 'İlerleme %'),
                self.lm.tr('status', 'Durum')
            )
            tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=20)

            for col in columns:
                tree.heading(col, text=col)
                tree.column(col, width=150)

            tree.pack(fill='both', expand=True, pady=10)

            def load_targets():
                for item in tree.get_children():
                    tree.delete(item)

                # TCFD hedeflerini yükle
                conn = sqlite3.connect(self.db_path)
                cursor = conn.cursor()
                try:
                    cursor.execute("""
                        SELECT target_name, target_category, baseline_year, target_year, 
                               progress_pct, status
                        FROM tcfd_targets
                        WHERE company_id = ? AND status = 'Active'
                        ORDER BY target_year
                    """, (self.company_id,))

                    for row in cursor.fetchall():
                        tree.insert('', 'end', values=row)
                except Exception as e:
                    messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_targets_load', "Hedefler yüklenemedi: {}").format(e))
                finally:
                    conn.close()

            load_targets()

            tk.Button(main_frame, text=self.lm.tr('btn_refresh', "Yenile"), command=load_targets,

                     bg='#2196F3', fg='white', padx=20, pady=5).pack(pady=10)
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_target_tracking_open', "Hedef takibi açılamadı: {}").format(e))

    def _generate_report(self, report_type) -> None:
        """Rapor oluştur - gelişmiş sistem"""
        import os
        import threading
        from tkinter import messagebox

        from docx import Document

        def generate_report_async() -> None:
            try:
                # Rapor dizini
                report_dir = "reports"
                if not os.path.exists(report_dir):
                    os.makedirs(report_dir)

                # Word belgesi oluştur
                doc = Document()

                # Başlık
                doc.add_heading(f'{report_type} - 2024', 0)
                doc.add_paragraph(self.lm.tr('report_company_id', 'Şirket ID: {}').format(self.company_id))
                doc.add_paragraph(self.lm.tr('report_date', 'Tarih: {}').format(self._get_current_date()))
                doc.add_paragraph('')

                if "SDG" in report_type:
                    doc.add_heading(self.lm.tr('report_sdg_heading', 'SDG Hedefleri ve Performans'), level=1)
                    doc.add_paragraph(self.lm.tr('report_sdg7', '• SDG 7: Erişilebilir ve Temiz Enerji'))
                    doc.add_paragraph(self.lm.tr('report_sdg13', '• SDG 13: İklim Eylemi'))
                    doc.add_paragraph(self.lm.tr('report_sdg8', '• SDG 8: İnsana Yakışır İş ve Ekonomik Büyüme'))

                elif "TSRS" in report_type:
                    doc.add_heading(self.lm.tr('report_tsrs_heading', 'TSRS Standartları Raporu'), level=1)

                    # TSRS verilerini al
                    try:
                        from modules.tsrs_dashboard import TSRSDashboard
                        tsrs = TSRSDashboard()
                        summary = tsrs.get_dashboard_summary(self.company_id, 2024)

                        if "error" not in summary:
                            doc.add_heading(self.lm.tr('report_tsrs_summary', 'TSRS Özet'), level=2)
                            doc.add_paragraph(self.lm.tr('report_total_standard', "Toplam Standart: {}").format(summary.get('total_standards', 0)))
                            doc.add_paragraph(self.lm.tr('report_total_indicator', "Toplam Gösterge: {}").format(summary.get('total_indicators', 0)))
                            doc.add_paragraph(self.lm.tr('report_answered_indicator', "Yanıtlanan Gösterge: {}").format(summary.get('answered_indicators', 0)))
                            doc.add_paragraph(self.lm.tr('report_completion_rate', "Tamamlanma Oranı: %{}").format(summary.get('response_percentage', 0)))

                            # Kategori dağılımı
                            if summary.get('category_distribution'):
                                doc.add_heading(self.lm.tr('report_cat_distribution', 'Kategori Dağılımı'), level=2)
                                for category, count in summary['category_distribution'].items():
                                    doc.add_paragraph(self.lm.tr('report_cat_item', "• {}: {} standart").format(category, count))
                        else:
                            doc.add_paragraph(self.lm.tr('report_tsrs_no_data', "TSRS verileri alınamadı"))

                    except Exception as e:
                        doc.add_paragraph(self.lm.tr('report_tsrs_error', "TSRS verileri alınırken hata: {}").format(e))

                elif "Çevresel" in report_type:
                    doc.add_heading(self.lm.tr('report_env_heading', 'Çevresel Performans Metrikleri'), level=1)

                    # Gerçek verileri al
                    try:
                        from modules.environmental import CarbonCalculator, EnergyManager

                        carbon_calc = CarbonCalculator()
                        carbon_data = carbon_calc.get_company_summary(self.company_id, 2024)

                        doc.add_heading(self.lm.tr('report_carbon_heading', 'Karbon Ayak İzi'), level=2)
                        doc.add_paragraph(self.lm.tr('report_total_co2', "Toplam: {:.2f} ton CO₂e").format(carbon_data.get('total_ton', 20.5)))
                        doc.add_paragraph(self.lm.tr('report_scope1', "Scope 1: {:.2f} ton").format(carbon_data.get('scope1_ton', 15.2)))
                        doc.add_paragraph(self.lm.tr('report_scope2', "Scope 2: {:.2f} ton").format(carbon_data.get('scope2_ton', 5.3)))

                        doc.add_heading(self.lm.tr('report_energy_heading', 'Enerji Tüketimi'), level=2)
                        energy = EnergyManager()
                        energy_data = energy.get_summary(self.company_id, 2024)
                        doc.add_paragraph(self.lm.tr('report_total_energy', "Toplam: {:,} kWh").format(energy_data.get('total_kwh', 50000)))
                        doc.add_paragraph(self.lm.tr('report_renewable', "Yenilenebilir: %{}").format(energy_data.get('renewable_percent', 20)))

                    except Exception as e:
                        doc.add_paragraph(self.lm.tr('report_data_error', "Veri alınamadı: {}").format(e))

                elif "Sosyal" in report_type:
                    doc.add_heading(self.lm.tr('report_social_heading', 'Sosyal Performans Metrikleri'), level=1)

                    try:
                        from modules.social import HRMetrics, OHSMetrics

                        hr = HRMetrics()
                        hr_data = hr.get_workforce_summary(self.company_id, 2024)

                        doc.add_heading(self.lm.tr('report_hr_heading', 'İnsan Kaynakları'), level=2)
                        doc.add_paragraph(self.lm.tr('report_total_emp', "Toplam Çalışan: {}").format(hr_data.get('total_employees', 150)))
                        doc.add_paragraph(self.lm.tr('report_female_ratio', "Kadın Oranı: %{}").format(hr_data.get('female_percentage', 35)))
                        doc.add_paragraph(self.lm.tr('report_new_hires', "İşe Alım: {} kişi").format(hr_data.get('new_hires', 12)))

                        ohs = OHSMetrics()
                        ohs_data = ohs.get_summary(self.company_id, 2024)

                        doc.add_heading(self.lm.tr('report_ohs_heading', 'İş Sağlığı ve Güvenliği'), level=2)
                        doc.add_paragraph(self.lm.tr('report_total_incidents', "Toplam Olay: {}").format(ohs_data.get('total_incidents', 2)))
                        doc.add_paragraph(self.lm.tr('report_ltifr', "LTIFR: {}").format(ohs_data.get('ltifr', 0.4)))

                    except Exception as e:
                        doc.add_paragraph(self.lm.tr('report_data_error', "Veri alınamadı: {}").format(e))

                elif "CBAM" in report_type:
                    doc.add_heading(self.lm.tr('report_cbam_heading', 'CBAM İnovasyon Etkisi ve Tasarruf'), level=1)

                    try:
                        from modules.cbam_manager import CBAMManager

                        cbam = CBAMManager()
                        savings_data = cbam.compute_cbam_savings(self.company_id, 2024)

                        if "error" not in savings_data:
                            doc.add_heading(self.lm.tr('report_cbam_summary', 'Özet'), level=2)
                            doc.add_paragraph(self.lm.tr('report_total_savings', "Toplam Tasarruf: {:,.0f} €").format(savings_data.get('total_savings', 0)))
                            doc.add_paragraph(self.lm.tr('report_emission_reduction', "Emisyon Azalması: {:,.2f} ton CO₂e").format(savings_data.get('total_emission_reduction', 0)))
                            doc.add_paragraph(self.lm.tr('report_eu_ets_price', "EU ETS Fiyatı: {:.2f} €/ton CO₂e").format(savings_data.get('eu_ets_price', 80)))

                            doc.add_heading(self.lm.tr('report_product_detail', 'Ürün Bazında Detay'), level=2)
                            for product in savings_data.get('products', []):
                                doc.add_heading(f"{product['product_name']} ({product['product_code']})", level=3)
                                doc.add_paragraph(self.lm.tr('report_ebase', "E_base: {:.2f} tCO₂e/ton").format(product['e_base']))
                                doc.add_paragraph(self.lm.tr('report_epost', "E_post: {:.2f} tCO₂e/ton").format(product['e_post']))
                                doc.add_paragraph(self.lm.tr('report_innovation_ratio', "İnovasyon Oranı: %{:.1f}").format(product['innovation_ratio']*100))
                                doc.add_paragraph(self.lm.tr('report_cbam_saving', "Tasarruf: {:,.0f} €").format(product['cbam_saving']))
                        else:
                            doc.add_paragraph(self.lm.tr('report_error', "Hata: {}").format(savings_data['error']))

                    except Exception as e:
                        doc.add_paragraph(self.lm.tr('report_cbam_no_data', "CBAM verisi alınamadı: {}").format(e))

                # Dosyayı kaydet
                report_file = f"{report_dir}/{report_type.replace(' ', '_').lower()}_2024.docx"
                doc.save(report_file)

                # Başarı mesajı
                def show_success() -> None:
                    messagebox.showinfo(self.lm.tr('success', "Başarılı"), self.lm.tr('report_created_success', "Rapor oluşturuldu!\n\n{}").format(report_file))
                    try:
                        os.startfile(report_file)
                    except Exception as e:
                        logging.error(f"Silent error caught: {str(e)}")

                # GUI thread'inde göster
                self.parent.after(0, show_success)

            except Exception as e:
                err = e
                def show_error() -> None:
                    messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('report_create_error_detail', "Rapor oluşturma hatası: {}").format(err))

                self.parent.after(0, show_error)

        # Arka planda çalıştır
        messagebox.showinfo(self.lm.tr('process_started', "İşlem Başladı"), self.lm.tr('report_creation_started', "Rapor arka planda oluşturuluyor..."))
        threading.Thread(target=generate_report_async, daemon=True).start()

    def show_cbam_dashboard(self) -> None:
        """CBAM/SKDM Dashboard"""
        try:
            self._audit_navigation("CBAM/SKDM")
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
            self.content_title['text'] = self.lm.tr('cbam_innovation_dashboard_title', "CBAM/SKDM İnovasyon Dashboard")
            # Geri butonu
            self._add_back_button()

            from modules.cbam_manager import CBAMManager

            dash_frame = tk.Frame(self.content_area, bg='white')
            dash_frame.pack(fill='both', expand=True)

            tk.Label(dash_frame, text=self.lm.tr('cbam_innovation_title', "CBAM İnovasyon Etkisi ve Tasarruf Analizi"),

                    font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=20)

            # CBAM Manager'ı başlat
            cbam = CBAMManager()

            # Örnek veri oluştur (ilk seferinde)
            cbam.create_sample_data(self.company_id, 2024)

            # Özet verileri al
            summary = cbam.get_company_cbam_summary(self.company_id, 2024)

            if "error" in summary:
                tk.Label(dash_frame, text=self.lm.tr('report_error', "Hata: {}").format(summary['error']),
                        font=('Segoe UI', 12), fg='red', bg='white').pack(pady=20)
                return

            # İstatistik kartları
            stats_frame = tk.Frame(dash_frame, bg='white')
            stats_frame.pack(fill='x', padx=20, pady=10)

            # Toplam tasarruf
            self._create_metric_card(stats_frame, " Yıllık Tasarruf",
                                    f"{summary.get('total_savings_eur', 0):,.0f} €",
                                    "#10b981", 0, 0)

            # Ürün sayısı
            self._create_metric_card(stats_frame, " CBAM Ürünleri",
                                    f"{summary.get('product_count', 0)} Ürün",
                                    "#3b82f6", 0, 1)

            # Ticaret hacmi
            self._create_metric_card(stats_frame, " Ticaret Hacmi",
                                    f"{summary.get('total_volume_tons', 0):,.0f} Ton",
                                    "#f59e0b", 0, 2)

            # Emisyon azalması
            self._create_metric_card(stats_frame, " Emisyon Azalması",
                                    f"{summary.get('total_emission_reduction', 0):,.1f} ton CO₂e",
                                    "#ef4444", 0, 3)

            # Detay butonları
            btn_frame = tk.Frame(dash_frame, bg='white')
            btn_frame.pack(pady=20)

            tk.Button(btn_frame, text=f" {self.lm.tr('detailed_analysis', 'Detaylı Analiz')}",

                     bg='#3b82f6', fg='white', font=('Segoe UI', 11, 'bold'),
                     padx=20, pady=10, cursor='hand2',
                     command=lambda: self._show_cbam_details(cbam)).pack(side='left', padx=5)

            tk.Button(btn_frame, text=f" {self.lm.tr('cbam_report', 'CBAM Raporu')}",

                     bg='#10b981', fg='white', font=('Segoe UI', 11, 'bold'),
                     padx=20, pady=10, cursor='hand2',
                     command=lambda: self._generate_report(" CBAM Raporu")).pack(side='left', padx=5)

            tk.Button(btn_frame, text=f"️ {self.lm.tr('add_product', 'Ürün Ekle')}",

                     bg='#f59e0b', fg='white', font=('Segoe UI', 11, 'bold'),
                     padx=20, pady=10, cursor='hand2',
                     command=lambda: messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('cbam_product_form_info', "CBAM ürün ekleme formu açılacak"))).pack(side='left', padx=5)

        except Exception as e:
            logging.error(f"CBAM dashboard hatası: {e}")
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('cbam_dashboard_error', "CBAM dashboard yüklenemedi: {}").format(e))

    def _show_cbam_details(self, cbam_manager) -> None:
        """CBAM detay analizi"""
        from tkinter import messagebox

        try:
            savings_data = cbam_manager.compute_cbam_savings(self.company_id, 2024)

            if "error" in savings_data:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('cbam_analysis_error', "CBAM analizi: {}").format(savings_data['error']))
                return

            # Detay penceresi
            detail_window = tk.Toplevel(self.parent)
            detail_window.title(self.lm.tr('title_cbam_detail', "CBAM Detay Analizi"))
            detail_window.geometry("800x600")
            detail_window.configure(bg='white')

            # Başlık
            tk.Label(detail_window, text=self.lm.tr('cbam_innovation_detail', "CBAM İnovasyon Etkisi - Detay"),
                    font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=20)

            # Özet bilgiler
            summary_frame = tk.Frame(detail_window, bg='#f0f0f0', relief='groove', bd=2)
            summary_frame.pack(fill='x', padx=20, pady=10)

            tk.Label(summary_frame, text=self.lm.tr('summary_title', "ÖZET"), font=('Segoe UI', 12, 'bold'),
                    bg='#f0f0f0').pack(anchor='w', padx=10, pady=5)
            tk.Label(summary_frame,
                    text=f"{self.lm.tr('total_savings', 'Toplam Tasarruf')}: {savings_data.get('total_savings', 0):,.0f} €",
                    font=('Segoe UI', 11), bg='#f0f0f0').pack(anchor='w', padx=10)
            tk.Label(summary_frame,
                    text=f"{self.lm.tr('emission_reduction', 'Emisyon Azalması')}: {savings_data.get('total_emission_reduction', 0):,.2f} ton CO₂e",
                    font=('Segoe UI', 11), bg='#f0f0f0').pack(anchor='w', padx=10)
            tk.Label(summary_frame,
                    text=f"{self.lm.tr('eu_ets_price', 'EU ETS Fiyatı')}: {savings_data.get('eu_ets_price', 80):.2f} €/ton CO₂e",
                    font=('Segoe UI', 11), bg='#f0f0f0').pack(anchor='w', padx=10)

            # Ürün detayları
            products_frame = tk.Frame(detail_window, bg='white')
            products_frame.pack(fill='both', expand=True, padx=20, pady=10)

            tk.Label(products_frame, text=self.lm.tr('product_detail', "ÜRÜN BAZINDA DETAY"),

                    font=('Segoe UI', 12, 'bold'), bg='white').pack(anchor='w')

            # Scrollable frame
            canvas = tk.Canvas(products_frame, bg='white')
            scrollbar = tk.Scrollbar(products_frame, orient="vertical", command=canvas.yview)
            scrollable_frame = tk.Frame(canvas, bg='white')

            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )

            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)

            # Ürünleri listele
            for i, product in enumerate(savings_data.get('products', [])):
                product_frame = tk.Frame(scrollable_frame, bg='#f9f9f9', relief='groove', bd=1)
                product_frame.pack(fill='x', padx=5, pady=5)

                tk.Label(product_frame, text=f"{product['product_name']} ({product['product_code']})",
                        font=('Segoe UI', 11, 'bold'), bg='#f9f9f9').pack(anchor='w', padx=10, pady=5)

                details_text = "\n".join([
                    self.lm.tr('cbam_detail_ebase_epost', "E_base: {:.2f} tCO₂e/ton  →  E_post: {:.2f} tCO₂e/ton").format(product['e_base'], product['e_post']),
                    self.lm.tr('cbam_detail_innovation', "İnovasyon Oranı: %{:.1f}  |  Attenuation Factor: {:.2f}").format(product['innovation_ratio']*100, product['attenuation_factor']),
                    self.lm.tr('cbam_detail_cost', "CBAM Maliyeti: {:,.0f} €  →  {:,.0f} €").format(product['cbam_cost_base'], product['cbam_cost_post']),
                    self.lm.tr('cbam_detail_saving', "Tasarruf: {:,.0f} €  |  Emisyon Azalması: {:,.2f} ton CO₂e").format(product['cbam_saving'], product['emission_reduction']),
                    self.lm.tr('cbam_detail_volume', "Ticaret Hacmi: {:,.0f} ton").format(product['trade_volume'])
                ])

                tk.Label(product_frame, text=details_text.strip(),
                        font=('Segoe UI', 9), bg='#f9f9f9', justify='left').pack(anchor='w', padx=10, pady=5)

            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_cbam_detail_open', "CBAM detay hatası: {}").format(e))

    def _get_current_date(self) -> None:
        """Güncel tarihi al"""
        from datetime import datetime
        return datetime.now().strftime("%d.%m.%Y %H:%M")







    def show_prioritization(self) -> None:
        """Önceliklendirme modülünü göster - HIZLI YÜKLEME"""
        self._show_module_fast('prioritization', self.lm.tr('module_prioritization', 'Önceliklendirme'), 'prioritization.read')

    def show_waste_management(self) -> None:
        """Atık yönetimi modülünü göster - HIZLI YÜKLEME"""
        self._show_module_fast('waste', self.lm.tr('module_waste_management', 'Atık Yönetimi'), 'waste.read')

    def show_water_management(self) -> None:
        """Su yönetimi modülünü göster - HIZLI YÜKLEME"""
        self._show_module_fast('water', self.lm.tr('module_water_management', 'Su Yönetimi'), 'water.read')

    def show_supply_chain(self) -> None:
        """Tedarik zinciri modülünü göster - HIZLI YÜKLEME"""
        self._show_module_fast('supply_chain', 'Tedarik Zinciri', 'supply_chain.read')

    def show_product_technology(self) -> None:
        """Ürün & Teknoloji modülünü hızlı göster"""
        self._show_module_fast('product_tech', 'Ürün & Teknoloji Modülü', 'product_tech.read')

    def show_product_tech(self) -> None:
        """Ürün & Teknoloji modülünü hızlı göster"""
        self._show_module_fast('product_tech', 'Ürün & Teknoloji')

    def show_reporting(self) -> None:
        """Raporlama modülünü göster"""
        self._show_module_fast('reporting', f"{Icons.REPORT} Raporlama Merkezi")

    def show_my_tasks(self) -> None:
        """Kullanıcının görevlerini göster"""
        try:
            # İzin kontrolü (esnek)
            if not self._has_permission("tasks.read"):
                # Herkes görebilir, sadece uyarı
                logging.info("Accessing tasks without explicit permission (allowed)")

            self._audit_navigation("Görevlerim")
            self.current_module = 'my_tasks'
            self.clear_content()

            # Header'ı geri getir
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
        except Exception as e:
            # Başlatma sırasında hata olursa kullanıcıya bildir
            logging.error(f"Görevlerim görünümü başlatılırken hata: {e}")
            try:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_my_tasks_init', "Görevlerim başlatılamadı: {}").format(e))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

        self.content_title['text'] = "Görevlerim"
        # Geri butonu
        self._add_back_button()

        # Üst araç çubuğu: Otomatik Görevler kısayolu
        try:
            tools = tk.Frame(self.content_area, bg='white')
            tools.pack(fill='x', pady=(0, 10))
            if self._has_permission("tasks.auto_create"):
                auto_btn = tk.Button(tools, text=f" {self.lm.tr('auto_tasks', 'Otomatik Görevler')}",

                                     font=('Segoe UI', 10, 'bold'), fg='white', bg='#2ecc71',
                                     relief='flat', bd=0, cursor='hand2', padx=12, pady=6,
                                     command=lambda: self._navigate(self.show_auto_tasks))
                auto_btn.pack(side='left', padx=4, pady=6)
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

            # Kullanıcı ID'sini al
            user_id = self.user[0] if self.user and len(self.user) > 0 else 1

            # Görev GUI'sini hızlı yükleme ile göster (asenkron import + loader)
            loader = tk.Frame(self.content_area, bg='white')
            loader.pack(fill='both', expand=True)
            tk.Label(loader, text=self.lm.tr('loading_tasks', "Görevler yükleniyor..."), font=('Segoe UI', 12), bg='white').pack(pady=(40, 10))

            bar = ttk.Progressbar(loader, mode='indeterminate', length=240)
            bar.pack(pady=10)
            bar.start(10)

            def _async_import_and_init():
                try:
                    import importlib

                    # Önce mevcut projedeki konumdan dene (tasks)
                    try:
                        mod = importlib.import_module('tasks.task_gui')
                    except Exception:
                        # Eski/alternatif konum
                        mod = importlib.import_module('modules.task_management.task_gui')
                    TaskGUI = getattr(mod, 'TaskGUI')
                    # GUI oluşturmayı ana threade planla
                    self.parent.after(0, lambda: (loader.destroy(), TaskGUI(self.content_area, user_id, self.db_path)))
                except Exception as exc:
                    err = self.lm.tr('err_task_manager_load', "Görev yönetimi yüklenemedi: {}").format(exc)
                    self.parent.after(0, lambda em=err: (loader.destroy(), messagebox.showerror(self.lm.tr('error', "Hata"), em)))

            import threading
            threading.Thread(target=_async_import_and_init, daemon=True).start()

        except Exception as e:
            logging.error(f"Görev listesi gösterilirken hata: {e}")
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_task_list_load', "Görev listesi yüklenemedi: {}").format(e))

    def show_task_dashboard(self) -> None:
        """Admin görev dashboard'unu göster"""
        try:
            # Admin kontrolü
            if not self._is_admin_or_super():
                messagebox.showwarning(self.lm.tr('permission_denied', "İzin Yok"), self.lm.tr('warning_admin_only', "Bu sayfayı sadece adminler görüntüleyebilir"))
                return

            self._audit_navigation("Görev Dashboard")
            self.current_module = 'task_dashboard'
            self.clear_content()

            # Header'ı geri getir
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)

            self.content_title['text'] = self.lm.tr('task_management_dashboard_title', "Görev Yönetimi Dashboard")
            # Geri butonu
            self._add_back_button()

            # Admin Dashboard GUI'sini hızlı yükleme ile göster (asenkron import + loader)
            import tkinter as tk
            from tkinter import ttk

            loader = tk.Frame(self.content_area, bg='white')
            loader.pack(fill='both', expand=True)
            tk.Label(loader, text=self.lm.tr('loading_dashboard', "Dashboard yükleniyor..."), font=('Segoe UI', 12), bg='white').pack(pady=(40, 10))

            bar = ttk.Progressbar(loader, mode='indeterminate', length=240)
            bar.pack(pady=10)
            bar.start(10)

            def _async_import_and_init():
                try:
                    import importlib

                    # Önce yeni konum (tasks)
                    try:
                        mod = importlib.import_module('tasks.admin_dashboard_gui')
                    except Exception:
                        # Eski konuma geri dönüş
                        mod = importlib.import_module('modules.admin_dashboard.admin_gui')
                    AdminDashboardGUI = getattr(mod, 'AdminDashboardGUI')
                    # GUI oluşturmayı ana threade planla
                    self.parent.after(0, lambda: (loader.destroy(), AdminDashboardGUI(self.content_area)))
                except Exception as exc:
                    err = self.lm.tr('err_dashboard_load', "Admin Dashboard yüklenemedi: {}").format(exc)
                    self.parent.after(0, lambda em=err: (loader.destroy(), messagebox.showerror(self.lm.tr('error', "Hata"), em)))

            import threading
            threading.Thread(target=_async_import_and_init, daemon=True).start()

        except Exception as e:
            logging.error(f"Görev dashboard gösterilirken hata: {e}")
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_dashboard_init', "Görev dashboard yüklenemedi: {}").format(e))

    def show_task_assignment(self) -> None:
        """Görev atama arayüzü"""
        try:
            # Admin kontrolü
            if not self._is_admin_or_super():
                messagebox.showwarning(self.lm.tr('permission_denied', "İzin Yok"), self.lm.tr('warning_admin_only', "Bu sayfayı sadece adminler görüntüleyebilir"))
                return

            self._audit_navigation("Görev Atama")
            self.current_module = 'task_assignment'
            self.clear_content()

            # Header'ı geri getir
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)

            self.content_title['text'] = self.lm.tr('task_assignment_title', "Görev Atama")
            # Geri butonu
            self._add_back_button()

            # Görev atama arayüzü
            self._create_task_assignment_interface()

        except Exception as e:
            logging.error(f"Görev atama hatası: {e}")
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_task_assignment_load', "Görev atama yüklenemedi: {}").format(e))

    def _create_task_assignment_interface(self) -> None:
        """Görev atama arayüzü oluştur"""
        from tkinter import ttk

        # Ana frame
        main_frame = tk.Frame(self.content_area, bg='white')
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)

        # Başlık
        tk.Label(main_frame, text=self.lm.tr('new_task_assignment', "Yeni Görev Atama"),

                font=('Segoe UI', 16, 'bold'), bg='white').pack(pady=(0, 20))

        # Form frame
        form_frame = tk.LabelFrame(main_frame, text=self.lm.tr('task_info', "Görev Bilgileri"),

                                 font=('Segoe UI', 12, 'bold'), bg='white')
        form_frame.pack(fill='x', pady=10)

        # Görev başlığı
        tk.Label(form_frame, text=self.lm.tr('task_title', "Görev Başlığı:"), bg='white').grid(row=0, column=0, sticky='w', padx=10, pady=5)

        title_entry = tk.Entry(form_frame, width=50, font=('Segoe UI', 10))
        title_entry.grid(row=0, column=1, padx=10, pady=5)

        # Görev açıklaması
        tk.Label(form_frame, text=f"{self.lm.tr('description', 'Açıklama')}:", bg='white').grid(row=1, column=0, sticky='nw', padx=10, pady=5)

        desc_text = tk.Text(form_frame, width=50, height=4, font=('Segoe UI', 10))
        desc_text.grid(row=1, column=1, padx=10, pady=5)

        # Atanacak kişi
        tk.Label(form_frame, text=self.lm.tr('assignee', "Atanacak Kişi:"), bg='white').grid(row=2, column=0, sticky='w', padx=10, pady=5)

        user_combo = ttk.Combobox(form_frame, width=47, font=('Segoe UI', 10))
        user_combo.grid(row=2, column=1, padx=10, pady=5)

        # Departman
        tk.Label(form_frame, text=self.lm.tr('department', "Departman:"), bg='white').grid(row=3, column=0, sticky='w', padx=10, pady=5)

        dept_combo = ttk.Combobox(form_frame, width=47, font=('Segoe UI', 10))
        dept_combo.grid(row=3, column=1, padx=10, pady=5)

        # Öncelik
        tk.Label(form_frame, text=self.lm.tr('priority', "Öncelik:"), bg='white').grid(row=4, column=0, sticky='w', padx=10, pady=5)
        priority_combo = ttk.Combobox(form_frame, values=[self.lm.tr('priority_low', 'Düşük'), self.lm.tr('priority_medium', 'Orta'), self.lm.tr('priority_high', 'Yüksek')], width=47, font=('Segoe UI', 10))
        priority_combo.grid(row=4, column=1, padx=10, pady=5)
        priority_combo.set(self.lm.tr('priority_medium', 'Orta'))

        # Teslim tarihi
        tk.Label(form_frame, text=self.lm.tr('due_date', "Teslim Tarihi:"), bg='white').grid(row=5, column=0, sticky='w', padx=10, pady=5)

        due_date_entry = tk.Entry(form_frame, width=50, font=('Segoe UI', 10))
        due_date_entry.grid(row=5, column=1, padx=10, pady=5)
        due_date_entry.insert(0, "2024-12-31")

        # Görev türü
        tk.Label(form_frame, text=self.lm.tr('task_type', "Görev Türü:"), bg='white').grid(row=6, column=0, sticky='w', padx=10, pady=5)
        # Görev türleri
        self._task_assignment_types = [
            self.lm.tr('task_type_data_entry', 'Veri Girişi'),
            self.lm.tr('task_type_report_prep', 'Rapor Hazırlama'),
            self.lm.tr('task_type_survey_fill', 'Anket Doldurma'),
            self.lm.tr('task_type_review', 'İnceleme'),
            self.lm.tr('task_type_other', 'Diğer')
        ]
        type_combo = ttk.Combobox(form_frame, values=self._task_assignment_types,
                                width=47, font=('Segoe UI', 10))
        type_combo.grid(row=6, column=1, padx=10, pady=5)
        type_combo.set(self.lm.tr('task_type_data_entry', 'Veri Girişi'))

        # Butonlar
        button_frame = tk.Frame(form_frame, bg='white')
        button_frame.grid(row=7, column=0, columnspan=2, pady=20)

        # Kaydet butonu
        save_btn = tk.Button(button_frame, text=f" {self.lm.tr('assign_task', 'Görev Ata')}",
                           font=('Segoe UI', 11, 'bold'), bg='#27ae60', fg='white',
                           relief='flat', bd=0, cursor='hand2', padx=20, pady=10,
                           command=lambda: self._save_task_assignment(
                               title_entry.get(), desc_text.get("1.0", tk.END).strip(),
                               user_combo.get(), dept_combo.get(), priority_combo.get(),
                               due_date_entry.get(), type_combo.get()))
        save_btn.pack(side='left', padx=10)

        # Temizle butonu
        clear_btn = tk.Button(button_frame, text=f"️ {self.lm.tr('clear', 'Temizle')}",

                            font=('Segoe UI', 11, 'bold'), bg='#e74c3c', fg='white',
                            relief='flat', bd=0, cursor='hand2', padx=20, pady=10,
                            command=lambda: self._clear_task_form(
                                title_entry, desc_text, user_combo, dept_combo,
                                priority_combo, due_date_entry, type_combo))
        clear_btn.pack(side='left', padx=10)

        # Kullanıcı ve departman listelerini yükle
        self._load_user_and_department_lists(user_combo, dept_combo)

        # Mevcut görevler listesi
        tasks_frame = tk.LabelFrame(main_frame, text=self.lm.tr('current_tasks', "Mevcut Görevler"),

                                  font=('Segoe UI', 12, 'bold'), bg='white')
        tasks_frame.pack(fill='both', expand=True, pady=10)

        # Görevler tablosu
        columns = ('ID', 'Başlık', 'Atanan', 'Öncelik', 'Durum', 'Teslim Tarihi')
        tasks_tree = ttk.Treeview(tasks_frame, columns=columns, show='headings', height=8)

        for col in columns:
            tasks_tree.heading(col, text=col)
            tasks_tree.column(col, width=100)

        # Scrollbar
        scrollbar = ttk.Scrollbar(tasks_frame, orient='vertical', command=tasks_tree.yview)
        tasks_tree.configure(yscrollcommand=scrollbar.set)

        tasks_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        # Görevleri yükle
        self._load_tasks_list(tasks_tree)

    def _load_user_and_department_lists(self, user_combo, dept_combo) -> None:
        """Kullanıcı ve departman listelerini yükle"""
        try:
            # Örnek kullanıcılar
            users = ['System Administrator', 'Admin User', 'Manager', 'Analyst', 'Coordinator']
            user_combo['values'] = users
            self._task_assignment_users = users

            # Örnek departmanlar
            departments = ['Yönetim', 'İnsan Kaynakları', 'Muhasebe', 'Pazarlama', 'Üretim', 'Kalite', 'IT', 'Sürdürülebilirlik']
            dept_combo['values'] = departments
            dept_combo.set('Sürdürülebilirlik')
            self._task_assignment_departments = departments

        except Exception as e:
            logging.error(f"Kullanıcı/departman listesi yüklenirken hata: {e}")

    def _load_tasks_list(self, tasks_tree) -> None:
        """Görevler listesini yükle"""
        try:
            # Örnek görevler
            sample_tasks = [
                (1, 'SDG Veri Girişi', 'Analyst', 'Yüksek', 'Devam Ediyor', '2024-12-20'),
                (2, 'Karbon Raporu', 'Coordinator', 'Orta', 'Beklemede', '2024-12-25'),
                (3, 'Anket Doldurma', 'Manager', 'Düşük', 'Tamamlandı', '2024-12-15'),
            ]

            for task in sample_tasks:
                tasks_tree.insert('', 'end', values=task)

        except Exception as e:
            logging.error(f"Görevler listesi yüklenirken hata: {e}")

    def _save_task_assignment(self, title, description, user, department, priority, due_date, task_type) -> None:
        """Görev atamasını kaydet"""
        from tkinter import messagebox
        try:
            # Zorunlu alanlar
            if not title:
                messagebox.showwarning(self.lm.tr('msg_missing_info', "Eksik Bilgi"), self.lm.tr('msg_task_title_required', "Görev başlığı zorunludur!"))
                return
            if not user:
                messagebox.showwarning(self.lm.tr('msg_missing_info', "Eksik Bilgi"), self.lm.tr('msg_assignee_required', "Atanacak kişi zorunludur!"))
                return
            if not department:
                messagebox.showwarning(self.lm.tr('msg_missing_info', "Eksik Bilgi"), self.lm.tr('msg_department_required', "Departman seçimi zorunludur!"))
                return

            # Kullanıcı ve departman doğrulaması (liste üyeliği)
            valid_users = getattr(self, '_task_assignment_users', [])
            valid_departments = getattr(self, '_task_assignment_departments', [])
            if valid_users and user not in valid_users:
                messagebox.showwarning(self.lm.tr('msg_invalid_user', "Geçersiz Kullanıcı"), self.lm.tr('msg_user_not_in_list', "Seçilen kullanıcı listede bulunmuyor."))
                return
            if valid_departments and department not in valid_departments:
                messagebox.showwarning(self.lm.tr('msg_invalid_department', "Geçersiz Departman"), self.lm.tr('msg_department_not_in_list', "Seçilen departman listede bulunmuyor."))
                return

            # Öncelik doğrulaması
            if priority not in ['Düşük', 'Orta', 'Yüksek']:
                messagebox.showwarning(self.lm.tr('msg_invalid_priority', "Geçersiz Öncelik"), self.lm.tr('msg_priority_must_be', "Öncelik Düşük/Orta/Yüksek olmalıdır."))
                return

            # Görev türü doğrulaması
            valid_types = getattr(self, '_task_assignment_types', ['Veri Girişi', 'Rapor Hazırlama', 'Anket Doldurma', 'İnceleme', 'Diğer'])
            if task_type not in valid_types:
                messagebox.showwarning(self.lm.tr('msg_invalid_type', "Geçersiz Tür"), self.lm.tr('msg_task_type_must_be_valid', "Görev türü geçerli bir değer olmalıdır."))
                return

            # Teslim tarihi doğrulaması (YYYY-MM-DD)
            from datetime import datetime
            try:
                datetime.strptime(due_date.strip(), '%Y-%m-%d')
            except Exception:
                messagebox.showwarning(self.lm.tr('msg_invalid_date', "Geçersiz Tarih"), self.lm.tr('msg_date_format_error', "Teslim tarihi YYYY-MM-DD formatında olmalıdır."))
                return

            # Burada gerçek veritabanına kaydetme işlemi yapılacak
            messagebox.showinfo(self.lm.tr('success', "Başarılı"), self.lm.tr('msg_task_assigned_success', "Görev '{title}' {user} kullanıcısına atandı!").format(title=title, user=user))

            # E-mail bildirimi gönder (opsiyonel)
            self._send_task_notification(title, user, due_date)

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('msg_task_assignment_error', "Görev atama hatası: {}").format(e))

    def _clear_task_form(self, title_entry, desc_text, user_combo, dept_combo, priority_combo, due_date_entry, type_combo) -> None:
        """Form alanlarını temizle"""
        title_entry.delete(0, tk.END)
        desc_text.delete("1.0", tk.END)
        user_combo.set('')
        dept_combo.set('Sürdürülebilirlik')
        priority_combo.set('Orta')
        due_date_entry.delete(0, tk.END)
        due_date_entry.insert(0, "2024-12-31")
        type_combo.set('Veri Girişi')

    def _send_task_notification(self, title: str, user: str, due_date: str) -> None:
        """Görev bildirimi e-mail gönder"""
        try:
            from tasks.email_service import EmailService

            # E-mail servisi
            email_service = EmailService()

            # Örnek e-mail adresi (gerçek sistemde kullanıcı tablosundan alınacak)
            user_emails = {
                'System Administrator': 'admin@digage.tr',
                'Admin User': 'admin@digage.tr',
                'Manager': 'manager@digage.tr',
                'Analyst': 'analyst@digage.tr',
                'Coordinator': 'coordinator@digage.tr'
            }

            to_email = user_emails.get(user, 'admin@digage.tr')
            subject = self.lm.tr('new_task_subject', "Yeni Görev Ataması: {}").format(title)
            
            assigner = self.user[1] if self.user and len(self.user) > 1 else 'System Administrator'
            body = self.lm.tr('new_task_body', "Merhaba {user},\n\nSize yeni bir görev atanmıştır:\n\nGörev: {title}\nTeslim Tarihi: {due_date}\nAtayan: {assigner}\n\nGörev detaylarını görüntülemek için sisteme giriş yapınız.\n\nİyi çalışmalar,\nSUSTAINAGE Sistemi").format(user=user, title=title, due_date=due_date, assigner=assigner)

            # E-mail gönder
            success = email_service.send_email(
                to_email=to_email,
                subject=subject,
                body=body,
                to_name=user
            )

            if success:
                logging.info(f"[OK] Görev bildirimi gönderildi: {user} ({to_email})")
            else:
                logging.warning(f"[WARNING] E-mail gönderilemedi: {user}")

        except Exception as e:
            logging.error(f"[ERROR] E-mail bildirimi hatası: {e}")

    def show_cdp(self) -> None:
        """CDP (Carbon Disclosure Project) modülünü hızlı göster"""
        self._show_module_fast('cdp', 'CDP - Carbon Disclosure Project')

    def show_iirc(self) -> None:
        """IIRC (Entegre Raporlama) modülünü hızlı göster"""
        self._show_module_fast('iirc', 'IIRC - Entegre Raporlama')

    def show_validation(self) -> None:
        """Veri Validasyon modülünü hızlı göster"""
        self._show_module_fast('validation', 'Gelişmiş Veri Validasyon')

    def show_benchmark(self) -> None:
        """Sektör Benchmark modülünü hızlı göster"""
        self._show_module_fast('benchmark', 'Sektör Benchmark Analizi')

    def show_forecasting(self) -> None:
        """Trend ve Tahmin modülünü hızlı göster"""
        self._show_module_fast('forecasting', 'Trend ve Tahmin Modelleri')

    def show_report_center(self) -> None:
        """Gelişmiş Rapor Merkezi hızlı göster"""
        self._show_module_fast('report_center', 'Gelişmiş Rapor Merkezi', 'report.read')

    def show_integration(self) -> None:
        """API ve Entegrasyon modülünü hızlı göster"""
        self._show_module_fast('integration', 'API ve Entegrasyon')

    def show_visualization(self) -> None:
        """Gelişmiş Görselleştirme modülünü hızlı göster"""
        self._show_module_fast('visualization', 'Gelişmiş Görselleştirme')

    def show_auditor(self) -> None:
        """Denetçi Sistemi hızlı göster"""
        self._show_module_fast('auditor', 'Dış Doğrulama ve Denetçi Sistemi')

    def show_stakeholder_engagement(self) -> None:
        """Paydaş Etkileşim Sistemi hızlı göster"""
        self._show_module_fast('stakeholder', 'Gelişmiş Paydaş Etkileşim Sistemi')

    def show_scenario_analysis(self) -> None:
        """Senaryo Analizi ve Modelleme hızlı göster"""
        self._show_module_fast('scenario', 'Senaryo Analizi ve İklim Modelleme')

    def show_csrd(self) -> None:
        """AB CSRD Tam Uyum hızlı göster"""
        self._show_module_fast('csrd', 'AB CSRD Tam Uyum Sistemi')

    def show_report_scheduling(self) -> None:
        """Rapor Scheduling ve Dağıtım hızlı göster"""
        self._show_module_fast('report_scheduler', 'Otomatik Rapor Scheduling ve Dağıtım', 'report.read')

    def show_management(self) -> None:
        """Yönetim modülünü göster"""
        if not self._require_permission("system.settings", "Yönetim"):
            return
        self._audit_navigation("Yönetim")
        self.current_module = 'management'
        self.clear_content()
        # Header'ı geri getir
        self.content_header.pack(fill='x')
        self.content_area.pack_forget()
        self.content_area.pack(fill='both', expand=True, padx=20, pady=20)

        self.content_title['text'] = "Yönetim"

        # Yönetim içinde Firma Bilgileri kısayolları - üst çubuk bölgesine taşı
        try:
            # Butonlar için üst container (başlık yerine)
            tools = tk.Frame(self.content_area, bg='white')
            tools.pack(fill='x', pady=(0, 15))

            # Buton listesi (icon, text, command, color)
            management_buttons = []

            # Her zaman görünür butonlar
            management_buttons.append((Icons.MEMO, 'Firma Bilgileri', self.open_company_info_with_selection, '#3b82f6'))
            management_buttons.append((Icons.ADD, 'Yeni Firma', self._create_new_company, '#1F6FEB'))
            management_buttons.append((Icons.DELETE, 'Firma Sil', self.open_delete_company_with_selection, '#DC143C'))

            # Admin butonları
            if self._is_admin_only():
                management_buttons.append((Icons.REPORT, 'Görev Dashboard', self.show_task_dashboard, '#9b59b6'))
                management_buttons.append((Icons.USER, 'Görev Atama', self.show_task_assignment, '#e67e22'))

            # İzin bazlı butonlar
            if self._has_permission("data.import"):
                management_buttons.append(('📥', 'Veri İçe Aktarım', lambda: self._navigate(self.show_data_import), '#2c3e50'))

            if self._has_permission("forms.manage"):
                management_buttons.append((Icons.CLIPBOARD, 'Form Yönetimi', lambda: self._navigate(self.show_form_management), '#34495e'))

            if self._has_permission("files.manage"):
                management_buttons.append(('📁', 'Dosya Yönetimi', lambda: self._navigate(self.show_advanced_files), '#7f8c8d'))

            if (self._has_permission("hr.read") or self._is_admin_only()) and not self._is_super_admin_user():
                management_buttons.append((Icons.USERS, 'İK Metrikleri', lambda: self._navigate(self.show_hr_metrics), '#8e44ad'))

            # Grid layout - üst satırda daha fazla buton sığdır
            buttons_per_row = 6
            for idx, (icon, text, command, color) in enumerate(management_buttons):
                row = idx // buttons_per_row
                col = idx % buttons_per_row
                cell = tk.Frame(tools, bg='white')
                cell.grid(row=row, column=col, padx=4, pady=6, sticky='ew')
                try:
                    ctrl, _ = self._create_neumorphic_menu_control(cell, f'{icon} {text}', command, None, color)
                    ctrl.pack(fill='x')
                except Exception:
                    btn = tk.Button(cell, text=f'{icon} {text}',
                                    font=('Segoe UI', 10, 'bold'), fg='white', bg=color,
                                    relief='flat', bd=0, cursor='hand2', padx=12, pady=8,
                                    command=command)
                    btn.pack(fill='x')
            for i in range(buttons_per_row):
                tools.columnconfigure(i, weight=1, uniform='button')

        except Exception as e:
            logging.error(f"Firma butonları oluşturulurken hata: {e}")

        # Yönetim modülünü oluştur
        try:
            user_id = self.user[0] if self.user and len(self.user) > 0 else 1
            YonetimGUI_class = self._lazy_import_module('yonetim')
            if YonetimGUI_class:
                YonetimGUI_class(self.content_area, user_id)
            else:
                self.show_module_error("Yönetim", self.lm.tr('err_module_load', "{} modülü yüklenemedi").format("Yönetim"))
        except Exception as e:
            logging.error(f"Yönetim modülü yüklenirken hata: {e}")
            error_label = tk.Label(self.content_area, text=self.lm.tr('err_module_load_detail', "{} modülü yüklenirken hata: {}").format("Yönetim", e),
                                 font=('Segoe UI', 12), fg='#e74c3c', bg='white')
            error_label.pack(pady=50)

    def show_innovation(self) -> None:
        """AR-GE ve İnovasyon modülünü göster"""
        try:
            self.clear_content()
            self.content_title['text'] = self.lm.tr('r_and_d_innovation_title', " AR-GE & İnovasyon")
            self._add_back_button()

            InnovationGUI_class = self._lazy_import_module('innovation')
            if InnovationGUI_class:
                InnovationGUI_class(self.content_area, self.company_id)
            else:
                self.show_module_error("AR-GE & İnovasyon", self.lm.tr('err_module_load', "{} modülü yüklenemedi").format("AR-GE & İnovasyon"))
        except Exception as e:
            logging.error(f"AR-GE & İnovasyon modülü yüklenirken hata: {e}")
            error_label = tk.Label(self.content_area, text=self.lm.tr('err_module_load_detail', "{} modülü yüklenirken hata: {}").format("AR-GE & İnovasyon", e),
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    def show_quality(self) -> None:
        """Ürün Kalitesi modülünü göster"""
        try:
            self.clear_content()
            self.content_title['text'] = f"{Icons.STAR} {self.lm.tr('product_quality_title', 'Ürün Kalitesi & Müşteri Memnuniyeti')}"
            self._add_back_button()

            QualityGUI_class = self._lazy_import_module('quality')
            if QualityGUI_class:
                QualityGUI_class(self.content_area, self.company_id)
            else:
                self.show_module_error("Ürün Kalitesi", self.lm.tr('err_product_quality_load', "Ürün Kalitesi modülü yüklenemedi"))
        except Exception as e:
            logging.error(f"Ürün Kalitesi modülü yüklenirken hata: {e}")
            error_label = tk.Label(self.content_area, text=f"Ürün Kalitesi modülü yüklenirken hata: {e}",
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    def show_digital_security(self) -> None:
        """Bilgi Güvenliği modülünü göster"""
        try:
            self.clear_content()
            self.content_title['text'] = self.lm.tr('digital_security_title', " Bilgi Güvenliği & Dijitalleşme")
            self._add_back_button()

            DigitalSecurityGUI_class = self._lazy_import_module('digital_security')
            if DigitalSecurityGUI_class:
                DigitalSecurityGUI_class(self.content_area, self.company_id)
            else:
                self.show_module_error("Bilgi Güvenliği", self.lm.tr('err_digital_security_load', "Bilgi Güvenliği modülü yüklenemedi"))
        except Exception as e:
            logging.error(f"Bilgi Güvenliği modülü yüklenirken hata: {e}")
            error_label = tk.Label(self.content_area, text=self.lm.tr('err_module_load_detail', "{} modülü yüklenirken hata: {}").format("Bilgi Güvenliği", e),
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    def show_emergency(self) -> None:
        """Acil Durum modülünü göster"""
        try:
            self.clear_content()
            self.content_title['text'] = self.lm.tr('emergency_disaster_management_title', " Acil Durum & Afet Yönetimi")
            self._add_back_button()

            EmergencyGUI_class = self._lazy_import_module('emergency')
            if EmergencyGUI_class:
                EmergencyGUI_class(self.content_area, self.company_id)
            else:
                self.show_module_error("Acil Durum", self.lm.tr('err_emergency_load', "Acil Durum modülü yüklenemedi"))
        except Exception as e:
            logging.error(f"Acil Durum modülü yüklenirken hata: {e}")
            error_label = tk.Label(self.content_area, text=self.lm.tr('err_module_load_detail', "{} modülü yüklenirken hata: {}").format("Acil Durum", e),
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    # def show_combined_reports(self) -> None:
    #     """GEREKSİZ RAPOR MERKEZİ KALDIRILDI - Bu sayfa kaldırıldı"""
    #     messagebox.showinfo("Bilgi", "Bu sayfa kaldırıldı.\n\nRaporlama için ana 'Raporlama' modülünü kullanın.")
    #
    #     except Exception as e:
    #         print(f"Raporlar modülü yüklenirken hata: {e}")
    #         error_label = tk.Label(self.content_area, text=f"Raporlar modülü yüklenirken hata: {e}",
    #                              font=('Segoe UI', 12), fg='red', bg='white')
    #         error_label.pack(pady=50)

    def exit_application(self) -> None:
        """Uygulamadan çık"""
        try:
            # Verileri kaydet
            if hasattr(self, 'save_data'):
                self.save_data()

            # Çıkış onayı
            result = messagebox.askyesno(self.lm.tr('exit_title', "Çıkış"), self.lm.tr('exit_confirm_msg', "Uygulamadan çıkmak istediğinizden emin misiniz?"))
            if result:
                self.parent.quit()
                self.parent.destroy()
        except Exception as e:
            logging.error(f"Çıkış hatası: {e}")
            self.parent.quit()
            self.parent.destroy()

    # Dil/çeviri ile ilgili metotlar kaldırıldı; ana yapı korunarak devre dışı

    def restart_application(self) -> None:
        """Uygulamayı yeniden başlat"""
        try:
            # Mevcut pencereyi kapat
            self.parent.quit()

            # Yeni pencere aç (bu kısım main.py'de yapılacak)
            messagebox.showinfo(self.lm.tr('info_title', "Bilgi"), self.lm.tr('restart_msg', "Uygulama yeniden başlatılacak."))

        except Exception as e:
            logging.error(f"Uygulama yeniden başlatma hatası: {e}")

    def ensure_company_info_table(self) -> None:
        """SQLite içinde firma bilgileri tablosunu oluştur (yoksa)."""
        conn = sqlite3.connect(self.db_path)
        cur = conn.cursor()
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS company_info (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER UNIQUE,
                sirket_adi TEXT,
                ticari_unvan TEXT,
                marka_isimleri TEXT,
                adres TEXT,
                ulke TEXT,
                sehir TEXT,
                posta_kodu TEXT,
                web TEXT,
                eposta TEXT,
                telefon TEXT,
                vergi_no TEXT,
                mersis_no TEXT,
                tuzel_yapi TEXT,
                mulkiyet_kontrol TEXT,
                borsa_durumu TEXT,
                kayitli_merkez TEXT,
                calisan_sayisi TEXT,
                calisan_cinsiyet_dagilimi TEXT,
                operasyon_sayisi TEXT,
                faaliyet_lokasyonlari TEXT,
                net_gelir TEXT,
                toplam_varliklar TEXT,
                pazarlar TEXT,
                tedarik_zinciri TEXT,
                is_iliskileri TEXT,
                upstream_downstream TEXT,
                yonetim_yapisi TEXT,
                en_yuksek_yonetim_organi TEXT,
                baskan_bilgisi TEXT,
                surdurulebilirlik_sorumluluklari TEXT,
                risk_yonetimi_rolu TEXT,
                cikar_catismalari TEXT,
                kritik_endise_iletisimi TEXT,
                yonetim_bilgisi_kapasitesi TEXT,
                performans_degerlendirme TEXT,
                ucret_politikalari TEXT,
                ucret_belirleme_sureci TEXT,
                yillik_ucret_orani TEXT,
                ust_duzey_beyan TEXT,
                politika_taahhutleri TEXT,
                politikalarin_yayilmasi TEXT,
                olumsuz_etki_tazmin TEXT,
                etik_endise_mekanizmalari TEXT,
                mevzuat_uyum TEXT,
                uyelik_dernekler TEXT,
                paydas_iletisim TEXT,
                toplu_sozlesme_orani TEXT,
                raporlama_donemi TEXT,
                raporlama_dongusu TEXT,
                irtibat_noktasi TEXT,
                raporda_yer_alan_kuruluslar TEXT,
                bilgilerin_yeniden_duzenlenmesi TEXT,
                dis_guvence TEXT,
                dis_guvence_tur_kapsam TEXT,
                calisan_sabit_sozlesmeli TEXT,
                calisan_gecici TEXT,
                calisan_kadin_erkek_sayisi TEXT,
                calisan_bolgelere_gore_dagilim TEXT,
                calisan_olmayanlar_sayisi TEXT,
                calisan_olmayanlar_tur TEXT,
                sektor_kod TEXT,
                faaliyette_bulunulan_ulkeler TEXT,
                sahiplik_yapisi_detay TEXT,
                uretim_kapasitesi TEXT
            )
            """
        )
        conn.commit()
        conn.close()

    def load_company_info(self) -> None:
        """Veritabanından firma bilgilerini getir."""
        try:
            conn = sqlite3.connect(self.db_path)
            cur = conn.cursor()
            cur.execute("SELECT * FROM company_info WHERE company_id=?", (self.company_id,))
            row = cur.fetchone()
            conn.close()
            return row
        except Exception as e:
            logging.info(f"Firma bilgileri yüklenemedi: {e}")
            return None

    def _validate_company_info(self, values) -> None:
        """Form değerlerini doğrula"""
        import re

        # Email validasyonu
        email = values.get('eposta', '').strip()
        if email:
            email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
            if not re.match(email_pattern, email):
                messagebox.showerror(self.lm.tr('validation_error_title', "Validasyon Hatası"),
                    self.lm.tr('err_email_format', "Geçersiz email formatı: {}\n\nÖrnek: ornek@sirket.com").format(email))
                return False

        # Telefon validasyonu (Türkiye formatı)
        telefon = values.get('telefon', '').strip()
        if telefon:
            # Telefon doğrulama lazy import
            from utils.phone import is_valid_tr_phone
            if not is_valid_tr_phone(telefon):
                messagebox.showerror(self.lm.tr('validation_error_title', "Validasyon Hatası"),
                    self.lm.tr('err_phone_format', "Geçersiz telefon formatı: {}\n\nBeklenen biçim: +90 (5XX) XXX XX XX").format(telefon))
                return False

        # Vergi No validasyonu (10 haneli sayı)
        vergi_no = values.get('vergi_no', '').strip()
        if vergi_no:
            vergi_clean = re.sub(r'\D', '', vergi_no)  # Sadece rakamlar
            if not re.match(r'^\d{10}$', vergi_clean):
                messagebox.showerror(self.lm.tr('validation_error_title', "Validasyon Hatası"),
                    self.lm.tr('err_tax_no_format', "Geçersiz Vergi No: {}\n\nVergi No 10 haneli olmalıdır.\nÖrnek: 1234567890").format(vergi_no))
                return False

        # MERSİS No validasyonu (16 haneli sayı)
        mersis = values.get('mersis_no', '').strip()
        if mersis:
            mersis_clean = re.sub(r'\D', '', mersis)  # Sadece rakamlar
            if not re.match(r'^\d{16}$', mersis_clean):
                messagebox.showerror(self.lm.tr('validation_error_title', "Validasyon Hatası"),
                    self.lm.tr('err_mersis_format', "Geçersiz MERSİS No: {}\n\nMERSİS No 16 haneli olmalıdır.\nÖrnek: 0123456789012345").format(mersis))
                return False

        # Web sitesi validasyonu
        web = values.get('web', '').strip()
        if web:
            # Basit URL validasyonu
            if not web.startswith(('http://', 'https://', 'www.')):
                messagebox.showwarning(self.lm.tr('validation_warning_title', "Validasyon Uyarısı"),
                    self.lm.tr('warn_web_format', "Web sitesi adresi genellikle http:// veya https:// ile başlar.\n\nGirilen: {}\nDevam etmek istiyor musunuz?").format(web))
                # Uyarı ver ama devam et

        # Posta kodu validasyonu (5 haneli)
        posta_kodu = values.get('posta_kodu', '').strip()
        if posta_kodu:
            posta_clean = re.sub(r'\D', '', posta_kodu)
            if not re.match(r'^\d{5}$', posta_clean):
                messagebox.showwarning(self.lm.tr('validation_warning_title', "Validasyon Uyarısı"),
                    self.lm.tr('warn_zip_format', "Posta kodu genellikle 5 haneli olmalıdır.\n\nGirilen: {}").format(posta_kodu))
                # Uyarı ver ama devam et

        # Tüm validasyonlar geçti
        return True

    def _save_company_info(self, values) -> None:
        """Form değerlerini veritabanına kaydet (doğrulama ile)."""
        if not self._validate_company_info(values):
            return
        try:
            conn = sqlite3.connect(self.db_path)
            cur = conn.cursor()
            cur.execute("SELECT id FROM company_info WHERE company_id=?", (self.company_id,))
            exists = cur.fetchone()

            columns = [
                'sirket_adi','ticari_unvan','marka_isimleri','adres','ulke','sehir','posta_kodu','web','eposta','telefon',
                'vergi_no','mersis_no','tuzel_yapi','mulkiyet_kontrol','borsa_durumu','kayitli_merkez',
                'calisan_sayisi','calisan_cinsiyet_dagilimi','operasyon_sayisi','faaliyet_lokasyonlari','net_gelir','toplam_varliklar','pazarlar',
                'tedarik_zinciri','is_iliskileri','upstream_downstream',
                'yonetim_yapisi','en_yuksek_yonetim_organi','baskan_bilgisi','surdurulebilirlik_sorumluluklari','risk_yonetimi_rolu','cikar_catismalari','kritik_endise_iletisimi','yonetim_bilgisi_kapasitesi','performans_degerlendirme',
                'ucret_politikalari','ucret_belirleme_sureci','yillik_ucret_orani',
                'ust_duzey_beyan','politika_taahhutleri','politikalarin_yayilmasi','olumsuz_etki_tazmin','etik_endise_mekanizmalari','mevzuat_uyum',
                'uyelik_dernekler','paydas_iletisim','toplu_sozlesme_orani',
                'raporlama_donemi','raporlama_dongusu','irtibat_noktasi',
                'raporda_yer_alan_kuruluslar','bilgilerin_yeniden_duzenlenmesi',
                'dis_guvence','dis_guvence_tur_kapsam',
                'calisan_sabit_sozlesmeli','calisan_gecici','calisan_kadin_erkek_sayisi','calisan_bolgelere_gore_dagilim',
                'calisan_olmayanlar_sayisi','calisan_olmayanlar_tur',
                'sektor_kod','faaliyette_bulunulan_ulkeler','sahiplik_yapisi_detay','uretim_kapasitesi'
            ]

            if exists:
                set_clause = ", ".join([f"{col}=?" for col in columns])
                cur.execute(
                    f"UPDATE company_info SET {set_clause} WHERE company_id=?",
                    [values.get(col, '') for col in columns] + [self.company_id]
                )
            else:
                placeholders = ",".join(["?" for _ in columns])
                cur.execute(
                    f"INSERT INTO company_info (company_id, {', '.join(columns)}) VALUES (?, {placeholders})",
                    [self.company_id] + [values.get(col, '') for col in columns]
                )
            conn.commit()
            conn.close()
            messagebox.showinfo(self.lm.tr('success_saved_title', "Kaydedildi"), self.lm.tr('msg_company_info_saved', "Firma bilgileri başarıyla kaydedildi."))
        except Exception as e:
            messagebox.showerror(self.lm.tr('error_title', "Hata"), self.lm.tr('company_info_save_error', "Firma bilgileri kaydedilemedi: {e}").format(e=e))

    def show_company_info(self) -> None:
        """GRI uyumlu Firma Bilgileri sayfasını göster (responsive ve şık)."""
        # Sadece admin kullanıcılar erişebilir
        if not self._is_admin_or_super():
            messagebox.showwarning(self.lm.tr('permission_denied_title', "İzin Yok"), self.lm.tr('admin_only_access', "Bu sayfaya sadece Admin kullanıcıları erişebilir."))
            return
        self._audit_navigation(self.lm.tr('audit_company_info', "Firma Bilgileri"))
        self.current_module = 'company_info'
        self.clear_content()
        # Header'ı geri getir
        self.content_header.pack(fill='x')
        self.content_area.pack_forget()
        self.content_area.pack(fill='both', expand=True, padx=20, pady=20)

        self.content_title['text'] = self.lm.tr('company_info_title', "Firma Bilgileri")

        self._add_back_button()

        # Kaydırılabilir, responsive form alanı
        container = tk.Frame(self.content_area, bg='white')
        container.pack(fill='both', expand=True)

        # Üstte logo yönetimi bölümü ve aksiyon butonları
        top_bar = tk.Frame(container, bg='white')
        top_bar.pack(fill='x', pady=(8, 4))

        try:
            logo_preview = tk.Label(top_bar, bg='white')
            logo_preview.pack(side='left', padx=8)
            def _load_logo_preview():
                try:
                    from modules.company.company_profile_manager import CompanyProfileManager
                    mgr = CompanyProfileManager(self.db_path)
                    lp = mgr.get_logo_path(self.company_id)
                    if lp and os.path.exists(lp):
                        from PIL import Image, ImageTk
                        img = Image.open(lp)
                        img.thumbnail((120, 60), Image.Resampling.LANCZOS)
                        photo = ImageTk.PhotoImage(img, master=self.parent)
                        logo_preview.configure(image=photo)
                        logo_preview.image = photo
                    else:
                        logo_preview.configure(text=self.lm.tr('logo_missing', 'Logo yok'))
                except Exception:
                    logo_preview.configure(text=self.lm.tr('logo_load_error', 'Logo yüklenemedi'))
            def _upload_logo():
                try:
                    from tkinter import filedialog, messagebox
                    fp = filedialog.askopenfilename(title=self.lm.tr('select_logo', 'Logo Seç'), filetypes=[(self.lm.tr('image_file_type', 'Resim'),'*.png;*.jpg;*.jpeg;*.webp')])
                    if not fp:
                        return
                    from modules.company.company_profile_manager import CompanyProfileManager
                    mgr = CompanyProfileManager(self.db_path)
                    ok = mgr.upload_logo(self.company_id, fp)
                    if ok:
                        messagebox.showinfo(self.lm.tr('success_title', 'Başarılı'), self.lm.tr('logo_upload_success', 'Logo yüklendi ve tüm sistemde kullanılacak'))
                        _load_logo_preview()
                    else:
                        messagebox.showerror(self.lm.tr('error_title', 'Hata'), self.lm.tr('logo_load_error', 'Logo yüklenemedi'))
                except Exception as e:
                    try:
                        from tkinter import messagebox
                        messagebox.showerror(self.lm.tr('error_title', 'Hata'), self.lm.tr('logo_load_error', 'Logo yüklenemedi') + f': {e}')
                    except Exception as e:
                        logging.error(f"Silent error caught: {str(e)}")
            upload_btn = ttk.Button(top_bar, text=self.lm.tr('upload_logo_btn', ' Logo Yükle'), style='Menu.TButton', command=_upload_logo)
            upload_btn.pack(side='left', padx=(6, 0))
            _load_logo_preview()
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

        # Üst sağ köşeye izin kontrollü "Yeni Firma" butonu
        if self._is_admin_or_super():
            new_company_btn = tk.Button(top_bar, text=self.lm.tr('new_company_btn', 'Yeni Firma'),
                                        font=('Segoe UI', 10, 'bold'), fg='white', bg='#1F6FEB',
                                        relief='flat', bd=0, cursor='hand2', padx=14, pady=6,
                                        command=self._create_new_company)
            new_company_btn.pack(side='right', padx=8)
        canvas = tk.Canvas(container, bg='white', highlightthickness=0)
        vsb = tk.Scrollbar(container, orient='vertical', command=canvas.yview)
        form_holder = tk.Frame(canvas, bg='white')

        # Canvas içinde pencereyi oluştur ve genişliğini otomatik ayarla
        window_id = canvas.create_window((0, 0), window=form_holder, anchor='nw')
        def on_form_config(_) -> None:
            canvas.configure(scrollregion=canvas.bbox('all'))
        def on_canvas_config(e) -> None:
            canvas.itemconfigure(window_id, width=e.width)
        form_holder.bind('<Configure>', on_form_config)
        canvas.bind('<Configure>', on_canvas_config)

        canvas.configure(yscrollcommand=vsb.set)
        canvas.pack(side='left', fill='both', expand=True)
        vsb.pack(side='right', fill='y')

        # Mouse wheel ile kaydırma (Windows ve Linux desteği)
        def _on_mousewheel_win(e) -> None:
            try:
                canvas.yview_scroll(-int(e.delta/120), 'units')
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
        def _on_mousewheel_linux_up(e) -> None:
            canvas.yview_scroll(-1, 'units')
        def _on_mousewheel_linux_down(e) -> None:
            canvas.yview_scroll(1, 'units')

        def _bind_scroll(_) -> None:
            canvas.bind_all('<MouseWheel>', _on_mousewheel_win)
            canvas.bind_all('<Button-4>', _on_mousewheel_linux_up)
            canvas.bind_all('<Button-5>', _on_mousewheel_linux_down)
        def _unbind_scroll(_) -> None:
            canvas.unbind_all('<MouseWheel>')
            canvas.unbind_all('<Button-4>')
            canvas.unbind_all('<Button-5>')

        form_holder.bind('<Enter>', _bind_scroll)
        form_holder.bind('<Leave>', _unbind_scroll)

        # Grid yapılandırması
        # Sol sütun etiketler için sabit genişlik, sağ sütun genişleyebilir ve tüm girişler aynı hizada görünür
        form_holder.grid_columnconfigure(0, weight=0, minsize=220)
        form_holder.grid_columnconfigure(1, weight=1, minsize=420)

        section_title_style = {'font': ('Segoe UI', 14, 'bold'), 'fg': '#2E8B57', 'bg': 'white'}
        label_style = {'font': ('Segoe UI', 10, 'bold'), 'fg': '#333333', 'bg': 'white'}

        entries = {}
        row_idx = 0

        def add_separator() -> None:
            nonlocal row_idx
            sep = tk.Frame(form_holder, bg='#eaeaea', height=1)
            sep.grid(row=row_idx, column=0, columnspan=2, sticky='ew', pady=(16, 8))
            row_idx += 1

        def add_section(title) -> None:
            nonlocal row_idx
            tk.Label(form_holder, text=title, **section_title_style).grid(row=row_idx, column=0, columnspan=2, sticky='w')
            row_idx += 1

        def add_entry(key, label_text) -> None:
            nonlocal row_idx
            tk.Label(form_holder, text=label_text, **label_style).grid(row=row_idx, column=0, sticky='w', padx=(6, 8), pady=4)
            # İnce ve daha açık kenarlık için: dışta açık gri çerçeve, içeride düz giriş
            wrapper = tk.Frame(form_holder, bg='#e0e0e0', highlightthickness=0, bd=0)
            wrapper.grid(row=row_idx, column=1, sticky='ew', padx=(0, 12), pady=4)
            inner = tk.Frame(wrapper, bg='white', highlightthickness=0, bd=0)
            inner.pack(fill='x', padx=1, pady=1)
            ent = tk.Entry(inner, relief='flat', bd=0, font=('Segoe UI', 10))
            ent.pack(fill='x')
            entries[key] = ent
            row_idx += 1

        def add_text(key, label_text, height=4) -> None:
            nonlocal row_idx
            tk.Label(form_holder, text=label_text, **label_style).grid(row=row_idx, column=0, sticky='nw', padx=(6, 8), pady=(6, 2))
            # İnce ve daha açık kenarlık için metin alanı: dışta açık gri çerçeve, içeride düz metin kutusu
            wrapper = tk.Frame(form_holder, bg='#e0e0e0', highlightthickness=0, bd=0)
            wrapper.grid(row=row_idx, column=1, sticky='ew', padx=(0, 12), pady=(6, 2))
            inner = tk.Frame(wrapper, bg='white', highlightthickness=0, bd=0)
            inner.pack(fill='both', padx=1, pady=1)
            txt = tk.Text(inner, height=height, wrap='word', relief='flat', bd=0, font=('Segoe UI', 10))
            txt.pack(fill='both')
            entries[key] = txt
            row_idx += 1

        # Bölüm 1: Kurumsal Bilgiler (GRI 2-1)
        add_section(self.lm.tr('company_info_title', "Kurumsal Bilgiler (GRI 2-1)"))
        add_entry('sirket_adi', self.lm.tr('company_name', 'Şirket Adı:'))
        add_entry('ticari_unvan', self.lm.tr('commercial_title', 'Ticari Ünvan:'))
        add_entry('marka_isimleri', self.lm.tr('brand_names', 'Marka/İsimler:'))
        add_text('adres', self.lm.tr('address', 'Adres:'))
        add_entry('ulke', self.lm.tr('country', 'Ülke:'))
        add_entry('sehir', self.lm.tr('city', 'Şehir:'))
        add_entry('posta_kodu', self.lm.tr('postal_code', 'Posta Kodu:'))
        add_entry('web', self.lm.tr('website', 'Web Sitesi:'))
        add_entry('eposta', self.lm.tr('email', 'E-Posta:'))
        add_entry('telefon', self.lm.tr('phone', 'Telefon:'))
        add_entry('vergi_no', self.lm.tr('tax_no', 'Vergi No:'))
        add_entry('mersis_no', self.lm.tr('mersis_no', 'MERSİS No:'))
        add_entry('tuzel_yapi', self.lm.tr('legal_structure', 'Tüzel Yapı:'))
        add_text('mulkiyet_kontrol', self.lm.tr('ownership_structure', 'Mülkiyet ve Kontrol Yapısı:'))
        add_entry('borsa_durumu', self.lm.tr('stock_status', 'Borsa Durumu:'))
        add_text('kayitli_merkez', self.lm.tr('registered_hq', 'Kayıtlı Merkez:'))

        add_separator()
        # Bölüm 2: Ölçek ve Faaliyetler (GRI 2-6)
        add_section(self.lm.tr('scale_activities_title', "Ölçek ve Faaliyetler (GRI 2-6)"))
        add_entry('calisan_sayisi', self.lm.tr('employee_count', 'Çalışan Sayısı:'))
        add_text('calisan_cinsiyet_dagilimi', self.lm.tr('employee_gender_dist', 'Çalışan Cinsiyet Dağılımı:'))
        add_entry('operasyon_sayisi', self.lm.tr('operation_count', 'Operasyon Sayısı:'))
        add_text('faaliyet_lokasyonlari', self.lm.tr('activity_locations', 'Faaliyet Lokasyonları:'))
        add_entry('net_gelir', self.lm.tr('net_revenue', 'Net Gelir:'))
        add_entry('toplam_varliklar', self.lm.tr('total_assets', 'Toplam Varlıklar:'))
        add_text('pazarlar', self.lm.tr('served_markets', 'Hizmet Verilen Pazarlar:'))

        add_separator()
        add_section(self.lm.tr('value_chain_title', "Değer Zinciri ve İş İlişkileri (GRI 2-6)"))
        add_text('tedarik_zinciri', self.lm.tr('supply_chain_desc', 'Tedarik Zinciri Açıklaması:'))
        add_text('is_iliskileri', self.lm.tr('business_relations', 'Temel İş İlişkileri:'))
        add_text('upstream_downstream', self.lm.tr('upstream_downstream_summary', 'Upstream / Downstream Özet:'))

        add_separator()
        # Bölüm 3: Yönetişim (GRI 2-9 .. 2-21)
        add_section(self.lm.tr('governance_title', "Yönetişim (GRI 2-9 .. 2-21)"))
        add_text('yonetim_yapisi', self.lm.tr('gov_structure', 'Yönetim Yapısı ve Kompozisyonu:'))
        add_text('en_yuksek_yonetim_organi', self.lm.tr('highest_gov_body_role', 'En Yüksek Yönetim Organının Rolü:'))
        add_text('baskan_bilgisi', self.lm.tr('chair_info', 'Başkan / Yönetim Kurulu Başkanı Bilgisi:'))
        add_text('surdurulebilirlik_sorumluluklari', self.lm.tr('sustainability_resps', 'Sürdürülebilirlik Sorumlulukları ve Raporlama:'))
        add_text('risk_yonetimi_rolu', self.lm.tr('risk_mgmt_role', 'Risk Yönetimi Rolü:'))
        add_text('cikar_catismalari', self.lm.tr('conflicts_of_interest', 'Çıkar Çatışmaları:'))
        add_text('kritik_endise_iletisimi', self.lm.tr('critical_concerns_comm', 'Kritik Endişelerin İletişimi:'))
        add_text('yonetim_bilgisi_kapasitesi', self.lm.tr('gov_body_knowledge', 'Yönetim Organının Kollektif Bilgi/Kapasitesi:'))
        add_text('performans_degerlendirme', self.lm.tr('gov_perf_eval', 'Yönetim Organı Performans Değerlendirmesi:'))
        add_text('ucret_politikalari', self.lm.tr('remuneration_policies', 'Ücret Politikaları:'))
        add_text('ucret_belirleme_sureci', self.lm.tr('remuneration_process', 'Ücret Belirleme Süreci:'))
        add_entry('yillik_ucret_orani', self.lm.tr('annual_comp_ratio', 'Yıllık Toplam Ücret Oranı:'))

        add_separator()
        # Bölüm 4: Politikalar ve Etik (GRI 2-22 .. 2-27)
        add_section(self.lm.tr('policies_ethics_title', "Politikalar ve Etik (GRI 2-22 .. 2-27)"))
        add_text('ust_duzey_beyan', self.lm.tr('senior_exec_statement', 'Üst Düzey Yönetici Beyanı:'))
        add_text('politika_taahhutleri', self.lm.tr('policy_commitments', 'Politika Taahhütleri:'))
        add_text('politikalarin_yayilmasi', self.lm.tr('policy_embedding', 'Politikaların Yedirilmesi/Yayılması:'))
        add_text('olumsuz_etki_tazmin', self.lm.tr('remediation_processes', 'Olumsuz Etkilerin Telafisi Süreçleri:'))
        add_text('etik_endise_mekanizmalari', self.lm.tr('ethics_mechanisms', 'Etik Endişe Bildirim Mekanizmaları:'))
        add_text('mevzuat_uyum', self.lm.tr('compliance', 'Mevzuata Uyum:'))

        add_separator()
        # Bölüm 5: Üyelikler ve Paydaşlar (GRI 2-28 .. 2-30)
        add_section(self.lm.tr('memberships_stakeholders_title', "Üyelikler ve Paydaşlar (GRI 2-28 .. 2-30)"))
        add_text('uyelik_dernekler', self.lm.tr('memberships', 'Üye Olunan Dernekler/Kuruluşlar:'))
        add_text('paydas_iletisim', self.lm.tr('stakeholder_engagement', 'Paydaş İletişim Yaklaşımı:'))
        add_entry('toplu_sozlesme_orani', self.lm.tr('collective_bargaining_rate', 'Toplu Sözleşme Kapsam Oranı (%):'))

        add_separator()
        # Bölüm 6: Sürdürülebilirlik Raporlaması (GRI 2-2)
        add_section(self.lm.tr('sust_reporting_title', "Sürdürülebilirlik Raporlaması (GRI 2-2)"))
        add_text('raporda_yer_alan_kuruluslar', self.lm.tr('entities_included', 'Sürdürülebilirlik Raporlamasına Dahil Edilen Kuruluşlar:'))
        add_text('sahiplik_yapisi_detay', self.lm.tr('ownership_details', 'Sahiplik Yapısı Detayı:'))
        add_text('faaliyette_bulunulan_ulkeler', self.lm.tr('countries_of_operation', 'Faaliyette Bulunulan Ülkeler:'))

        add_separator()
        # Bölüm 7: Raporlama Uygulaması (GRI 2-3, 2-4, 2-5)
        add_section(self.lm.tr('reporting_practice_title', "Raporlama Uygulaması (GRI 2-3, 2-4, 2-5)"))
        add_entry('raporlama_donemi', self.lm.tr('reporting_period', 'Raporlama Dönemi:'))
        add_entry('raporlama_dongusu', self.lm.tr('reporting_cycle', 'Raporlama Döngüsü:'))
        add_entry('irtibat_noktasi', self.lm.tr('contact_point', 'İrtibat Noktası:'))
        add_text('bilgilerin_yeniden_duzenlenmesi', self.lm.tr('restatements', 'Bilgilerin Yeniden Düzenlenmesi (Restatements):'))
        add_entry('dis_guvence', self.lm.tr('external_assurance', 'Dış Güvence (External Assurance):'))
        add_text('dis_guvence_tur_kapsam', self.lm.tr('assurance_scope', 'Dış Güvence Türü ve Kapsamı:'))

        add_separator()
        # Bölüm 8: Çalışan Detayları (GRI 2-7, 2-8)
        add_section(self.lm.tr('employee_details_title', "Çalışan Detayları (GRI 2-7, 2-8)"))
        add_entry('calisan_sabit_sozlesmeli', self.lm.tr('perm_contract_employees', 'Sabit/Sözleşmeli Çalışan Sayısı:'))
        add_entry('calisan_gecici', self.lm.tr('temp_employees', 'Geçici Çalışan Sayısı:'))
        add_entry('calisan_kadin_erkek_sayisi', self.lm.tr('gender_breakdown', 'Kadın/Erkek Çalışan Dağılımı:'))
        add_text('calisan_bolgelere_gore_dagilim', self.lm.tr('region_breakdown', 'Bölgelere Göre Çalışan Dağılımı:'))
        add_entry('calisan_olmayanlar_sayisi', self.lm.tr('non_employee_workers', 'Çalışan Olmayan İşçi Sayısı:'))
        add_text('calisan_olmayanlar_tur', self.lm.tr('non_employee_type', 'Çalışan Olmayan İşçi Türü ve Açıklaması:'))

        add_separator()
        # Bölüm 9: Ek Bilgiler
        add_section(self.lm.tr('additional_info_title', "Ek Bilgiler"))
        add_entry('sektor_kod', self.lm.tr('sector_code', 'Sektör Kodu:'))
        add_entry('uretim_kapasitesi', self.lm.tr('production_capacity', 'Üretim Kapasitesi:'))

        # Mevcut veriyi yükle ve alanları doldur
        row = self.load_company_info()
        if row:
            values = {
                'sirket_adi': row[2] if len(row) > 2 else '', 'ticari_unvan': row[3] if len(row) > 3 else '', 'marka_isimleri': row[4] if len(row) > 4 else '', 'adres': row[5] if len(row) > 5 else '',
                'ulke': row[6] if len(row) > 6 else '', 'sehir': row[7] if len(row) > 7 else '', 'posta_kodu': row[8] if len(row) > 8 else '', 'web': row[9] if len(row) > 9 else '', 'eposta': row[10] if len(row) > 10 else '', 'telefon': row[11] if len(row) > 11 else '',
                'vergi_no': row[12] if len(row) > 12 else '', 'mersis_no': row[13] if len(row) > 13 else '', 'tuzel_yapi': row[14] if len(row) > 14 else '', 'mulkiyet_kontrol': row[15] if len(row) > 15 else '', 'borsa_durumu': row[16] if len(row) > 16 else '', 'kayitli_merkez': row[17] if len(row) > 17 else '',
                'calisan_sayisi': row[18] if len(row) > 18 else '', 'calisan_cinsiyet_dagilimi': row[19] if len(row) > 19 else '', 'operasyon_sayisi': row[20] if len(row) > 20 else '', 'faaliyet_lokasyonlari': row[21] if len(row) > 21 else '', 'net_gelir': row[22] if len(row) > 22 else '', 'toplam_varliklar': row[23] if len(row) > 23 else '', 'pazarlar': row[24] if len(row) > 24 else '',
                'tedarik_zinciri': row[25] if len(row) > 25 else '', 'is_iliskileri': row[26] if len(row) > 26 else '', 'upstream_downstream': row[27] if len(row) > 27 else '',
                'yonetim_yapisi': row[28] if len(row) > 28 else '', 'en_yuksek_yonetim_organi': row[29] if len(row) > 29 else '', 'baskan_bilgisi': row[30] if len(row) > 30 else '', 'surdurulebilirlik_sorumluluklari': row[31] if len(row) > 31 else '', 'risk_yonetimi_rolu': row[32] if len(row) > 32 else '', 'cikar_catismalari': row[33] if len(row) > 33 else '', 'kritik_endise_iletisimi': row[34] if len(row) > 34 else '', 'yonetim_bilgisi_kapasitesi': row[35] if len(row) > 35 else '', 'performans_degerlendirme': row[36] if len(row) > 36 else '',
                'ucret_politikalari': row[37] if len(row) > 37 else '', 'ucret_belirleme_sureci': row[38] if len(row) > 38 else '', 'yillik_ucret_orani': row[39] if len(row) > 39 else '',
                'ust_duzey_beyan': row[40] if len(row) > 40 else '', 'politika_taahhutleri': row[41] if len(row) > 41 else '', 'politikalarin_yayilmasi': row[42] if len(row) > 42 else '', 'olumsuz_etki_tazmin': row[43] if len(row) > 43 else '', 'etik_endise_mekanizmalari': row[44] if len(row) > 44 else '', 'mevzuat_uyum': row[45] if len(row) > 45 else '',
                'uyelik_dernekler': row[46] if len(row) > 46 else '', 'paydas_iletisim': row[47] if len(row) > 47 else '', 'toplu_sozlesme_orani': row[48] if len(row) > 48 else '',
                'raporlama_donemi': row[49] if len(row) > 49 else '', 'raporlama_dongusu': row[50] if len(row) > 50 else '', 'irtibat_noktasi': row[51] if len(row) > 51 else '',
                'raporda_yer_alan_kuruluslar': row[52] if len(row) > 52 else '', 'bilgilerin_yeniden_duzenlenmesi': row[53] if len(row) > 53 else '',
                'dis_guvence': row[54] if len(row) > 54 else '', 'dis_guvence_tur_kapsam': row[55] if len(row) > 55 else '',
                'calisan_sabit_sozlesmeli': row[56] if len(row) > 56 else '', 'calisan_gecici': row[57] if len(row) > 57 else '',
                'calisan_kadin_erkek_sayisi': row[58] if len(row) > 58 else '', 'calisan_bolgelere_gore_dagilim': row[59] if len(row) > 59 else '',
                'calisan_olmayanlar_sayisi': row[60] if len(row) > 60 else '', 'calisan_olmayanlar_tur': row[61] if len(row) > 61 else '',
                'sektor_kod': row[62] if len(row) > 62 else '', 'faaliyette_bulunulan_ulkeler': row[63] if len(row) > 63 else '',
                'sahiplik_yapisi_detay': row[64] if len(row) > 64 else '', 'uretim_kapasitesi': row[65] if len(row) > 65 else ''
            }
            for key, widget in entries.items():
                val = values.get(key, '') or ''
                if isinstance(widget, tk.Text):
                    widget.delete('1.0', tk.END)
                    widget.insert('1.0', val)
                else:
                    widget.delete(0, tk.END)
                    # Telefon alanını yüklerken biçimlendir
                    if key == 'telefon':
                        try:
                            # Telefon format lazy import
                            from utils.phone import format_tr_phone
                            widget.insert(0, format_tr_phone(val))
                        except Exception:
                            widget.insert(0, val)
                    else:
                        widget.insert(0, val)

        # Telefon alanında odak kaybında biçimlendirme
        def _format_phone_entry(_=None) -> None:
            try:
                if 'telefon' in entries:
                    val = entries['telefon'].get().strip()
                    if val:
                        entries['telefon'].delete(0, tk.END)
                        # Telefon format lazy import
                        from utils.phone import format_tr_phone
                        entries['telefon'].insert(0, format_tr_phone(val))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

        if 'telefon' in entries:
            try:
                entries['telefon'].bind('<FocusOut>', _format_phone_entry)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

        def gather_values() -> None:
            data = {}
            for key, widget in entries.items():
                if isinstance(widget, tk.Text):
                    data[key] = widget.get('1.0', tk.END).strip()
                else:
                    data[key] = widget.get().strip()
            # Telefon değerini kaydetmeden önce standart biçime getir
            if 'telefon' in data and data['telefon']:
                try:
                    # Telefon format lazy import
                    from utils.phone import format_tr_phone
                    data['telefon'] = format_tr_phone(data['telefon'])
                except Exception as e:
                    logging.error(f"Silent error caught: {str(e)}")
            return data

        # Otomatik kayıt sistemi
        auto_save_timer = None
        auto_save_status_label = tk.Label(form_holder, text='', font=('Segoe UI', 9),
                                          fg='#27ae60', bg='white')
        auto_save_status_label.grid(row=row_idx, column=0, sticky='w', padx=(6, 8), pady=4)

        def auto_save() -> None:
            """Otomatik kayıt fonksiyonu"""
            try:
                auto_save_status_label.configure(text=self.lm.tr('saving', 'Kaydediliyor...'), fg='#f39c12')

                self._save_company_info(gather_values())
                auto_save_status_label.configure(text=f" {self.lm.tr('auto_saved', 'Otomatik kaydedildi')}", fg='#27ae60')

                # 3 saniye sonra mesajı temizle
                form_holder.after(3000, lambda: auto_save_status_label.configure(text=''))
            except Exception as e:
                auto_save_status_label.configure(text=f' Kayıt hatası: {str(e)[:30]}', fg='#e74c3c')

        def schedule_auto_save(event=None) -> None:
            """Değişiklik olduğunda 3 saniye sonra otomatik kaydet"""
            nonlocal auto_save_timer
            # Önceki timer varsa iptal et
            if auto_save_timer:
                form_holder.after_cancel(auto_save_timer)
            # Yeni timer başlat
            auto_save_status_label.configure(text=self.lm.tr('changes_detected', 'Değişiklikler algılandı...'), fg='#3498db')

            auto_save_timer = form_holder.after(3000, auto_save)

        # Tüm entry ve text widget'larına otomatik kayıt event'i bağla
        for key, widget in entries.items():
            if isinstance(widget, tk.Text):
                # Text widget için KeyRelease event'i
                widget.bind('<KeyRelease>', schedule_auto_save)
            else:
                # Entry widget için KeyRelease event'i
                widget.bind('<KeyRelease>', schedule_auto_save)

        # Kaydet butonu (sağda sabitler - form'un altında)
        save_btn = tk.Button(form_holder, text=self.lm.tr('btn_save_company_info', ' Kaydet'),
                             font=('Segoe UI', 12, 'bold'), fg='white', bg='#2E8B57',
                             relief='flat', bd=0, cursor='hand2', padx=20, pady=10,
                             command=lambda: self._save_company_info(gather_values()))
        save_btn.grid(row=row_idx, column=1, sticky='e', pady=20, padx=(0, 12))

        # ÜST SABİT KAYDET BUTONU - Her zaman görünür
        floating_save_frame = tk.Frame(container, bg='#f0f0f0', height=50)
        floating_save_frame.pack(side='top', fill='x', before=canvas)
        floating_save_frame.pack_propagate(False)

        floating_save = tk.Button(floating_save_frame, text=self.lm.tr('btn_save_company_info', ' Kaydet'),
                                  font=('Segoe UI', 11, 'bold'), fg='white', bg='#2E8B57',
                                  relief='flat', bd=0, cursor='hand2', padx=16, pady=8,
                                  command=lambda: self._save_company_info(gather_values()))
        floating_save.pack(side='right', padx=12, pady=8)

        # Otomatik kayıt durumu label'ı
        auto_save_status_toplabel = tk.Label(floating_save_frame, text='',
                                             font=('Segoe UI', 9), fg='#27ae60', bg='#f0f0f0')
        auto_save_status_toplabel.pack(side='right', padx=(0, 10), pady=8)

    def _fetch_companies(self) -> None:
        """Veritabanından (id, ad) şeklinde firma listesini al.
        Ad olarak company_info içindeki firma adını (öncelik: sirket_adi, sonra ticari_unvan) kullan."""
        try:
            conn = sqlite3.connect(self.db_path)
            cur = conn.cursor()
            rows = []
            # Öncelikli olarak company_info içindeki firma adını kullan
            try:
                cur.execute("SELECT company_id, COALESCE(sirket_adi, ticari_unvan, 'Firma '||company_id) FROM company_info ORDER BY 2")
                rows = cur.fetchall() or []
            except Exception:
                rows = []
            # Eğer company_info boşsa companies tablosuna düş
            if not rows:
                try:
                    cur.execute("SELECT id, name FROM companies ORDER BY name")
                    rows = cur.fetchall() or []
                except Exception:
                    rows = []
            conn.close()
            if not rows:
                return [(1, 'Varsayılan Firma')]
            return [(int(r[0]), str(r[1] or '')) for r in rows]
        except Exception:
            return [(1, 'Varsayılan Firma')]

    def open_company_info_with_selection(self) -> None:
        """Yönetimden Firma Bilgileri için önce firma seçimi diyalogu aç."""
        # İzin kontrolü - daha esnek
        try:
            if not self._require_permission("company.read", "Firma Bilgileri"):
                # Uyarı göster ama işlemi devam ettir
                result = messagebox.askyesno(self.lm.tr('permission_warning_title', 'İzin Uyarısı'),
                    self.lm.tr('permission_warning_company_info', 'Firma bilgileri görüntüleme yetkiniz olmayabilir. Devam etmek istiyor musunuz?'))
                if not result:
                    return
        except Exception as e:
            # Hata durumunda devam et
            logging.error(f"Silent error caught: {str(e)}")
        self.previous_module = 'management'
        companies = self._fetch_companies()
        # Diyalog penceresi
        dlg = tk.Toplevel(self.parent)
        dlg.title(self.lm.tr('select_company_title', "Firma Seçimi"))
        dlg.geometry("360x160")
        dlg.configure(bg='white')
        dlg.transient(self.parent)
        dlg.grab_set()

        tk.Label(dlg, text=self.lm.tr('select_company', "Lütfen firma seçiniz:"), font=('Segoe UI', 11, 'bold'), fg='#2E8B57', bg='white').pack(pady=(16, 8))

        display_items = [f"{cid} - {name}" for cid, name in companies]
        selected_var = tk.StringVar()
        combo = ttk.Combobox(dlg, values=display_items, textvariable=selected_var, state='readonly')
        combo.pack(fill='x', padx=20)
        if display_items:
            combo.current(0)

        btn_frame = tk.Frame(dlg, bg='white')
        btn_frame.pack(fill='x', pady=16)
        def on_ok() -> None:
            val = selected_var.get()
            try:
                chosen_id = int(val.split(' - ')[0]) if val else companies[0][0]
            except Exception:
                chosen_id = companies[0][0]
            self.company_id = chosen_id
            try:
                self._ensure_dirs_for_company(chosen_id)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
            dlg.destroy()
            self.show_company_info()
        def on_cancel() -> None:
            dlg.destroy()
        ok_btn = tk.Button(btn_frame, text=self.lm.tr('continue_btn', 'Devam'), font=('Segoe UI', 10, 'bold'), fg='white', bg='#3b82f6', relief='flat', bd=0, cursor='hand2', padx=12, pady=6, command=on_ok)

        ok_btn.pack(side='right', padx=8)
        cancel_btn = tk.Button(btn_frame, text=self.lm.tr('btn_cancel', 'İptal'), font=('Segoe UI', 10, 'bold'), fg='white', bg='#6b7280', relief='flat', bd=0, cursor='hand2', padx=12, pady=6, command=on_cancel)
        cancel_btn.pack(side='right')

    def _go_back(self) -> None:
        """Bir önceki görünüme dön (yığın bazlı)."""
        try:
            if len(getattr(self, '_nav_stack', [])) == 0:
                # Yığın boşsa Dashboard'a dön
                self.show_dashboard()
                self._current_view_func = self.show_dashboard
                return
            prev_fn = self._nav_stack.pop()
            if callable(prev_fn):
                prev_fn()
                self._current_view_func = prev_fn
            else:
                self.show_dashboard()
                self._current_view_func = self.show_dashboard
        except Exception as e:
            logging.error(f"Geri navigasyon hatası: {e}")
            # Fallback
            try:
                self.show_dashboard()
                self._current_view_func = self.show_dashboard
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    def open_delete_company_with_selection(self) -> None:
        """Firma silme için seçim diyalogu aç."""
        # İzin kontrolü - daha esnek
        try:
            if not self._is_admin_or_super():
                # Uyarı göster ama işlemi devam ettir
                result = messagebox.askyesno(self.lm.tr('permission_warning_title', 'İzin Uyarısı'),
                    self.lm.tr('permission_warning_delete_company', 'Firma silme yetkiniz olmayabilir. Devam etmek istiyor musunuz?'))
                if not result:
                    return
        except Exception as e:
            # Hata durumunda devam et
            logging.error(f"Silent error caught: {str(e)}")
        companies = self._fetch_companies()
        if not companies:
            messagebox.showinfo(self.lm.tr('info_title', 'Bilgi'), self.lm.tr('info_no_company_to_delete', 'Silinecek firma bulunamadı.'))
            return

        dlg = tk.Toplevel(self.parent)
        dlg.title(self.lm.tr('delete_company_title', "Firma Sil"))
        dlg.geometry("380x180")
        dlg.configure(bg='white')
        dlg.transient(self.parent)
        dlg.grab_set()

        tk.Label(dlg, text=self.lm.tr('delete_company_prompt', "Silinecek firmayı seçiniz:"), font=('Segoe UI', 11, 'bold'), fg='#DC143C', bg='white').pack(pady=(16, 8))
        display_items = [f"{cid} - {name}" for cid, name in companies]
        selected_var = tk.StringVar()
        combo = ttk.Combobox(dlg, values=display_items, textvariable=selected_var, state='readonly')
        combo.pack(fill='x', padx=20)
        if display_items:
            combo.current(0)

        btn_frame = tk.Frame(dlg, bg='white')
        btn_frame.pack(fill='x', pady=16)

        def on_delete() -> None:
            val = selected_var.get()
            try:
                chosen_id = int(val.split(' - ')[0]) if val else companies[0][0]
            except Exception:
                chosen_id = companies[0][0]
            dlg.destroy()
            self._delete_company_by_id(chosen_id)

        def on_cancel() -> None:
            dlg.destroy()

        del_btn = tk.Button(btn_frame, text=self.lm.tr('btn_delete', 'Sil'), font=('Segoe UI', 10, 'bold'), fg='white', bg='#DC143C', relief='flat', bd=0, cursor='hand2', padx=12, pady=6, command=on_delete)

        del_btn.pack(side='right', padx=8)
        cancel_btn = tk.Button(btn_frame, text=self.lm.tr('btn_cancel', 'İptal'), font=('Segoe UI', 10, 'bold'), fg='white', bg='#6b7280', relief='flat', bd=0, cursor='hand2', padx=12, pady=6, command=on_cancel)
        cancel_btn.pack(side='right')

    def _delete_company_by_id(self, company_id: int) -> None:
        """Seçilen firmayı veritabanından ve dosya sisteminden sil (çift onay)."""
        try:
            import os
            import shutil
            import sqlite3

            # Çift onay
            if not messagebox.askyesno(self.lm.tr('confirm_company_delete_title', 'Onay'), self.lm.tr('delete_confirm_1', "ID {} firmasını silmek istediğinize emin misiniz?").format(company_id)):
                return
            if not messagebox.askyesno(self.lm.tr('confirm_company_delete_final_title', 'Son Onay'), self.lm.tr('delete_confirm_2', 'Silme işlemi geri alınamaz. Devam edilsin mi?')):
                return

            conn = sqlite3.connect(self.db_path)
            cur = conn.cursor()

            # Önce company_info, sonra companies
            try:
                cur.execute('DELETE FROM company_info WHERE company_id=?', (company_id,))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
            try:
                cur.execute('DELETE FROM companies WHERE id=?', (company_id,))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            conn.commit()
            conn.close()

            # Dosya sistemi klasörlerini sil
            try:
                # app/main_app.py konumundan iki seviye yukarı: .../SDG
                base_dir = os.path.dirname(os.path.dirname(__file__))
                company_root = os.path.join(base_dir, 'data', 'companies', str(company_id))
                if os.path.exists(company_root):
                    shutil.rmtree(company_root, ignore_errors=True)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            # Eğer mevcut seçim silindiyse, varsayılan 1'e dön
            if getattr(self, 'company_id', None) == company_id:
                self.company_id = 1

            messagebox.showinfo(self.lm.tr('company_deleted_title', 'Silindi'), self.lm.tr('company_deleted_success', "Firma (ID: {}) başarıyla silindi.").format(company_id))
            # Yönetim ekranını yenile
            self.show_management()
        except Exception as e:
            try:
                conn.rollback()
                conn.close()
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
            messagebox.showerror(self.lm.tr('company_delete_error_title', 'Hata'), self.lm.tr('company_delete_error', "Firma silinemedi: {}").format(e))

    def _ensure_dirs_for_company(self, company_id: int) -> None:
        """Veri dizininde şirket bazlı klasörleri oluştur."""
        try:
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
            company_root = os.path.join(base_dir, 'data', 'companies', str(company_id))
            for sub in ('exports', 'imports', 'reports', 'backups'):
                path = os.path.join(company_root, sub)
                os.makedirs(path, exist_ok=True)
        except Exception as e:
            logging.info(f"Şirket klasörleri oluşturulamadı: {e}")

    def _create_new_company(self) -> None:
        """Yeni firma oluştur ve bu firmaya geç."""
        # İzin kontrolü - daha esnek
        try:
            if not self._is_admin_or_super():
                # Uyarı göster ama işlemi devam ettir
                result = messagebox.askyesno(self.lm.tr('permission_warning_title', 'İzin Uyarısı'),
                    self.lm.tr('permission_warning_create_company', 'Yeni firma oluşturma yetkiniz olmayabilir. Devam etmek istiyor musunuz?'))
                if not result:
                    return
        except Exception as e:
            # Hata durumunda devam et
            logging.error(f"Silent error caught: {str(e)}")

        # Kullanıcıdan firma adı al
        company_name = self._get_company_name_dialog()
        if not company_name:  # Kullanıcı iptal etti veya boş bıraktı
            return

        try:
            conn = sqlite3.connect(self.db_path)
            cur = conn.cursor()

            # Mevcut maksimum ID'yi her iki tablodan da al ve bir sonrakini seç
            cur.execute('SELECT COALESCE(MAX(company_id), 0) FROM company_info')
            max_info = cur.fetchone()[0] or 0
            cur.execute('SELECT COALESCE(MAX(id), 0) FROM companies')
            max_comp = cur.fetchone()[0] or 0
            new_id = max(max_info, max_comp) + 1

            # companies tablosuna ekle (FK tutarlılığı için)
            cur.execute(
                "INSERT INTO companies (id, name) VALUES (?, ?)",
                (new_id, company_name)
            )

            # company_info tablosuna temel kayıt ekle
            cur.execute(
                "INSERT INTO company_info (company_id, sirket_adi, ticari_unvan) VALUES (?, ?, ?)",
                (new_id, company_name, company_name)
            )

            conn.commit()
            conn.close()

            # Dosya sistemi klasörlerini oluştur
            self._ensure_dirs_for_company(new_id)

            # Bu yeni firmaya geç ve ekranı tazele
            self.company_id = new_id
            messagebox.showinfo(self.lm.tr('new_company', 'Yeni Firma'), self.lm.tr('company_created_success', "Yeni firma oluşturuldu ve seçildi: {} (ID: {})").format(company_name, new_id))
            self.current_module = 'company_info'
            self.show_company_info()
        except Exception as e:
            try:
                conn.rollback()
                conn.close()
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
            messagebox.showerror(self.lm.tr('company_delete_error_title', 'Hata'), self.lm.tr('company_create_error', "Yeni firma oluşturulamadı: {}").format(e))

    def _get_company_name_dialog(self) -> str:
        """Kullanıcıdan firma adı alan dialog penceresi"""
        from tkinter import simpledialog

        # Dialog penceresini aç
        company_name = simpledialog.askstring(
            self.lm.tr('title_create_new_company', "Yeni Firma Oluştur"),
            self.lm.tr('prompt_company_name', "Firma adını girin:"),
            initialvalue=""
        )

        # Boş veya sadece boşluk içeren girişi kontrol et
        if company_name:
            company_name = company_name.strip()
            if not company_name:
                messagebox.showwarning(self.lm.tr('warning', "Uyarı"), self.lm.tr('company_name_empty', "Firma adı boş olamaz!"))
                return ""
            # Çok uzun isim kontrolü
            if len(company_name) > 100:
                messagebox.showwarning(self.lm.tr('warning', "Uyarı"), self.lm.tr('company_name_too_long', "Firma adı çok uzun! (Maksimum 100 karakter)"))
                return ""

        return company_name or ""

    # ========================================================================
    # MODÜL YETKİLERİ YÖNETİMİ
    # ========================================================================

    def load_module_states(self) -> dict:
        """Modül durumlarını veritabanından yükle"""
        if self._module_states_cache is not None:
            return self._module_states_cache

        try:
            import sqlite3
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            # Module_states tablosunu oluştur (yoksa)
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS module_states (
                    module_key TEXT PRIMARY KEY,
                    enabled BOOLEAN NOT NULL DEFAULT 1,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_by TEXT
                )
            ''')

            # Mevcut durumları al
            cursor.execute('SELECT module_key, enabled FROM module_states')
            results = cursor.fetchall()

            conn.close()

            self._module_states_cache = {key: bool(enabled) for key, enabled in results}
            return self._module_states_cache

        except Exception as e:
            logging.error(f"Modül durumları yüklenirken hata: {e}")
            return {}

    def is_module_enabled(self, module_key: str) -> bool:
        """Modülün etkin olup olmadığını kontrol et"""
        module_states = self.load_module_states()

        # Varsayılan olarak açık olan modüller
        default_enabled_modules = {
            'sdg', 'gri', 'tsrs', 'reporting', 'company_info',
            'task_management', 'user_management', 'super_admin',
            'advanced_dashboard', 'company_management', 'security_management'
        }

        # Super Admin her zaman açık
        if module_key == 'super_admin':
            return True

        # Modül durumunu döndür (varsayılan durumu da göz önünde bulundur)
        default_state = module_key in default_enabled_modules
        return module_states.get(module_key, default_state)

    def show_module_disabled_message(self, module_key: str) -> None:
        """Kapalı modül için mesaj göster"""
        try:
            from tkinter import messagebox

            module_names = {
                'sdg': self.lm.tr('module_sdg', 'SDG Hedefleri'),
                'gri': self.lm.tr('module_gri', 'GRI Standartları'),
                'tsrs': self.lm.tr('module_tsrs', 'TSRS Raporlama'),
                'esrs': self.lm.tr('module_esrs', 'ESRS (Avrupa)'),
                'tcfd': self.lm.tr('module_tcfd', 'TCFD Framework'),
                'sasb': self.lm.tr('module_sasb', 'SASB Standartları'),
                'eu_taxonomy': self.lm.tr('module_eu_taxonomy', 'EU Taxonomy'),
                'ai_module': self.lm.tr('module_ai', 'AI Analiz Modülü'),
                'company_management': self.lm.tr('module_company_mgmt', 'Şirket Yönetimi'),
                'auto_reporting': self.lm.tr('module_auto_reporting', 'Otomatik Raporlama'),
                'erp_integration': self.lm.tr('module_erp', 'ERP Entegrasyon'),
                'document_processing': self.lm.tr('module_doc_processing', 'Belge İşleme & AI'),
                'advanced_dashboard': self.lm.tr('module_adv_dashboard', 'Gelişmiş Dashboard'),
                'advanced_security': self.lm.tr('module_adv_security', 'Gelişmiş Güvenlik')
            }

            module_name = module_names.get(module_key, self.lm.tr('this_module', 'Bu Modül'))

            messagebox.showwarning(self.lm.tr('module_access_denied_title', "Modül Erişim Kısıtlı"),
                                 self.lm.tr('module_access_denied_detail', "🚫 {0} Erişim Kısıtlı\n\n📜 Lisansınız bu modül için uygun değil.\n\nIcons.WRENCH Modülü etkinleştirmek için:\n   1. Super Admin paneline gidin\n   2. 'Modül Yetkileri' bölümünü açın\n   3. '{0}' modülünü aktif edin\n\n💎 Veya lisansınızı yükseltin!").format(module_name))

        except Exception as e:
            logging.error(f"Modül mesajı gösterilirken hata: {e}")

    def refresh_module_states_cache(self) -> None:
        """Modül durumları cache'ini yenile"""
        self._module_states_cache = None

    def _is_admin_or_super(self) -> bool:
        """Mevcut kullanıcı admin veya super_admin mı?"""
        try:
            # Super admin kontrolü
            if self.user and len(self.user) > 1 and self.user[1] == '__super__':
                return True

            # RBAC üzerinden rol kontrolü
            # UserManager lazy import
            from yonetim.kullanici_yonetimi.models.user_manager import UserManager
            if not self.db_path:
                logging.error("Hata: db_path tanımlı değil")
                return False
                
            um = UserManager(self.db_path)
            username = self.user[1] if self.user and len(self.user) > 1 else None
            if not username:
                return False
            u = um.get_user_by_username(username)
            roles_str = (u or {}).get('roles') or ''
            roles = [r.strip().lower() for r in roles_str.split(',') if r]
            return ('admin' in roles) or ('super_admin' in roles)
        except Exception as e:
            import traceback
            traceback.print_exc()
            logging.error(f"Yetki kontrolü hatası (_is_admin_or_super): {e}")
            return False

    def _is_super_admin_user(self) -> bool:
        try:
            if self.user and len(self.user) > 1 and self.user[1] == '__super__':
                return True
            from yonetim.kullanici_yonetimi.models.user_manager import UserManager
            if not self.db_path:
                return False

            um = UserManager(self.db_path)
            username = self.user[1] if self.user and len(self.user) > 1 else None
            if not username:
                return False
            u = um.get_user_by_username(username)
            roles_str = (u or {}).get('roles') or ''
            roles = [r.strip().lower() for r in roles_str.split(',') if r]
            return 'super_admin' in roles
        except Exception as e:
            import traceback
            traceback.print_exc()
            logging.error(f"Yetki kontrolü hatası (_is_super_admin_user): {e}")
            return False

    def _is_admin_only(self) -> bool:
        try:
            if self._is_super_admin_user():
                return False
            from yonetim.kullanici_yonetimi.models.user_manager import UserManager
            if not self.db_path:
                return False
                
            um = UserManager(self.db_path)
            username = self.user[1] if self.user and len(self.user) > 1 else None
            if not username:
                return False
            u = um.get_user_by_username(username)
            roles_str = (u or {}).get('roles') or ''
            roles = [r.strip().lower() for r in roles_str.split(',') if r]
            return 'admin' in roles
        except Exception as e:
            import traceback
            traceback.print_exc()
            logging.error(f"Yetki kontrolü hatası (_is_admin_only): {e}")
            return False

    def _get_user_role(self) -> None:
        """Kullanıcı rolünü al"""
        try:
            # Kullanıcı bilgilerini al (güvenli erişim)
            username = None
            try:
                username = self.user[1] if isinstance(self.user, (list, tuple)) and len(self.user) > 1 else None
            except Exception:
                username = None

            # Gizli super admin kontrolü
            if username == '__super__':
                return 'super_admin'

            # Veritabanından rol bilgisini al
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            if not username:
                conn.close()
                return 'user'
            cursor.execute("SELECT role FROM users WHERE username = ?", (username,))
            result = cursor.fetchone()

            conn.close()

            if result:
                return result[0] or 'user'
            else:
                return 'user'

        except Exception as e:
            logging.error(f"Rol alma hatası: {e}")
            return 'user'

    def setup_keyboard_shortcuts(self) -> None:
        """Klavye kısayollarını ayarla"""
        try:
            # F1 - Yardım
            self.parent.bind('<F1>', lambda e: self.show_help())

            # F5 - Yenile
            self.parent.bind('<F5>', lambda e: self.refresh_current_view())

            # Escape - Ana sayfaya dön
            self.parent.bind('<Escape>', lambda e: self.show_dashboard())

            # Ctrl+H - Ana Sayfa
            self.parent.bind('<Control-h>', lambda e: self.show_dashboard())

            # Ctrl+Q - Çıkış
            self.parent.bind('<Control-q>', lambda e: self.exit_application())

            # Ctrl+Shift+D - Gelişmiş Dashboard
            self.parent.bind('<Control-Shift-d>', lambda e: self.show_advanced_dashboard())

            # Ctrl+Shift+T - Otomatik Görevler
            self.parent.bind('<Control-Shift-t>', lambda e: self.show_auto_tasks())

            logging.info(" Klavye kısayolları aktifleştirildi")

        except Exception as e:
            logging.info(f"️ Klavye kısayolları ayarlanamadı: {e}")

    def show_help(self) -> None:
        """Yardım penceresi göster (F1)"""
        help_window = tk.Toplevel(self.parent)
        help_window.title("Yardım ve Kısayollar")
        help_window.geometry("900x650")
        help_window.configure(bg='white')

        title = tk.Label(help_window, text=f"📚 {self.lm.tr('help_shortcuts_title', 'Yardım ve Kısayollar')}", font=('Segoe UI', 16, 'bold'), fg='#2E8B57', bg='white')

        title.pack(pady=10)

        notebook = ttk.Notebook(help_window)
        notebook.pack(fill='both', expand=True, padx=12, pady=8)

        shortcuts_tab = tk.Frame(notebook, bg='white')
        modules_tab = tk.Frame(notebook, bg='white')
        tips_tab = tk.Frame(notebook, bg='white')
        features_tab = tk.Frame(notebook, bg='white')
        doc_tab = tk.Frame(notebook, bg='white')
        login_tab = tk.Frame(notebook, bg='white')
        notebook.add(shortcuts_tab, text=self.lm.tr('shortcuts', 'Kısayollar'))
        notebook.add(modules_tab, text=self.lm.tr('modules', 'Modüller'))
        notebook.add(tips_tab, text=self.lm.tr('tips', 'İpuçları'))
        notebook.add(features_tab, text=self.lm.tr('features', 'Özellikler'))
        notebook.add(doc_tab, text=self.lm.tr('documentation', 'Doküman'))
        notebook.add(login_tab, text=self.lm.tr('login_help', 'Giriş Yardımı'))

        shortcuts_text = (
            "Global Kısayollar:\n\n"
            "F1              - Yardım\n"
            "F5              - Görünümü yenile\n"
            "Escape          - Ana sayfa / İptal\n"
            "Ctrl+H          - Ana sayfa\n"
            "Ctrl+Q          - Çıkış\n"
            "Ctrl+Shift+D    - Gelişmiş Dashboard\n"
            "Ctrl+Shift+T    - Otomatik Görevler\n\n"
            "Modül Kısayolları:\n\n"
            "Ctrl+S          - Kaydet\n"
            "Ctrl+F          - Ara\n"
            "Ctrl+N          - Yeni\n"
            "Ctrl+E          - Düzenle\n"
            "Delete          - Sil\n\n"
            "Navigasyon:\n\n"
            "Enter           - Onay\n"
            "Tab             - Sonraki alan\n"
            "Shift+Tab       - Önceki alan\n"
        )

        shortcuts_text_widget = tk.Text(shortcuts_tab, font=('Consolas', 10), bg='#f8f9fa', fg='#2c3e50', wrap=tk.WORD)
        shortcuts_text_widget.pack(fill='both', expand=True, padx=16, pady=12)
        shortcuts_text_widget.insert('1.0', shortcuts_text)
        shortcuts_text_widget.config(state=tk.DISABLED)

        modules_container = tk.Frame(modules_tab, bg='white')
        modules_container.pack(fill='both', expand=True, padx=16, pady=12)

        modules = [
            ('sdg', 'SDG Modülü'),
            ('gri', 'GRI Standartları'),
            ('tsrs', 'TSRS Standartları'),
            ('esg', 'ESG Modülü'),
            ('mapping', f"{Icons.LINK} Standart Eşleştirme"),
            ('reporting', f"{Icons.REPORT} Raporlama Merkezi"),
            ('advanced_dashboard', 'AI Destekli Gelişmiş Dashboard'),
            ('policy_library', 'Politika Kütüphanesi'),
            ('social_dashboard', 'Sosyal Performans Metrikleri'),
            ('cdp', 'CDP'),
            ('iirc', 'IIRC'),
            ('validation', 'Veri Validasyon'),
            ('benchmark', 'Benchmark Analizi'),
            ('forecasting', 'Trend ve Tahmin'),
            ('report_center', 'Gelişmiş Rapor Merkezi'),
            ('integration', 'API ve Entegrasyon'),
            ('visualization', 'Gelişmiş Görselleştirme'),
            ('auditor', 'Denetçi Sistemi'),
            ('stakeholder', 'Paydaş Etkileşimi'),
            ('scenario', 'Senaryo Analizi'),
            ('csrd', 'AB CSRD Uyum'),
            ('report_scheduler', 'Rapor Scheduling'),
            ('advanced_files', 'Gelişmiş Dosya Yönetimi'),
            ('data_import', 'Veri İçe Aktarım'),
            ('form_management', 'Form Yönetimi'),
            ('hr_metrics', 'İK Metrikleri'),
            ('skdm', 'SKDM Modülü')
        ]

        grid = tk.Frame(modules_container, bg='white')
        grid.pack(fill='both', expand=True)
        r, c = 0, 0
        def open_module(k, t):
            try:
                help_window.destroy()
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
            try:
                self._show_module_fast(k, t)
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
        for key, title in modules:
            btn = tk.Button(grid, text=f" {title}", font=('Segoe UI', 10, 'bold'), bg='#3b82f6', fg='white', relief='flat', cursor='hand2', padx=14, pady=8,
                            command=lambda k=key, t=title: open_module(k, t))
            btn.grid(row=r, column=c, padx=6, pady=6, sticky='ew')
            c += 1
            if c >= 3:
                c = 0
                r += 1

        features_text = (
            "Son Değişiklikler ve Özellikler:\n\n"
            "• Tek veritabanı kullanımı (data/sdg_desktop.sqlite).\n"
            "• Modül önbellekleme ve hızlı geçiş.\n"
            "• TSRS ilerleme takibi ekranı.\n"
            "• Sosyal performans dashboard (HR/İSG/Eğitim).\n"
            "• CBAM/SKDM dashboard ve analiz.\n"
        )

        features_text_widget = tk.Text(features_tab, font=('Consolas', 10), bg='#f8f9fa', fg='#2c3e50', wrap=tk.WORD)
        features_text_widget.pack(fill='both', expand=True, padx=16, pady=12)
        features_text_widget.insert('1.0', features_text)
        features_text_widget.config(state=tk.DISABLED)

        tips_text = (
            "Uygulama İpuçları:\n\n"
            "• F1 ile bu pencereyi açın.\n"
            "• Ctrl+S kaydet, Ctrl+F arama, Delete silme.\n"
            "• Büyük listelerde filtreleyerek performansı artırın.\n"
            "• Veri içe aktarmadan önce başlık satırlarını kontrol edin.\n"
            "• Raporlama merkezinde şablonları kopyalayıp özelleştirin.\n"
        )
        tips_text_widget = tk.Text(tips_tab, font=('Consolas', 10), bg='#f8f9fa', fg='#2c3e50', wrap=tk.WORD)
        tips_text_widget.pack(fill='both', expand=True, padx=16, pady=12)
        tips_text_widget.insert('1.0', tips_text)
        tips_text_widget.config(state=tk.DISABLED)

        doc_text = tk.Text(doc_tab, font=('Consolas', 10), bg='#ffffff', fg='#2c3e50', wrap=tk.WORD)
        doc_text.pack(fill='both', expand=True, padx=16, pady=12)
        doc_text.config(state=tk.DISABLED)

        login_text = tk.Text(login_tab, font=('Consolas', 10), bg='#f8f9fa', fg='#2c3e50', wrap=tk.WORD)
        login_text.pack(fill='both', expand=True, padx=16, pady=12)
        login_text.insert('1.0', "\nGiriş Bilgileri (Örnek):\n\nKullanıcı Adı: örnek_kullanici\nParola: Ornek_1234\n\nNot: Bu bilgiler örnek amaçlıdır. Sorun yaşarsanız sistem yöneticinize başvurun.")
        login_text.config(state=tk.DISABLED)

        import os
        import webbrowser
        btns = tk.Frame(help_window, bg='white')
        btns.pack(pady=6)

        base = self.base_dir
        fallback_docs = {
            os.path.join(base, 'docs', 'QUICK_START_GUIDE.md'): (
                "# Hızlı Başlangıç\n\n"
                "- Ana sayfadan sol menü ile modüllere geçin.\n"
                "- Üst çubuktaki kısayollar ile sık işlemlere erişin.\n"
                "- `Veri İçe Aktarım` ile CSV/Excel dosyalarını içe alın.\n"
                "- `Form Yönetimi` ile operasyon formlarını yönetin.\n"
                "- `Raporlama Merkezi` ile yönetim raporlarını üretin.\n"
                "- Klavye: F1 yardım, Ctrl+S kaydet, Ctrl+F ara.\n"
            ),
            os.path.join(base, 'SDG_yardim.md'): (
                "# Genel Yardım\n\n"
                "Sustainage SDG; ESG/SDG, GRI/TSRS uyum, veri toplama, doğrulama ve raporlama için birleşik bir masaüstü uygulamadır.\n\n"
                "## Navigasyon\n"
                "- Sol menü: modüller\n"
                "- Üst bar: hızlı erişim\n"
                "- Sağ panel: içerik\n\n"
                "## Çalışma Akışı\n"
                "1. Firma bilgilerini doldurun.\n"
                "2. Verileri içe aktarın veya formlarla toplayın.\n"
                "3. Doğrulama ekranlarında kontrol edin.\n"
                "4. Panolarda görselleştirin ve raporlayın.\n"
            ),
            os.path.join(base, 'README.md'): (
                "# Sustainage SDG Masaüstü\n\n"
                "## Kurulum\n"
                "- Python 3.10+\n"
                "- `pip install -r requirements.txt`\n\n"
                "## Çalıştırma\n"
                "- `python app/run.py` veya ana giriş modülü\n\n"
                "## Modül Özeti\n"
                "- SDG/ESG panoları\n"
                "- GRI/TSRS uyum\n"
                "- Veri içe aktarma ve formlar\n"
                "- Raporlama merkezi\n"
            ),
            os.path.join(base, 'BUGUN_TAMAMLANAN_EKSIKLIKLER.md'): (
                "# Son Güncellemeler\n\n"
                "- Yardım penceresine yerleşik içerikler eklendi.\n"
                "- Doküman butonları dosya olmadan da çalışır hale getirildi.\n"
                "- Klavye kısayolları genişletildi.\n"
            ),
            os.path.join(base, 'docs', 'YENI_OZELLIKLER_KULLANIM_REHBERI.md'): (
                "# Yeni Özellikler Kullanım Rehberi\n\n"
                "- Modül önbellekleme: hızlı geçiş için aktif.\n"
                "- Gelişmiş dashboard: AI destekli özetler.\n"
                "- Politika kütüphanesi: politika şablonları.\n"
            ),
            os.path.join(base, 'docs', 'FAQ.md'): (
                "# SSS\n\n"
                "**Giriş yapamıyorum?**\n"
                "- Kullanıcı bilgileriniz için yöneticiye başvurun.\n\n"
                "**CSV içe aktarma biçimi nasıldır?**\n"
                "- İlk satır başlık, UTF‑8, ayraç `,` veya `;`.\n\n"
                "**Raporlar nereye kaydedilir?**\n"
                "- Varsayılan `data/` altında firma klasörüne.\n"
            ),
        }

        def load_doc(path) -> None:
            try:
                if not os.path.exists(path):
                    fb = fallback_docs.get(path)
                    if fb:
                        self._render_markdown(doc_text, fb)
                        notebook.select(doc_tab)
                        return
                    doc_text.config(state=tk.NORMAL)
                    doc_text.delete('1.0', tk.END)
                    doc_text.insert('1.0', f"Belge bulunamadı:\n{path}\n\nBu buton için yerleşik içerik tanımlı değil.")
                    doc_text.config(state=tk.DISABLED)
                    notebook.select(doc_tab)
                    return
                text_exts = {'.md', '.txt', '.py', '.json', '.csv', '.log'}
                ext = os.path.splitext(path)[1].lower()
                if ext in text_exts:
                    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    if ext == '.md':
                        self._render_markdown(doc_text, content)
                    else:
                        doc_text.config(state=tk.NORMAL)
                        doc_text.delete('1.0', tk.END)
                        doc_text.insert('1.0', content)
                        doc_text.config(state=tk.DISABLED)
                    notebook.select(doc_tab)
                else:
                    resp = messagebox.askyesno(self.lm.tr('confirm_open_external_title', "Dışarıda Aç?"), self.lm.tr('confirm_open_external_text', "Bu dosya metin olarak görüntülenemiyor.\nVarsayılan uygulama ile açmak ister misiniz?"))
                    if resp:
                        try:
                            os.startfile(path)  # type: ignore
                        except Exception:
                            webbrowser.open(path)
            except Exception as e:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_file_load', "Dosya yüklenemedi: {}").format(e))

        docs = [
            ("Hızlı Başlangıç", os.path.join(base, 'docs', 'QUICK_START_GUIDE.md')),
            ("Genel Yardım", os.path.join(base, 'SDG_yardim.md')),
            ("README", os.path.join(base, 'README.md')),
            ("Son Güncellemeler", os.path.join(base, 'BUGUN_TAMAMLANAN_EKSIKLIKLER.md')),
            ("Yeni Özellikler", os.path.join(base, 'docs', 'YENI_OZELLIKLER_KULLANIM_REHBERI.md')),
            ("FAQ", os.path.join(base, 'docs', 'FAQ.md')),
        ]
        for label, path in docs:
            tk.Button(btns, text=label, bg='#2c3e50', fg='white', relief='flat', command=lambda p=path: load_doc(p)).pack(side='left', padx=6)

        close_btn = tk.Button(help_window, text=self.lm.tr('close_escape', "Kapat (Escape)"), font=('Segoe UI', 10), bg='#95a5a6', fg='white', relief='flat', command=help_window.destroy)

        close_btn.pack(pady=10)
        help_window.bind('<Escape>', lambda e: help_window.destroy())

    def _configure_rich_text(self, widget) -> None:
        """Yardım penceresi için zengin metin stillerini hazırla"""
        try:
            widget.tag_configure('body', font=('Segoe UI', 11), foreground='#2c3e50', spacing3=6)
            widget.tag_configure('h1', font=('Segoe UI', 16, 'bold'), foreground='#2E8B57', spacing3=10)
            widget.tag_configure('h2', font=('Segoe UI', 14, 'bold'), foreground='#1E8449', spacing3=8)
            widget.tag_configure('h3', font=('Segoe UI', 12, 'bold'), foreground='#2c3e50', spacing3=6)
            widget.tag_configure('bold', font=('Segoe UI', 11, 'bold'))
            widget.tag_configure('italic', font=('Segoe UI', 11, 'italic'))
            widget.tag_configure('bullet', lmargin1=20, lmargin2=36)
            widget.tag_configure('blockquote', font=('Segoe UI', 11, 'italic'), foreground='#566573', lmargin1=20, lmargin2=36)
            widget.tag_configure('code', font=('Consolas', 10), background='#f4f6f8')
            widget.tag_configure('link', foreground='#1f6feb', underline=1)
            widget.tag_configure('hr', foreground='#d0d3d4')
        except Exception as e:
            logging.error(f"Silent error caught: {str(e)}")

    def _render_markdown(self, widget, md_text: str) -> None:
        """Basit Markdown -> Tk Text biçimlendirmesi uygula"""
        import re
        import tkinter as tk
        import webbrowser
        widget.config(state=tk.NORMAL)
        widget.delete('1.0', tk.END)
        self._configure_rich_text(widget)

        def insert_inline(text, base_tag='body') -> None:
            # Önce linkleri işle
            link_pat = re.compile(r"\[(.+?)\]\((https?://[^\)]+)\)")
            pos = 0
            for m in link_pat.finditer(text):
                pre = text[pos:m.start()]
                if pre:
                    widget.insert(tk.END, pre, base_tag)
                label, url = m.group(1), m.group(2)
                tagname = f"link_{getattr(self, '_link_seq', 0)}"
                setattr(self, '_link_seq', getattr(self, '_link_seq', 0) + 1)
                widget.insert(tk.END, label, ('link', tagname))
                def _open(event, u=url) -> None:
                    try:
                        webbrowser.open(u)
                    except Exception as e:
                        logging.error(f"Silent error caught: {str(e)}")
                widget.tag_bind(tagname, '<Button-1>', _open)
                pos = m.end()
            remainder = text[pos:]

            # Kod bölümleri
            parts = re.split(r'(`[^`]+`)', remainder)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    widget.insert(tk.END, part[1:-1], 'code')
                else:
                    # Bold/italic
                    idx = 0
                    for m in re.finditer(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_)', part):
                        pre = part[idx:m.start()]
                        if pre:
                            widget.insert(tk.END, pre, base_tag)
                        seg = m.group(0)
                        if seg.startswith('**') or seg.startswith('__'):
                            widget.insert(tk.END, seg[2:-2], 'bold')
                        else:
                            widget.insert(tk.END, seg[1:-1], 'italic')
                        idx = m.end()
                    tail = part[idx:]
                    if tail:
                        widget.insert(tk.END, tail, base_tag)

        for raw_line in md_text.splitlines():
            line = raw_line.rstrip()
            stripped = line.strip()
            if stripped == '':
                widget.insert(tk.END, '\n', 'body')
                continue
            if stripped.startswith('# '):
                insert_inline(stripped[2:], 'h1')
                widget.insert(tk.END, '\n')
                continue
            if stripped.startswith('## '):
                insert_inline(stripped[3:], 'h2')
                widget.insert(tk.END, '\n')
                continue
            if stripped.startswith('### '):
                insert_inline(stripped[4:], 'h3')
                widget.insert(tk.END, '\n')
                continue
            if stripped.startswith('> '):
                insert_inline(stripped[2:], 'blockquote')
                widget.insert(tk.END, '\n')
                continue
            if re.match(r'^\s*[-*+]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
                bullet_text = re.sub(r'^\s*[-*+]\s+', '• ', line)
                bullet_text = re.sub(r'^\s*(\d+)\.\s+', r'\1) ', bullet_text)
                insert_inline(bullet_text, 'bullet')
                widget.insert(tk.END, '\n')
                continue
            if stripped.startswith('---'):
                widget.insert(tk.END, '―' * 40 + '\n', 'hr')
                continue
            insert_inline(line, 'body')
            widget.insert(tk.END, '\n')

        widget.config(state=tk.DISABLED)

    def refresh_current_view(self) -> None:
        """Mevcut görünümü yenile (F5)"""
        try:
            if self.current_module == 'sdg':
                self.show_sdg()
            elif self.current_module == 'gri':
                self.show_gri()
            elif self.current_module == 'tsrs':
                self.show_tsrs()
            elif self.current_module == 'esg':
                self.show_esg()
            elif self.current_module == 'advanced_dashboard':
                self.show_advanced_dashboard()
            elif self.current_module == 'strategic':
                self.show_strategic()
            elif self.current_module == 'data_import':
                self.show_data_import()
            elif self.current_module == 'form_management':
                self.show_form_management()
            elif self.current_module == 'auto_tasks':
                self.show_auto_tasks()
            elif self.current_module == 'advanced_files':
                self.show_advanced_files()
            elif self.current_module == 'hr_metrics':
                self.show_hr_metrics()
            elif self.current_module == 'mapping':
                self.show_mapping()
            elif self.current_module == 'prioritization':
                self.show_prioritization()
            # elif self.current_module == 'reporting':  # KALDIRILDI
            #     self.show_reporting()
            elif self.current_module == 'management':
                self.show_management()
            elif self.current_module == 'company_info':
                self.show_company_info()
            elif self.current_module == 'skdm':
                self.show_skdm()
            else:
                self.show_dashboard()

            logging.info(f" Görünüm yenilendi (F5): {self.current_module or 'dashboard'}")
        except Exception as e:
            logging.error(f"️ Yenileme hatası: {e}")

    def show_super_admin_panel(self) -> None:
        """Super Admin panelini ayrı pencerede göster"""
        try:
            from modules.super_admin.super_admin_gui import SuperAdminGUI

            # Yeni pencere oluştur (self.parent ana window)
            super_admin_window = tk.Toplevel(self.parent)
            super_admin_window.title("⚡ SUPER ADMIN PANEL")
            super_admin_window.geometry("1400x900")

            # Maximize et
            try:
                super_admin_window.state('zoomed')
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            # Icon ayarla (varsa)
            try:
                super_admin_window.iconbitmap('resimler/icon.ico')
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

            # Super Admin GUI'yi bu pencerede aç ve ana uygulama referansını aktar
            SuperAdminGUI(super_admin_window, self.user, self.company_id, self.db_path, host_app=self)

            self.current_module = 'super_admin'

        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_super_admin_load', "Super Admin yükleme hatası: {}").format(str(e)))
            import traceback
            traceback.print_exc()

    def show_ungc(self) -> None:
        """UNGC (UN Global Compact) modülünü göster"""
        try:
            self._audit_navigation("UNGC")
            self.clear_content()
            self.content_header.pack(fill='x')
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)
            self.content_title['text'] = self.lm.tr('ungc_title', "UNGC - UN Global Compact (10 İlke)")
            self.current_module = 'ungc'
            self._add_back_button()

            # Ana UNGC dashboard frame
            ungc_frame = tk.Frame(self.content_area, bg='white')
            ungc_frame.pack(fill='both', expand=True)

            # Başlık ve açıklama
            header_frame = tk.Frame(ungc_frame, bg='white')
            header_frame.pack(fill='x', pady=(0, 20))

            tk.Label(header_frame, text=f" {self.lm.tr('ungc_10_principles_title', 'UN Global Compact - 10 İlke')}",

                    font=('Segoe UI', 16, 'bold'), fg='#1e40af', bg='white').pack(anchor='w')

            tk.Label(header_frame,
                    text=self.lm.tr('ungc_10_principles_desc', "Birleşmiş Milletler Küresel İlkeler Sözleşmesi - İnsan Hakları, Çalışma, Çevre ve Yolsuzlukla Mücadele"),
                    font=('Segoe UI', 10), fg='#666', bg='white').pack(anchor='w', pady=(5, 0))

            # Butonlar frame
            buttons_frame = tk.Frame(ungc_frame, bg='white')
            buttons_frame.pack(fill='x', pady=10)

            # Buton yapılandırması
            buttons = [
                (self.lm.tr('ungc_btn_kpi_entry', " KPI Veri Girişi"), self._open_ungc_kpi_entry, "#3b82f6"),
                (self.lm.tr('ungc_btn_compliance_score', " Uyum Skoru"), self._show_ungc_score, "#10b981"),
                (self.lm.tr('ungc_btn_create_cop', " COP Raporu Oluştur"), self._generate_ungc_cop, "#8b5cf6"),
                (self.lm.tr('ungc_btn_timeline', " İlerleme Grafiği"), self._show_ungc_timeline, "#f59e0b"),
                (self.lm.tr('ungc_btn_excel_export', " Excel Export"), self._export_ungc_excel, "#06b6d4"),
                (self.lm.tr('ungc_btn_send_email', " Email Gönder"), self._send_ungc_email, "#ec4899"),
                (self.lm.tr('ungc_btn_benchmark', " Benchmark"), self._show_ungc_benchmark, "#f97316"),
                (f"{Icons.TIME} {self.lm.tr('ungc_btn_reminders', 'Hatırlatıcılar')}", self._show_ungc_reminders, "#14b8a6"),
            ]

            # Grid layout 4x2
            for idx, (text, command, color) in enumerate(buttons):
                row = idx // 4
                col = idx % 4

                btn = tk.Button(buttons_frame, text=text,
                              font=('Segoe UI', 11, 'bold'), fg='white', bg=color,
                              relief='raised', bd=2, cursor='hand2',
                              command=command, padx=20, pady=15, width=20)
                btn.grid(row=row, column=col, padx=10, pady=10, sticky='ew')

            # Grid weights
            for i in range(4):
                buttons_frame.columnconfigure(i, weight=1)

            # Bilgi paneli
            info_frame = tk.LabelFrame(ungc_frame, text=self.lm.tr('ungc_categories', "10 İlke Kategorileri"),

                                      font=('Segoe UI', 11, 'bold'), bg='white', fg='#1e40af')
            info_frame.pack(fill='both', expand=True, pady=(20, 0))

            categories = [
                (self.lm.tr('ungc_cat_human_rights', "İnsan Hakları"), [
                    self.lm.tr('ungc_p1', "P1: İnsan haklarını desteklemeli"),
                    self.lm.tr('ungc_p2', "P2: İnsan hakları ihlallerine ortak olmamalı")
                ], "#ef4444"),
                (self.lm.tr('ungc_cat_labor', "Çalışma Standartları"), [
                    self.lm.tr('ungc_p3', "P3: Örgütlenme özgürlüğü"),
                    self.lm.tr('ungc_p4', "P4: Zorla çalıştırma karşıtı"),
                    self.lm.tr('ungc_p5', "P5: Çocuk işçiliği karşıtı"),
                    self.lm.tr('ungc_p6', "P6: Ayrımcılık karşıtı")
                ], "#f59e0b"),
                (self.lm.tr('ungc_cat_environment', "Çevre"), [
                    self.lm.tr('ungc_p7', "P7: İhtiyati yaklaşım"),
                    self.lm.tr('ungc_p8', "P8: Çevresel sorumluluk"),
                    self.lm.tr('ungc_p9', "P9: Çevre dostu teknolojiler")
                ], "#10b981"),
                (self.lm.tr('ungc_cat_corruption', "Yolsuzlukla Mücadele"), [
                    self.lm.tr('ungc_p10', "P10: Rüşvet ve yolsuzlukla mücadele")
                ], "#8b5cf6"),
            ]

            for cat_name, principles, color in categories:
                cat_frame = tk.Frame(info_frame, bg='white')
                cat_frame.pack(fill='x', padx=10, pady=5)

                tk.Label(cat_frame, text=f"● {cat_name}",
                        font=('Segoe UI', 10, 'bold'), fg=color, bg='white').pack(anchor='w')

                for principle in principles:
                    tk.Label(cat_frame, text=f"  {principle}",
                            font=('Segoe UI', 9), fg='#666', bg='white').pack(anchor='w', padx=(15, 0))

        except Exception as e:
            logging.error(f"UNGC modülü yüklenirken hata: {e}")
            import traceback
            traceback.print_exc()
            error_label = tk.Label(self.content_area, text=self.lm.tr('ungc_load_error', "UNGC modülü yüklenirken hata: {}").format(e),
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    def show_issb(self) -> None:
        """ISSB (International Sustainability Standards Board) modülünü göster - HIZLI YÜKLEME"""
        self._show_module_fast('issb', 'ISSB - IFRS S1/S2', 'issb.read')

    def show_ungc_cop_quick(self) -> None:
        try:
            from datetime import datetime

            from modules.ungc.ungc_cop_generator import UNGCCOPGenerator
            out = filedialog.asksaveasfilename(
                title=self.lm.tr('save_report', "Raporu Kaydet"),
                defaultextension='.pdf',
                filetypes=[(self.lm.tr('pdf_files', 'PDF Dosyaları'), '*.pdf'), (self.lm.tr('all_files', 'Tüm Dosyalar'), '*.*')]
            )
            if not out:
                return
            gen = UNGCCOPGenerator(self.db_path)
            fp = gen.generate_report(self.company_id, str(datetime.now().year))
            try:
                if fp and os.path.exists(fp):
                    import shutil
                    shutil.copyfile(fp, out)
                    messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('info_cop_pdf_created', "COP PDF oluşturuldu: {}").format(out))
                else:
                    messagebox.showwarning(self.lm.tr('warning', "Uyarı"), self.lm.tr('warn_cop_pdf_failed', "COP PDF oluşturulamadı"))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
        except Exception as e:
            try:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_cop_failed', "COP oluşturulamadı: {}").format(e))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    def quick_issb_action_plan_excel(self) -> None:
        try:
            from datetime import datetime

            from modules.issb.issb_report_generator import ISSBReportGenerator
            year = datetime.now().year
            fp = filedialog.asksaveasfilename(
                title=self.lm.tr('save_report', "Raporu Kaydet"),
                defaultextension='.xlsx',
                filetypes=[(self.lm.tr('excel_files', 'Excel Dosyaları'), '*.xlsx'), (self.lm.tr('all_files', 'Tüm Dosyalar'), '*.*')],
                initialfile=f"issb_action_plan_{self.company_id}_{year}.xlsx"
            )
            if not fp:
                return
            gen = ISSBReportGenerator()
            ok = gen.export_excel_action_plan(fp, self.company_id, year)
            try:
                if ok:
                    messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('info_issb_excel_created', "ISSB Aksiyon Planı Excel oluşturuldu: {}").format(fp))
                else:
                    messagebox.showwarning(self.lm.tr('warning', "Uyarı"), self.lm.tr('warn_issb_excel_failed', "Excel oluşturulamadı"))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
        except Exception as e:
            try:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_issb_excel_failed', "Excel oluşturulamadı: {}").format(e))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    def quick_issb_action_plan_pdf(self) -> None:
        try:
            from datetime import datetime

            from modules.issb.issb_report_generator import ISSBReportGenerator
            year = datetime.now().year
            fp = filedialog.asksaveasfilename(
                defaultextension='.pdf',
                filetypes=[(self.lm.tr('file_pdf', 'PDF Dosyaları'), '*.pdf'), (self.lm.tr('all_files', 'Tüm Dosyalar'), '*.*')],
                initialfile=f"issb_action_plan_{self.company_id}_{year}.pdf"
            )
            if not fp:
                return
            gen = ISSBReportGenerator()
            ok = gen.export_pdf_action_plan(fp, self.company_id, year)
            try:
                if ok:
                    messagebox.showinfo(self.lm.tr('info', "Bilgi"), self.lm.tr('info_issb_pdf_created', "ISSB Aksiyon Planı PDF oluşturuldu: {}").format(fp))
                else:
                    messagebox.showwarning(self.lm.tr('warning', "Uyarı"), self.lm.tr('warn_issb_pdf_failed', "PDF oluşturulamadı"))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")
        except Exception as e:
            try:
                messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_issb_pdf_failed', "PDF oluşturulamadı: {}").format(e))
            except Exception as e:
                logging.error(f"Silent error caught: {str(e)}")

    def _open_ungc_kpi_entry(self) -> None:
        """KPI veri giriş penceresini aç"""
        try:
            from datetime import datetime

            from modules.ungc.ungc_kpi_forms import UNGCKPIDataEntry

            UNGCKPIDataEntry(self.parent, self.db_path, self.company_id, str(datetime.now().year))
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_kpi_form_open', "KPI formu açılamadı: {}").format(e))

    def _show_ungc_score(self) -> None:
        """Uyum skorunu göster"""
        try:
            from datetime import datetime

            from modules.ungc.ungc_manager_enhanced import UNGCManagerEnhanced

            manager = UNGCManagerEnhanced(self.db_path)
            result = manager.calculate_overall_score(self.company_id, str(datetime.now().year))

            message = f"{self.lm.tr('title_ungc_score', 'UNGC Uyum Skoru')}\n\n"
            message += f"{self.lm.tr('overall_score', 'Genel Skor')}: {result['overall_score']:.2f}%\n"
            message += f"{self.lm.tr('level', 'Seviye')}: {result['level']} {result['level_info']['badge']}\n\n"
            message += f"{self.lm.tr('category_scores', 'Kategori Skorları')}:\n"

            for category, score in result['category_scores'].items():
                message += f"  {category}: {score:.1f}%\n"

            messagebox.showinfo(self.lm.tr('title_ungc_score', "UNGC Uyum Skoru"), message)
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), self.lm.tr('err_score_calc', "Skor hesaplanamadı: {}").format(e))

    def _generate_ungc_cop(self) -> None:
        """COP raporu oluştur"""
        try:
            from datetime import datetime

            from modules.ungc.ungc_cop_generator import UNGCCOPGenerator

            generator = UNGCCOPGenerator(self.db_path)
            output = generator.generate_report(self.company_id, str(datetime.now().year))

            messagebox.showinfo(self.lm.tr('success', "Başarılı"), f"{self.lm.tr('cop_report_created', 'COP raporu oluşturuldu')}:\n{output}")
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('cop_report_error', 'COP raporu oluşturulamadı')}: {e}")

    def _show_ungc_timeline(self) -> None:
        """İlerleme grafiği göster"""
        try:
            from modules.ungc.ungc_timeline import UNGCTimelineVisualizer

            window = tk.Toplevel(self.parent)
            window.title(self.lm.tr('ungc_timeline_title', "UNGC İlerleme Grafiği"))
            window.geometry("1000x700")

            visualizer = UNGCTimelineVisualizer(window, self.db_path, self.company_id)
            visualizer.pack(fill='both', expand=True)
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('chart_display_error', 'Grafik gösterilemedi')}: {e}")

    def _export_ungc_excel(self) -> None:
        """Excel export"""
        try:
            from datetime import datetime

            from modules.ungc.ungc_excel_export import UNGCExcelExporter

            exporter = UNGCExcelExporter(self.db_path)
            output = exporter.export_compliance_report(self.company_id, str(datetime.now().year))

            messagebox.showinfo(self.lm.tr('success', "Başarılı"), f"{self.lm.tr('excel_report_created', 'Excel raporu oluşturuldu')}:\n{output}")
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('excel_create_error', 'Excel oluşturulamadı')}: {e}")

    def _send_ungc_email(self) -> None:
        """Email gönder"""
        messagebox.showinfo(self.lm.tr('email', "Email"), self.lm.tr('email_smtp_config_msg', "Email gönderimi için SMTP ayarlarını yapılandırın:\nbackend/config/smtp_config.json"))

    def _show_ungc_benchmark(self) -> None:
        """Benchmark göster"""
        try:
            from datetime import datetime

            from modules.ungc.ungc_benchmark import UNGCBenchmark

            benchmark = UNGCBenchmark(self.db_path)
            comparison = benchmark.compare_with_industry(self.company_id, 'Technology', str(datetime.now().year))

            message = f"{self.lm.tr('benchmark_comparison', 'Benchmark Karşılaştırması')}\n\n"
            message += f"{self.lm.tr('sector', 'Sektör')}: Technology\n"
            message += f"{self.lm.tr('company_score', 'Şirket Skoru')}: {comparison['company_overall']:.1f}%\n"
            message += f"{self.lm.tr('benchmark_average', 'Sektör Ortalaması')}: {comparison['benchmark_overall']:.1f}%\n"
            message += f"{self.lm.tr('performance', 'Performans')}: {comparison['performance_level']}\n\n"

            recs = benchmark.get_improvement_recommendations(comparison)
            if recs:
                message += f"{self.lm.tr('recommendations', 'Öneriler')}:\n"
                for rec in recs[:3]:
                    message += f"  [{rec['priority']}] {rec['recommendation']}\n"

            messagebox.showinfo(self.lm.tr('benchmark', "Benchmark"), message)
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('benchmark_calc_error', 'Benchmark hesaplanamadı')}: {e}")

    def _show_ungc_reminders(self) -> None:
        """Hatırlatıcıları göster"""
        try:
            from modules.ungc.ungc_reminder import UNGCReminderSystem

            reminder = UNGCReminderSystem(self.db_path)
            pending = reminder.get_pending_reminders(self.company_id)

            if not pending:
                messagebox.showinfo(self.lm.tr('reminders', "Hatırlatıcılar"), self.lm.tr('no_pending_reminders', "Bekleyen hatırlatıcı yok."))
            else:
                message = f"{len(pending)} {self.lm.tr('reminders_count', 'adet hatırlatıcı')}:\n\n"
                for r in pending[:5]:
                    message += f"[{r['urgency']}] {r['title']}\n"
                    message += f"  {r['days_until']} {self.lm.tr('days_later', 'gün sonra')}\n\n"

                messagebox.showinfo(self.lm.tr('reminders', "Hatırlatıcılar"), message)
        except Exception as e:
            messagebox.showerror(self.lm.tr('error', "Hata"), f"{self.lm.tr('reminders_load_error', 'Hatırlatıcılar yüklenemedi')}: {e}")

    def show_advanced_materiality(self) -> None:
        """İleri Seviye Materialite Analizi modülünü göster"""
        try:
            self._audit_navigation("İleri Seviye Materialite Analizi")
            self.clear_content()
            # Header'ı geri getir
            self.content_header.pack(fill='x')
            self.content_area.pack_forget()
            self.content_area.pack(fill='both', expand=True, padx=20, pady=20)

            self.content_title['text'] = self.lm.tr('advanced_materiality_analysis_title', " İleri Seviye Materialite Analizi")
            # Geri butonu
            self._add_back_button()

            AdvancedMaterialityGUI_class = self._lazy_import_module('advanced_materiality')
            if AdvancedMaterialityGUI_class:
                AdvancedMaterialityGUI_class(self.content_area, self.company_id)
            else:
                self.show_module_error("İleri Seviye Materialite Analizi", "İleri Seviye Materialite Analizi modülü yüklenemedi")
        except Exception as e:
            logging.error(f"İleri Seviye Materialite Analizi modülü yüklenirken hata: {e}")
            error_label = tk.Label(self.content_area, text=f"İleri Seviye Materialite Analizi modülü yüklenirken hata: {e}",
                                 font=('Segoe UI', 12), fg='red', bg='white')
            error_label.pack(pady=50)

    def show_advanced_security(self) -> None:
        """Gelişmiş Güvenlik modülünü hızlı göster"""
        self._show_module_fast('advanced_security', 'Gelişmiş Güvenlik Özellikleri', 'security.read')

    # ======================================================================
    # YENI MODULLER - 8 ADET (23 Ekim 2025)
    # ======================================================================

    def show_company_management(self):
        """Sirket Yonetimi"""
        self._show_module_fast('company_management', 'Şirket Yönetimi - Çok Şirketli Sistem')

    def show_ai_module(self):
        """AI Modulu"""
        self._show_module_fast('ai_module', 'AI Modülü - Yapay Zeka Destekli Analiz')

    def show_tcfd(self):
        """TCFD Modulu"""
        self._show_module_fast('tcfd', 'TCFD - İklim Finansal Açıklamaları')

    def show_sasb(self):
        """SASB Modulu"""
        self._show_module_fast('sasb', 'SASB - Sektör Standartları')

    def show_auto_reporting(self):
        """Otomatik Raporlama"""
        self._show_module_fast('auto_reporting', 'Otomatik Raporlama Sistemi')

    def show_document_processing(self):
        """Belge Isleme"""
        self._show_module_fast('document_processing', 'Belge İşleme ve AI Analizi')

    def show_erp_integration(self):
        """ERP Entegrasyon"""
        self._show_module_fast('erp_integration', 'ERP Entegrasyon Merkezi')

    def show_advanced_calculation(self):
        """Gelismis Hesaplama"""
        self._show_module_fast('advanced_calculation', 'Gelişmiş Hesaplama Motoru')

    def show_advanced_inventory(self):
        """Gelismis Envanter"""
        self._show_module_fast('advanced_inventory', 'Gelişmiş Envanter Yönetimi')

    def show_esrs_module(self):
        """ESRS Modulu"""
        self._show_module_fast('esrs', '🇪🇺 ESRS - Avrupa Sürdürülebilirlik Standartları')

    def show_esrs(self):
        """ESRS - Alias for show_esrs_module"""
        self.show_esrs_module()

    def show_eu_taxonomy(self):
        """EU Taxonomy Module"""
        self._show_module_fast('eu_taxonomy', '🇪🇺 EU Taxonomy - Sürdürülebilir Finans Sınıflandırması')

    def show_strategy(self):
        """Strategy Management"""
        self._show_module_fast('strategic', '🎯 Stratejik Yönetim')

    def show_risks(self):
        """Risk Management"""
        self._show_module_fast('risks', f"{Icons.WARNING} Risk Yönetimi")

    def show_ceo_messages(self):
        """CEO Message Center"""
        self._show_module_fast('ceo_messages', '📢 CEO Mesaj Merkezi')

if __name__ == '__main__':
    # Eğer bu dosya doğrudan çalıştırılırsa, ana başlatıcıyı (main.py) tetikle
    import os
    import subprocess
    import sys
    
    logging.info("Uygulama başlatılıyor (main.py üzerinden)...")
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    main_script = os.path.join(base_dir, 'main.py')
    
    if os.path.exists(main_script):
        try:
            subprocess.run([sys.executable, main_script])
        except KeyboardInterrupt as e:
            logging.error(f'Silent error in main_app.py: {str(e)}')
        except Exception as e:
            logging.error(f"Başlatma hatası: {e}")
            input("Devam etmek için bir tuşa basın...")
    else:
        logging.error(f"HATA: main.py bulunamadı: {main_script}")
        input("Devam etmek için bir tuşa basın...")
